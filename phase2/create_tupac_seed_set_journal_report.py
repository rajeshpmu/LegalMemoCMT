"""Create a journal/report document for a Tupac human-review example."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments/LegalMemoCMT_Tupac_Manual_Review_Seed_Set_Journal_Report.docx"


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_heading("LegalMemoCMT: Human Review Seed-Set Learning Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Worked Tupac example: lOcZ_IJbM3I_turn08614")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Purpose: document how one machine-assisted Tupac utterance is reviewed, corrected, "
        "and converted into a human-reviewed training example. The example demonstrates why "
        "basic emotion and courtroom interactional affect must be represented as separate tasks."
    )

    doc.add_heading("1. Executive Summary", 1)
    doc.add_paragraph(
        "The Phase 1 trimodal checkpoint predicted anger with confidence 0.701481 for this "
        "1.481-second witness clip. That prediction is retained as provenance, but it is not "
        "accepted as the human label. The transcript contains no explicit emotional language, "
        "DeBERTa identified NO_EMOTION_CONTENT at confidence 0.658731, and SpeechBrain produced "
        "neutral with confidence 1.0. Odyssey produced mid-range valence, arousal, and dominance. "
        "After audiovisual review, the defensible basic-emotion label is neutral, while the "
        "presentation may be described separately as assertive or emphatic."
    )
    doc.add_paragraph(
        "This is precisely the type of example to include in the initial 500-1,000 utterance "
        "human-reviewed seed set. It teaches an adapted LegalMemoCMT model not to conflate "
        "assertiveness with anger."
    )

    doc.add_heading("2. Utterance and Provenance", 1)
    add_table(
        doc,
        ["Field", "Observed value", "Why it matters"],
        [
            ["utterance_id", "lOcZ_IJbM3I_turn08614", "Stable key used to join every modality and review decision."],
            ["youtube_id", "lOcZ_IJbM3I", "Identifies the raw courtroom source."],
            ["speaker_cluster_id", "SPEAKER_10", "Connects the turn to the diarized voice cluster."],
            ["speaker_role", "Witness", "Confirms that the row belongs to the witness-speaking corpus."],
            ["split", "train", "Determines the eventual training partition."],
            ["start / end", "00:29:36.488 / 00:29:37.969", "Allows the reviewer to reproduce the exact video interval."],
            ["duration", "1.481 seconds", "Very short; visual and prosodic evidence must be interpreted cautiously."],
            ["transcript", "He was the one that got them there.", "Declarative content without an explicit emotion term."],
        ],
    )
    doc.add_paragraph(
        "The reviewer should open both the MP4 and WAV referenced by the manifest. The "
        "utterance ID, source, and timestamps are not merely descriptive fields: together they "
        "form the reproducibility link from the final label back to the raw evidence."
    )

    doc.add_heading("3. Machine Evidence Before Human Review", 1)
    add_table(
        doc,
        ["Evidence source", "Output", "Technical interpretation"],
        [
            ["Phase 1 trimodal MELD model", "anger, 0.701481", "A weak-label prediction learned from MELD-style emotion correlations."],
            ["SpeechBrain audio SER", "neu, 1.0", "Categorical acoustic evidence supports neutral; confidence is not literal certainty."],
            ["Odyssey valence", "0.558519", "Near the middle of the 0-1 evidence range; not strongly negative."],
            ["Odyssey arousal", "0.518461", "Moderate activation; not sufficient by itself to indicate anger."],
            ["Odyssey dominance", "0.527564", "Only slightly above midpoint; weak support for firm delivery."],
            ["DeBERTa target scope", "NO_EMOTION_CONTENT, 0.658731", "The statement does not contain a meaningful explicit emotional target."],
            ["DeBERTa temporal scope", "NOT_APPLICABLE", "Correct because no emotional content was selected."],
            ["Courtroom-affect heuristic", "ASSERTIVE, 0.60", "A candidate behavioral description, not a gold label."],
        ],
    )
    doc.add_paragraph(
        "The important point is that the models answer different questions. Phase 1 estimates "
        "which MELD emotion class the multimodal pattern resembles. SpeechBrain estimates an "
        "audio emotion class from an IEMOCAP-trained model. Odyssey supplies continuous acoustic "
        "attributes. DeBERTa evaluates the semantic meaning and target of the text. None of these "
        "outputs alone is a human annotation for courtroom testimony."
    )

    doc.add_heading("4. Why the Anger Prediction Is Not Accepted", 1)
    add_bullets(
        doc,
        [
            "The text is factual and declarative: it identifies who brought people somewhere.",
            "There is no explicit anger word, insult, threat, correction, or disagreement phrase.",
            "SpeechBrain predicts neutral, providing independent acoustic evidence against a categorical anger interpretation.",
            "Odyssey valence is 0.558519, not strongly negative, and arousal is only moderate at 0.518461.",
            "Dominance at 0.527564 can support an emphatic delivery hypothesis, but it is too close to midpoint to prove anger.",
            "The clip lasts only 1.481 seconds, so a single facial movement or prosodic moment should not be overinterpreted.",
        ],
    )
    doc.add_paragraph(
        "The correct conclusion is not that the Phase 1 model is useless. Its error is valuable: "
        "it identifies a domain-shift pattern in which assertive courtroom delivery can resemble "
        "anger in a MELD-trained classifier. Preserving the prediction allows this failure mode "
        "to be measured during later fine-tuning and evaluation."
    )

    doc.add_heading("5. Separating Basic Emotion from Courtroom Affect", 1)
    add_table(
        doc,
        ["Dimension", "Recommended value", "Meaning"],
        [
            ["Basic emotion", "neutral", "The witness does not provide sufficient evidence of anger as an expressed categorical emotion."],
            ["Courtroom affect", "ASSERTIVE or emphatic candidate", "The witness appears engaged and firm in delivering a factual answer."],
            ["Affect intensity", "1 (mild), subject to review", "The available evidence is limited and the segment is very short."],
            ["Target scope", "NO_EMOTION_CONTENT", "The semantic content does not identify an emotional state."],
            ["Temporal scope", "NOT_APPLICABLE", "There is no emotional state for which a time reference is needed."],
        ],
    )
    doc.add_paragraph(
        "These labels are not contradictory. A witness can be basic_emotion=neutral while "
        "courtroom_affect=ASSERTIVE. Basic emotion describes expressed affective state; courtroom "
        "affect describes interactional presentation during examination. Keeping the dimensions "
        "separate prevents the training target from treating every firm answer as anger."
    )

    doc.add_heading("6. Recommended Human Review Record", 1)
    add_table(
        doc,
        ["Field", "Value", "Status"],
        [
            ["human_basic_emotion", "neutral", "Human decision"],
            ["human_basic_emotion_review_status", "CONFIRMED", "Allowed project status"],
            ["human_basic_emotion_confidence", "0.85", "Reviewer confidence; not model probability"],
            ["human_basic_emotion_reviewer", "your-name", "Replace with reviewer identifier"],
            ["human_basic_emotion_notes", "Phase 1 anger conflicts with neutral audio and no-emotion transcript; delivery is assertive rather than angry.", "Audit explanation"],
            ["human_courtroom_affect", "ASSERTIVE", "Recommended extension if human affect fields are enabled"],
            ["human_courtroom_affect_confidence", "0.75", "Separate affect confidence"],
            ["human_affect_intensity", "1", "Mild, because the clip is short and evidence is limited"],
        ],
    )
    doc.add_paragraph(
        "The current basic-emotion review workflow accepts CONFIRMED, REJECTED, and DEFERRED. "
        "Do not use HUMAN_SINGLE or REVIEWED in human_basic_emotion_review_status unless the schema "
        "is explicitly changed. Machine fields such as phase1_basic_emotion and proposed_basic_emotion "
        "must remain unchanged."
    )

    doc.add_heading("7. How This Becomes a Training Example", 1)
    doc.add_paragraph(
        "After the reviewer enters the fields above, the merge stage gives the confirmed human "
        "label priority over the machine prediction. The common training target becomes:"
    )
    doc.add_paragraph("training_label=neutral", style="Intense Quote")
    add_bullets(
        doc,
        [
            "training_label_source=HUMAN_GOLD",
            "training_label_status=CONFIRMED",
            "training_label_is_human_gold=YES",
            "phase1_basic_emotion=anger remains available for error analysis.",
            "The courtroom-affect decision remains a separate target and should not be encoded as anger.",
        ],
    )
    doc.add_paragraph(
        "The row can therefore be used to fine-tune the basic-emotion task with neutral as the "
        "reviewed target, while also contributing to a separate courtroom-affect task if the "
        "affect annotation is confirmed. It should not be used as a final test example if it was "
        "included in training."
    )

    doc.add_heading("8. Why This Belongs in the 500-1,000 Row Seed Set", 1)
    add_bullets(
        doc,
        [
            "It is a clear Phase 1 versus human disagreement: anger at 0.701481 versus reviewed neutral.",
            "It contains independent audio evidence: SpeechBrain neutral at 1.0.",
            "It contains semantic evidence: DeBERTa NO_EMOTION_CONTENT at 0.658731.",
            "It demonstrates that moderate arousal and dominance do not automatically imply anger.",
            "It separates factual courtroom assertiveness from categorical emotional expression.",
            "It is short and therefore useful for testing whether a model is overreacting to brief emphatic segments.",
            "It provides an auditable error case for the final report, not just a corrected label.",
        ],
    )
    doc.add_paragraph(
        "The seed set should be stratified rather than selected only by model confidence. Include "
        "high-confidence disagreements, neutral/non-neutral conflicts, scope disagreements, each "
        "emotion class, each courtroom affect, different speakers and sources, and duration bands. "
        "This gives the adapted model examples of the failure modes it must learn to correct."
    )

    doc.add_heading("9. Generalizable Learning for LegalMemoCMT", 1)
    doc.add_paragraph(
        "This example supports a central design principle for the project: the emotional meaning "
        "of courtroom language and the behavior of the current speaker are not always the same. "
        "A witness may discuss a violent event, another person's distress, or a disputed fact in a "
        "controlled voice. Conversely, a firm answer may be assertive without being angry."
    )
    add_bullets(
        doc,
        [
            "Preserve original machine predictions for domain-shift measurement.",
            "Use transcript scope to prevent emotion-target leakage.",
            "Use audio and video to validate expressed speaker behavior, not to assign text target scope.",
            "Treat SpeechBrain confidence as evidence, not certainty, because it is out of domain.",
            "Use Odyssey attributes as low-level acoustic evidence rather than direct courtroom labels.",
            "Keep basic emotion and courtroom affect as separate prediction tasks.",
            "Use human-confirmed rows for gold evaluation and clearly marked machine-assisted rows only for provisional training experiments.",
        ],
    )

    doc.add_heading("10. Reproducible Review Workflow", 1)
    doc.add_paragraph("The row was produced through this evidence chain:")
    for step in [
        "Tupac raw video and subtitles were converted into turn-level clips.",
        "Pyannote diarization assigned SPEAKER_10 to the turn.",
        "Manual role mapping identified the cluster as a Witness.",
        "The row passed the witness-speaking and duration filters.",
        "ViT face-crop features were extracted for the trimodal Phase 1 checkpoint.",
        "The Phase 1 checkpoint generated the original weak emotion prediction.",
        "Odyssey and SpeechBrain generated independent audio evidence.",
        "DeBERTa and rule-based scope logic evaluated the transcript meaning.",
        "A human reviewed the transcript, audio, and video and assigned the final basic-emotion label.",
        "The merge stage created the common training_label only after CONFIRMED review status.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("11. Conclusion", 1)
    doc.add_paragraph(
        "For lOcZ_IJbM3I_turn08614, the recommended human basic-emotion label is neutral, not "
        "anger. The evidence does not show a meaningful emotional target in the words, the audio "
        "cross-check is neutral, and the continuous acoustic attributes are moderate rather than "
        "strongly negative. The delivery can still be recorded as assertive or emphatic in the "
        "separate courtroom-affect dimension."
    )
    doc.add_paragraph(
        "This is not merely a single-row correction. It is a training-design example showing why "
        "LegalMemoCMT needs provenance-preserving multimodal review, separate label dimensions, "
        "and a human-reviewed seed set. Similar cases should be deliberately collected so the "
        "adapted model learns the distinction between what a witness says, how the witness says it, "
        "and whose emotion the text describes."
    )
    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
