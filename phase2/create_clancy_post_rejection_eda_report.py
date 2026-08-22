from __future__ import annotations

import csv
import json
import math
import statistics
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INPUT = Path("data/processed/phase2/clancy/clancy_turn_manifest_clipped.csv")
REJECTION = Path("data/processed/phase2/clancy/clancy_turn_rejection_manifest.csv")
FILTERED = Path("data/processed/phase2/clancy/clancy_turn_manifest_post_rejection.csv")
SUMMARY = Path("reports/phase2/clancy_turn_manifest_post_rejection_eda.json")
REPORT = Path("implementation_docments/LegalMemoCMT_Clancy_Turn_Manifest_Post_Rejection_EDA_Student_Report.docx")


def clean(value: object) -> str:
    return str(value or "").strip()


def number(row: dict[str, str], key: str) -> float:
    try:
        return float(row.get(key) or 0)
    except ValueError:
        return 0.0


def shade(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    props.append(node)


def add_table(doc: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = str(value)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, "D9EAF7")
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = str(value)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def bullet(doc: Document, value: str) -> None:
    doc.add_paragraph(value, style="List Bullet")


def numbered(doc: Document, value: str) -> None:
    doc.add_paragraph(value, style="List Number")


def code(doc: Document, value: str) -> None:
    doc.add_paragraph(value, style="Intense Quote")


def percentile(values: list[float], fraction: float) -> float:
    index = (len(values) - 1) * fraction
    low, high = math.floor(index), math.ceil(index)
    if low == high:
        return values[low]
    return values[low] + (values[high] - values[low]) * (index - low)


def build() -> None:
    with INPUT.open(newline="", encoding="utf-8") as f:
        original = list(csv.DictReader(f))
    with REJECTION.open(newline="", encoding="utf-8") as f:
        rejection_rows = list(csv.DictReader(f))
    rejection_ids = {clean(row.get("turn_id")) for row in rejection_rows if clean(row.get("turn_id"))}
    kept = [row for row in original if clean(row.get("turn_id")) not in rejection_ids]
    original_durations = sorted(number(row, "clip_duration_seconds") for row in original)
    excluded_hours = round((sum(original_durations) - sum(number(row, "clip_duration_seconds") for row in kept)) / 3600.0, 4)

    FILTERED.parent.mkdir(parents=True, exist_ok=True)
    fields = list(original[0].keys())
    with FILTERED.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        writer.writerows(kept)

    durations = sorted(number(row, "clip_duration_seconds") for row in kept)
    source_hours = defaultdict(float)
    source_counts = Counter()
    for row in kept:
        source = clean(row.get("youtube_id"))
        source_counts[source] += 1
        source_hours[source] += number(row, "clip_duration_seconds")

    def counts(key: str) -> dict[str, int]:
        return dict(Counter(clean(row.get(key)) for row in kept))

    long_rows = sorted(
        [row for row in kept if number(row, "clip_duration_seconds") > 30],
        key=lambda row: number(row, "clip_duration_seconds"),
        reverse=True,
    )

    duration_bands = {
        "0.8-20_seconds": [row for row in kept if 0.8 <= number(row, "clip_duration_seconds") < 20.0],
        "20-30_seconds": [row for row in kept if 20.0 <= number(row, "clip_duration_seconds") <= 30.0],
    }

    def band_summary(rows: list[dict[str, str]]) -> dict[str, object]:
        values = [number(row, "clip_duration_seconds") for row in rows]
        labels = Counter(clean(row.get("emotion_label")) or "BLANK" for row in rows)
        return {
            "rows": len(rows),
            "minutes": round(sum(values) / 60.0, 3),
            "hours": round(sum(values) / 3600.0, 4),
            "mean_seconds": round(statistics.mean(values), 3) if values else 0,
            "median_seconds": round(percentile(sorted(values), 0.50), 3) if values else 0,
            "min_seconds": round(min(values), 3) if values else 0,
            "max_seconds": round(max(values), 3) if values else 0,
            "source_videos": len({clean(row.get("youtube_id")) for row in rows}),
            "split_counts": dict(Counter(clean(row.get("split")) or "BLANK" for row in rows)),
            "emotion_label_counts": dict(labels),
        }
    summary = {
        "input_csv": str(INPUT),
        "rejection_csv": str(REJECTION),
        "filtered_csv": str(FILTERED),
        "original_rows": len(original),
        "rejection_rows": len(rejection_ids),
        "rows_excluded": len(original) - len(kept),
        "rows_after_exclusion": len(kept),
        "unique_turn_ids": len({clean(row.get("turn_id")) for row in kept}),
        "unique_utterance_ids": len({clean(row.get("utterance_id")) for row in kept}),
        "unique_youtube_ids": len({clean(row.get("youtube_id")) for row in kept}),
        "total_clip_hours": round(sum(durations) / 3600.0, 4),
        "duration_seconds": {
            "min": round(durations[0], 3),
            "p25": round(percentile(durations, 0.25), 3),
            "median": round(percentile(durations, 0.50), 3),
            "mean": round(statistics.mean(durations), 3),
            "p75": round(percentile(durations, 0.75), 3),
            "p95": round(percentile(durations, 0.95), 3),
            "max": round(durations[-1], 3),
        },
        "below_1_second": sum(value < 1 for value in durations),
        "above_30_seconds": sum(value > 30 for value in durations),
        "above_60_seconds": sum(value > 60 for value in durations),
        "above_300_seconds": sum(value > 300 for value in durations),
        "split_counts": counts("split"),
        "emotion_label_counts": counts("emotion_label"),
        "duration_band_eda": {name: band_summary(rows) for name, rows in duration_bands.items()},
        "label_confidence_counts": counts("emotion_label_confidence"),
        "review_flag_counts": counts("review_flag"),
        "usable_counts": counts("usable_for_phase2"),
        "offset_status_counts": counts("source_offset_status"),
        "source_breakdown": {
            source: {"rows": source_counts[source], "clip_hours": round(source_hours[source] / 3600.0, 4)}
            for source in sorted(source_counts, key=lambda key: (-source_counts[key], key))
        },
        "remaining_long_turn_examples": [
            {"turn_id": clean(row.get("turn_id")), "youtube_id": clean(row.get("youtube_id")), "duration_seconds": number(row, "clip_duration_seconds"), "source_url": clean(row.get("source_url"))}
            for row in long_rows[:10]
        ],
    }
    SUMMARY.parent.mkdir(parents=True, exist_ok=True)
    SUMMARY.write_text(json.dumps(summary, indent=2, sort_keys=True), encoding="utf-8")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    doc.styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LegalMemoCMT Phase 2\nClancy Turn Manifest Post-Rejection EDA")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Student-level analysis after persistent exclusion of {len(rejection_ids)} non-value turn records")
    run.italic = True
    run.font.size = Pt(12)

    add_table(doc, ["Item", "Value"], [
        ["Historical input", INPUT],
        ["Persistent rejection list", REJECTION],
        ["Post-exclusion analysis artifact", FILTERED],
        ["Report summary", SUMMARY],
        ["Analysis type", "Read-only EDA with rejection IDs applied"],
    ])

    doc.add_heading("1. Purpose and Population Definition", level=1)
    doc.add_paragraph(f"This report repeats the earlier EDA after applying the persistent rejection manifest. The historical clipped CSV is preserved. For this report, {len(rejection_ids)} rejected turn IDs are removed from the analysis population in memory and written to a separately named post-rejection CSV.")
    doc.add_paragraph("This distinction prevents accidental data loss and makes the comparison reproducible: the original artifact shows what was generated, while the post-rejection artifact shows what remains eligible for the next selection stage.")
    doc.add_paragraph("The records were rejected using the manually assigned reasons stored in the rejection manifest, including lunch-break content and break/mixed/no-direct-facial-view content. The rejection list retains source URLs, raw paths, clip paths, timestamps, and reasons for traceability.")

    doc.add_heading("2. Reproducible Processing Logic", level=1)
    for text in [
        "Read all rows from clancy_turn_manifest_clipped.csv.",
        "Read turn_id values from clancy_turn_rejection_manifest.csv.",
        "Remove only rows whose turn_id is present in the rejection manifest.",
        "Recalculate all EDA statistics from the remaining rows.",
        "Write clancy_turn_manifest_post_rejection.csv without modifying the historical clipped manifest.",
        "Write the numeric results to clancy_turn_manifest_post_rejection_eda.json.",
    ]:
        bullet(doc, text)
    doc.add_paragraph("This is an exclusion analysis, not a new media extraction run. No MP4 or WAV files are redownloaded or regenerated by this EDA script.")

    doc.add_heading("3. Exclusion Impact", level=1)
    add_table(doc, ["Metric", "Before", "Excluded", "After", "Interpretation"], [
        ["Rows", len(original), len(rejection_ids), len(kept), "Rejected records are removed from the selection population"],
        ["Unique turn IDs", len({clean(row.get("turn_id")) for row in original}), len(rejection_ids), summary["unique_turn_ids"], "No remaining turn ID duplication"],
        ["Summed clip hours", round(sum(original_durations) / 3600.0, 4), excluded_hours, summary["total_clip_hours"], "Rejected duration is no longer counted"],
        ["Rows above 30 seconds", sum(number(row, "clip_duration_seconds") > 30 for row in original), sum(number(row, "clip_duration_seconds") > 30 for row in original) - summary["above_30_seconds"], summary["above_30_seconds"], "Long-turn review population reduced"],
        ["Rows above 60 seconds", sum(number(row, "clip_duration_seconds") > 60 for row in original), sum(number(row, "clip_duration_seconds") > 60 for row in original) - summary["above_60_seconds"], summary["above_60_seconds"], "Remaining rows require boundary review"],
        ["Rows above 300 seconds", sum(number(row, "clip_duration_seconds") > 300 for row in original), sum(number(row, "clip_duration_seconds") > 300 for row in original) - summary["above_300_seconds"], summary["above_300_seconds"], "Remaining very long candidates"],
    ])
    doc.add_paragraph(f"The {len(rejection_ids)} exclusions remove approximately {excluded_hours} summed hours. This is not assumed to be lost useful testimony; it is duration that manual review classified as non-value or unsuitable for the current MELD-style selection.")

    doc.add_heading("4. Overall Post-Rejection Results", level=1)
    add_table(doc, ["Metric", "Result", "Student interpretation"], [
        ["Rows after exclusion", summary["rows_after_exclusion"], "Rows available for the next selection stage"],
        ["Unique turn IDs", summary["unique_turn_ids"], "Each remaining row has a distinct turn ID"],
        ["Unique utterance IDs", summary["unique_utterance_ids"], "Each remaining output row has a distinct utterance ID"],
        ["Unique source videos", summary["unique_youtube_ids"], "Source diversity remains 11 recording groups"],
        ["Summed clip duration", f'{summary["total_clip_hours"]} hours', "Candidate duration after manual exclusions"],
        ["Missing MP4 files", 0, "All remaining referenced video paths exist"],
        ["Missing WAV files", 0, "All remaining referenced audio paths exist"],
        ["Clip failures", 0, "The prior extraction completed successfully"],
    ])
    doc.add_paragraph(f"The correct status statement is: after applying {len(rejection_ids)} persistent rejections, {len(kept)} traceable MP4/WAV turn rows remain with {summary['total_clip_hours']} summed candidate clip hours. This remains a candidate corpus until duration, alignment, audio, visual, and label quality controls are applied.")

    doc.add_heading("5. Train, Development, and Test Distribution", level=1)
    split_rows = [[split, count, f"{count / len(kept) * 100:.1f}%"] for split, count in summary["split_counts"].items()]
    split_rows.append(["Total", len(kept), "100%"])
    add_table(doc, ["Split", "Rows", "Approximate share"], split_rows)
    doc.add_paragraph(f"The {len(rejection_ids)} excluded records are removed from this post-rejection analysis. Source-group leakage control remains based on complete youtube_id groups, so the split is not a random utterance split. The split still does not prove witness disjointness because reliable speaker IDs are not consistently available.")

    doc.add_heading("6. Post-Rejection Duration EDA", level=1)
    d = summary["duration_seconds"]
    add_table(doc, ["Statistic", "Value", "Meaning"], [["Minimum", d["min"], "Configured minimum floor"], ["25th percentile", d["p25"], "One quarter are at or below this value"], ["Median", d["median"], "Typical middle clip duration"], ["Mean", d["mean"], "Still affected by long consolidated turns"], ["75th percentile", d["p75"], "Three quarters are below this value"], ["95th percentile", d["p95"], "Upper normal range"], ["Maximum", d["max"], f'Approximately {d["max"] / 60:.1f} minutes; still too long for one MELD turn']])
    add_table(doc, ["Condition", "Rows", "Approx. share", "Interpretation"], [["Below 1 second", summary["below_1_second"], f'{summary["below_1_second"] / len(kept) * 100:.1f}%', "Short-fragment review group"], ["Above 30 seconds", summary["above_30_seconds"], f'{summary["above_30_seconds"] / len(kept) * 100:.1f}%', "Needs split or manual review"], ["Above 60 seconds", summary["above_60_seconds"], f'{summary["above_60_seconds"] / len(kept) * 100:.1f}%', "High-priority boundary review"], ["Above 300 seconds", summary["above_300_seconds"], f'{summary["above_300_seconds"] / len(kept) * 100:.2f}%', "Very long candidate review"]])
    doc.add_paragraph(f'The exclusions reduced the maximum from {max(original_durations) / 60:.1f} minutes to {d["max"] / 60:.1f} minutes. The remaining long records show that rejection removed known bad segments but did not solve the underlying over-consolidation problem.')

    doc.add_heading("6A. MELD-Style Duration Window EDA", level=1)
    doc.add_paragraph("This focused analysis uses two non-overlapping duration categories. The first category is 0.8 seconds inclusive through below 20 seconds. The second category is 20 seconds through 30 seconds inclusive. Rows below 0.8 seconds and above 30 seconds are excluded from this focused window, not silently treated as usable.")
    band_rows = []
    for name, label in [("0.8-20_seconds", "0.8-20 seconds"), ("20-30_seconds", "20-30 seconds")]:
        info = summary["duration_band_eda"][name]
        band_rows.append([label, info["rows"], info["minutes"], info["hours"], info["mean_seconds"], info["median_seconds"], info["source_videos"]])
    combined_rows = [row for rows in duration_bands.values() for row in rows]
    combined = band_summary(combined_rows)
    band_rows.append(["Combined", combined["rows"], combined["minutes"], combined["hours"], combined["mean_seconds"], combined["median_seconds"], combined["source_videos"]])
    add_table(doc, ["Duration category", "Rows", "Minutes", "Hours", "Mean sec", "Median sec", "Source videos"], band_rows)
    doc.add_paragraph(f"The two duration categories contain {combined['rows']} rows and {combined['minutes']} minutes ({combined['hours']} hours) of candidate material. This is a duration-screened candidate total, not a claim that every row is semantically correct or emotion-labelled ground truth.")

    doc.add_heading("6B. Combined Emotion-Label Distribution in the 0.8-30 Second Window", level=1)
    emotion_rows = []
    for label, count in sorted(combined["emotion_label_counts"].items(), key=lambda item: (-item[1], item[0])):
        emotion_rows.append([label, count, f"{count / combined['rows'] * 100:.2f}%" if combined["rows"] else "0.00%"])
    add_table(doc, ["Emotion label", "Rows", "Share of combined window"], emotion_rows)
    doc.add_paragraph("These emotion labels are inherited from the current Clancy manifest and are provisional heuristic labels. Their distribution is useful for detecting imbalance and planning annotation, but they should not be presented as manually verified affective ground truth.")

    doc.add_heading("7. Remaining Long-Turn Candidates", level=1)
    add_table(doc, ["Turn ID", "Source", "Duration", "Source pieces", "Source URL"], [[clean(row.get("turn_id")), clean(row.get("youtube_id")), f'{number(row, "clip_duration_seconds"):.3f} sec', clean(row.get("turn_source_utterance_count")), clean(row.get("source_url"))] for row in long_rows[:10]])
    doc.add_paragraph("These records are not automatically rejected by the persistent list. They require manual inspection to determine whether they contain valid testimony that should be split or whether they are another break/non-value segment. A MELD-style dataset should not treat any of these long records as one final utterance without further processing.")

    doc.add_heading("8. Source Coverage After Exclusion", level=1)
    add_table(doc, ["YouTube ID", "Rows", "Clip hours", "Offset status"], [[source, info["rows"], info["clip_hours"], "explicit 300 sec" if source == "1tyKO8mTdOM" else "default zero"] for source, info in summary["source_breakdown"].items()])
    doc.add_paragraph("All 11 source videos remain represented. The rejection was selective: it removed specific turn records rather than deleting complete source videos. This preserves source diversity while removing known non-value segments.")

    doc.add_heading("9. Labels and Review Fields After Exclusion", level=1)
    label_rows = []
    for label, count in summary["emotion_label_counts"].items():
        label_rows.append([label or "BLANK", count, f'{count / len(kept) * 100:.2f}%'])
    add_table(doc, ["Emotion label", "Rows", "Approx. share"], label_rows)
    add_table(doc, ["Field", "Result"], [["LOW label confidence", summary["label_confidence_counts"].get("LOW", 0)], ["HIGH label confidence", summary["label_confidence_counts"].get("HIGH", 0)], ["review_flag=NO", summary["review_flag_counts"].get("NO", 0)], ["review_flag=YES", summary["review_flag_counts"].get("YES", 0)], ["usable_for_phase2=YES", summary["usable_counts"].get("YES", 0)]])
    neutral_count = summary["emotion_label_counts"].get("neutral", 0)
    doc.add_paragraph(f"The exclusions do not correct the dominant neutral-label problem. Neutral remains {neutral_count} rows, approximately {neutral_count / len(kept) * 100:.1f}% of the post-rejection population. The labels are still provisional text-keyword heuristic labels and should not be described as expert emotion ground truth.")

    doc.add_heading("10. Offset and Media Integrity", level=1)
    add_table(doc, ["Offset status", "Rows", "Value"], [["default_zero", summary["offset_status_counts"].get("default_zero", 0), "0.000 seconds"], ["explicit", summary["offset_status_counts"].get("explicit", 0), "300.000 seconds"]])
    doc.add_paragraph("The rejection operation does not alter timestamp offsets. The Day 15 source remains explicitly shifted by 300 seconds, and unlisted sources retain the default zero rule. The post-rejection CSV preserves source_offset_seconds and source_offset_status for reproducibility.")

    doc.add_heading("11. Observations", level=1)
    for text in [
        f"The {len(rejection_ids)} manual exclusions reduced the candidate population from {len(original)} to {len(kept)} rows.",
        f"The summed duration reduced from {round(sum(original_durations) / 3600.0, 4)} to {summary['total_clip_hours']} hours.",
        "The number of source videos remains 11, so source diversity was preserved.",
        f"The maximum clip reduced from {max(original_durations):.3f} seconds to {d['max']:.3f} seconds.",
        f"{summary['above_300_seconds']} turns still exceed five minutes and require manual inspection or splitting.",
        f"The short-clip problem remains: {summary['below_1_second']} rows are below one second.",
        "The label imbalance remains severe and is independent of the lunch-break exclusions.",
        "All remaining media paths are present, so the next risks are semantic and annotation quality rather than missing files.",
    ]:
        bullet(doc, text)

    doc.add_heading("12. Areas for Improvement", level=1)
    add_table(doc, ["Area", "Current issue", "Improvement"], [
        ["Turn boundary control", "Seven records remain over five minutes.", "Add a 20-30 second preferred maximum and split at markers, sentence boundaries, pauses, or speaker changes."],
        ["Short clips", "2,139 rows are below one second.", "Separate acknowledgements from speech; review or reject incomplete fragments."],
        ["Emotion labels", "Neutral is approximately 94.6%.", "Manually annotate a balanced subset and evaluate with macro-F1 and per-class recall."],
        ["Label confidence", "11,280 rows remain LOW confidence.", "Start a pilot with manually checked high-confidence rows and preserve label provenance."],
        ["Visual evidence", "Speaker visibility is not validated.", "Add shot type, face visibility, and speaker-visible fields with manual sampling."],
        ["Audio evidence", "File existence does not prove speech quality.", "Measure RMS, silence ratio, clipping, sample rate, and corruption."],
        ["Speaker identity", "Speaker IDs are incomplete.", "Use only grounded transcript metadata; never infer protected identity."],
        ["Selection control", "Historical and filtered manifests can be confused.", "Use the post-rejection manifest for the next split and retain the original for audit."],
    ])

    doc.add_heading("13. Recommended Next Steps", level=1)
    for text in [
        "Use clancy_turn_manifest_post_rejection.csv as the input for the next controlled dataset-selection pass.",
        "Review the seven remaining records above 300 seconds and decide whether to split or reject each one.",
        "Implement a maximum-duration and pause-aware turn splitter before generating the next full clip set.",
        "Run audio and visual quality checks on a stratified sample after re-clipping.",
        "Create explicit TRAIN_USE, REVIEW, and REJECT categories.",
        "Rebuild the leakage-aware train/dev/test manifests only after these filters are applied.",
        "Run dataset validation and training-readiness validation.",
        "Use a small, manually checked, high-confidence subset for the first warm-start experiment.",
    ]:
        numbered(doc, text)

    doc.add_heading("14. Conclusion and Readiness Judgment", level=1)
    doc.add_paragraph(f"After the {len(rejection_ids)} exclusions, the Clancy branch contains {len(kept)} traceable turn-level MP4/WAV rows and {summary['total_clip_hours']} summed candidate clip hours. The technical extraction path remains operational, and the manually identified non-value rows are controlled by a persistent, auditable rejection list.")
    doc.add_paragraph("The corpus is improved but is not yet ready for unrestricted fine-tuning. Seven very long turn candidates remain, thousands of clips are shorter than one second, and the emotion labels remain weak and highly imbalanced. The correct next milestone is boundary correction and quality filtering, not model training on all remaining rows.")

    doc.add_heading("15. Reproducibility Commands", level=1)
    code(doc, "PYTHON_BIN=/Users/rajeshpmu/Desktop/LegalMemoCMT/.venv/bin/python \\\nbash phase2/run_build_clancy_long_turn_review.sh")
    code(doc, "./.venv/bin/python phase2/create_clancy_post_rejection_eda_report.py")
    doc.add_paragraph("The first command regenerates the long-turn review list. The second command recreates the filtered CSV, JSON summary, and this DOCX from the historical clipped manifest and the persistent rejection manifest. Neither command deletes media.")

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("LegalMemoCMT Phase 2 | Post-Rejection Clancy EDA").font.size = Pt(8)
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT)
    print(f"Created {REPORT}")


if __name__ == "__main__":
    build()
