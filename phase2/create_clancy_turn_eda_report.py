from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("implementation_docments/LegalMemoCMT_Clancy_Turn_Manifest_EDA_Student_Report.docx")


def shade(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    props.append(node)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = str(value)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, "D9EAF7")
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = str(value)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def code(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Intense Quote")


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    doc.styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LegalMemoCMT Phase 2\nClancy Turn Manifest EDA Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Student-level technical analysis of the generated turn-level multimodal corpus")
    run.italic = True
    run.font.size = Pt(12)

    add_table(doc, ["Item", "Value"], [
        ["Input artifact", "data/processed/phase2/clancy/clancy_turn_manifest_clipped.csv"],
        ["Corpus branch", "Clancy primary benchmark branch of Phase 2"],
        ["Preparation level", "Turn-level MP4 and WAV clips with transcript metadata"],
        ["EDA type", "Read-only structural, media-path, duration, source, split, offset, and label analysis"],
        ["Interpretation", "Technical clip generation passed; scientific training readiness still requires improvement"],
    ])

    doc.add_heading("1. Purpose of This EDA", level=1)
    doc.add_paragraph(
        "This report examines the output of the Clancy turn-level clip-generation stage. The purpose is to determine what was actually created, whether the manifest is internally consistent, how much media was generated, and whether the rows are suitable for the next Phase 2 dataset-preparation stage."
    )
    doc.add_paragraph(
        "The analysis separates two questions. First, did the software create the requested files and record their locations correctly? Second, are the resulting samples good training examples? A PASS on the first question does not automatically mean a PASS on the second question."
    )
    doc.add_paragraph(
        "This distinction is important for a student explanation: the pipeline can complete without errors while still producing samples that require review because a turn is too long, a label is weak, or the visible speaker is unknown."
    )

    doc.add_heading("2. Input and Pipeline Context", level=1)
    doc.add_paragraph(
        "The input manifest is produced by the Clancy turn-manifest builder. The builder groups suitable adjacent subtitle fragments into turn-level records. The turn clipper then uses the start and end times to create a turn-level MP4 and a matching WAV. It also carries forward source URL, subtitle path, emotion-label fields, split fields, and source-offset fields."
    )
    doc.add_paragraph("The processing sequence is:")
    for text in [
        "Raw Clancy video, audio, and subtitle files are already downloaded.",
        "Subtitle cues are converted into turn-level metadata.",
        "The source-specific offset is added to subtitle-relative timestamps.",
        "ffmpeg extracts a turn-level MP4 and WAV.",
        "The output manifest records both source paths and generated clip paths.",
        "This EDA checks the resulting rows before dataset splitting and fine-tuning readiness validation.",
    ]:
        bullet(doc, text)
    doc.add_paragraph(
        "The source offset rule is: an explicit valid offset is applied when present; a missing, blank, invalid, or unlisted offset uses 0.000 seconds. The current explicit Day 15 offset is 300.000 seconds."
    )

    doc.add_heading("3. How the EDA Was Calculated", level=1)
    doc.add_paragraph("The analysis used the following definitions:")
    add_table(doc, ["Measure", "Calculation", "Why it matters"], [
        ["Rows", "Number of CSV data rows excluding the header", "Number of generated turn records"],
        ["Unique turn IDs", "Count of distinct turn_id values", "Checks whether one turn is duplicated"],
        ["Source videos", "Count of distinct youtube_id values", "Measures source diversity and split granularity"],
        ["Clip hours", "sum(clip_duration_seconds) / 3600", "Total generated duration, not automatically usable hours"],
        ["Missing media", "Count of missing video_path or audio_path files", "Checks manifest-to-file integrity"],
        ["Split leakage", "Check whether one youtube_id appears in multiple partitions", "Prevents source context leaking into evaluation"],
        ["Offset status", "Count of explicit versus default_zero rows", "Shows how timestamp correction was applied"],
    ])

    doc.add_heading("4. Overall Corpus Results", level=1)
    add_table(doc, ["Metric", "Result", "Student interpretation"], [
        ["Turn rows", "11,926", "11,926 turn-level records were written"],
        ["Unique turn IDs", "11,926", "No turn ID duplication was detected"],
        ["Unique utterance IDs", "11,926", "Each output row has a distinct output identifier"],
        ["Unique source videos", "11", "The corpus uses 11 YouTube source groups"],
        ["Total generated clip duration", "31.4766 hours", "Sum of extracted turn durations; requires quality filtering"],
        ["MP4 files missing", "0", "Every row points to an existing video file"],
        ["WAV files missing", "0", "Every row points to an existing audio file"],
        ["Clip failures", "0", "ffmpeg extraction completed for all rows"],
    ])
    doc.add_paragraph(
        "The defensible statement is: the turn clip-generation stage created 11,926 traceable MP4/WAV pairs representing 31.4766 summed clip hours across 11 sources, with no missing referenced files. This is not the same as claiming 31.4766 hours of clean, validated training data."
    )

    doc.add_heading("5. Train, Development, and Test Distribution", level=1)
    add_table(doc, ["Split", "Rows", "Approximate share"], [
        ["Train", "8,858", "74.3%"], ["Development", "1,491", "12.5%"], ["Test", "1,577", "13.2%"], ["Total", "11,926", "100%"],
    ])
    doc.add_paragraph(
        "The split counts are approximately consistent with a 70/15/15 target, but exact percentages depend on whole-source grouping. A complete YouTube source is kept in one partition rather than splitting individual turns randomly. This is important because adjacent turns share room acoustics, camera position, legal vocabulary, and background speakers."
    )
    doc.add_paragraph(
        "The split is leakage-aware at source-group level. It does not prove speaker identity disjointness because the current Clancy transcript metadata does not reliably identify every speaker."
    )

    doc.add_heading("6. Clip-Duration EDA", level=1)
    add_table(doc, ["Statistic", "Value", "Meaning"], [
        ["Minimum", "0.600 seconds", "Configured minimum-duration floor"],
        ["25th percentile", "1.633 seconds", "One quarter of clips are at or below this duration"],
        ["Median", "4.249 seconds", "Half of clips are shorter than approximately 4.25 seconds"],
        ["Mean", "9.502 seconds", "Raised by several very long consolidated turns"],
        ["75th percentile", "9.328 seconds", "Three quarters are below approximately 9.33 seconds"],
        ["95th percentile", "26.455 seconds", "Most clips are within a practical short-clip range"],
        ["Maximum", "5,131.167 seconds", "One clip is approximately 85.5 minutes and is not one valid turn"],
    ])
    add_table(doc, ["Duration condition", "Rows", "Approx. share", "Interpretation"], [
        ["Below 1 second", "2,139", "17.9%", "May be acknowledgements or overly short fragments; review needed"],
        ["Above 30 seconds", "477", "4.0%", "May be valid speech, but check for over-merging"],
        ["Above 60 seconds", "144", "1.2%", "High-priority review group"],
    ])
    doc.add_paragraph(
        "The maximum and upper-tail values are the most important finding. A turn-level sample should represent one coherent speaker turn, not a long sequence of courtroom events. The presence of an 85.5-minute record strongly suggests that the consolidation heuristic merged cues across pauses, recesses, or speaker changes."
    )

    doc.add_heading("7. Examples of Over-Merged Turns", level=1)
    add_table(doc, ["Turn ID", "Source", "Duration", "Pieces", "Why it needs review"], [
        ["D6skXbDJJj8_turn05552", "D6skXbDJJj8", "5,131.167 sec", "17", "Approximately 85.5 minutes; not one utterance"],
        ["D6skXbDJJj8_turn05148", "D6skXbDJJj8", "1,865.292 sec", "9", "Approximately 31.1 minutes; likely multiple segments"],
        ["93sfJyLzhqM_turn02128", "93sfJyLzhqM", "1,865.031 sec", "9", "Approximately 31.1 minutes; likely over-consolidated"],
        ["D6skXbDJJj8_turn04400", "D6skXbDJJj8", "1,613.380 sec", "15", "Approximately 26.9 minutes; boundary inspection required"],
    ])
    doc.add_paragraph(
        "These examples show why a successful ffmpeg run is not enough. The files exist and can be played, but their semantic unit is too large for a MELD-style utterance corpus. The rows should be split at reliable subtitle turn markers, sentence boundaries, long pauses, or detected speaker changes before training."
    )

    doc.add_heading("8. Source Coverage and Generated Duration", level=1)
    add_table(doc, ["YouTube ID", "Rows", "Clip hours", "Offset"], [
        ["peYoD-JkAXE", "3,043", "5.8210", "0.000 sec"], ["93sfJyLzhqM", "1,924", "5.9103", "0.000 sec"],
        ["D6skXbDJJj8", "1,615", "6.3446", "0.000 sec"], ["1tyKO8mTdOM", "1,284", "3.6559", "300.000 sec"],
        ["DCBWoWhsTpA", "1,136", "2.9646", "0.000 sec"], ["D5L_c9Mla1U", "1,077", "2.8428", "0.000 sec"],
        ["ZJd1Lk4w-qo", "992", "2.0472", "0.000 sec"], ["sHUdRcABC-Q", "291", "0.6010", "0.000 sec"],
        ["ma6ni-5zuW0", "209", "0.2826", "0.000 sec"], ["VQQb9BbgwHg", "208", "0.7482", "0.000 sec"],
        ["LGYerNiA8kI", "147", "0.2584", "0.000 sec"],
    ])
    doc.add_paragraph(
        "The source breakdown shows recording-level diversity, but it is not equivalent to 11 independent cases or 11 independent witness populations. Each YouTube ID is a source group, and source-level grouping is currently the strongest available leakage-control unit."
    )

    doc.add_heading("9. Emotion Label EDA", level=1)
    add_table(doc, ["Label", "Rows", "Approx. share", "Interpretation"], [
        ["neutral", "11,288", "94.65%", "Dominant heuristic class"], ["confidence", "360", "3.02%", "Small minority class"],
        ["anger", "114", "0.96%", "Small class requiring review"], ["stress", "86", "0.72%", "Small class requiring review"],
        ["fear", "65", "0.54%", "Very small class"], ["sadness", "13", "0.11%", "Extremely sparse class"],
    ])
    doc.add_paragraph(
        "The labels are provisional text-keyword heuristic labels. They are not equivalent to expert human emotion annotations or validated facial-expression labels. The extreme neutral imbalance means that a model could achieve superficially good accuracy by predicting neutral too often."
    )
    add_table(doc, ["Field", "Result"], [
        ["LOW label confidence", "11,288 rows"], ["HIGH label confidence", "638 rows"],
        ["review_flag=NO", "11,288 rows"], ["review_flag=YES", "638 rows"], ["usable_for_phase2=YES", "11,926 rows"],
    ])
    doc.add_paragraph(
        "The usable flag is not a final scientific quality label. It is an operational selection flag inherited from the manifest and must be refined using duration, alignment, audio, video visibility, and label-review rules."
    )

    doc.add_heading("10. Source Offset EDA", level=1)
    add_table(doc, ["Status", "Rows", "Value", "Explanation"], [
        ["default_zero", "10,642", "0.000 seconds", "No exclusive verified offset configured"],
        ["explicit", "1,284", "300.000 seconds", "Day 15 correction for 1tyKO8mTdOM"],
    ])
    doc.add_paragraph(
        "This field is important for reproducibility. The manifest records both the numerical offset and whether it came from an explicit configuration or the default rule."
    )

    doc.add_heading("11. Technical Observations", level=1)
    for text in [
        "All 11,926 requested turn rows produced an MP4 and a WAV.",
        "The manifest retains source paths, clipped paths, source IDs, turn IDs, timestamps, duration fields, and offset fields.",
        "There are no duplicate turn IDs or duplicate utterance IDs in this output.",
        "The split is leakage-aware by YouTube source group.",
        "The duration distribution shows that consolidation sometimes creates samples much larger than one turn.",
        "The labels are severely imbalanced and mostly low-confidence heuristic labels.",
        "The generated 31.4766 hours is raw candidate clip duration, not final validated training duration.",
        "Reliable witness facial visibility cannot be claimed from this manifest alone.",
    ]:
        bullet(doc, text)

    doc.add_heading("12. Areas Requiring Improvement", level=1)
    add_table(doc, ["Priority", "Problem", "Why it matters", "How to improve"], [
        ["1", "Over-merged long turns", "One row may contain multiple speakers or events.", "Add a 20-30 second preferred maximum; split at markers, sentences, long pauses, or speaker changes."],
        ["2", "Very short clips", "2,139 rows are below one second.", "Keep acknowledgements separately; require a minimum duration for initial training."],
        ["3", "Label imbalance", "94.65% are neutral; accuracy can mislead.", "Manually annotate a balanced subset and report macro-F1 and per-class recall."],
        ["4", "Low confidence labels", "11,288 rows are LOW confidence.", "Use HIGH-confidence rows first and manually review minority labels."],
        ["5", "Speaker uncertainty", "The speaking or visible participant is not always known.", "Use grounded subtitle metadata and manual review; do not invent identities."],
        ["6", "Visual validation missing", "A valid MP4 may show a judge, counsel, or wide shot.", "Add face, shot, and speaker-visibility checks."],
        ["7", "Audio quality screening", "Media may contain silence or clipping.", "Measure RMS, silence ratio, clipping ratio, sample rate, and corruption."],
        ["8", "Adjacent spillover", "Neighboring turns may repeat context.", "Check overlap and apply a controlled boundary policy."],
        ["9", "Source diversity", "11 videos do not mean 11 cases or witness populations.", "Track case, hearing, witness, and source counts separately."],
        ["10", "Manifest naming", "Cue and turn branches can be confused.", "Keep explicit turn-level filenames and separate output directories."],
    ])

    doc.add_heading("13. Recommended Improvement Sequence", level=1)
    for text in [
        "Implement and test a maximum turn-duration rule, beginning with a 30-second preferred maximum.",
        "Rebuild the turn manifest and clips into a separate pilot output directory for comparison.",
        "Rerun the duration EDA and compare over-30-second and over-60-second counts.",
        "Create TRAIN_USE, REVIEW, and REJECT categories instead of relying only on usable_for_phase2=YES.",
        "Run audio and video validation on the filtered rows.",
        "Manually inspect short clips, long clips, every source video, every split, and minority labels.",
        "Create the leakage-aware train/dev/test manifest only after filtering.",
        "Run dataset validation and training-readiness validation.",
        "Start with a small warm-start experiment using the highest-confidence manually checked subset.",
        "Report macro-F1, per-class recall, and a confusion matrix before expanding the dataset.",
    ]:
        numbered(doc, text)

    doc.add_heading("14. Student-Level Readiness Explanation", level=1)
    doc.add_paragraph(
        "From a software-engineering perspective, this stage is successful because the requested files were produced, the row-to-file references are valid, and ffmpeg reported no failures. From a machine-learning perspective, the stage is not fully ready because some rows do not represent a clean single turn and the labels are weak and highly imbalanced."
    )
    doc.add_paragraph(
        "The accurate explanation to a guide is that the end-to-end multimodal extraction path has been demonstrated at scale, and the EDA has identified the quality controls required before the corpus becomes a defensible fine-tuning dataset."
    )

    doc.add_heading("15. Conclusion", level=1)
    doc.add_paragraph(
        "The Clancy turn-level manifest contains 11,926 generated multimodal samples across 11 source videos and 31.4766 summed clip hours. Every row has a corresponding MP4 and WAV, and there are no detected duplicate IDs or missing media paths. This confirms that the source-to-clip engineering pipeline is operational."
    )
    doc.add_paragraph(
        "The main limitation is semantic quality rather than file generation. The extreme maximum duration shows that turn consolidation can merge multiple courtroom events into one record. The label distribution is also dominated by low-confidence neutral heuristic labels. The next milestone is quality-controlled turn splitting and annotation validation, followed by a new leakage-aware dataset manifest and a cautious warm-start experiment."
    )

    doc.add_heading("16. Reproducibility References", level=1)
    for text in [
        "Input artifact: data/processed/phase2/clancy/clancy_turn_manifest_clipped.csv",
        "Turn manifest builder: phase2/build_clancy_turn_manifest.py",
        "Turn clip builder: phase2/build_clancy_turn_clips.py",
        "Turn manifest wrapper: phase2/run_build_clancy_turn_manifest.sh",
        "Turn clip wrapper: phase2/run_build_clancy_turn_clips.sh",
        "Offset configuration: data/processed/phase2/clancy/clancy_source_offsets.csv",
        "Offset rule: missing, blank, invalid, or unlisted source offsets use 0.000 seconds.",
        "This report was generated from a read-only EDA and does not modify the input CSV or media files.",
    ]:
        doc.add_paragraph(text)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("LegalMemoCMT Phase 2 | Clancy Turn Manifest EDA").font.size = Pt(8)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Created {OUT}")


if __name__ == "__main__":
    build()
