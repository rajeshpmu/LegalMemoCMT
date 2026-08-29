"""Append the completed 200-row audio-SER pilot result to guidance DOCX files."""

from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
MARKER = "AUDIO_SER_PILOT_RESULT_V1"
TARGETS = [
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Witness_Speaking_Pipeline_Student_Speaking_Guide.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Phase2_Corpus_Build_Code_Level_Student_Guide.docx",
]


def main() -> None:
    changed = []
    for path in TARGETS:
        doc = Document(path)
        if any(MARKER in p.text for p in doc.paragraphs):
            continue
        backup = path.with_name(path.name + ".before_audio_ser_pilot_result.docx")
        if not backup.exists():
            shutil.copy2(path, backup)
        doc.add_heading("Completed 200-Row Audio-SER Pilot Result", level=1)
        p = doc.add_paragraph(); p.add_run(MARKER).font.size = Pt(8)
        doc.add_paragraph(
            "The pilot processed 200 rows successfully. Odyssey completed 200/200 rows, "
            "SpeechBrain completed 200/200 rows, all 200 audio paths existed, and all rows "
            "were status=OK. The output is data/processed/phase2/clancy/"
            "audio_ser_evidence_200.csv and the audit summary is reports/phase2/"
            "clancy_audio_ser_evidence_200.json."
        )
        doc.add_paragraph(
            "The sample is not corpus-wide: all 200 rows come from DCBWoWhsTpA and are in "
            "the dev split. Therefore this run proves the adapter and provenance path, but "
            "it does not estimate performance across all Clancy sources, speakers, or splits."
        )
        doc.add_paragraph(
            "SpeechBrain produced neu=151, ang=30, hap=17, and sad=2. Its confidence mean "
            "was approximately 0.993 and the median was 1.0, indicating saturation rather "
            "than calibrated courtroom confidence. The labels remain cross-check evidence. "
            "Odyssey produced continuous valence, arousal, and dominance-like outputs with "
            "means approximately 0.457, 0.372, and 0.446 respectively. These are model "
            "outputs in an approximate 0-to-1 scale, not physical units or gold labels."
        )
        doc.add_paragraph(
            "The next controlled step is to run a stratified sample across multiple source "
            "videos and splits, inspect disagreement with Phase 1 predictions and transcript "
            "scope, and calibrate or threshold the evidence before using it to prioritize "
            "human annotation. The full run should not overwrite canonical labels."
        )
        doc.save(path)
        changed.append(str(path))
    print({"updated": changed, "marker": MARKER})


if __name__ == "__main__":
    main()
