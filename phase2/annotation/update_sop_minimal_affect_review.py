"""Append the minimal courtroom-affect review strategy to the Clancy SOP."""
from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt

ROOT=Path(__file__).resolve().parents[2]
SOP=ROOT/"implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx"
MARKER="MINIMAL_AFFECT_REVIEW_QUEUE_V1"

def main():
    doc=Document(SOP)
    if any(MARKER in p.text for p in doc.paragraphs):
        print('Already updated',SOP); return
    backup=SOP.with_name(SOP.name+'.before_minimal_affect_review.docx')
    if not backup.exists(): shutil.copy2(SOP,backup)
    doc.add_heading("Minimal Courtroom-Affect Review Queue",level=1)
    p=doc.add_paragraph(); p.add_run(MARKER).font.size=Pt(8)
    doc.add_paragraph("Heuristics can reduce the number of clips reviewed first, but they cannot replace human validation of courtroom affect. The purpose of this queue is active-learning triage: review rare, conflicting, or weakly supported cases first, then use a small representative sample of the common class to check systematic bias.")
    doc.add_heading("Priority order",level=2)
    for text in [
        "Review every non-neutral candidate: ASSERTIVE, GUARDED, DEFENSIVE, TENSE, DISTRESSED, and AGITATED.",
        "Review UNKNOWN rows when disagreement is HIGH, arousal is elevated, or valence is strongly negative.",
        "Review all low-confidence NEUTRAL_CALM rows, especially rows with missing low/moderate-arousal evidence.",
        "Review a stratified sample of high-confidence NEUTRAL_CALM rows rather than reviewing every common row immediately.",
        "Use the remaining neutral rows later for calibration and quality-control sampling.",
    ]: doc.add_paragraph(text,style='List Bullet')
    doc.add_heading("Command",level=2)
    doc.add_paragraph("The first queue can be generated with:")
    p=doc.add_paragraph(style='Intense Quote'); p.add_run("./.venv/bin/python phase2/annotation/inspect_clancy_courtroom_affect_candidates.py \\\n  --input-csv data/processed/phase2/clancy/courtroom_affect_candidates_200.csv \\\n  --output-csv data/processed/phase2/clancy/review_non_neutral_affect.csv \\\n  --summary-json reports/phase2/clancy_review_non_neutral_affect.json \\\n  --affect ASSERTIVE,GUARDED,DEFENSIVE,TENSE,DISTRESSED,AGITATED \\\n  --print-rows 200")
    doc.add_paragraph("A lower-confidence neutral queue can be generated with --affect NEUTRAL_CALM --min-confidence 0.60 --print-rows 30. The output should be watched, listened to, and compared with the transcript before a human_basic_emotion, courtroom_affect, or affect_intensity value is entered.")
    doc.add_heading("Interpretation guardrail",level=2)
    doc.add_paragraph("A candidate confidence is rule-support strength, not emotional intensity and not probability of correctness. A small review queue is acceptable for active learning, but a row must not become a gold label simply because it was not selected for the first queue. The final human review record should retain the machine candidate, supporting evidence, missing evidence, reviewer decision, confidence, and notes. No deception, credibility, truthfulness, or reliability inference is permitted.")
    doc.save(SOP); print('Updated',SOP)

if __name__=='__main__': main()
