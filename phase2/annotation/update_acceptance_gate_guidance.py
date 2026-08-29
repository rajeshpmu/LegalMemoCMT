"""Append the automatic acceptance-gate procedure to Phase 2 DOCX guides."""
from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
MARKER = "AUTOMATIC_ACCEPTANCE_GATE_V1"
DOCS = [
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Witness_Speaking_Pipeline_Student_Speaking_Guide.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Phase2_Corpus_Build_Code_Level_Student_Guide.docx",
]

def update(path: Path) -> bool:
    if not path.exists(): return False
    doc = Document(path)
    if any(MARKER in p.text for p in doc.paragraphs): return False
    backup = path.with_name(path.name + ".before_acceptance_gate_v1.docx")
    if not backup.exists(): shutil.copy2(path, backup)
    doc.add_heading("Automatic Acceptance Gate for Machine-Assisted Labels", level=1)
    p = doc.add_paragraph(); p.add_run(MARKER).font.size = Pt(8)
    doc.add_paragraph(
        "The candidate generator and the acceptance gate are separate stages. The generator "
        "proposes basic-emotion and courtroom-affect candidates while preserving Phase 1 output. "
        "The gate decides whether a row is sufficiently supported for a machine-assisted SILVER "
        "tier, or whether it must remain UNRESOLVED for human review. It does not create gold labels."
    )
    doc.add_heading("Gate rule", level=2)
    doc.add_paragraph(
        "A row is AUTO_ADJUDICATED with annotation_tier=SILVER only when "
        "basic_emotion_review_candidate_confidence >= 0.70, "
        "proposed_courtroom_affect_confidence >= 0.60, both candidates are resolved, and no "
        "critical conflict is detected. Otherwise it receives annotation_status=UNRESOLVED and "
        "annotation_tier=WEAK. The gate writes final_basic_emotion and final_courtroom_affect "
        "only for accepted rows; unresolved rows retain a usable affect candidate where supported "
        "but set final_basic_emotion=UNRESOLVED."
    )
    doc.add_heading("Worked examples", level=2)
    doc.add_paragraph(
        "DCBWoWhsTpA_turn06801 passes: basic-emotion review candidate neutral at 0.75, "
        "CALM_COMPOSED at 0.60, QUOTED_SPEECH scope, and no critical conflict. It becomes "
        "AUTO_ADJUDICATED/SILVER for machine-assisted training experiments, not human gold data."
    )
    doc.add_paragraph(
        "DCBWoWhsTpA_turn06570 does not pass: its basic candidate remains the Phase 1 disgust "
        "candidate at only 0.47, while HESITANT_UNCERTAIN is the courtroom-affect candidate at "
        "0.75. It becomes UNRESOLVED/WEAK, with final_basic_emotion=UNRESOLVED and the affect "
        "candidate retained for review. This prevents a strong affect score from hiding weak basic "
        "emotion evidence."
    )
    doc.add_heading("Reproducible command", level=2)
    doc.add_paragraph(
        "./.venv/bin/python phase2/annotation/apply_clancy_annotation_acceptance_gate.py \\\n+  --input-csv data/processed/phase2/clancy/courtroom_affect_candidates_200_v4.csv \\\n+  --output-csv data/processed/phase2/clancy/courtroom_affect_candidates_200_v4_gated.csv \\\n+  --summary-json reports/phase2/clancy_courtroom_affect_candidates_200_v4_gated.json \\\n+  --basic-threshold 0.70 --affect-threshold 0.60"
    )
    doc.add_paragraph(
        "The output must be reported with counts of AUTO_ADJUDICATED/SILVER, UNRESOLVED/WEAK, "
        "and critical conflicts. Thresholds are review-policy parameters, not calibrated model "
        "probabilities; they should be evaluated against a human-reviewed sample before scaling."
    )
    doc.save(path); return True

def main():
    print({"updated_documents": [str(p) for p in DOCS if update(p)], "marker": MARKER})

if __name__ == "__main__": main()
