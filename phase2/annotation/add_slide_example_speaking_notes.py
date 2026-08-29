"""Append detailed speaking notes for the Clancy pipeline example callouts."""
from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt

ROOT=Path(__file__).resolve().parents[2]
TARGETS=[
 ROOT/"implementation_docments/LegalMemoCMT_Clancy_Witness_Speaking_Pipeline_Student_Speaking_Guide.docx",
 ROOT/"implementation_docments/LegalMemoCMT_Phase2_Corpus_Build_Code_Level_Student_Guide.docx",
]
MARKER="SLIDE_EXAMPLE_SPEAKING_NOTES_V1"

def main():
 changed=[]
 for path in TARGETS:
  doc=Document(path)
  if any(MARKER in p.text for p in doc.paragraphs): continue
  backup=path.with_name(path.name+".before_slide_example_speaking_notes.docx")
  if not backup.exists(): shutil.copy2(path,backup)
  doc.add_heading("Detailed Speaking Notes for Slides 13-18", level=1)
  p=doc.add_paragraph(); p.add_run(MARKER).font.size=Pt(8)
  doc.add_paragraph("The following notes explain the example callout on each slide. I should explain the field, its source, its meaning, and its limitation rather than merely reading the value from the slide.")
  sections=[
   ("Slide 13: Independent Audio SER Evidence", "Example: DCBWoWhsTpA_turn06575 has audio-SER fields appended without changing the Phase 1 label; Odyssey arousal=0.37817.", [
    "utterance_id identifies one exact witness turn and is the join key across manifests.",
    "audio-SER fields were produced from the corresponding WAV clip, not from the transcript alone.",
    "audio_arousal=0.37817 is a model-specific continuous output on an approximate 0-to-1 scale. It is evidence of activation in the voice, not a physical measurement.",
    "The phrase without changing the Phase 1 label means phase1_basic_emotion remains preserved. The independent model is an additional evidence column, not a replacement label.",
    "I should not claim that low arousal proves neutral emotion; it only motivates a controlled-delivery hypothesis for review."
   ]),
   ("Slide 14: Conservative Fusion and Review", "Example: DCBWoWhsTpA_turn06575 shows Phase 1=fear, SpeechBrain=neu, and scope=OTHER_PERSON_DESCRIBED.", [
    "Phase 1 fear is the original trimodal MELD-transfer prediction. It reflects the checkpoint output and must remain auditable.",
    "SpeechBrain neu is an audio-only categorical cross-check. It is not identical to a human neutral label and its confidence is not assumed calibrated.",
    "OTHER_PERSON_DESCRIBED means words such as hopeless and suicidal ideation refer to another person in the testimony, not necessarily to the speaking witness.",
    "The disagreement is useful because it reveals semantic-content leakage: the text is negative while the speaker's delivery can remain controlled.",
    "The correct action is human multimodal review of transcript, WAV, and MP4, not automatic deletion of the fear prediction."
   ]),
   ("Slide 15: Corrected Active-Review Policy", "Example: DCBWoWhsTpA_turn06057 has Phase 1=fear at 0.897316, audio=neu, and scope=UNCLEAR.", [
    "The high Phase 1 confidence does not make fear ground truth; it only tells me the source checkpoint is confident.",
    "UNCLEAR means the transparent scope heuristic found no reliable target. It prevents me from claiming that the emotion belongs to another person.",
    "Because there is no strong leakage evidence, this row is a disagreement-review example rather than a justified neutral correction.",
    "The corrected policy therefore allows an UNRESOLVED outcome when used with --policy conservative. The legacy policy may still produce a neutral candidate, but that candidate is explicitly provisional."
   ]),
   ("Slide 16: Scope-Aware Candidate Rules", "Example: DCBWoWhsTpA_turn06575 has audio=neu, arousal=0.37817, and proposed neutral at 0.720000.", [
    "The neutral proposal is supported by multiple conditions: comparable neutral audio, low/moderate arousal, and a non-self emotion target.",
    "The confidence 0.720000 is a rule-based proposal confidence, not a calibrated probability from a trained classifier.",
    "proposed_basic_emotion=neutral is kept separate from human_basic_emotion. A human reviewer must confirm or reject it.",
    "proposed_courtroom_affect=CALM_COMPOSED describes the witness's controlled presentation. It can coexist with sadness or another basic emotion and does not imply credibility."
   ]),
   ("Slide 17: Agreement, Disagreement, and Confidence", "Example: DCBWoWhsTpA_turn05916 has Phase 1=neutral and SpeechBrain=ang; neither output is overwritten.", [
    "The two predictions disagree after mapping SpeechBrain ang to anger. This is a comparable categorical conflict.",
    "The transcript is a short identification response, so audio or recording artifacts may influence the classifier. The example must be watched and listened to, not interpreted from labels alone.",
    "Preserving both values allows later evaluation of which signal was more useful and prevents circular relabeling.",
    "The priority score is for ordering manual work. It is not an accuracy score and does not establish that either prediction is correct."
   ]),
   ("Slide 18: Scope-Aware Pilot EDA", "Example: the pilot has 200 rows, 178 proposed neutral, 132 NEUTRAL_CANDIDATE, and 2 UNRESOLVED outcomes; all rows are from DCBWoWhsTpA/dev.", [
    "200 rows is the number processed in the pilot, not the size of the full Clancy corpus.",
    "All rows share one source and one split, so this EDA cannot measure generalization across videos, witnesses, or train/test partitions.",
    "178 proposed neutral shows the effect of the rule layer, not a human prevalence estimate. The original Phase 1 neutral count was 163.",
    "132 NEUTRAL_CANDIDATE rows require interpretation according to policy; they are not automatically accepted for fine-tuning.",
    "The 2 UNRESOLVED rows demonstrate that the pipeline can refuse to make a categorical recommendation when the evidence is insufficient.",
    "The conclusion I should state is that this is a reproducible active-review artifact and a leakage diagnostic, not a finished gold-label dataset."
   ]),
  ]
  for title, example, points in sections:
   doc.add_heading(title, level=2); doc.add_paragraph(example)
   for point in points: doc.add_paragraph(point, style="List Bullet")
  doc.add_heading("How to explain the overall method", level=2)
  doc.add_paragraph("I can explain the evidence chain as: one utterance_id joins the original Phase 1 prediction, the audio-only models, transcript scope heuristics, and the proposed review outcome. The system separates what the witness says, how the witness sounds, what the models predict, and what a human finally decides. This separation is the key protection against treating negative courtroom subject matter as the witness's own emotion.")
  doc.save(path); changed.append(str(path))
 print({"updated":changed,"marker":MARKER})

if __name__=="__main__": main()
