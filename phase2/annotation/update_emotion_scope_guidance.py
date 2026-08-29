"""Document the transcript-scope and cross-modal review layer."""
from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
MARKER = "EMOTION_SCOPE_REVIEW_V1"
TARGETS = [
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Witness_Speaking_Pipeline_Student_Speaking_Guide.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Phase2_Corpus_Build_Code_Level_Student_Guide.docx",
]


def main() -> None:
    changed = []
    for path in TARGETS:
        doc = Document(path)
        if any(MARKER in p.text for p in doc.paragraphs):
            continue
        backup = path.with_name(path.name + ".before_emotion_scope_review.docx")
        if not backup.exists():
            shutil.copy2(path, backup)
        doc.add_heading("Emotion Target Scope and Cross-Modal Review", level=1)
        p = doc.add_paragraph(); p.add_run(MARKER).font.size = Pt(8)
        doc.add_paragraph(
            "build_clancy_emotion_scope_review.py joins the Phase 1 pseudo-label manifest "
            "and the independent audio-SER manifest by utterance_id. It detects explicit "
            "emotion words, external-person references, reported speech, and event terms. "
            "It then compares the Phase 1 emotion with SpeechBrain and Odyssey evidence."
        )
        doc.add_paragraph(
            "The output fields are semantic_emotion_present, emotion_target_scope, "
            "modality_disagreement_score, semantic_leakage_risk, annotation_priority, "
            "review_flag, and review_reason. It also writes proposed_basic_emotion and "
            "proposed_courtroom_affect as explicitly provisional review recommendations. "
            "These fields never overwrite the original model outputs or human labels."
        )
        doc.add_paragraph(
            "The conservative rule is important: when the transcript contains negative "
            "emotion language about another person and the audio is neutral with low or "
            "moderate arousal, the script proposes neutral and CALM_COMPOSED for review. "
            "This addresses emotion-target confusion, where a witness describes someone "
            "else's distress in a controlled professional voice. The rule is not a gold-label "
            "generator; HIGH and MEDIUM priority rows require human multimodal review."
        )
        doc.add_paragraph(
            "Pilot command: ./.venv/bin/python phase2/annotation/"
            "build_clancy_emotion_scope_review.py --phase1-csv "
            "data/processed/phase2/clancy/phase1_trimodal_pseudo_labels_200.csv "
            "--audio-ser-csv data/processed/phase2/clancy/audio_ser_evidence_200.csv "
            "--output-csv data/processed/phase2/clancy/emotion_scope_review_200.csv "
            "--summary-json reports/phase2/clancy_emotion_scope_review_200.json --max-rows 200"
        )
        doc.add_paragraph(
            "The 200-row review artifact found 24 HIGH, 151 MEDIUM, and 25 LOW priority "
            "rows. The example DCBWoWhsTpA_turn06575 is classified as "
            "OTHER_PERSON_DESCRIBED with HIGH semantic leakage risk: Phase 1 predicted fear, "
            "SpeechBrain predicted neu, and the conservative recommendation is neutral plus "
            "CALM_COMPOSED. This is retained as a disagreement case for annotation."
        )
        doc.save(path)
        changed.append(str(path))
    print({"updated": changed, "marker": MARKER})


if __name__ == "__main__":
    main()
