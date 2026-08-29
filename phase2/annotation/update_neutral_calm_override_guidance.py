"""Document the narrow neutral-calm auto-adjudication override."""
from pathlib import Path
import shutil

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
MARKER = "NEUTRAL_CALM_OVERRIDE_V1_GUIDANCE"
DOCS = [
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Witness_Speaking_Pipeline_Student_Speaking_Guide.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Phase2_Corpus_Build_Code_Level_Student_Guide.docx",
]

def update(path: Path) -> bool:
    if not path.exists(): return False
    doc = Document(path)
    if any(MARKER in p.text for p in doc.paragraphs): return False
    backup = path.with_name(path.name + ".before_neutral_calm_override_v1.docx")
    if not backup.exists(): shutil.copy2(path, backup)
    doc.add_heading("Narrow Neutral-Calm Auto-Adjudication Rule", level=1)
    p = doc.add_paragraph(); p.add_run(MARKER).font.size = Pt(8)
    doc.add_paragraph(
        "A conservative special case is now available for weak non-neutral Phase 1 predictions. "
        "It is not a general rule that converts low-confidence emotions to neutral. It applies "
        "only when Phase 1 is non-neutral with confidence below 0.60, negative_activation_candidate "
        "is NO, distress_corroboration_present is NO, speaker_emotion_evidence_present is NO, "
        "the courtroom-affect candidate is CALM_COMPOSED, and its confidence is at least 0.80."
    )
    doc.add_paragraph(
        "When every condition is true, the gate writes final_basic_emotion=neutral, "
        "annotation_status=AUTO_ADJUDICATED, and annotation_tier=SILVER. The original Phase 1 "
        "emotion and confidence remain unchanged, so the domain-shift disagreement remains auditable."
    )
    doc.add_heading("Why the rule is safe", level=2)
    for text in [
        "It requires positive evidence of a calm/composed courtroom presentation, not merely a weak Phase 1 prediction.",
        "It excludes negative activation and distress corroboration, reducing the risk of hiding a genuinely activated delivery.",
        "It requires no independent evidence that the current speaker is expressing the emotion, which is appropriate for reported or quoted content.",
        "It creates a SILVER machine-assisted label, not a human gold annotation; the gate thresholds still require later calibration against reviewed data.",
    ]:
        doc.add_paragraph(text, style="List Bullet")
    doc.add_heading("Important non-example", level=2)
    doc.add_paragraph(
        "DCBWoWhsTpA_turn06570 should not pass this override. Although its Phase 1 disgust confidence "
        "is below 0.60, its affect candidate is HESITANT_UNCERTAIN rather than CALM_COMPOSED and its "
        "basic candidate confidence is only 0.47. It therefore remains UNRESOLVED/WEAK."
    )
    doc.add_heading("Reproducible gate command", level=2)
    doc.add_paragraph(
        "./.venv/bin/python phase2/annotation/apply_clancy_annotation_acceptance_gate.py \\\n+  --input-csv data/processed/phase2/clancy/courtroom_affect_candidates_200_v4.csv \\\n+  --output-csv data/processed/phase2/clancy/courtroom_affect_candidates_200_v5_gated.csv \\\n+  --summary-json reports/phase2/clancy_courtroom_affect_candidates_200_v5_gated.json \\\n+  --basic-threshold 0.70 --affect-threshold 0.60"
    )
    doc.save(path); return True

def main():
    print({"updated_documents": [str(p) for p in DOCS if update(p)], "marker": MARKER})

if __name__ == "__main__": main()
