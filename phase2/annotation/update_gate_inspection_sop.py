"""Append acceptance-gate inspection commands to the Clancy SOP."""
from pathlib import Path
import shutil

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
SOP = ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx"
MARKER = "ACCEPTANCE_GATE_INSPECTION_SCRIPT_V1"


def main() -> None:
    doc = Document(SOP)
    if any(MARKER in paragraph.text for paragraph in doc.paragraphs):
        print("Already updated", SOP)
        return
    backup = SOP.with_name(SOP.name + ".before_gate_inspection_v1.docx")
    if not backup.exists():
        shutil.copy2(SOP, backup)

    doc.add_heading("Inspecting Acceptance-Gate Review Queues", level=1)
    marker = doc.add_paragraph(); marker.add_run(MARKER).font.size = Pt(8)
    doc.add_paragraph(
        "After applying the automatic acceptance gate, the gated CSV contains both machine-"
        "accepted SILVER rows and unresolved WEAK rows. Before any human annotation or training, "
        "inspect the unresolved rows and the critical conflicts as readable review records. This "
        "does not alter the gated CSV and does not turn a machine candidate into a gold label."
    )
    doc.add_heading("Reusable inspection script", level=2)
    doc.add_paragraph(
        "phase2/annotation/inspect_clancy_acceptance_gate.py reads the gated CSV and prints the "
        "utterance identity, source, time range, transcript, Phase 1 prediction, basic-emotion "
        "candidate, courtroom-affect candidate, scope, corroboration fields, final fields, and "
        "the acceptance-gate reason. It supports three modes."
    )
    doc.add_heading("Commands", level=2)
    doc.add_paragraph(
        "Print all unresolved rows, expected to be 75 in the 200-row pilot:"
    )
    doc.add_paragraph(
        "./.venv/bin/python phase2/annotation/inspect_clancy_acceptance_gate.py \\\n+  --input-csv data/processed/phase2/clancy/courtroom_affect_candidates_200_v4_gated.csv \\\n+  --mode unresolved --print-rows 75",
        style="Intense Quote",
    )
    doc.add_paragraph(
        "Print all critical conflicts, expected to be 7 in the 200-row pilot:"
    )
    doc.add_paragraph(
        "./.venv/bin/python phase2/annotation/inspect_clancy_acceptance_gate.py \\\n+  --input-csv data/processed/phase2/clancy/courtroom_affect_candidates_200_v4_gated.csv \\\n+  --mode critical --print-rows 7",
        style="Intense Quote",
    )
    doc.add_paragraph("Print the union of unresolved and critical rows:")
    doc.add_paragraph(
        "./.venv/bin/python phase2/annotation/inspect_clancy_acceptance_gate.py \\\n+  --input-csv data/processed/phase2/clancy/courtroom_affect_candidates_200_v4_gated.csv \\\n+  --mode both --print-rows 200",
        style="Intense Quote",
    )
    doc.add_heading("Save a review CSV", level=2)
    doc.add_paragraph(
        "For spreadsheet or manual annotation, save the unresolved queue without changing the "
        "source gated file:"
    )
    doc.add_paragraph(
        "./.venv/bin/python phase2/annotation/inspect_clancy_acceptance_gate.py \\\n+  --input-csv data/processed/phase2/clancy/courtroom_affect_candidates_200_v4_gated.csv \\\n+  --mode unresolved --print-rows 75 \\\n+  --output-csv data/processed/phase2/clancy/unresolved_75_for_review.csv",
        style="Intense Quote",
    )
    doc.add_paragraph(
        "The expected pilot counts are 75 UNRESOLVED rows and 7 rows with critical_conflict=YES. "
        "A critical row may also be unresolved, so do not add the two counts mechanically. Review "
        "the acceptance_gate_reason and the transcript/audio/video evidence before entering human "
        "labels. SILVER means machine-assisted acceptance for controlled experiments, not gold "
        "annotation. No deception, credibility, truthfulness, or reliability label is inferred."
    )
    doc.save(SOP)
    print("Updated", SOP)


if __name__ == "__main__":
    main()
