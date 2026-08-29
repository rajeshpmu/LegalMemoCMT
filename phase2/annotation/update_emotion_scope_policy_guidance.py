"""Document the selectable legacy and conservative emotion-scope policies."""
from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
MARKER = "EMOTION_SCOPE_POLICY_V1"
TARGETS = [
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Witness_Speaking_Pipeline_Student_Speaking_Guide.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Phase2_Corpus_Build_Code_Level_Student_Guide.docx",
]


def main() -> None:
    changed = []
    for path in TARGETS:
        doc = Document(path)
        if any(MARKER in p.text for p in doc.paragraphs):
            continue
        backup = path.with_name(path.name + ".before_emotion_scope_policy.docx")
        if not backup.exists(): shutil.copy2(path, backup)
        doc.add_heading("Policy Selection: Legacy Review versus Conservative Review", level=1)
        p = doc.add_paragraph(); p.add_run(MARKER).font.size = Pt(8)
        doc.add_paragraph(
            "The emotion-scope script supports two explicit policies. The default legacy "
            "policy is retained for the current active-review experiment because it proposes "
            "neutral when SpeechBrain says neutral and Odyssey arousal is low, even when the "
            "scope is unclear. This is a review candidate only and must not be treated as an "
            "automatic correction or gold label."
        )
        doc.add_paragraph(
            "The conservative policy is stricter: it proposes neutral only when the transcript "
            "contains emotion language about another person, and it writes UNRESOLVED for an "
            "ordinary Phase 1/audio conflict. Use --policy conservative when reducing false "
            "neutralization is more important than producing a broad review queue."
        )
        doc.add_paragraph(
            "Legacy pilot command: add --policy legacy. Conservative pilot command: add "
            "--policy conservative. The chosen policy is recorded in the JSON summary. The "
            "legacy 200-row rerun produced neutral=180, while original Phase 1 outputs remain "
            "unchanged. Human multimodal review is still required before training."
        )
        doc.save(path); changed.append(str(path))
    print({"updated": changed, "marker": MARKER})


if __name__ == "__main__": main()
