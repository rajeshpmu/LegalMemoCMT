"""Append the standardized DeBERTa scope-inference procedure to guidance DOCX files."""
from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
MARKER = "DEBERTA_SCOPE_INFERENCE_V1"
DOCS = [
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Witness_Speaking_Pipeline_Student_Speaking_Guide.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Phase2_Corpus_Build_Code_Level_Student_Guide.docx",
]

def update(path: Path) -> bool:
    if not path.exists(): return False
    doc = Document(path)
    if any(MARKER in p.text for p in doc.paragraphs): return False
    backup = path.with_name(path.name + ".before_deberta_scope_v1.docx")
    if not backup.exists(): shutil.copy2(path, backup)
    doc.add_heading("Standardized DeBERTa Scope Inference", level=1)
    p = doc.add_paragraph(); p.add_run(MARKER).font.size = Pt(8)
    doc.add_paragraph(
        "LegalMemoCMT standardizes on MoritzLaurer/deberta-v3-large-zeroshot-v2.0 for optional "
        "transcript-based scope suggestions. It is used as a zero-shot NLI classifier: the input "
        "statement is compared with natural-language hypotheses. It is not part of the original "
        "Phase 1 checkpoint and its outputs are not gold labels."
    )
    doc.add_paragraph(
        "The new implementation classifies sequentially. First it predicts emotion_target_scope "
        "using hypotheses such as 'The witness is describing their own emotional state' and "
        "'The witness is reporting or quoting another person's emotional statement.' Then it "
        "conditions temporal-scope inference on the selected target interpretation and predicts "
        "CURRENT, PAST_SELF, PAST_OTHER, HYPOTHETICAL, or UNCLEAR. Previous question context is "
        "included when available so pronouns and courtroom interaction are not treated as isolated sentences."
    )
    doc.add_heading("Reproducible pilot command", level=2)
    doc.add_paragraph(
        "./.venv/bin/python phase2/annotation/run_deberta_scope_inference.py \\\n+  --input-csv data/processed/phase2/clancy/emotion_scope_review_200_scope_aware.csv \\\n+  --output-csv data/processed/phase2/clancy/emotion_scope_review_200_deberta.csv \\\n+  --summary-json reports/phase2/clancy_emotion_scope_review_200_deberta.json \\\n+  --max-rows 200 --device -1"
    )
    doc.add_paragraph(
        "The output writes deberta_target_scope, deberta_target_scope_confidence, score maps, "
        "deberta_temporal_scope, temporal confidence, context, model name, and "
        "deberta_scope_annotation_status=AUTO_SUGGESTED. Existing emotion_target_scope and other "
        "canonical fields are preserved for comparison. Review disagreements before updating a human label."
    )
    doc.add_paragraph(
        "For example, 'I had just found out what had happened and I was pretty much in shock' "
        "should be examined as SELF_EXPRESSED followed by PAST_SELF. This means the witness is "
        "describing an earlier personal state; it does not automatically describe the witness's "
        "current courtroom emotion. No credibility, truthfulness, deception, or reliability label is inferred."
    )
    doc.save(path); return True

def main():
    print({"updated_documents": [str(p) for p in DOCS if update(p)], "marker": MARKER})

if __name__ == "__main__": main()
