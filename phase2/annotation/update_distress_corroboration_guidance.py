"""Append the negative-activation and distress-corroboration correction to guidance DOCX files."""
from pathlib import Path
import shutil

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
MARKER = "DISTRESS_CORROBORATION_RULE_V3"
DOCS = [
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Witness_Speaking_Pipeline_Student_Speaking_Guide.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Phase2_Corpus_Build_Code_Level_Student_Guide.docx",
]


def update(path: Path) -> bool:
    if not path.exists():
        return False
    doc = Document(path)
    if any(MARKER in paragraph.text for paragraph in doc.paragraphs):
        return False
    backup = path.with_name(path.name + ".before_distress_corroboration_v3.docx")
    if not backup.exists():
        shutil.copy2(path, backup)
    doc.add_heading("Heuristic Correction: Negative Activation Is Not Distress", level=1)
    marker = doc.add_paragraph()
    marker.add_run(MARKER).font.size = Pt(8)
    doc.add_paragraph(
        "The affect generator now separates low-level acoustic evidence from a courtroom-affect "
        "class. audio_valence <= 0.30 together with audio_arousal >= 0.45 is recorded as "
        "negative_activation_candidate=YES. It is not sufficient to label DISTRESSED because "
        "the same acoustic pattern can occur in tense, serious, controlled, or explanatory speech."
    )
    doc.add_heading("Distress promotion rule", level=2)
    doc.add_paragraph(
        "DISTRESSED is exposed as a machine candidate only when negative valence and elevated "
        "arousal are supported by at least one corroborating cue: a negative categorical audio "
        "class, distress-specific language, or a separately recorded visual-distress indicator. "
        "If no corroboration exists, distress_corroboration_present=NO and the distress score is "
        "suppressed. The row remains available for human review instead of being falsely promoted."
    )
    doc.add_heading("Worked example: DCBWoWhsTpA_turn06801", level=2)
    doc.add_paragraph(
        "This 10.123-second witness clip has Phase 1 sadness at 0.544674, neutral SpeechBrain "
        "evidence, arousal 0.452778, and valence 0.270715. Its transcript reports a statement "
        "made by another person, so the scope is QUOTED_SPEECH. The revised output records "
        "negative_activation_candidate=YES but distress_corroboration_present=NO. It therefore "
        "does not automatically call the witness distressed. The machine candidate is "
        "NEUTRAL_CALM at 0.60 from neutral audio and verified visual matching, and the separate "
        "basic-emotion review candidate is neutral at 0.75. A human reviewer may record "
        "CALM_COMPOSED and affect_intensity=1 after inspecting the clip, but those are not "
        "created automatically by this heuristic."
    )
    doc.add_heading("Why this improves Phase 2", level=2)
    for text in [
        "Phase 1 emotion, audio-SER output, V/A/D evidence, machine affect candidate, and human annotation remain separate fields.",
        "Negative subject matter and reported distress are not treated as the witness's own emotion without scope and delivery evidence.",
        "Candidate confidence describes rule support, not the probability that the label is correct and not affect intensity.",
        "The absence of distress corroboration is recorded as a missing-evidence condition, not silently converted into a neutral gold label.",
    ]:
        doc.add_paragraph(text, style="List Bullet")
    doc.save(path)
    return True


def main() -> None:
    print({"updated_documents": [str(path) for path in DOCS if update(path)], "marker": MARKER})


if __name__ == "__main__":
    main()
