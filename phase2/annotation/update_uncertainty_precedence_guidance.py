"""Append the uncertainty-over-tension correction to Phase 2 guidance documents."""
from pathlib import Path
import shutil

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
MARKER = "UNCERTAINTY_PRECEDENCE_AFFECT_V2"
DOCS = [
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Witness_Speaking_Pipeline_Student_Speaking_Guide.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Phase2_Corpus_Build_Code_Level_Student_Guide.docx",
]


def update(path: Path) -> bool:
    if not path.exists():
        return False
    doc = Document(path)
    if any(MARKER in paragraph.text for paragraph in doc.paragraphs):
        return False
    backup = path.with_name(path.name + ".before_uncertainty_precedence_v2.docx")
    if not backup.exists():
        shutil.copy2(path, backup)

    doc.add_heading("Heuristic Correction: Uncertainty Before Generic Tension", level=1)
    marker = doc.add_paragraph()
    marker.add_run(MARKER).font.size = Pt(8)
    doc.add_paragraph(
        "The courtroom-affect proposal logic was revised after reviewing "
        "DCBWoWhsTpA_turn06570. The old pattern arousal >= 0.45 plus valence <= 0.45 "
        "plus a hesitation marker could label a short qualified answer as TENSE. "
        "That is too deterministic: hesitation and epistemic qualification can indicate "
        "uncertainty without visible or vocal tension."
    )
    doc.add_heading("Revised rule order", level=2)
    for text in [
        "Detect hesitation markers such as Um, uh, er, and well, and epistemic markers such as I do not believe, I am not sure, I do not know, I think, maybe, or perhaps.",
        "When hesitation and epistemic uncertainty occur together, give HESITANT_UNCERTAIN precedence over the generic TENSE score.",
        "Use arousal, valence, dominance, and visual evidence as supporting evidence. Do not treat one arousal threshold as a ground-truth emotion decision.",
        "Require additional tension evidence before a future rule promotes TENSE: vocal strain, strong prosodic activation, visible bodily tension, repeated interruption, or similar evidence.",
        "Treat bracketed markers such as [snorts], [sighs], [laughs], and [inaudible] as context only. They must not be directly mapped to disgust or another emotion.",
        "Keep Phase 1 predictions unchanged. The new fields are machine review candidates, not human labels; affect intensity remains UNKNOWN until human review.",
    ]:
        doc.add_paragraph(text, style="List Bullet")
    doc.add_heading("Student-level worked example", level=2)
    doc.add_paragraph(
        "For DCBWoWhsTpA_turn06570, the transcript is 'Um, [snorts] I don't believe "
        "she shared that specific with me.' The speaker is a human-verified witness. "
        "The audio candidate is neu, arousal is 0.494551, valence is 0.373194, and "
        "the Phase 1 prediction is disgust at 0.472638. The word 'Um' supplies hesitation "
        "and 'I don't believe' supplies epistemic uncertainty. The rule therefore proposes "
        "HESITANT_UNCERTAIN at 0.75, with response_stance_candidate=UNCERTAIN_RESPONSE. "
        "It proposes neutral as the basic-emotion review candidate at 0.75 because the "
        "non-neutral Phase 1 prediction is low confidence and conflicts with neutral audio "
        "and qualified wording. This is not an automatic human label. A reviewer may record "
        "human_basic_emotion=neutral, courtroom_affect=HESITANT_UNCERTAIN, and affect_intensity=1 "
        "only after listening to and viewing the clip."
    )
    doc.add_heading("Reproducible command", level=2)
    doc.add_paragraph(
        "Run the revised 200-row pilot with:"
    )
    command = (
        "./.venv/bin/python phase2/annotation/propose_clancy_courtroom_affect.py \\\n+"
        "  --input-csv data/processed/phase2/clancy/emotion_scope_review_200_scope_aware.csv \\\n+"
        "  --output-csv data/processed/phase2/clancy/courtroom_affect_candidates_200_v2.csv \\\n+"
        "  --summary-json reports/phase2/clancy_courtroom_affect_candidates_200_v2.json \\\n+"
        "  --max-rows 200"
    )
    doc.add_paragraph(command, style="Intense Quote")
    doc.add_paragraph(
        "The pilot output recorded 18 HESITANT_UNCERTAIN, 134 NEUTRAL_CALM, 11 ASSERTIVE, "
        "2 GUARDED, 2 DISTRESSED, and 33 UNKNOWN candidates. All 200 remain review-required. "
        "These counts show how the rule changes the review queue; they do not prove that the "
        "candidate labels are correct."
    )
    doc.save(path)
    return True


def main() -> None:
    updated = [str(path) for path in DOCS if update(path)]
    print({"updated_documents": updated, "marker": MARKER})


if __name__ == "__main__":
    main()
