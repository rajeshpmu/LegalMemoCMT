"""Append the clean basic-emotion/courtroom-affect taxonomy update to DOCX guides."""
from pathlib import Path
import shutil

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
MARKER = "CLEAN_AFFECT_TAXONOMY_V3"
DOCS = [
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Witness_Speaking_Pipeline_Student_Speaking_Guide.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Phase2_Corpus_Build_Code_Level_Student_Guide.docx",
]

def update(path: Path) -> bool:
    if not path.exists():
        return False
    doc = Document(path)
    if any(MARKER in p.text for p in doc.paragraphs):
        return False
    backup = path.with_name(path.name + ".before_clean_affect_taxonomy_v3.docx")
    if not backup.exists():
        shutil.copy2(path, backup)
    doc.add_heading("Clean Taxonomy: Basic Emotion Versus Courtroom Affect", level=1)
    p = doc.add_paragraph(); p.add_run(MARKER).font.size = Pt(8)
    doc.add_paragraph(
        "Courtroom affect now uses CALM_COMPOSED rather than NEUTRAL_CALM. This keeps "
        "basic emotion and courtroom behavior as separate dimensions: a witness can be "
        "basic_emotion=neutral and courtroom_affect=CALM_COMPOSED, but can also be "
        "basic_emotion=sadness while presenting in a calm, controlled manner."
    )
    doc.add_heading("Speaker-emotion evidence field", level=2)
    doc.add_paragraph(
        "speaker_emotion_evidence_present records whether the available transcript scope "
        "supports an emotion expressed by the current speaker. The machine proposal uses "
        "NO for OTHER_PERSON_DESCRIBED, EVENT_DESCRIBED, and QUOTED_SPEECH; YES for "
        "SELF_EXPRESSED; and UNKNOWN when scope is unresolved. This is evidence about "
        "attribution, not a final emotion label. A quoted report such as 'She said...' "
        "should not automatically transfer the quoted person's emotion to the witness."
    )
    doc.add_heading("Worked example: DCBWoWhsTpA_turn06801", level=2)
    doc.add_paragraph(
        "The revised record retains Phase 1 sadness at 0.544674, proposes neutral as the "
        "basic-emotion review candidate at 0.75, and uses courtroom_affect=CALM_COMPOSED "
        "rather than NEUTRAL_CALM. It records negative_activation_candidate=YES because "
        "valence and arousal are negative/elevated, but distress_corroboration_present=NO. "
        "Because the scope is QUOTED_SPEECH, speaker_emotion_evidence_present=NO. "
        "The output is still a machine candidate and requires human review; affect intensity "
        "is not inferred automatically."
    )
    doc.add_heading("Interpretation guardrail", level=2)
    doc.add_paragraph(
        "negative_activation_candidate=YES and basic_emotion_review_candidate=neutral are "
        "not contradictory. The first describes low-level acoustic evidence; the second is "
        "an integrated categorical review suggestion. Neither field is a gold annotation, "
        "and no deception, credibility, truthfulness, or reliability label is inferred."
    )
    doc.save(path)
    return True

def main() -> None:
    print({"updated_documents": [str(p) for p in DOCS if update(p)], "marker": MARKER})

if __name__ == "__main__":
    main()
