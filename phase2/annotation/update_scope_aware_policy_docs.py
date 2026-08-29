"""Append the scope-aware policy and pilot result to Phase 2 DOCX guidance."""
from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
MARKER = "SCOPE_AWARE_POLICY_V1"
TARGETS = [
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Clancy_Witness_Speaking_Pipeline_Student_Speaking_Guide.docx",
    ROOT / "implementation_docments/LegalMemoCMT_Phase2_Corpus_Build_Code_Level_Student_Guide.docx",
]


def main() -> None:
    changed = []
    for path in TARGETS:
        doc = Document(path)
        if any(MARKER in p.text for p in doc.paragraphs):
            continue
        backup = path.with_name(path.name + ".before_scope_aware_policy.docx")
        if not backup.exists(): shutil.copy2(path, backup)
        doc.add_heading("Scope-Aware Neutral Candidate Policy", level=1)
        p = doc.add_paragraph(); p.add_run(MARKER).font.size = Pt(8)
        doc.add_paragraph(
            "The emotion-scope review script supports --policy scope_aware in addition to "
            "the retained legacy and conservative policies. This policy treats CALM_COMPOSED "
            "as behavioral presentation, not as proof of neutral emotion. It requires neutral "
            "SpeechBrain evidence, arousal below 0.45, no explicit contrary visual/audio "
            "failure, and then evaluates the transcript target scope."
        )
        doc.add_paragraph(
            "For OTHER_PERSON_DESCRIBED, EVENT_DESCRIBED, or QUOTED_SPEECH, it creates a "
            "NEUTRAL_CANDIDATE at 0.72 confidence for human review. For SELF_EXPRESSED it "
            "keeps the outcome UNRESOLVED. For UNCLEAR it creates a lower-confidence neutral "
            "candidate at 0.58 with review required. Low arousal alone and SpeechBrain neutral "
            "alone never create a final label."
        )
        doc.add_paragraph(
            "Run with: ./.venv/bin/python phase2/annotation/"
            "build_clancy_emotion_scope_review.py --phase1-csv "
            "data/processed/phase2/clancy/phase1_trimodal_pseudo_labels_200.csv "
            "--audio-ser-csv data/processed/phase2/clancy/audio_ser_evidence_200.csv "
            "--output-csv data/processed/phase2/clancy/emotion_scope_review_200_scope_aware.csv "
            "--summary-json reports/phase2/clancy_emotion_scope_review_200_scope_aware.json "
            "--max-rows 200 --policy scope_aware"
        )
        doc.add_paragraph(
            "The pilot result for DCBWoWhsTpA_turn06575 is Phase 1 fear at 0.704363, "
            "SpeechBrain neu, Odyssey arousal 0.378170, OTHER_PERSON_DESCRIBED, and the "
            "scope-aware recommendation neutral/CALM_COMPOSED at 0.72. This remains a "
            "machine proposal requiring human multimodal review."
        )
        doc.save(path); changed.append(str(path))
    print({"updated": changed, "marker": MARKER})


if __name__ == "__main__": main()
