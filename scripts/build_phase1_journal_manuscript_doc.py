#!/usr/bin/env python3
from __future__ import annotations

import json
import subprocess
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "implementation_docments" / "LegalMemoCMT_Phase1_IEEE_Style_Project_Report.docx"
FIG_DIR = ROOT / "implementation_docments" / "figures" / "phase1_journal"
FIG_DIR.mkdir(parents=True, exist_ok=True)

PIPELINE_SVG = FIG_DIR / "phase1_replication_pipeline.svg"
PIPELINE_PNG = FIG_DIR / "phase1_replication_pipeline.png"
VIT_SVG = FIG_DIR / "vit_support_pipeline.svg"
VIT_PNG = FIG_DIR / "vit_support_pipeline.png"
METRIC_PNG = FIG_DIR / "phase1_metric_comparison.png"
FOLD2_CONF_PNG = ROOT / "results" / "paper_aligned_meld_cv" / "cmt_min" / "fold_2" / "analysis_test" / "confusion_matrix.png"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11.5)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_num(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], font_size: float = 10.2) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(font_size)
    doc.add_paragraph()


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.italic = True


def add_picture(doc: Document, path: Path, width: float = 6.4, caption: str | None = None) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        if caption:
            add_caption(doc, caption)


def render_mermaid(code: str, svg_path: Path, png_path: Path) -> None:
    mmd_path = png_path.with_suffix(".mmd")
    mmd_path.write_text(code, encoding="utf-8")
    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(mmd_path), "-o", str(svg_path), "-b", "white"],
        check=True,
        cwd=str(ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(mmd_path), "-o", str(png_path), "-b", "white"],
        check=True,
        cwd=str(ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )


def metric_chart() -> None:
    labels = ["Paper-aligned\nMELD mean", "Face-crop\ngated Fold 2", "Face-crop\ngated Fold 4", "Face-crop+\naux Fold 2", "Face-crop+\naux Fold 4"]
    accuracy = [0.6247, 0.6222, 0.5973, 0.6054, 0.5992]
    macro = [0.4395, 0.4191, 0.4156, 0.4351, 0.4330]
    weighted = [0.6195, 0.6109, 0.5968, 0.6022, 0.6056]

    x = range(len(labels))
    fig, ax = plt.subplots(figsize=(11, 4.7))
    ax.bar([i - 0.23 for i in x], accuracy, width=0.22, label="Accuracy", color="#2F5D8A")
    ax.bar(list(x), weighted, width=0.22, label="Weighted F1", color="#4E8D5D")
    ax.bar([i + 0.23 for i in x], macro, width=0.22, label="Macro F1", color="#D28E39")
    ax.set_ylim(0, 0.75)
    ax.set_ylabel("Score")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels)
    ax.set_title("Phase 1 paper-aligned baseline and ViT-support comparison")
    ax.legend(frameon=False, ncol=3, loc="upper center", bbox_to_anchor=(0.5, 1.12))
    ax.grid(axis="y", linestyle="--", alpha=0.25)
    fig.tight_layout()
    fig.savefig(METRIC_PNG, dpi=220, bbox_inches="tight")
    plt.close(fig)


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def build_assets() -> None:
    render_mermaid(
        """flowchart LR
  A["MELD raw clips and annotations"] --> B["Build raw and fold manifests"]
  B --> C["Train paper-aligned CMT + MIN"]
  C --> D["Evaluate held-out MELD test folds"]
  D --> E["Export metrics, predictions, confusion matrix"]
  E --> F["Interpret imbalance and class confusion"]
""",
        PIPELINE_SVG,
        PIPELINE_PNG,
    )
    render_mermaid(
        """flowchart LR
  A["Raw .mp4 utterance"] --> B["Sample RGB frames"]
  B --> C["Face-crop or full-frame preprocessing"]
  C --> D["Pretrained ViT"]
  D --> E["Cached .npy visual embeddings"]
  E --> F["Warm-start multimodal Phase 1 model"]
  F --> G["Prediction, confidence, top-3 probabilities"]
""",
        VIT_SVG,
        VIT_PNG,
    )
    metric_chart()


def build_doc() -> Document:
    build_assets()
    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LegalMemoCMT Phase 1 IEEE-Style Project Report")
    r.font.name = "Times New Roman"
    r.font.size = Pt(22)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Paper replication with supported ViT facial-cue experiments")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12.5)
    r.italic = True

    add_para(
        doc,
        "This Phase 1 project report documents the LegalMemoCMT work as a paper-aligned replication and implementation study. The main claim is the reproduction of a strong conversational multimodal baseline on MELD using pretrained text and audio encoders with cross-modal transformer fusion and MIN pooling. The ViT facial-cue branch is reported as a supporting extension that helps analyze raw-video behavior, but it is not the central benchmark claim.",
    )

    doc.add_heading("Abstract", level=1)
    add_para(
        doc,
        "Emotion recognition in conversations requires modeling context, modality interaction, and class imbalance. This Phase 1 project reproduces the core MemoCMT-style conversational benchmark path on MELD and evaluates a supporting facial-cue extension based on Vision Transformer (ViT) features extracted from raw video. The paper-aligned MELD workflow uses BERT for text, HuBERT for audio, a cross-modal transformer fusion block, and MIN pooling. Five-fold MELD cross-validation yields a mean accuracy of 0.6247, mean weighted F1 of 0.6195, and mean macro F1 of 0.4395. The ViT facial-cue experiments provide additional evidence that raw video can influence predictions, especially for boundary cases, but they do not replace the paper-aligned conversational baseline as the primary result. The report also presents fold-level results, confusion patterns, and limitations that arise from class imbalance and emotionally close labels.",
    )
    add_para(doc, "Keywords: emotion recognition in conversation; MELD; MemoCMT; cross-modal transformer; HuBERT; BERT; Vision Transformer; facial cues; cross-validation; imbalance-aware learning.")

    doc.add_heading("1. Introduction", level=1)
    add_para(
        doc,
        "Emotion recognition in conversation is difficult because each utterance is shaped by dialogue context, speaker identity, and modality-specific noise. In MELD, the labels are conversational and imbalanced, so a model can score well on accuracy while still failing on the minority emotions. The goal of Phase 1 was therefore not only to obtain a working model, but to replicate the structure of the base paper closely enough that the results can be interpreted in the same research frame.",
    )
    add_para(
        doc,
        "The project first establishes a paper-aligned conversational baseline on MELD using pretrained BERT and HuBERT encoders, cross-modal transformer fusion, and MIN pooling. After the baseline is in place, a ViT-based facial-cue path is added as a supporting extension so the system can also process raw video and surface visual evidence during analysis and demo runs. The ViT branch is useful for understanding what the model sees in a clip, but it is not treated as a substitute for the main conversational result.",
    )
    add_para(
        doc,
        "The document is written as an implementation-oriented IEEE-style project report: it explains the dataset, pipeline, model architecture, training and evaluation procedure, the observed results, and the remaining limitations. The emphasis is on reproducibility and honest interpretation rather than on claiming a new state-of-the-art result.",
    )

    doc.add_heading("2. Related Work", level=1)
    add_bullets(
        doc,
        [
            "MELD established a widely used multimodal conversational benchmark with text, audio, visual information, and both emotion and sentiment labels.",
            "DialogueRNN and DialogueGCN showed that conversational context and speaker structure matter for emotion recognition in dialogue.",
            "BERT and HuBERT provide pretrained language and speech representations that transfer well to downstream classification tasks.",
            "Vision Transformer provides a strong visual backbone for extracting frame-level representations from face crops or full frames.",
            "Focal loss and weighted cross-entropy are standard imbalance-aware objectives that help when some emotion classes dominate the dataset.",
            "Recent surveys on multimodal emotion recognition in conversations emphasize the same core challenge observed in this project: complementary information is useful, but class imbalance and ambiguous emotional boundaries remain difficult.",
        ],
    )
    add_para(
        doc,
        "The related work is relevant here because the project follows the same design logic: use pretrained encoders, model conversational context, and inspect whether adding another modality really changes the error profile rather than simply increasing model complexity.",
    )

    doc.add_heading("3. Materials and Methods", level=1)
    doc.add_heading("3.1 Dataset and Manifests", level=2)
    add_para(
        doc,
        "The core benchmark is MELD, a multi-party conversational dataset from Friends with seven emotion classes. The repository organizes MELD through manifests so that training and evaluation are reproducible. The MELD raw manifest and fold manifests allow the model to be trained on dialogue-grouped splits rather than on random utterance splits, which reduces leakage across train and validation partitions.",
    )
    add_bullets(
        doc,
        [
            "The paper-aligned MELD run uses five folds built from the MELD train/dev dialogue pool.",
            "Held-out test evaluation is carried out per fold, producing predictions and metrics for each fold.",
            "The ViT facial-cue paths use raw MELD mp4 clips and cache extracted facial embeddings as .npy files.",
            "The same manifest-driven organization is reused so that the sample ID, label, transcript, and media paths remain traceable.",
        ],
    )
    add_para(
        doc,
        "A second emotion corpus, CREMA-D, is also supported in the project to validate the audio-emotion pipeline and the cross-validation machinery. In this report, however, MELD is the primary dataset because it is the closest fit to the paper-aligned conversational benchmark story.",
    )

    doc.add_heading("3.2 Baseline Model", level=2)
    add_para(
        doc,
        "The paper-aligned baseline follows the MemoCMT-style design: BERT encodes text, HuBERT encodes audio, the cross-modal transformer mixes modality features, and MIN pooling aggregates the aligned representations before classification. The model is trained with the project’s existing training loop and evaluated using accuracy, weighted F1, macro F1, and related summary outputs.",
    )
    add_para(
        doc,
        "The implementation goal is not to replace the paper architecture with a new invention. It is to keep the baseline close enough to the MemoCMT design that the observed behavior is meaningful, then test whether the additional visual support actually helps on the kinds of clips that matter in MELD.",
    )

    add_picture(
        doc,
        PIPELINE_PNG,
        width=6.3,
        caption="Figure 1. Paper-aligned Phase 1 replication workflow: manifests -> CMT + MIN training -> test evaluation -> predictions and confusion analysis.",
    )

    doc.add_heading("3.3 ViT Facial-Cue Support Path", level=2)
    add_para(
        doc,
        "The facial-cue extension processes raw video clips by sampling RGB frames, passing them through a pretrained ViT model, and caching the resulting embeddings as .npy files. The project supports both face-crop and full-frame variants. In the face-crop setting, the video signal is narrowed to the speaker’s face, which is more suitable for emotionally expressive courtroom-style testimony. In the full-frame setting, the model sees broader context but also more background noise.",
    )
    add_para(
        doc,
        "The ViT branch is treated as a support mechanism for the project’s later stages. It provides a concrete way to inspect video behavior and to study whether face-focused visual evidence can improve the model’s handling of neutral-heavy or emotionally close clips. At the same time, the results show that adding visual support alone does not automatically outperform the strong paper-aligned text-plus-audio baseline.",
    )
    add_picture(
        doc,
        VIT_PNG,
        width=6.25,
        caption="Figure 2. Raw-mp4 facial-cue path: video -> frame sampling -> face crop/full frame -> ViT -> cached embeddings -> multimodal inference.",
    )

    doc.add_heading("3.4 Training and Evaluation Protocol", level=2)
    add_num(
        doc,
        [
            "Train the MELD folds with the paper-aligned configuration and save the best validation checkpoint per fold.",
            "Evaluate the resulting checkpoint on the held-out test split.",
            "Export per-sample predictions, metrics.json, confusion matrices, and top-confusion tables.",
            "Compare fold 2 and fold 4 behavior to understand whether the model is stable across different dialogue partitions.",
            "Use the ViT-based runs as supporting experiments, not as the sole basis for the main benchmark claim.",
        ],
    )
    add_para(
        doc,
        "For evaluation, the report uses accuracy, weighted F1, macro F1, and unweighted accuracy. Accuracy and weighted F1 tend to look stronger than macro F1 under class imbalance, so macro F1 is essential for understanding minority-class behavior.",
    )

    doc.add_heading("4. Results", level=1)
    add_para(
        doc,
        "The strongest and most stable result in the repository is the paper-aligned MELD CMT + MIN baseline. Across five folds, it achieves a mean accuracy of 0.6247 with a standard deviation of 0.0072, a mean weighted F1 of 0.6195, and a mean macro F1 of 0.4395. The unweighted accuracy is lower at 0.4417, which confirms that class imbalance remains a meaningful issue.",
    )
    add_table(
        doc,
        ["Run", "Accuracy", "Weighted F1", "Macro F1", "Unweighted Accuracy"],
        [
            ["Paper-aligned MELD 5-fold mean", "0.6247", "0.6195", "0.4395", "0.4417"],
            ["Paper-aligned Fold 2", "0.6375", "0.6254", "0.4430", "0.4369"],
            ["Paper-aligned Fold 4", "0.6199", "0.6194", "0.4606", "0.4672"],
            ["Face-crop gated Fold 2", "0.6222", "0.6109", "0.4191", "0.4173"],
            ["Face-crop gated Fold 4", "0.5973", "0.5968", "0.4156", "0.4254"],
            ["Face-crop gated + aux Fold 2", "0.6054", "0.6022", "0.4351", "0.4492"],
            ["Face-crop gated + aux Fold 4", "0.5992", "0.6056", "0.4330", "0.4638"],
        ],
    )
    add_para(
        doc,
        "The table shows the core conclusion clearly. The paper-aligned conversational baseline is stronger overall than the facial-cue branches that were added later. The ViT support path is useful, but it is not a universal improvement. Its value is mainly in helping with certain boundary cases and in making the visual branch observable for analysis and review.",
    )

    add_picture(
        doc,
        METRIC_PNG,
        width=6.6,
        caption="Figure 3. Phase 1 metric comparison across the main paper-aligned baseline and the ViT-support runs used in the project.",
    )

    doc.add_heading("4.1 Fold 2 Error Analysis", level=2)
    add_para(
        doc,
        "Fold 2 is useful because it shows a balanced mix of strengths and weaknesses. In the paper-aligned baseline, the model reaches 0.6375 accuracy and 0.6254 weighted F1, while macro F1 remains lower at 0.4430. The confusion matrix shows a recurring tendency to confuse neutral with emotionally close classes, which is consistent with the class imbalance in MELD.",
    )
    add_picture(
        doc,
        FOLD2_CONF_PNG,
        width=5.9,
        caption="Figure 4. Paper-aligned Fold 2 confusion matrix, showing the dominant neutral-heavy error structure.",
    )
    add_para(
        doc,
        "The fold 2 pattern matters because it explains why the project does not frame the result as simple accuracy maximization. The model learns useful signal, but it still confuses the classes that are emotionally close or underrepresented. That is exactly why macro F1 is reported alongside accuracy and weighted F1.",
    )

    doc.add_heading("4.2 What the ViT Support Adds", level=2)
    add_para(
        doc,
        "The ViT-based facial-cue runs do not replace the baseline. Instead, they show how raw video can influence the decision path. In several examples, the face-crop branch changes a near-miss neutral example into a correct neutral prediction, while some emotionally ambiguous clips remain confidently wrong. This is still useful because it reveals where the model relies on visual evidence and where it does not.",
    )
    add_bullets(
        doc,
        [
            "Face-crop support helps expose facial evidence when the clip is visually informative.",
            "Neutral-heavy cases can improve when the correct face region is emphasized.",
            "Ambiguous high-arousal clips can still be misclassified with high confidence.",
            "The facial branch should therefore be read as an analytical support path, not a blanket performance upgrade.",
        ],
    )

    doc.add_heading("5. Discussion", level=1)
    add_para(
        doc,
        "The report’s main technical conclusion is that the paper-aligned conversational baseline is already strong enough to serve as the reference model for Phase 1. The ViT support branch adds interpretability and a new visual input path, but it does not consistently outperform the baseline. That is not a failure of the project; it is a realistic result for a class-imbalanced conversational dataset where text and audio already carry most of the predictive signal.",
    )
    add_para(
        doc,
        "The model’s behavior also fits the qualitative error analysis. Neutral examples remain central because they dominate the dataset, and minority emotions such as disgust, fear, and surprise are more sensitive to sampling and context. This is why the confusion matrix and top-confusion tables are more informative than accuracy alone. A higher accuracy without better class separation would not be a genuine improvement for a conversational emotion system.",
    )
    add_para(
        doc,
        "The ViT support path is still worth keeping. It creates a bridge between the conversational benchmark and the future courtroom-testimony work, where facial information may matter more strongly. In that sense, the current visual experiments are a methodological stepping stone: they show how a raw mp4 can be turned into a cached feature representation and then analyzed inside the same multimodal framework.",
    )

    doc.add_heading("6. Limitations", level=1)
    add_bullets(
        doc,
        [
            "The facial-cue branch is supplementary and does not consistently outperform the text-plus-audio baseline.",
            "MELD remains class imbalanced, so macro F1 stays substantially lower than weighted F1.",
            "Some clips are still confidently misclassified, which means the model can be certain and wrong at the same time.",
        "The project uses CREMA-D as a supporting benchmark, but the paper-aligned MELD result is the primary Phase 1 result.",
        ],
    )
    add_para(
        doc,
        "These limitations are important because they keep the interpretation honest. The current work is strong enough to document a reproducible benchmark implementation, but it is not a claim that the model has solved conversational emotion recognition in general.",
    )

    doc.add_heading("7. Conclusion", level=1)
    add_para(
        doc,
        "Phase 1 now contains a reproducible paper-aligned MELD baseline and a supporting ViT facial-cue path. The baseline demonstrates that the core MemoCMT-style design can be implemented successfully and evaluated consistently. The ViT branch demonstrates that raw video can be converted into usable embeddings and inspected in the same workflow, but it should be understood as a supporting extension rather than a replacement for the conversational benchmark. This distinction is the most important technical conclusion from the current implementation.",
    )

    doc.add_heading("References", level=1)
    refs = [
        "1. Khan, M., Tran, P.-N., Pham, N. T., El Saddik, A., & Othmani, A. (2025). MemoCMT: multimodal emotion recognition using cross-modal transformer-based feature fusion. Scientific Reports, 15(1), 5473. https://doi.org/10.1038/s41598-025-89202-x",
        "2. Poria, S., Hazarika, D., Majumder, N., Naik, G., Cambria, E., & Mihalcea, R. (2019). MELD: A Multimodal Multi-Party Dataset for Emotion Recognition in Conversations. ACL 2019. https://aclanthology.org/P19-1050/",
        "3. Cao, H., Cooper, D., Keutmann, M., Gur, R., Nenkova, A., & Verma, R. (2014). CREMA-D: Crowd-sourced Emotional Multimodal Actors Dataset. IEEE Transactions on Affective Computing. https://pmc.ncbi.nlm.nih.gov/articles/PMC4313618/",
        "4. Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. NAACL 2019 / arXiv. https://arxiv.org/abs/1810.04805",
        "5. Hsu, W.-N., Bolte, B., Tsai, Y.-H. H., Lakhotia, K., Salakhutdinov, R., & Mohamed, A. (2021). HuBERT: Self-Supervised Speech Representation Learning by Masked Prediction of Hidden Units. arXiv. https://arxiv.org/abs/2106.07447",
        "6. Dosovitskiy, A., Beyer, L., Kolesnikov, A., et al. (2021). An Image is Worth 16x16 Words: Transformers for Image Recognition at Scale. ICLR 2021 / arXiv. https://arxiv.org/abs/2010.11929",
        "7. Schneider, S., Baevski, A., Collobert, R., & Auli, M. (2019). wav2vec: Unsupervised Pre-training for Speech Recognition. arXiv. https://arxiv.org/abs/1904.05862",
        "8. Baevski, A., Zhou, H., Mohamed, A., & Auli, M. (2020). wav2vec 2.0: A Framework for Self-Supervised Learning of Speech Representations. NeurIPS 2020 / arXiv. https://arxiv.org/abs/2006.11477",
        "9. Lin, T.-Y., Goyal, P., Girshick, R., He, K., & Dollár, P. (2018). Focal Loss for Dense Object Detection. ICCV 2017 / arXiv. https://arxiv.org/abs/1708.02002",
        "10. Majumder, N., Poria, S., Hazarika, D., Mihalcea, R., Gelbukh, A., & Cambria, E. (2019). DialogueRNN: An Attentive RNN for Emotion Detection in Conversations. AAAI 2019. https://arxiv.org/abs/1811.00405",
        "11. Ghosal, D., Majumder, N., Poria, S., Chhaya, N., & Gelbukh, A. (2020). DialogueGCN: A Graph Convolutional Neural Network for Emotion Recognition in Conversation. EMNLP-IJCNLP 2019. https://aclanthology.org/D19-1015/",
        "12. Ghosh, S., Ramaneswaran, S., Tyagi, U., Srivastava, H., Lepcha, S., Sakshi, S., & Manocha, D. (2022). M-MELD: A Multilingual Multi-Party Dataset for Emotion Recognition in Conversations. arXiv. https://arxiv.org/abs/2203.16799",
        "13. Wu, C., Cai, Y., Liu, Y., Zhu, P., Xue, Y., Gong, Z., Hirschberg, J., & Ma, B. (2025). Multimodal Emotion Recognition in Conversations: A Survey of Methods, Trends, Challenges and Prospects. arXiv. https://arxiv.org/abs/2505.20511",
        "14. Shou, Y., Meng, T., Ai, W., Yin, N., & Li, K. (2023). A Comprehensive Survey on Multi-modal Conversational Emotion Recognition with Deep Learning. arXiv. https://arxiv.org/abs/2312.05735",
        "15. Zhang, X., et al. (2024). Emotion Recognition in Conversations: A Survey Focusing on Multimodal Approaches. Electronics, 12(22), 4714. https://www.mdpi.com/2079-9292/12/22/4714",
        "16. TelME: Teacher-leading Multimodal Fusion Network for Emotion Recognition in Conversation. NAACL 2024. https://aclanthology.org/2024.naacl-long.5/",
        "17. Multimodal Emotion Recognition Calibration in Conversations. 2024. https://www.atailab.cn/seminar2024Fall/pdf/2024_MM_Multimodal%20Emotion%20Recognition%20Calibration%20in%20Conversations.pdf",
    ]
    for ref in refs:
        add_para(doc, ref)

    add_para(
        doc,
        "Reproducibility note: the main scripts behind this Phase 1 project report are scripts/run_paper_aligned_meld_cv.sh, scripts/analyze_meld_vit_facecue_fold2.sh, scripts/analyze_meld_vit_facecue_fold4.sh, scripts/run_demo_paper_aligned_raw_mp4.sh, and scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build_doc()
