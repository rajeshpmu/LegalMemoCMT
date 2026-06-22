from __future__ import annotations

import csv
import json
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_MELD_CV_Paper_Aligned_Report_From_Submission_First.docx"

SUMMARY_JSON = ROOT / "results" / "paper_aligned_meld_cv" / "cmt_min" / "summary.json"
FOLD_DIR = ROOT / "results" / "paper_aligned_meld_cv" / "cmt_min"
FOLD2_ANALYSIS_DIR = FOLD_DIR / "fold_2" / "analysis_test"
FOLD2_PRED_PATH = FOLD_DIR / "fold_2" / "predictions_test.csv"
FOLD2_METRICS = FOLD_DIR / "fold_2" / "metrics.json"
FOLD2_REPORT = FOLD2_ANALYSIS_DIR / "report.txt"
FIRST20_CSV = FOLD2_ANALYSIS_DIR / "predicted_vs_actual_first20.csv"
TOP_CONFUSIONS_CSV = FOLD2_ANALYSIS_DIR / "top_confusions.csv"
CONFUSION_MATRIX_CSV = FOLD2_ANALYSIS_DIR / "confusion_matrix.csv"
SUBMISSION_FIRST = ROOT / "submission_first"
FIG_DIR = ROOT / "implementation_docments" / "figures"
FOLD2_CONFUSION_PNG = FIG_DIR / "fold2_confusion_matrix.png"
FIRST20_CONFUSION_PNG = FIG_DIR / "fold2_first20_confusion_matrix.png"


def configure_document(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def read_summary() -> dict:
    return json.loads(SUMMARY_JSON.read_text(encoding="utf-8"))


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def metric_value(path: Path, key: str) -> float:
    return float(json.loads(path.read_text(encoding="utf-8"))[key])


def _clean_label(label: str) -> str:
    if ":" in label:
        return label.split(":", 1)[1]
    return label


def build_confusion_image(matrix_csv: Path, output_png: Path, title: str) -> None:
    output_png.parent.mkdir(parents=True, exist_ok=True)
    df = pd.read_csv(matrix_csv, index_col=0)
    df.index = [_clean_label(str(v)) for v in df.index]
    df.columns = [_clean_label(str(v)) for v in df.columns]

    fig, ax = plt.subplots(figsize=(9, 7))
    im = ax.imshow(df.values, cmap="Blues", aspect="auto")
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            ax.text(j, i, f"{int(df.values[i, j])}", ha="center", va="center", fontsize=8)
    ax.set_xticks(range(len(df.columns)))
    ax.set_xticklabels(df.columns, rotation=45, ha="right")
    ax.set_yticks(range(len(df.index)))
    ax.set_yticklabels(df.index)
    ax.set_title(title, fontsize=13)
    ax.set_xlabel("Predicted label")
    ax.set_ylabel("Actual label")
    ax.set_xlim(-0.5, len(df.columns) - 0.5)
    ax.set_ylim(len(df.index) - 0.5, -0.5)
    fig.tight_layout()
    fig.savefig(output_png, dpi=200, bbox_inches="tight")
    plt.close(fig)


def build_first20_confusion_image(rows: list[dict[str, str]], output_png: Path, title: str) -> None:
    output_png.parent.mkdir(parents=True, exist_ok=True)
    labels = ["neutral", "joy", "surprise", "sadness", "anger", "fear", "disgust"]
    label_to_idx = {name: idx for idx, name in enumerate(labels)}
    matrix = [[0 for _ in labels] for _ in labels]
    for row in rows:
        actual = row["actual_name"]
        pred = row["predicted_name"]
        if actual in label_to_idx and pred in label_to_idx:
            matrix[label_to_idx[actual]][label_to_idx[pred]] += 1

    df = pd.DataFrame(matrix, index=labels, columns=labels)
    fig, ax = plt.subplots(figsize=(9, 7))
    im = ax.imshow(df.values, cmap="Oranges", aspect="auto")
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            ax.text(j, i, f"{int(df.values[i, j])}", ha="center", va="center", fontsize=8)
    ax.set_xticks(range(len(df.columns)))
    ax.set_xticklabels(df.columns, rotation=45, ha="right")
    ax.set_yticks(range(len(df.index)))
    ax.set_yticklabels(df.index)
    ax.set_title(title, fontsize=13)
    ax.set_xlabel("Predicted label")
    ax.set_ylabel("Actual label")
    ax.set_xlim(-0.5, len(df.columns) - 0.5)
    ax.set_ylim(len(df.index) - 0.5, -0.5)
    fig.tight_layout()
    fig.savefig(output_png, dpi=200, bbox_inches="tight")
    plt.close(fig)


def build_doc() -> Document:
    doc = Document()
    configure_document(doc)

    summary = read_summary()
    fold_metrics = []
    for fold_idx in range(5):
        metrics_path = FOLD_DIR / f"fold_{fold_idx}" / "metrics.json"
        metrics = json.loads(metrics_path.read_text(encoding="utf-8"))
        fold_metrics.append((fold_idx, metrics))

    first20_rows = read_csv_rows(FIRST20_CSV)
    top_confusions = read_csv_rows(TOP_CONFUSIONS_CSV)
    build_confusion_image(CONFUSION_MATRIX_CSV, FOLD2_CONFUSION_PNG, "Fold 2 Confusion Matrix")
    build_first20_confusion_image(first20_rows, FIRST20_CONFUSION_PNG, "Fold 2 First 20 Predicted vs Actual Matrix")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LegalMemoCMT MELD CV Paper-Aligned Report")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "A student-oriented technical report explaining the paper-aligned MELD 5-fold CV run, the actions performed by scripts/run_paper_aligned_meld_cv.sh, and the Fold 2 error analysis."
    )
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)

    doc.add_paragraph()
    add_para(
        doc,
        "Purpose: this document explains the MELD paper-aligned workflow in detail, from the state of the project after the submission_first snapshot through the implementation of the current bidirectional CMT training path, the 5-fold MELD cross-validation run, the resulting metrics, and the error analysis that followed. It is written as a training document for a student who needs to understand both what was executed and why each step mattered.",
    )
    add_para(
        doc,
        "Scope: the report covers the main bash script scripts/run_paper_aligned_meld_cv.sh, the artefacts it created, the exact outcomes of each step, the reason Fold 2 was chosen for deeper analysis, the predicted-versus-actual inspection, the top confusions, and the relationship of this work to the earlier submission_first snapshot.",
    )
    add_para(
        doc,
        "Caution: this report does not overstate the results. It describes what the script actually produced, what metrics were observed, and where the model still fails. It also distinguishes the completed fold analysis from the separate paper-exact scripts that were prepared but not executed.",
    )

    doc.add_heading("1. Context Since submission_first", level=1)
    add_para(
        doc,
        f"The folder {SUBMISSION_FIRST} is the earlier submission snapshot. Compared with that snapshot, the current repository now includes a more complete paper-aligned MELD workflow, paper-exact script templates, improved model wiring for pretrained encoders, raw-media manifest builders, 5-fold dialogue-based MELD cross-validation, and richer analysis and documentation artifacts.",
    )
    add_para(
        doc,
        "In practical terms, the work since submission_first moved from a generic multimodal scaffold toward a more explicit MemoCMT-style pipeline. The main additions were not just new scripts, but also the supporting data flow needed to make the paper-aligned runs meaningful: raw MELD manifests, fold manifests, pretrained text/audio encoding, and a bidirectional cross-attention fusion block with MIN pooling.",
    )
    add_bullets(
        doc,
        [
            "Raw MELD manifest generation so the paper-aligned path can consume clips directly.",
            "Dialogue-grouped MELD 5-fold cross-validation manifests.",
            "Bidirectional text/audio cross-attention CMT for pretrained mode.",
            "Fine-tuning support for pretrained BERT and HuBERT backbones.",
            "Aggregation and analysis scripts that convert fold predictions into reportable outputs.",
            "Dedicated paper-exact script templates and documentation, kept separate from the currently runnable workflow.",
        ],
    )

    doc.add_heading("2. What scripts/run_paper_aligned_meld_cv.sh Does", level=1)
    add_para(
        doc,
        "The bash script is the paper-aligned MELD main workflow. It builds the raw manifest, creates five MELD folds from the train/dev pool, trains the paper-aligned CMT + MIN model on each fold, evaluates each fold on the held-out MELD test split, and then aggregates the fold metrics. The goal is not to compare every pooling type. The goal is to obtain the main paper-aligned MELD result for CMT + MIN under a cross-validation workflow.",
    )
    add_numbered(
        doc,
        [
            "Set runtime environment variables that reduce tokenizers parallelism and make the CPU-based run less fragile on macOS.",
            "Resolve the Python interpreter from the active environment.",
            "Rebuild the raw MELD manifest with clip paths, transcripts, labels, and split assignments.",
            "Build five train/validation fold CSVs from the MELD train/dev dialogue pool.",
            "For each fold, train a model on the fold-specific training manifest and validate on the fold-specific validation manifest.",
            "After each fold is trained, evaluate the resulting checkpoint on the official MELD test split and save predictions plus metrics.",
            "After all five folds finish, aggregate the fold metrics into a JSON summary and a markdown summary.",
        ],
    )

    doc.add_heading("3. Detailed Outcome of Every Action in the Script", level=1)
    add_para(
        doc,
        "This section maps each major action to its concrete output and the final outcome observed in the workspace. The objective is to make the script understandable as a sequence of state changes, not just as a list of commands.",
    )

    action_rows = [
        [
            "Build raw manifest",
            "scripts/build_meld_raw_manifest.py",
            "Wrote data/manifests/meld_raw.csv with 13,707 rows and skipped 1 unreadable sample (dev dia110_utt7 missing under data/MELD/raw).",
            "The raw MELD manifest became available for all downstream fold building and evaluation steps.",
        ],
        [
            "Build 5 CV folds",
            "scripts/build_meld_cv_folds.py",
            "Created five train/validation manifest pairs under data/manifests/meld_cv and wrote meld_cv_summary.json plus meld_fold_assignments.csv.",
            "Five fold-specific dialogue-grouped splits were ready for training.",
        ],
        [
            "Fold 0 train/val",
            "src.train.train",
            "Completed 5 epochs and saved best_model.pt for the fold-0 run.",
            "The first fold completed training successfully.",
        ],
        [
            "Fold 0 test eval",
            "src.train.evaluate",
            "Wrote fold_0/metrics.json and fold_0/predictions_test.csv.",
            "Fold 0 produced a complete test set result.",
        ],
        [
            "Fold 1 train/val",
            "src.train.train",
            "Completed 5 epochs and saved best_model.pt for the fold-1 run.",
            "The second fold completed training successfully.",
        ],
        [
            "Fold 1 test eval",
            "src.train.evaluate",
            "Wrote fold_1/metrics.json and fold_1/predictions_test.csv.",
            "Fold 1 produced a complete test set result.",
        ],
        [
            "Fold 2 train/val",
            "src.train.train",
            "Completed 5 epochs and saved best_model.pt for the fold-2 run.",
            "Fold 2 completed successfully and became the representative analysis fold.",
        ],
        [
            "Fold 2 test eval",
            "src.train.evaluate",
            "Wrote fold_2/metrics.json and fold_2/predictions_test.csv.",
            "Fold 2 produced the analysis set used in this report.",
        ],
        [
            "Fold 3 train/val",
            "src.train.train",
            "Completed 5 epochs and saved best_model.pt for the fold-3 run.",
            "Fold 3 completed successfully.",
        ],
        [
            "Fold 3 test eval",
            "src.train.evaluate",
            "Wrote fold_3/metrics.json and fold_3/predictions_test.csv.",
            "Fold 3 produced a complete test set result.",
        ],
        [
            "Fold 4 train/val",
            "src.train.train",
            "Completed 5 epochs and saved best_model.pt for the fold-4 run.",
            "Fold 4 completed successfully.",
        ],
        [
            "Fold 4 test eval",
            "src.train.evaluate",
            "Wrote fold_4/metrics.json and fold_4/predictions_test.csv.",
            "Fold 4 produced a complete test set result.",
        ],
        [
            "Aggregate fold metrics",
            "scripts/aggregate_fold_metrics.py",
            "Wrote summary.json and summary.md with mean/std over the five folds.",
            "The final cross-validation summary became available for reporting.",
        ],
    ]
    add_table(
        doc,
        ["Action", "Tool", "Observed final outcome", "Why it matters"],
        action_rows,
    )

    doc.add_heading("4. Final Fold-Level Results", level=1)
    add_para(
        doc,
        "The cross-validation summary is the correct way to read the MELD result because each fold uses a different validation dialogue set. The table below lists the actual fold metrics written to the individual metrics.json files and aggregated into summary.json.",
    )

    fold_rows = []
    for fold_idx, metrics in fold_metrics:
        fold_rows.append(
            [
                str(fold_idx),
                f"{metrics['accuracy']:.4f}",
                f"{metrics['weighted_accuracy']:.4f}",
                f"{metrics['unweighted_accuracy']:.4f}",
                f"{metrics['macro_f1']:.4f}",
                f"{metrics['weighted_f1']:.4f}",
            ]
        )
    add_table(
        doc,
        ["Fold", "Accuracy", "Weighted Acc", "Unweighted Acc", "Macro F1", "Weighted F1"],
        fold_rows,
    )

    add_para(
        doc,
        f"Aggregate summary across five folds: mean accuracy {summary['metrics']['accuracy']['mean']:.4f} with std {summary['metrics']['accuracy']['std']:.4f}; mean macro F1 {summary['metrics']['macro_f1']['mean']:.4f} with std {summary['metrics']['macro_f1']['std']:.4f}; mean weighted F1 {summary['metrics']['weighted_f1']['mean']:.4f} with std {summary['metrics']['weighted_f1']['std']:.4f}. The spread is modest, which means the model is reasonably stable across folds, but it is not solving the minority classes cleanly.",
    )

    doc.add_heading("5. Why Fold 2 Was Chosen for Deeper Analysis", level=1)
    add_para(
        doc,
        "Fold 2 was selected for a detailed error analysis because it had the strongest overall held-out accuracy and weighted F1 among the five folds. That makes it a reasonable representative of the model's best overall test behavior. It was not chosen because it had the highest macro F1; in fact, another fold had higher macro F1. The intent was not to cherry-pick the most flattering fold, but to inspect a strong, representative fold that also had a complete predictions file ready for analysis.",
    )
    add_para(
        doc,
        "This choice matters because in an imbalanced dataset like MELD, accuracy and weighted F1 can be high even when minority-class performance remains weak. So the fold-2 selection should be understood as a practical choice for examining the best overall performance profile, not as a claim that fold 2 is the best possible fold on every metric.",
    )
    add_bullets(
        doc,
        [
            "Fold 2 had the highest accuracy among the five folds.",
            "Fold 2 also had the highest weighted F1 among the five folds.",
            "A complete predictions file was available for fold 2, which made analysis straightforward.",
            "The fold still exhibited the same neutral-collapse behavior seen in other folds, so it was representative rather than anomalous.",
        ],
    )

    doc.add_heading("6. Fold 2 Error Analysis", level=1)
    add_para(
        doc,
        f"Fold 2 test metrics were accuracy {metric_value(FOLD2_METRICS, 'accuracy'):.4f}, weighted accuracy {metric_value(FOLD2_METRICS, 'weighted_accuracy'):.4f}, unweighted accuracy {metric_value(FOLD2_METRICS, 'unweighted_accuracy'):.4f}, macro F1 {metric_value(FOLD2_METRICS, 'macro_f1'):.4f}, and weighted F1 {metric_value(FOLD2_METRICS, 'weighted_f1'):.4f}. The separate error-analysis files were created in {FOLD2_ANALYSIS_DIR}.",
    )
    add_para(
        doc,
        "The most important finding from the fold-2 error analysis is that the model still collapses many emotion classes into neutral or into a nearby emotion class. The error pattern is not random. It is structured and consistent with an imbalance-driven bias toward neutral predictions.",
    )

    top_rows = []
    for row in top_confusions[:20]:
        top_rows.append(
            [
                f"{row['actual_name']} ({row['actual_label']})",
                f"{row['predicted_name']} ({row['predicted_label']})",
                row["count"],
            ]
        )
    add_table(
        doc,
        ["Actual class", "Predicted class", "Count"],
        top_rows,
    )

    add_para(
        doc,
        "The top confusion matrix pairs show a clear pattern: neutral is the dominant sink class, and sadness, joy, anger, surprise, fear, and disgust are often pushed toward neutral or neighboring emotions. The most frequent specific confusion in the fold-2 analysis was neutral being predicted as joy, while the most common minority-class failure was sadness being predicted as neutral. Joy and anger are also repeatedly confused with each other and with neutral, which suggests that the model has learned coarse affect but not the fine-grained boundaries among the MELD labels.",
    )
    add_para(
        doc,
        "Three rows are especially useful to read carefully. The first is neutral -> joy with count 97. This means the model often sees neutral utterances as if they were positive or expressive, which is a sign that the neutral class is not cleanly separated from emotionally marked speech. The second is sadness -> neutral with count 86. This is a classic imbalance failure: the model simplifies a low-frequency negative emotion into the dominant baseline class. The third is joy -> neutral with count 83. Even when the true label is a common positive emotion, the model still often falls back to neutral instead of preserving the emotional signal.",
    )
    add_para(
        doc,
        "A second set of rows reinforces the same point. Anger -> neutral appears 71 times, and neutral -> surprise appears 64 times. Together these show that the classifier is not only missing minority emotions, but is also uncertain about the boundary between emotionally loaded and emotionally ordinary utterances. The model does not collapse every emotion into one single class; instead, it alternates between neutral, joy, anger, and surprise, which is consistent with learning rough affective direction but not precise category boundaries.",
    )
    doc.add_picture(str(FOLD2_CONFUSION_PNG), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        doc,
        "Figure: fold-2 confusion matrix heatmap built from the full test predictions. Darker cells indicate more frequent confusions. The dominant band around neutral is visible, as are repeated off-diagonal errors toward joy, anger, and surprise.",
    )

    doc.add_heading("7. Predicted vs Actual: First 20 Examples", level=1)
    add_para(
        doc,
        f"The first 20 predictions file at {FIRST20_CSV} is useful because it shows example-level behavior, not just aggregate metrics. The table below reproduces the first 20 rows from that file.",
    )

    first20_rows_out = []
    for row in first20_rows:
        first20_rows_out.append(
            [
                row["sample_id"],
                row["actual_name"],
                row["predicted_name"],
                f"{float(row['confidence']):.3f}",
                row["correct"],
            ]
        )
    add_table(
        doc,
        ["Sample", "Actual", "Predicted", "Confidence", "Correct"],
        first20_rows_out,
    )

    add_para(
        doc,
        "A few concrete examples make the behavior easy to see. test_dia0_utt0 is surprise but predicted neutral. test_dia0_utt1 is anger but predicted neutral. test_dia1_utt2 is joy but predicted anger. test_dia1_utt5 is joy but predicted neutral. test_dia2_utt2 is sadness but predicted neutral. These examples are not isolated errors; they are consistent with the confusion matrix and with the model's broader tendency to flatten emotion into neutral or nearby affective classes.",
    )
    add_para(
        doc,
        "The first example, test_dia0_utt0, has actual surprise and predicted neutral with confidence 0.434. That is a low-to-moderate confidence mistake and shows that the model has not learned a strong surprise signature in this fold. The second example, test_dia0_utt1, has actual anger but predicted neutral with confidence 0.600. That is a more confident error and shows the model can be wrong while still sounding fairly sure. The third example, test_dia1_utt2, has actual joy but predicted anger with confidence 0.516. This is a useful reminder that some confusions are not only neutral collapses; the model also confuses nearby emotionally active classes with each other.",
    )
    add_para(
        doc,
        "If you look at test_dia1_utt5, the model predicts neutral for a joy example with confidence 0.899. That is one of the clearest signs of the dominant neutral bias: the model becomes highly confident when it falls back to the neutral class. By contrast, correct joy predictions such as test_dia1_utt1 and test_dia1_utt4 show much lower confidence values. The row-by-row table therefore shows that the model is not just making random errors; it is preferring neutral very strongly and only occasionally preserving the intended emotion.",
    )
    doc.add_picture(str(FIRST20_CONFUSION_PNG), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        doc,
        "Figure: confusion heatmap built from the first 20 predicted-vs-actual rows. This image is not the full test confusion matrix; it is a compact visual of the first 20 examples so that the table entries can be read as a small example-level error pattern rather than as aggregate counts.",
    )

    doc.add_heading("8. What the Fold 2 Results Mean Technically", level=1)
    add_para(
        doc,
        "The results show that the model has learned a reasonable amount of conversational emotion structure, but it still exhibits a dominant bias toward neutral. That is visible both in the aggregate confusion matrix and in individual predictions. The weighted metrics are higher than the macro metrics because the model is doing better on frequent classes than on minority classes.",
    )
    add_bullets(
        doc,
        [
            "Neutral is the main sink class for misclassifications.",
            "Sadness and joy are frequently collapsed into neutral.",
            "Anger is frequently confused with neutral, surprise, and joy.",
            "Surprise is unstable and moves across neutral, joy, and anger.",
            "Disgust and fear remain the hardest classes because they are rare and often under-predicted.",
        ],
    )
    add_para(
        doc,
        "This is exactly why macro-F1 matters in addition to accuracy. The model is not failing in a random way; it is failing most strongly on the class balance problem. The fold-2 analysis demonstrates that the chosen architecture is usable and stable, but it does not yet provide a clean separation of all MELD emotions.",
    )

    doc.add_heading("9. Non-Fatal Warnings and Execution Notes", level=1)
    add_para(
        doc,
        "During the run, the audio loader emitted PySoundFile warnings and fell back to audioread for some clips. Those warnings were not fatal. They indicated that the pipeline could still decode the audio, but via a different backend. The run also produced the expected Hugging Face weight-loading messages for BERT and HuBERT, which were normal for the pretrained encoders.",
    )
    add_para(
        doc,
        "An earlier shell invocation produced a transient 'n: command not found' at the very end of the wrapper script. That error did not invalidate the fold outputs already produced. The missing final aggregation was then completed manually with scripts/aggregate_fold_metrics.py, which wrote the summary JSON and markdown files. In other words, the training and per-fold evaluation were finished; only the summary step needed to be rerun.",
    )
    add_bullets(
        doc,
        [
            "Non-fatal audio fallback warnings were observed.",
            "Pretrained weight loading messages were expected.",
            "Per-fold metrics and prediction files were successfully written.",
            "The aggregate summary was later generated manually and completed successfully.",
        ],
    )

    doc.add_heading("10. What Has Been Implemented Since submission_first", level=1)
    add_para(
        doc,
        "The work since submission_first is broader than one run. The project now includes a reproducible paper-aligned MELD path and the supporting documentation needed to explain it. The most important implementation additions are listed below so that a student can see the project's evolution as a sequence of engineering decisions.",
    )
    add_table(
        doc,
        ["Area", "Addition since submission_first", "Purpose"],
        [
            ["Model", "Bidirectional text/audio cross-attention CMT", "Implements the paper's fusion idea more directly."],
            ["Model control", "--fine-tune-backbones", "Allows pretrained BERT and HuBERT to be updated instead of frozen."],
            ["Data", "build_meld_raw_manifest.py", "Creates raw-media MELD manifests for pretrained runs."],
            ["CV", "build_meld_cv_folds.py", "Produces dialogue-grouped 5-fold MELD train/validation splits."],
            ["Training", "run_paper_aligned_meld_cv.sh", "Runs the paper-aligned MELD primary workflow."],
            ["Analysis", "analyze_predictions.py outputs for fold 2", "Produces confusion matrix and example-level error analysis."],
            ["Documentation", "MELD vs CREMA-D deep dive and paper exact README", "Explains benchmark choice and exact protocol separately."],
            ["Documentation", "Student index updates", "Helps the reader find the right document quickly."],
        ],
    )

    doc.add_heading("11. What This Does Not Claim", level=1)
    add_bullets(
        doc,
        [
            "It does not claim perfect reproduction of every paper training detail unless the exact paper schedule is also used.",
            "It does not claim fold 2 is globally the best fold on every metric; it was selected for a strong overall test profile and available analysis outputs.",
            "It does not claim the model has solved minority-class recognition on MELD.",
            "It does not claim CREMA-D should replace MELD; the report uses CREMA-D only as a secondary benchmark.",
            "It does not treat the one-shell 'n: command not found' artifact as a substantive model failure, because the fold outputs were already produced and the summary was later generated successfully.",
        ],
    )

    doc.add_heading("12. Student-Focused Interpretation", level=1)
    add_para(
        doc,
        "If you are reading this as a student, the most important lesson is that the pipeline is now structured enough to separate benchmark design from model design. MELD remains the right primary dataset because it tests conversational multimodal emotion recognition. CREMA-D is useful but secondary. The paper-aligned path should therefore be read as a benchmark comparison story, not as a claim that one dataset should replace the other.",
    )
    add_para(
        doc,
        "The second important lesson is that a cross-validation workflow gives a more reliable estimate than a single split. The fold-level metrics show that the model is fairly stable, but also that its minority-class behavior still needs work. That is a much more useful conclusion than simply quoting the highest accuracy number.",
    )

    doc.add_heading("13. Final Summary", level=1)
    add_para(
        doc,
        "The MELD paper-aligned cross-validation workflow completed successfully. The raw manifest was built, five dialogue-grouped folds were created, each fold was trained and evaluated, and the aggregate summary was produced. Fold 2 was then analyzed in detail because it had the strongest overall test-set accuracy and weighted F1 among the folds, while still showing the same neutral-collapse behavior that limits the model's class-balanced performance.",
    )
    add_para(
        doc,
        "The final technical conclusion is restrained: the current paper-aligned CMT + MIN model is stable and useful, but it still leans heavily toward neutral and does not cleanly separate all MELD emotions. That is why the report emphasizes both the fold summary and the error analysis. The summary shows the model's average performance; the confusion matrix and predicted-vs-actual table show how the model actually behaves.",
    )

    doc.add_heading("Appendix A. Files Produced by the Run", level=1)
    add_bullets(
        doc,
        [
            "results/paper_aligned_meld_cv/cmt_min/summary.json",
            "results/paper_aligned_meld_cv/cmt_min/summary.md",
            "results/paper_aligned_meld_cv/cmt_min/fold_0/metrics.json and predictions_test.csv",
            "results/paper_aligned_meld_cv/cmt_min/fold_1/metrics.json and predictions_test.csv",
            "results/paper_aligned_meld_cv/cmt_min/fold_2/metrics.json, predictions_test.csv, and analysis_test/*",
            "results/paper_aligned_meld_cv/cmt_min/fold_3/metrics.json and predictions_test.csv",
            "results/paper_aligned_meld_cv/cmt_min/fold_4/metrics.json and predictions_test.csv",
        ],
    )

    doc.add_heading("Appendix B. Representative Report Files", level=1)
    add_bullets(
        doc,
        [
            f"Fold 2 report text: {FOLD2_REPORT}",
            f"Fold 2 predictions: {FOLD2_PRED_PATH}",
            f"Fold 2 confusion matrix: {CONFUSION_MATRIX_CSV}",
            f"Fold 2 top confusions: {TOP_CONFUSIONS_CSV}",
            f"Fold 2 first-20 table: {FIRST20_CSV}",
        ],
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
