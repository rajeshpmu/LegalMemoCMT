from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOP = ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx"


def code(document: Document, value: str) -> None:
    p = document.add_paragraph(style="No Spacing")
    r = p.add_run(value)
    r.font.name = "Courier New"
    r.font.size = Pt(8)


document = Document(SOP)
document.add_section(WD_SECTION.NEW_PAGE)
document.add_heading("Runtime Correction After First Pilot Attempt", level=1)
document.add_paragraph(
    "The first attempted 200-row pilot terminated with macOS Segmentation fault: 11 rather than a "
    "normal Python exception. A faulthandler trace showed that the crash occurred while the Hugging Face "
    "stack attempted to import TensorFlow through transformers. TensorFlow is not needed because this "
    "inference path loads a PyTorch checkpoint. The wrapper now exports USE_TF=0 and "
    "TRANSFORMERS_NO_TF=1 before Python starts, preventing that unnecessary TensorFlow import."
)
document.add_paragraph(
    "The wrapper default device is now CPU instead of automatic accelerator selection. This is a "
    "conservative reproducibility choice for the current macOS environment. DEVICE=auto or DEVICE=mps "
    "can be tested later, but CPU is the documented pilot device because it avoids an additional native "
    "accelerator failure mode."
)
document.add_paragraph(
    "The checkpoint also reports model_cfg.use_video=false. Therefore MODALITIES=text,audio,video is a "
    "request to the wrapper, but the effective model input is text,audio. The script records both "
    "requested and effective modalities and emits a warning. The Clancy video paths are preserved for "
    "future video-enabled inference, but this particular Phase 1 checkpoint does not use them. It would "
    "be technically incorrect to describe this run as trimodal model inference."
)
document.add_heading("Corrected Pilot Command", level=2)
document.add_paragraph("Rerun the pilot with the updated wrapper. The explicit DEVICE=cpu is shown for clarity, although CPU is now the default.")
code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python \\
INPUT_CSV=$PWD/data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest.csv \\
OUTPUT_CSV=$PWD/data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_phase1_pseudo_200.csv \\
CHECKPOINT=$PWD/results/paper_aligned_meld_cv/cmt_min/fold_2/best_model.pt \\
SUMMARY_JSON=$PWD/reports/phase2/clancy_duration_0_8_to_20_phase1_pseudo_200.json \\
MAX_ROWS=200 \\
BATCH_SIZE=4 \\
MODALITIES=text,audio,video \\
DEVICE=cpu \\
bash phase2/run_pseudo_label_clancy_with_phase1.sh''')
document.add_heading("What to Verify After the Rerun", level=2)
for item in [
    "The command exits normally and writes the CSV and JSON summary.",
    "The summary reports requested_modalities as audio,text,video and effective_modalities as audio,text.",
    "The summary notes that the checkpoint has use_video=false.",
    "The output contains phase1_basic_emotion, confidence, entropy, probabilities, checkpoint, source, and modality provenance fields.",
    "The original emotion_label fields remain unchanged.",
    "No courtroom_affect, deception, truthfulness, credibility, or reliability labels are created.",
]:
    document.add_paragraph(item, style="List Bullet")
document.add_paragraph(
    "A one-row CPU smoke test completed successfully after this correction. The full 200-row pilot has "
    "not been run automatically; run the corrected command manually, inspect the output, and only then "
    "decide whether to process the full duration pools."
)
document.save(SOP)
print(f"Updated {SOP}")
