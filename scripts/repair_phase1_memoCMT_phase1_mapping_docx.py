from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph


DOCX_PATH = Path("/Users/rajeshpmu/Desktop/LegalMemoCMT/implementation_docments/LegalMemoCMT_Phase1_ESA_Reading_Script.docx")


REMOVE_PREFIXES = [
    "Slide 6: MemoCMT vs Phase 1 Implementation Mapping",
    "Use this slide to explain the exact relationship between the MemoCMT paper and the Phase 1 implementation.",
    "Say that the paper’s CMT is written mathematically with explicit Wq, Wk, and Wv projections, but Phase 1 implements the same idea using PyTorch’s built-in nn.MultiheadAttention inside BidirectionalCrossAttentionCMT.",
    "Then explain the important N/A case. The base MemoCMT paper is text-plus-audio, so video is not part of the original design.",
    "This means the comparison is not paper versus unrelated model.",
    "Code references to mention aloud: src/models/model.py for the paper CMT and video branches, scripts/run_paper_aligned_meld_cv.sh for the paper-aligned text+audio baseline, and scripts/run_phase1_raw_mp4_demo.sh for the raw-mp4 demo route used by the visual experiments.",
]


def style_run(run, *, size=11, bold=False, color="000000"):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold


def remove_paragraph(paragraph):
    el = paragraph._element
    el.getparent().remove(el)


def insert_paragraph_before(anchor_paragraph, text, *, size=11, bold=False):
    new_p = OxmlElement("w:p")
    anchor_paragraph._p.addprevious(new_p)
    paragraph = Paragraph(new_p, anchor_paragraph._parent)
    paragraph.text = text
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.space_before = Pt(0)
    for run in paragraph.runs:
        style_run(run, size=size, bold=bold)
    return paragraph


def insert_paragraph_after(anchor_paragraph, text, *, size=11, bold=False):
    new_p = OxmlElement("w:p")
    anchor_paragraph._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor_paragraph._parent)
    paragraph.text = text
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.space_before = Pt(0)
    for run in paragraph.runs:
        style_run(run, size=size, bold=bold)
    return paragraph


def remove_existing_section(doc: Document):
    # remove the stray inserted section if present
    for para in list(doc.paragraphs):
        txt = para.text.strip()
        if any(txt.startswith(prefix) for prefix in REMOVE_PREFIXES):
            remove_paragraph(para)
    # remove the matching table if present
    for table in list(doc.tables):
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if headers == ["Area", "Paper", "Phase 1", "Exact code refs"]:
            tbl = table._tbl
            tbl.getparent().remove(tbl)


def add_table_after_paragraph(doc: Document, anchor_paragraph, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    tbl = table._tbl
    anchor_paragraph._p.addnext(tbl)
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    return table


def shift_headings(doc: Document):
    for para in doc.paragraphs:
        txt = para.text.strip()
        m = re.match(r"^Slide (\d+): (.+)$", txt)
        if m and int(m.group(1)) >= 6:
            para.text = f"Slide {int(m.group(1)) + 1}: {m.group(2)}"


def insert_new_section(doc: Document):
    anchor = None
    for para in doc.paragraphs:
        if para.text.strip().startswith("Slide 7: Literature Review 2: MELD"):
            anchor = para
            break
    if anchor is None:
        raise RuntimeError("Could not find the shifted MELD heading.")

    heading = insert_paragraph_before(anchor, "Slide 6: MemoCMT vs Phase 1 Implementation Mapping", size=14, bold=True)
    p1 = insert_paragraph_after(
        heading,
        "Use this slide to explain the exact relationship between the MemoCMT paper and the Phase 1 implementation. The safest explanation is that Phase 1 reproduces the paper’s text-plus-audio CMT baseline in paper mode, then extends the system with a video pathway and gated fusion for later experiments.",
        size=11.2,
    )
    p2 = insert_paragraph_after(
        p1,
        "Say that the paper’s CMT is written mathematically with explicit Wq, Wk, and Wv projections, but Phase 1 implements the same idea using PyTorch’s built-in nn.MultiheadAttention inside BidirectionalCrossAttentionCMT. That is a practical implementation choice, not a change in the underlying fusion principle.",
        size=11.2,
    )
    p3 = insert_paragraph_after(
        p2,
        "Then explain the important N/A case. The base MemoCMT paper is text-plus-audio, so video is not part of the original design. In Phase 1, video is added separately through SequenceEncoder and then combined with either legacy_fusion or gated_fusion, with an optional video_aux_classifier to support the auxiliary-loss branch.",
        size=11.2,
    )
    p4 = insert_paragraph_after(
        p3,
        "This means the comparison is not paper versus unrelated model. It is paper baseline versus a faithful Phase 1 reproduction plus a controlled visual extension. That distinction is important because the paper-aligned MELD result should be judged on text and audio, while the video branch should be discussed as a Phase 1 extension.",
        size=11.2,
    )
    p5 = insert_paragraph_after(
        p4,
        "Code references to mention aloud: src/models/model.py for the paper CMT and video branches, scripts/run_paper_aligned_meld_cv.sh for the paper-aligned text+audio baseline, and scripts/run_phase1_raw_mp4_demo.sh for the raw-mp4 demo route used by the visual experiments.",
        size=11.0,
        bold=True,
    )

    headers = ["Area", "Paper", "Phase 1", "Exact code refs"]
    rows = [
        ["CMT fusion", "Explicit Wq/Wk/Wv formulas", "BidirectionalCrossAttentionCMT + MultiheadAttention", "src/models/model.py:160-246"],
        ["Text/audio", "BERT + HuBERT", "PretrainedBackboneEncoder in paper mode", "src/models/model.py:249-316"],
        ["Pooling", "CLS / MEAN / MAX / MIN", "Same options; MELD baseline uses min", "src/models/model.py:190-208; scripts/run_paper_aligned_meld_cv.sh:42-62"],
        ["Video", "N/A", "SequenceEncoder + legacy/gated fusion + aux head", "src/models/model.py:333-445"],
    ]
    table = add_table_after_paragraph(doc, p5, headers, rows)
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                style_run(run, size=9.2, bold=True)


def main():
    doc = Document(str(DOCX_PATH))
    remove_existing_section(doc)
    insert_new_section(doc)
    doc.save(str(DOCX_PATH))
    print(f"Updated {DOCX_PATH}")


if __name__ == "__main__":
    main()
