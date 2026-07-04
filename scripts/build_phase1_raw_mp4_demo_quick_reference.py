#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "Phase1_Raw_MP4_Demo_Quick_Reference.docx"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def add_para(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)


def add_code(doc: Document, text: str) -> None:
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(10.2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Phase 1 Raw MP4 Demo Quick Reference")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Use this page during the demo for exact commands, key files, and fast fallback steps")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    add_para(doc, "This is the short version of the full SOP. It is intended to sit beside you during the review.")

    doc.add_heading("1. Key Files", level=1)
    add_table(
        doc,
        ["What", "Path"],
        [
            ["Manifest", "data/manifests/meld_test.csv"],
            ["Baseline checkpoint", "results/paper_aligned_meld_cv/cmt_min/fold_2/best_model.pt"],
            ["Gated + aux checkpoint", "results/facial_cues/meld_vit_facecrop_gated_video_aux/fold_4/best_model.pt"],
            ["Demo clips", "data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/*.mp4"],
            ["Demo bundle", "results/phase1_review_demo/fold2_baseline/"],
        ],
    )

    doc.add_heading("2. Main Demo Commands", level=1)
    add_para(doc, "Baseline:")
    add_code(
        doc,
        r"""DEVICE=cuda \
bash scripts/run_demo_paper_aligned_raw_mp4.sh \
  test_dia279_utt9 \
  data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia279_utt9.mp4""",
    )
    add_para(doc, "Gated + aux:")
    add_code(
        doc,
        r"""DEVICE=cuda \
CHECKPOINT=results/facial_cues/meld_vit_facecrop_gated_video_aux/fold_4/best_model.pt \
bash scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh \
  test_dia279_utt9 \
  data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia279_utt9.mp4""",
    )
    add_para(doc, "Near-miss clip for comparison:")
    add_code(
        doc,
        r"""DEVICE=cuda \
bash scripts/run_demo_paper_aligned_raw_mp4.sh \
  test_dia278_utt5 \
  data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia278_utt5.mp4""",
    )

    doc.add_heading("3. Bundle / Batch", level=1)
    add_code(
        doc,
        r"""bash scripts/run_phase1_review_demo_bundle.sh
bash scripts/run_phase1_raw_mp4_demo_batch.sh demo_pairs.csv""",
    )

    doc.add_heading("4. Unseen Reviewer Video", level=1)
    add_para(doc, "Use this only if the reviewer brings a clip not already in MELD.")
    add_code(
        doc,
        r"""# Option A: temporary one-row manifest, then reuse the normal demo script
DEVICE=cuda \
CHECKPOINT=results/paper_aligned_meld_cv/cmt_min/fold_2/best_model.pt \
bash scripts/run_phase1_raw_mp4_demo.sh \
  reviewer_video_001 \
  /path/to/reviewer_new_clip.mp4

# Option B: prediction-only use
# Do not claim dataset accuracy unless ground truth is known.""",
    )

    doc.add_heading("5. Speaking Order", level=1)
    add_para(doc, "1. Show the clip.")
    add_para(doc, "2. Run baseline.")
    add_para(doc, "3. Run gated + aux on the same clip.")
    add_para(doc, "4. Show confidence and top-3.")
    add_para(doc, "5. Show metrics and confusion matrix.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
