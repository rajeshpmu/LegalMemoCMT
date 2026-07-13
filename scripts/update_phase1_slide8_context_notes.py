from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "implementation_docments" / "LegalMemoCMT_Phase1_ESA_Prep_Guide.docx"


def insert_paragraph_before(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def remove_section(doc: Document, start_heading: str, next_heading_prefix: str) -> None:
    paras = list(doc.paragraphs)
    start = next((i for i, p in enumerate(paras) if p.text.strip() == start_heading), None)
    if start is None:
        return
    end = next((i for i in range(start + 1, len(paras)) if paras[i].text.strip().startswith(next_heading_prefix)), len(paras))
    for i in range(start, end):
        el = paras[i]._element
        el.getparent().remove(el)


def update_docx() -> None:
    doc = Document(str(DOCX_PATH))
    remove_section(doc, "2.3 Slide 8 MELD Contextual Analysis", "2.4 Literature Review Slides")

    anchor = next((p for p in doc.paragraphs if p.text.strip().startswith("2.4 Literature Review Slides")), None)
    if anchor is None:
        raise RuntimeError("Could not locate the literature review anchor.")

    h = insert_paragraph_before(anchor, style="Heading 2")
    h.add_run("2.3 Slide 8 MELD Contextual Analysis")

    bullets = [
        "Speaker context is handled at the data-split level in scripts/build_meld_cv_folds.py: the script extracts a dialogue key from each MELD sample_id, groups utterances by dialogue, and assigns whole dialogues to train or validation folds. That prevents utterances from the same conversation from being split across folds and leaking context into validation.",
        "The model handles conversational signals only indirectly in src/models/model.py: the text and audio sequences are encoded by pretrained branches, then combined with cross-modal fusion or gated fusion. This means the project can use dialogue context through learned modality interactions, but it does not yet have a separate speaker-memory module like DialogueRNN.",
        "Imbalance is handled in src/train/train.py: compute_class_weights(...) builds class weights from the training subset, and FocalLoss(...) is available when the loss type is focal. This helps because MELD is neutral-heavy and simple accuracy would overstate performance on the majority class.",
        "Error analysis is handled in src/train/evaluate.py and the analysis scripts: the evaluator writes per-sample predictions, metrics, confusion matrices, and top-confusion tables. Those outputs let you inspect whether the model is failing because of context shift, speaker interaction, or long-range conversational dependence.",
    ]
    for t in bullets:
        p = insert_paragraph_before(anchor)
        r = p.add_run(t)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11.3)
        p.alignment = 3

    h2 = insert_paragraph_before(anchor, style="Heading 3")
    h2.add_run("Code paths to mention in the viva")
    code_lines = [
        "scripts/build_meld_cv_folds.py:13-102 - builds dialogue-safe MELD folds using dialogue_key grouping.",
        "src/models/model.py:303-444 - defines the text/audio/video encoders, cross-modal fusion, gated fusion, and the classifier forward pass.",
        "src/train/train.py:215-297 - selects the loss function, computes class weights, and trains / validates the model.",
        "src/train/evaluate.py:17-145 - loads checkpoints and writes predictions, metrics, and confusion outputs for review.",
    ]
    for t in code_lines:
        p = insert_paragraph_before(anchor)
        r = p.add_run(t)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11.0)
        p.alignment = 3

    doc.save(str(DOCX_PATH))


if __name__ == "__main__":
    update_docx()
    print(f"Updated {DOCX_PATH}")
