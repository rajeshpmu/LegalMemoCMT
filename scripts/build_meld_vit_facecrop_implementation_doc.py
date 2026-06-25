from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "implementation_docments" / "LegalMemoCMT_MELD_Face_Crop_ViT_Implementation_Plan.docx"
FIG_DIR = ROOT / "implementation_docments" / "figures" / "meld_vit_facecrop_plan"
FIG_DIR.mkdir(parents=True, exist_ok=True)
PIPELINE_SVG = FIG_DIR / "facecrop_pipeline.svg"
PIPELINE_PNG = FIG_DIR / "facecrop_pipeline.png"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)


def add_para(doc: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.italic = italic
    r.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc: Document, code: str) -> None:
    for line in code.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(9.2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value


def render_mermaid(code: str, svg_path: Path, png_path: Path) -> None:
    mmd_path = svg_path.with_suffix(".mmd")
    mmd_path.write_text(code, encoding="utf-8")
    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(mmd_path), "-o", str(svg_path), "-b", "white"],
        check=True,
        cwd=str(ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(mmd_path), "-o", str(png_path), "-b", "white"],
        check=True,
        cwd=str(ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )


def build_doc() -> Document:
    render_mermaid(
        """flowchart LR
  A[Raw MELD MP4 clip] --> B[Sample 16 full RGB frames]
  B --> C[Detect speaker face in each sampled frame]
  C --> D[Crop face region with padding]
  D --> E[Resize crop to 224x224]
  E --> F[Pretrained ViT-Base]
  F --> G[(.npy face embedding)]
  G --> H[Face-crop manifest row]
  H --> I[MELD fold CSVs]
  I --> J[Warm-start training]
  J --> K[Metrics + confusion matrix]
""",
        PIPELINE_SVG,
        PIPELINE_PNG,
    )

    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    run = title.add_run("LegalMemoCMT MELD Face-Crop ViT Implementation Plan")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(
        "A student-level explanation of why the next facial-cue step should crop the face before ViT, how the new scripts are executed, and what outputs should be expected."
    )
    run.italic = True
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"

    add_para(
        doc,
        "Purpose: this document explains the next facial-cue experiment after the full-frame ViT run. The goal is not to change the model backbone or the MELD fold logic. The only planned change is the visual input path: instead of sending full sampled frames to ViT, the pipeline first crops the face region from each frame and then extracts embeddings from those face crops.",
    )
    add_para(
        doc,
        "Why this matters: the full-frame approach proved that the visual branch works, but it still leans heavily toward neutral predictions. For a courtroom-testimony setting, the face is the more direct source of emotional evidence, so face-crop-before-ViT is the better controlled next ablation.",
    )

    doc.add_heading("1. Why Face Crop Before ViT", level=1)
    add_bullets(
        doc,
        [
            "The face carries the most direct emotion cues: brows, eyes, mouth shape, and tension.",
            "Full frames include background noise, clothing, courtroom objects, and camera framing effects.",
            "Cropping the face makes the ViT input more focused and easier to interpret.",
            "The change is still small enough to be a controlled experiment because the backbone, labels, and folds stay the same.",
        ],
    )
    add_para(
        doc,
        "Student interpretation: the question is not whether video matters. The question is whether ViT should study the whole frame or the speaker face. The face crop path is a better answer for the legal-domain research direction because it emphasizes the human expression instead of the environment around the human.",
    )

    doc.add_heading("2. How This Differs From the Full-Frame Path", level=1)
    add_bullets(
        doc,
        [
            "Full-frame path: sample frames, resize them, send the whole image to ViT.",
            "Face-crop path: sample frames, detect face, crop the face, resize the crop, send only the face region to ViT.",
            "Training path: the same warm-start MELD checkpoint is reused; only the visual features change.",
            "Evaluation path: the same MELD fold splits and the same test analysis scripts are reused.",
        ],
    )
    add_code(
        doc,
        """full frame: clip -> sample frames -> resize -> ViT -> .npy
face crop: clip -> sample frames -> detect face -> crop face -> resize -> ViT -> .npy""",
    )
    add_para(
        doc,
        "The important experimental rule is fairness. If the model gets better or worse after the crop change, that difference should be attributed mainly to the visual input strategy, not to a new loss, a new fold split, or a different checkpoint source.",
    )

    doc.add_heading("3. Script-by-Script Execution Guide", level=1)
    add_table(
        doc,
        ["Script", "What it does", "What it produces"],
        [
            ["scripts/run_meld_vit_facecrop_prepare.sh", "Builds the face-crop manifest and the paper-aligned control MELD CSVs", "data/manifests/meld_vit_facecrop.csv and meld_vit_facecrop_control_cv/"],
            ["scripts/build_meld_vit_facecrop_manifest.py", "Samples frames, crops faces, extracts ViT embeddings, writes .npy files", "data/processed/MELD_VIT_FACECROP/.../*.npy"],
            ["scripts/run_meld_vit_facecrop_video_only_control_fold2.sh", "Runs the dedicated video-only face-crop control for Fold 2", "A video-only control checkpoint and test report for Fold 2"],
            ["scripts/analyze_meld_vit_facecrop_video_only_control_fold2.sh", "Exports video-only control Fold 2 predictions and analysis", "Video-only control confusion matrix and error table"],
            ["scripts/run_meld_vit_facecrop_text_audio_fold2.sh", "Runs the text+audio-only reference on the same face-crop folds", "A non-visual reference checkpoint for Fold 2"],
            ["scripts/analyze_meld_vit_facecrop_text_audio_fold2.sh", "Exports text+audio Fold 2 predictions and analysis", "Text+audio confusion matrix and error table"],
            ["scripts/run_meld_vit_facecrop_fold2.sh", "Runs the full warm-start tri-modal face-crop training for MELD Fold 2", "results/facial_cues/meld_vit_facecrop/fold_2/best_model.pt"],
            ["scripts/analyze_meld_vit_facecrop_fold2.sh", "Exports Fold 2 tri-modal predictions and analysis", "metrics.json, predictions_test.csv, confusion_matrix.csv"],
            ["scripts/run_meld_vit_facecrop_video_only_control_fold4.sh", "Runs the dedicated video-only face-crop control for Fold 4", "A video-only control checkpoint and test report for Fold 4"],
            ["scripts/analyze_meld_vit_facecrop_video_only_control_fold4.sh", "Exports video-only control Fold 4 predictions and analysis", "Video-only control confusion matrix and error table"],
            ["scripts/run_meld_vit_facecrop_text_audio_fold4.sh", "Runs the text+audio-only reference on the same face-crop folds", "A non-visual reference checkpoint for Fold 4"],
            ["scripts/analyze_meld_vit_facecrop_text_audio_fold4.sh", "Exports text+audio Fold 4 predictions and analysis", "Text+audio confusion matrix and error table"],
            ["scripts/run_meld_vit_facecrop_fold4.sh", "Runs the full warm-start tri-modal face-crop training for MELD Fold 4", "results/facial_cues/meld_vit_facecrop/fold_4/best_model.pt"],
            ["scripts/analyze_meld_vit_facecrop_fold4.sh", "Exports Fold 4 tri-modal predictions and analysis", "metrics.json, predictions_test.csv, confusion_matrix.csv"],
            ["scripts/run_meld_vit_facecrop_suite.sh", "Runs the full prepare/train/analyze sequence", "A complete face-crop experiment run"],
            ["scripts/run_meld_vit_facecrop_video_only_control_suite.sh", "Runs the dedicated video-only control sequence", "A standalone video-only control experiment run"],
        ],
    )
    add_para(
        doc,
        "The training engine itself is still the existing warm-start code path. That is deliberate. The model architecture and the paper-aligned checkpoint remain stable so the effect of cropping faces can be isolated cleanly. The fold assignment now uses the paper-aligned MELD split as the reference control, while the visual branch is redirected to the face-crop embeddings.",
    )
    add_para(
        doc,
        "Checkpoint selection is now weighted toward MELD's class imbalance rather than being driven only by accuracy. The new default is weighted F1, with macro F1 reported alongside it so the student can see whether minority emotions really improved.",
    )

    doc.add_heading("4. What the Face-Crop Manifest Builder Does", level=1)
    add_para(
        doc,
        "The manifest builder is the key new script. It starts from the MELD annotation CSV files, finds the raw utterance clip, samples the same 16 frames as the full-frame path, detects the speaker face inside each sampled frame, crops that region, and then sends the face crop to the pretrained ViT model. The ViT output for each frame is saved as a .npy embedding file. The manifest row then points to that file.",
    )
    add_bullets(
        doc,
        [
            "The dataset label column is still the emotion label, not sentiment.",
            "The output feature file is still a .npy array of shape (16, 768).",
            "If no face is detected in a frame, the script falls back to a center crop so the run can continue safely.",
            "The manifest keeps the raw clip path so the source utterance is still traceable.",
        ],
    )
    add_code(
        doc,
        """frames = sample_video_frames(video_path, cfg)
face_frames = [crop_face_frame(frame, cfg, cascade) for frame in frames]
inputs = image_processor(images=list(batch), return_tensors="pt")
outputs = vit_model(**inputs)
hidden = outputs.last_hidden_state[:, 0, :]
np.save(video_feat_path, vit_embeddings)""",
    )
    add_para(
        doc,
        "Student takeaway: the role of face-crop preprocessing is to make the visual signal more focused before ViT sees it. ViT is still the feature extractor. The new script only changes what the feature extractor is looking at.",
    )

    doc.add_heading("5. Why Warm-Start Is Still the Right Training Strategy", level=1)
    add_bullets(
        doc,
        [
            "Start from the best weighted-CE MELD checkpoint so the text-audio backbone is preserved.",
            "Add the face-crop visual branch on top of that baseline.",
            "Keep the same fold splits and pooling choice so the comparison remains fair.",
            "Use the new face-crop embeddings to test whether better visual focus reduces neutral bias.",
        ],
    )
    add_para(
        doc,
        "This is a transfer-learning style experiment. The model is not starting from scratch. It is being extended from a strong text+audio backbone so the effect of the new visual representation can be measured cleanly.",
    )

    doc.add_heading("6. Expected Outputs and How to Read Them", level=1)
    add_bullets(
        doc,
        [
            "The .npy file is the learned visual representation for one utterance after face-crop ViT processing.",
            "The fold CSV files decide which utterances are used for training and validation.",
            "best_model.pt is the best checkpoint by validation accuracy.",
            "metrics.json summarizes accuracy, weighted F1, macro F1, and related scores.",
            "predictions_test.csv shows each utterance's predicted and actual label.",
            "confusion_matrix.csv shows which emotion classes are still being confused.",
        ],
    )
    add_para(
        doc,
        "What to look for: a useful result would reduce the tendency to predict neutral too often and improve minority emotion recall, especially sadness, anger, fear, and disgust. If face crops make the visual branch more focused, the confusion matrix should become less neutral-dominated than the full-frame run.",
    )

    doc.add_heading("6.1 Best Checkpoint Selection and Ablation Order", level=2)
    add_para(
        doc,
        "For this MELD experiment, the checkpoint should not be chosen only by validation accuracy. MELD is imbalanced, so accuracy can hide poor minority-class behavior. The safer default is to select the best checkpoint by weighted F1, and to inspect macro F1 as the class-balance check.",
    )
    add_bullets(
        doc,
        [
            "Weighted F1 is the best default for selecting the checkpoint because it stays sensitive to the overall class distribution.",
            "Macro F1 should be reported alongside it because it reveals whether minority emotions improved.",
            "Accuracy is still useful, but it should not be the only selection rule for an imbalanced MELD experiment.",
        ],
    )
    add_para(
        doc,
        "The next ablation sequence should be read as a small controlled study, not as three unrelated runs. First, run a video-only face-crop branch to see whether the visual signal has independent predictive value. Second, run the text+audio baseline to keep a clean non-visual reference. Third, run the full text+audio+video face-crop model to test whether the visual branch actually adds value on top of the already strong multimodal backbone. The doc should present the runs in that order so the comparison story is easy to explain.",
    )
    add_para(
        doc,
        "Recommended execution order for the current video-signal check: run the face-crop preparation step once, then run the dedicated video-only control Fold 2 and Fold 4 scripts, then analyze those two outputs, and finally run the tri-modal Fold 2 and Fold 4 scripts if you want to compare visual-only and multimodal behavior. The control-suite wrapper can be used only when you want the full video-only control pass in one command.",
    )
    add_table(
        doc,
        ["Run", "Modalities", "Question answered"],
        [
            ["Video-only control face crop", "video", "Does the cropped face stream carry useful emotion signal by itself?"],
            ["Text+audio-only reference", "text,audio", "What is the non-visual reference level for this same manifest/fold setup?"],
            ["Text+audio+video face crop", "text,audio,video", "Does the face-crop branch improve over the text+audio reference?"],
        ],
    )
    add_para(
        doc,
        "If video-only has signal but the full tri-modal model does not improve much, the next technical step should be to make the fusion more selective. In that case, a gated fusion block or a small auxiliary video loss would be more informative than just adding more epochs.",
    )

    doc.add_heading("6.2 Gated Fusion Follow-Up", level=2)
    add_para(
        doc,
        "The gated-fusion run is the next refinement after the weak video-only control. The idea is simple: do not force the model to trust the video branch all the time. Instead, let the model learn a gate that decides how much each modality should contribute to the final fused representation. This is useful when video has some signal, but not enough signal to dominate by itself.",
    )
    add_bullets(
        doc,
        [
            "The text and audio branches still provide the stable conversational backbone.",
            "The video branch contributes a separate face-crop representation from ViT embeddings.",
            "The gate learns a weight for each modality so the model can trust video only when it is actually helpful.",
            "This is safer than an auxiliary loss at this stage because the visual branch is still weaker than the text/audio path.",
            "The goal is not to force video to dominate. The goal is to make the model selective about when to use it.",
        ],
    )
    add_para(
        doc,
        "Student explanation: think of the gate as a learned switchboard. If the utterance has strong facial evidence, the gate can raise the video contribution. If the face crop is noisy or unhelpful, the gate can down-weight it and let text/audio carry more of the decision. That is why gated fusion is a better next step than an auxiliary loss right now.",
    )
    add_code(
        doc,
        """text_repr, audio_repr, video_repr -> gated fusion block
gate scores -> softmax weights
weighted modality mix -> fused vector
fused vector -> classifier -> emotion label""",
    )
    add_para(
        doc,
        "The important scientific point is that gated fusion tests whether the visual branch can contribute conditionally rather than unconditionally. In other words, the model should learn when to trust the face crop, not just whether the face crop exists.",
    )
    add_para(
        doc,
        "Error-pattern summary for the gated Fold 2 run: the confusion matrix still shows a strong pull toward neutral, but it is less collapsed than the video-only control. The strongest correct cells are neutral, surprise, fear, and joy, while sadness, anger, fear, and disgust still leak into neutral or nearby emotions. In practical terms, the gate is helping the model keep the useful facial signal, but it is not yet strong enough to eliminate neutral bias.",
    )
    add_bullets(
        doc,
        [
            "Neutral remains the dominant prediction, which means the class-imbalance problem is still present.",
            "Surprise, fear, and joy are learned better than in the video-only control, which suggests the gate is adding useful selectivity.",
            "Sadness, anger, and disgust still suffer from confusion with neutral and with each other, so the decision boundary is not yet sharp.",
            "The matrix is consistent with a model that has useful facial signal but still needs better calibration or fusion control.",
        ],
    )
    add_para(
        doc,
        "Student interpretation: the gate does not magically solve the imbalance problem. What it does is prevent the video branch from being ignored completely. That is why the matrix improves compared with video-only, but still shows a strong neutral bias. The model is learning better than before, but it still needs careful tuning or a stronger video-aware fusion design to separate the lower-frequency emotions cleanly.",
    )
    add_bullets(
        doc,
        [
            "Use gated fusion first because the video-only control showed usable but weak signal.",
            "Do not jump to auxiliary losses until you know the gate itself is useful.",
            "Keep the paper-aligned warm-start checkpoint so the comparison stays controlled.",
            "Keep the same MELD control folds so the only major change is the fusion mechanism.",
        ],
    )
    add_para(
        doc,
        "Recommended gated-fusion execution order: run the gated Fold 2 training script, analyze Fold 2, run the gated Fold 4 training script, analyze Fold 4, and only then compare the gated result to the video-only and tri-modal legacy face-crop runs. If you want a single wrapper, use the gated suite script after you have verified that the control folds and face-crop features already exist.",
    )
    add_table(
        doc,
        ["Script", "Purpose", "Why it is in the order"],
        [
            ["scripts/run_meld_vit_facecrop_gated_fold2.sh", "Train gated fusion on Fold 2", "First gated test of whether the face-crop branch can help selectively."],
            ["scripts/analyze_meld_vit_facecrop_gated_fold2.sh", "Export Fold 2 predictions and confusion matrix", "Shows whether the gate improves neutral bias and class balance."],
            ["scripts/run_meld_vit_facecrop_gated_fold4.sh", "Train gated fusion on Fold 4", "Checks whether the same gate behavior generalizes across folds."],
            ["scripts/analyze_meld_vit_facecrop_gated_fold4.sh", "Export Fold 4 predictions and confusion matrix", "Lets you compare Fold 4 against Fold 2 and the baseline."],
            ["scripts/run_meld_vit_facecrop_gated_suite.sh", "Run both folds in sequence", "Convenience wrapper only after the single-fold behavior is understood."],
        ],
    )

    doc.add_heading("6.3 Next Experiment: Gated Fusion + Video Auxiliary Loss", level=2)
    add_para(
        doc,
        "The next planned experiment keeps the same face-crop ViT pipeline and the same warm-start backbone, but adds a small auxiliary loss on the video branch. This is a stricter test than gated fusion alone because it does two things at once: it lets the model gate the modalities, and it also asks the video branch to solve the emotion task on its own with a small amount of extra training pressure.",
    )
    add_bullets(
        doc,
        [
            "Face-crop ViT stays as the visual feature source, so the visual representation is still face-centered.",
            "Gated fusion stays in place so the model can decide how much to trust video per utterance.",
            "The video auxiliary head adds a second prediction path from the video representation itself.",
            "The auxiliary loss weight is small, lambda = 0.1, so the video branch is helped without overpowering the main multimodal classifier.",
            "The learning rate is lowered to 2e-5 because the current gated run already shows useful signal and should now be refined more gently.",
            "Training now uses up to 8 epochs with early stopping patience = 2 so the experiment can stop as soon as validation weighted F1 stops improving.",
            "The best checkpoint is selected by validation weighted F1 because MELD is still imbalanced and weighted F1 is the safer checkpoint criterion than accuracy alone.",
        ],
    )
    add_para(
        doc,
        "Student explanation: think of the model as having two jobs at once. The main job is still to predict emotion from text, audio, and video after gated fusion. The extra job is to make the video branch itself produce a useful emotion prediction. The video auxiliary loss does not replace the main classifier; it nudges the visual branch to stay emotionally informative instead of becoming a passive passenger behind the text/audio backbone.",
    )
    add_para(
        doc,
        "Why this is worth trying now: the video-only control was weak, but the gated run showed that the face-crop signal is not useless. That means the visual branch has some signal, but it still needs stronger structure. The auxiliary loss gives the branch a direct training signal while the gate prevents the multimodal classifier from trusting video too much when it is not reliable.",
    )
    add_code(
        doc,
        """main logits from gated multimodal fusion
video logits from auxiliary video head
main loss = weighted CE on fused prediction
aux loss = weighted CE on video prediction
total loss = main loss + 0.1 * aux loss""",
    )
    add_para(
        doc,
        "The important optimization idea is that the auxiliary loss should stay small. If lambda is too large, the model may overfocus on video and hurt the text/audio backbone. If lambda is too small, the auxiliary branch will not matter. A value like 0.1 is a cautious middle ground for the first test. The 8-epoch cap with patience 2 keeps the run from wasting time after validation weighted F1 has already stopped improving.",
    )
    add_bullets(
        doc,
        [
            "Warm-start stays enabled, so the training begins from the paper-aligned MELD checkpoint rather than from random weights.",
            "Checkpoint selection is based on validation weighted F1, not validation accuracy.",
            "Early stopping watches validation weighted F1 and stops after 2 epochs without improvement.",
        ],
    )
    add_bullets(
        doc,
        [
            "The model still starts from the paper-aligned weighted-CE checkpoint, so the experiment remains warm-started.",
            "The new visual branch is not trained from scratch; it is refined on top of the existing baseline.",
            "The fold splits stay the same, so any result change should come from the new training objective rather than a new data partition.",
            "If the confusion matrix improves, that will suggest the visual branch is now contributing more directly and more consistently.",
        ],
    )
    add_para(
        doc,
        "Recommended execution order for the new Phase 1 experiment: first make sure the face-crop control folds already exist, then run the Fold 2 gated + video-aux script, analyze the Fold 2 predictions, then run the Fold 4 gated + video-aux script, and analyze the Fold 4 predictions. Only after that should you compare the confusion matrices against the earlier gated-only and video-only runs.",
    )
    add_table(
        doc,
        ["Script", "Role in the experiment", "Expected interpretation"],
        [
            ["scripts/run_meld_vit_facecrop_gated_video_aux_fold2.sh", "Train Fold 2 with gated fusion and a video auxiliary loss", "Checks whether a small extra video loss improves the weak facial cue branch."],
            ["scripts/analyze_meld_vit_facecrop_gated_video_aux_fold2.sh", "Export Fold 2 confusion matrix and predictions", "Shows whether the auxiliary loss reduces neutral bias or improves minority classes."],
            ["scripts/run_meld_vit_facecrop_gated_video_aux_fold4.sh", "Train Fold 4 with the same gated + auxiliary setup", "Checks whether the improvement generalizes on a more balanced fold."],
            ["scripts/analyze_meld_vit_facecrop_gated_video_aux_fold4.sh", "Export Fold 4 confusion matrix and predictions", "Lets you compare Fold 4 against Fold 2 and decide whether the method is stable."],
            ["scripts/run_meld_vit_facecrop_gated_video_aux_suite.sh", "Run both folds in sequence", "Convenience wrapper once the single-fold behavior is understood."],
        ],
    )
    add_para(
        doc,
        "What to expect if the experiment works: the model should keep the useful early gated behavior, but the video branch should become less passive. In practical terms, the confusion matrix should ideally move a little away from neutral-only fallback and show better separation for surprise, fear, joy, and possibly sadness or anger. If it does not, then the auxiliary loss is probably too weak or the visual representation still needs better design.",
    )

    doc.add_heading("7. Where This Fits in the Thesis Story", level=1)
    add_para(
        doc,
        "This face-crop branch is the last sensible Phase 1 visual ablation before moving the remaining effort into the courtroom-testimony novelty. It is a clean, scientifically defensible test because it changes only one variable: what part of the frame ViT sees.",
    )
    add_bullets(
        doc,
        [
            "If face crops help, freeze them as the final MELD visual design.",
            "If face crops do not help, keep the full-frame path as the simpler baseline and move on.",
            "Either outcome gives a stronger thesis argument than guessing which visual strategy is best.",
        ],
    )
    add_para(
        doc,
        "Short thesis summary: the project has already shown that full-frame facial cues are technically working but still biased toward neutral. The next controlled step is to ask whether face-crop ViT produces a more discriminative and more courtroom-relevant visual representation.",
    )

    return doc


def main() -> None:
    doc = build_doc()
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
