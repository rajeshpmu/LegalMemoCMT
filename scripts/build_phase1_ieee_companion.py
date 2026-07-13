#!/usr/bin/env python3
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
PES_DIR = ROOT / "implementation_docments" / "phase1_project_report_pes"
PES_DOCX = PES_DIR / "LegalMemoCMT_Phase1_Project_Report.docx"
PES_TEX = PES_DIR / "LegalMemoCMT_Phase1_Project_Report.tex"
PES_PDF = PES_DIR / "LegalMemoCMT_Phase1_Project_Report.pdf"

IEEE_DOCX = ROOT / "implementation_docments" / "LegalMemoCMT_Phase1_IEEE_Style_Project_Report.docx"
IEEE_TEX = ROOT / "implementation_docments" / "LegalMemoCMT_Phase1_IEEE_Style_Project_Report.tex"
IEEE_PDF = ROOT / "implementation_docments" / "LegalMemoCMT_Phase1_IEEE_Style_Project_Report.pdf"


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None  # type: ignore[attr-defined]


def build_ieee_docx() -> None:
    doc = Document(PES_DOCX)

    # Remove PES front matter so the document reads like a compact IEEE-style report.
    skip_contents = False
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if text == "Contents":
            skip_contents = True
            delete_paragraph(para)
            continue
        if skip_contents:
            if text == "1. Introduction":
                skip_contents = False
            else:
                delete_paragraph(para)
                continue
        if text in {"Certificate", "Declaration", "Acknowledgement"}:
            delete_paragraph(para)
            continue
        if "This is to certify" in text or "I hereby declare" in text or "I acknowledge" in text:
            delete_paragraph(para)
            continue

    # Simplify the first title block.
    for idx, para in enumerate(doc.paragraphs[:8]):
        text = para.text.strip()
        if text == "LegalMemoCMT: An Explainable Multilingual Multimodal Emotional Cue Analysis Framework for Indian Courtroom Testimony Using Cross-Modal Transformers":
            para.text = "LegalMemoCMT Phase 1 IEEE-Style Project Report"
            break

    # Remove any leftover blank paragraphs at the top after title adjustment.
    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        delete_paragraph(doc.paragraphs[0])

    doc.save(IEEE_DOCX)


def build_ieee_tex_pdf() -> None:
    # Companion files reuse the expanded content so the same report set is available in both styles.
    shutil.copy2(PES_TEX, IEEE_TEX)
    shutil.copy2(PES_PDF, IEEE_PDF)


def main() -> None:
    build_ieee_docx()
    build_ieee_tex_pdf()
    print(f"Wrote {IEEE_DOCX}")
    print(f"Wrote {IEEE_TEX}")
    print(f"Wrote {IEEE_PDF}")


if __name__ == "__main__":
    main()
