from __future__ import annotations

import csv
import json
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_Phase2_UCR_Enriched_Ledger_Explainer.docx"
LEDGER = ROOT / "data" / "phase2" / "source_manifests" / "case_candidate_ledger_ucr_enriched.csv"
VERIFIED_INVENTORY = ROOT / "data" / "processed" / "phase2" / "verified_case_inventory.csv"
HEARING_MANIFEST = ROOT / "data" / "processed" / "phase2" / "hearing_manifest.csv"
WITNESS_MANIFEST = ROOT / "data" / "phase2" / "source_manifests" / "witness_harvest_manifest_resolved.csv"
LEGALMELD_METADATA = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated" / "legalmeld_metadata_validated.csv"
ALIGNMENT_REVIEW = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated" / "alignment_review_sample.csv"
QUALITY_SUMMARY = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated" / "dataset_quality_summary.json"
TRAIN_CSV = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated" / "train.csv"
DEV_CSV = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated" / "dev.csv"
TEST_CSV = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated" / "test.csv"
VERIFICATION = ROOT / "phase2" / "ucr_case_verification.py"
INVENTORY = ROOT / "phase2" / "build_ucr_case_inventory.py"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    doc.add_paragraph()


def load_ledger_rows() -> list[dict[str, str]]:
    if not LEDGER.exists():
        return []
    with LEDGER.open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def load_inventory_rows() -> list[dict[str, str]]:
    if not VERIFIED_INVENTORY.exists():
        return []
    with VERIFIED_INVENTORY.open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def load_hearing_rows() -> list[dict[str, str]]:
    if not HEARING_MANIFEST.exists():
        return []
    with HEARING_MANIFEST.open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def load_witness_rows() -> list[dict[str, str]]:
    if not WITNESS_MANIFEST.exists():
        return []
    with WITNESS_MANIFEST.open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def load_legalmeld_rows() -> list[dict[str, str]]:
    if not LEGALMELD_METADATA.exists():
        return []
    with LEGALMELD_METADATA.open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def load_alignment_review_rows() -> list[dict[str, str]]:
    if not ALIGNMENT_REVIEW.exists():
        return []
    with ALIGNMENT_REVIEW.open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def load_json(path: Path) -> dict[str, object]:
    if not path.exists():
        return {}
    return json.loads(path.read_text())


def count_csv_rows(path: Path) -> int:
    if not path.exists():
        return 0
    with path.open(newline="", encoding="utf-8-sig") as f:
        return sum(1 for _ in csv.DictReader(f))


def normalize_count(value: object) -> str:
    try:
        return f"{int(value):,}"
    except Exception:
        return "0"


def summarize_rows(rows: list[dict[str, str]]) -> dict[str, str]:
    counts = Counter(row.get("verification_status", "") for row in rows)
    case_page = Counter(row.get("case_page_resolved", "") for row in rows)
    identity = Counter(row.get("case_identity_verified", "") for row in rows)
    return {
        "total_rows": normalize_count(len(rows)),
        "verified_rows": normalize_count(counts.get("verified", 0)),
        "unresolved_rows": normalize_count(counts.get("unresolved", 0)),
        "invalid_rows": normalize_count(counts.get("invalid_case_number", 0)),
        "resolved_yes": normalize_count(case_page.get("yes", 0)),
        "resolved_no": normalize_count(case_page.get("no", 0)),
        "identity_yes": normalize_count(identity.get("yes", 0)),
        "identity_no": normalize_count(identity.get("no", 0)),
        "video_yes": normalize_count(sum(1 for row in rows if (row.get("has_videos") or "").strip().lower() == "yes")),
        "transcript_yes": normalize_count(sum(1 for row in rows if (row.get("has_transcripts") or "").strip().lower() == "yes")),
        "court_recording_yes": normalize_count(sum(1 for row in rows if (row.get("has_court_recordings") or "").strip().lower() == "yes")),
    }


def first_row(rows: list[dict[str, str]], status: str) -> dict[str, str]:
    for row in rows:
        if (row.get("verification_status") or "").strip().lower() == status.lower():
            return row
    return {}


def first_hearing_row(rows: list[dict[str, str]], pairing_status: str) -> dict[str, str]:
    for row in rows:
        if (row.get("pairing_status") or "").strip().lower() == pairing_status.lower():
            return row
    return {}


def first_any_row(rows: list[dict[str, str]]) -> dict[str, str]:
    return rows[0] if rows else {}


def group_counts(rows: list[dict[str, str]], key: str) -> list[tuple[str, int]]:
    counts: Counter[str] = Counter((row.get(key) or "").strip() for row in rows)
    return counts.most_common()


def build_doc() -> Document:
    rows = load_ledger_rows()
    inventory_rows = load_inventory_rows()
    hearing_rows = load_hearing_rows()
    witness_rows = load_witness_rows()
    legalmeld_rows = load_legalmeld_rows()
    review_rows = load_alignment_review_rows()
    summary = load_json(QUALITY_SUMMARY)
    stats = summarize_rows(rows)
    verified = first_row(rows, "verified")
    unresolved = first_row(rows, "unresolved")
    invalid = first_row(rows, "invalid_case_number")

    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LegalMemoCMT Phase 2: How to Read case_candidate_ledger_ucr_enriched.csv")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Student-level explainer for the UCR verification layer and the promotion gate into the verified inventory.")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12.5)

    add_para(
        doc,
        "This document explains what the enriched UCR ledger does, how it differs from the human-curated candidate list, and how to explain it in a Phase 2 guidance call. The important idea is that the file is a verification layer, not a corpus-size estimate and not a download manifest.",
    )

    doc.add_heading("1. What This File Is", level=1)
    add_para(
        doc,
        "The file `data/phase2/source_manifests/case_candidate_ledger_ucr_enriched.csv` is the corrected UCR enrichment output. It is built to answer one question only: does the UCR page resolve to actual case-specific records, or is it just a generic page shell or a placeholder case number?",
    )
    add_bullets(
        doc,
        [
            "It keeps the original candidate row as a starting point.",
            "It adds a verification pass against actual UCR API responses.",
            "It records whether the case page resolved, whether the case identity was verified, and how many real records were found.",
            "It blocks placeholder and non-specific case numbers before any UCR request is made.",
        ],
    )

    doc.add_heading("2. Current Snapshot", level=1)
    add_table(
        doc,
        ["Metric", "Current value"],
        [
            ["Total rows", stats["total_rows"]],
            ["Verified rows", stats["verified_rows"]],
            ["Unresolved rows", stats["unresolved_rows"]],
            ["Invalid placeholder rows", stats["invalid_rows"]],
            ["Rows with case_page_resolved=yes", stats["resolved_yes"]],
            ["Rows with case_identity_verified=yes", stats["identity_yes"]],
            ["Rows with has_transcripts=yes", stats["transcript_yes"]],
            ["Rows with has_court_recordings=yes", stats["court_recording_yes"]],
            ["Rows with has_videos=yes", stats["video_yes"]],
        ],
    )
    add_para(
        doc,
        "This matters because the file now separates three different outcomes: verified real cases, unresolved real candidates, and invalid placeholder rows. That is the core discipline needed before anything is promoted into the inventory or the hearing manifest.",
    )

    doc.add_heading("3. How To Read The Columns", level=1)
    add_table(
        doc,
        ["Column", "Student-level meaning"],
        [
            ["case_page_resolved", "Whether the UCR case detail endpoint returned a real case detail response."],
            ["case_identity_verified", "Whether the returned records were grounded in actual case-specific metadata."],
            ["actual_record_count", "How many case-traceable records were extracted from the API responses."],
            ["transcript_record_count", "How many of those records were transcript records."],
            ["court_recording_count", "How many of those records were court recording or TAP-style records."],
            ["video_record_count", "How many records point to actual video media or downloadable media URLs."],
            ["tap_count", "How many records were TAP-type records."],
            ["first_record_date", "The earliest record date found among extracted records."],
            ["last_record_date", "The latest record date found among extracted records."],
            ["verification_status", "The final status: verified, unresolved, or invalid_case_number."],
            ["verification_notes", "Short human-readable reason for the status that was assigned."],
        ],
    )

    doc.add_heading("4. What The Status Values Mean", level=1)
    add_numbered(
        doc,
        [
            "Verified means the case page resolved and real case-specific records were extracted.",
            "Unresolved means a real candidate was checked, but the API did not return usable case detail or case-traceable records.",
            "Invalid_case_number means the row was a placeholder or non-specific label, so the request was rejected before it could be treated as evidence.",
        ],
    )

    doc.add_heading("5. Example Rows", level=1)
    example_rows: list[list[str]] = []
    for label, row in [
        ("Verified", verified),
        ("Unresolved", unresolved),
        ("Invalid placeholder", invalid),
    ]:
        example_rows.append(
            [
                label,
                row.get("tribunal", ""),
                row.get("case_family", ""),
                row.get("case_number", ""),
                row.get("verification_status", ""),
                row.get("case_page_resolved", ""),
                row.get("case_identity_verified", ""),
                row.get("actual_record_count", ""),
                row.get("verification_notes", ""),
            ]
        )
    add_table(
        doc,
        [
            "Row type",
            "Tribunal",
            "Case family",
            "Case number",
            "Verification status",
            "Page resolved",
            "Identity verified",
            "Actual records",
            "Notes",
        ],
        example_rows,
    )
    add_para(
        doc,
        "The verified row shows the positive case: the UCR endpoint returned grounded records. The unresolved row shows a real candidate that did not resolve to usable case detail. The invalid placeholder row shows the pre-request rejection logic working correctly.",
    )

    doc.add_heading("6. How The Verification Works In Code", level=1)
    add_para(
        doc,
        "The verification logic lives in `phase2/ucr_case_verification.py`. The most important function is `resolve_case_verification()`. It does not rely on the page shell alone. It uses API responses and then filters the returned records so only case-traceable items remain.",
    )
    add_numbered(
        doc,
        [
            "First, `is_placeholder_case_number()` rejects empty strings, `TBD`, `TO_BE_FILLED`, `ICTR case family`, and `MICT cases` before making any request.",
            "Next, `fetch_case_detail()` calls `/api/Summary/ByCaseDetail` to see whether the case detail actually resolves.",
            "If case detail exists, `fetch_case_docs_by_lang()` and `fetch_case_related_docs()` query the underlying record data.",
            "The returned rows are filtered so only case-family-matching, traceable records stay in the result.",
            "Counts are then derived from the actual filtered records, not from static page headings like `SELECTED DOCUMENTS`.",
            "Finally, `verification_status` is set to `verified`, `unresolved`, or `invalid_case_number` based on the actual outcome.",
        ],
    )
    add_para(
        doc,
        "The key student-level point is that page headings such as transcripts, videos, or court recordings are not used as proof by themselves. Only actual record metadata and record-level URLs count.",
    )

    doc.add_heading("7. How Promotion Into The Inventory Works", level=1)
    add_para(
        doc,
        "The inventory builder in `phase2/build_ucr_case_inventory.py` is the next gate. It reads the enriched ledger and keeps only rows that are actually verified unless the command is explicitly run with an override flag.",
    )
    add_bullets(
        doc,
        [
            "Verified rows are promoted into `verified_case_inventory.csv`.",
            "Unresolved rows are skipped because they do not yet prove usable case metadata.",
            "Invalid placeholders are not promoted at all.",
            "This is how the pipeline avoids false positives from generic UCR page templates.",
        ],
    )
    add_para(
        doc,
        "So the enriched ledger is not the final inventory. It is the proof layer that decides what is safe to move forward.",
    )

    doc.add_heading("8. Why The Current File Is Better Than The Older Flag-Based View", level=1)
    add_bullets(
        doc,
        [
            "Old page text flags could be fooled by generic headings.",
            "The new file checks actual UCR API outputs.",
            "The new file records resolution and traceability explicitly.",
            "The new file allows negative controls to stay negative.",
            "The new file keeps placeholder rows from inflating case counts or media counts.",
        ],
    )
    add_para(
        doc,
        "In plain language: the old approach could confuse a page template for evidence. The current enriched ledger prevents that mistake.",
    )

    doc.add_heading("9. How To Explain This In A Guidance Call", level=1)
    add_numbered(
        doc,
        [
            "Say that the ledger is a verification layer, not a download estimate.",
            "Say that only rows with real case-specific metadata are promoted.",
            "Say that placeholders are rejected before they can pollute the corpus pipeline.",
            "Say that unresolved rows are kept separate from verified rows.",
            "Say that the verified inventory is the file that should drive the next stage.",
        ],
    )
    add_para(
        doc,
        "A concise way to say it is: 'The enriched UCR ledger tells us whether a case really resolves to record-level evidence. If it does, we can promote it. If it does not, we leave it out. That keeps the inventory grounded in real records rather than generic page text.'",
    )

    doc.add_heading("10. What This Means For Phase 2", level=1)
    add_bullets(
        doc,
        [
            "The verified inventory remains the canonical source for later manifest building.",
            "The enriched ledger remains useful for traceability and validation.",
            "Placeholder or unresolved rows should not be used to estimate corpus size.",
            "The same discipline later carries over to Indian-source acquisition packs.",
        ],
    )
    add_para(
        doc,
        "This is a small but important part of the overall Phase 2 implementation: it protects the dataset from false positives before any hearing, witness, or utterance logic starts.",
    )

    doc.add_heading("10.1 How To Explain verified_case_inventory.csv", level=2)
    add_para(
        doc,
        "If you want to explain `data/processed/phase2/verified_case_inventory.csv` in a guidance call, you can say: 'This file keeps only cases with grounded metadata and removes placeholder-style false positives. In other words, the row has already passed the UCR verification layer, so it is safe to treat as the canonical case inventory for the next stage.'",
    )
    add_bullets(
        doc,
        [
            "Grounded metadata means the case actually resolved to real record-level evidence.",
            "Placeholder-style false positives are rows like `TO_BE_FILLED`, `TBD`, `ICTR case family`, and `MICT cases`.",
            "The verified inventory is the file the later manifest builders should trust.",
            "The enriched ledger helps explain how the row was verified, but the verified inventory is what moves forward.",
        ],
    )
    add_para(
        doc,
        "The verified inventory is more detailed than the ledger. The ledger tells you whether a case resolved; the inventory lists each grounded record that belongs to that verified case, so it is the case-level bridge into hearing-level work.",
    )

    doc.add_heading("10.2 What Each Column In Verified Inventory Means", level=2)
    add_table(
        doc,
        ["Column", "Meaning"],
        [
            ["inventory_id", "Unique ID for one inventory row."],
            ["source_record_id", "Source candidate row that produced the inventory row."],
            ["requested_case_number", "Case number requested from UCR."],
            ["resolved_case_number", "Canonical case number returned by UCR."],
            ["case_name", "Human-readable case family name."],
            ["case_number", "Case number attached to the verified row."],
            ["case_description", "Short tribunal-provided description of the case."],
            ["case_page_resolved", "Whether the case page resolved to actual case detail."],
            ["case_identity_verified", "Whether real case-specific records were extracted."],
            ["actual_record_count", "Number of grounded records extracted for the case."],
            ["transcript_record_count", "Number of extracted transcript records."],
            ["court_recording_count", "Number of extracted court-recording or TAP-style records."],
            ["video_record_count", "Number of records pointing to real video media or downloadable media URLs."],
            ["tap_count", "Number of TAP records in the verified case."],
            ["first_record_date", "Earliest record date seen in the verified row set."],
            ["last_record_date", "Latest record date seen in the verified row set."],
            ["document_title", "Title of the specific record copied into the inventory row."],
            ["document_type", "Record type such as transcript, decision, order, schedule, or TAP."],
            ["doc_signature_date", "Date associated with the specific record."],
            ["doc_source_desc", "Short source or origin description for the record."],
            ["document_path", "Normalized document path or URL for the record."],
            ["is_video", "Whether the row points to video media."],
            ["source_status", "Verification status copied from the UCR verification layer."],
            ["verification_status", "Final row status, usually verified for valid inventory rows."],
            ["record_case_number", "Case number recorded on the individual document row."],
            ["index_management_id", "UCR index-management identifier when present."],
            ["document_id", "UCR document identifier when present."],
            ["record_id", "UCR record identifier when present."],
            ["source_endpoint", "Which UCR endpoint produced the row."],
            ["verification_notes", "Why the row was accepted and how it was resolved."],
        ],
    )

    doc.add_heading("10.3 Examples Of Different Inventory Groups", level=2)
    add_para(
        doc,
        "The verified inventory is grouped by real case number and then repeated across many record types. That is why one verified case can produce thousands of rows. The groups below are examples from the current inventory.",
    )
    add_table(
        doc,
        ["Group example", "What it looks like in the inventory", "Student-level interpretation"],
        [
            [
                "ICTY / Karadzic Trial / IT-95-5/18",
                "4,327 rows in the current inventory; common document types include `Dec`, `TRA`, `Ord`, and `Sch`.",
                "This is one verified case with many grounded records, so the inventory is not one row per case but one row per record.",
            ],
            [
                "ICTY / Mladic Trial / IT-09-92",
                "2,589 rows in the current inventory; transcript records dominate, with decisions and orders also present.",
                "A single case family can contribute a large verified record set with mixed record types.",
            ],
            [
                "ICTY / Popovic et al. Trial / IT-05-88",
                "3,237 rows in the current inventory; the row set includes transcripts and many supporting records.",
                "The inventory preserves traceable evidence for later hearing grouping, not just a high-level case label.",
            ],
            [
                "ICTY / Perisic Trial / IT-04-81",
                "1,320 rows in the current inventory; the set still includes transcript and non-transcript record types.",
                "This shows the inventory can cover smaller verified cases too, not only the largest trial families.",
            ],
            [
                "ICTR / Bagosora et al. Trial / ICTR-98-41",
                "4,210 rows in the current inventory.",
                "This demonstrates tribunal diversity: the verified inventory is not limited to ICTY alone.",
            ],
        ],
    )
    add_para(
        doc,
        "Another useful way to explain the groups is by record type. Some inventory rows are transcripts, some are decisions, some are orders, and some are other record categories. The inventory keeps them because they are all grounded in the verified case metadata, even though only some of them later become hearing/video candidates.",
    )
    add_bullets(
        doc,
        [
            "Case grouping means all rows trace back to one verified case number.",
            "Record grouping means each row represents one record, not one hearing clip.",
            "Transcript-heavy groups are useful for alignment and utterance extraction.",
            "Decision-heavy or order-heavy groups still matter for provenance, but they are not automatically tri-modal training samples.",
        ],
    )

    doc.add_heading("10.4 How To Explain hearing_manifest.csv", level=2)
    add_para(
        doc,
        "The hearing manifest is the next bridge after the verified inventory. If the verified inventory says 'this case is real and has grounded records,' the hearing manifest says 'these records belong to a real hearing, and here is how the transcript row and video row line up.' It turns case-level evidence into hearing-level evidence with pairing metadata.",
    )
    add_bullets(
        doc,
        [
            "The hearing manifest does not invent new evidence.",
            "It groups verified case records into hearing/session units.",
            "It checks whether transcript and video records both exist for the same hearing.",
            "It keeps track of pairing confidence so only strong matches are treated as tri-modal candidates.",
            "It is the bridge from verified case metadata into utterance-level dataset building.",
        ],
    )

    doc.add_heading("10.5 What Each Column In hearing_manifest.csv Means", level=2)
    add_table(
        doc,
        ["Column", "Meaning"],
        [
            ["hearing_id", "Unique identifier for one hearing row."],
            ["tribunal", "Tribunal family such as ICTY or ICTR."],
            ["case_family", "Human-readable case family name."],
            ["case_number", "Canonical verified case number."],
            ["hearing_date", "Date associated with the hearing or record."],
            ["session_number", "Session or courtroom number if available."],
            ["record_title", "Title of the transcript or video record used to form the hearing row."],
            ["tap_or_record_id", "TAP or record identifier tied to the hearing."],
            ["record_detail_url", "Resolved record page or media URL for the hearing."],
            ["transcript_record_id", "Identifier for the transcript record, when present."],
            ["transcript_url", "Transcript URL, when present."],
            ["video_record_id", "Identifier for the video record, when present."],
            ["video_url", "Video URL, when present."],
            ["video_language", "Language code on the video record."],
            ["transcript_language", "Language code on the transcript record."],
            ["expected_duration_minutes", "Estimated duration for the paired hearing."],
            ["witness_name_or_code", "Public witness name, protected code, or unresolved witness label."],
            ["witness_identity_status", "How confident the pipeline is about the witness identity."],
            ["examination_type", "Direct, cross, redirect, recross, court questioning, or unknown."],
            ["video_verified", "Whether an actual video record was found."],
            ["transcript_verified", "Whether an actual transcript record was found."],
            ["pairing_status", "paired, video_only, transcript_only, or unpaired."],
            ["pairing_confidence", "High or medium pairing confidence where supported by the evidence."],
            ["eligible_for_trimodal_dataset", "YES only when transcript and video are both grounded and the hearing is a usable testimony row."],
            ["notes", "Short evidence summary for how the row was formed."],
        ],
    )

    doc.add_heading("10.6 Examples Of Different Hearing Groups", level=2)
    add_para(
        doc,
        "The hearing manifest shows three different kinds of rows in the current pipeline: paired hearings, transcript-only hearings, and video-only hearings. That is important because not every verified case record becomes a tri-modal training sample.",
    )
    hearing_examples: list[list[str]] = []
    for status in ["paired", "transcript_only", "video_only"]:
        row = first_hearing_row(hearing_rows, status) if hearing_rows else {}
        hearing_examples.append(
            [
                status,
                row.get("tribunal", ""),
                row.get("case_family", ""),
                row.get("case_number", ""),
                row.get("hearing_date", ""),
                row.get("record_title", ""),
                row.get("tap_or_record_id", ""),
                row.get("transcript_record_id", ""),
                row.get("video_record_id", ""),
                row.get("pairing_confidence", ""),
                row.get("eligible_for_trimodal_dataset", ""),
                row.get("notes", ""),
            ]
        )
    add_table(
        doc,
        [
            "Row type",
            "Tribunal",
            "Case family",
            "Case number",
            "Hearing date",
            "Record title",
            "TAP / record ID",
            "Transcript ID",
            "Video ID",
            "Pairing confidence",
            "Tri-modal eligible",
            "Notes",
        ],
        hearing_examples,
    )
    add_para(
        doc,
        "A paired hearing is the strongest kind of row because it has both transcript and video evidence. A transcript-only hearing still matters for text coverage, but it is not yet a tri-modal sample. A video-only hearing still shows media availability, but without the transcript it does not yet support MELD-style alignment.",
    )
    add_bullets(
        doc,
        [
            "Paired hearings show the full transcript-plus-video bridge.",
            "Transcript-only hearings show grounded text evidence but no matching video yet.",
            "Video-only hearings show media evidence but no aligned transcript yet.",
            "Eligible rows are only the paired rows that meet the testimony and verification checks.",
        ],
    )
    add_para(
        doc,
        "In the current manifest, the strongest examples are paired rows such as `ICTR-98-41` on `01/06/2007` and `IT-95-5/18` on `24/03/2016`, where both transcript and video records are present and the pairing confidence is high.",
    )

    doc.add_heading("10.7 How To Explain witness_harvest_manifest_resolved.csv", level=2)
    add_para(
        doc,
        "This file turns the grounded hearing rows into witness-level rows. Its job is to carry forward the witness or protected-code information only when that information can be supported by the hearing metadata, not by guessing or by using placeholder rows as corpus data.",
    )
    add_bullets(
        doc,
        [
            "The file is built from the hearing manifest, not from raw guesses.",
            "If a witness name or protected code cannot be recovered, the safe fallback is `UNRESOLVED_WITNESS`.",
            "The manifest still keeps the row because hearing-level evidence may still be useful later.",
            "Placeholder cases are not turned into witness rows at all.",
        ],
    )
    add_table(
        doc,
        ["Column", "Meaning"],
        [
            ["manifest_id", "Unique ID for one witness-manifest row."],
            ["tribunal", "Tribunal family such as ICTY or ICTR."],
            ["case_name", "Human-readable case family name."],
            ["case_number", "Canonical case number."],
            ["hearing_id", "The hearing row this witness row came from."],
            ["hearing_date", "Hearing date for the witness row."],
            ["witness_name_or_code", "Public witness name, protected code, or `UNRESOLVED_WITNESS`."],
            ["witness_type", "public_witness, protected_witness, or unresolved_witness."],
            ["speaker_role", "Usually Witness in this manifest."],
            ["examination_type", "Direct, cross, redirect, recross, court questioning, or unknown."],
            ["transcript_url", "Transcript URL when a transcript record exists."],
            ["video_url", "Video URL when a video record exists."],
            ["expected_duration_minutes", "Estimated duration inherited from the hearing manifest."],
            ["download_status", "Whether both modalities are resolved or only one is present."],
            ["annotation_status", "Current annotation stage, usually not started at this point."],
            ["utterance_count", "Reserved for later utterance-level counting."],
            ["emotion_label_status", "Placeholder status for later emotion annotation."],
            ["credibility_label_status", "Placeholder status for later credibility annotation."],
            ["source_record_id", "The source TAP or record ID from the hearing manifest."],
            ["pairing_confidence", "How strong the transcript-video pairing evidence is."],
            ["eligible_for_trimodal_dataset", "YES only if the row is usable for tri-modal work."],
            ["notes", "Short traceability note copied from the hearing layer."],
        ],
    )
    witness_example = first_any_row(witness_rows)
    add_para(
        doc,
        "In the current resolved witness manifest, most rows are still `UNRESOLVED_WITNESS`, which is the correct conservative behavior. The pipeline does not invent names when a witness identity is not safely grounded.",
    )
    add_table(
        doc,
        ["Example", "What it shows"],
        [
            ["Current row count", f"{count_csv_rows(WITNESS_MANIFEST):,} witness rows."],
            ["Witness type pattern", "Mostly `unresolved_witness`, with `public_witness` or `protected_witness` only when the evidence supports it."],
            ["Paired examples", "Rows with `pairing_confidence=high` and `eligible_for_trimodal_dataset=YES` carry forward the strongest hearing-level evidence."],
            ["Representative row", f"{witness_example.get('tribunal', '')} / {witness_example.get('case_number', '')} / {witness_example.get('witness_name_or_code', '')} / {witness_example.get('eligible_for_trimodal_dataset', '')}"],
        ],
    )
    add_para(
        doc,
        "So if you want to explain this file in one sentence, you can say: it converts each grounded hearing into a witness-level row, but it keeps `UNRESOLVED_WITNESS` when the witness identity is not strong enough to name safely.",
    )

    doc.add_heading("10.8 How To Explain legalmeld_metadata_validated.csv", level=2)
    add_para(
        doc,
        "This is the MELD-style master table. Each row is one aligned utterance with text, audio, video, speaker, witness, and quality metadata. It is the first file that really looks like a courtroom equivalent of MELD because it is utterance-level rather than hearing-level.",
    )
    add_table(
        doc,
        ["Column", "Meaning"],
        [
            ["utterance_id", "Unique ID for one utterance row."],
            ["hearing_id", "The hearing that the utterance belongs to."],
            ["tribunal", "Tribunal family such as ICTY or ICTR."],
            ["case_number", "Canonical case number."],
            ["case_family", "Human-readable case family name."],
            ["hearing_date", "Hearing date."],
            ["witness_id", "Witness code or witness name used for the row."],
            ["speaker_role", "Who is speaking: Prosecutor, Witness, Judge, or Defence."],
            ["speaker_name", "Speaker label taken from the transcript."],
            ["examination_type", "Direct, cross, redirect, recross, court questioning, or unknown."],
            ["utterance_text", "The transcript text for the utterance."],
            ["start_time", "Aligned start timestamp."],
            ["end_time", "Aligned end timestamp."],
            ["duration_ms", "Utterance duration in milliseconds."],
            ["video_clip", "Relative path to the video clip."],
            ["audio_clip", "Relative path to the audio clip."],
            ["transcript_source", "Original transcript source row or identifier."],
            ["alignment_status", "matched, fuzzy, or fallback."],
            ["alignment_score", "Numeric score used to rank the alignment."],
            ["alignment_method", "Exact, fuzzy, or fallback alignment method."],
            ["alignment_confidence", "HIGH, MEDIUM, or LOW alignment confidence."],
            ["asr_text", "The ASR transcript used during alignment."],
            ["transcript_text_normalized", "Normalized transcript used for matching."],
            ["text_similarity", "Similarity score between transcript and ASR text."],
            ["clip_duration_seconds", "Length of the exported clip."],
            ["word_timestamp_count", "Number of aligned word timestamps."],
            ["manual_review_required", "YES when the row should be checked manually."],
            ["split_group_id", "Group ID used to keep related utterances together."],
            ["split_strategy", "The splitting rule used to avoid leakage."],
            ["quality_tier", "A, B, C, or REJECT-style quality grouping."],
            ["audio_present", "Whether audio was found for the clip."],
            ["audio_rms", "Audio loudness metric."],
            ["silence_ratio", "How much of the clip is silent."],
            ["clipping_ratio", "How much clipping is present in the audio."],
            ["sample_rate", "Audio sample rate."],
            ["audio_validation_status", "valid, silent, clipped, corrupt, or missing."],
            ["face_detected", "Whether a face was detected in the clip."],
            ["face_visible_ratio", "How much of the clip shows a visible face."],
            ["shot_type", "Courtroom shot description."],
            ["speaker_visible", "Whether the speaker is visibly present."],
            ["video_quality_status", "Whether the video clip passed the video check."],
            ["emotion", "Reserved for later annotation."],
            ["credibility", "Reserved for later annotation."],
            ["split", "train, dev, test, or review."],
        ],
    )
    if legalmeld_rows:
        lm_example = first_any_row(legalmeld_rows)
        add_table(
            doc,
            ["Example row", "Interpretation"],
            [
                ["Matched / high", "A row like a witness answer with `alignment_status=matched`, `alignment_confidence=HIGH`, `quality_tier=B`, and valid audio/video indicates the cleanest utterance-level sample."],
                ["Fuzzy / medium", "A row like a witness or prosecutor utterance with `alignment_status=fuzzy` and `alignment_confidence=MEDIUM` is still usable, but it is less exact than a matched row."],
                ["Fallback / low", "A row like `Please continue.` with `alignment_status=fallback`, `alignment_confidence=LOW`, and `quality_tier=REJECT` is sent to review instead of being treated as a strong training sample."],
                ["Current dataset size", f"{count_csv_rows(LEGALMELD_METADATA):,} utterance rows across {summary.get('cases_represented', 0)} cases and {summary.get('hearings_represented', 0)} hearings."],
                ["Representative row", f"{lm_example.get('case_number', '')} / {lm_example.get('speaker_role', '')} / {lm_example.get('alignment_status', '')} / {lm_example.get('split', '')}"],
            ],
        )
    add_bullets(
        doc,
        [
            f"Current alignment confidence counts: {summary.get('alignment_confidence_counts', {})}.",
            f"Current quality-tier counts: {summary.get('quality_tier_counts', {})}.",
            f"Current split counts: {summary.get('split_counts', {})}.",
            f"Current speaker-role counts: {summary.get('speaker_role_counts', {})}.",
        ],
    )
    add_para(
        doc,
        "In the current run, the best explanation is that this file is already the MELD-style bridge, but it is still small and still dominated by medium-confidence aligned rows. That is useful for proving the method, but it is not yet the final large-scale corpus.",
    )

    doc.add_heading("10.9 How To Explain alignment_review_sample.csv", level=2)
    add_para(
        doc,
        "This file is the manual-check shortlist. It contains the rows that are not strong enough to trust blindly, usually because the alignment fell back or the confidence is too low. It is not a training file. It is a debugging and improvement file.",
    )
    add_table(
        doc,
        ["Column", "Meaning"],
        [
            ["All major metadata columns", "The review sample keeps the same columns as the main LegalMELD metadata so you can inspect the row in context."],
            ["manual_review_required", "Usually YES for the rows placed here."],
            ["alignment_status", "Often fallback in the review sample."],
            ["alignment_confidence", "Often LOW in the review sample."],
            ["quality_tier", "Often REJECT or another low-confidence tier."],
            ["video_quality_status", "Used to see whether the media itself is the problem."],
            ["audio_validation_status", "Used to see whether the audio itself is the problem."],
        ],
    )
    review_example = first_any_row(review_rows)
    add_table(
        doc,
        ["Example", "What it shows"],
        [
            ["Current review size", f"{count_csv_rows(ALIGNMENT_REVIEW):,} rows."],
            ["Current pattern", "The review set is small and is mainly used to expose weak utterance boundaries and fallback alignments."],
            ["Representative row", f"{review_example.get('utterance_id', '')} / {review_example.get('alignment_status', '')} / {review_example.get('alignment_confidence', '')} / {review_example.get('quality_tier', '')}"],
        ],
    )
    add_para(
        doc,
        "So the right student-level way to say it is: this file shows which utterances need review so the pipeline can be improved before scaling. It is the list you inspect when you want to fix alignment quality instead of just collecting more data.",
    )

    doc.add_heading("10.10 How To Explain dataset_quality_summary.json", level=2)
    add_para(
        doc,
        "This JSON file is the aggregate scorecard for the current validated run. It does not contain clips or utterances themselves. It tells you how strong the batch was overall and how much of it is safe to use immediately.",
    )
    add_table(
        doc,
        ["Field", "Current meaning"],
        [
            ["alignment_confidence_counts", str(summary.get("alignment_confidence_counts", {}))],
            ["quality_tier_counts", str(summary.get("quality_tier_counts", {}))],
            ["split_counts", str(summary.get("split_counts", {}))],
            ["speaker_role_counts", str(summary.get("speaker_role_counts", {}))],
            ["audio_valid_count", str(summary.get("audio_valid_count", 0))],
            ["video_valid_count", str(summary.get("video_valid_count", 0))],
            ["manual_review_count", str(summary.get("manual_review_count", 0))],
            ["split_leakage_violations", str(summary.get("split_leakage_violations", 0))],
            ["total_clip_hours", str(summary.get("total_clip_hours", 0))],
            ["average_clip_duration_seconds", str(summary.get("average_clip_duration_seconds", 0))],
            ["distinct_witnesses", str(summary.get("distinct_witnesses", 0))],
            ["hearings_represented", str(summary.get("hearings_represented", 0))],
            ["cases_represented", str(summary.get("cases_represented", 0))],
        ],
    )
    add_bullets(
        doc,
        [
            f"High-confidence alignments: {summary.get('high_confidence_alignments', 0)}.",
            f"Medium-confidence alignments: {summary.get('alignment_confidence_counts', {}).get('MEDIUM', 0) if isinstance(summary.get('alignment_confidence_counts'), dict) else 0}.",
            f"Low-confidence alignments: {summary.get('alignment_confidence_counts', {}).get('LOW', 0) if isinstance(summary.get('alignment_confidence_counts'), dict) else 0}.",
            f"Manual-review rows: {summary.get('manual_review_count', 0)}.",
            f"Split leakage violations: {summary.get('split_leakage_violations', 0)}.",
        ],
    )
    add_para(
        doc,
        "In the current run, the summary says the audio and video files are valid across the batch, but the alignment quality is mixed. That is why the dataset is useful as a bootstrap set but still needs improvement before final fine-tuning.",
    )

    doc.add_heading("10.11 How To Explain train.csv / dev.csv / test.csv", level=2)
    add_para(
        doc,
        "These are the leakage-aware split files. They are derived from the validated utterance metadata and separated by group so the same hearing or closely related utterance cluster does not leak across train, development, and test.",
    )
    add_bullets(
        doc,
        [
            "train.csv is the main learning partition.",
            "dev.csv is the tuning and model-selection partition.",
            "test.csv is the held-out evaluation partition.",
            "The split should be understood as group-based, not random row shuffling.",
        ],
    )
    add_table(
        doc,
        ["Column", "Meaning"],
        [
            ["Dialogue_ID", "Hearing-level ID for the utterance cluster."],
            ["Utterance_ID", "Unique utterance ID."],
            ["Utterance", "The aligned utterance text."],
            ["Speaker", "Speaker label taken from the transcript."],
            ["SpeakerRole", "Prosecutor, Witness, Judge, or Defence."],
            ["WitnessID", "Witness code or witness name used in the split row."],
            ["CaseNumber", "Canonical case number."],
            ["CaseFamily", "Human-readable case family name."],
            ["HearingDate", "Date of the hearing."],
            ["AlignmentStatus", "matched, fuzzy, or fallback."],
            ["AlignmentConfidence", "HIGH, MEDIUM, or LOW."],
            ["VideoPath", "Path to the utterance clip."],
            ["AudioPath", "Path to the utterance audio."],
        ],
    )
    add_table(
        doc,
        ["Split file", "Current row count", "Student-level meaning"],
        [
            ["train.csv", f"{count_csv_rows(TRAIN_CSV):,}", "Main learning set with the largest number of utterances."],
            ["dev.csv", f"{count_csv_rows(DEV_CSV):,}", "Small tuning set used to compare model settings."],
            ["test.csv", f"{count_csv_rows(TEST_CSV):,}", "Held-out set used only after the model is fixed."],
        ],
    )
    add_para(
        doc,
        "The key thing to say in a guidance call is that these files are leakage-aware partitions, not random slices. That is why the split leakage count in the summary is zero: the same grouped evidence is not duplicated across partitions.",
    )
    add_bullets(
        doc,
        [
            f"Current split sizes: train {count_csv_rows(TRAIN_CSV)}, dev {count_csv_rows(DEV_CSV)}, test {count_csv_rows(TEST_CSV)}.",
            f"Current leakage violations: {summary.get('split_leakage_violations', 0)}.",
            "The split files are the evaluation-ready view of the validated utterance dataset.",
        ],
    )

    doc.add_heading("11. Short Answer You Can Use Verbally", level=1)
    add_para(
        doc,
        "Yes, the enriched ledger now shows whether the UCR page really resolves to case-specific records. Verified rows are grounded in actual API-returned metadata; unresolved rows are real candidates that did not resolve cleanly; placeholder rows are rejected up front. That is why the file is safe to use as a validation layer, but not as a download estimate source.",
    )

    doc.add_heading("12. Code Paths To Mention", level=1)
    add_table(
        doc,
        ["Script or function", "Role"],
        [
            ["`phase2/ucr_case_verification.py` -> `resolve_case_verification()`", "Does the actual UCR verification and record filtering."],
            ["`phase2/ucr_case_verification.py` -> `validation_summary()`", "Returns the structured verification fields that go into the enriched CSV."],
            ["`phase2/build_ucr_case_inventory.py`", "Promotes only verified rows into the canonical inventory."],
            ["`phase2/run_enrich_case_ledger_from_ucr_site.sh`", "Wrapper that regenerates the enriched ledger."],
            ["`phase2/run_build_ucr_case_inventory.sh`", "Wrapper that regenerates the verified inventory."],
        ],
    )

    doc.add_heading("13. Bottom-Line Judgment", level=1)
    add_para(
        doc,
        "The file is doing the right job now. It tells you whether the UCR case really resolves to case-specific records before anything is promoted. That is the correct behavior for Phase 2, because it prevents generic page shells from being mistaken for real corpus evidence.",
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
