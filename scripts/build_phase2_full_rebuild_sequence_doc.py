from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_Phase2_Full_Rebuild_Sequence_Student_Guide.docx"
VALIDATED_ROOT = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated"
SUMMARY_JSON = VALIDATED_ROOT / "dataset_quality_summary.json"
VERIFIED_INVENTORY = ROOT / "data" / "processed" / "phase2" / "verified_case_inventory.csv"
HEARING_MANIFEST = ROOT / "data" / "processed" / "phase2" / "hearing_manifest.csv"
WITNESS_MANIFEST = ROOT / "data" / "phase2" / "source_manifests" / "witness_harvest_manifest_resolved.csv"
LEGALMELD_METADATA = VALIDATED_ROOT / "legalmeld_metadata_validated.csv"
ALIGNMENT_REVIEW = VALIDATED_ROOT / "alignment_review_sample.csv"
TRAIN_CSV = VALIDATED_ROOT / "train.csv"
DEV_CSV = VALIDATED_ROOT / "dev.csv"
TEST_CSV = VALIDATED_ROOT / "test.csv"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    doc.add_paragraph()


def read_summary() -> dict[str, object]:
    if not SUMMARY_JSON.exists():
        return {}
    try:
        return json.loads(SUMMARY_JSON.read_text(encoding="utf-8"))
    except Exception:
        return {}


def fmt_int(value: object) -> str:
    try:
        return f"{int(value):,}"
    except Exception:
        return "0"


def fmt_float(value: object, digits: int = 2) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return f"{0.0:.{digits}f}"


def add_stage_table(doc: Document) -> None:
    add_table(
        doc,
        ["Order", "Script", "Main input", "Main output", "Student-level meaning"],
        [
            [
                "1",
                "phase2/run_deduplicate_case_ledger.sh",
                "Merged candidate ledgers and manual case shortlists",
                "Cleaned deduplicated candidate ledger",
                "Removes repeated case rows so the next steps operate on one clean candidate list instead of duplicates.",
            ],
            [
                "2",
                "phase2/run_crawl_official_ucr_case_pages.sh",
                "Official UCR case pages and source URLs",
                "Crawled page evidence cache",
                "Collects page-level evidence from the tribunal portal before any record is promoted into the corpus pipeline.",
            ],
            [
                "3",
                "phase2/run_enrich_case_ledger_from_ucr_site.sh",
                "Deduplicated candidate ledger plus crawled UCR pages",
                "case_candidate_ledger_ucr_enriched.csv and validation report",
                "Checks whether each row really resolves to case-specific data or is only a generic page shell.",
            ],
            [
                "4",
                "phase2/run_build_tribunal_manifest_from_ledger.sh",
                "Curated ledger rows",
                "Tribunal manifest",
                "Converts the source shortlist into a tribunal-oriented planning layer that still stays grounded in verified rows.",
            ],
            [
                "5",
                "phase2/run_build_witness_manifest_from_ledger.sh",
                "Curated ledger rows and witness clues from transcripts",
                "Witness manifest",
                "Extracts witness-level rows while preserving protected codes and unresolved identities instead of guessing.",
            ],
            [
                "6",
                "phase2/run_build_ucr_case_inventory.sh",
                "Enriched ledger",
                "verified_case_inventory.csv",
                "Creates the canonical case inventory and drops placeholder-style false positives before any later download logic.",
            ],
            [
                "7",
                "phase2/run_split_ucr_inventory_by_media_type.sh",
                "Verified inventory",
                "Video-bearing and transcript-only manifests",
                "Separates media-rich cases from transcript-only cases so the pipeline does not treat every row as tri-modal.",
            ],
            [
                "8",
                "phase2/run_build_hearing_witness_manifests.sh",
                "Verified inventory plus UCR media evidence",
                "hearing_manifest.csv and witness_harvest_manifest_resolved.csv",
                "Turns case-level proof into hearing-level and witness-level rows that are actually usable downstream.",
            ],
            [
                "9",
                "phase2/run_validate_trimodal_corpus.sh",
                "Resolved hearing and witness manifests",
                "Validated hearing/witness reports",
                "Applies quality checks so only grounded transcript-video pairs can proceed to utterance extraction.",
            ],
            [
                "10",
                "phase2/run_build_legalmeld_dataset_validated.sh --skip-existing",
                "Validated hearing manifest, media manifest, selection manifest, cached media",
                "legalmeld_metadata_validated.csv, train/dev/test CSVs, review sample, quality summary",
                "Builds the MELD-style utterance dataset with transcript parsing, ASR alignment, clip extraction, and split-aware grouping.",
            ],
            [
                "11",
                "phase2/run_phase2_split_manifest.sh",
                "Phase 2 dataset CSV",
                "Split manifest",
                "Writes the split column explicitly so train/dev/test assignment can be inspected separately from the clip builder.",
            ],
            [
                "12",
                "phase2/run_phase2_sanitize_manifest.sh",
                "Split manifest",
                "Clean split manifest",
                "Removes rows that should not be used for training and prepares the data for audio extraction.",
            ],
            [
                "13",
                "phase2/run_phase2_extract_audio.sh",
                "Clean split manifest and video clips",
                "Tri-modal manifest with audio paths",
                "Creates the audio branch of the multimodal sample by extracting WAV files from the video sources.",
            ],
        ],
    )


def add_file_tree(doc: Document) -> None:
    add_code = """
data/processed/phase2/
  verified_case_inventory.csv
  hearing_manifest.csv
  legalmeld_validated/
    legalmeld_metadata_validated.csv
    alignment_review_sample.csv
    dataset_quality_summary.json
    train.csv
    dev.csv
    test.csv
data/phase2/source_manifests/
  case_candidate_ledger.csv
  case_candidate_ledger_ucr_enriched.csv
  witness_harvest_manifest_resolved.csv
"""
    for line in add_code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)


def build_doc() -> Document:
    s = read_summary()
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LegalMemoCMT Phase 2 Full Rebuild Sequence")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Student guide for rerunning the entire Phase 2 tribunal bootstrap pipeline from the candidate ledger to the utterance-level LegalMELD dataset."
    )
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12.5)

    add_para(
        doc,
        "This guide explains the full Phase 2 rebuild as a chain of dependent steps. The main idea is that each stage adds one layer of evidence and one layer of structure. The process starts with a human-curated candidate ledger, then it verifies UCR case pages, then it turns verified cases into grounded hearing and witness manifests, and finally it builds the utterance-level LegalMELD dataset with aligned text, audio, and video clips. The tribunal corpus is the bootstrap dataset; the same logic can later be reused for the Indian courtroom adaptation layer.",
    )

    doc.add_heading("What You Are Rebuilding", level=1)
    add_para(
        doc,
        "When you rerun Phase 2, you are not just downloading files again. You are rebuilding the evidence chain that proves each sample is grounded in a real hearing record. That chain is important because the downstream multimodal dataset should only contain rows that come from actual tribunal metadata, actual transcript segments, and actual clip extraction from validated media.",
    )
    add_bullets(
        doc,
        [
            "The candidate ledger is a shortlist, not proof.",
            "The enriched ledger is a verification layer, not a corpus estimate.",
            "The verified inventory is the first canonical source of truth.",
            "The hearing and witness manifests turn records into grounded corpora units.",
            "The LegalMELD exporter turns grounded hearing material into utterance-level training rows.",
        ],
    )

    doc.add_heading("Current Output Snapshot", level=1)
    add_table(
        doc,
        ["Artifact", "Path", "Why it matters"],
        [
            ["verified_case_inventory.csv", str(VERIFIED_INVENTORY.relative_to(ROOT)), "Canonical case-level inventory used to decide what can enter the later stages."],
            ["hearing_manifest.csv", str(HEARING_MANIFEST.relative_to(ROOT)), "Hearing-level grounding layer with transcript and video pairing evidence."],
            ["witness_harvest_manifest_resolved.csv", str(WITNESS_MANIFEST.relative_to(ROOT)), "Witness-level layer that keeps protected or unresolved identities traceable."],
            ["legalmeld_metadata_validated.csv", str(LEGALMELD_METADATA.relative_to(ROOT)), "The utterance-level master CSV that contains the final aligned rows."],
            ["alignment_review_sample.csv", str(ALIGNMENT_REVIEW.relative_to(ROOT)), "Diagnostic rows that need manual checking before scaling the corpus."],
            ["dataset_quality_summary.json", str(SUMMARY_JSON.relative_to(ROOT)), "Machine-readable summary of how many rows were kept, rejected, and split."],
            ["train.csv / dev.csv / test.csv", str(VALIDATED_ROOT.relative_to(ROOT)), "Leakage-aware split files used for model training and evaluation."],
        ],
    )
    if s:
        add_para(
            doc,
            f"Current validated snapshot: {fmt_int(s.get('total_utterances', 0))} total utterances, {fmt_int(s.get('high_confidence_alignments', 0))} high-confidence alignments, {fmt_int(s.get('fuzzy_alignments', 0))} medium-confidence alignments, {fmt_int(s.get('rejected_utterances', 0))} rejected utterances, and {fmt_float(s.get('total_clip_hours', 0.0), 3)} total clip hours.",
        )

    doc.add_heading("Full Rebuild Order", level=1)
    add_para(
        doc,
        "Use the steps below in order if you want to rebuild the whole Phase 2 tribunal bootstrap path from the source manifests up to the final utterance-level dataset. The student-level explanation tells you what each step is doing, why it exists, and what file to inspect after the step finishes.",
    )
    add_stage_table(doc)

    doc.add_heading("How To Understand Each Stage", level=1)
    sections = [
        (
            "1. Deduplicate the case ledger",
            "The deduplication step is the cleanup pass for the candidate shortlist. Student-wise, this means you first remove repeated or overlapping rows so every case appears once in the working ledger. If you skip this step, the same case may be crawled more than once, which makes the later validation output harder to interpret.",
            [
                "Input: merged or manually curated candidate rows.",
                "Output: a cleaned ledger with repeated case rows removed.",
                "Inspect: `data/phase2/source_manifests/case_candidate_ledger.csv` and the deduplicated ledger output.",
            ],
        ),
        (
            "2. Crawl official UCR case pages",
            "This is the discovery pass. The crawler visits official tribunal pages and records the page evidence that belongs to each candidate case. This is still not corpus data; it is only the evidence that will later help decide whether the row really resolves to real records. The main idea is to separate page discovery from dataset trust.",
            [
                "Input: deduplicated candidate ledger and source URLs.",
                "Output: a crawl cache or page evidence cache.",
                "Inspect: the crawl logs and the source-site links stored in the candidate ledger.",
            ],
        ),
        (
            "3. Enrich the ledger from UCR",
            "The enrichment step is where the pipeline becomes strict. It checks whether the UCR page really resolves to case-specific metadata or whether it is just a generic template. This stage is critical because page headings alone are not evidence. The code should reject placeholders and unresolved family labels before promotion.",
            [
                "Input: crawled UCR page evidence and the cleaned candidate ledger.",
                "Output: `case_candidate_ledger_ucr_enriched.csv` plus the validation report.",
                "Inspect: the enrichment status fields and the negative-control rows.",
            ],
        ),
        (
            "4. Build the tribunal manifest",
            "The tribunal manifest is a planning bridge. It takes the verified case evidence and organizes it around tribunal and case structure so the rest of the corpus logic can operate at a higher level. Think of it as the map that says, 'these are the grounded tribunal rows we can work with next.'",
            [
                "Input: curated or verified case rows.",
                "Output: tribunal-level manifest rows.",
                "Inspect: whether the case family, tribunal, and case number fields stay grounded.",
            ],
        ),
        (
            "5. Build the witness manifest",
            "The witness manifest is where witness identity logic starts to matter. The goal is not to deanonymize protected witnesses, but to keep public witness names and protected codes traceable enough for downstream alignment. If the identity cannot be resolved, the pipeline should preserve `UNRESOLVED_WITNESS` rather than inventing a name.",
            [
                "Input: the verified case and hearing evidence.",
                "Output: resolved witness manifest rows.",
                "Inspect: witness code, speaker role, and the evidence notes.",
            ],
        ),
        (
            "6. Build the verified inventory",
            "This is the most important case-level file. It is the canonical inventory because it only keeps cases with grounded metadata. In student terms, this is the file where the pipeline stops trusting the rough shortlist and starts trusting only verified records. Placeholder-style false positives are removed here.",
            [
                "Input: `case_candidate_ledger_ucr_enriched.csv`.",
                "Output: `verified_case_inventory.csv`.",
                "Inspect: that the row still has grounded case metadata and actual record evidence.",
            ],
        ),
        (
            "7. Split the inventory by media type",
            "This stage separates cases that can support video-based work from cases that are transcript only. That matters because a multimodal project should not count every hearing as tri-modal by default. The split also keeps the later download logic efficient and prevents text-only rows from being treated as video rows.",
            [
                "Input: verified inventory.",
                "Output: media-specific manifests.",
                "Inspect: whether video-bearing rows and transcript-only rows are separated correctly.",
            ],
        ),
        (
            "8. Build hearing and witness manifests",
            "This step converts case-level proof into record-level structure. A hearing row should only appear when the record is grounded in real transcript and video evidence. A witness row should only be created when the witness or protected code can be linked to that hearing without turning placeholders into corpus data.",
            [
                "Input: verified inventory plus resolved media evidence.",
                "Output: hearing manifest and resolved witness manifest.",
                "Inspect: pairing evidence, case numbers, hearing dates, and witness identities.",
            ],
        ),
        (
            "9. Validate the tri-modal corpus",
            "Validation is the safety gate before the utterance-level dataset is built. The purpose is to make sure that the hearing and witness layers are internally consistent and that the audio/video/transcript evidence really belongs together. This prevents the final model dataset from inheriting bad assumptions from the source manifests.",
            [
                "Input: hearing and witness manifests.",
                "Output: validation results and rejected examples.",
                "Inspect: the validation summary and any rows flagged for review.",
            ],
        ),
        (
            "10. Build the LegalMELD dataset",
            "This is the main multimodal transformation. The builder reads transcript text, segments it into utterances, aligns each utterance with ASR word timestamps, extracts the clip window with ffmpeg, and writes one row per utterance. This is where the dataset stops being a hearing list and starts becoming a training dataset.",
            [
                "Input: validated hearing manifest, media manifest, and selection manifest.",
                "Output: `legalmeld_metadata_validated.csv`, `train.csv`, `dev.csv`, `test.csv`, `alignment_review_sample.csv`, and `dataset_quality_summary.json`.",
                "Inspect: the alignment confidence, text similarity, clip duration, and split assignment.",
            ],
        ),
        (
            "11. Write the split manifest",
            "The split manifest makes the train/dev/test assignment explicit. That helps you inspect the split before model training and check for leakage across witnesses or hearings. It also makes the work easier to explain during a guidance call because you can point to the exact rows that will train, validate, or test the model.",
            [
                "Input: Phase 2 dataset CSV.",
                "Output: split-bearing manifest.",
                "Inspect: whether group-based splitting kept related utterances together.",
            ],
        ),
        (
            "12. Sanitize the split manifest",
            "Sanitization is a protection step. It removes or separates rows that should not be used for training and keeps the cleaned outputs ready for the audio extraction stage. The important idea is that cleaning here should not invent new data; it should only remove or normalize rows that are not ready yet.",
            [
                "Input: split manifest.",
                "Output: cleaned split manifest.",
                "Inspect: which rows were kept and which rows were removed.",
            ],
        ),
        (
            "13. Extract audio from video",
            "The audio extraction step creates the WAV branch of the tri-modal sample. This is the point where the same hearing clip becomes useful for audio modeling. In simple terms, the pipeline is turning one courtroom video into a synchronized audio file, a synchronized video clip, and a text row that all point to the same utterance.",
            [
                "Input: cleaned split manifest and video files.",
                "Output: tri-modal manifest with audio paths.",
                "Inspect: whether the extracted WAV is valid and aligned with the clip.",
            ],
        ),
    ]

    for heading, explanation, bullets in sections:
        doc.add_heading(heading, level=2)
        add_para(doc, explanation)
        add_bullets(doc, bullets)

    doc.add_heading("Why The Order Matters", level=1)
    add_numbered(
        doc,
        [
            "If you verify too late, you can accidentally treat page text as proof of a record.",
            "If you build hearing rows too early, you may preserve placeholder-style false positives.",
            "If you build utterance rows before the hearing layer is grounded, the clip builder inherits bad inputs.",
            "If you split too early, leakage checks become meaningless because the groups may not reflect the true hearing or witness structure.",
            "If you extract audio before clip validation, you can waste time producing WAV files from clips that should have been rejected.",
        ],
    )

    doc.add_heading("How To Run It Manually", level=1)
    add_para(
        doc,
        "For a manual rebuild, it is best to think in two stages. First rebuild the corpus grounding layers, then rebuild the utterance-level dataset. If you only changed the alignment logic, you usually only need the last stage. If you changed the case inventory or hearing manifests, you need the earlier stages too. In this workspace, the validated utterance builder must be run with `PYTHON_BIN=/usr/bin/python3` because that interpreter can import `faster_whisper`, while the conda interpreter currently does not. The alignment backend can also be selected through `ALIGNMENT_BACKEND=auto`, `whisperx`, or `heuristic`, but the run command still needs the correct Python binary first.",
    )
    add_para(
        doc,
        "For a new witness-testimony candidate, the manual example I would use is `Nahimana et al. (ICTR-99-52)`. The official ICTR case page shows that this case has a large witness record, and the UCR downloader can enumerate multiple public video recordings for it. A safe student-level way to explain the workflow is: first probe a single recording with `python3 phase2/download_ucr_case_video.py --case-number ICTR-99-52 --date 01/03/2002 --index 1 --verify`; if the probe returns a real MP4, then add the case row to the curated candidate ledger or source manifest, rerun the enrichment and inventory builders, split the verified inventory by media type, rebuild the hearing and witness manifests, and finally rerun the validated LegalMELD builder. If the probe returns no recordings, stop at transcript-only and do not promote the case into the tri-modal branch.",
    )
    add_code_block = """
python3 phase2/download_ucr_case_video.py --case-number ICTR-99-52 --date 01/03/2002 --index 1 --verify
bash phase2/run_deduplicate_case_ledger.sh
bash phase2/run_crawl_official_ucr_case_pages.sh
bash phase2/run_enrich_case_ledger_from_ucr_site.sh
bash phase2/run_build_tribunal_manifest_from_ledger.sh
bash phase2/run_build_witness_manifest_from_ledger.sh
bash phase2/run_build_ucr_case_inventory.sh
bash phase2/run_split_ucr_inventory_by_media_type.sh
bash phase2/run_build_hearing_witness_manifests.sh
bash phase2/run_validate_trimodal_corpus.sh
PYTHON_BIN=/usr/bin/python3 ALIGNMENT_BACKEND=auto bash phase2/run_build_legalmeld_dataset_validated.sh --skip-existing
bash phase2/run_phase2_split_manifest.sh
bash phase2/run_phase2_sanitize_manifest.sh
bash phase2/run_phase2_extract_audio.sh
"""
    for line in add_code_block.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)

    doc.add_heading("What To Inspect After A Rebuild", level=1)
    add_table(
        doc,
        ["File", "What it tells you", "Student interpretation"],
        [
            [str(VERIFIED_INVENTORY.relative_to(ROOT)), "Which cases were actually kept", "Confirms the case-level truth set is grounded."],
            [str(HEARING_MANIFEST.relative_to(ROOT)), "Which hearings were paired with real transcript and video evidence", "Shows the bridge from case-level data to hearing-level data."],
            [str(WITNESS_MANIFEST.relative_to(ROOT)), "Which witnesses or protected codes were resolved", "Shows who the speech is associated with, without deanonymizing protected identities."],
            [str(LEGALMELD_METADATA.relative_to(ROOT)), "All utterance-level metadata and alignment decisions", "Shows the actual training rows used by the model."],
            [str(ALIGNMENT_REVIEW.relative_to(ROOT)), "Rows that need manual review", "Tells you where the alignment pipeline is still weak."],
            [str(SUMMARY_JSON.relative_to(ROOT)), "The roll-up statistics", "Shows the dataset size, quality, and split health in one place."],
            [str(TRAIN_CSV.relative_to(ROOT)), "Train partition", "Rows used for fitting the model."],
            [str(DEV_CSV.relative_to(ROOT)), "Development partition", "Rows used for tuning decisions during development."],
            [str(TEST_CSV.relative_to(ROOT)), "Test partition", "Rows reserved for final evaluation."],
        ],
    )

    doc.add_heading("Student-Level Summary", level=1)
    add_para(
        doc,
        "The full Phase 2 rebuild is a chain of evidence. First you collect and clean candidate cases, then you verify them against the official UCR records, then you convert verified records into hearing and witness manifests, and finally you build the utterance-level LegalMELD dataset that contains text, audio, and video for one aligned speech segment at a time. This is the right order because the later stages depend on the earlier stages being grounded and trustworthy.",
    )
    add_bullets(
        doc,
        [
            "The bootstrap corpus proves the method.",
            "The verified inventory proves the case records.",
            "The hearing and witness manifests prove the pairing logic.",
            "The LegalMELD export proves the utterance alignment logic.",
            "The split files prove the final model-ready organization.",
        ],
    )
    if s:
        add_para(
            doc,
            f"With the current validated outputs, you are rebuilding a corpus that already reports {fmt_int(s.get('hearings_represented', 0))} hearings represented and {fmt_int(s.get('distinct_witnesses', 0))} distinct witnesses in the validated snapshot. That makes the current run useful as a controlled bootstrap, but still not large enough to claim final readiness for all target settings.",
        )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
