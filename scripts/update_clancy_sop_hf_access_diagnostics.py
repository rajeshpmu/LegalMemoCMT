from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOP = ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx"


def code(document: Document, value: str) -> None:
    paragraph = document.add_paragraph(style="No Spacing")
    run = paragraph.add_run(value)
    run.font.name = "Courier New"
    run.font.size = Pt(8)


document = Document(SOP)
document.add_section(WD_SECTION.NEW_PAGE)
document.add_heading("Diagnosing Hugging Face Model Access", level=1)
document.add_paragraph(
    "After the Hub compatibility correction, a result such as `FAIL: Hugging Face returned no diarization "
    "pipeline` should be treated as an access problem until repository access is tested separately. The "
    "preflight checker now calls Hugging Face `model_info` before loading pyannote. This distinguishes a "
    "bad or missing token, a token without access to the gated model, and a genuine local package error."
)
document.add_heading("Run the Diagnostic", level=2)
code(document, '''export HF_TOKEN="<your Hugging Face read token>"
./.venv-diarization/bin/python phase2/check_clancy_diarization_prerequisites.py \\
  --model pyannote/speaker-diarization-3.1 \\
  --load-model''')
document.add_heading("Interpretation", level=2)
for item in [
    "`HF_TOKEN=present` only means a non-empty environment variable exists; it does not prove that the value is a valid token.",
    "`model_repo_access=FAIL` means first verify the token, its read permission, and the Hugging Face model terms.",
    "Open `https://huggingface.co/pyannote/speaker-diarization-3.1` while signed in and accept the model conditions if prompted. Also accept any linked gated dependency conditions requested by the model page.",
    "Create or use a Hugging Face token with read access. Export it in the same terminal session that runs the checker; never put it in a repository file.",
    "`model_repo_access=PASS` followed by `model_access=PASS` means the environment and model access are ready for diarization.",
    "A LibreSSL warning from urllib3 is not the cause of this specific gated-model failure; it is a warning from the Python SSL build unless an actual HTTPS request fails.",
]:
    document.add_paragraph(item, style="List Bullet")
document.add_heading("Safe Token Check", level=2)
code(document, '''./.venv-diarization/bin/python - <<'PY'
import os
from huggingface_hub import HfApi
token = os.environ.get("HF_TOKEN") or os.environ.get("HUGGINGFACE_TOKEN")
if not token:
    raise SystemExit("HF_TOKEN is not set")
user = HfApi().whoami(token=token)
print("Hugging Face identity check passed:", user.get("name") or user.get("user", "unknown"))
PY''')
document.add_paragraph(
    "This prints the account identity, not the secret token. If this identity check fails, fix the token "
    "before retrying pyannote. Once access is fixed, rerun the preflight and only then start the diarization "
    "pilot."
)
document.save(SOP)
print(f"Updated {SOP}")
