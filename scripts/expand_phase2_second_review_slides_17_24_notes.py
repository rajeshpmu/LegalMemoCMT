"""Append detailed student-level speaking notes for Slides 17-24."""
from pathlib import Path
import shutil

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "implementation_docments/LegalMemoCMT_Phase2_Second_Review_Call_Speaking_Doc.docx"
MARKER = "EXPANDED_SLIDES_17_24_TECHNICAL_NOTES_V1"


def main() -> None:
    doc = Document(DOCX)
    if any(MARKER in p.text for p in doc.paragraphs):
        print("Already updated", DOCX)
        return
    backup = DOCX.with_name(DOCX.name + ".before_expanded_slides_17_24.docx")
    if not backup.exists():
        shutil.copy2(DOCX, backup)
    doc.add_heading("Expanded Technical Speaking Notes: Slides 17-24", level=1)
    marker = doc.add_paragraph(); marker.add_run(MARKER).font.size = Pt(8)

    sections = [
        ("Slide 17: Example SILVER Auto-Adjudication", [
            "This slide is a worked example of the acceptance gate, not a claim that the system has discovered the true emotional state. I should begin with provenance: the row originated in the scope-aware Clancy manifest, was enriched by the courtroom-affect candidate generator, and was then passed through apply_clancy_annotation_acceptance_gate.py.",
            "The Phase 1 fields are historical machine evidence. phase1_basic_emotion=sadness and phase1_basic_emotion_confidence=0.544674 are not deleted when the integrated review candidate becomes neutral. Keeping both values lets me measure domain shift and inspect cases where MELD-trained semantics do not transfer directly to courtroom testimony.",
            "The basic-emotion candidate and courtroom-affect candidate answer different questions. basic_emotion_review_candidate=neutral is a proposal about the witness's own categorical emotion. proposed_courtroom_affect=CALM_COMPOSED is a proposal about the witness's observed delivery or interactional presentation. A witness can therefore be sad but calm, or neutral but defensive; the fields must not be treated as synonyms.",
            "The row has emotion_target_scope=QUOTED_SPEECH and speaker_emotion_evidence_present=NO. In practical terms, the witness reports what another person said. The emotional words belong to the quoted content unless the video or voice supplies independent evidence that the current witness is expressing that emotion.",
            "negative_activation_candidate=YES is a low-level audio result. It means the valence/excitement pattern is negative or activated; it is not equivalent to basic_emotion=sadness and is not equivalent to courtroom_affect=DISTRESSED. distress_corroboration_present=NO blocks an automatic distressed interpretation.",
            "The gate accepts this row because the basic candidate confidence is 0.75, the affect candidate confidence is 0.60, both candidates are resolved, and critical_conflict=NO. The final fields are therefore final_basic_emotion=neutral and final_courtroom_affect=CALM_COMPOSED, with annotation_status=AUTO_ADJUDICATED and annotation_tier=SILVER. I should explain SILVER as machine-assisted acceptance for controlled experiments, not human gold truth.",
        ]),
        ("Slide 18: Example Weak / Unresolved Candidate", [
            "This slide demonstrates why a strong score in one dimension cannot compensate for weak evidence in another. DCBWoWhsTpA_turn06570 has a courtroom-affect candidate HESITANT_UNCERTAIN at 0.75, but its basic-emotion candidate is the preserved Phase 1 disgust result at only 0.47.",
            "The transcript contains 'Um' and 'I don't believe', which are hesitation and epistemic-uncertainty evidence. The revised heuristic gives HESITANT_UNCERTAIN precedence over generic TENSE. The bracketed [snorts] marker is retained as context only; a transcription token for a vocalization is not a reliable disgust label.",
            "Because the basic confidence is below 0.70, the row fails the standard acceptance gate. The output sets final_basic_emotion=UNRESOLVED and annotation_status=UNRESOLVED, annotation_tier=WEAK, but retains final_courtroom_affect=HESITANT_UNCERTAIN as a review candidate because its affect confidence meets the affect threshold.",
            "This is not a discarded row. It is a valuable active-review example because it exposes disagreement between semantic/basic emotion evidence and interactional delivery evidence. A human reviewer should inspect the transcript, audio, and video before creating final human labels.",
        ]),
        ("Slide 19: How the Gated CSV Was Produced", [
            "The diagram shows a two-stage architecture. First, propose_clancy_courtroom_affect.py reads transcript text, scope, audio-SER fields, and visual verification fields. It calculates interpretable scores for CALM_COMPOSED, HESITANT_UNCERTAIN, GUARDED, DEFENSIVE, ASSERTIVE, TENSE, DISTRESSED, and AGITATED.",
            "The candidate generator does not replace the original Phase 1 prediction. It appends fields such as proposed_courtroom_affect, proposed_courtroom_affect_confidence, courtroom_affect_evidence, negative_activation_candidate, distress_corroboration_present, speaker_emotion_evidence_present, and basic_emotion_review_candidate.",
            "Second, apply_clancy_annotation_acceptance_gate.py evaluates the candidate record. The standard rule is approximately: accepted = basic_confidence >= 0.70 AND affect_confidence >= 0.60 AND candidates resolved AND critical_conflict=NO. Accepted rows receive AUTO_ADJUDICATED/SILVER; other rows receive UNRESOLVED/WEAK.",
            "The diagram is useful because it prevents a common implementation mistake: merging heuristic generation and final acceptance into one opaque function. Separating them means I can change a scoring rule, rerun the gate, and compare outputs while preserving the original evidence.",
        ]),
        ("Slide 20: Why the SILVER Row Passed", [
            "I should explain this slide as a row-level audit. The input evidence is not just the final label: it includes the source utterance ID, transcript, speaker role, emotion scope, Phase 1 output, audio evidence, visual verification, proposed candidates, confidence values, and the gate reason.",
            "For the worked row, the quotation scope and absence of speaker-emotion evidence reduce the risk of assigning the reported person's sadness to the witness. CALM_COMPOSED is a behavioral description, not a claim that the witness feels no sadness.",
            "The gate's output fields are derived fields. final_basic_emotion comes from the accepted basic review candidate, final_courtroom_affect comes from the accepted affect candidate, and acceptance_gate_reason records why the branch was selected. The source and candidate fields remain available for audit.",
            "The correct student explanation is 'this row meets a transparent machine-acceptance policy', not 'the model proved the witness was neutral'. Calibration against human-reviewed rows is still required before calling the tier reliable.",
        ]),
        ("Slide 21: Why the Row Stayed WEAK", [
            "This slide shows conservative rejection of automatic finalization. The affect proposal is usable as a hypothesis, but the basic-emotion evidence is below threshold. Therefore the gate refuses to manufacture a final basic emotion.",
            "The final_basic_emotion=UNRESOLVED value is important because it tells later training code not to treat disgust as an accepted label. The preserved phase1_basic_emotion=disgust remains available for error analysis, and final_courtroom_affect=HESITANT_UNCERTAIN remains available to guide manual review.",
            "I should distinguish weak from wrong. WEAK means the evidence is insufficient for automatic acceptance; it does not prove that every candidate is incorrect. Human review can later replace unresolved values with explicitly sourced human annotations.",
        ]),
        ("Slide 22: Acceptance Gate - Two Critical-Failure Checks", [
            "The first critical check catches a non-neutral Phase 1 prediction that changes to neutral without safe attribution evidence and CALM_COMPOSED delivery. This is a material disagreement about whose emotion is being represented, so it receives focused review.",
            "The second critical check catches proposed_courtroom_affect=DISTRESSED without distress corroboration. Negative valence plus elevated excitement is only negative activation. Corroboration must come from a categorical audio cue, distress-specific language, or separately recorded visual evidence.",
            "These checks are implemented after candidate generation, so they can inspect the complete evidence record. They are safeguards against semantic leakage and against treating continuous audio dimensions as direct psychological labels.",
        ]),
        ("Slide 23: Critical Conflicts Versus Ordinary Weak Rows", [
            "In the current pilot there are 75 UNRESOLVED/WEAK rows and 7 critical conflicts. The 7 are a subset of the 75, not an additional group. I should say this explicitly so the counts are not incorrectly added to 82.",
            "Critical rows have a specific ambiguity requiring focused comparison of transcript scope, audio delivery, visual behavior, and role context. Ordinary weak rows usually fail because a confidence value is below threshold or a candidate is UNKNOWN; they may not contain a direct contradiction.",
            "The acceptance_gate_reason field records the explanation. This is important for a student implementation because a reviewer can sort the queue by reason instead of treating all unresolved rows as the same type of error.",
        ]),
        ("Slide 24: Decision and Next Action", [
            "The pilot demonstrates that the pipeline can produce a traceable machine-assisted annotation tier, but it does not establish final courtroom affect prevalence or human-label accuracy. The result is a controlled review state.",
            "The immediate next action is to inspect all critical rows and the remaining unresolved rows, then manually validate a stratified sample of SILVER rows. I should compare accepted labels with human decisions and calculate precision or agreement before expanding the automatic gate to the full Clancy corpus.",
            "The DeBERTa scope pilot is a complementary next step. It should produce target and temporal scope suggestions using natural-language hypotheses, but it must be compared against existing scope fields and reviewed before canonical fields are overwritten.",
            "Only after these checks should approved human labels be merged into the annotation manifest and used for supervised or weakly supervised training. No pipeline output should be described as deception, credibility, truthfulness, or reliability evidence.",
        ]),
    ]
    for heading, paragraphs in sections:
        doc.add_heading(heading, level=2)
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
    doc.save(DOCX)
    print("Updated", DOCX)


if __name__ == "__main__":
    main()
