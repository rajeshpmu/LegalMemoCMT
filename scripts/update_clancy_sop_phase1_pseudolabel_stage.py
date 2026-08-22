from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOP = ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx"


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)


def add_code(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="No Spacing")
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def main() -> None:
    document = Document(SOP)
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Phase 1 MELD Pseudo-Labeling Stage", level=1)
    document.add_paragraph(
        "Purpose. This stage transfers the seven-class basic-emotion vocabulary learned in Phase 1 "
        "to the validated Clancy courtroom pools as weak, provenance-preserving labels. It is not a "
        "replacement for human annotation and it does not create courtroom-specific affect, credibility, "
        "truthfulness, or deception labels. The outputs are hypotheses that can support later review, "
        "baseline experiments, and controlled adaptation."
    )

    document.add_heading("Where This Stage Fits", level=2)
    document.add_paragraph(
        "Run pseudo-labeling only after the source, clipping, duration filtering, group split, and split "
        "validation stages have passed. The two current Clancy pools are structurally ready: Pool A "
        "(0.8 to less than 20 seconds) has 9,073 unique rows and Pool B (20 to 30 seconds) has 551 "
        "unique rows; both have zero reported split-leakage groups. This means the rows are suitable "
        "for a controlled inference pilot, not that the model predictions are ground truth."
    )
    add_table(
        document,
        ["Order", "Stage", "Main input", "Main output", "Student interpretation"],
        [
            ["1", "Duration and rejection filtering", "post-rejection turn manifest", "two duration manifests", "Remove known unsuitable material and keep comparable clip lengths."],
            ["2", "Group-aware split", "duration manifest", "train/dev/test CSVs", "Keep complete YouTube source groups in one split to reduce leakage."],
            ["3", "Split validation", "dataset manifest", "validation JSON and issues CSV", "Check unique IDs and ensure no group crosses partitions."],
            ["4", "Phase 1 inference pilot", "validated Pool A or B CSV", "200-row pseudo-label CSV and JSON", "Test whether the checkpoint can produce usable weak labels."],
            ["5", "Manual pilot review", "pseudo-label CSV plus clips", "accept/revise decision", "Inspect confidence, text, audio, and video before scaling."],
            ["6", "Full-pool inference", "validated Pool A and Pool B CSVs", "full provenance CSVs", "Generate weak labels for analysis, not automatic annotation truth."],
        ],
    )

    document.add_heading("Why Use the Phase 1 Checkpoint", level=2)
    document.add_paragraph(
        "Phase 1 was trained for the MELD-compatible seven-class basic-emotion task: neutral, anger, "
        "disgust, fear, joy, sadness, and surprise. The checkpoint is therefore appropriate only for "
        "transferring that basic-emotion task into Clancy as a baseline. It does not understand courtroom "
        "roles, testimony context, witness examination, or the meaning of a legal statement as a trained "
        "courtroom-affect annotator would. A prediction such as sadness is a model output, not evidence "
        "that a witness is distressed, deceptive, truthful, credible, or unreliable."
    )
    document.add_paragraph(
        "The current primary checkpoint is `results/paper_aligned_meld_cv/cmt_min/fold_2/best_model.pt`. "
        "It was selected because fold 2 had the strongest weighted F1 (0.6254) among the reported folds. "
        "Fold 4 had the strongest macro F1 (0.4606), so it remains a reasonable sensitivity comparison, "
        "but the current reproducible primary command uses fold 2."
    )

    document.add_heading("Python Environment Requirement", level=2)
    document.add_paragraph(
        "The repository `.venv` is the normal Phase 2 environment, but on the current machine it does "
        "not contain PyTorch. `/opt/anaconda3/bin/python` contains the required torch installation and "
        "was used to verify the inference script help output. The explicit PYTHON_BIN assignment is "
        "therefore required for this stage. This does not change the source data or download media again."
    )
    add_code(document, 'PYTHON_BIN=/opt/anaconda3/bin/python')
    add_code(document, '"$PYTHON_BIN" -c "import torch; print(torch.__version__)"')
    add_code(document, '"$PYTHON_BIN" phase2/pseudo_label_clancy_with_phase1.py --help')

    document.add_heading("Exact Execution Order", level=2)
    document.add_paragraph(
        "The following commands are intentionally ordered from a small pilot to full processing. "
        "Do not start with the full 9,073-row pool. The pilot creates an auditable checkpoint before "
        "the same operation is expanded. MAX_ROWS=0 means all rows; a positive value limits the pilot."
    )
    document.add_paragraph("Step 1 - Run a 200-row Pool A pilot", style="List Number")
    add_code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python \\
INPUT_CSV=$PWD/data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest.csv \\
OUTPUT_CSV=$PWD/data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_phase1_pseudo_200.csv \\
CHECKPOINT=$PWD/results/paper_aligned_meld_cv/cmt_min/fold_2/best_model.pt \\
SUMMARY_JSON=$PWD/reports/phase2/clancy_duration_0_8_to_20_phase1_pseudo_200.json \\
MAX_ROWS=200 \\
BATCH_SIZE=4 \\
MODALITIES=text,audio,video \\
bash phase2/run_pseudo_label_clancy_with_phase1.sh''')
    document.add_paragraph("Step 2 - Inspect the pilot output", style="List Number")
    add_code(document, 'head -n 2 data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_phase1_pseudo_200.csv')
    add_code(document, 'cat reports/phase2/clancy_duration_0_8_to_20_phase1_pseudo_200.json')
    document.add_paragraph("Step 3 - After the pilot gate passes, process all Pool A rows", style="List Number")
    add_code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python \\
INPUT_CSV=$PWD/data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest.csv \\
OUTPUT_CSV=$PWD/data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_phase1_pseudo.csv \\
CHECKPOINT=$PWD/results/paper_aligned_meld_cv/cmt_min/fold_2/best_model.pt \\
SUMMARY_JSON=$PWD/reports/phase2/clancy_duration_0_8_to_20_phase1_pseudo.json \\
MAX_ROWS=0 \\
BATCH_SIZE=4 \\
MODALITIES=text,audio,video \\
bash phase2/run_pseudo_label_clancy_with_phase1.sh''')
    document.add_paragraph("Step 4 - Process all Pool B rows separately", style="List Number")
    add_code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python \\
INPUT_CSV=$PWD/data/processed/phase2/clancy/duration_20_to_30/clancy_dataset_manifest.csv \\
OUTPUT_CSV=$PWD/data/processed/phase2/clancy/duration_20_to_30/clancy_dataset_manifest_phase1_pseudo.csv \\
CHECKPOINT=$PWD/results/paper_aligned_meld_cv/cmt_min/fold_2/best_model.pt \\
SUMMARY_JSON=$PWD/reports/phase2/clancy_duration_20_to_30_phase1_pseudo.json \\
MAX_ROWS=0 \\
BATCH_SIZE=4 \\
MODALITIES=text,audio,video \\
bash phase2/run_pseudo_label_clancy_with_phase1.sh''')

    document.add_heading("What the Script Does Internally", level=2)
    add_bullets(document, [
        "Reads the selected Clancy manifest without downloading or recreating the source media.",
        "Loads the checkpoint model configuration and verifies that it is a seven-class MELD model.",
        "Builds the same text, audio, and video sample representation expected by the Phase 1 model.",
        "Runs inference in evaluation mode and converts logits to seven-class softmax probabilities.",
        "Selects the highest-probability class as the weak basic-emotion prediction.",
        "Calculates confidence as the highest class probability and entropy as a measure of distributional uncertainty.",
        "Writes the original row plus checkpoint, modality, model configuration, class, confidence, entropy, and probability provenance.",
        "Writes a JSON summary containing row counts, prediction counts, confidence statistics, and limitations.",
    ])
    document.add_paragraph(
        "The script preserves the original `emotion_label`, `emotion_label_source`, and "
        "`emotion_label_confidence` fields. It adds fields with the `phase1_` prefix instead of silently "
        "overwriting existing values. This makes it possible to compare an older label with the new "
        "prediction and to reproduce which checkpoint produced it."
    )

    document.add_heading("Output Fields and Meaning", level=2)
    add_table(
        document,
        ["Field", "Meaning", "How to use it"],
        [
            ["phase1_basic_emotion", "One of the seven MELD classes selected by maximum probability.", "Use as a weak-label candidate or baseline target after review."],
            ["phase1_basic_emotion_confidence", "Probability assigned to the selected class, from 0 to 1.", "Prioritize high-confidence rows for initial inspection; do not call it accuracy."],
            ["phase1_basic_emotion_entropy", "Uncertainty of the complete seven-class probability distribution.", "Higher values indicate a less decisive prediction and more review need."],
            ["phase1_basic_emotion_probabilities", "All seven class probabilities stored as JSON text.", "Inspect close competing classes and preserve the full prediction evidence."],
            ["phase1_basic_emotion_checkpoint", "Exact checkpoint path used for inference.", "Provides experiment provenance and supports reruns."],
            ["phase1_basic_emotion_source", "Fixed source marker for Phase 1 pseudo-labeling.", "Distinguishes weak labels from human labels."],
            ["phase1_basic_emotion_modalities", "Modalities supplied to the model, normally text,audio,video.", "Confirms whether the result was trimodal or a missing-modality run."],
            ["phase1_basic_emotion_model_config", "Serialized model configuration read from the checkpoint.", "Records architecture settings needed to interpret the output."],
        ],
    )

    document.add_heading("Pilot Review Gate", level=2)
    document.add_paragraph("Before processing all rows, manually review the 200-row pilot using the CSV, summary JSON, and a sample of linked clips.")
    add_table(
        document,
        ["Check", "Pass condition", "If it fails"],
        [
            ["Row preservation", "200 output rows, one unique utterance/turn ID per input row.", "Stop and inspect manifest loading or ID mapping."],
            ["Provenance", "Checkpoint, source, modalities, probabilities, confidence, and entropy are populated.", "Do not scale until provenance is complete."],
            ["Media linkage", "Text, audio, and video paths point to the same utterance/turn record.", "Return to clipping or manifest validation."],
            ["Prediction behavior", "Labels are plausible enough to support review; no obvious systematic failure dominates.", "Compare modalities, inspect low-confidence rows, and reconsider checkpoint or preprocessing."],
            ["Label discipline", "No courtroom-affect, credibility, truthfulness, or deception field is created.", "Treat as a pipeline error and do not use the output."],
        ],
    )
    document.add_paragraph(
        "A pilot pass does not prove that the labels are correct. It only shows that the inference path is "
        "technically functioning and that the predictions are sufficiently interpretable to justify a full "
        "weak-label run. Low-confidence or ambiguous rows should remain available for manual annotation or "
        "later quality analysis rather than being silently discarded."
    )

    document.add_heading("Reproducibility and Guardrails", level=2)
    add_bullets(document, [
        "Keep the input duration manifests unchanged and write pseudo-labels to new output filenames.",
        "Record the checkpoint path, model configuration, modality set, batch size, and summary JSON for every run.",
        "Use the same group-aware train/dev/test partitions; do not randomly reshuffle rows after labeling.",
        "Do not use pseudo-labels as human gold labels in evaluation. Evaluation requires verified or manually reviewed labels.",
        "Do not map basic emotions into courtroom affect automatically. NEUTRAL_CALM, HESITANT_UNCERTAIN, GUARDED, DEFENSIVE, ASSERTIVE, TENSE, DISTRESSED, and AGITATED require a separate annotation policy.",
        "Do not infer deceptive, truthful, lying, credible, or unreliable behavior from emotion, face, audio, or model confidence.",
        "The script does not download media, create new clips, or alter the original Clancy source manifests.",
    ])

    document.add_heading("Current Implementation Status", level=2)
    document.add_paragraph(
        "Completed: duration pools were created, both split validations returned PASS, the inference script "
        "and shell wrapper were created, syntax/help checks passed, and this SOP now documents the controlled "
        "execution order. Pending: run the 200-row inference pilot, inspect its predictions and provenance, "
        "then decide whether to process Pool A and Pool B in full. No Phase 1 inference was executed as part "
        "of this SOP update."
    )

    document.save(SOP)
    print(f"Updated {SOP}")


if __name__ == "__main__":
    main()
