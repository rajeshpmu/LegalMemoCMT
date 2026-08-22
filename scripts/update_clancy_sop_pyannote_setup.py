from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOP = ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx"


def code(document: Document, value: str) -> None:
    p = document.add_paragraph(style="No Spacing")
    r = p.add_run(value)
    r.font.name = "Courier New"
    r.font.size = Pt(8)


document = Document(SOP)
document.add_section(WD_SECTION.NEW_PAGE)
document.add_heading("Hugging Face and pyannote Diarization Setup", level=1)
document.add_paragraph(
    "This section explains how to configure the optional speaker-diarization dependency on the local "
    "machine. Diarization is not a legal-role classifier. The Hugging Face model separates voices into "
    "anonymous clusters; the later manual review maps those clusters to Witness, Prosecutor, Defence, "
    "Judge, Other, or UNKNOWN."
)
document.add_heading("Environment and Token Safety", level=2)
document.add_paragraph(
    "Use the same Python environment that will run the diarization script. The current repository uses "
    "`/opt/anaconda3/bin/python` for PyTorch and Hugging Face inference. The Hugging Face access token "
    "must be supplied through the shell environment as `HF_TOKEN` or `HUGGINGFACE_TOKEN`; never write it "
    "into a CSV, shell script, DOCX, or committed file. Accept the terms for the selected pyannote model "
    "on Hugging Face before attempting model loading."
)
code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python
export HF_TOKEN="<your Hugging Face token>"
bash phase2/setup_clancy_diarization.sh''')
document.add_paragraph(
    "The setup wrapper installs `pyannote.audio>=3.3,<4`. If no token is present, it stops after package "
    "installation and tells the user to set the token. When a token is present, it also loads the default "
    "`pyannote/speaker-diarization-3.1` model to test access. Model files are cached under the repository "
    "`.cache/huggingface` directory by the diarization wrapper unless `HF_HOME` is overridden."
)
document.add_heading("Preflight Checks", level=2)
code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python
export HF_TOKEN="<your Hugging Face token>"
"$PYTHON_BIN" phase2/check_clancy_diarization_prerequisites.py \\
  --model pyannote/speaker-diarization-3.1 \\
  --load-model''')
document.add_paragraph(
    "A successful preflight prints the active Python path, PyTorch version, pyannote.audio version, "
    "`HF_TOKEN=present`, model access PASS, and status PASS. A missing token, unavailable package, or "
    "permission error must be fixed before running the diarizer."
)
document.add_heading("Pilot Diarization", level=2)
document.add_paragraph(
    "Run one source recording first. The script uses `source_audio_path` when available, so all short "
    "turns from one recording share the same diarization cluster namespace. `--max-sources 1` limits the "
    "pilot and avoids processing the full corpus before the output format is understood."
)
code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python \\
HF_TOKEN="$HF_TOKEN" \\
bash phase2/run_clancy_diarization.sh \\
  --input-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_subtitle_evidence.csv \\
  --output-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_diarized.csv \\
  --segments-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_diarization_segments_200.csv \\
  --device cpu \\
  --max-sources 1''')
document.add_paragraph(
    "For a CUDA or Apple Silicon experiment, use `--device cuda` or `--device mps` only after verifying "
    "the corresponding torch backend. CPU is the reproducible default. Optional `--min-speakers` and "
    "`--max-speakers` can constrain the diarizer when a hearing has a known approximate participant count, "
    "but they should not be used to force a legal-role count."
)
document.add_heading("Full Source-Level Diarization", level=2)
code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python \\
HF_TOKEN="$HF_TOKEN" \\
bash phase2/run_clancy_diarization.sh \\
  --input-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_subtitle_evidence.csv \\
  --output-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_diarized.csv \\
  --segments-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_diarization_segments_200.csv \\
  --device cpu''')
document.add_paragraph(
    "The diarizer writes one segment file containing source path, anonymous cluster ID, start time, end "
    "time, and model name. It also writes an enriched manifest by selecting the diarization segment with "
    "the greatest overlap with each turn. It does not assign `speaker_role`, and it does not decide whether "
    "the visible person is the speaker."
)
document.add_heading("Manual Role Review Remains Required", level=2)
document.add_paragraph(
    "After diarization, run the cluster review-sheet script, listen to the example audio clips, inspect the "
    "matching video, and complete the role map. Only the manually completed map can set "
    "`speaker_role_source=manual_verified_diarization`. The final application script validates role values "
    "and preserves unresolved clusters as UNKNOWN."
)
document.add_heading("Troubleshooting", level=2)
for item in [
    "HF_TOKEN is missing: export it in the current shell and rerun preflight.",
    "Model access denied: accept the pyannote model terms and confirm the token has access.",
    "pyannote.audio import fails: rerun setup with the same PYTHON_BIN used for diarization.",
    "Source audio is missing: fix the manifest or exclude the row; do not diarize isolated turn clips as a substitute.",
    "Too many speakers: inspect the source recording and optionally test min/max speaker bounds; do not interpret cluster count as role count.",
    "Native runtime instability: use CPU before trying MPS or CUDA and keep the model cache unchanged for reproducibility.",
]:
    document.add_paragraph(item, style="List Bullet")
document.add_heading("Implementation Status", level=2)
document.add_paragraph(
    "Completed: setup wrapper, preflight checker, device controls, Hugging Face model selection, cache "
    "configuration, and SOP instructions. The deterministic subtitle stage has been run on 200 rows. "
    "Diarization itself remains pending until pyannote.audio is installed and a valid Hugging Face token is "
    "available."
)
document.save(SOP)
print(f"Updated {SOP}")
