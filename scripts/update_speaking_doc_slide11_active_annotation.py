from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "implementation_docments/LegalMemoCMT_Phase2_First_Guidance_Call_Speaking_Doc.docx"
BACKUP = ROOT / "implementation_docments/LegalMemoCMT_Phase2_First_Guidance_Call_Speaking_Doc.before_slide11_active_annotation.docx"
MARKER = "Slide 11. Planned Next Actions: Active Annotation Preparation"


def code(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="No Spacing")
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)


document = Document(DOCX)
if any(MARKER in paragraph.text for paragraph in document.paragraphs):
    print("Slide 11 speaking section already exists")
    raise SystemExit(0)
if not BACKUP.exists():
    shutil.copy2(DOCX, BACKUP)

document.add_section(WD_SECTION.NEW_PAGE)
document.add_heading(MARKER, level=1)
document.add_paragraph(
    "This slide presents the immediate implementation sequence after the visible-witness and Pool A design. "
    "The important point is that active annotation is not a shortcut around corpus construction. It is a later "
    "stage that depends on reliable speaker evidence, role mapping, transcript alignment, and valid tri-modal clips. "
    "I will explain the six boxes from left to right and then explain the decision gate below them."
)

document.add_heading("1. Complete the One-Source Pyannote Pilot", level=2)
document.add_paragraph(
    "The first box means I run Pyannote on one source audio file rather than the entire Clancy corpus. This limits "
    "the experiment so that I can inspect the output manually and identify problems before scaling. The diarization "
    "pipeline detects speech intervals and assigns anonymous cluster IDs such as SPEAKER_00. It does not identify "
    "Witness, Defence, Prosecutor, or Judge. The output is therefore an intermediate audio-evidence layer."
)
code(document, '''PYTHON_BIN=$PWD/.venv-diarization/bin/python \\
HF_TOKEN="$HF_TOKEN" \\
bash phase2/run_clancy_diarization.sh \\
  --input-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_subtitle_evidence.csv \\
  --output-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_diarized.csv \\
  --segments-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_diarization_segments_200.csv \\
  --device cpu \\
  --max-sources 1''')
document.add_paragraph(
    "The two expected artifacts have different purposes. The segments CSV is the raw audit record containing source "
    "audio path, cluster ID, segment start, and segment end. The diarized manifest retains the utterance rows and "
    "adds the strongest overlapping cluster and speaker_cluster_overlap_seconds. I must not treat an output file as "
    "complete until the process exits and the files can be opened and counted."
)

document.add_heading("2. Save Raw Segments and the Diarized Manifest", level=2)
document.add_paragraph(
    "The second box is a provenance requirement. I retain the raw Pyannote intervals instead of saving only the final "
    "role-aware row. This allows a reviewer to ask why an utterance received a particular cluster. The utterance-level "
    "assignment is a temporal join: if a transcript/turn interval overlaps one or more Pyannote intervals, the current "
    "implementation assigns the candidate with the greatest overlap. This is useful evidence, but it is not proof of "
    "identity. Short utterances and overlapping speech can produce weak assignments."
)

document.add_heading("3. Map Anonymous Clusters to Courtroom Roles", level=2)
document.add_paragraph(
    "The third box introduces human verification. I create a cluster review sheet with representative audio/video "
    "clips and the total number of rows associated with each cluster. I listen to the samples and map clusters to "
    "Witness, Prosecutor, Defence, Judge, or Other. The role decision is stored separately from the machine-generated "
    "cluster ID. This distinction is important because Pyannote knows acoustic similarity, not legal semantics."
)
code(document, '''./.venv/bin/python phase2/create_clancy_speaker_role_review.py \\
  --input-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_diarized.csv \\
  --output-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_speaker_role_review.csv \\
  --examples-per-cluster 1''')
document.add_paragraph(
    "For each reviewed cluster I record role_label, role_confidence, witness_in_segment, witness_speaking_status, "
    "visual_target_role, visual_speaker_match, and review notes. A high role confidence means that the sample was "
    "clear enough for the reviewer to make a role decision; it does not mean that the diarization model identified a "
    "person with high certainty."
)

document.add_heading("4. Align Clusters to Turns and Validate Witness Clips", level=2)
document.add_paragraph(
    "The fourth box combines the speaker evidence with the transcript and video checks. A witness-speaking row should "
    "satisfy speaker_role=Witness and witness_speaking_status=SPEAKING. Visual review must separately confirm that the "
    "witness is visible if the row is intended for facial-emotion modeling. If the camera shows the witness while a "
    "lawyer is speaking, the correct interpretation is witness_in_segment=YES and witness_speaking_status=LISTENING_OR_ADDRESSED, "
    "not witness speech."
)
document.add_paragraph(
    "I also verify that the transcript text, WAV, and MP4 refer to the same time interval; that audio is present and "
    "not silent; that the video is decodable; and that the alignment confidence is acceptable. This prevents a speaker "
    "label from making an otherwise incorrect clip appear valid."
)

document.add_heading("5. Extend the Schema and Select a Diverse Seed", level=2)
document.add_paragraph(
    "After the role-aware manifest is reliable, I run the additive annotation schema extension. It adds canonical basic "
    "emotion and courtroom-affect fields, controlled annotation statuses, model suggestion fields, observable cue "
    "fields, active-learning priority fields, and iteration provenance. Existing emotion_label fields are preserved for "
    "compatibility. Unsupported placeholders such as emotion_label=unknown remain untrusted and are not promoted to a "
    "canonical basic emotion."
)
code(document, '''PYTHON_BIN=$PWD/.venv/bin/python \\
bash phase2/annotation/run_extend_annotation_schema.sh \\
  --input-csv <role-aware-manifest.csv> \\
  --output-csv <schema-extended-manifest.csv>

PYTHON_BIN=$PWD/.venv/bin/python \\
bash phase2/annotation/run_build_seed_annotation_set.sh \\
  --input-csv <schema-extended-manifest.csv> \\
  --output-csv <seed.csv> \\
  --summary-json <seed-summary.json> \\
  --config configs/annotation/active_learning.yaml''')
document.add_paragraph(
    "The seed selector is deliberately diversity-oriented. It can spread rows across source videos, witnesses, and "
    "examination phases. The target is initially approximately 500–1000 rows, including neutral, high-affect, and "
    "ambiguous examples. It should not select only easy high-confidence rows because that would make the seed look "
    "better than the real corpus."
)

document.add_heading("6. Export the Human Review Queue", level=2)
document.add_paragraph(
    "The sixth box creates a reviewer-facing CSV. It contains the utterance ID, source and witness metadata, original "
    "clip paths, transcript context, previous question where available, examination phase, role evidence, model "
    "suggestions, priority score, and priority reason. The annotator must be able to inspect the original audio/video "
    "clip, not only a text prediction."
)
code(document, '''PYTHON_BIN=$PWD/.venv/bin/python \\
bash phase2/annotation/run_compute_annotation_priority.sh \\
  --input-csv <seed-or-suggestion-manifest.csv> \\
  --output-csv <priority-manifest.csv> \\
  --summary-json <priority-summary.json> \\
  --config configs/annotation/active_learning.yaml

PYTHON_BIN=$PWD/.venv/bin/python \\
bash phase2/annotation/run_build_annotation_queue.sh \\
  --input-csv <priority-manifest.csv> \\
  --output-csv <annotation_queue.csv> \\
  --max-rows 100''')
document.add_paragraph(
    "The priority score is not a label-quality score. It is an operational selection score. It can combine predictive "
    "entropy, top-two margin, disagreement between modalities, class rarity, and quality flags. The weights are in "
    "configs/annotation/active_learning.yaml so I can inspect or change them without hiding decisions inside source code."
)

document.add_heading("Decision Gate: Why Pseudo-Labeling Does Not Start Immediately", level=2)
document.add_paragraph(
    "The slide's decision gate is the most important explanation. I do not begin pseudo-labeling simply because a "
    "pretrained model is available. First, the audio must be diarized, the clusters must be mapped to courtroom roles, "
    "the witness turns must be aligned to transcript and clips, and the text/audio/video rows must pass validation. "
    "Otherwise a model may learn from lawyer speech, incorrect timestamps, duplicated context, or a video that does not "
    "show the intended witness."
)
document.add_paragraph(
    "Once the evidence gate passes, I still create a human-reviewed seed before accepting pseudo-labels. Model outputs "
    "are stored as suggestions or pseudo-labels with source, checkpoint, confidence, and iteration fields. They never "
    "replace human labels. The held-out benchmark remains witness-disjoint and is not used for active selection, "
    "threshold tuning, or pseudo-label training."
)

document.add_heading("Student-Level Summary for the Guidance Call", level=2)
document.add_paragraph(
    "My next action is not to train immediately. I will complete one controlled Pyannote run, preserve its raw and "
    "utterance-level outputs, map voice clusters to courtroom roles, verify witness-speaking tri-modal clips, extend "
    "the annotation schema, select a diverse seed, and export a review queue. Only after those artifacts are reliable "
    "will I use pretrained models to suggest labels and plan warm-start active-learning iterations."
)

document.save(DOCX)
print(f"Updated {DOCX}")
