from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_RunPod_SOP.docx"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def build_doc() -> Document:
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("LegalMemoCMT RunPod SOP")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(
        "Student-level step-by-step guide for setting up a RunPod pod, pulling the LegalMemoCMT repo, downloading MELD, and running the paper-aligned audio + text 5-fold workflow."
    )
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12.5)

    add_para(
        doc,
        "This SOP is written for a student who has a clean RunPod pod with the LegalMemoCMT code already pulled by Git, but who still needs to install the environment, download the MELD raw dataset, build the manifests and folds, and then run the paper-aligned audio + text benchmark path. The goal is to make the workflow executable one step at a time, without assuming that the raw data already exists on the pod.",
    )
    add_para(
        doc,
        "The document is intentionally operational. It explains what each step does, why it matters, what the expected output should look like, and what you should check before moving to the next step. That makes it suitable for first-time execution on a cloud pod.",
    )

    doc.add_heading("1. What This SOP Is For", level=1)
    add_bullets(
        doc,
        [
            "Prepare the RunPod environment for LegalMemoCMT.",
            "Download MELD raw data and annotations onto the pod.",
            "Build the MELD manifest and 5-fold cross-validation splits.",
            "Run the paper-aligned audio + text model path.",
            "Evaluate the saved checkpoints and inspect the outputs.",
            "Keep the execution order simple enough for student use.",
        ],
    )

    doc.add_heading("2. What You Should Already Have on the Pod", level=1)
    add_bullets(
        doc,
        [
            "The LegalMemoCMT repository cloned from Git.",
            "A working Python environment or conda environment.",
            "Sufficient disk space for raw MELD data, processed features, and results.",
            "Internet access from the pod for dataset and package downloads.",
            "A shell where you can run the repository scripts directly.",
        ],
    )

    doc.add_heading("3. Recommended Directory Layout", level=1)
    add_table(
        doc,
        ["Folder", "Purpose"],
        [
            ["data/MELD/", "Raw MELD archive, annotations, and extracted files"],
            ["data/processed/MELD/", "Extracted audio/video features and processed artifacts"],
            ["data/manifests/", "CSV manifests and fold splits"],
            ["results/paper_aligned_meld_cv/", "MELD 5-fold cross-validation checkpoints and metrics"],
            ["results/paper_aligned_case_study/", "Optional case-study runs and pooling comparisons"],
        ],
    )
    add_para(
        doc,
        "Keeping the directories separate matters because raw data, processed features, and training outputs are different stages of the workflow. If these are mixed together, debugging becomes harder and accidental overwrites become more likely.",
    )

    doc.add_heading("4. Environment Setup", level=1)
    add_numbered(
        doc,
        [
            "Clone or verify the LegalMemoCMT repository on the pod.",
            "Create or activate the Python environment used for the project.",
            "Install the required packages from the project requirements or environment file.",
            "Confirm that Python can import the project modules and that the smoke test passes.",
        ],
    )
    add_para(
        doc,
        "At this stage the important idea is not the exact shell syntax but the validation chain. You want to confirm that the pod can run the repository code before you spend time downloading MELD. A failed environment is much easier to fix before the data download begins.",
    )

    doc.add_heading("5. Step 1: Verify the Repository and Environment", level=1)
    add_bullets(
        doc,
        [
            "Make sure you are inside the LegalMemoCMT repository root on the pod.",
            "Run the repository smoke test after activation of the environment.",
            "Confirm that the model imports, the data loader imports, and the training code all resolve correctly.",
        ],
    )
    add_para(
        doc,
        "If the smoke test fails, do not start the MELD download yet. Fix the environment first, because a broken package setup will make every later step more confusing.",
    )

    doc.add_heading("6. Step 2: Download MELD", level=1)
    add_para(
        doc,
        "MELD is the primary benchmark for the paper-aligned case study. On the pod, you need the raw archive and the annotation CSVs before any manifest can be built. The repository already includes a MELD download helper script, so the correct workflow is to use the script rather than manually assembling the data layout.",
    )
    add_bullets(
        doc,
        [
            "Download the raw MELD archive into data/MELD/.",
            "Download the MELD annotation CSV files into the annotation directory.",
            "Extract the archive if requested by the helper script.",
            "Check that the expected train/dev/test media layout exists after extraction.",
        ],
    )
    add_para(
        doc,
        "This is the first truly important data step. The model cannot train on MELD until the raw clips and annotation CSVs are present. If the raw data is incomplete, the manifest builder will either skip samples or produce an empty or partially empty manifest.",
    )

    doc.add_heading("7. Step 3: Build the MELD Manifest", level=1)
    add_bullets(
        doc,
        [
            "Convert the raw MELD annotations into a row-based manifest.",
            "Extract or derive audio and video feature files where needed.",
            "Write split-aware CSVs into data/manifests/.",
            "Validate that the manifest rows have usable transcript, audio, and video paths.",
        ],
    )
    add_para(
        doc,
        "The manifest is the bridge between raw benchmark data and model training. It tells the training loop which sample belongs to which split, what the label is, and where to find the modality files. Without a valid manifest, the training code has nothing reliable to read.",
    )

    doc.add_heading("8. Step 4: Build 5-Fold CV Splits", level=1)
    add_para(
        doc,
        "The paper-aligned MELD path uses five folds so that the model is evaluated more than once under slightly different training partitions. That matters because one split can overstate or understate performance by accident. Five folds give you a more stable picture of how the model behaves.",
    )
    add_bullets(
        doc,
        [
            "Create fold CSVs for the train and validation partitions.",
            "Keep the test split held out for final evaluation.",
            "Preserve the same split discipline across every fold.",
            "Use the fold outputs later for confusion analysis and report generation.",
        ],
    )

    doc.add_heading("9. Step 5: Run the Paper-Aligned Audio + Text Model", level=1)
    add_para(
        doc,
        "For your current goal, the important model path is the paper-aligned audio + text base model. This means using the pretrained text backbone and the pretrained speech backbone, then combining them through the cross-modal transformer fusion block. In this phase, video is not the core requirement for the benchmark run.",
    )
    add_bullets(
        doc,
        [
            "Use the pretrained paper mode.",
            "Keep the modalities to text and audio for the base benchmark run.",
            "Use the selected pooling setting from the paper-aligned setup.",
            "Train fold by fold so each fold produces its own checkpoint and metrics.",
        ],
    )
    add_para(
        doc,
        "This run is the one that matters if your immediate goal is a clean five-fold MELD result using the audio + text checkpoint path. The student should think of it as the main benchmark line, not as a side experiment.",
    )

    doc.add_heading("10. Step 6: Evaluate the Saved Checkpoints", level=1)
    add_bullets(
        doc,
        [
            "Evaluate each saved best_model.pt on the held-out MELD test split.",
            "Save the metrics JSON file for each fold.",
            "Export per-sample predictions so the predictions can be inspected later.",
            "Generate confusion matrices and fold summaries from the saved outputs.",
        ],
    )
    add_para(
        doc,
        "Evaluation is a separate stage from training because the saved checkpoint is the thing you want to inspect, compare, and report. The student should not confuse the training loss printed during learning with the final benchmark score printed after evaluation.",
    )

    doc.add_heading("11. Check Whether Raw Data Is Still Needed", level=1)
    add_para(
        doc,
        "Before downloading the full MELD raw archive, verify whether the manifest already points to files that exist on the pod. If the audio, video, and transcript references all exist, then the raw archive is not needed for that run. If any referenced file is missing, the raw archive or the missing processed features are still required.",
    )
    add_bullets(
        doc,
        [
            "Use the manifest-assets checker to inspect file existence without starting training.",
            "Treat the manifest as sufficient only when all referenced files are present.",
            "If the manifest points to raw clips or raw features that are absent, download the missing data first.",
            "If transcripts are stored inline in the CSV, the transcript text may be enough even when the raw archive is absent.",
        ],
    )

    doc.add_heading("12. What You Should Expect at Each Stage", level=1)
    add_table(
        doc,
        ["Stage", "Expected output"],
        [
            ["Environment check", "Smoke test passes and the imports resolve"],
            ["MELD download", "Raw archive and annotations appear under data/MELD/"],
            ["Manifest build", "CSV manifests are written to data/manifests/"],
            ["5-fold split build", "Fold-specific CSVs are created"],
            ["Training", "A best_model.pt checkpoint and epoch logs are written"],
            ["Evaluation", "Metrics JSON and predictions CSV are written"],
            ["Analysis", "Confusion matrices and summary files are generated"],
        ],
    )

    doc.add_heading("13. Common Problems and How to Think About Them", level=1)
    add_bullets(
        doc,
        [
            "If the dataset download fails, check the network and the target URL before retrying.",
            "If the manifest is empty, the raw files are likely not in the expected MELD layout.",
            "If training crashes early, confirm that the dependencies and the manifest columns are correct.",
            "If evaluation fails to load the checkpoint, make sure the checkpoint path matches the fold output directory.",
            "If the scores look weak, inspect the confusion matrix before assuming the model is broken.",
        ],
    )
    add_para(
        doc,
        "A student should treat these as workflow checks, not as signs of failure. A cloud pod workflow often needs a few corrections before the first complete run becomes stable.",
    )

    doc.add_heading("14. What Not To Do", level=1)
    add_bullets(
        doc,
        [
            "Do not start training before confirming the MELD manifest exists.",
            "Do not overwrite raw data with processed outputs.",
            "Do not mix one-off experiments with the benchmark outputs.",
            "Do not skip evaluation and go straight from training logs to reporting.",
            "Do not treat the pod as disposable; keep the file structure organized from the start.",
        ],
    )

    doc.add_heading("15. Practical Student Sequence", level=1)
    add_numbered(
        doc,
        [
            "Clone the repo and verify the environment on RunPod.",
            "Run the smoke test.",
            "Download MELD raw data and annotations.",
            "Build and validate the MELD manifest.",
            "Generate the 5-fold cross-validation splits.",
            "Train the paper-aligned audio + text model fold by fold.",
            "Evaluate the best checkpoint for each fold.",
            "Export predictions and inspect the confusion matrix.",
            "Keep the outputs for report writing and mentor discussion.",
        ],
    )

    doc.add_heading("16. Final Advice", level=1)
    add_para(
        doc,
        "If your immediate goal is the MELD 5-fold CV audio + text checkpoint path, treat the SOP as a staged checkpoint list. Do not move to the next step until the current step has produced the expected artifact. That is the safest way to use a RunPod pod for a research project, because it makes each stage auditable and repeatable.",
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
