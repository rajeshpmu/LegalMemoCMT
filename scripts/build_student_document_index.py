from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "implementation_docments"
OUTPUT = DOC_DIR / "LegalMemoCMT_Student_Document_Index.docx"


CORE_DOCS: list[tuple[str, str, str]] = [
    (
        "LegalMemoCMT_Fusion_Teaching_Guide.docx",
        "Primary teaching guide for the current implementation. Explains how speech, text, and video are encoded, fused, pooled, and evaluated.",
        "Start here if you want to understand the whole pipeline before looking at the review documents.",
    ),
    (
        "LegalMemoCMT_MELD_CV_5Fold_Analysis_Report.docx",
        "Fold-by-fold MELD cross-validation report with per-fold metrics, confusion matrices, and fold 2 versus fold 4 interpretation.",
        "Use this when you want the detailed results-level report for the MELD CV experiment.",
    ),
    (
        "LegalMemoCMT_CREMA_D_Analysis_Report.docx",
        "Held-out CREMA-D analysis report with confusion matrix and technical interpretation of the current available output.",
        "Use this if you want the existing CREMA-D analysis summarized in a clean report form.",
    ),
    (
        "LegalMemoCMT_MELD_vs_CREMA_D_Training_Strategy_Deep_Dive.docx",
        "Dedicated explanation of why MELD remains the main paper-aligned benchmark, how imbalance should be handled, and where CREMA-D fits as a secondary benchmark.",
        "Read this when you want a detailed answer to the MELD-versus-CREMA-D training question.",
    ),
    (
        "LegalMemoCMT_Paper_Aligned_Technical_Details_20Page.docx",
        "Long-form technical explanation that cross-checks the MemoCMT base paper against the current Phase 1 implementation.",
        "Use this for in-depth theoretical understanding of the paper-aligned architecture, encoders, fusion, and evaluation logic.",
    ),
    (
        "LegalMemoCMT_Paper_Aligned_Technical_Details_Expanded.docx",
        "Expanded technical version that explains the key concepts in a longer teaching format.",
        "Useful if you want a step-by-step discussion before reading the 20-page comparison document.",
    ),
    (
        "LegalMemoCMT_Paper_Aligned_Technical_Details.docx",
        "Shorter technical summary of the paper-aligned implementation.",
        "Best as a quick overview when you do not need the full long-form explanation.",
    ),
]


OPERATIONAL_DOCS: list[tuple[str, str, str]] = [
    (
        "LegalMemoCMT_Benchmark_Execution_and_Analysis_Guide.docx",
        "Explains how the MELD 5-fold CV and CREMA-D analysis were executed, evaluated, and turned into the report documents.",
        "Use this if you want the most direct script-to-report explanation for the completed benchmark analyses.",
    ),
    (
        "LegalMemoCMT_Phase1_Completion_And_Gap_Report.docx",
        "Explains how much of Phase 1 is implemented, what the remaining gaps are, and what would be needed to reach a full paper-adaptation completion.",
        "Use this when you want a detailed student-level assessment of completeness and the remaining steps to 100%.",
    ),
    (
        "LegalMemoCMT_Guidance_Call_Comparison_Report.docx",
        "Guidance-call comparison document that summarizes the current MELD and CREMA-D results against the MemoCMT base paper and lists the next-step mentor questions.",
        "Use this before a mentor call when you need a concise but deep comparison of what has been done, how it matches the paper, and what should improve next.",
    ),
    (
        "LegalMemoCMT_Improvement_Step1_Class_Balanced_Focal_Loss_Guide.docx",
        "Student-level guide for the first improvement step: class-balanced focal loss, the focused MELD Fold 2/Fold 4 reruns, and the CREMA-D CV improvement workflow.",
        "Use this when you want a deep explanation of the new training, evaluation, and analysis scripts that follow the guidance-call recommendations.",
    ),
    (
        "LegalMemoCMT_Focal_Loss_Dataflow_Student_Guide.docx",
        "End-to-end student guide for the focal-loss path, including the raw data, manifests, fold CSVs, weighted-CE baseline, focal-loss reruns, warm-start refinement, and resume-capable Fold 2 continuation.",
        "Use this when you want a single document that explains how the focal-loss improvement pipeline works from raw dataset to confusion-matrix analysis.",
    ),
    (
        "LegalMemoCMT_MELD_Facial_Cues_ViT_Implementation_Plan.docx",
        "Student-level implementation plan for improving the weighted-CE MELD baseline with ViT-based facial cues, checkpoint warm-starting, and the new face-cue scripts.",
        "Use this when you want to understand the next multimodal expansion step: baseline reuse, face-cue preparation, and the visual branch upgrade.",
    ),
    (
        "LegalMemoCMT_MELD_Face_Crop_ViT_Implementation_Plan.docx",
        "Student-level implementation plan for the next visual ablation: cropping the face before ViT, then reusing the same warm-start MELD training and analysis flow.",
        "Use this when you want to understand why face crops may be better than full frames for courtroom-style facial-cue modeling and how the new scripts should run.",
    ),
    (
        "Guidance_Call_Advanced_AI_ML_Recommendations.docx",
        "Guidance-call technical report that explains the advanced next-step AI/ML recommendations, supported by MELD Fold 2, MELD Fold 4, and CREMA-D confusion matrices and metrics.",
        "Use this when you want a student-level deep dive into the recommended loss, contrastive, context, calibration, and stabilization upgrades for the next mentor discussion.",
    ),
    (
        "LegalMemoCMT_MELD_CV_Paper_Aligned_Report_From_Submission_First.docx",
        "Detailed student report explaining the MELD paper-aligned 5-fold CV run, the results of each action in the bash workflow, and the Fold 2 error analysis.",
        "Use this when you want a full execution narrative and model-behavior analysis for the MELD paper-aligned run.",
    ),
    (
        "LegalMemoCMT_Paper_Aligned_Runner_Guide.docx",
        "Explains the two main paper-aligned scripts: the MELD-only case study and the full MELD + CREMA-D suite.",
        "Use this when you want to understand which runner script to execute and what each one produces.",
    ),
    (
        "LegalMemoCMT_MemoCMT_Style_Replication_Plan.docx",
        "Describes the MemoCMT-style replication and case-study strategy for the Phase 1 experiments.",
        "Use this to understand how the project is organized around benchmark runs and pooling studies.",
    ),
    (
        "LegalMemoCMT_Adapted_Objective_Summary.docx",
        "Explains how the project is framed as an adapted multimodal emotion recognition study using LegalMemoCMT-style design.",
        "Useful for understanding the project goal, scope, and reporting language.",
    ),
    (
        "LegalMemoCMT_Pipeline_Operations_SOP.docx",
        "Operational guide showing which scripts create datasets, split data, export predictions, and analyze outputs.",
        "Use this when you want to run the project in the correct order.",
    ),
    (
        "LegalMemoCMT_Cloud_SOP.docx",
        "Cloud migration and workflow SOP covering how to move the project to shared or remote execution.",
        "Useful if you want to run the project on Google Drive, Colab, or a cloud VM.",
    ),
    (
        "Implementation_Start_Guide_LegalMemoCMT.docx",
        "Early implementation guide that explains how to begin building the project from the initial scaffold.",
        "Useful for students who want to understand the setup process from the beginning.",
    ),
    (
        "LegalMemoCMT_Phase1_Code_Explanation.docx",
        "Explains the implemented code structure, data flow, fusion modules, training loop, evaluation path, and current scaffold.",
        "Useful if you want a code-level explanation of what exists in the repository, including the custom fusion blocks used in the Phase 1 implementation.",
    ),
    (
        "LegalMemoCMT_Student_Document_Index.docx",
        "Parent index that tells a student which documents to read first and how the collection is organized.",
        "This is the document you are reading now.",
    ),
]


ADAPTATION_AND_POLICY_DOCS: list[tuple[str, str, str]] = [
    (
        "Second_Review_LegalMemoCMT.docx",
        "Second-review discussion document that now includes the post-first-review technical progress, the new benchmark split, and the remaining work toward a stronger final state.",
        "Use this after the first review when you want a detailed student-level continuation of the project story and the next mentor discussion points.",
    ),
    (
        "Phase2_Methodology_LegalMemoCMT.docx",
        "Methodology document for the later legal-domain adaptation phase.",
        "Use this when you want to understand how Phase 1 supports the courtroom-adaptation direction.",
    ),
    (
        "LegalMemoCMT_Video_Modality_Policy.docx",
        "Explains why video is supported in code but treated as an ablation-only stream in the core paper-aligned comparison.",
        "Read this if you want to understand the project policy around the video stream.",
    ),
    (
        "Zeroth_Review_LegalMemoCMT.docx",
        "Early review document covering the initial research direction, motivation, and proposed system framing.",
        "Useful for understanding the starting point of the project.",
    ),
    (
        "First_Review_LegalMemoCMT.docx",
        "First review document covering architecture design, algorithms, expected outcomes, hardware requirements, and the proposed Phase 2 courtroom-use framing.",
        "Useful for seeing how the project evolved after the zeroth review and how the later legal-domain use cases are being described.",
    ),
]


ARCHIVE_AND_DRAFTS: list[tuple[str, str, str]] = [
    (
        "LegalMemoCMT_Multimodal_Learning_Teaching_Draft.docx",
        "Older multimodal teaching draft used during early document creation.",
        "Archive/legacy draft. Read only if you want to see the earlier broad teaching version.",
    ),
    (
        "First_Review_LegalMemoCMT copy.docx",
        "Duplicate copy of the first review document.",
        "Archive copy. Not needed for normal reading.",
    ),
    (
        "Zeroth_Review_LegalMemoCMT copy.docx",
        "Duplicate copy of the zeroth review document.",
        "Archive copy. Not needed for normal reading.",
    ),
    (
        "Zeroth_Review_LegalMemoCMT.docx",
        "Original zeroth review document.",
        "Archive/early-stage review. Useful for history, but not a first-read document.",
    ),
    (
        "First_Review_LegalMemoCMT.docx",
        "Original first review document.",
        "Archive/early-stage review. Useful for history, but not a first-read document.",
    ),
]


def configure_document(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            styles[style_name].font.name = "Times New Roman"

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def curated_and_auto_detected() -> tuple[list[tuple[str, str, str]], list[str]]:
    curated_names = {name for name, _, _ in (CORE_DOCS + OPERATIONAL_DOCS + ADAPTATION_AND_POLICY_DOCS + ARCHIVE_AND_DRAFTS)}
    entries = list(CORE_DOCS + OPERATIONAL_DOCS + ADAPTATION_AND_POLICY_DOCS + ARCHIVE_AND_DRAFTS)
    extras = []
    for path in sorted(DOC_DIR.glob("*.docx")):
        if path.name.startswith("~$"):
            continue
        if path.name == OUTPUT.name:
            continue
        if path.name not in curated_names:
            extras.append(path.name)
    return entries, extras


def build_doc() -> Document:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    run = title.add_run("LegalMemoCMT Student Document Index")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(
        "A parent document that lists the most important files a student should read, in the recommended order."
    )
    run.italic = True
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"

    doc.add_paragraph(
        "Purpose: this index helps a student understand which documents matter most, why they matter, and in what order they should be read. It is intended to stay current as the project grows."
    )
    doc.add_paragraph(
        "Maintenance rule: whenever a new important document is created, rerun this generator script so the index is refreshed."
    )
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("1. Recommended Reading Order", level=1)
    steps = [
        "Read the Fusion Teaching Guide first to understand how the model combines modalities.",
        "Then read the 20-page technical document to compare the current implementation with the MemoCMT base paper.",
        "Then read the expanded and short technical documents if you want a simpler second pass.",
        "Read the Paper-Aligned Runner Guide and Pipeline Operations SOP to connect the ideas to the actual scripts and source files.",
        "Read the MemoCMT-style replication plan and objective summary to understand the report framing.",
        "Read the Phase 2 and Video Policy documents only after Phase 1 is clear.",
        "Use the review documents last, since they explain the early evolution of the project rather than the final student-ready structure.",
    ]
    for step in steps:
        add_bullet(doc, step)

    doc.add_heading("2. Most Important Documents", level=1)
    add_table(
        doc,
        ["Document", "Why a student should read it", "Reading priority"],
        [[name, purpose, note] for name, purpose, note in CORE_DOCS + OPERATIONAL_DOCS + ADAPTATION_AND_POLICY_DOCS],
    )

    doc.add_heading("3. What Each Document Covers", level=1)
    doc.add_heading("3.1 Core Teaching Documents", level=2)
    for name, purpose, note in CORE_DOCS:
        doc.add_heading(name, level=2)
        doc.add_paragraph(purpose)
        doc.add_paragraph(note)

    doc.add_heading("3.2 Operational and Project-Use Documents", level=2)
    for name, purpose, note in OPERATIONAL_DOCS:
        doc.add_heading(name, level=2)
        doc.add_paragraph(purpose)
        doc.add_paragraph(note)

    doc.add_heading("3.3 Adaptation, Review, and Policy Documents", level=2)
    for name, purpose, note in ADAPTATION_AND_POLICY_DOCS:
        doc.add_heading(name, level=2)
        doc.add_paragraph(purpose)
        doc.add_paragraph(note)

    doc.add_heading("3.4 Archive and Draft Documents", level=2)
    for name, purpose, note in ARCHIVE_AND_DRAFTS:
        doc.add_heading(name, level=2)
        doc.add_paragraph(purpose)
        doc.add_paragraph(note)

    doc.add_heading("4. Documents That Are Helpful but Secondary", level=1)
    add_bullet(
        doc,
        "These documents are still useful, but they are not the first files a new student should read. They are better after the main reading order is complete.",
    )
    secondary = [
        "Implementation notes, scripts, and generated analysis files under the results directory.",
        "Any future review decks or supporting notes that explain a narrow experiment or one-off analysis.",
    ]
    for item in secondary:
        add_bullet(doc, item)

    doc.add_heading("5. Auto-Detected New Documents", level=1)
    doc.add_paragraph(
        "Any document created later that is not in the curated list below will appear here after the generator is rerun."
    )
    _, extras = curated_and_auto_detected()
    if extras:
        add_table(doc, ["New document"], [[name] for name in extras])
    else:
        doc.add_paragraph("No additional unlisted documents were detected at generation time.")

    doc.add_heading("6. Student-Facing Update Rule", level=1)
    add_bullet(
        doc,
        "Whenever a new important document is created, add it to the curated list in scripts/build_student_document_index.py if you want it to appear in the top section.",
    )
    add_bullet(
        doc,
        "Rerun the script so the parent index is regenerated and the auto-detected section is refreshed.",
    )
    add_bullet(
        doc,
        "If the new document is only temporary or experimental, leave it out of the curated list and let it appear only in the auto-detected section.",
    )

    doc.add_heading("7. Reference for the Student", level=1)
    doc.add_paragraph(
        "This index is meant to reduce confusion. A student should be able to start from this page, follow the reading order, and then drill down into the technical, operational, and review documents only when needed."
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
