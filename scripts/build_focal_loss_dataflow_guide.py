from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "implementation_docments" / "LegalMemoCMT_Focal_Loss_Dataflow_Student_Guide.docx"
FIG_DIR = ROOT / "implementation_docments" / "figures" / "focal_loss_dataflow"
FIG_DIR.mkdir(parents=True, exist_ok=True)

PIPELINE_SVG = FIG_DIR / "focal_loss_dataflow.svg"
PIPELINE_PNG = FIG_DIR / "focal_loss_dataflow.png"


def render_mermaid(code: str, svg_path: Path, png_path: Path) -> None:
    mmd_path = png_path.with_suffix(".mmd")
    mmd_path.write_text(code, encoding="utf-8")
    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(mmd_path), "-o", str(svg_path), "-b", "white"],
        check=True,
        cwd=str(ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(mmd_path), "-o", str(png_path), "-b", "white"],
        check=True,
        cwd=str(ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build_doc() -> Document:
    render_mermaid(
        """sequenceDiagram
  participant A as Raw MELD/CREMA-D manifest
  participant B as Fold CSVs
  participant C as src/train/train.py
  participant D as Weighted CE or Focal Loss
  participant E as best_model.pt
  participant F as Evaluate on held-out data
  participant G as metrics.json + predictions_test.csv
  participant H as analyze_predictions.py
  A->>B: build train/val splits
  B->>C: load rows from manifest
  C->>D: choose weighted CE or focal loss
  D->>E: save best checkpoint
  E->>F: run evaluation
  F->>G: write metrics and predictions
  G->>H: build confusion matrix and error analysis
""",
        PIPELINE_SVG,
        PIPELINE_PNG,
    )

    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    r = title.add_run("LegalMemoCMT Focal Loss Dataflow Student Guide")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    r = subtitle.add_run(
        "A student-level explanation of how the focal-loss improvement scripts work, what manifests they use, and how the outputs move from raw data to validation metrics and confusion analysis."
    )
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)

    add_para(
        doc,
        "This guide is meant to show the full focal-loss improvement path in one place. It explains the raw datasets, the manifest files, the fold CSVs, the weighted-CE baseline, the focal-loss reruns, the warm-start refinement, and the resume-capable Fold 2 continuation.",
    )

    doc.add_heading("1. Which datasets are used", level=1)
    add_para(
        doc,
        "The focal-loss improvement path uses MELD as the conversational benchmark and CREMA-D as the speaker-independent emotional benchmark. The data is not fed directly into training. Instead, it is converted into manifests and fold CSVs so the scripts can control the train/validation/test split precisely.",
    )
    add_bullets(
        doc,
        [
            "MELD raw manifest: data/manifests/meld_raw.csv",
            "MELD fold manifests: data/manifests/meld_cv/meld_fold_*.csv",
            "CREMA-D manifest: data/manifests/crema_d.csv",
            "CREMA-D fold manifests: data/manifests/crema_d_cv/crema_d_fold_*.csv",
        ],
    )
    add_para(
        doc,
        "The key student point is that focal loss does not change the dataset itself. It changes how the model is rewarded during training when the dataset is imbalanced.",
    )

    doc.add_heading("2. What the paper-aligned baseline did", level=1)
    add_para(
        doc,
        "Before focal loss, the project used the paper-aligned baseline scripts. These scripts trained the same model family with weighted cross-entropy and stored the results under the paper_aligned results tree. Those outputs are the reference point for the improvement experiment.",
    )
    add_bullets(
        doc,
        [
            "Weighted CE is the baseline loss.",
            "MELD paper-aligned CV uses five dialogue folds.",
            "CREMA-D paper-aligned analysis uses the existing benchmark outputs.",
            "The improvement scripts must be compared against those baseline metrics and confusion matrices.",
        ],
    )

    doc.add_heading("3. How focal loss changes training", level=1)
    add_para(
        doc,
        "Focal loss is still built on cross-entropy, but it down-weights easy examples and focuses more on hard examples. In a class-imbalanced dataset, that can help the model pay more attention to minority emotions.",
    )
    add_bullets(
        doc,
        [
            "Easy examples contribute less to the loss.",
            "Hard examples contribute more to the loss.",
            "Class weights are still used so the rare classes matter.",
            "Gamma controls how aggressively the model focuses on hard cases.",
        ],
    )
    add_para(
        doc,
        "In this repository, the trainer already supports focal loss. The improvement scripts simply activate it by passing `--loss-type focal` and `--focal-gamma 2.0`.",
    )
    add_code(
        doc,
        """loss_fn = FocalLoss(gamma=args.focal_gamma, weight=class_weights)
ce = F.cross_entropy(logits, targets, weight=self.weight, reduction="none")
pt = torch.exp(-ce)
loss = ((1 - pt) ** self.gamma) * ce""",
    )

    doc.add_heading("4. The main improvement scripts", level=1)
    add_table(
        doc,
        ["Script", "What it does", "What to inspect"],
        [
            ["scripts/run_improvement_class_balanced_focal_meld_selected.sh", "Runs MELD Fold 2 and Fold 4 with focal loss", "metrics.json, predictions_test.csv, confusion matrices"],
            ["scripts/analyze_improvement_class_balanced_focal_meld_selected.sh", "Exports predictions and builds MELD error analysis", "analysis_test/confusion_matrix.csv"],
            ["scripts/run_improvement_class_balanced_focal_crema_d_cv.sh", "Retrains all CREMA-D folds with focal loss", "fold metrics and summary.json"],
            ["scripts/analyze_improvement_class_balanced_focal_crema_d_cv.sh", "Exports CREMA-D fold predictions and analysis", "analysis_val/confusion_matrix.csv"],
        ],
    )
    add_para(
        doc,
        "These scripts are the focal-loss version of the older paper-aligned pipeline. The main idea is to hold the architecture fixed and change the optimization objective so we can see whether the confusion pattern improves.",
    )

    doc.add_heading("5. The warm-start focal path", level=1)
    add_para(
        doc,
        "The warm-start focal scripts take the next step beyond the from-scratch focal reruns. Instead of learning from initialization, they begin from the already trained weighted-CE checkpoint and then continue with focal loss at a smaller learning rate.",
    )
    add_bullets(
        doc,
        [
            "scripts/run_warmstart_focal_meld_selected.sh starts from the weighted-CE checkpoint and fine-tunes with focal loss.",
            "scripts/run_warmstart_focal_crema_d_cv.sh does the same idea for CREMA-D.",
            "This is a refinement experiment, not a restart experiment.",
            "The purpose is to see whether focal loss can improve an already sensible boundary rather than build one from scratch.",
        ],
    )
    add_para(
        doc,
        "Student interpretation: warm-start focal is a controlled way to ask whether focal loss can polish the baseline instead of fighting against a newly initialized model.",
    )

    doc.add_heading("6. The resume-capable Fold 2 path", level=1)
    add_para(
        doc,
        "The resume-capable Fold 2 script goes one step further than the warm-start focal script. It explicitly loads the weighted-CE Fold 2 checkpoint and then continues training for a few extra epochs with focal loss and a smaller learning rate.",
    )
    add_bullets(
        doc,
        [
            "Wrapper: bash scripts/run_resume_warmstart_focal_meld_fold2.sh",
            "Python entrypoint: scripts/resume_warmstart_focal_meld_fold2.py",
            "Base checkpoint: results/improvement/warmstart_focal/meld_selected/cmt_min/fold_2/base_weighted_ce_checkpoint.pt",
            "Output directory: results/improvement/warmresume_focal/meld_fold_2",
        ],
    )
    add_para(
        doc,
        "This is the narrowest and most targeted focal-loss experiment in the repository. It isolates a single question: if we take the good weighted-CE Fold 2 checkpoint and continue from there with focal loss, do we get a better validation and test result?",
    )
    add_code(
        doc,
        """python3 scripts/run_resume_warmstart_focal_meld_fold2.sh --epochs 5
python3 scripts/analyze_resume_warmstart_focal_meld_fold2.sh""",
    )

    doc.add_heading("7. The manifest flow", level=1)
    add_para(
        doc,
        "The manifest flow is the part students should understand first. The raw annotation CSVs become manifest rows, the fold builder splits those rows into train and validation CSVs, and the training script reads those fold CSVs one fold at a time.",
    )
    add_bullets(
        doc,
        [
            "Raw MELD/CREMA-D data is converted into manifest rows.",
            "The fold CSVs are the actual inputs to the training script.",
            "Training does not change the split unless you deliberately rebuild it.",
            "The manifest rows carry the sample_id, split, label, clip path, and transcript or modality features.",
        ],
    )
    add_code(
        doc,
        """data/manifests/meld_raw.csv
data/manifests/meld_cv/meld_fold_2_train.csv
data/manifests/meld_cv/meld_fold_2_val.csv
data/manifests/crema_d.csv
data/manifests/crema_d_cv/crema_d_fold_0_train.csv""",
    )

    doc.add_heading("8. What the training output folders contain", level=1)
    add_table(
        doc,
        ["Output file", "Meaning", "Why it matters"],
        [
            ["best_model.pt", "Best checkpoint based on validation performance", "This is what later evaluation loads"],
            ["metrics.json", "Accuracy, weighted accuracy, unweighted accuracy, macro F1, weighted F1", "This is the summary of the run"],
            ["predictions_test.csv / predictions_val.csv", "Per-sample predictions", "This enables confusion analysis"],
            ["analysis_test/confusion_matrix.csv or analysis_val/confusion_matrix.csv", "Class-by-class confusion counts", "This reveals the actual error pattern"],
            ["summary.json / summary.md", "Aggregated fold results", "This is useful for fold-level reporting"],
        ],
    )
    add_para(
        doc,
        "The student takeaway is that a model checkpoint is not the final answer. The checkpoint is then turned into evaluation files, and those evaluation files are what you read to understand the effect of focal loss.",
    )

    doc.add_heading("9. How to read the focal-loss results", level=1)
    add_para(
        doc,
        "The most important comparison is not just accuracy. The model might improve or worsen in ways that accuracy does not reveal, especially when the dataset is imbalanced.",
    )
    add_bullets(
        doc,
        [
            "Accuracy tells the total fraction of correct predictions.",
            "Weighted F1 reflects the score after accounting for class frequency.",
            "Macro F1 tells whether the rare classes improved or not.",
            "The confusion matrix tells which classes the model keeps mixing up.",
        ],
    )
    add_para(
        doc,
        "When focal loss works, the confusion matrix should become less collapsed and minority classes should start appearing as distinct predictions instead of being absorbed into the dominant label.",
    )

    doc.add_heading("10. Expected comparison story", level=1)
    add_numbered(
        doc,
        [
            "Weighted CE remains the main baseline to beat.",
            "Focal loss from scratch tests whether the new objective alone helps.",
            "Warm-start focal tests whether focal loss can refine a good baseline.",
            "Resume-capable Fold 2 tests the cleanest refinement path on the most informative fold.",
            "CREMA-D reruns test whether the same objective helps a more speaker-independent benchmark.",
        ],
    )
    add_para(
        doc,
        "If focal loss improves macro F1 and reduces collapse in the confusion matrix, then it is worth keeping as the first improvement step. If not, the evidence says the problem is deeper than the loss function alone.",
    )

    doc.add_heading("11. End-to-end summary for students", level=1)
    add_para(
        doc,
        "The full pipeline is: raw data -> manifest -> fold CSVs -> model training -> best checkpoint -> validation metrics -> test predictions -> confusion matrix -> comparison against the paper-aligned baseline. That is the complete student story for the focal-loss implementation.",
        italic=True,
    )
    doc.add_picture(str(PIPELINE_PNG), width=Inches(6.9))

    return doc


def add_code(doc: Document, code: str) -> None:
    for line in code.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(9.2)


def main() -> None:
    build_doc().save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
