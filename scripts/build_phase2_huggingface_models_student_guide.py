"""Build a student-level technical guide for the Phase 2 model stack."""
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "implementation_docments/LegalMemoCMT_Phase2_HuggingFace_Models_Student_Deep_Dive.docx"

def code(doc, value):
    p = doc.add_paragraph(); r = p.add_run(value); r.font.name = "Courier New"; r.font.size = Pt(9)
    return p

def bullets(doc, values):
    for value in values: doc.add_paragraph(value, style="List Bullet")

def model_section(doc, number, name, purpose, status, input_text, output_text, use_text, interpretation, limits, example):
    doc.add_heading(f"{number}. {name}", level=1)
    doc.add_paragraph(f"Purpose: {purpose}")
    doc.add_paragraph(f"Phase 2 status: {status}")
    doc.add_heading("What the model receives", level=2); doc.add_paragraph(input_text)
    doc.add_heading("What the model produces", level=2); doc.add_paragraph(output_text)
    doc.add_heading("How Phase 2 uses the output", level=2); doc.add_paragraph(use_text)
    doc.add_heading("Technical interpretation", level=2); doc.add_paragraph(interpretation)
    doc.add_heading("Limitations and safeguards", level=2); bullets(doc, limits)
    doc.add_heading("Worked Phase 2 example", level=2); doc.add_paragraph(example)

def main():
    doc=Document()
    sec=doc.sections[0]; sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65)
    doc.add_heading("LegalMemoCMT Phase 2: Hugging Face Models", 0)
    doc.add_paragraph("Student-Level Technical Deep Dive for Slide 22: Phase 2 Hugging Face Models")
    doc.add_paragraph("This document explains the six model components listed in the Phase 2 review presentation. It is deliberately careful about terminology: a pretrained model supplies evidence or a candidate prediction; it does not create human gold labels, courtroom truth, credibility, or deception judgments.")
    doc.add_heading("1. How to read the model stack", level=1)
    doc.add_paragraph("The models solve different problems in the pipeline. Whisper answers what was said and when. Pyannote answers which acoustic speaker cluster was active and when. ViT converts sampled video faces into numerical features. Audio emotion models describe vocal evidence. DeBERTa tests semantic hypotheses about who or when an emotion is being described. The Phase 1 MELD checkpoint supplies the original basic-emotion baseline. These outputs are joined by utterance_id, source video, and timestamps, then passed to explicit rules and human review.")
    doc.add_paragraph("The central design principle is evidence separation. For example, a transcript may mention that another person was hopeless, while the witness speaks calmly. The text model may detect negative meaning, the audio model may detect low excitement, and the video representation may show a stable face. The integrated annotation layer must preserve all of these signals rather than allowing one signal to overwrite the others.")
    doc.add_heading("2. Model inventory from Slide 22", level=1)
    t=doc.add_table(rows=1, cols=4); t.style="Table Grid"
    for cell, value in zip(t.rows[0].cells,["Model / checkpoint","Modality","Phase 2 output","Status"]): cell.text=value
    rows=[
        ["faster-whisper tiny.en","Audio + transcript","ASR text and word timestamps","Used in alignment pipeline"],
        ["pyannote/speaker-diarization-3.1","Audio","Speech intervals and local speaker clusters","Used for diarization"],
        ["google/vit-base-patch16-224-in21k","Video","768-dimensional face-crop embeddings","Used for video features"],
        ["3loi/SER-Odyssey-Baseline-WavLM-Multi-Attributes","Audio","Valence, excitement, dominance evidence","Evidence adapter / pilot"],
        ["speechbrain/emotion-recognition-wav2vec2-IEMOCAP","Audio","Categorical audio emotion candidate","Evidence cross-check"],
        ["MoritzLaurer/deberta-v3-large-zeroshot-v2.0","Text","NLI scope hypotheses","Planned/controlled inference"],
    ]
    for row in rows:
        cells=t.add_row().cells
        for c,v in zip(cells,row): c.text=v
    doc.add_paragraph("A model being installed, named in a configuration, or supported by a script does not prove that its output has already been validated on the full Clancy corpus. I report operational status separately from planned use.")

    model_section(doc,1,"faster-whisper tiny.en","Create an automatic speech-recognition transcript and word-level timing from courtroom audio.","Operational in the validated LegalMELD and Clancy alignment stages. It is an ASR model, not an emotion model.","A WAV segment or hearing audio. The audio should be readable, sufficiently loud, and associated with the correct source video.","A sequence of recognized words with start and end times, plus an audio duration. The pipeline also retains the original subtitle/transcript text.","The builder compares transcript text with ASR text and uses word timestamps to estimate utterance boundaries. It writes fields such as asr_text, transcript_text_normalized, word_timestamp_count, alignment_method, alignment_confidence, and text_similarity.","A simplified matching score is a normalized text similarity between the transcript string and recognized words. If W is the recognized word sequence and T is the normalized transcript, a sequence matcher can be represented as S(T,W) in [0,1]. Exact or strong matching is classified HIGH; weaker matching is sent to review. These are quality indicators, not probabilities that the transcript is legally true.",[
        "tiny.en may misrecognize names, accents, overlapping speakers, legal terminology, or quiet speech.",
        "A high text similarity does not prove that the clip begins and ends at the correct utterance.",
        "The source offset must be applied consistently when raw video begins before the subtitle timeline.",
        "Low-confidence alignment must not enter initial fine-tuning without review.",
    ],"For a witness answer, the subtitle may say ‘For seven and a half years’ while ASR hears unrelated nearby speech. The similarity becomes low, the row is marked for review/rejection, and the original text and ASR text are both preserved for diagnosis.")

    model_section(doc,2,"pyannote/speaker-diarization-3.1","Find speech activity intervals and group similar voices into anonymous local speaker clusters.","Operational in the Clancy diarization environment `.venv-diarization` after compatible Torch and torchaudio versions and Hugging Face access were configured.","A source WAV or extracted audio from a raw courtroom video. The audio should contain the same time base used by the turn manifest.","Segments with start/end times and labels such as SPEAKER_07. The label means a local acoustic cluster, not a named person or legal role.","The pipeline intersects diarization intervals with turn intervals and writes cluster IDs, overlap evidence, and segment-to-turn mappings. Manual review maps clusters to Witness, Prosecutor, Defence, Judge, or Other.","For each diarization segment d=[s_d,e_d] and turn t=[s_t,e_t], overlap seconds are max(0,min(e_d,e_t)-max(s_d,s_t)). A positive overlap supports an association, but multiple clusters overlapping one turn indicate crosstalk, boundary problems, or subtitle grouping. Diarization does not itself determine witness speaking.",[
        "Cluster IDs are local to a source recording and are not automatically consistent across different videos.",
        "Diarization identifies voice activity, not legal role, identity, facial visibility, or emotion.",
        "Parallel workers preserve source-local IDs but do not solve cross-video speaker identity.",
        "Manual cluster-role mappings and visual verification remain necessary.",
    ],"If SPEAKER_07 is manually verified as Defence in one source, the role map can label matching rows in that source. It should not be copied to another source unless an independent identity link exists.")

    model_section(doc,3,"google/vit-base-patch16-224-in21k","Represent sampled face crops as fixed-length video features for downstream multimodal models.","Operational for Clancy video feature extraction. The current extractor uses a Haar largest-face crop with padding and a center-crop fallback.","Frames sampled from each utterance or source video, after face detection/cropping and image preprocessing to the ViT input size.","A 768-dimensional float32 embedding per sampled frame, stored in `.npy` files. A typical row can have an array shaped approximately [N,768], where N is the number of sampled frames.","The embedding path is added to the manifest as the feature path consumed by the Phase 1 video loader. It supplies visual representation; it does not directly output a courtroom-affect label.","For frame embedding x in R^768, the model produces a vector. A downstream model may pool frame vectors, for example mean pooling v_bar=(1/N) sum_i x_i, or use a temporal encoder. Similarity may be measured with cosine similarity cos(x,y)=x·y/(||x||||y||), but the current feature extraction step does not claim that similarity equals emotion.",[
        "A face crop is not proof that the speaking witness is the detected face.",
        "The center-crop fallback may contain a person, background, or courtroom view when detection fails.",
        "The current features are compatible with the Phase 1 loader only when the expected dimensionality is correct.",
        "Visual features should be combined with timestamps and role evidence before selecting witness clips.",
    ],"A row may have `visual_speaker_match=YES` after human review and a 768-D `.npy` feature path. The vector is then available to a model, while the human verification field records why the row is eligible.")

    model_section(doc,4,"3loi/SER-Odyssey-Baseline-WavLM-Multi-Attributes","Provide independent dimensional acoustic evidence about vocal presentation.","Used as an evidence source in the pilot annotation workflow; not a courtroom-affect classifier and not a gold-label generator.","The utterance WAV, normally mono at the configured sample rate and linked to the same utterance_id as the video and transcript.","Continuous valence, excitement, and dominance-like values. In the current terminology, excitement is the activation dimension; it is not a replacement for a basic emotion label.","The values are joined to Phase 1 predictions and transcript-scope fields. Low or moderate excitement can support a controlled presentation, but it cannot alone prove CALM_COMPOSED, neutral, or absence of distress.","A normalized dimensional output can be treated as a point (V,E,D). Distance from a reference or threshold can be computed, but threshold rules are calibration heuristics. For example, `E < 0.45` is a configured screening threshold, not a universal psychological law. It must be checked against reviewed courtroom examples.",[
        "The model is trained on general speech emotion conditions and may not match courtroom speech.",
        "Low excitement does not exclude sadness, fear, or serious subject matter.",
        "Negative valence does not automatically mean distress, anger, or tension.",
        "Use dimensional outputs as low-level evidence and require corroboration for strong affect labels.",
    ],"For the Clancy review example, negative valence and moderate excitement indicate negative activation, but the witness appears controlled and the distress corroboration field is NO. The safe interpretation is not automatic DISTRESSED.")

    model_section(doc,5,"speechbrain/emotion-recognition-wav2vec2-IEMOCAP","Provide a categorical acoustic cross-check against the Phase 1 MELD prediction.","Used as an audio evidence source in the pilot. It is not trained on the LegalMemoCMT courtroom taxonomy.","An utterance WAV with valid audio. The waveform is encoded by wav2vec2 and classified by the SpeechBrain head.","A categorical candidate and confidence, commonly including `neu`, `ang`, `hap`, and `sad`. The Phase 2 mapping is `neu -> neutral`, `ang -> anger`, `hap -> joy`, and `sad -> sadness`.","The candidate is preserved as audio_ser or SpeechBrain evidence and compared with the Phase 1 basic-emotion prediction. Agreement can increase confidence in a candidate; disagreement increases review priority.","The classifier produces logits z_k for categories k and probabilities p_k=exp(z_k)/sum_j exp(z_j). The reported confidence is usually max_k p_k. A value of 1.0 is model confidence, not literal certainty. Because the label spaces differ, missing MELD classes such as fear, disgust, and surprise cannot be treated as absent emotions.",[
        "The four-class IEMOCAP vocabulary is not equivalent to the seven-class MELD vocabulary.",
        "A confidence of 1.0 can be overconfident out of domain.",
        "Audio SER describes vocal expression, not the emotion mentioned in the transcript.",
        "It must not create deception, credibility, or truthfulness labels.",
    ],"If Phase 1 predicts fear but SpeechBrain predicts neutral and the voice is controlled, the disagreement is useful evidence of semantic leakage. The original fear prediction is preserved and human review can propose neutral without overwriting the machine output.")

    model_section(doc,6,"MoritzLaurer/deberta-v3-large-zeroshot-v2.0","Use natural-language inference to propose emotion-target and temporal-scope labels from transcript context.","Planned or controlled enhancement. It should be run only after the Transformers dependency and model access are available, then compared with existing fields.","Normalized transcript text and optionally nearby context. The input should include enough context to resolve pronouns and quotation boundaries.","Scores for natural-language hypotheses such as ‘The witness is describing another person's emotional state’ or ‘The witness is describing their own earlier emotional state’. The highest-scoring hypothesis becomes a machine suggestion with confidence and review metadata.","The output can populate separate DeBERTa evidence fields for target scope and temporal scope. It should support review prioritization and semantic-leakage analysis, not overwrite Phase 1 emotion or human annotations.","For hypotheses h_k, the model produces entailment-related scores. A practical zero-shot score is obtained by comparing entailment and contradiction logits, for example p_k=softmax([z_entail(h_k), z_contradict(h_k)])_entail, then normalizing across candidate labels if configured. The margin m=p_top-p_second and entropy H=-sum_k p_k log p_k are useful uncertainty measures. These formulas describe model evidence, not legal certainty.",[
        "NLI hypotheses can be sensitive to wording, context length, quotation marks, and speaker attribution.",
        "The model can classify semantic scope without knowing who is visibly speaking.",
        "It may confuse reported speech with the witness's own emotion unless context is supplied.",
        "Its output must remain a candidate and be compared with audio, video, role, and human review.",
    ],"For ‘Yes. She said ...’, a suitable output is QUOTED_SPEECH and possibly PAST_OTHER. That prevents the emotional content of the quoted person from being automatically assigned to the current witness.")

    doc.add_heading("3. How the outputs are integrated", level=1)
    code(doc,"row -> provenance check\n     -> ASR/transcript alignment\n     -> diarization and cluster-to-turn overlap\n     -> visual feature extraction\n     -> Phase 1 basic-emotion baseline\n     -> audio SER evidence\n     -> DeBERTa scope hypotheses (planned)\n     -> courtroom-affect candidate rules\n     -> acceptance gate\n     -> SILVER or UNRESOLVED/WEAK")
    doc.add_paragraph("The integration key is the row identity and time interval, not just the filename. Every result should retain utterance_id, source video ID, start_time, end_time, audio path, video path, and transcript source. This makes it possible to inspect the same clip when models disagree.")
    doc.add_heading("4. Core formulas used in interpretation", level=1)
    formulas=[
        ("Softmax", "p_i = exp(z_i) / sum_j exp(z_j)", "Converts logits into a categorical distribution; confidence is not ground truth."),
        ("Prediction confidence", "c = max_i p_i", "Used as a screening value for acceptance gates."),
        ("Entropy", "H(p) = - sum_i p_i log(p_i)", "Higher entropy means the model is less decisive."),
        ("Margin", "m = p_top - p_second", "Small margins identify ambiguous predictions."),
        ("Diarization overlap", "o = max(0, min(e_d,e_t)-max(s_d,s_t))", "Measures segment/turn temporal intersection in seconds."),
        ("Duration", "duration_ms = 1000 * (end_time - start_time)", "Supports MELD-style duration filtering and outlier review."),
        ("Acceptance gate", "basic_confidence >= 0.70 AND affect_confidence >= 0.60 AND NOT critical_conflict", "Creates AUTO_ADJUDICATED/SILVER candidates, not gold labels."),
    ]
    t=doc.add_table(rows=1,cols=3); t.style="Table Grid"
    for c,v in zip(t.rows[0].cells,["Concept","Formula","Phase 2 meaning"]): c.text=v
    for row in formulas:
        cells=t.add_row().cells
        for c,v in zip(cells,row): c.text=v
    doc.add_heading("5. What the models must never be used to claim", level=1)
    bullets(doc,[
        "A diarization cluster is not automatically a verified witness identity.",
        "A basic-emotion prediction is not a courtroom-affect label.",
        "Negative transcript content is not proof that the current witness feels that emotion.",
        "Low excitement is not proof of neutral emotion or absence of sadness.",
        "A face embedding is not proof that the visible person is speaking.",
        "No model output is evidence that a person is deceptive, truthful, credible, or unreliable.",
    ])
    doc.add_heading("6. Student explanation of the Phase 2 outcome", level=1)
    doc.add_paragraph("The correct final interpretation is a structured evidence record. For each row, I can say what the transcript says, which voice was detected, whether the witness role was manually established, what the audio and video models supplied, what semantic scope was proposed, and why the acceptance gate accepted or withheld the row. This is more defensible than selecting one model's label and calling it the final courtroom emotion.")
    doc.add_paragraph("The Phase 2 objective is therefore not to find a pretrained model that understands courtroom behavior automatically. It is to use pretrained models as reproducible evidence generators inside an auditable pipeline, then use explicit rules and human review to create a trustworthy courtroom-specific dataset.")
    doc.add_heading("7. Reading and implementation references", level=1)
    refs=[
        "https://huggingface.co/guillaumekln/faster-whisper",
        "https://huggingface.co/pyannote/speaker-diarization-3.1",
        "https://huggingface.co/google/vit-base-patch16-224-in21k",
        "https://huggingface.co/3loi/SER-Odyssey-Baseline-WavLM-Multi-Attributes",
        "https://huggingface.co/speechbrain/emotion-recognition-wav2vec2-IEMOCAP",
        "https://huggingface.co/MoritzLaurer/deberta-v3-large-zeroshot-v2.0",
    ]
    for r in refs: doc.add_paragraph(r, style="List Bullet")
    doc.save(OUT); print(OUT)

if __name__=='__main__': main()
