from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "implementation_docments" / "LegalMemoCMT_Benchmark_Execution_and_Analysis_Guide.docx"
MELD_ROOT = ROOT / "results" / "paper_aligned_meld_cv" / "cmt_min"
CREMA_ROOT = ROOT / "results" / "paper_aligned_crema_d" / "cmt_min"


def style(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.95)
    sec.right_margin = Inches(0.95)


def p(doc: Document, text: str, italic: bool = False) -> None:
    para = doc.add_paragraph(text)
    if italic:
        for run in para.runs:
            run.italic = True


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def number(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def load_json(path: Path) -> dict[str, object]:
    return json.loads(path.read_text(encoding="utf-8"))


def fold_metrics(root: Path) -> list[dict[str, object]]:
    out = []
    for fold in range(5):
        f = root / f"fold_{fold}"
        if (f / "metrics.json").exists():
            out.append({"fold": fold, "metrics": load_json(f / "metrics.json")})
    return out


def build_doc() -> Document:
    doc = Document()
    style(doc)

    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p0.add_run("LegalMemoCMT Benchmark Execution and Analysis Guide")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "Times New Roman"

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p1.add_run("How the paper-aligned MELD and CREMA-D benchmarks were executed, evaluated, and analyzed from saved outputs")
    r.italic = True
    r.font.size = Pt(13)
    r.font.name = "Times New Roman"

    p(
        doc,
        "This guide explains the script-to-result chain that produced the MELD 5-fold CV analysis and the CREMA-D held-out analysis. It is meant to help a student explain to a mentor how the benchmark runs were carried out, how evaluation was done, and how the final analysis documents were built from the saved outputs.",
    )
    p(
        doc,
        "Important framing: this guide describes completed outputs already present in the workspace. It does not describe a new rerun. The analysis is based on existing metrics.json files, predictions CSVs, confusion matrix CSVs, and the generated report documents.",
    )

    doc.add_heading("1. Benchmark Scope", level=1)
    bullets(
        doc,
        [
            "MELD is treated as the primary paper-aligned benchmark because it is dialogue-based and closer to the MemoCMT case-study setting.",
            "CREMA-D is treated as a secondary benchmark and held-out analysis reference because it is publicly available and helps test the same fusion logic on a different emotion distribution.",
            "The analysis documents explain the results from the saved outputs; they are not just strategy notes.",
        ],
    )

    doc.add_heading("2. MELD 5-Fold Cross-Validation Workflow", level=1)
    p(
        doc,
        "The MELD workflow is organized as a fold-based evaluation. The script builds a raw manifest, generates five folds, trains one model per fold, evaluates each fold on the held-out MELD test split, exports predictions, and then aggregates the fold metrics into a summary report."
    )
    number(
        doc,
        [
            "Build the MELD raw manifest from the dataset files.",
            "Split the train and dev dialogues into five cross-validation folds.",
            "Train the paper-aligned CMT + MIN model on each fold using the fold-specific train and validation manifests.",
            "Evaluate each fold on the common MELD test split and save metrics.json plus predictions_test.csv.",
            "Analyze each fold's saved predictions to create fold-specific confusion matrices and first-20 summary tables.",
            "Aggregate the five fold metrics into a summary table and a fold-summary file.",
        ],
    )
    table(
        doc,
        ["Artifact", "What it contains", "Why it matters"],
        [
            ["fold_*/metrics.json", "Accuracy, weighted accuracy, unweighted accuracy, macro F1, weighted F1, sample count", "Gives the numeric result per fold"],
            ["fold_*/predictions_test.csv", "Per-sample actual vs predicted values, confidence, and correctness", "Enables error inspection without rerunning evaluation"],
            ["fold_*/analysis_test/confusion_matrix.csv", "Actual-vs-predicted class matrix", "Shows which emotions are confused with which others"],
            ["fold_*/analysis_test/top_confusions.csv", "Largest off-diagonal confusions", "Provides a compact error summary"],
            ["fold_summary.md/json", "Aggregated fold-level metrics", "Gives the overall fold comparison"],
        ],
    )

    doc.add_heading("3. MELD Fold-Level Findings", level=1)
    p(
        doc,
        "The completed MELD CV outputs show a tight band of performance across all five folds. Fold 2 has the best weighted aggregate performance, while fold 4 has the best class balance as measured by macro F1 and unweighted accuracy. That trade-off is why fold 2 is a strong candidate if the goal is a representative fold with the highest weighted score, but fold 4 is important if the report emphasizes class balance."
    )
    table(
        doc,
        ["Fold", "Accuracy", "Weighted Acc.", "Unweighted Acc.", "Macro F1", "Weighted F1"],
        [
            [f"fold_{row['fold']}", f"{row['metrics']['accuracy']:.4f}", f"{row['metrics']['weighted_accuracy']:.4f}", f"{row['metrics']['unweighted_accuracy']:.4f}", f"{row['metrics']['macro_f1']:.4f}", f"{row['metrics']['weighted_f1']:.4f}"]
            for row in fold_metrics(MELD_ROOT)
        ],
    )
    bullets(
        doc,
        [
            "The recurring MELD confusion pattern is neutral, joy, sadness, and anger being mixed.",
            "This means the model is learning broad emotional polarity more reliably than fine-grained class boundaries.",
            "The fold-by-fold confusion matrices make it clear that the errors are stable and not random.",
        ],
    )

    doc.add_heading("4. How the MELD Analysis Report Was Built", level=1)
    number(
        doc,
        [
            "Use the saved fold predictions_test.csv files.",
            "Generate per-fold confusion matrices and predicted-vs-actual tables from those saved CSVs.",
            "Render each confusion matrix as an image for the Word report.",
            "Write a fold-by-fold explanation of the dominant confusions.",
            "Add a small interpretation section explaining why fold 2 and fold 4 are the key comparison points.",
        ],
    )
    p(
        doc,
        "This is why the MELD report is not merely a training log. It is a post-hoc analysis document built from already completed evaluation outputs. The report emphasizes how the model confuses classes, not just what score it achieved."
    )

    doc.add_heading("5. CREMA-D Held-Out Analysis Workflow", level=1)
    p(
        doc,
        "The CREMA-D workflow is narrower than the MELD fold workflow. It uses the saved held-out predictions from the existing paper-aligned CMT + MIN run, then generates a confusion matrix, a top-confusions table, and a first-20 example table. The analysis document turns those saved outputs into a mentor-readable interpretation."
    )
    number(
        doc,
        [
            "Take the already saved CREMA-D held-out predictions and metrics.",
            "Generate the confusion matrix and top-confusions CSV from the predictions.",
            "Render the confusion matrix as an image.",
            "Create the report that explains why the model is collapsing toward fear in this run.",
        ],
    )
    table(
        doc,
        ["Artifact", "Meaning", "Interpretation"],
        [
            ["metrics.json", "Low accuracy / low macro F1 / low weighted F1", "The model is weak on this run"],
            ["confusion_matrix.csv/png", "Actual-vs-predicted class counts", "Shows a strong collapse into fear"],
            ["top_confusions.csv", "Largest off-diagonal errors", "Shows which emotions are most often misread"],
            ["predicted_vs_actual_first20.csv", "First sample-level predictions", "Useful for concrete explanation"],
        ],
    )

    doc.add_heading("6. CREMA-D Technical Reading", level=1)
    bullets(
        doc,
        [
            "The model is not making random mistakes; it is over-predicting a narrow emotion class.",
            "That kind of confusion suggests a learned bias or representation mismatch, not just noise.",
            "The report explicitly warns that this is not a strong benchmark result.",
            "The value of the report is diagnostic: it shows how the model fails on a specific dataset under the current configuration.",
        ],
    )

    doc.add_heading("7. Why the Two Analyses Are Different", level=1)
    bullets(
        doc,
        [
            "MELD analysis is fold-based and reports stability across five splits.",
            "CREMA-D analysis is a single held-out result and is therefore more of a diagnostic snapshot.",
            "MELD lets you compare fold 2 and fold 4; CREMA-D lets you discuss a collapse pattern around fear.",
            "The two reports serve different teaching purposes: one is about robustness across folds, the other is about a failure mode on a saved held-out run.",
        ],
    )

    doc.add_heading("8. How to Explain This to a Mentor", level=1)
    p(
        doc,
        "A concise mentor explanation is: the MELD benchmark was executed as a 5-fold paper-aligned cross-validation study, and the saved fold predictions were then turned into per-fold confusion matrices and a fold-summary report. The CREMA-D benchmark was evaluated on an existing held-out run, and that output was converted into a confusion-matrix report that shows a strong class-collapse pattern. Both reports are analysis documents built from saved outputs, not from fresh reruns."
    )
    p(
        doc,
        "If the mentor asks why fold 2 was highlighted, the correct answer is that fold 2 has the strongest weighted performance, while fold 4 is the better balanced fold. If the mentor asks why CREMA-D is discussed differently, the answer is that the current saved CREMA-D output is a diagnostic held-out analysis rather than a full 5-fold study."
    )

    doc.add_heading("9. Relevant Scripts and Outputs", level=1)
    table(
        doc,
        ["Script / File", "Role"],
        [
            ["scripts/run_paper_aligned_meld_cv.sh", "Defines the MELD 5-fold CV experiment and saves fold outputs"],
            ["scripts/analyze_saved_meld_cv_folds.py", "Builds per-fold MELD analysis folders and aggregate fold summary"],
            ["scripts/build_meld_cv_analysis_report.py", "Creates the MELD fold report docx from the saved fold outputs"],
            ["scripts/run_paper_aligned_crema_d_analysis.sh", "Defines the CREMA-D held-out analysis export path"],
            ["scripts/build_crema_d_analysis_report.py", "Creates the CREMA-D analysis report docx from the saved held-out output"],
            ["scripts/render_confusion_matrix_images.py", "Renders confusion matrix CSVs as PNG figures"],
        ],
    )

    doc.add_heading("10. Output Locations", level=1)
    bullets(
        doc,
        [
            "MELD fold analyses: results/paper_aligned_meld_cv/cmt_min/fold_*/analysis_test/",
            "MELD fold summary: results/paper_aligned_meld_cv/cmt_min/fold_summary.md and fold_summary.json",
            "CREMA-D analysis: results/paper_aligned_crema_d/cmt_min/analysis_test/",
            "Reports: implementation_docments/LegalMemoCMT_MELD_CV_5Fold_Analysis_Report.docx and LegalMemoCMT_CREMA_D_Analysis_Report.docx",
        ],
    )

    return doc


def main() -> None:
    if not MELD_ROOT.exists():
        raise FileNotFoundError(MELD_ROOT)
    if not CREMA_ROOT.exists():
        raise FileNotFoundError(CREMA_ROOT)
    doc = build_doc()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
