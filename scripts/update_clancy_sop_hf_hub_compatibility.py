from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOP = ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx"


def code(document: Document, value: str) -> None:
    paragraph = document.add_paragraph(style="No Spacing")
    run = paragraph.add_run(value)
    run.font.name = "Courier New"
    run.font.size = Pt(8)


document = Document(SOP)
document.add_section(WD_SECTION.NEW_PAGE)
document.add_heading("Hugging Face Hub Compatibility Fix", level=1)
document.add_paragraph(
    "A diarization preflight initially failed with `TypeError: hf_hub_download() got an unexpected "
    "keyword argument 'use_auth_token'`. This was a dependency mismatch: pyannote.audio 3.4.0 uses "
    "the legacy authentication keyword internally, while the installed Hugging Face Hub 1.x API no "
    "longer accepts it. This error does not prove that the token is invalid or that the diarization "
    "model is unavailable."
)
document.add_heading("Reproducible Correction", level=2)
document.add_paragraph(
    "The repository now pins `huggingface_hub==0.13.4` together with torch 2.2.2, torchaudio 2.2.2, "
    "and pyannote.audio 3.4.0 in the isolated `.venv-diarization` environment. Hub 0.13.4 retains the "
    "compatibility layer that translates pyannote's `use_auth_token` call to the supported token request. "
    "Do not install pyannote into the main Anaconda environment for this stage."
)
code(document, '''bash
phase2/setup_clancy_diarization.sh

./.venv-diarization/bin/python - <<'PY'
import huggingface_hub, pyannote.audio, torch, torchaudio
print("huggingface_hub", huggingface_hub.__version__)
print("pyannote.audio", pyannote.audio.__version__)
print("torch", torch.__version__)
print("torchaudio", torchaudio.__version__)
print("AudioMetaData", hasattr(torchaudio, "AudioMetaData"))
PY''')
document.add_heading("Model Access Test", level=2)
code(document, '''export HF_TOKEN="<your Hugging Face read token>"
./.venv-diarization/bin/python phase2/check_clancy_diarization_prerequisites.py \\
  --model pyannote/speaker-diarization-3.1 \\
  --load-model''')
document.add_paragraph(
    "A successful dependency test should show Hub 0.13.4, pyannote.audio 3.4.0, matching Torch audio "
    "versions, and `AudioMetaData=True`. If model loading then reports that the pipeline is gated or "
    "returns no pipeline, accept the model conditions on Hugging Face and retry with a real read token. "
    "Never write the token into a CSV, source file, DOCX, shell script, or git history."
)
document.add_paragraph(
    "The preflight scripts retain a compatibility fallback for `Pipeline.from_pretrained`: they try the "
    "newer `token` argument first and use `use_auth_token` when required by the installed pyannote API. "
    "The Hub pin is still necessary because pyannote also calls `hf_hub_download` internally."
)
document.save(SOP)
print(f"Updated {SOP}")
