#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "Phase1_Raw_MP4_Demo_SOP_Student_Guide.docx"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_code(doc: Document, text: str) -> None:
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def build_doc() -> Document:
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Phase 1 Raw MP4 Demo SOP")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Student-level guide to running the MELD raw-video demo, explaining the scripts, the outputs, and the reasons predictions can be wrong")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)

    add_para(
        doc,
        "Purpose: this document is written as a speaking and defense guide. It tells you exactly how to run the live raw-mp4 demo from Runpod, what each script does, what the model output means, and where in the code you should look if a prediction is wrong. The goal is not just to run the demo, but to be able to explain it under cross-examination.",
    )

    doc.add_heading("1. What This Demo Proves", level=1)
    add_bullets(
        doc,
        [
            "The model can start from a raw .mp4 clip and produce a final emotion prediction.",
            "The pipeline can show the ground-truth label, the predicted label, the confidence, and the top-3 alternatives.",
            "The same live pipeline can be reused for the baseline checkpoint and for the later gated/aux checkpoints.",
            "The demo can also show the metrics summary, the confusion matrix, and the top confusions so the reviewer sees system-level behavior, not just one sample.",
        ],
    )
    add_para(
        doc,
        "Student explanation: the demo is a full input-to-output story. The reviewer gives a video clip, the pipeline extracts features, the trained model predicts an emotion, and the analysis tables show whether the behavior is stable or biased.",
    )

    doc.add_heading("2. Files and Artifacts You Must Have", level=1)
    add_table(
        doc,
        ["Artifact", "Typical path", "Why it is needed"],
        [
            ["Raw video clips", "data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/*.mp4", "These are the clips you present live during the demo."],
            ["Manifest", "data/manifests/meld_test.csv", "This tells the pipeline the sample_id, label, transcript, and file metadata."],
            ["Baseline checkpoint", "results/paper_aligned_meld_cv/cmt_min/fold_2/best_model.pt", "This is the stable Phase 1 reference model."],
            ["Gated + aux checkpoint", "results/facial_cues/meld_vit_facecrop_gated_video_aux/fold_4/best_model.pt", "This is the later facial-cue improvement model, if present on the pod."],
            ["Analysis bundle", "results/phase1_review_demo/fold2_baseline/", "This stores the demo examples, metrics summary, confusion matrix, and top confusions."],
        ],
    )
    add_para(
        doc,
        "Important safety rule: keep the raw clips separate from the model outputs. The clips are inputs, while the checkpoint and analysis folders are derived artifacts. That separation helps you explain the pipeline cleanly and prevents accidental editing of source data.",
    )

    doc.add_heading("3. The Scripts and Their Roles", level=1)
    add_table(
        doc,
        ["Script", "Role", "What to say in the viva"],
        [
            ["scripts/git_add_demo_raw_test_clips.sh", "Stages only the five approved raw mp4 clips.", "This is a controlled allowlist; it prevents accidentally pushing the full dataset."],
            ["scripts/run_demo_paper_aligned_raw_mp4.sh", "Wrapper for the baseline checkpoint demo.", "This is the stable reference demo run."],
            ["scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh", "Wrapper for the gated+aux checkpoint demo.", "This shows the later improved model if the checkpoint is available."],
            ["scripts/run_phase1_raw_mp4_demo.sh", "Main raw-mp4-to-prediction runner.", "This is the real engine that extracts features and runs inference."],
            ["scripts/predict_phase1_raw_mp4_demo.py", "Python inference logic.", "This is where the clip is read, features are extracted, and probabilities are produced."],
            ["scripts/run_phase1_raw_mp4_demo_batch.sh", "Runs multiple clips from a CSV.", "This is useful when you want a repeatable batch demo and logs."],
            ["scripts/run_phase1_review_demo_bundle.sh", "Builds the reviewer bundle from saved outputs.", "This prepares the summary files, example table, confusion matrix, and top confusions."],
        ],
    )

    doc.add_heading("4. Why the Surprise Example Was Predicted as Anger", level=1)
    add_para(
        doc,
        "The clip `test_dia278_utt5` is a useful teaching example because it is a near-miss rather than a completely failed prediction. The ground truth is surprise, but the model predicts anger. The top-3 probabilities are close enough to show that the model is uncertain and is moving around the correct emotional region.",
    )
    add_table(
        doc,
        ["Field", "Value", "What it suggests"],
        [
            ["Ground truth", "surprise", "This is the correct MELD label for the clip."],
            ["Prediction", "anger", "The fused features leaned slightly toward a nearby high-arousal class."],
            ["Top-1 confidence", "0.2123", "The decision is not strongly confident."],
            ["Top-3 surprise score", "0.1692", "Surprise remains close to the decision boundary."],
            ["Top-3 neutral score", "0.1770", "Neutral is also nearby, which often happens when the visual signal is brief or weak."],
        ],
    )
    add_bullets(
        doc,
        [
            "Surprise and anger are both high-arousal emotions, so their facial signals can overlap.",
            "The pipeline samples a fixed number of frames, so it may miss the exact peak moment of surprise.",
            "If the sampled frames show tension or open-mouthed motion without enough context, the model may lean toward anger.",
            "A face crop that is slightly off or too tight can weaken the visual cue before ViT.",
            "MELD class imbalance also makes the minority emotions harder to separate cleanly.",
        ],
    )
    add_para(
        doc,
        "Student explanation: this is not a random mistake. It shows that the model has learned something about the emotional space, because surprise is still near the top. But it has not perfectly separated similar high-arousal emotions, so the final vote drifted toward anger.",
    )
    add_para(
        doc,
        "If the reviewer asks where the error can come from in the code, explain the pipeline in order: frame sampling in src/data/preprocessing.py, face-crop or full-frame feature extraction in the ViT branch, modality fusion in src/models/model.py, and checkpoint selection in src/train/train.py. That is the correct technical path for explaining a wrong prediction.",
    )
    add_para(
        doc,
        "A good viva answer is: 'This is a near-miss rather than a collapse. The model still ranked surprise in the top-3, but the sampled frames and the learned fusion weights leaned slightly toward anger. That tells me the model is learning meaningful emotion structure, but it still confuses similar high-arousal expressions under MELD imbalance.'",
    )

    doc.add_heading("5.1 Labeled Example: Confidently Wrong Prediction", level=2)
    add_para(
        doc,
        "The clip `test_dia278_utt5` is also useful as a second kind of error example, because the model is not only wrong, it is very confident: the ground truth is surprise, but the prediction is joy with confidence 0.9435. This is more severe than a small near-miss, because the model is not hovering near the decision boundary; it has strongly committed to the wrong class.",
    )
    add_table(
        doc,
        ["Field", "Value", "What it means"],
        [
            ["Ground truth", "surprise", "The MELD annotation says the clip should be surprise."],
            ["Prediction", "joy", "The fused model representation strongly favored joy."],
            ["Confidence", "0.9435", "The model was very certain, but certainty does not guarantee correctness."],
            ["Top-2 class", "surprise 0.0284", "Surprise was far behind the winning class."],
            ["Top-3 class", "anger 0.0092", "A third class still appeared, but with very low probability."],
        ],
    )
    add_bullets(
        doc,
        [
            "The transcript itself can bias the text branch toward joy: `That’s not true, there are great pictures of us!` sounds positive or playful.",
            "The model may therefore learn a happy or approving tone even if the label is surprise.",
            "The face-crop frames may also show smiling or friendly facial cues instead of the brief surprise peak.",
            "Because the fusion block combines modalities, one strong modality can dominate the final decision.",
            "High confidence here means the model’s internal score gap was large; it does not mean the answer is correct.",
        ],
    )
    add_para(
        doc,
        "Student explanation: this is a confident misclassification, not a random error. The text and possibly the sampled facial frames may both look more like joy than surprise, so the fused classifier locked onto joy. The important lesson is that softmax confidence measures certainty, not truth. If the transcript, facial expression, and sampling all lean in one direction, the model can become very sure and still be wrong.",
    )
    add_para(
        doc,
        "If the reviewer asks where to inspect this in the code, explain the same stack in order: `src/data/preprocessing.py` for frame sampling, the face-crop or full-frame extraction branch in the ViT scripts, `src/models/model.py` for fusion, and `src/train/train.py` for how the model was trained and selected. A confident wrong prediction often means the model learned the wrong association too strongly, not that the code crashed.",
    )

    doc.add_heading("5.2 How to Diagnose Modal Bias, Label Ambiguity, and Frame Sampling", level=2)
    add_para(
        doc,
        "To decide why this example failed, do not jump to one explanation. Check three things in order: modal bias, label ambiguity, and frame sampling. That gives you a defensible student-style diagnosis instead of a guess.",
    )
    add_numbered(
        doc,
        [
            "First, test whether the prediction changes when you reduce the active modalities.",
            "Second, inspect the transcript and ask whether the label could be context-dependent or ambiguous.",
            "Third, inspect the cached video embedding and ask whether the sampled frames likely missed the peak expression.",
        ],
    )
    add_table(
        doc,
        ["Diagnostic question", "How to test it", "What would support it"],
        [
            ["Is the model text/audio biased?", "Run the same clip with `MODALITIES=text,audio` and compare it with `MODALITIES=text,audio,video`.", "If the prediction stays close to joy or changes strongly when video is removed, the fusion may be leaning on text/audio."],
            ["Is the label ambiguous?", "Read the transcript and compare it to the MELD label in the manifest.", "If the text sounds playful or positive while the label is surprise, the example may be context-dependent and hard to separate."],
            ["Did frame sampling miss the peak?", "Inspect the cached `.npy` shape and remember that the script samples a fixed number of frames.", "If the expression is brief, the model may see smiling or neutral-looking frames instead of the true surprise moment."],
        ],
    )
    add_code(
        doc,
        r"""# 1) Compare the full multimodal prediction with a text+audio-only run
DEVICE=cuda \
MODALITIES=text,audio \
bash scripts/run_demo_paper_aligned_raw_mp4.sh \
  test_dia278_utt5 \
  data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia278_utt5.mp4

# 2) Compare with video-only
DEVICE=cuda \
MODALITIES=video \
bash scripts/run_phase1_raw_mp4_demo.sh \
  test_dia278_utt5 \
  data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia278_utt5.mp4

# 3) Inspect the cached visual embedding shape
python3 scripts/read_meld_vit_facecue_npy.py \
  --file results/phase1_review_demo/raw_mp4_cache/test_dia278_utt5_facecrop.npy""",
    )
    add_para(
        doc,
        "What to look for in the `.npy` inspection: the shape should match the ViT-based cache, the values should be finite, and the embedding should not be empty. This does not prove the prediction is correct, but it does prove the feature-extraction step produced a usable visual tensor.",
    )
    add_para(
        doc,
        "Student explanation: if the text+audio-only run also prefers joy, then the transcript and audio are likely biasing the model. If the video-only run behaves differently, that means the visual branch contains a different signal but it is being overruled in the fused model. If the `.npy` embedding exists with the expected ViT size and finite values, the feature extraction step is not broken; the problem is then more likely the content of the sampled frames or the fusion decision itself.",
    )
    add_para(
        doc,
        "A very strong viva answer is: 'I diagnose this as a combination of modal bias and label ambiguity. The transcript is positive-sounding, so the text branch can lean toward joy. Surprise is a context-sensitive label in MELD, so the example is inherently ambiguous. The raw video is sampled into a fixed number of frames, so if the peak surprise moment is short, the ViT branch may not capture it strongly enough. The final classifier then becomes confidently wrong because multiple weak or misleading signals point to joy.'",
    )

    doc.add_heading("5. Recommended Run Order", level=1)
    add_para(
        doc,
        "Use this order when you are about to present the demo on Runpod. The first steps prepare the evidence, the middle steps run the live clip predictions, and the last step shows the summary analysis.",
    )
    add_numbered(
        doc,
        [
            "Check that the raw demo clips are already available on Runpod at the expected MELD test paths.",
            "Check that the checkpoint files you want to compare are present on the pod.",
            "Build or refresh the reviewer bundle so you have the metrics summary, confusion matrix, and top confusions ready.",
            "Run the baseline paper-aligned demo first.",
            "Run the gated+aux demo second on the same clip.",
            "Optionally repeat the same pair on one more clip so the reviewer sees a correct case and a near-miss case.",
            "Use the batch runner only if you want a prepared CSV-based bundle of multiple clips.",
        ],
    )
    add_para(
        doc,
        "Exact command sequence for the common live demo:",
    )
    add_code(
        doc,
        r"""# 1) Prepare the reviewer bundle from saved outputs
bash scripts/run_phase1_review_demo_bundle.sh

# 2) Run the baseline paper-aligned checkpoint on one clip
DEVICE=cuda \
bash scripts/run_demo_paper_aligned_raw_mp4.sh \
  test_dia279_utt9 \
  data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia279_utt9.mp4

# 3) Run the gated+aux checkpoint on the same clip
DEVICE=cuda \
CHECKPOINT=results/facial_cues/meld_vit_facecrop_gated_video_aux/fold_4/best_model.pt \
bash scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh \
  test_dia279_utt9 \
  data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia279_utt9.mp4

# 4) Repeat on a near-miss example if needed
DEVICE=cuda \
bash scripts/run_demo_paper_aligned_raw_mp4.sh \
  test_dia278_utt5 \
  data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia278_utt5.mp4""",
    )
    add_para(
        doc,
        "If you prefer a repeatable list of clips, use the sample CSV at `implementation_docments/phase1_raw_mp4_demo_pairs_sample.csv` or create your own CSV with `sample_id` and `video_path` columns, then run `scripts/run_phase1_raw_mp4_demo_batch.sh` on that file.",
    )

    doc.add_heading("6. The Exact Live Demo Command Pattern", level=1)
    add_para(
        doc,
        "For the baseline checkpoint, the simplest live demo command on Runpod is:",
    )
    add_code(
        doc,
        r"""DEVICE=cuda \
bash scripts/run_demo_paper_aligned_raw_mp4.sh \
  test_dia279_utt9 \
  data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia279_utt9.mp4""",
    )
    add_para(
        doc,
        "For the gated+aux checkpoint, use the same raw clip but point at the gated wrapper:",
    )
    add_code(
        doc,
        r"""DEVICE=cuda \
CHECKPOINT=results/facial_cues/meld_vit_facecrop_gated_video_aux/fold_4/best_model.pt \
bash scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh \
  test_dia279_utt9 \
  data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia279_utt9.mp4""",
    )
    add_para(
        doc,
        "Student explanation: the wrapper scripts are just convenience entry points. The actual prediction work happens in scripts/run_phase1_raw_mp4_demo.sh and scripts/predict_phase1_raw_mp4_demo.py.",
    )

    doc.add_heading("7. What Happens Inside the Raw MP4 Pipeline", level=1)
    add_numbered(
        doc,
        [
            "The script checks that the manifest, checkpoint, and raw mp4 all exist.",
            "It loads the MELD row for the chosen sample_id.",
            "It builds a preprocess config with 224x224 frames and a fixed number of sampled frames.",
            "It loads a pretrained ViT image processor and ViT encoder.",
            "It extracts either face-crop features or full-frame features from the raw mp4.",
            "It saves the extracted embeddings as a cached .npy file so the clip does not need to be processed again.",
            "It builds a one-sample dataset entry with the manifest metadata and the new video features.",
            "It loads the trained LegalMemoCMTPhase1 checkpoint.",
            "It applies modality masking so the model uses the chosen modalities.",
            "It computes logits, converts them to softmax probabilities, and prints the predicted label and top-3 classes.",
        ],
    )
    add_para(
        doc,
        "This order matters. If you understand it in sequence, you can explain the pipeline from the raw video all the way to the final label.",
    )

    doc.add_heading("8. What the Main Python Script Does", level=1)
    add_para(
        doc,
        "The main inference logic is in scripts/predict_phase1_raw_mp4_demo.py. That script is the best one to study if the reviewer asks, 'What exactly happens to my video?'",
    )
    add_table(
        doc,
        ["Code responsibility", "Where to inspect", "Student explanation"],
        [
            ["Load the manifest row", "load_manifest(args.manifest)", "This recovers the transcript, split, and ground-truth label for the selected sample_id."],
            ["Select facecrop or fullframe mode", "if args.vision_mode == 'facecrop': ...", "This decides whether the model sees cropped faces or full sampled frames before ViT."],
            ["Extract ViT embeddings", "extract_vit_facecrop_embeddings(...) / extract_vit_face_embeddings(...)", "This is the visual feature extraction step that converts image frames into reusable numeric embeddings."],
            ["Cache the extracted embeddings", "np.save(cached_feature_path, video_features)", "This makes the demo faster the next time you run the same clip."],
            ["Build a one-sample batch", "ManifestDataset(...) and collate_samples(...)", "This makes the single clip look like one normal training example."],
            ["Load the model", "LegalMemoCMTPhase1(model_cfg).to(device)", "This recreates the trained architecture before loading its learned weights."],
            ["Run inference", "logits = model(...)", "This computes the class scores for the emotions."],
            ["Turn logits into probabilities", "torch.softmax(logits, dim=-1)", "This gives the top-3 output probabilities used in the demo."],
        ],
    )
    add_para(
        doc,
        "The top-3 values are softmax probabilities, not accuracy and not F1. They tell you how much the model 'leans' toward each class for that specific example.",
    )

    doc.add_heading("9. Why Predictions Can Be Wrong", level=1)
    add_para(
        doc,
        "When a prediction is wrong, you should not answer vaguely. You should connect the error to one of the concrete pipeline stages below.",
    )
    add_table(
        doc,
        ["Possible cause", "Where it appears in code", "What it means in practice"],
        [
            ["Bad or incomplete frame sampling", "src/data/preprocessing.py", "The sampled frames may miss the facial expression or the relevant moment."],
            ["Face crop too tight or too loose", "scripts/build_meld_vit_facecrop_manifest.py and the face-crop branch in the raw demo", "Important visual context may be removed, or the face may be too small to interpret well."],
            ["Model dominated by text/audio", "src/train/train.py and src/models/model.py", "The fusion setup may let the stronger text/audio branch overpower the video signal."],
            ["Class imbalance bias", "training losses and fold behavior in src/train/train.py", "The model may over-predict neutral or other frequent classes."],
            ["Checkpoint selection not aligned with the best metric", "train.py validation checkpoint logic", "If the best checkpoint is chosen by accuracy only, the model may not be the best for macro F1 or minority classes."],
            ["Wrong sample or label mapping", "load_manifest() and manifest rows", "The clip can be correct, but the wrong row or label can make the explanation misleading."],
        ],
    )
    add_para(
        doc,
        "Student explanation: when the model gets a clip wrong, the most disciplined answer is to say which stage is responsible. That makes your explanation technical instead of emotional or speculative.",
    )

    doc.add_heading("10. How to Check Whether a Checkpoint Actually Uses Video", level=1)
    add_para(
        doc,
        "A checkpoint can only be called video-aware if the model configuration and the learned weights both support the video branch. In this repository, the raw demo can still accept video features even when the checkpoint was not trained strongly on video, so you need to check both the script settings and the saved state dictionary.",
    )
    add_bullets(
        doc,
        [
            "Check the active modalities in the demo command.",
            "Check whether the checkpoint was trained with the video branch enabled.",
            "Check whether the loaded state_dict contains video-related weights.",
            "Check whether the checkpoint was saved from a run that actually used `--modalities text,audio,video` or a video-only / gated video experiment.",
        ],
    )
    add_table(
        doc,
        ["Check", "Where to look", "What it tells you"],
        [
            ["Modalities flag", "scripts/run_phase1_raw_mp4_demo.sh and scripts/predict_phase1_raw_mp4_demo.py", "If video is removed here, the model is not using video at inference time."],
            ["Model config", "model_cfg.use_video = 'video' in modalities", "This turns the video branch on or off in the model."],
            ["Checkpoint weights", "model.load_state_dict(checkpoint['model_state'], strict=False)", "Missing or unexpected keys show whether the checkpoint matches the current architecture."],
            ["Training recipe", "scripts/run_* and scripts/analyze_* for the checkpoint family", "This tells you whether the checkpoint came from a text+audio run, a gated run, or an aux-loss run."],
        ],
    )
    add_para(
        doc,
        "A practical rule is: if the checkpoint came from a text-and-audio-only run, then face crop or full frame at demo time will not suddenly make it a video-trained model. The script may still run, but the model will not have learned to use the video signal meaningfully.",
    )
    add_para(
        doc,
        "A stronger video-aware checkpoint will usually show video-related parameters in its state dict and will come from a training script that explicitly enabled the video branch. When the log prints missing keys for a new auxiliary or gated module, that often means the architecture has changed and the checkpoint is only partially warm-started. That is normal as long as the intended base weights loaded and the new branch is intentionally initialized.",
    )
    add_para(
        doc,
        "For viva purposes, you can say: 'I know the checkpoint is video-aware because the training command included the video modality, the model config enabled the video branch, and the saved weights include the corresponding video module parameters. If any of those are missing, the run is not a true video-trained comparison.'",
    )

    doc.add_heading("11. How to Explain the Output Fields", level=1)
    add_table(
        doc,
        ["Field", "Meaning", "How to explain it"],
        [
            ["sample_id", "The unique MELD utterance identifier.", "This identifies the exact clip and the exact manifest row."],
            ["split", "train/dev/test partition.", "This tells you which dataset partition the clip belongs to."],
            ["video_path", "Path to the raw mp4 file.", "This is the original evidence clip used by the live demo."],
            ["cached_video_features", "The .npy file produced by feature extraction.", "This is the numeric ViT output saved for reuse."],
            ["ground_truth_label", "The true emotion class from MELD.", "This is the label the model should match."],
            ["predicted_label", "The class selected by argmax.", "This is the model’s final decision."],
            ["confidence", "The top-1 softmax probability.", "This shows how strongly the model preferred its chosen label."],
            ["top_k", "The highest-probability classes and scores.", "This is not accuracy; it is the model’s probability ranking for that sample."],
        ],
    )

    doc.add_heading("12. How to Defend the Demo in a Viva", level=1)
    add_bullets(
        doc,
        [
            "If asked why the baseline demo is valid, say that it uses a frozen trained checkpoint and a real MELD test clip.",
            "If asked why the same clip can be reused, say that the input stays fixed while the checkpoint changes, which isolates the effect of the model stage.",
            "If asked why the top-3 classes matter, say that they reveal uncertainty and nearby class confusion.",
            "If asked why a prediction is wrong, walk through preprocessing, feature extraction, fusion, and checkpoint selection in that order.",
            "If asked whether a wrong prediction means the whole system failed, say no: the confusion matrix and top confusions show the error is structured, not random.",
        ],
    )
    add_para(
        doc,
        "A strong student answer is: 'The clip is correct as input, the pipeline extracted the visual features properly, but the final fusion layer still favored another class because the model is biased by the training distribution and the nearby emotions are hard to separate.'",
    )

    doc.add_heading("13. Common Demo Debugging Issue: Video Feature Shape Mismatch", level=1)
    add_para(
        doc,
        "If the raw demo crashes with a message like 'mat1 and mat2 shapes cannot be multiplied (32x128 and 768x256)', that means the demo built the wrong video tensor shape for the model. In this project, the ViT-based raw mp4 pipeline produces 768-dimensional visual embeddings, so the model expects 768-dimensional video features, not a 128-dimensional placeholder tensor.",
    )
    add_table(
        doc,
        ["What happened", "Why it happened", "How to explain it"],
        [
            ["ViT extracted (32, 768) embeddings", "That is the correct shape for the face-crop ViT branch.", "This is the real video signal the model should see."],
            ["Dataset returned a (32, 128) placeholder", "The raw demo accidentally told ManifestDataset not to load the extracted video.", "The model received the wrong feature dimension."],
            ["Linear layer multiplication failed", "The model’s video encoder expected 768 input features but got 128.", "This is a code-path bug, not a dataset-label bug."],
        ],
    )
    add_bullets(
        doc,
        [
            "The fix is to ensure the raw demo keeps the extracted ViT embeddings.",
            "In code, the key issue was the dataset wrapper being created with `load_video=False`.",
            "After changing it to `load_video=True`, the `.npy` cache is passed into the model instead of being replaced by zeros.",
        ],
    )
    add_para(
        doc,
        "If the reviewer asks what this means, say: 'The ViT branch was extracting the correct 768-dimensional features, but the demo wrapper accidentally replaced them with a 128-dimensional placeholder tensor. That created a shape mismatch at the video encoder. The fix is to keep the extracted .npy features inside the dataset object so the model sees the correct visual embedding dimension.'",
    )

    doc.add_heading("14. How to Test a New Video That Is Not in the Dataset", level=1)
    add_para(
        doc,
        "Sometimes a reviewer will show you a new mp4 clip that is not already part of MELD. In that case, the standard raw demo script cannot be used completely unchanged, because it expects a manifest row with a sample_id, split, label, transcript, and file paths. The cleanest way to handle this is a temporary demo flow.",
    )
    add_para(
        doc,
        "Worked example: use the MELD test clip `test_dia244_utt14` as the new-video example. In the document folder, a temporary manifest example is provided at `implementation_docments/phase1_unseen_video_temp_manifest_example.csv`. That file shows the exact fields the script needs.",
    )
    add_bullets(
        doc,
        [
            "If the reviewer gives you a new clip, copy it to the pod or local machine first.",
            "If you know the transcript, keep it in a temporary CSV or text note.",
            "If you know the label, use it only for checking correctness; if not, treat the run as prediction-only.",
            "Create a temporary one-row manifest if you want to reuse the existing manifest-driven raw demo script.",
            "Or use a prediction-only wrapper that accepts the raw mp4 directly and prints the model output without requiring a ground-truth row.",
        ],
    )
    add_table(
        doc,
        ["Situation", "What you should do", "Why"],
        [
            ["New clip with known label", "Create a temporary one-row manifest and run the raw demo script.", "This preserves the same code path and allows truth-vs-prediction checking."],
            ["New clip with no label", "Use a prediction-only wrapper or a temporary manifest without presenting it as evaluated accuracy.", "The model can still predict, but you should not claim a measured score."],
            ["New clip with transcript available", "Include the transcript in the temporary row or wrapper input.", "The text branch can still contribute to the fused result."],
            ["New clip without transcript", "Leave the text blank and explain that the run is limited to the available modalities.", "This is the honest way to present the evidence."],
        ],
    )
    add_table(
        doc,
        ["Temporary field", "Example value", "Why it is needed"],
        [
            ["sample_id", "reviewer_video_001", "This is the temporary id the script uses to find the row."],
            ["split", "test", "The demo should treat the clip like a test-time example."],
            ["label", "0", "Use the true label only if you want to check correctness; otherwise explain it as optional."],
            ["transcript", "Ross , foot on the floor or come over no more!", "This lets the text branch still contribute."],
            ["audio_path", "data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia244_utt14.mp4", "The demo loader uses this to recover the audio stream."],
            ["video_path", "data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia244_utt14.mp4", "The same clip is used for visual feature extraction."],
        ],
    )
    add_para(
        doc,
        "Student-level technical detail: the current scripts/predict_phase1_raw_mp4_demo.py starts by loading a manifest row using sample_id. That is why an unseen video needs either a temporary manifest entry or a new wrapper. After that, the same pipeline still applies: extract audio from the mp4, sample frames, build ViT features, load the checkpoint, and produce softmax probabilities.",
    )
    add_para(
        doc,
        "Exact command pattern for an unseen video:",
    )
    add_code(
        doc,
        r"""# 1) Create a temporary manifest row using the example file
# implementation_docments/phase1_unseen_video_temp_manifest_example.csv

# 2) Then run the normal raw demo script on the raw mp4
DEVICE=cuda \
CHECKPOINT=results/paper_aligned_meld_cv/cmt_min/fold_2/best_model.pt \
bash scripts/run_phase1_raw_mp4_demo.sh \
  reviewer_video_001 \
  data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia244_utt14.mp4

# 3) If the reviewer clip is prediction-only, present the output as a live prediction,
# not as an official dataset score.""",
    )
    add_bullets(
        doc,
        [
            "If the new clip is not in MELD, do not claim it is part of the official test score.",
            "If you know the true label, compare it against the prediction as a sanity check.",
            "If you do not know the true label, present the result as a live prediction demo only.",
            "If you want the exact same code path as the dataset examples, create a temporary manifest row first.",
        ],
    )
    add_para(
        doc,
        "If you want to explain the limitation clearly in the viva, say: 'For a new external video, I can still use the same feature-extraction and inference path, but I must either create a temporary manifest row or use a prediction-only wrapper. The prediction is valid, but it is not a dataset score unless I also know the ground-truth label.'",
    )

    doc.add_heading("15. Suggested Talking Order During the Demo", level=1)
    add_numbered(
        doc,
        [
            "Introduce the raw clip and the expected emotion.",
            "Run the baseline raw-mp4 demo and show the prediction.",
            "Point out the confidence and the top-3 alternatives.",
            "Repeat with a known near-miss clip.",
            "Switch to the gated+aux checkpoint if available on Runpod.",
            "Show the metrics summary and confusion matrix.",
            "Conclude by explaining what the model learned and what it still confuses.",
        ],
    )

    doc.add_heading("16. Practical Troubleshooting", level=1)
    add_table(
        doc,
        ["Symptom", "Likely fix", "Where to look"],
        [
            ["Missing checkpoint", "Copy the checkpoint to the pod or point CHECKPOINT to the correct path.", "Wrapper scripts and the results/ directory."],
            ["Missing raw video", "Make sure the mp4 is uploaded or mounted at the same path.", "data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/"],
            ["transformers ImportError", "Install the project requirements on the pod.", "requirements-phase1.txt"],
            ["Low GPU use", "Use DEVICE=cuda and verify the pod actually has a GPU.", "Runpod environment and script env vars."],
            ["Slow demo", "Reuse the cached .npy feature file or run one clip at a time.", "results/phase1_review_demo/raw_mp4_cache/"],
        ],
    )
    add_para(
        doc,
        "One important note: a prediction being wrong does not necessarily mean the raw clip or the code is broken. It may simply mean the model is seeing a hard example, a class imbalance effect, or a frame-sampling limitation. Your job is to explain which of those is most likely for the given clip.",
    )

    doc.add_heading("17. Short Closing Statement", level=1)
    add_para(
        doc,
        "If you need a short closing statement for the review, say: 'This SOP shows the complete Phase 1 demo path from raw mp4 to emotion prediction, plus the metrics and error analysis needed to defend the result. I can explain the wrappers, the inference script, the ViT feature extraction, the checkpoint loading, the softmax probabilities, and the likely causes of wrong predictions. That means I can demonstrate the model clearly and also defend its behavior technically.'",
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
