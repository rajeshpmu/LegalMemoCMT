"""Append speaking notes for the gated-manifest flow and Slides 17/18."""
from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt

ROOT=Path(__file__).resolve().parents[1]
DOCX=ROOT/"implementation_docments/LegalMemoCMT_Phase2_Second_Review_Call_Speaking_Doc.docx"
MARKER="GATED_MANIFEST_SPEAKING_NOTES_V1"

def main():
    doc=Document(DOCX)
    if any(MARKER in p.text for p in doc.paragraphs): print('Already updated'); return
    backup=DOCX.with_name(DOCX.name+'.before_gated_flow_v1.docx')
    if not backup.exists(): shutil.copy2(DOCX,backup)
    doc.add_heading('Gated Manifest Flow: Student Speaking Notes',1)
    p=doc.add_paragraph(); p.add_run(MARKER).font.size=Pt(8)
    doc.add_paragraph('The Mermaid diagram explains that courtroom-affect candidate generation and acceptance are two different operations. The input is emotion_scope_review_200_scope_aware.csv, which already preserves Phase 1 predictions and contains transcript scope, audio-SER, and visual evidence. propose_clancy_courtroom_affect.py reads those fields and produces interpretable candidates such as CALM_COMPOSED or HESITANT_UNCERTAIN. apply_clancy_annotation_acceptance_gate.py then checks confidence thresholds and critical conflicts. This separation makes it possible to improve the heuristic without silently changing the original model output.')
    doc.add_paragraph('The gate is not a psychological truth detector. A row becomes AUTO_ADJUDICATED/SILVER only when the basic-emotion review confidence is at least 0.70, courtroom-affect confidence is at least 0.60, both candidates are resolved, and no critical conflict is recorded. Otherwise the row remains UNRESOLVED/WEAK. The figure is therefore a data-lineage diagram: each arrow identifies a transformation and each branch identifies a review decision.')
    doc.add_heading('Slide 17: SILVER example explanation',2)
    doc.add_paragraph('For DCBWoWhsTpA_turn06801, I explain the result in order. First, Phase 1 sadness at 0.544674 is retained as historical machine evidence. Next, the transcript scope is QUOTED_SPEECH and speaker-emotion evidence is NO, so the emotional content in the quotation is not automatically attributed to the witness. Negative activation is recorded separately and distress corroboration is NO. The candidate layer therefore proposes neutral for basic emotion and CALM_COMPOSED for courtroom presentation. The gate accepts the row because the candidates meet the configured thresholds and no critical conflict is present. The final machine-assisted fields are neutral and CALM_COMPOSED, with AUTO_ADJUDICATED/SILVER status. This is suitable for controlled weak-supervision experiments, not a human gold label.')
    doc.add_heading('Slide 18: WEAK example explanation',2)
    doc.add_paragraph('For DCBWoWhsTpA_turn06570, I explain why the row is not automatically finalized even though the affect candidate has confidence 0.75. The basic candidate remains the Phase 1 disgust result with confidence 0.47, which is below the 0.70 basic threshold. The transcript contains hesitation and epistemic qualification, so HESITANT_UNCERTAIN is more appropriate than generic TENSE, but the speaker-emotion scope is unresolved. The gate therefore sets final_basic_emotion=UNRESOLVED and annotation_tier=WEAK while retaining HESITANT_UNCERTAIN as a review candidate. The [snorts] marker is context only; it is not evidence of disgust.')
    doc.add_heading('What to say about the generated CSV',2)
    doc.add_paragraph('The CSV is not just a list of labels. It is a provenance record containing the original Phase 1 prediction, candidate evidence, gate thresholds, conflict status, final machine-assisted fields, and acceptance reason. This lets me show exactly how a final field was derived and lets another researcher rerun the same two scripts. Before human annotation, I inspect the 75 unresolved rows and 7 critical-conflict rows; before scaling, I manually validate a sample of the 125 SILVER rows.')
    doc.save(DOCX); print('Updated',DOCX)

if __name__=='__main__': main()
