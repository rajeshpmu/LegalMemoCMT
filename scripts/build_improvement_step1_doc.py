from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "implementation_docments" / "LegalMemoCMT_Improvement_Step1_Class_Balanced_Focal_Loss_Guide.docx"

FOLD2_METRICS = ROOT / "results" / "paper_aligned_meld_cv" / "cmt_min" / "fold_2" / "metrics.json"
FOLD4_METRICS = ROOT / "results" / "paper_aligned_meld_cv" / "cmt_min" / "fold_4" / "metrics.json"
CREMA_METRICS = ROOT / "results" / "paper_aligned_crema_d" / "cmt_min" / "metrics.json"


def style_document(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)


def add_para(doc: Document, text: str, *, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def read_metric(path: Path, key: str) -> float:
    import json

    return float(json.loads(path.read_text(encoding="utf-8"))[key])


def build_doc() -> Document:
    fold2_acc = read_metric(FOLD2_METRICS, "accuracy")
    fold2_uf1 = read_metric(FOLD2_METRICS, "unweighted_accuracy")
    fold2_mf1 = read_metric(FOLD2_METRICS, "macro_f1")
    fold2_wf1 = read_metric(FOLD2_METRICS, "weighted_f1")

    fold4_acc = read_metric(FOLD4_METRICS, "accuracy")
    fold4_uf1 = read_metric(FOLD4_METRICS, "unweighted_accuracy")
    fold4_mf1 = read_metric(FOLD4_METRICS, "macro_f1")
    fold4_wf1 = read_metric(FOLD4_METRICS, "weighted_f1")

    crema_acc = read_metric(CREMA_METRICS, "accuracy")
    crema_uf1 = read_metric(CREMA_METRICS, "unweighted_accuracy")
    crema_mf1 = read_metric(CREMA_METRICS, "macro_f1")
    crema_wf1 = read_metric(CREMA_METRICS, "weighted_f1")

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    run = title.add_run("LegalMemoCMT Improvement Step 1: Class-Balanced Focal Loss Guide")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(
        "A student-level explanation of the new improvement scripts, how they differ from the paper-aligned runs, and what we expect to change in the MELD and CREMA-D error patterns."
    )
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)

    add_para(
        doc,
        "This document describes the first improvement step that follows the guidance-call recommendation order: replace the current weighted cross-entropy objective with class-balanced focal loss, then re-run selected MELD folds and the CREMA-D cross-validation track. The goal is not to change the model architecture first. The goal is to change the training signal so that the model stops being rewarded for over-predicting the dominant classes.",
    )
    add_para(
        doc,
        "The new scripts are separate from the existing paper-aligned scripts. That means the previous paper-aligned results remain intact, and the improvement experiments can be compared cleanly against them.",
    )

    doc.add_heading("1. Why This Change Comes First", level=1)
    add_para(
        doc,
        "The current error analysis shows a clear pattern: MELD is still too neutral-heavy, while CREMA-D is still too close to class collapse. Those are exactly the kinds of errors that a better loss function should improve before we attempt larger architectural changes such as supervised contrastive learning or dialogue context modeling.",
    )
    add_bullets(
        doc,
        [
            "MELD Fold 2 shows the strongest weighted performance, but still confuses neutral with joy, anger, and surprise.",
            "MELD Fold 4 shows the strongest class-balanced behavior, which means the model can improve on minority classes when the metric rewards it properly.",
            "CREMA-D is still near collapse, which usually means the optimization signal is not strong enough for minority class separation.",
        ],
    )
    add_para(
        doc,
        "Class-balanced focal loss is the smallest change that directly attacks all three problems at once. It is therefore the most defensible first improvement.",
    )

    doc.add_heading("2. What the Old Scripts Did", level=1)
    add_para(
        doc,
        "The previous paper-aligned workflow used weighted cross-entropy and the paper-style text-plus-audio path. For MELD, the script trained five folds and evaluated each fold on the held-out MELD test split. For CREMA-D, the analysis script exported predictions from the trained checkpoint and generated confusion analysis files.",
    )
    add_bullets(
        doc,
        [
            "scripts/run_paper_aligned_meld_cv.sh: built MELD raw manifests, split train/dev dialogues into 5 folds, trained each fold with weighted cross-entropy, and evaluated each fold on MELD test.",
            "scripts/run_paper_aligned_crema_d_analysis.sh: exported predictions from the already trained CREMA-D checkpoint and generated confusion matrices and top-confusion tables.",
            "Those scripts gave us the baseline paper-aligned results that we are now trying to improve.",
        ],
    )

    doc.add_heading("3. What the New Improvement Scripts Change", level=1)
    add_table(
        doc,
        ["Aspect", "Paper-aligned scripts", "New improvement scripts"],
        [
            ["Loss", "Weighted cross-entropy", "Class-balanced focal loss"],
            ["MELD folds", "All 5 folds", "Only Fold 2 and Fold 4"],
            ["CREMA-D", "Held-out analysis of the existing trained model", "Full 5-fold CV rerun with focal loss"],
            ["Purpose", "Paper-aligned baseline", "First targeted improvement step"],
            ["Output directories", "results/paper_aligned_*", "results/improvement/class_balanced_focal/*"],
        ],
    )
    add_para(
        doc,
        "The new scripts are designed to isolate the effect of the loss function. By keeping the model architecture, the pretrained text/audio path, and the min pooling unchanged, we can see whether the error pattern changes because of the new objective rather than because of unrelated implementation drift. By default, the improvement scripts reuse the existing MELD and CREMA-D fold CSVs that were already built for the paper-aligned runs, so the comparison stays controlled and the split logic does not change unless you intentionally rebuild it.",
    )
    add_para(
        doc,
        "After the first focal-loss run was interrupted, the next recommended path is the warm-start version. That path starts from the already trained weighted-CE checkpoints for the selected MELD folds, lowers the learning rate, switches the loss to focal loss, and fine-tunes only a few more epochs. This is a refinement experiment rather than a from-scratch experiment, so it tests whether focal loss can improve an already reasonable solution instead of trying to build the whole decision boundary under focal loss from zero.",
    )

    doc.add_heading("4. How the New MELD Improvement Script Works", level=1)
    add_para(
        doc,
        "The MELD improvement script is targeted rather than exhaustive. It reruns only Fold 2 and Fold 4 because those two folds are the most informative for the next-step discussion. It reuses the existing fold CSVs from `data/manifests/meld_cv`, which were already built during the paper-aligned MELD run, so the only experiment change is the loss function.",
    )
    add_bullets(
        doc,
        [
            "Fold 2 is the strongest weighted/overall anchor, so it is the best place to see whether the new loss preserves the best paper-comparison behavior.",
            "Fold 4 is the strongest class-balanced anchor, so it is the best place to see whether the new loss improves minority-class learning without hurting balance.",
        ],
    )
    add_para(
        doc,
        "The script does not rebuild the fold splits unless you deliberately uncomment the rebuild lines. In the default path it simply points to the existing `data/manifests/meld_cv` CSVs, then trains each selected fold using `--loss-type focal` instead of `--loss-type weighted-ce`. The evaluation step still uses the MELD test split, so the final result is directly comparable to the previous paper-aligned test analysis.",
    )
    add_numbered(
        doc,
        [
            "Build a raw MELD manifest from the annotation CSVs and clip paths.",
            "Rebuild the 5 MELD dialogue folds from that manifest so the selected folds stay identical in structure.",
            "Train Fold 2 with class-balanced focal loss.",
            "Train Fold 4 with class-balanced focal loss.",
            "Evaluate each trained checkpoint on the held-out MELD test split.",
            "Export predictions and generate confusion/error analysis for each selected fold.",
        ],
    )

    doc.add_heading("5. How the New CREMA-D Improvement Script Works", level=1)
    add_para(
        doc,
        "The CREMA-D improvement script reruns the full speaker-independent 5-fold CV track, but with focal loss instead of weighted cross-entropy. It reuses the existing `data/manifests/crema_d_cv` split files by default, so again the only experiment variable is the loss. This is important because CREMA-D is currently the dataset with the clearest collapse problem, so it is the best place to test whether the new objective stabilizes training.",
    )
    add_bullets(
        doc,
        [
            "The script rebuilds the CREMA-D CV folds so the speaker-independent separation stays intact.",
            "Each fold trains on the remaining speakers and validates on the held-out speakers.",
            "Each fold is trained with `--loss-type focal` and `--focal-gamma 2.0`.",
            "After training, the fold is evaluated on its validation split and then aggregated across folds.",
        ],
    )
    add_para(
        doc,
        "Unlike the old CREMA-D analysis script, this new script is not just an inspection wrapper. It is a full re-training experiment. That makes it suitable for answering the question: does a better imbalance-aware loss improve the actual speaker-independent CV behavior?",
    )

    doc.add_heading("6. How Training, Evaluation, and Analysis Fit Together", level=1)
    add_para(
        doc,
        "The workflow is the same in spirit as the earlier scripts, but the loss function and the run scope are different. Student-wise, it helps to think of the pipeline as three separate layers.",
    )
    add_table(
        doc,
        ["Stage", "What it does", "New improvement meaning"],
        [
            ["Training", "Fits the model weights on train folds", "Now uses focal loss to focus on hard examples"],
            ["Evaluation", "Runs the trained checkpoint on held-out data", "Still reports accuracy, macro F1, weighted F1, and class balance behavior"],
            ["Analysis", "Exports predictions and confusion matrices", "Lets us see whether neutral collapse or class collapse has reduced"],
        ],
    )
    add_para(
        doc,
        "The key point is that the analysis files are not the training themselves. They are the evidence produced after training. That evidence is what we compare against the old paper-aligned runs.",
    )

    doc.add_heading("7. What Focal Loss Is Doing", level=1)
    add_para(
        doc,
        "Focal loss starts from cross-entropy but multiplies each sample loss by a factor that becomes smaller when the model is already confident and correct. In other words, easy examples contribute less and hard examples contribute more. This is useful when the dataset is imbalanced or when the model keeps learning the easy dominant class too well.",
    )
    add_bullets(
        doc,
        [
            "If the model already predicts a sample correctly with high confidence, the loss contribution shrinks.",
            "If the model is uncertain or wrong, the loss contribution stays large.",
            "Class weights can still be added so minority emotions matter more than majority emotions.",
            "For this project, that should reduce neutral over-prediction and improve minority-class recall.",
        ],
    )
    add_para(
        doc,
        "In the code, the trainer already supports focal loss. The new scripts simply switch the loss type from weighted-ce to focal so we can isolate the effect of the new objective without rewriting the training loop.",
    )
    add_para(
        doc,
        "For the next manual step, the most important path is the resume-capable Fold 2 run: it starts from the weighted-CE Fold 2 checkpoint, keeps the same paper-aligned architecture and min pooling, and then continues training with focal loss for a few extra epochs at a smaller learning rate.",
    )

    doc.add_heading("8. Why Only MELD Folds 2 and 4?", level=1)
    add_para(
        doc,
        "This is a targeted ablation, not a full rerun. Fold 2 and Fold 4 were chosen because they summarize the two most informative sides of the MELD story.",
    )
    add_table(
        doc,
        ["Fold", "Why it matters", "Current baseline signal"],
        [
            ["Fold 2", "Best weighted anchor", f"Acc {fold2_acc:.4f}, W-F1 {fold2_wf1:.4f}, macro F1 {fold2_mf1:.4f}"],
            ["Fold 4", "Best balanced-class anchor", f"Acc {fold4_acc:.4f}, UW-Acc {fold4_uf1:.4f}, macro F1 {fold4_mf1:.4f}"],
        ],
    )
    add_para(
        doc,
        "By rerunning only those two folds with focal loss, we can check whether the new training objective improves the model where it matters most: one fold that already performs well overall, and one fold that is strongest on class balance. If the new loss helps both, that is a good sign that the next stage of work is worthwhile.",
    )

    doc.add_heading("9. Why Full CREMA-D CV Is Repeated", level=1)
    add_para(
        doc,
        "CREMA-D is not being treated as a single-checkpoint analysis here. The improvement script reruns the full 5-fold speaker-independent CV because that is the cleanest way to see whether focal loss improves generalization across speakers.",
    )
    add_para(
        doc,
        "The current CREMA-D baseline is very weak, so a single-fold result would be too easy to misread. A full CV rerun gives a more trustworthy answer about whether the new objective helps the model avoid collapse and spread predictions more evenly across classes.",
    )

    doc.add_heading("10. What We Expect to See", level=1)
    add_para(
        doc,
        "The expected result is not that every metric automatically jumps. The expected result is that the confusion pattern changes in the right direction.",
    )
    add_bullets(
        doc,
        [
            "MELD Fold 2 should keep much of its overall strength while reducing neutral-heavy confusions.",
            "MELD Fold 4 should preserve balanced-class behavior and ideally improve minority-class separation.",
            "CREMA-D should move away from near-collapse and show a wider distribution of correct predictions across classes.",
            "Macro F1 and UW-Acc should be watched carefully, because they reveal minority-class improvement more clearly than accuracy alone.",
        ],
    )
    add_para(
        doc,
        "A good improvement is one where the confusion matrix becomes less lopsided, even if the top-line accuracy changes only a little. That is often the first sign that the model is actually learning the right boundaries instead of just exploiting the easiest class.",
    )

    doc.add_heading("11. How This Differs from the Paper-Aligned Baseline", level=1)
    add_bullets(
        doc,
        [
            "The old paper-aligned scripts used weighted cross-entropy; the new ones use focal loss.",
            "The old MELD script ran all 5 folds; the new MELD improvement script runs only the two folds that are most informative for the next-step decision.",
            "The old CREMA-D analysis script only inspected a trained checkpoint; the new CREMA-D improvement script retrains the model across 5 folds.",
            "The old workflow was about reporting the paper-aligned baseline; the new workflow is about improving the baseline in a controlled way.",
            "The warm-start focal workflow continues from the weighted-CE fold checkpoints and uses a smaller learning rate for a short refinement run.",
            "The resume-capable Fold 2 path goes one step further by explicitly loading the weighted-CE Fold 2 checkpoint before continuing with focal loss.",
        ],
    )

    doc.add_heading("12. What to Look For After You Run the New Scripts", level=1)
    add_numbered(
        doc,
        [
            "Check whether the MELD Fold 2 confusion matrix still over-predicts neutral or whether the confusion is more evenly distributed.",
            "Check whether Fold 4 keeps its stronger balanced-class behavior while improving overall confidence calibration.",
            "Check whether CREMA-D predictions stop collapsing into one class and begin spreading across the label set more sensibly.",
            "Compare macro F1 and UW-Acc against the earlier paper-aligned results, not just accuracy.",
            "Use the error tables to identify whether the confusing pairs have become smaller or shifted to different emotion pairs.",
        ],
    )
    add_para(
        doc,
        "If those signs move in the right direction, then the next recommendation in the guidance plan becomes more justified. If they do not, then the next stage should likely be supervised contrastive learning or a stronger imbalance-aware objective such as LDAM-DRW.",
    )

    doc.add_heading("13. Student Summary", level=1)
    add_para(
        doc,
        "The main idea is simple: do not change everything at once. Keep the paper-aligned architecture, change the loss first, and measure whether the confusion matrices improve in the exact places they are currently weak. That gives you a clean before-and-after story that is easy to explain in a guidance call.",
    )
    add_para(
        doc,
        "In short, these new scripts are a controlled experiment. They are meant to answer a narrow but important question: if we keep the same model and the same paper-aligned fusion path, does class-balanced focal loss improve the MELD and CREMA-D class confusion behavior enough to justify the next step?",
    )

    doc.add_heading("14. Current Evidence to Keep in Mind", level=1)
    add_table(
        doc,
        ["Source", "Acc / W-Acc", "UW-Acc", "Macro F1", "Weighted F1", "Interpretation"],
        [
            ["MELD Fold 2", f"{fold2_acc:.4f}", f"{fold2_uf1:.4f}", f"{fold2_mf1:.4f}", f"{fold2_wf1:.4f}", "Best weighted anchor; still neutral-heavy"],
            ["MELD Fold 4", f"{fold4_acc:.4f}", f"{fold4_uf1:.4f}", f"{fold4_mf1:.4f}", f"{fold4_wf1:.4f}", "Best balanced-class anchor"],
            ["CREMA-D", f"{crema_acc:.4f}", f"{crema_uf1:.4f}", f"{crema_mf1:.4f}", f"{crema_wf1:.4f}", "Near-collapse into one class"],
        ],
    )
    add_para(
        doc,
        "These are the reference points the new focal-loss experiments should be compared against. If the new scripts improve these patterns, the next recommendation in the guidance call becomes stronger. If they do not, that is also useful evidence, because it tells us the problem is deeper than the loss function alone.",
    )

    doc.add_heading("14.1 Resume-Capable Fold 2 Manual Command Sequence", level=2)
    add_para(
        doc,
        "If you want to start with the most targeted next step, use the resume-capable Fold 2 path first. This is the safest place to begin because it directly tests whether focal loss can refine the already learned weighted-CE Fold 2 decision boundary instead of rebuilding the model from scratch.",
    )
    add_numbered(
        doc,
        [
            "bash scripts/run_resume_warmstart_focal_meld_fold2.sh",
            "bash scripts/analyze_resume_warmstart_focal_meld_fold2.sh",
        ],
    )
    add_para(
        doc,
        "After those two commands finish, you can inspect the Fold 2 confusion matrix and decide whether to extend the same resume-style idea to Fold 4 or move on to the CREMA-D warm-start run. If you want the older from-scratch warm-start workflow instead, the suite wrapper is still available, but it is no longer the first choice for the manual next step.",
    )

    doc.add_heading("15. Exact Manual Command Sequence", level=1)
    add_para(
        doc,
        "For the current manual next step, run the resume-capable Fold 2 path first. That gives you the cleanest before-and-after comparison against the weighted-CE baseline because it keeps the architecture and split fixed while changing only the training continuation strategy.",
    )
    add_numbered(
        doc,
        [
            "bash scripts/run_resume_warmstart_focal_meld_fold2.sh",
            "bash scripts/analyze_resume_warmstart_focal_meld_fold2.sh",
        ],
    )
    add_para(
        doc,
        "If the Fold 2 resume result looks promising, the next follow-up would be to mirror the same resume-style setup for Fold 4 and then revisit CREMA-D. If you prefer a single chained command for the older from-scratch improvement workflow, you can still use `bash scripts/run_improvement_class_balanced_focal_suite.sh`, but that is separate from the resume-capable Fold 2 path you should run first now.",
    )

    doc.save(OUT)
    return doc


def main() -> None:
    build_doc()
    print(OUT)


if __name__ == "__main__":
    main()
