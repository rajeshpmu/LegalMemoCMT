from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "implementation_docments/LegalMemoCMT_Phase2_Corpus_Build_Code_Level_Student_Guide.docx"


def style_document(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix); r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)


def code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = "Menlo"; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(35, 35, 35)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, value in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = value; cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.color.rgb = RGBColor(20, 48, 87)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths): row.cells[i].width = Inches(width)


def add_script_section(doc: Document, script: str, purpose: str, input_text: str, output_text: str, functions: list[str], explanation: str, command: str | None = None) -> None:
    heading(doc, script, level=2)
    para(doc, purpose, "Purpose: ")
    para(doc, input_text, "Inputs: ")
    para(doc, output_text, "Outputs: ")
    para(doc, explanation, "Student explanation: ")
    if command:
        para(doc, "Manual command:", "Manual command:")
        code(doc, command)
    para(doc, "Important functions and code responsibilities:", "Important functions and code responsibilities:")
    for item in functions: bullet(doc, item)


def build() -> None:
    doc = Document(); style_document(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("LegalMemoCMT Phase 2\nCorpus Build: Code-Level Student Guide")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(20, 48, 87)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Execution order, Python responsibilities, outputs, and quality gates").italic = True
    para(doc, "This guide explains how the current Phase 2 corpus is constructed from public source material. It is written so a student can explain the implementation to a mentor without claiming that every generated row is automatically training-ready.")
    para(doc, "The current primary benchmark is the Lindsay Clancy courtroom branch. IRMCT/ICTY is retained as a secondary tribunal bootstrap branch. Tupac/Keffe D and Indian-SIM are acquisition or adaptation branches until they have verified media, transcript evidence, and processed utterance rows.")

    heading(doc, "1. The Main Idea")
    para(doc, "The corpus is built as a chain of traceable transformations. A public URL becomes a source record, a source record becomes transcript or subtitle cues, cues become turn rows, turn timestamps become MP4 and WAV clips, and quality filters decide which rows can be considered for training.")
    table(doc, ["Layer", "Question answered", "Main artifact"], [
        ["Source", "Where did the recording or transcript come from?", "clancy_source_manifest.csv / UCR manifests"],
        ["Turn", "Which piece of the transcript is one speaker turn?", "clancy_turn_manifest.csv"],
        ["Media", "Which exact video and audio interval represents that turn?", "turn_clips/video and turn_clips/audio"],
        ["Quality", "Should this row be used, reviewed, or rejected?", "filtered_rows, EDA, rejection manifest"],
        ["Dataset", "How are rows assigned without leakage?", "train.csv, dev.csv, test.csv"],
    ], [1.0, 3.3, 2.8])

    heading(doc, "2. Exact Execution Order")
    para(doc, "The order below is the safest rerunnable order for the Clancy branch. The commands use the repository-local virtual environment so that yt-dlp, python-docx, faster-whisper-related packages, and ffmpeg support are consistent.")
    table(doc, ["Step", "Run", "Purpose", "Main output"], [
        ["1", "run_build_clancy_source_manifest.sh", "Probe the URL shortlist and record source metadata.", "clancy_source_manifest.csv"],
        ["2", "run_build_clancy_corpus.sh or existing subtitle/media preparation", "Create or register the weak training rows from downloaded source material.", "clancy_training_manifest_weak.csv"],
        ["3", "run_build_clancy_weak_labels.sh", "Attach provisional text-based emotion labels and label provenance.", "clancy_training_manifest_weak.csv / weak label outputs"],
        ["4", "run_build_clancy_turn_manifest.sh", "Group subtitle cues into turn-level records and apply persistent exclusions.", "clancy_turn_manifest.csv"],
        ["5", "run_build_clancy_turn_clips.sh", "Use timestamps and source offsets to extract MP4 and WAV clips.", "clancy_turn_manifest_clipped.csv and turn_clips/"],
        ["6", "create_clancy_post_rejection_eda_report.py", "Remove persistent rejected IDs from the analysis population and describe what remains.", "post-rejection CSV, JSON, DOCX"],
        ["7", "run_build_clancy_duration_outlier_report.sh", "Flag short and long duration candidates for review.", "clancy_duration_outliers.csv and JSON"],
        ["8", "run_filter_legalmeld_rows_by_use.sh", "Create usable, review, reject, confidence, audio, video, and split subsets.", "filtered_rows/*.csv"],
        ["9", "run_build_clancy_dataset_split.sh", "Assign complete source groups to train/dev/test.", "clancy_dataset_manifest.csv and train/dev/test.csv"],
        ["10", "run_build_clancy_dataset_validation.sh", "Check paths, required fields, split values, and leakage.", "validation summary and issues CSV"],
        ["11", "run_build_clancy_training_readiness.sh", "Check whether labels, paths, splits, and required metadata are present.", "training readiness summary and issues CSV"],
    ], [0.4, 2.5, 3.4, 2.5])
    para(doc, "Not every stage must download again. Source-manifest probing records metadata; turn-manifest rebuilding does not download media; turn clipping creates derived clips; and validation only reads the produced dataset. Using --skip-existing or the script defaults prevents unnecessary repeated work where supported.")

    heading(doc, "3. Stage-by-Stage Code Explanation")
    add_script_section(doc, "phase2/build_clancy_source_manifest.py", "Builds a structured source inventory from the URL shortlist.", "data/clancy_urls.txt, yt-dlp, and optionally browser cookies for public source probing.", "data/processed/phase2/clancy/clancy_source_manifest.csv and reports/phase2/clancy_source_manifest_summary.json.", [
        "_read_urls() removes blank lines and duplicate URLs so the same source is not probed twice.",
        "_probe_title() calls yt-dlp metadata extraction without treating the output as proof that testimony is suitable.",
        "The CSV writer records URL, title, video ID, source category, and media/subtitle fields for later stages.",
        "The summary JSON records how many URLs were inspected and which sources could be resolved.",
    ], "This stage answers where the source came from. It does not decide whether a courtroom segment is useful, whether a witness is present, or whether an emotion label is correct.", "PYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_source_manifest.sh")

    add_script_section(doc, "phase2/build_clancy_corpus.py and phase2/run_build_clancy_corpus.sh", "Prepares the source-level media and subtitle inventory used to create weak training rows.", "The source URL list, downloaded media/subtitle files, and the configured Clancy output directory.", "Source and weak-training manifests under data/processed/phase2/clancy/.", [
        "Argument parsing keeps the input URL file and output paths configurable.",
        "The downloader uses stable source IDs and filenames so reruns can identify existing media.",
        "Media and subtitle paths are written into rows rather than inferred later from filenames.",
    ], "This stage produces source evidence and registration records. It should not be described as final alignment because subtitles can still contain cue boundaries that do not equal clean speaker turns.", "PYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_corpus.sh")

    add_script_section(doc, "phase2/build_clancy_weak_labels.py and run_build_clancy_weak_labels.sh", "Adds provisional labels used for exploration and label-distribution EDA.", "Weak training or utterance rows containing text.", "Rows with emotion_label, emotion_label_source, and emotion_label_confidence.", [
        "Text normalization makes keyword checks case-insensitive and robust to punctuation.",
        "Keyword or phrase rules map text patterns to provisional labels such as neutral, sadness, anger, fear, stress, or confidence.",
        "The source and confidence fields make it clear that these labels are heuristic, not expert ground truth.",
    ], "The label generator is useful for discovering imbalance and selecting review samples. It cannot reliably infer facial emotion from text alone, so the current labels must not be presented as final supervised truth.", "PYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_weak_labels.sh")

    add_script_section(doc, "phase2/build_clancy_turn_manifest.py", "Converts subtitle/utterance rows into consolidated turn-level rows.", "clancy_training_manifest_weak.csv plus the persistent rejection manifest.", "data/processed/phase2/clancy/clancy_turn_manifest.csv and issue/summary reports.", [
        "_time_to_seconds() converts subtitle time strings into numeric seconds.",
        "_split_turn_segments() separates explicit markers or cue fragments where a single source row contains multiple pieces.",
        "_allocate_segment_spans() assigns timestamp spans to split text pieces while retaining traceability to the original cue.",
        "_consolidate_turns() merges adjacent compatible pieces from the same source group when the text and sequence indicate continuity.",
        "_apply_rejections() reads clancy_turn_rejection_manifest.csv and removes known lunch breaks, news segments, breaks, and other manually excluded turns.",
    ], "The important design decision is that rejection is persistent. A manually rejected turn is not silently regenerated into the next manifest. The original clipped file remains available for audit, while future manifests omit its turn ID.", "PYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_turn_manifest.sh")

    add_script_section(doc, "phase2/build_clancy_turn_clips.py", "Creates one derived video clip and one derived WAV for each turn row.", "clancy_turn_manifest.csv, raw MP4/WAV paths, and clancy_source_offsets.csv.", "turn_clips/video/<youtube_id>/<turn_id>.mp4, turn_clips/audio/<youtube_id>/<turn_id>.wav, and clancy_turn_manifest_clipped.csv.", [
        "_load_source_offsets() reads an explicit offset per source video.",
        "_source_offset() applies the explicit offset or the documented default of 0 seconds when the source is unlisted.",
        "_run_ffmpeg() passes start time and duration to ffmpeg. The same interval is used for video and audio.",
        "The main loop adds optional start/end padding, clamps times, records failures, and writes the exact derived paths.",
        "The output records source_offset_seconds and source_offset_status so a clip can be audited back to the raw recording.",
    ], "This is the core modality-construction stage. A matching filename is not enough: the manifest must show the source path, timestamp interval, offset rule, and both derived paths.", "PYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_turn_clips.sh")

    add_script_section(doc, "phase2/create_clancy_post_rejection_eda_report.py", "Creates the current analysis population after applying persistent manual exclusions.", "clancy_turn_manifest_clipped.csv and clancy_turn_rejection_manifest.csv.", "clancy_turn_manifest_post_rejection.csv, clancy_turn_manifest_post_rejection_eda.json, and the post-rejection EDA DOCX.", [
        "The rejection IDs are loaded into a set for efficient membership checks.",
        "Rows whose turn_id occurs in that set are excluded from the filtered analysis CSV.",
        "Duration statistics, source coverage, split counts, emotion-label counts, and long-turn examples are calculated from the filtered rows.",
        "The report distinguishes candidate duration from validated training duration.",
    ], "This report does not delete media. It creates a reproducible view of what remains after human review. That distinction preserves auditability and allows a different selection policy later.", "./.venv/bin/python phase2/create_clancy_post_rejection_eda_report.py")

    add_script_section(doc, "phase2/build_clancy_duration_outlier_report.py", "Finds duration values needing manual review before MELD-style training selection.", "The post-rejection turn manifest.", "clancy_duration_outliers.csv and clancy_duration_outlier_summary.json.", [
        "The script calculates quartiles and an upper IQR fence.",
        "Rows below 0.8 seconds are short-fragment candidates.",
        "Rows over 30 seconds are review candidates because a MELD-style sample is normally an utterance-sized segment, not a long courtroom block.",
        "The report flags candidates; it does not automatically reject every statistical outlier.",
    ], "A long turn can be valid testimony, but it may contain multiple sentences or multiple speaker turns. The correct action can be splitting, manual review, or rejection depending on the content.", "PYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_duration_outlier_report.sh")

    add_script_section(doc, "phase2/filter_legalmeld_rows_by_use.py", "Creates operational subsets so training, review, and rejection decisions are explicit.", "Validated LegalMELD metadata CSV.", "filtered_rows/legalmeld_rows_usable.csv, review.csv, reject.csv, confidence subsets, and media-valid subsets.", [
        "classify_row() reads split, quality tier, alignment confidence, manual-review flags, and media-validation fields.",
        "reason_for_row() records why a row was placed in a category instead of only writing a boolean.",
        "build_rows() allows one row to occur in more than one analytical category, such as usable and audio_valid.",
    ], "The usable file is a selection view, not a new source of truth. A row is only useful if its transcript, audio, video, alignment, duration, and split evidence agree.", "PYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_filter_legalmeld_rows_by_use.sh")

    add_script_section(doc, "phase2/build_clancy_dataset_split.py", "Creates leakage-aware train, development, and test partitions.", "clancy_utterance_manifest.csv or the selected turn-level manifest plus the persistent exclusion list.", "clancy_dataset_manifest.csv, train.csv, dev.csv, test.csv, and split summary JSON.", [
        "_group_rows() groups rows by the selected group column, normally youtube_id for source-level separation.",
        "_split_groups() assigns complete groups to train/dev/test according to target ratios rather than splitting rows randomly.",
        "The exclusion loader prevents persistent rejected IDs from entering the selected dataset.",
        "Each output row retains split_group_id and split_strategy for later audit.",
    ], "If adjacent utterances from the same source recording appear in both training and test, the model can learn the courtroom background rather than general emotion cues. Group splitting reduces this leakage risk.", "PYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_dataset_split.sh")

    add_script_section(doc, "phase2/build_clancy_dataset_validation.py and build_clancy_training_readiness.py", "Check whether the selected manifest is structurally and operationally ready for model work.", "The split-aware dataset manifest and the weak-label training manifest.", "Validation/readiness summary JSON files and issue CSV files.", [
        "Validation checks required IDs, split values, non-empty paths, and whether a source group occurs across multiple splits.",
        "Training readiness checks label presence, confidence, source grouping, and required media/text fields.",
        "Issue CSVs preserve row-level failures for manual correction instead of hiding them in aggregate counts.",
    ], "A PASS from a structural validator does not prove that the emotion label is correct or that the witness face is visible. It only proves that the row satisfies the checks implemented by that validator.", "PYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_dataset_validation.sh")

    heading(doc, "4. Tribunal Branch: Same Principle, Different Source Layer")
    para(doc, "The tribunal branch uses a different source-discovery layer but the same final design: verified metadata, hearing rows, witness rows, transcript/video pairing, timestamp alignment, clip extraction, validation, and leakage-aware splits.")
    table(doc, ["Tribunal stage", "Script / module", "Reason it exists"], [
        ["Case verification", "run_enrich_case_ledger_from_ucr_site.sh and build_ucr_case_inventory.py", "Reject placeholder or generic UCR pages and retain actual case-specific records."],
        ["Hearing construction", "build_hearing_manifest.py and hearing_witness_manifest_utils.py", "Create one row per real hearing or record and pair transcript/video evidence."],
        ["Witness construction", "build_witness_manifest.py and resolve_witnesses_from_transcripts.py", "Extract public names or protected codes without deanonymization."],
        ["Utterance export", "build_legalmeld_dataset.py", "Transcribe/align, extract MP4/WAV clips, and write LegalMELD metadata."],
        ["Validation", "run_validate_trimodal_corpus.sh and trimodal_validation_utils.py", "Check media playability, duration, pairing, and traceability."],
    ], [1.3, 3.0, 3.8])
    para(doc, "The tribunal output is currently a bootstrap pilot. Its existence demonstrates the method; it should not be described as equivalent in scale or quality to the primary Clancy branch without checking the current summary reports.")

    heading(doc, "5. What the Important Fields Mean")
    table(doc, ["Field", "Student-level meaning", "Why it matters"], [
        ["turn_id", "Stable identifier for one derived turn row.", "Links the manifest, MP4, WAV, rejection list, and review notes."],
        ["turn_text / utterance_text", "Text intended to represent the turn.", "The language modality and alignment target."],
        ["turn_start_time / turn_end_time", "Original subtitle or cue boundaries.", "The interval passed through offset and padding logic."],
        ["source_offset_seconds", "Correction from source timeline to raw media timeline.", "Prevents clips being cut at the wrong place when a livestream starts late."],
        ["clip_video_path / clip_audio_path", "Derived MP4 and WAV files.", "The actual multimodal inputs for a row."],
        ["emotion_label", "Current provisional target label.", "Used for EDA and pilot experiments; requires stronger annotation for claims."],
        ["quality_tier / manual_review_required", "Decision about confidence and review.", "Prevents weak rows entering the first training run silently."],
        ["split / split_group_id", "Dataset partition and grouping key.", "Controls leakage and makes evaluation more credible."],
    ], [1.7, 3.3, 3.1])

    heading(doc, "6. What Is Actually Training-Ready?")
    para(doc, "The pipeline can create a technically complete row, but technical completeness is not the same as scientific readiness. For a first controlled fine-tuning pilot, select rows that have valid text, matching audio and video, acceptable duration, no persistent rejection, no unresolved overlap, acceptable alignment confidence, and a label that is either manually verified or explicitly treated as weak supervision.")
    bullet(doc, "Start with the Clancy 0.8–30 second candidate window, but inspect examples from every source video.")
    bullet(doc, "Treat rows over 30 seconds as split/review candidates, not automatic training rows.")
    bullet(doc, "Treat rows below 0.8 seconds as a separate short-fragment review group.")
    bullet(doc, "Do not call heuristic emotion labels ground truth. Report the source and confidence of every label.")
    bullet(doc, "Use the group-based train/dev/test partitions and record leakage results in the experiment report.")

    heading(doc, "7. Rerun and Reproducibility Rules")
    code(doc, "cd /Users/rajeshpmu/Desktop/LegalMemoCMT\nPYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_source_manifest.sh\nPYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_turn_manifest.sh\nPYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_turn_clips.sh\n./.venv/bin/python phase2/create_clancy_post_rejection_eda_report.py\nPYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_duration_outlier_report.sh\nPYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_filter_legalmeld_rows_by_use.sh\nPYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_dataset_split.sh\nPYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_dataset_validation.sh\nPYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_training_readiness.sh")
    para(doc, "Before a full rerun, inspect the input and output paths in each wrapper. The Clancy scripts use persistent rejection and source-offset manifests, so those files are part of reproducibility. If a manual rejection or offset changes, regenerate the downstream manifest and reports rather than editing only the final CSV.")

    heading(doc, "8. Current Limitations to Explain Honestly")
    bullet(doc, "The Clancy emotion labels are heavily imbalanced and largely heuristic; they are useful for pipeline testing but not sufficient as expert emotion ground truth.")
    bullet(doc, "A valid MP4 proves that the file can be decoded, not that the speaking witness is visible or that the transcript text matches the speech.")
    bullet(doc, "The duration window is a selection rule, not proof that each turn is a semantically complete utterance.")
    bullet(doc, "Tupac/Keffe D and Indian-SIM require source acquisition and verification before their EDA can be compared numerically with Clancy.")
    bullet(doc, "The scripts are reproducible only when source URLs, downloaded files, subtitles, offset tables, rejection lists, and environment versions are retained.")

    heading(doc, "9. One-Sentence Explanation for the Mentor")
    para(doc, "I built the corpus by moving from verified source metadata to transcript turns, applying explicit timeline corrections, extracting synchronized MP4/WAV clips, filtering and reviewing quality, and finally creating leakage-aware MELD-style partitions; the current Clancy branch is the primary benchmark, while tribunal data proves the reusable method and the other branches remain controlled expansion plans.")

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
