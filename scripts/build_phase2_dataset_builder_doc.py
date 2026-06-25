from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_Phase2_Dataset_Builder_Implementation_Guide.docx"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def build_doc() -> Document:
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("LegalMemoCMT Phase 2 Dataset Builder Implementation Guide")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(
        "Student-focused explanation of the manifest-driven courtroom dataset pipeline, weak supervision, and MELD warm-start adaptation."
    )
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12.5)

    add_para(
        doc,
        "This guide explains how the Phase 2 dataset builder is organized around two input CSVs: a tribunal-source planning file and a witness-harvest manifest. The purpose is to turn public legal records into a traceable, incremental, and reviewable dataset that can later feed the LegalMemoCMT fine-tuning pipeline.",
    )

    doc.add_heading("1. Input Files", level=1)
    add_table(
        doc,
        ["File", "Role", "Key fields"],
        [
            ["tribunal_sources_target_dataset.csv", "High-level source planning", "tribunal, case_family, content_type, source_url, target_video_hours, target_witnesses"],
            ["witness_harvest_manifest.csv", "Primary acquisition manifest", "manifest_id, tribunal, case_name, witness_name_or_code, source_evidence_url, transcript_url_to_resolve, video_url_to_resolve, target_duration_minutes, target_utterance_count, priority"],
        ],
    )
    add_para(
        doc,
        "The first file says what to collect. The second file says which witness-level records to harvest. Together they define the project target before any downloading starts.",
    )

    doc.add_heading("2. Stage 1: Manifest Loading and Validation", level=1)
    add_bullets(
        doc,
        [
            "Load the source CSVs with pandas.",
            "Validate required columns and detect duplicate rows.",
            "Check for missing URLs, invalid durations, and invalid utterance counts.",
            "Create stable manifest identifiers when one is not already present.",
            "Preserve protected witness identifiers exactly as provided.",
        ],
    )

    doc.add_heading("3. Stage 2: Record Discovery", level=1)
    add_bullets(
        doc,
        [
            "Resolve transcript URLs from source evidence pages when the manifest only contains a search target.",
            "Resolve video URLs only when public media is actually available.",
            "Write the intermediate resolved table to data/resolved_manifest.csv.",
            "Keep unresolved rows for manual follow-up rather than deleting them.",
        ],
    )

    doc.add_heading("4. Stage 3: Transcript Acquisition", level=1)
    add_bullets(
        doc,
        [
            "Download transcripts to data/raw/transcripts/.",
            "Support PDF, HTML, and plain-text records.",
            "Extract visible text into a clean text payload for segmentation.",
            "Keep the raw file and the extracted text side by side for traceability.",
        ],
    )

    doc.add_heading("5. Stage 4: Video and Audio Acquisition", level=1)
    add_bullets(
        doc,
        [
            "Download public video files into data/raw/videos/.",
            "Extract audio from video files into data/raw/audio/.",
            "Support mp4, webm, mp3, and wav inputs.",
            "Keep video and audio paths in the final manifest even if one modality is missing.",
        ],
    )

    doc.add_heading("6. Stage 5: Utterance Segmentation", level=1)
    add_para(
        doc,
        "Transcript segmentation is a research step, not a perfect legal parser. The implementation should identify speaker turns where possible, label the most likely role, and preserve the original text so the student can inspect how the segments were formed.",
    )
    add_bullets(
        doc,
        [
            "Witness",
            "Judge",
            "Prosecutor",
            "Defense counsel",
            "Clarification / question turns",
        ],
    )

    doc.add_heading("7. Final Dataset Format", level=1)
    add_table(
        doc,
        ["Column", "Meaning"],
        [
            ["utterance_id", "Unique utterance identifier"],
            ["manifest_id", "Witness manifest identifier"],
            ["tribunal", "ICTY, ICTR, or IRMCT"],
            ["case_name", "Case family or hearing case"],
            ["speaker_role", "Witness, judge, prosecutor, defense counsel"],
            ["speaker_name", "Witness code or protected name"],
            ["utterance_text", "Segment text"],
            ["start_time / end_time", "Optional timing fields"],
            ["video_path / audio_path", "Resolved media paths"],
            ["emotion_label / credibility_label / question_type", "Weak labels for review or supervised training"],
            ["cross_examination_flag", "Whether the turn is part of adversarial questioning"],
        ],
    )

    doc.add_heading("8. Weak Supervision", level=1)
    add_bullets(
        doc,
        [
            "Emotion labels: neutral, fear, anger, sadness, stress, confidence.",
            "Question types: open, closed, leading, clarification, cross examination.",
            "Credibility indicators: consistent, uncertain, contradictory.",
            "Store the weak labels separately so they can be reviewed and improved later.",
        ],
    )

    doc.add_heading("9. Dashboard and Status Reporting", level=1)
    add_bullets(
        doc,
        [
            "Witnesses collected",
            "Hours collected",
            "Utterances collected",
            "Transcript coverage",
            "Video coverage",
            "Annotation coverage",
        ],
    )

    doc.add_heading("10. Adaptation Path", level=1)
    add_bullets(
        doc,
        [
            "Use the best MELD fold checkpoint as the warm-start model.",
            "Reset the classifier head if the Phase 2 label space differs from MELD.",
            "Freeze backbones initially if the courtroom dataset is small.",
            "Fine-tune on the Phase 2 manifest and evaluate the saved checkpoint.",
        ],
    )

    doc.add_heading("11. Scope and Safety", level=1)
    add_bullets(
        doc,
        [
            "Use only publicly available records.",
            "Preserve witness identifiers exactly as provided.",
            "Do not attempt deanonymization.",
            "Do not predict guilt, innocence, or deception.",
            "Keep all intermediate artifacts for traceability.",
        ],
    )

    add_para(
        doc,
        "This guide is intended to help a student understand the dataset builder from input CSVs to the final dataset table and dashboard. It is deliberately modular so that new witnesses and new public records can be added incrementally.",
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()

