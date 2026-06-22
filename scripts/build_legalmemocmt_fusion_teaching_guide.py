from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_Fusion_Teaching_Guide.docx"
ARCH_IMG = ROOT / "artifacts" / "mermaid" / "paper_aligned_architecture.png"
FUSION_IMG = ROOT / "artifacts" / "mermaid" / "paper_aligned_fusion_study.png"


def style_document(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, cell in enumerate(row):
            cells[i].text = cell
    doc.add_paragraph()


def add_image(doc: Document, path: Path, width: float = 6.5) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"[Image missing: {path.name}]")


def add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LegalMemoCMT Fusion Teaching Guide")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "How speech, text, and video are encoded, fused, pooled, and interpreted in the current Phase 1 implementation"
    )
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)

    doc.add_paragraph()
    add_para(
        doc,
        "Purpose: to teach students how the current LegalMemoCMT system combines individual modality encoders with a cross-modal transformer fusion block, and how that design relates to the MemoCMT base paper while still remaining faithful to the current MELD and CREMA-D implementation.",
    )
    add_para(
        doc,
        "Scope: modality-specific preprocessing, legacy versus pretrained encoder modes, fusion tensor shapes, pooling strategies, training and evaluation flow, prediction export, and the role of video as an optional ablation stream.",
    )
    add_para(
        doc,
        "Audience: students who need a technical but readable guide to understand what the code does, why each module exists, and how the pieces fit together end to end.",
    )
    add_para(
        doc,
        "Important framing: this guide explains the implemented system as it exists in the repository. It does not claim that the system reproduces the base paper exactly; instead, it shows how the architecture and ideas are adapted for the current project.",
    )
    doc.add_paragraph()


def add_intro(doc: Document) -> None:
    doc.add_heading("1. Why Fusion Matters", level=1)
    add_para(
        doc,
        "Fusion is the step that turns three separate modality pipelines into one multimodal emotion recognition model. Speech can capture prosody and voice energy, text can capture semantics and discourse meaning, and video can capture facial and body cues. Each modality is informative, but each also has blind spots. The fusion block is the part of the system that learns how to combine those partial views into a single prediction.",
    )
    add_para(
        doc,
        "A student should think of fusion as evidence integration. The model does not simply append three vectors and hope for the best. It first encodes each stream into a compact representation, then learns how much those streams should influence one another, and finally reduces the combined representation to an emotion label. That sequence is important because it separates representation learning from interaction learning and from classification.",
    )
    add_para(
        doc,
        "In the current implementation, the fusion block is a cross-modal transformer. That means the modality vectors are treated like a very short sequence, the transformer learns inter-modality interactions, and a pooling step converts the transformed sequence into one vector for classification. This is a clean and teachable pattern because it resembles the way sequence transformers work in text, but at the level of modalities instead of words.",
    )
    add_para(
        doc,
        "The same design can be understood at two levels. At the engineering level, it is a set of tensors and modules in PyTorch. At the conceptual level, it is a way to ask which modality is most useful at a given moment, whether the modalities agree, and whether a conflict between modalities is itself informative. The teaching value is in showing both levels clearly.",
    )
    doc.add_paragraph()


def add_base_paper_alignment(doc: Document) -> None:
    doc.add_heading("2. Base Paper Alignment", level=1)
    add_para(
        doc,
        "The MemoCMT base paper is the conceptual anchor for the current Phase 1 design. The paper's core idea is to use a speech branch and a text branch, then fuse them with a cross-modal transformer. The paper reports that the choice of fusion pooling matters and compares CLS, mean, max, and min aggregation. The current implementation follows that design philosophy while adapting it to the datasets and workflow used in this project.",
    )
    add_para(
        doc,
        "The base paper's encoder choices are important for teaching because they explain why pretrained models are used. For text, BERT provides a contextual representation rather than a bag of words. For speech, HuBERT or related self-supervised audio encoders provide a learned acoustic representation rather than a handcrafted feature vector. The fusion block then operates on modality-level vectors, not on raw tokens or raw waveforms.",
    )
    add_para(
        doc,
        "The current project keeps the same structural logic: one encoder per modality, one fusion block, one classifier. That is why the guide focuses on the path from input to encoded vector to fused vector to predicted emotion. The student can understand the model by following that path and by seeing where the code uses the same ideas as the paper and where it extends them.",
    )
    add_table(
        doc,
        ["Paper idea", "Current implementation"],
        [
            ["Speech branch", "Pretrained audio encoder in paper mode; legacy audio sequence encoder in baseline mode"],
            ["Text branch", "BERT tokenizer + pretrained text encoder in paper mode; legacy token encoder in baseline mode"],
            ["Fusion", "Cross-modal transformer with learnable modality tokens"],
            ["Pooling", "CLS, mean, max, or min"],
            ["Output", "Emotion class logits"],
            ["Video", "Implemented as an optional branch, but excluded from the core paper-aligned comparison"],
        ],
    )
    add_para(
        doc,
        "This table is the easiest way to teach the project: the architecture is not a black box. It is a concrete sequence of known modules with known responsibilities.",
    )


def add_data_flow(doc: Document) -> None:
    doc.add_heading("3. Data Flow From Manifest to Model Input", level=1)
    add_para(
        doc,
        "The model does not read raw files directly in the training loop. Instead, a CSV manifest describes each sample. The manifest contains sample_id, split, label, transcript, audio_path, and video_path. The dataset layer reads the manifest and prepares tensors that match the selected encoder mode.",
    )
    add_para(
        doc,
        "There are two main operating modes. In legacy mode, the text branch consumes integer token-like values and the audio/video branches consume padded feature sequences. In pretrained or paper mode, the text branch uses a HuggingFace tokenizer and a pretrained text backbone, while the audio branch uses waveform input and a pretrained audio backbone. The video branch remains a sequence encoder in both modes.",
    )
    add_para(
        doc,
        "This separation is pedagogically useful because students can see the difference between a feature-based pipeline and a backbone-based pipeline. The feature-based pipeline is simpler and easier to debug. The backbone-based pipeline is closer to the MemoCMT paper and usually gives a better foundation for comparing fusion strategies.",
    )
    add_code_block(
        doc,
        [
            "Manifest row -> dataset loader -> batch collation -> modality tensors",
            "text transcript -> tokenizer or token encoder",
            "audio file -> waveform or feature sequence",
            "video file -> frame feature sequence",
            "encoder outputs -> modality vectors",
            "modality vectors -> cross-modal transformer -> pooled fused vector",
            "fused vector -> classifier -> emotion logits",
        ],
    )
    add_para(
        doc,
        "A student should understand that the manifest is not just bookkeeping. It is the interface between raw data and the learning system. It defines which sample belongs to which split, which label is used as ground truth, and which file paths must be valid before training can begin.",
    )


def add_modality_section(doc: Document) -> None:
    doc.add_heading("4. The Individual Modality Encoders", level=1)
    add_para(
        doc,
        "The current implementation uses one encoder per modality. This modular design is central to multimodal learning because each modality has a different signal structure. Text is discrete and symbolic. Speech is continuous and temporal. Video is spatial-temporal. The encoders translate those different forms into a common latent dimension so the fusion module can combine them.",
    )

    doc.add_heading("4.1 Speech Emotion Recognition Encoder", level=2)
    add_para(
        doc,
        "In the paper-aligned path, speech is processed through a pretrained audio model such as HuBERT or a similar self-supervised speech encoder. The goal of the speech encoder is to transform a waveform into a high-level representation that still retains the emotional content of the voice: pitch movement, speaking rate, pauses, stress, voice quality, and rhythm. The output is a single vector per sample, not a class label by itself.",
    )
    add_para(
        doc,
        "In the legacy path, audio is represented as a time sequence of extracted features and then passed through a transformer encoder. This is a simpler teaching path and helps students understand how a sequence model pools information over time before fusion. The main lesson is the same in both cases: the speech encoder compresses a longer temporal signal into a latent speech emotion vector.",
    )
    add_bullets(
        doc,
        [
            "Input: raw waveform or a feature sequence.",
            "Hidden goal: preserve emotionally meaningful temporal structure.",
            "Output: a fixed-size speech representation, typically shaped like [batch, fusion_dim].",
            "Why it matters: speech often carries emotional tone that is not present in the transcript.",
        ],
    )
    add_para(
        doc,
        "For students, it helps to think of the speech encoder as a lens. It looks at the utterance through the voice. A calm sentence can sound angry, tired, hesitant, or anxious depending on how it is spoken. The encoder's job is to make that signal available to the fusion layer.",
    )

    doc.add_heading("4.2 Text Emotion Recognition Encoder", level=2)
    add_para(
        doc,
        "In the paper-aligned path, the text encoder uses a BERT tokenizer and a pretrained BERT-style backbone such as bert-base-uncased or a similar model. The tokenizer converts raw transcript text into token IDs and attention masks. The backbone then produces contextual token embeddings, which are pooled into one sentence-level vector. This representation captures semantics, negation, discourse cues, and contextual word meaning.",
    )
    add_para(
        doc,
        "In the legacy path, text is represented by a simpler token-like sequence and passed through a transformer encoder. The legacy path is useful for debugging and for showing students how self-attention works even before introducing a full pretrained model. However, the paper-aligned path is the better choice when the goal is to follow MemoCMT more closely.",
    )
    add_bullets(
        doc,
        [
            "Input: tokenized transcript with attention mask.",
            "Hidden goal: capture context, wording, and discourse meaning.",
            "Output: a fixed-size text representation, again shaped like [batch, fusion_dim].",
            "Why it matters: the transcript tells the model what was said, but not how it was said.",
        ],
    )
    add_para(
        doc,
        "Students should learn that the text encoder is not a bag-of-words classifier. In a BERT-style encoder, the meaning of each token depends on its neighbors. This is important because emotional meaning often depends on word order, contrast, negation, and the relationship between clauses.",
    )

    doc.add_heading("4.3 Video Emotion Representation", level=2)
    add_para(
        doc,
        "The video branch processes raw visual data or pre-extracted visual features. In the current implementation, the video stream is encoded by a sequence encoder that learns from frame-level or feature-level inputs. That means the model can handle a raw video file if preprocessing is enabled, or it can consume saved feature arrays if the data have already been extracted.",
    )
    add_para(
        doc,
        "The important teaching point is that video is not just a single image. Emotion can be visible through micro-expression, gaze direction, head movement, facial tension, body posture, and motion across frames. The encoder therefore treats the frames as a sequence and produces a video representation that can be fused with speech and text.",
    )
    add_bullets(
        doc,
        [
            "Input: frame sequence or frame features.",
            "Hidden goal: capture visible emotion cues over time.",
            "Output: a fixed-size video representation.",
            "Why it matters: visual cues can confirm, soften, or contradict speech and text.",
        ],
    )
    add_para(
        doc,
        "In the paper-aligned comparison, video is available in the code but is not part of the core MemoCMT-style result table. It is better treated as an ablation or extension stream. That makes the main fusion study cleaner and more faithful to the base paper's audio-text emphasis, while still keeping the code ready for a broader multimodal extension.",
    )


def add_fusion_section(doc: Document) -> None:
    doc.add_heading("5. How Cross-Modal Fusion Works", level=1)
    add_para(
        doc,
        "This is the key section for students. After the modality encoders finish, the model has three vectors, one for speech, one for text, and one for video. Each vector already summarizes one channel of information. Fusion is the stage where these vectors are made to interact so the model can learn relationships such as agreement, contradiction, reinforcement, and complementarity.",
    )
    add_para(
        doc,
        "In the current code, fusion begins by stacking the modality vectors into a short sequence. If the text vector is z_text, the speech vector is z_audio, and the video vector is z_video, then the stack is [z_text, z_audio, z_video]. The stack has shape [batch, 3, fusion_dim]. This makes the modality vectors look like tokens in a sequence, which is exactly what a transformer can process well.",
    )
    add_para(
        doc,
        "The fusion block also adds learnable modality tokens. These are not the same as text tokens or frame tokens. They are modality identity embeddings that tell the transformer which position corresponds to which modality. In other words, they make it easier for the model to distinguish text evidence from speech evidence from visual evidence.",
    )
    add_code_block(
        doc,
        [
            "z_text  = Enc_text(text_input)",
            "z_audio = Enc_audio(audio_input)",
            "z_video = Enc_video(video_input)",
            "",
            "S = stack([z_text, z_audio, z_video])                # [B, 3, D]",
            "S = S + modality_embeddings                           # identity for each stream",
            "F = TransformerEncoder(S)                             # cross-modal interaction",
            "h = Pool(F)                                           # CLS / mean / max / min",
            "y = Classifier(h)                                     # emotion logits",
        ],
    )
    add_para(
        doc,
        "This sequence is the conceptual heart of the project. The encoders turn raw inputs into vectors. The fusion transformer turns vectors into a joint multimodal representation. The pooling step turns the joint representation into one final summary. The classifier maps that summary to an emotion label.",
    )
    add_para(
        doc,
        "The transformer is useful here because self-attention can model interactions between the three streams without forcing a fixed, hand-designed rule. For example, if the transcript is neutral but the voice sounds tense and the face looks strained, the transformer can learn that this combination is a stronger emotion cue than any single stream alone.",
    )
    add_para(
        doc,
        "Students should notice the difference between concatenation and fusion. Concatenation simply places features side by side. Fusion, in the transformer sense, lets the features condition each other. That interaction is the main reason the model is called a cross-modal transformer rather than a feature concatenation baseline.",
    )

    doc.add_heading("5.1 Shape-Level Explanation", level=2)
    add_table(
        doc,
        ["Stage", "Typical shape", "Interpretation"],
        [
            ["Text encoder output", "[B, D]", "One vector per sample summarizing transcript meaning"],
            ["Speech encoder output", "[B, D]", "One vector per sample summarizing acoustic emotion cues"],
            ["Video encoder output", "[B, D]", "One vector per sample summarizing visual emotion cues"],
            ["Stacked fusion input", "[B, 3, D]", "A sequence of three modality tokens"],
            ["Fusion transformer output", "[B, 3, D]", "Contextualized modality vectors after cross-modal interaction"],
            ["Pooled fused vector", "[B, D]", "One summary vector per sample"],
            ["Classifier output", "[B, C]", "Logits for C emotion classes"],
        ],
    )
    add_para(
        doc,
        "This shape table is very useful for teaching because students often understand the model only after they understand the tensor dimensions. The whole pipeline can be seen as a reduction from three modality vectors to one fused vector to one prediction vector.",
    )

    doc.add_heading("5.2 What the Learnable Modality Tokens Do", level=2)
    add_para(
        doc,
        "The modality tokens are a small but important design choice. If the transformer only saw three anonymous vectors, it would have to infer the role of each vector from context alone. By adding a learnable embedding to each stream, the model gets an explicit clue about which vector corresponds to which modality. This is similar in spirit to positional encoding in a standard transformer.",
    )
    add_para(
        doc,
        "In teaching terms, the modality token says: this is text, this is audio, this is video. That label does not replace learning, but it stabilizes learning. It helps the model know that a vector at position 0 is not just another arbitrary token, but the speech-derived or text-derived signal whose identity matters in the fusion stage.",
    )
    add_para(
        doc,
        "The current code uses a very compact version of this idea: a 3-by-D parameter matrix is broadcast across the batch and added to the stacked modality sequence. The transformer then processes a modality-aware sequence rather than a bare vector set.",
    )


def add_pooling_section(doc: Document) -> None:
    doc.add_heading("6. Why Pooling Strategies Matter", level=1)
    add_para(
        doc,
        "After the cross-modal transformer, the model still has a short sequence of three fused modality tokens. A classifier cannot usually consume that sequence directly in the current design, so the sequence has to be pooled into one vector. The pooling choice is not a cosmetic detail. It changes how the model summarizes the fused evidence.",
    )
    add_para(
        doc,
        "The current implementation supports four pooling strategies: CLS, mean, max, and min. These are intentionally simple, because they let the student compare different ways of reducing a multimodal sequence without changing the rest of the architecture. That makes the pooling study interpretable and close in spirit to the MemoCMT paper's fusion comparison.",
    )
    add_bullets(
        doc,
        [
            "CLS pooling: take the first fused token as the sequence summary.",
            "Mean pooling: average the fused tokens dimension-wise.",
            "Max pooling: keep the strongest activation per dimension.",
            "Min pooling: keep the lowest activation per dimension.",
        ],
    )
    add_para(
        doc,
        "For students, the important question is not which pooling sounds mathematically best, but what assumption each pooling method makes. CLS assumes the first token can act as a learned summary. Mean assumes all modalities contribute evenly. Max assumes the strongest evidence should dominate. Min is less common, but it can reveal whether the model learns a conservative summary when one dimension is consistently suppressed.",
    )
    add_para(
        doc,
        "The base paper reports that pooling choice affects performance and includes a case study around this point. The current project preserves that idea by making pooling a configuration parameter and by exposing separate runs for CLS, mean, max, and min in the paper-aligned scripts.",
    )
    add_image(doc, FUSION_IMG, width=6.7)
    add_para(
        doc,
        "This diagram is useful for teaching the pooling study. It shows that the same encoders can produce different results simply by changing how the fused tokens are summarized before classification.",
    )


def add_paper_aligned_workflow(doc: Document) -> None:
    doc.add_heading("7. Paper-Aligned Workflow in the Repository", level=1)
    add_para(
        doc,
        "The repository contains two styles of experimentation. The first is a lightweight or legacy baseline path. The second is the paper-aligned path, which uses pretrained encoders and a clearer MemoCMT-style structure. The teaching guide should emphasize the paper-aligned workflow because that is the path that matches the base paper most closely.",
    )
    add_table(
        doc,
        ["Script or module", "Purpose"],
        [
            ["scripts/run_paper_aligned_case_study.sh", "Runs MELD pooling variants for the MemoCMT-style case study"],
            ["scripts/run_paper_aligned_suite.sh", "Runs the paper-aligned multimodal benchmark suite"],
            ["scripts/run_phase1_mandatory.sh", "Runs the project's mandatory validation workflow"],
            ["scripts/export_predictions.py", "Exports per-sample predicted vs actual values"],
            ["scripts/analyze_predictions.py", "Builds first-20 summaries and confusion matrices"],
            ["src/models/model.py", "Defines the encoders, fusion block, and classifier"],
            ["src/train/train.py", "Handles training, masking, and checkpoint saving"],
            ["src/train/evaluate.py", "Handles held-out evaluation and export"],
        ],
    )
    add_para(
        doc,
        "Students should see that the project is script-driven. The scripts are not decorative. They define the experiment protocol: what data are used, what modalities are included, what pooling is tested, what encoder mode is active, and what outputs are saved. That is how the codebase turns from a collection of modules into a reproducible study.",
    )
    add_para(
        doc,
        "The paper-aligned case study is especially important because it mirrors the way the base paper compares different fusion aggregations. The student can run one script and receive multiple model variants, each with a distinct pooling method, while keeping the rest of the pipeline stable. That is the correct way to teach controlled experimentation.",
    )


def add_video_section(doc: Document) -> None:
    doc.add_heading("8. Where Video Fits and Why It Is Kept Separate", level=1)
    add_para(
        doc,
        "Video is implemented in the codebase, but the core paper-aligned comparison keeps it separate. That is a deliberate design choice. The base MemoCMT paper is centered on speech and text fusion, so the cleanest comparison on the current datasets is to study speech only, text only, and speech-plus-text fusion before mixing in video as an additional variable.",
    )
    add_para(
        doc,
        "From a teaching perspective, video should be understood as a valid third evidence stream, but not as the main variable in the paper-aligned result table. The reason is experimental clarity. If video is included in the main comparison, the story changes from 'how does cross-modal text-speech fusion work?' to 'how does a full trimodal system behave?' Both are valid questions, but they are not the same question.",
    )
    add_bullets(
        doc,
        [
            "Video can be used for optional ablation studies.",
            "Video can be used to demonstrate the generality of the fusion block.",
            "Video can be used in later Phase 2 work if the project needs a broader multimodal setting.",
            "Video should not be mixed into the core paper-aligned comparison if the goal is to stay close to MemoCMT.",
        ],
    )
    add_para(
        doc,
        "In practice, the code still accepts video tensors, and the fusion module still expects three modality vectors. When the modality list excludes video, the code masks it out. That makes the stream available without forcing it into every experiment. This is a good teaching example of how a project can be modular while still having a clear primary research path.",
    )


def add_training_eval_section(doc: Document) -> None:
    doc.add_heading("9. Training, Evaluation, and Prediction Analysis", level=1)
    add_para(
        doc,
        "The training script is responsible for turning the architecture into an experiment. It loads the manifest, builds the dataset, creates train and validation splits, computes class weights if needed, and then trains the model for the selected number of epochs. The evaluation script reuses the same model and dataset logic but switches to held-out evaluation and optional export of metrics and per-sample predictions.",
    )
    add_para(
        doc,
        "This separation is very important pedagogically. Students often assume that training and evaluation are just two calls to the same function. In reality, training includes gradient updates, class imbalance handling, clipping, and checkpoint selection, while evaluation focuses on metrics, error analysis, and file export. Keeping those roles separate makes the workflow easier to reason about.",
    )
    add_para(
        doc,
        "The project also supports predicted-vs-actual inspection. That means the evaluation run can save a CSV with the sample ID, the split, the true label, the predicted label, the confidence, and whether the prediction is correct. This is the right way to teach students how to move from a single accuracy number to a concrete sample-level understanding of model behavior.",
    )
    add_bullets(
        doc,
        [
            "Train script: fit the model and save the checkpoint.",
            "Evaluate script: compute metrics on a held-out split.",
            "Export predictions: save per-sample actual vs predicted values.",
            "Analyze predictions: create first-20 summaries and confusion matrices.",
        ],
    )
    add_para(
        doc,
        "For teaching, the confusion matrix is often more informative than accuracy because it shows which emotions are being confused with which others. In the current project, that is especially useful for spotting dominant-class bias, such as a tendency to overpredict one emotion class when the model is uncertain.",
    )


def add_theory_section(doc: Document) -> None:
    doc.add_heading("10. The Theory Behind the Current Fusion Design", level=1)
    add_para(
        doc,
        "The current fusion design is not only an engineering choice. It is also a theory of multimodal emotion recognition. The theory is that each modality captures different evidence about affect, that those evidence streams can be projected into a common latent space, and that a transformer can learn the interaction pattern among them.",
    )
    add_para(
        doc,
        "From a representation learning perspective, the encoders reduce high-dimensional raw inputs into a shared dimension D. From an interaction perspective, the transformer lets each modality attend to the others. From a decision perspective, the classifier uses the pooled joint representation to predict one emotion class. If any one of those steps fails, the whole system becomes weak or unstable.",
    )
    add_para(
        doc,
        "The reason transformers work well for fusion is that attention is data-driven. The model does not need a hard-coded rule that says, for example, speech should dominate when audio energy is high. It can learn from the data whether speech, text, or video should dominate under certain patterns. That flexibility is valuable, but it also means the model needs strong regularization, good data, and careful evaluation.",
    )
    add_para(
        doc,
        "Pooling strategies then become the summarization theory. CLS is a learned summary token strategy. Mean pooling assumes an average consensus. Max pooling assumes the strongest signal should dominate. Min pooling assumes the lowest activation can be meaningful in contrastive cases. Teaching the student to reason about these assumptions makes the architecture much easier to understand.",
    )
    add_para(
        doc,
        "The biggest practical lesson is that fusion is not a single trick. It is a sequence of design decisions: representation, alignment, interaction, pooling, classification, and evaluation. A student who learns that sequence can understand many multimodal architectures, not just this project.",
    )


def add_quick_reference(doc: Document) -> None:
    doc.add_heading("11. Quick Reference for Students", level=1)
    add_table(
        doc,
        ["Question", "Short answer"],
        [
            ["What does SER do?", "Turns speech into a latent emotion vector."],
            ["What does TER do?", "Turns transcript text into a latent emotion vector."],
            ["What does the video branch do?", "Turns frames or visual features into a latent emotion vector."],
            ["What does fusion do?", "Lets the modality vectors interact and then summarizes them."],
            ["What does pooling do?", "Reduces the fused token sequence to one vector."],
            ["What does the classifier do?", "Maps the fused vector to emotion logits."],
            ["Why BERT?", "It gives contextual text representations instead of simple word counts."],
            ["Why HuBERT or wav2vec2?", "They give strong pretrained speech representations."],
            ["Why include modality tokens?", "They tell the transformer which stream is which."],
            ["Why keep video separate in the main comparison?", "To stay close to the MemoCMT-style speech-text study."],
        ],
    )
    add_para(
        doc,
        "If a student can answer the table above, they understand the structure of the current implementation at a useful level.",
    )


def add_diagrams(doc: Document) -> None:
    doc.add_heading("12. Architecture Diagrams", level=1)
    add_para(
        doc,
        "The next figure shows the paper-aligned architecture at a high level. It is the simplest visual explanation of how the three modality streams are turned into one fused prediction.",
    )
    add_image(doc, ARCH_IMG, width=6.8)
    add_para(
        doc,
        "The second figure shows the fusion study. It is the visual companion to the pooling discussion and is useful for teaching how the same encoder outputs can be summarized in different ways.",
    )
    add_image(doc, FUSION_IMG, width=6.8)


def add_comparison_section(doc: Document) -> None:
    doc.add_heading("13. Paper vs Phase 1 Implementation", level=1)
    add_para(
        doc,
        "The project is designed to be paper-aligned in spirit, but the exact implementation has been adapted to the available datasets and project goals. That is not a weakness; it is the honest description of the current state. Students should be able to distinguish between following the same architecture idea and reproducing the exact original setup.",
    )
    add_table(
        doc,
        ["Aspect", "Base paper", "Current Phase 1"],
        [
            ["Core modalities", "Speech and text", "Speech, text, and optional video"],
            ["Speech encoder", "HuBERT", "HuBERT or similar pretrained speech backbone in paper mode"],
            ["Text encoder", "BERT", "BERT tokenizer + BERT-style backbone in paper mode"],
            ["Fusion", "Cross-modal transformer", "Cross-modal transformer with modality tokens"],
            ["Pooling", "CLS, mean, max, min", "Same four pooling options"],
            ["Datasets", "Paper-specific benchmark choices", "MELD and CREMA-D for the current project"],
            ["Reporting", "Paper metrics and case study", "Held-out test metrics, confusion matrices, and sample-level exports"],
        ],
    )
    add_para(
        doc,
        "This comparison should be used carefully. It shows that the current work is modeled on the paper's logic, but it also shows where the current project is intentionally broader or more practical.",
    )


def add_appendix(doc: Document) -> None:
    doc.add_heading("14. Appendix: What Students Should Remember", level=1)
    add_numbered(
        doc,
        [
            "Each modality has its own encoder because each modality has a different data geometry.",
            "Fusion is about interaction, not just concatenation.",
            "Pooling is a modeling choice and should be evaluated, not ignored.",
            "Pretrained encoders are used in the paper-aligned path because they encode richer prior knowledge.",
            "Video is supported by the code, but it is not the main paper-aligned comparison stream.",
            "Evaluation should always include more than one metric and should include error analysis.",
            "Predicted-vs-actual tables are often the best way to teach what the model actually learned.",
            "A good multimodal system should be interpreted as evidence analysis, not as a legal judgment engine.",
        ],
    )
    add_para(
        doc,
        "If the student can explain these eight points in their own words, they understand the current implementation at a useful technical level.",
    )


def add_references(doc: Document) -> None:
    doc.add_heading("15. Reference Links", level=1)
    refs = [
        "MemoCMT base paper: https://www.nature.com/articles/s41598-025-89202-x.pdf",
        "Base paper-related implementation notes in the repo: src/models/model.py, src/train/train.py, src/train/evaluate.py",
        "Architecture diagram PNG: artifacts/mermaid/paper_aligned_architecture.png",
        "Fusion study diagram PNG: artifacts/mermaid/paper_aligned_fusion_study.png",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Bullet")


def main() -> None:
    doc = Document()
    style_document(doc)
    add_title_page(doc)
    add_intro(doc)
    add_page_break(doc)
    add_base_paper_alignment(doc)
    add_page_break(doc)
    add_data_flow(doc)
    add_page_break(doc)
    add_modality_section(doc)
    add_page_break(doc)
    add_fusion_section(doc)
    add_page_break(doc)
    add_pooling_section(doc)
    add_page_break(doc)
    add_paper_aligned_workflow(doc)
    add_page_break(doc)
    add_video_section(doc)
    add_page_break(doc)
    add_training_eval_section(doc)
    add_page_break(doc)
    add_theory_section(doc)
    add_page_break(doc)
    add_quick_reference(doc)
    add_page_break(doc)
    add_diagrams(doc)
    add_page_break(doc)
    add_comparison_section(doc)
    add_page_break(doc)
    add_appendix(doc)
    add_page_break(doc)
    add_references(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
