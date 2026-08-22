from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOP = ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx"


def code(document: Document, value: str) -> None:
    paragraph = document.add_paragraph(style="No Spacing")
    run = paragraph.add_run(value)
    run.font.name = "Courier New"
    run.font.size = Pt(8)


def table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = document.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, value in zip(t.rows[0].cells, headers):
        cell.text = value
    for row in rows:
        cells = t.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)


document = Document(SOP)
document.add_section(WD_SECTION.NEW_PAGE)
document.add_heading("Clancy ViT Face-Crop Feature Preparation", level=1)
document.add_paragraph(
    "Purpose. The selected Phase 1 trimodal checkpoints were trained with a visual input dimension of "
    "768. The Clancy source manifests currently point to raw MP4 clips, so they cannot be passed directly "
    "to these checkpoints through the existing generic 128-dimensional fallback extractor. This stage "
    "converts each Clancy video clip into a reusable NumPy feature file and records that file in a new "
    "manifest column. It does not download media, recreate clips, or modify the original text and audio."
)

document.add_heading("Existing Code Reused", level=2)
document.add_paragraph(
    "The implementation reuses the Phase 1 face-crop conventions rather than inventing a second visual "
    "format. `scripts/build_meld_vit_facecrop_manifest.py` supplies `sample_video_frames`, Haar face "
    "detection, padded largest-face cropping, ViT CLS extraction, and NumPy serialization. The new Clancy "
    "wrapper is `phase2/run_build_clancy_vit_facecrop_embeddings.sh`; the worker is "
    "`phase2/build_clancy_vit_facecrop_embeddings.py`."
)
table(document, ["Component", "Input", "Output", "Student meaning"], [
    ["Frame sampling", "Clancy MP4", "16 RGB frames", "Select a fixed visual sample from each clip."],
    ["Face crop", "Sampled frames", "224 x 224 face or center crop", "Focus the visual representation on the visible face when detected."],
    ["ViT encoder", "Cropped frames", "One 768-value CLS vector per frame", "Convert pixels into learned visual features."],
    ["NumPy save", "16 x 768 array", "`.npy` file", "Store the reusable numeric representation."],
    ["Manifest update", "Original CSV row", "`video_features_path`", "Connect the utterance to the exact visual feature file."],
])

document.add_heading("Manifest Traceability", level=2)
document.add_paragraph(
    "The output manifest deliberately keeps both source and processed paths. `video_path` remains the "
    "original MP4 path for human inspection. `raw_video_path` explicitly records that same source path. "
    "`video_features_path` points to the new `.npy` file and is the path used by the Phase 1 loader when "
    "present. `video_features_status` records whether the feature was created, reused, or failed. This "
    "separation prevents a feature file from being mistaken for the original courtroom recording."
)
table(document, ["Column", "Meaning", "Required check"], [
    ["video_path", "Original Clancy utterance/turn MP4", "The file must exist and remain traceable."],
    ["raw_video_path", "Explicit copy of original MP4 provenance", "Must match the source MP4 used for extraction."],
    ["video_features_path", "Saved ViT face-crop embedding file", "Must exist, load with NumPy, and have shape (*, 768)."],
    ["video_features_status", "created, reused, or failed", "No failed rows may enter trimodal inference."],
])

document.add_heading("Manual Execution Order", level=2)
document.add_paragraph(
    "Run the pilot first. The model download may occur on the first run if "
    "`google/vit-base-patch16-224-in21k` is not cached. CPU is the conservative default for reproducibility."
)
document.add_paragraph("Step 1 - Extract features for 200 Pool A rows", style="List Number")
code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python \\
INPUT_CSV=$PWD/data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest.csv \\
OUTPUT_CSV=$PWD/data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200.csv \\
OUTPUT_ROOT=$PWD/data/processed/phase2/clancy/vit_facecrop_embeddings \\
SUMMARY_JSON=$PWD/reports/phase2/clancy_vit_facecrop_200.json \\
bash phase2/run_build_clancy_vit_facecrop_embeddings.sh \\
  --max-rows 200 \\
  --batch-size 8 \\
  --device cpu''')
document.add_paragraph("Step 2 - Inspect the pilot and verify feature shape", style="List Number")
code(document, '''head -n 2 data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200.csv
cat reports/phase2/clancy_vit_facecrop_200.json

PYTHON_BIN=/opt/anaconda3/bin/python - <<'PY'
import csv
import numpy as np

p = "data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200.csv"
with open(p, newline="", encoding="utf-8") as f:
    rows = list(csv.DictReader(f))
for row in rows[:3]:
    arr = np.load(row["video_features_path"], allow_pickle=False)
    print(row["utterance_id"], arr.shape, arr.dtype, np.isfinite(arr).all())
PY''')
document.add_paragraph("Step 3 - Process the complete Pool A after the pilot passes", style="List Number")
code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python \\
INPUT_CSV=$PWD/data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest.csv \\
OUTPUT_CSV=$PWD/data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit.csv \\
OUTPUT_ROOT=$PWD/data/processed/phase2/clancy/vit_facecrop_embeddings \\
SUMMARY_JSON=$PWD/reports/phase2/clancy_vit_facecrop.json \\
bash phase2/run_build_clancy_vit_facecrop_embeddings.sh \\
  --batch-size 8 \\
  --device cpu \\
  --skip-existing''')
document.add_paragraph("Step 4 - Process the complete Pool B separately", style="List Number")
code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python \\
INPUT_CSV=$PWD/data/processed/phase2/clancy/duration_20_to_30/clancy_dataset_manifest.csv \\
OUTPUT_CSV=$PWD/data/processed/phase2/clancy/duration_20_to_30/clancy_dataset_manifest_vit.csv \\
OUTPUT_ROOT=$PWD/data/processed/phase2/clancy/vit_facecrop_embeddings \\
SUMMARY_JSON=$PWD/reports/phase2/clancy_vit_facecrop_20_to_30.json \\
bash phase2/run_build_clancy_vit_facecrop_embeddings.sh \\
  --batch-size 8 \\
  --device cpu \\
  --skip-existing''')

document.add_heading("How the Updated Loader Uses the Features", level=2)
document.add_paragraph(
    "`src/data/dataset.py` reads `video_features_path` into the sample object. During `__getitem__`, "
    "the loader chooses that path before `video_path`. Because the path ends in `.npy`, `_load_array` "
    "uses `numpy.load(..., allow_pickle=False)` and returns the stored numeric array. Older manifests "
    "without `video_features_path` continue to use their existing `video_path` behavior. This makes the "
    "change backward-compatible while allowing a true 768-dimensional visual input for the selected "
    "Phase 1 checkpoint."
)
document.add_paragraph(
    "The Phase 1 checkpoint must still be checked before inference. A video-enabled checkpoint should "
    "report `use_video=true` and `video_dim=768`. The previously selected text-audio fold-2 checkpoint "
    "under `results/paper_aligned_meld_cv` reports `use_video=false`; it must not be described as a "
    "trimodal run. Use a video-enabled checkpoint such as the face-crop gated checkpoint only after the "
    "`.npy` paths and shapes have passed validation."
)

document.add_heading("Quality Gate Before Pseudo-Labeling", level=2)
for item in [
    "The output row count equals the selected input row count.",
    "Every selected row has a non-empty existing `video_features_path`.",
    "Every array loads successfully with `allow_pickle=False`.",
    "Every array is finite, float32, and has second dimension 768.",
    "The MP4 remains available through `video_path` or `raw_video_path` for audit and manual review.",
    "The audio path still points to the matching WAV and the text remains in `utterance_text`.",
    "No failed or missing feature rows are passed to a trimodal checkpoint.",
]:
    document.add_paragraph(item, style="List Bullet")

document.add_heading("Current Status", level=2)
document.add_paragraph(
    "The extraction worker, shell wrapper, and loader update are implemented and syntax-checked. No Clancy "
    "768-dimensional embedding batch has been executed yet. The next action is the 200-row Pool A pilot; "
    "only after that pilot passes should the complete Pool A and Pool B feature manifests be created."
)

document.save(SOP)
print(f"Updated {SOP}")
