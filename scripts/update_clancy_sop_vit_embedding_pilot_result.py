from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOP = ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx"


def code(document: Document, value: str) -> None:
    paragraph = document.add_paragraph(style="No Spacing")
    run = paragraph.add_run(value)
    run.font.name = "Courier New"
    run.font.size = Pt(8)


document = Document(SOP)
document.add_section(WD_SECTION.NEW_PAGE)
document.add_heading("Recorded Result: Pool A ViT Pilot", level=1)
document.add_paragraph(
    "The 200-row Pool A pilot was executed successfully using the Clancy face-crop ViT extraction "
    "wrapper. The run created 200 embeddings, reused 0 existing embeddings, and failed 0 rows. This is "
    "a successful feature-extraction result, not yet a successful emotion-inference result. The next gate "
    "is to validate the saved arrays and then run a small trimodal Phase 1 inference pilot."
)
code(document, "output_csv = data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200.csv")
code(document, "output_root = data/processed/phase2/clancy/vit_facecrop_embeddings")
code(document, "summary_json = reports/phase2/clancy_vit_facecrop_200.json")
document.add_paragraph("Recorded extraction results:")
for item in [
    "Rows selected: 200",
    "Embeddings created: 200",
    "Embeddings reused: 0",
    "Failed rows: 0",
    "Embedding shape: (16, 768)",
    "Embedding dtype: float32",
    "Device: CPU",
    "ViT model: google/vit-base-patch16-224-in21k",
    "Face processing: Haar largest-face crop with padding and center-crop fallback",
]:
    document.add_paragraph(item, style="List Bullet")
document.add_paragraph(
    "The Transformers message about a slow image processor is informational. It means the installed "
    "fast processor was not selected; it is not a failed row or a feature-shape error. The pilot still "
    "completed with the expected numeric output."
)
document.add_heading("Post-Pilot Array Validation", level=2)
code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python - <<'PY'
import csv
import numpy as np

p = "data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200.csv"
with open(p, newline="", encoding="utf-8") as f:
    rows = list(csv.DictReader(f))
assert len(rows) == 200
for row in rows:
    arr = np.load(row["video_features_path"], allow_pickle=False)
    assert arr.shape == (16, 768), (row["utterance_id"], arr.shape)
    assert arr.dtype == np.float32, (row["utterance_id"], arr.dtype)
    assert np.isfinite(arr).all(), row["utterance_id"]
print("PASS", len(rows), "rows with finite float32 (16, 768) embeddings")
PY''')
document.add_heading("Next Inference Pilot", level=2)
document.add_paragraph(
    "Use a checkpoint whose stored configuration has `use_video=true` and `video_dim=768`. The fold-2 "
    "face-crop gated checkpoint is the current weighted-F1 primary trimodal candidate. Do not use the "
    "earlier `paper_aligned_meld_cv` fold-2 checkpoint for this trimodal step because its configuration "
    "has `use_video=false`."
)
code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python \\
INPUT_CSV=$PWD/data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200.csv \\
OUTPUT_CSV=$PWD/data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_phase1_pseudo_200.csv \\
CHECKPOINT=$PWD/results/facial_cues/meld_vit_facecrop_gated/fold_2/best_model.pt \\
SUMMARY_JSON=$PWD/reports/phase2/clancy_vit_phase1_pseudo_200.json \\
MAX_ROWS=0 \\
BATCH_SIZE=2 \\
MODALITIES=text,audio,video \\
DEVICE=cpu \\
bash phase2/run_pseudo_label_clancy_with_phase1.sh''')
document.add_paragraph(
    "For this pilot, MAX_ROWS=0 is correct because the input manifest already contains exactly 200 rows. "
    "The inference summary must report effective modalities containing audio, text, and video. If it "
    "reports that video is disabled, stop and inspect the checkpoint configuration rather than proceeding."
)
document.add_heading("Updated Status", level=2)
document.add_paragraph(
    "Feature extraction is now PASS for the 200-row pilot. Full Pool A and Pool B feature extraction, "
    "trimodal Phase 1 inference, and manual prediction review remain pending."
)
document.save(SOP)
print(f"Updated {SOP}")
