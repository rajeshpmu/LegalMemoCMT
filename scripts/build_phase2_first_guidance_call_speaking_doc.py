from __future__ import annotations

import csv
import json
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_Phase2_First_Guidance_Call_Speaking_Doc.docx"
CANDIDATE_LEDGER = ROOT / "data" / "phase2" / "source_manifests" / "case_candidate_ledger.csv"
SUMMARY_JSON = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated" / "dataset_quality_summary.json"
VERIFIED_INVENTORY = ROOT / "data" / "processed" / "phase2" / "verified_case_inventory.csv"
HEARING_MANIFEST = ROOT / "data" / "processed" / "phase2" / "hearing_manifest.csv"
WITNESS_MANIFEST = ROOT / "data" / "phase2" / "source_manifests" / "witness_harvest_manifest_resolved.csv"
LEGALMELD_METADATA = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated" / "legalmeld_metadata_validated.csv"
ALIGNMENT_REVIEW = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated" / "alignment_review_sample.csv"
TRAIN_CSV = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated" / "train.csv"
DEV_CSV = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated" / "dev.csv"
TEST_CSV = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated" / "test.csv"
EXPANDED_COMPARE_SUMMARY = ROOT / "reports" / "phase2" / "expanded_planning_vs_verified_inventory_summary.json"
EXPANDED_COMPARE_MISSING = ROOT / "data" / "processed" / "phase2" / "expanded_planning_missing_sources.csv"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def load_csv_rows(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def load_json(path: Path) -> dict[str, object]:
    if not path.exists():
        return {}
    return json.loads(path.read_text())


def count_rows(path: Path) -> int:
    return len(load_csv_rows(path))


def first_row(rows: list[dict[str, str]], key: str, value: str) -> dict[str, str]:
    for row in rows:
        if (row.get(key) or "").strip().lower() == value.lower():
            return row
    return {}


def format_count(value: object) -> str:
    try:
        return f"{int(value):,}"
    except Exception:
        return "0"


def summarize_counts() -> dict[str, str]:
    summary = load_json(SUMMARY_JSON)
    hearing_rows = load_csv_rows(HEARING_MANIFEST)
    witness_rows = load_csv_rows(WITNESS_MANIFEST)
    legalmeld_rows = load_csv_rows(LEGALMELD_METADATA)
    review_rows = load_csv_rows(ALIGNMENT_REVIEW)

    alignment_counts = summary.get("alignment_confidence_counts", {})
    quality_counts = summary.get("quality_tier_counts", {})
    split_counts = summary.get("split_counts", {})
    speaker_counts = summary.get("speaker_role_counts", {})

    if not isinstance(alignment_counts, dict):
        alignment_counts = {}
    if not isinstance(quality_counts, dict):
        quality_counts = {}
    if not isinstance(split_counts, dict):
        split_counts = {}
    if not isinstance(speaker_counts, dict):
        speaker_counts = {}

    return {
        "verified_cases": format_count(summary.get("cases_represented", 0)),
        "hearings_represented": format_count(summary.get("hearings_represented", 0)),
        "utterances": format_count(summary.get("total_utterances", summary.get("utterances_exported", 0))),
        "high_confidence": format_count(summary.get("high_confidence_alignments", 0)),
        "medium_confidence": format_count(alignment_counts.get("MEDIUM", 0)),
        "low_confidence": format_count(alignment_counts.get("LOW", 0)),
        "quality_b": format_count(quality_counts.get("B", 0)),
        "quality_reject": format_count(quality_counts.get("REJECT", 0)),
        "train_rows": format_count(count_rows(TRAIN_CSV)),
        "dev_rows": format_count(count_rows(DEV_CSV)),
        "test_rows": format_count(count_rows(TEST_CSV)),
        "metadata_rows": format_count(count_rows(LEGALMELD_METADATA)),
        "inventory_rows": format_count(count_rows(VERIFIED_INVENTORY)),
        "hearing_manifest_rows": format_count(count_rows(HEARING_MANIFEST)),
        "witness_manifest_rows": format_count(count_rows(WITNESS_MANIFEST)),
        "split_leakage": format_count(summary.get("split_leakage_violations", 0)),
        "distinct_witnesses": format_count(summary.get("distinct_witnesses", 0)),
        "witness_visible": format_count(summary.get("witness_visible_count", 0)),
        "audio_valid": format_count(summary.get("audio_valid_count", 0)),
        "video_valid": format_count(summary.get("video_valid_count", 0)),
        "manual_review": format_count(summary.get("manual_review_count", 0)),
        "alignment_counts": str(alignment_counts),
        "quality_counts": str(quality_counts),
        "split_counts": str(split_counts),
        "speaker_counts": str(speaker_counts),
        "hearing_rows": hearing_rows,
        "witness_rows": witness_rows,
        "legalmeld_rows": legalmeld_rows,
        "review_rows": review_rows,
    }


def add_section(doc: Document, title: str, purpose: str, explanation: str, bullets: list[str] | None = None) -> None:
    doc.add_heading(title, level=1)
    add_para(doc, "Purpose: ", bold=True)
    add_para(doc, purpose)
    add_para(doc, explanation)
    if bullets:
        add_bullets(doc, bullets)


def build_doc() -> Document:
    s = summarize_counts()
    candidate_rows = load_csv_rows(CANDIDATE_LEDGER)
    first_candidate = candidate_rows[0] if candidate_rows else {}
    first_verified = first_row(load_csv_rows(VERIFIED_INVENTORY), "verification_status", "verified")
    first_hearing_paired = first_row(load_csv_rows(HEARING_MANIFEST), "pairing_status", "paired")
    first_witness = first_row(load_csv_rows(WITNESS_MANIFEST), "witness_type", "unresolved_witness")
    first_review = s["review_rows"][0] if s["review_rows"] else {}
    first_legalmeld = s["legalmeld_rows"][0] if s["legalmeld_rows"] else {}
    compare_summary = load_json(EXPANDED_COMPARE_SUMMARY)
    compare_missing_rows = load_csv_rows(EXPANDED_COMPARE_MISSING)
    compare_counts = compare_summary.get("category_counts", {}) if isinstance(compare_summary.get("category_counts"), dict) else {}

    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LegalMemoCMT Phase 2 First Guidance Call Speaking Doc")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Technical speaking notes for explaining the current Phase 2 deck to the mentor without simply reading the slide text."
    )
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12.5)

    add_para(
        doc,
        "How to use this document: each section follows the slide order in the guidance-call deck. The goal is not to recite the slide bullets. The goal is to explain the logic behind each slide, the implementation status behind each claim, and the limits of what has actually been built so far.",
    )
    add_bullets(
        doc,
        [
            "Speak from the pipeline, not from the slide text.",
            "Use the tribunal corpus as a bootstrap proof, not as the final Indian adaptation result.",
            "Keep the claims conservative: verified inventory means grounded metadata, not finished fine-tuning.",
            "When a slide mentions Indian sources, explain them as a future adaptation layer, not as completed corpus acquisition.",
        ],
    )

    add_section(
        doc,
        "Slide 1. Title Slide",
        "Open the call by framing the project as a Phase 2 method and corpus-construction milestone, not as a finished model.",
        "The title slide should be used to anchor the conversation around the dissertation objective. The important thing here is to make it clear that Phase 2 is about building a grounded, reproducible pipeline for multilingual multimodal emotion analysis in a courtroom setting. The tribunal corpus is the bootstrap proof that the pipeline works. The Indian courtroom corpus is still the final adaptation target. If you begin by saying that, you immediately prevent the discussion from drifting into claims that the Indian corpus is already complete.",
        [
            "Explain that the deck is about the first Phase 2 guidance call, not the final defense.",
            "Make the distinction between bootstrap corpus and final adaptation corpus explicit.",
            "Use the phrase 'verification-first pipeline' early so the mentor understands the methodological emphasis.",
            "Do not start with model accuracy. Start with data construction and trust levels.",
        ],
    )

    add_section(
        doc,
        "Slide 2. Agenda",
        "Show the mentor the narrative order: Phase 1 review, Phase 2 novelty, implemented pipeline, Indian adaptation, and what remains.",
        "The agenda slide is not just a list of topics. It is your roadmap for how the argument flows. The mentor should hear that the presentation starts by connecting Phase 1 to Phase 2, then moves into the new data-construction method, then into the current code-level implementation, and finally into the Indian adaptation objective. This keeps the presentation logically progressive instead of jumping straight into technical details.",
        [
            "Use the agenda to say that the project has a structured ladder of evidence.",
            "Let the mentor see that you are not asking for a random expansion of scope.",
            "Signal that the Indian adaptation is already part of the Phase 2 plan, not an unrelated add-on.",
        ],
    )

    add_section(
        doc,
        "Slide 3. Phase 1 Review",
        "Explain Phase 1 as the multimodal baseline that proved the architecture, but not the utterance-level courtroom pipeline.",
        "When you explain Phase 1, do not present it as a failed stage. Present it as the baseline layer that established the multimodal backbone. The key limitation is unit size: Phase 1 works at the benchmark level, but courtroom speech is not naturally organized as one whole-video sample. A hearing contains many speakers, different speaking roles, and many small turns. That is why Phase 2 exists. It replaces coarse inputs with grounded utterance-level samples so the model learns from the right granularity.",
        [
            "Emphasize that Phase 1 proves the architecture can fuse text, audio, and video.",
            "Explain that whole hearings are too coarse for the final training unit.",
            "Say that Phase 2 exists because courtroom data needs utterance-level alignment, not only document-level understanding.",
        ],
    )

    add_section(
        doc,
        "Slide 4. Objective and Novelty of Phase 2",
        "Present Phase 2 as a data-construction contribution plus an adaptation strategy, not merely a larger dataset.",
        "The objective of Phase 2 is to build a MELD-style legal corpus from verified tribunal material and then reuse the same logic for Indian sources. The novelty is not a new neural architecture. The novelty is the disciplined pipeline: reject placeholders, verify actual case metadata, build hearing and witness manifests from grounded evidence, and only then export aligned utterances. That is a methodological contribution because it shows how to transform legal records into machine-learning samples without confusing page shells with evidence.",
        [
            "Say the novelty is in the data pipeline and traceability, not in a novel transformer block.",
            "Tie the tribunal bootstrap stage to the Indian adaptation objective.",
            "Use the phrase 'MELD-style utterance corpus' to connect the project to a known benchmark format.",
        ],
    )

    add_section(
        doc,
        "Slide 5. Title and Abstract for Phase 2",
        "Explain why the title stays close to the dissertation topic while the abstract stresses the bootstrap-and-reuse method.",
        "This slide should tell the mentor that the proposed title is still aligned with the project objective: Indian courtroom testimony with multilingual multimodal emotion analysis. The abstract should be explained as a method summary rather than a results summary. The important idea is that the tribunal corpus is the supervised bootstrap set, because public Indian courtroom testimony with synchronized video, audio, transcript, and witness metadata is not widely available. The abstract therefore needs to say that the tribunal data proves the method and the Indian data is the final domain adaptation target.",
        [
            "Do not make the abstract sound like the final Indian corpus already exists.",
            "Explain that the current contribution is a reproducible pipeline and a grounded bootstrap corpus.",
            "Say clearly that this project is about emotion and credibility analysis, not legal judgment or truth verification.",
        ],
    )

    add_section(
        doc,
        "Slide 6. Proposed System for Phase 2",
        "Walk the mentor through the pipeline as a chain of trust filters that turn case rows into utterance rows.",
        "The best way to explain this slide is to treat the system as a sequence of filters. First, a human-curated candidate is checked. Then the UCR verification layer decides whether the case resolves to real case-specific records. Then the verified inventory keeps only grounded rows. Then hearing and witness manifests organize the records at a usable semantic level. Finally, the utterance exporter aligns transcript turns with media and creates the training rows. The Indian adaptation strategy belongs here as the second branch of the same design, not as a separate project.",
        [
            "Explain that each stage exists to reduce uncertainty before the next stage begins.",
            "Make the distinction between case-level evidence and utterance-level training samples explicit.",
            "Mention that the same manifest logic is planned for Indian sources later.",
        ],
    )

    add_section(
        doc,
        "Slide 7. Indian Courtroom Adaptation Strategy",
        "Frame the Indian corpus as a realistic adaptation target that is methodologically necessary, but not easy to source at scale.",
        "This slide is where you connect the tribunal bootstrap to the dissertation's final Indian scope. The honest explanation is that there is no large public Indian courtroom corpus that cleanly gives you synchronized video, audio, transcript, and witness metadata at MELD scale. So the project has to use a hybrid strategy. Supreme Court and High Court livestreams help with Indian legal language and procedure. Mock trials help with direct and cross-examination structure. District court material is closest to real witness testimony but is not broadly available in public multimodal form. That is why the Indian phase is an adaptation corpus, not a simple download task.",
        [
            "Do not overstate Indian source availability.",
            "Say that Supreme Court and High Court material mainly supports legal-language adaptation.",
            "Explain that mock trials are useful because they preserve courtroom structure even if they are simulated.",
        ],
    )

    add_section(
        doc,
        "Slide 8. Proposed Architecture Diagram",
        "Use the two clearer diagrams to show the tribunal bootstrap path and the Indian reuse path.",
        "The diagram slide is the visual summary of the whole story, but the earlier combined version was too dense and too small to explain comfortably in a live call. So the architecture is now shown as two separate diagrams. The first diagram is the tribunal bootstrap flow: Phase 1 baseline, candidate ledger, UCR verification, verified inventory, hearing and witness manifests, transcript-first segmentation, forced alignment, clip extraction, media validation, LegalMELD export, and leakage-aware splitting. The second diagram is the Indian adaptation flow: Indian acquisition pack, Supreme Court and High Court livestreams, district court material, mock trials or academies, Indian candidate ledger, verification, verified inventory, hearing and witness manifests, alignment, and the Indian courtroom adaptation corpus. The important message is that the architecture is not just a classifier; it is a data-quality pipeline that ends in a training corpus, and the Indian branch reuses the same method later instead of being a separate idea. The alignment backend is just a selectable implementation detail inside that pipeline.",
        [
            "Walk the diagram from tribunal evidence to utterance-level samples.",
            "Point out that the Indian branch is the reuse path, not a separate model.",
            "Mention negative controls and positive controls as proof that verification is real, not cosmetic.",
            "Explain that the two-diagram layout is easier to read than one crowded diagram.",
        ],
    )

    add_section(
        doc,
        "Slide 9. Algorithms and Techniques Used",
        "Translate the pipeline into a small set of concrete techniques so the mentor can see the implementation is grounded.",
        "This slide should sound like a technical justification, not a random list of buzzwords. Placeholder rejection is there to stop bad inputs before any request is made. API-based verification is there because generic page headings are not evidence. Grouping and deduplication are there to keep hearing rows clean. Witness identity extraction is there because protected witnesses must not be deanonymized. ASR alignment, fuzzy matching, and fallback timing are there because transcript and audio do not always line up perfectly. ffmpeg and media validation are there because a clip that exists on disk is not automatically a usable training example. Group-based splitting is there because hearing-adjacent rows leak information if you split them randomly.",
        [
            "Explain each technique as a response to a real dataset problem.",
            "Stress that the pipeline reduces uncertainty stage by stage.",
            "Do not present fallback alignment as equivalent to matched alignment.",
        ],
    )

    add_section(
        doc,
        "Slide 10. Core Alignment Fields and Dataset Pipeline",
        "Use this section to explain how transcript text, ASR text, alignment confidence, and split discipline work together in the dataset.",
        "This is one of the most important technical slides in the whole deck because it explains how a courtroom utterance becomes a trustworthy training row. The pipeline keeps three textual views of the same utterance: `utterance_text` is the final exported transcript span that the model should learn from, `transcript_text_normalized` is the cleaned reference text used for matching, and `asr_text` is the machine-generated transcript from the audio clip. Those three fields do different jobs. `utterance_text` is the human-readable sample. `transcript_text_normalized` is the matching reference. `asr_text` is the noisy audio-side signal used to verify whether the clip boundaries and transcript span make sense. The alignment fields then turn that text comparison into a trust decision. `alignment_status` says whether the row was matched, fuzzy, or fallback. `alignment_method` says which alignment path was used. `alignment_confidence` collapses that into HIGH, MEDIUM, or LOW. `text_similarity` is a diagnostic score rather than a label. Together these fields decide whether a row is good enough for training, should be kept as a weaker but still usable sample, or should remain in the review bucket. In practice, the row may be aligned by a heuristic backend or by WhisperX-style alignment when the backend is available, but the meaning of the fields is the same.",
        [
            "Explain the three text views: `utterance_text`, `transcript_text_normalized`, and `asr_text`.",
            "Explain the alignment controls: `alignment_status`, `alignment_method`, `alignment_confidence`, and `text_similarity`.",
            "Explain that `split` is separate from alignment quality and protects train/dev/test leakage.",
            "Say that the most important novelty is verification-first promotion plus leakage-aware splitting and a review queue.",
        ],
    )

    add_table(
        doc,
        ["Field group", "What it means in this pipeline", "How to explain it simply"],
        [
            [
                "Text views",
                "`utterance_text`, `transcript_text_normalized`, and `asr_text` are three different views of the same courtroom turn.",
                "One is the final sample text, one is the cleaned reference, and one is the machine-generated check from audio.",
            ],
            [
                "Alignment controls",
                "`alignment_status`, `alignment_method`, `alignment_confidence`, and `text_similarity` decide how trustworthy the row is.",
                "These fields tell us whether the row was matched well, fuzzily, or only by fallback, and how much we should trust it.",
            ],
            [
                "Dataset usage",
                "`split`, `quality_tier`, and `manual_review_required` decide whether the row should train, validate, test, or stay in review.",
                "These fields keep noisy rows out of the clean training set and prevent leakage across the same witness or hearing.",
            ],
            [
                "Novel pipeline step",
                "Verification-first promotion, witness-disjoint splitting, and a review queue make the dataset auditable.",
                "This is the part that makes the corpus construction method stronger than a simple download pipeline.",
            ],
        ],
    )
    add_para(
        doc,
        "For the call, the key thing is not to read the slide as a list of fields. The key thing is to explain the logic chain: transcript text is cleaned, ASR text checks the audio, alignment fields decide trust, and split fields decide where the row belongs in the model workflow. That is why the pipeline is novel at the dataset level: it does not just export clips, it verifies them, scores them, and keeps leakage under control.",
    )

    add_section(
        doc,
        "Slide 11. Algorithms and Techniques: What To Read",
        "Connect each technical component to the literature so the mentor sees you know what is borrowed from the literature and what is project-specific engineering.",
        "This slide is where you show that the project is not improvising from scratch. MELD is the right reading for the utterance-level data design because it is the closest benchmark shape. Whisper is the right reading for robust ASR and noisy audio handling. WhisperX is the practical reference for word-level alignment when that dependency is available. NeMo Forced Aligner and Montreal Forced Aligner are the right references for word- and segment-level timestamping. BERT is the standard transcript encoder reference. ViT is the visual encoder reference. Focal Loss is the standard reading for class imbalance. ffmpeg and difflib are engineering tools, not research contributions, but they are still essential for the implementation. Explain that the papers justify the modeling ideas, while the pipeline code handles the courtroom-specific engineering.",
        [
            "Say which paper you would read first for each pipeline part.",
            "Clarify that some items are model papers and some are tool references.",
            "Use this slide to show that your implementation choices are technically grounded.",
        ],
    )

    add_table(
        doc,
        ["Reading", "How to explain its use in Phase 2"],
        [
            [
                "MELD: A Multimodal Multi-Party Dataset for Emotion Recognition in Conversations",
                "Use this to justify the utterance-level dataset format and the train/dev/test design.",
            ],
            [
                "Robust Speech Recognition via Large-Scale Weak Supervision (Whisper)",
                "Use this to explain why robust ASR and noisy transcript alignment matter in courtroom audio.",
            ],
            [
                "NeMo Forced Aligner docs + NeMo forced aligner paper + Montreal Forced Aligner paper",
                "Use these to explain how text and audio are converted into timestamps before clip extraction.",
            ],
            [
                "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding",
                "Use this to justify contextual transcript encoding.",
            ],
            [
                "An Image Is Worth 16x16 Words: Transformers for Image Recognition at Scale",
                "Use this to justify the vision branch for facial or frame-level cues.",
            ],
            [
                "Focal Loss for Dense Object Detection",
                "Use this to explain why imbalance-aware training matters when some classes dominate.",
            ],
        ],
    )
    add_para(
        doc,
        "For the call, the key thing is not to read the table line by line. The key thing is to show that each technique has a reason for being in the pipeline and that the reading order follows the implementation order.",
    )

    add_section(
        doc,
        "Slide 12. Source Sites And Download Flow",
        "Use this section to explain how public tribunal pages, UCR verification, and media downloads fit together.",
        "This is the first practical bridge from source discovery to corpus construction. The candidate ledger starts with public tribunal source pages and UCR inventory/search URLs. For example, the curated sources include an ICTY Karadzic trial hearing-list page, an ICTY Mladic witness-information PDF, the UCR case page for IT-04-81, and the ICTR Akayesu case page. The important idea is that crawling, searching, and downloading are three different steps. Crawling means starting from the public page or case-page link. Searching means checking the UCR inventory by case number, case family, or witness code. Downloading means retrieving the actual transcript or video asset only after the record is verified. The point of the section is to show that a generic tribunal page is not enough, and a page shell with headings is not corpus evidence.",
        [
            "Say that the candidate ledger points to public source pages, not to confirmed training rows.",
            "Explain that the UCR layer is used as the verification gate.",
            "Mention that downloads happen only after grounded record metadata is confirmed.",
            "Use the phrase 'generic headings are not proof of records' to keep the limitation clear.",
        ],
    )
    add_para(
        doc,
        "A concrete witness-testimony example is `Nahimana et al. (ICTR-99-52)`. The official ICTR page describes it as a completed media case with a large witness record, and the UCR downloader can list multiple downloadable video recordings for it. The practical story to tell is: first run a single-case probe with `python3 phase2/download_ucr_case_video.py --case-number ICTR-99-52 --date 01/03/2002 --index 1 --verify`, then use the existing enrichment, inventory, media-type split, hearing manifest, witness manifest, and LegalMELD scripts to carry the case into the corpus. If the probe returns no recordings, stop at transcript-only and do not present it as a tri-modal row.",
    )

    add_table(
        doc,
        ["Example source", "How to explain it in the call"],
        [
            [
                "https://www.icty.org/en/content/karadzic-trial-hearing-list",
                "A public tribunal source page that helps start the search, but does not itself prove the row is ready for corpus use.",
            ],
            [
                "https://www.icty.org/x/cases/mladic/custom11/en/mladic_otp_witness_info.pdf",
                "An official tribunal document that can support source discovery and case inspection before UCR verification.",
            ],
            [
                "https://ucr.irmct.org/scasedocs/case/IT-04-81",
                "A UCR case page that is useful for checking whether the case resolves to grounded record metadata.",
            ],
            [
                "https://unictr.irmct.org/en/cases/ictr-96-04",
                "A public tribunal case page that helps with search and traceability, but still needs UCR-grounded confirmation.",
            ],
        ],
    )

    add_section(
        doc,
        "Slide 13. What Is Implemented So Far",
        "Summarize the current code state as a real, reproducible bootstrap pipeline with validated outputs.",
        "This is the slide where you move from method to implementation status. The mentor should hear that the verified inventory exists, the hearing and witness manifest builders exist, and the LegalMELD export stage exists. The current output is small, but it is real. The current validated run contains 130 utterance rows, 123 paired hearings in the hearing manifest, and 0 split-leakage violations. The witness manifest is still conservative because it mostly contains unresolved witness rows, which is appropriate at this stage. The important thing is that the end-to-end path works and can be rerun. The builder is also configurable at the alignment step, so you can mention that the same export schema can be produced through a heuristic path or a WhisperX-style path.",
        [
            f"Verified case inventory currently holds {s['verified_cases']} case groups.",
            f"Hearing manifest currently has {s['hearing_manifest_rows']} rows and the current validated run includes {s['high_confidence']} high-confidence alignments and {s['medium_confidence']} medium-confidence alignments.",
            f"LegalMELD validated outputs currently include {s['metadata_rows']} utterance rows, {s['train_rows']} train rows, {s['dev_rows']} dev rows, and {s['test_rows']} test rows.",
            "The tribunal bootstrap layer is implemented; the Indian acquisition layer is still a planned extension.",
        ],
    )

    add_section(
        doc,
        "Slide 13A. Expanded Planning vs Verified Inventory",
        "Use this section to explain why the TAP shortlist is exhausted and how the next expansion decision is made.",
        "This is the comparison step that connects the planning layer to the verified inventory. I compared `phase2_expanded_planning_manifest.csv` against `verified_case_inventory.csv` and the summary shows 17 expanded rows in total. Of those, 4 rows are already covered by verified case names, 13 rows are still new expansion sources, and none are unresolved placeholders. The four covered rows are the two Karadzic Trial rows and the two Mladic Trial rows. In other words, the current TAP shortlist is not giving us genuinely new tribunal coverage any more; it is mostly repeating cases that the verified inventory already knows about. The 13 new rows are still important, but they are planning candidates, not proof of new corpus readiness. They include Popovic, Bagosora, Akayesu, Ntakirutimana/Ntahobali, IRMCT hearings, and witness planning rows that still need grounded verification before they can be promoted.",
        [
            f"The comparison summary reports {compare_summary.get('expanded_rows', 17)} expanded rows and {compare_summary.get('inventory_rows', 0):,} verified inventory rows.",
            f"{compare_counts.get('already_verified_by_case_name', 0)} rows are already covered by verified case names and {compare_counts.get('new_expansion_source', 0)} rows are still new expansion sources.",
            "The 4 covered rows are the Karadzic Trial and Mladic Trial tribunal rows, which already appear in the verified inventory.",
            "The next step is not another download attempt; it is to inspect the missing-source rows and broaden the tribunal planning inputs.",
        ],
    )
    add_table(
        doc,
        ["Category", "Rows", "What it means"],
        [
            [
                "already_verified_by_case_name",
                str(compare_counts.get("already_verified_by_case_name", 4)),
                "These rows already map to cases that exist in the verified inventory, so they do not add new corpus coverage.",
            ],
            [
                "already_verified_by_case_number",
                str(compare_counts.get("already_verified_by_case_number", 0)),
                "No row was matched this way in the current run, which means the name-based coverage signal was the useful one here.",
            ],
            [
                "new_expansion_source",
                str(compare_counts.get("new_expansion_source", 13)),
                "These rows are still missing from the verified inventory and are the rows that matter if we want broader tribunal growth.",
            ],
            [
                "unresolved_or_placeholder",
                str(compare_counts.get("unresolved_or_placeholder", 0)),
                "No placeholder or unresolved rows survived the comparison, which is good because it means the shortlist is at least grounded.",
            ],
        ],
    )
    add_para(
        doc,
        f"The missing-source file currently has {len(compare_missing_rows)} rows. That is the file I would inspect next, because those rows tell me which tribunal families and witness-planning entries still need broader source coverage. In practical terms, this step prevents me from re-running downloads on the same 4 already-covered rows. It also keeps the next expansion honest: I only broaden the planning manifests if the missing-source rows suggest a genuinely new tribunal family or a still-uncovered witness planning branch.",
    )

    add_section(
        doc,
        "Slide 14. Hearing Manifest EDA",
        "Use this section while showing hearing_manifest.csv in the guidance call.",
        "This is the cleanest file for showing how the hearing-level corpus is shaped before utterance-level export. The current manifest has 2,375 rows in total. Of those, 123 rows are paired with high confidence and eligible for the tri-modal branch, 2,240 rows are transcript-only, and 12 rows are video-only. All rows are still witness-conservative, because `witness_identity_status` remains unresolved across the current manifest. That is fine at this stage because the manifest is showing grounded hearing rows, not finalized witness labels.",
        [
            "Use one paired ICTR row to show the ideal grounded case: `hear_992268674c264fe8`, `ICTR-98-41`, `01/06/2007`.",
            "Use one paired ICTY row to show the same pattern in the other tribunal: `hear_c149355188328b0b`, `IT-95-5/18`, `24/03/2016`.",
            "Use one transcript-only row to explain why not every hearing becomes tri-modal: `hear_278afb925fbdcba5`, `ICTR-98-41`, `01/03/2006`.",
            "Use one video-only row to explain why the manifest does not invent transcript evidence: `hear_0a3fd1f78da95a16`, `ICTR-98-41`, `01/04/2011`.",
        ],
    )

    add_section(
        doc,
        "Slide 15. Validated LegalMELD EDA",
        "Use this section while showing the validated tribunal utterance dataset under `data/processed/phase2/legalmeld_validated/`.",
        "This is the utterance-level dataset layer, not the hearing manifest. The validated export is the MELD-style corpus that comes after transcript segmentation, alignment, media extraction, and validation. At this stage the dataset has 130 utterances in total, spread across 2 cases and 4 hearings, with 4 distinct witnesses represented in the summary. Every one of the 130 rows is present as text, audio, and video, so it is genuinely multimodal at the row level. The alignment quality is intentionally separated into tiers: 12 utterances are high-confidence, 111 are medium-confidence, and 7 are low-confidence review rows. The split structure is also already set up, with 75 train rows, 10 dev rows, 38 test rows, and 7 review rows. Audio validation is clean across all 130 rows, and video validation is also clean, so the main thing left to explain is that the dataset is still small and is not yet the final Indian adaptation corpus. For Phase 2, this matters because it shows the pipeline already produces a usable multimodal training set, even if it is still bootstrap-scale. The 45.2 usable minutes come from the high- and medium-confidence rows, which makes them the best lower-bound figure to cite when discussing early multimodal experimentation.",
        [
            "Explain that this is the actual utterance-level tribunal dataset.",
            "Point out that all 130 rows already have transcript, audio, and video paths.",
            "Use the high, medium, and review examples to show the quality tiers.",
            "Mention that zero split leakage matters because the same hearing or witness should not appear in multiple partitions.",
            "Say that the 45.2 usable minutes are the conservative Phase 2 figure for usable multimodal minutes.",
        ],
    )

    add_table(
        doc,
        ["EDA view", "How to explain it in the call"],
        [
            [
                "Overall export",
                "Say that the tribunal branch currently has 130 utterances, and all 130 rows are valid for text, audio, and video, with a reproducible split structure.",
            ],
            [
                "Usable minutes",
                "Say that the usable Phase 2 training size is 45.2 minutes, using the high- and medium-confidence rows rather than the low-confidence review rows.",
            ],
            [
                "High-confidence utterance",
                "Use `hear_7b2686bf1e5608e9_utt00002` to show a grounded row with matched alignment and media paths, which is the kind of row Phase 2 should trust most.",
            ],
            [
                "Medium-confidence utterance",
                "Use `hear_7b2686bf1e5608e9_utt00001` to show that the pipeline keeps usable rows even when alignment is fuzzy, which helps the corpus stay larger without pretending the rows are perfect.",
            ],
            [
                "Review utterance",
                "Use `hear_7b2686bf1e5608e9_utt00005` to show that fallback rows are surfaced for manual review instead of being hidden, which protects the Phase 2 dataset from noisy supervision.",
            ],
        ],
    )

    add_section(
        doc,
        "Slide 16. Witness Controlled Validation Comparison",
        "Use this compact section to explain the 4 anchor hearings versus the 6 manually promoted hearings.",
        "This slide is the bridge between the witness-only planning layer and the smaller controlled validation pass. The anchor hearings are the four hearings that are already validated in the witness utterance subset, so they are proof that the pipeline already works at the utterance level. The promoted hearings are six additional hearings selected manually from the broader witness hearing plan because they had strong discovery scores, traceable witness identities, and useful testimony coverage. The most important student-level point is that the promoted rows are not being called final training data. They are a controlled validation group. That means they are still part of the planning story, and the purpose is to inspect whether the witness-only path remains clean when we expand the set a little. Two promoted hearings needed label cleanup because the witness label mixed identity text with an examination heading or a public-name token. After normalization, the cleaned witness key is stored separately, so the hearing is kept but the label ambiguity is removed.",
        [
            "Say that the anchor hearings are already validated and therefore prove the pipeline.",
            "Say that the promoted hearings are manually selected, not automatically trusted.",
            "Use the phrase 'controlled validation pass' to show that this is a cautious expansion.",
            "Explain that the table is about prioritization and inspection, not about final corpus size.",
            "Mention that the two cleaned rows now carry a separate normalized witness key.",
        ],
    )
    add_table(
        doc,
        ["Group", "Hearing ID", "Date", "Witness", "Reason selected"],
        [
            ["Anchor", "hear_2ef83c852251d65c", "21/06/2004", "ANTIPAS NYANJWA", "Already validated in the utterance-level witness subset."],
            ["Anchor", "hear_7b2686bf1e5608e9", "03/05/2005", "DM190 | EMMANUEL NERETSE", "Already validated in the utterance-level witness subset."],
            ["Anchor", "hear_8280237faba9d96f", "16/09/2004", "FILIP REYNTJENS", "Already validated in the utterance-level witness subset."],
            ["Anchor", "hear_f5d485391c1d04cc", "19/01/2004", "ROMÉO DALLAIRE", "Already validated in the utterance-level witness subset."],
            ["Promoted", "hear_dea210cdb4c728e0", "03/02/2004", "BRENT BEARDSLEY | Exanination-in-chief by Mr. White", "Manually promoted from the broader hearing discovery plan for controlled validation."],
            ["Promoted", "hear_8a80539e19e2df44", "05/07/2005", "DK120 | MATHIEU NGIRUMPATSE", "Manually promoted from the broader hearing discovery plan for controlled validation."],
            ["Promoted", "hear_da77e01d6076db4b", "18/06/2003", "OMAR SERUSHAGO", "Manually promoted from the broader hearing discovery plan for controlled validation."],
            ["Promoted", "hear_b3fafce46640756e", "21/09/2006", "ALOYS NTABAKUZE", "Manually promoted from the broader hearing discovery plan for controlled validation."],
            ["Promoted", "hear_23dd766aab93f457", "26/01/2004", "ROMÉO DALLAIRE", "Manually promoted from the broader hearing discovery plan for controlled validation."],
            ["Promoted", "hear_ebd43a70140ec251", "27/01/2004", "ROMÉO DALLAIRE", "Manually promoted from the broader hearing discovery plan for controlled validation."],
        ],
    )

    add_section(
        doc,
        "Slide 17. Phase 2 Multimodal Summary Box",
        "Use this compact section to point at the validated artifact and the manifest file together.",
        "This slide is a quick summary of the Phase 2 tribunal bootstrap output. The artifact file is `data/processed/phase2/legalmeld_validated/legalmeld_metadata_validated.csv`, which is the validated row-level export that actually carries the multimodal samples. The manifest file is `data/processed/phase2/hearing_manifest.csv`, which is the grounded hearing layer that feeds the utterance-level export. The key message is that the export is not a random collection of clips. It is a verified pipeline output that starts from the hearing manifest and ends in a fully multimodal row-level artifact. The clean number to quote is 45.2 usable minutes, because that is the conservative training-quality subset from the high- and medium-confidence rows. The other clean number to quote is 130 valid rows, because every validated row has text, audio, and video available.",
        [
            "Say that the artifact file is the actual multimodal export.",
            "Say that the manifest file is the grounded hearing layer that feeds it.",
            "Use 45.2 minutes as the conservative usable figure for early Phase 2 work.",
            "Use 130 valid rows to show that all three modalities are present in the export.",
        ],
    )

    add_section(
        doc,
        "Slide 18. Validated Multimodal Quality EDA",
        "Use this section to explain which modality-quality columns are informative in the current export and which ones are not yet useful.",
        "This slide is only worth showing if you want to talk about quality, not just size. The useful fields right now are the audio and split fields. `audio_present` is YES for all 130 rows, `audio_validation_status` is valid for all 130 rows, `video_quality_status` is VALID for all 130 rows, and `sample_rate` is 16000 for all 130 rows. That means the export is genuinely multimodal and technically consistent. The audio quality metrics are also informative: `audio_rms` ranges from about 0.0037 to 0.0609, `silence_ratio` ranges from about 0.2928 to 0.9726, and `clipping_ratio` is 0.0 for every row. The split field is useful because it shows the leakage-aware partitioning: 75 train, 10 dev, 38 test, and 7 review. The face, emotion, and credibility columns are not useful yet in this export because they are still constant or blank, so they should be treated as future work rather than evidence of additional annotation.",
        [
            "Say that the useful fields are the audio metrics, the validation flags, and the split field.",
            "Say that the face, emotion, and credibility columns are not yet informative in the current export.",
            "Use the audio metric ranges to show that the clips are technically usable but still need careful quality interpretation.",
            "Mention that clipping is zero and the sample rate is stable at 16 kHz.",
        ],
    )

    add_table(
        doc,
        ["Field", "Current EDA result", "Why it matters now"],
        [
            ["audio_present", "YES for all 130 rows", "Confirms every validated row has audio."],
            ["audio_rms", "0.0037 to 0.0609, mean 0.0288", "Shows the energy range of the clips."],
            ["silence_ratio", "0.2928 to 0.9726, mean 0.6231", "Shows how much silence remains in each clip."],
            ["clipping_ratio", "0.0 for all 130 rows", "Confirms there is no clipping issue in the current export."],
            ["sample_rate", "16000 for all 130 rows", "Shows a consistent model-friendly sampling rate."],
            ["audio_validation_status", "valid for all 130 rows", "Confirms the audio side is usable for Phase 2."],
            ["video_quality_status", "VALID for all 130 rows", "Confirms the video side is usable for Phase 2."],
            ["split", "train 75, dev 10, test 38, review 7", "Shows how the export is partitioned."],
        ],
    )

    add_section(
        doc,
        "How The Controlled Validation Score Is Computed",
        "Explain this only as a prioritization score, not as a truth label.",
        "The score used for the controlled validation subset is a simple heuristic ranking function. It is not an annotation and it is not a model prediction. The purpose is to sort witness hearings so the strongest ones are easy to inspect first. The formula implemented in the subset builder is: score = estimated_testimony_minutes + 25 × discovery_row_count, then add 20 if the hearing has both protected-code and public-name evidence, add 15 if it has only protected-code evidence, add 10 if it has only public-name evidence, add 5 if discovery_row_count is greater than 1, and subtract 20 if estimated_testimony_minutes is zero or below. In the code, the result is rounded to two decimal places. You should describe this as a ranking heuristic for manual review, not as a scientific ground-truth score. The meaning is simple: longer hearings, stronger discovery evidence, and more traceable witness identity signals rise to the top.",
        [
            "Say that the score is only for prioritizing hearings.",
            "Say that it prefers hearings with more discovery rows and more testimony minutes.",
            "Explain that protected/public witness evidence adds confidence to the rank.",
            "Say that zero-minute hearings are penalized because they are less useful for controlled validation.",
        ],
    )

    add_table(
        doc,
        ["Formula part", "What it means in practice"],
        [
            ["estimated_testimony_minutes", "The main size term; larger hearings tend to be more useful for inspection."],
            ["25 × discovery_row_count", "Rewards hearings that produced more discovery rows in the broader plan."],
            ["+20 / +15 / +10 witness-status bonus", "Rewards more traceable witness identity patterns."],
            ["+5 for discovery_row_count > 1", "Gives a small extra preference to hearings with more than one discovery row."],
            ["-20 for zero-minute hearings", "Pushes weak or empty-looking entries down the list."],
            ["Rounded to 2 decimals", "Keeps the score readable in the CSV and the summary report."],
        ],
    )

    add_section(
        doc,
        "Manual Review Of The Controlled Validation Subset",
        "Explain how the hearing-by-hearing review was done and what was flagged.",
        "This is the manual inspection step that follows the controlled validation score. The important thing to say is that the inspection was metadata-based, not media-based. I read the controlled validation CSV row by row and compared the validation track, the witness label, the score, the discovery row count, and the manual note. For the four anchor hearings, the answer is simple: they are already validated in the utterance-level witness subset, so they were kept as verified anchors. For the six promoted hearings, I checked whether the witness label looked clean and whether the row still had a strong discovery signal. After cleanup, all six promoted hearings stayed in the controlled set, and the two mixed-label rows were normalized into a separate cleaned witness key instead of being discarded. That is the right kind of manual action at this stage: keep the hearing, keep the evidence, but separate the cleaned witness key from the original raw label so the dataset stays traceable.",
        [
            "Say that this manual review was done from the CSV fields first, not from the raw video.",
            "Say that the anchor rows were confirmed as already validated.",
            "Say that the two mixed labels were cleaned rather than dropped.",
            "Say that the hearing is still kept, but the cleaned witness key should be used for broader reuse.",
        ],
    )
    add_table(
        doc,
        ["Hearing ID", "Track", "Manual review status", "What I looked at", "Manual note"],
        [
            ["hear_2ef83c852251d65c", "anchor", "verified_anchor", "Validation track, witness row already proven, score, and notes", "Anchor hearing kept because the witness utterance subset already validates it."],
            ["hear_7b2686bf1e5608e9", "anchor", "verified_anchor", "Validation track, witness row already proven, score, and notes", "Anchor hearing kept because the witness utterance subset already validates it."],
            ["hear_8280237faba9d96f", "anchor", "verified_anchor", "Validation track, witness row already proven, score, and notes", "Anchor hearing kept because the witness utterance subset already validates it."],
            ["hear_f5d485391c1d04cc", "anchor", "verified_anchor", "Validation track, witness row already proven, score, and notes", "Anchor hearing kept because the witness utterance subset already validates it."],
            ["hear_dea210cdb4c728e0", "promoted", "flagged_for_label_cleanup", "Mixed witness label and examination-heading text", "Label includes an examination heading; clean the witness label before using it as a normalized witness key."],
            ["hear_8a80539e19e2df44", "promoted", "flagged_for_label_cleanup", "Mixed protected-code and public-name signal", "Witness label mixes a protected code and a public name; normalize the witness key carefully."],
            ["hear_da77e01d6076db4b", "promoted", "reviewed_promoted", "Score, witness label, discovery rows, and manual note", "Promoted hearing looks consistent and remains in the controlled validation set."],
            ["hear_b3fafce46640756e", "promoted", "reviewed_promoted", "Score, witness label, discovery rows, and manual note", "Promoted hearing looks consistent and remains in the controlled validation set."],
            ["hear_23dd766aab93f457", "promoted", "reviewed_promoted", "Score, witness label, discovery rows, and manual note", "Promoted hearing looks consistent and remains in the controlled validation set."],
            ["hear_ebd43a70140ec251", "promoted", "reviewed_promoted", "Score, witness label, discovery rows, and manual note", "Promoted hearing looks consistent and remains in the controlled validation set."],
        ],
    )
    add_para(
        doc,
        "How to explain the manual method: 'I did not treat the score as truth. I used it to sort the hearings, then I read the controlled-validation CSV by hearing ID and checked whether the witness label looked clean, whether the hearing already had anchor validation, and whether the manual note supported the selection. If the label mixed witness identity with exam-heading text or public-name text, I normalized it into a separate cleaned witness key but kept the hearing in the controlled set.'",
    )

    add_para(
        doc,
        "Field-by-field explanation: The purpose of this section is to help you explain the quality columns at a student level. The point is not only to say what the column contains, but why the column exists in the pipeline and how it should be used when deciding whether a row is good enough for multimodal training. Some columns are hard measurement fields, some are categorical validation labels, and some are future annotation placeholders. The current export is useful exactly because it separates these ideas instead of collapsing them into one vague score.",
    )
    add_table(
        doc,
        ["Field", "Unit / value type", "Technical meaning", "Practical use in Phase 2"],
        [
            [
                "audio_present",
                "Binary flag: YES / NO",
                "Tells you whether an audio track was successfully found and decoded for the clip.",
                "Use it as the first filter before any audio-based modeling. If this is NO, the row should not enter the multimodal set.",
            ],
            [
                "audio_rms",
                "Unitless normalized RMS amplitude, roughly 0 to 1",
                "RMS means root-mean-square energy. The code computes it from 16-bit PCM samples and divides by 32768, so it is a normalized loudness proxy rather than a physical unit like volts or decibels.",
                "Use it to compare loudness across clips and to spot clips that are almost silent or unusually strong. It is a quality indicator, not a label.",
            ],
            [
                "silence_ratio",
                "Fraction from 0 to 1",
                "The proportion of samples below a very small amplitude threshold. Higher values mean more of the clip is silence or near-silence.",
                "Use it to detect clips that are mostly silence or that contain only a short usable utterance. Very high values can mean the clip needs better trimming.",
            ],
            [
                "clipping_ratio",
                "Fraction from 0 to 1",
                "The proportion of samples near the maximum possible amplitude. Clipping happens when audio gets too loud and the waveform is flattened at the limits.",
                "Use it to reject or inspect distorted audio. In this export it is 0.0 everywhere, which is good because it means no clipping is visible.",
            ],
            [
                "sample_rate",
                "Hertz (samples per second)",
                "How many audio samples are stored each second. The current export uses 16000 Hz, which is a standard speech-processing rate.",
                "Use it to confirm that the audio is in a stable, model-friendly format. Mixed sample rates can complicate training and alignment.",
            ],
            [
                "audio_validation_status",
                "Categorical label",
                "The outcome of the audio-quality check, such as valid, silent, clipped, missing, corrupt, or unsupported format.",
                "Use it as the main gate for audio usability. A row with invalid audio should not be treated as clean multimodal training data even if text exists.",
            ],
            [
                "video_quality_status",
                "Categorical label",
                "The outcome of the video/media probe. In this export it is VALID for all rows, which means the video files are decodable and usable at a basic technical level.",
                "Use it as the first check that the video side is technically present. It does not yet prove that the witness face is visible.",
            ],
            [
                "split",
                "Categorical partition label",
                "Shows whether the row belongs to train, dev, test, or review. This is not a media-quality field; it is a dataset-usage field.",
                "Use it to keep leakage under control. Train/dev/test should stay separated by hearing or witness group, and review rows should stay out of the initial training set.",
            ],
            [
                "face_detected",
                "Categorical placeholder",
                "Would indicate whether a face detector found a face in the frame. In the current export it is UNKNOWN for all rows.",
                "Theoretical use: facial-cue quality control. Current use: not informative yet, so do not use it to claim witness visibility.",
            ],
            [
                "face_visible_ratio",
                "Fraction from 0 to 1, currently blank",
                "Would measure how much of the clip shows a visible face. A higher value would mean the face is visible for more of the clip.",
                "Useful later if the project wants emotion or credibility cues from facial expression. For now it is blank, so it should not be interpreted.",
            ],
            [
                "shot_type",
                "Categorical placeholder",
                "Would describe the camera framing, such as close-up, wide shot, or courtroom overview. In the current export it is UNKNOWN.",
                "Useful later to filter clips where the witness is on screen versus clips where the shot is too wide or the speaker is off camera.",
            ],
            [
                "speaker_visible",
                "Categorical placeholder",
                "Would tell you whether the current speaker is visibly on screen. The current export keeps this as UNKNOWN.",
                "Useful later for deciding whether a clip is suitable for facial or gesture-based modeling. Not useful as an evidence field yet.",
            ],
            [
                "emotion",
                "Annotation placeholder",
                "Currently blank. This is reserved for future emotion labels.",
                "Do not treat it as a real annotation field yet. It is intentionally empty until a human or model annotates it later.",
            ],
            [
                "credibility",
                "Annotation placeholder",
                "Currently blank. This is reserved for future credibility or trustworthiness labels.",
                "Do not use it in the current EDA as if it already contains ground truth. It is there so the schema can support later annotation.",
            ],
        ],
    )
    add_para(
        doc,
        "How to summarize this in one sentence: the audio fields tell you whether the sound is technically usable, the video field tells you whether the clip is technically decodable, the split field tells you whether the row is safe to use for training, and the placeholder face/emotion/credibility fields tell you what the schema is preparing for later rather than what is already ready today.",
    )

    add_section(
        doc,
        "Representative Row: hear_7b2686bf1e5608e9_utt00002",
        "Use this example to explain what a single validated row actually means in Phase 2.",
        "This row is one of the cleanest examples in the current export. It is a high-confidence tribunal utterance that connects the hearing-level manifest to a specific aligned audio-video-text sample. The important thing is that each field tells you something different. Some fields define where the row came from, some define how trustworthy the alignment is, some define the audio/video quality, and some define how the row should be used in training or evaluation. For this row, the values say the following: it comes from the ICTR tribunal, case ICTR-98-41, Bagosora et al. Trial, hearing date 03/05/2005, witness code DM190, and speaker role Witness. The row is aligned with `matched` status using `asr_exact`, with HIGH confidence. The clip starts at 00:36:39.310 and ends at 00:37:09.310, giving a 30-second segment. The audio exists, the audio is valid, the video is valid, the sample rate is 16000 Hz, the clipping ratio is 0.0, and the row is placed in the test split with `witness_disjoint` grouping. That means the row is suitable as a held-out evaluation example, not just a training row. The row is useful for Phase 2 because it shows the full pipeline working end to end: grounded case metadata, aligned transcript span, valid audio, valid video, and a leakage-aware split decision.",
        [
            "Use this row to show what a grounded multimodal training/evaluation sample looks like.",
            "Explain that the hearing is real, the witness code is preserved, and the row is not a placeholder.",
            "Point out that `matched` and `HIGH` make this a stronger row than the fuzzy or fallback cases.",
            "Mention that the row is in the `test` split, so it is especially important for evaluation discipline.",
        ],
    )

    add_table(
        doc,
        ["Field", "Value in this row", "How to explain it"],
        [
            ["utterance_id", "hear_7b2686bf1e5608e9_utt00002", "Unique identifier for the exact utterance row."],
            ["hearing_id", "hear_7b2686bf1e5608e9", "Links the utterance back to the parent hearing."],
            ["tribunal", "ICTR", "Shows the row came from the Rwanda tribunal branch."],
            ["case_number", "ICTR-98-41", "Grounded tribunal case number."],
            ["case_family", "Bagosora et al. Trial", "Human-readable case-family label."],
            ["hearing_date", "03/05/2005", "The date of the hearing that produced the utterance."],
            ["witness_id", "DM190", "Protected witness code used for grouping and traceability."],
            ["speaker_role", "Witness", "Shows this utterance is part of testimony, not argument."],
            ["speaker_name", "Answer", "Transcript-side label for the speaking turn."],
            ["examination_type", "unknown", "The pipeline did not confidently classify direct/cross/redirect here."],
            ["utterance_text", "Thank you, Mr. President. I entered or joined the Rwandan army when I entered the ESM in August 1975. I graduated from that school in 1978.", "The normalized transcript text for the utterance."],
            ["start_time", "00:36:39.310", "Clip start point in the source media."],
            ["end_time", "00:37:09.310", "Clip end point in the source media."],
            ["duration_ms", "30000", "Clip duration in milliseconds."],
            ["video_clip", "clips/hear_7b2686bf1e5608e9/utt00002.mp4", "Relative path to the extracted video clip."],
            ["audio_clip", "audio/hear_7b2686bf1e5608e9/utt00002.wav", "Relative path to the extracted audio clip."],
            ["transcript_source", "transcripts/hear_7b2686bf1e5608e9.txt", "Transcript file used to create the utterance."],
            ["alignment_status", "matched", "Shows the transcript and alignment step succeeded."],
            ["alignment_score", "0.346", "Internal alignment score; useful for comparing confidence, not as a label."],
            ["alignment_method", "asr_exact", "Alignment method used by the pipeline."],
            ["alignment_confidence", "HIGH", "The strongest confidence tier in the current export."],
            ["asr_text", "ASR text is present and differs from the transcript span", "The ASR output gives another view of the same clip for matching and debugging."],
            ["transcript_text_normalized", "Same transcript content normalized for matching", "Useful for matching against ASR and for preprocessing."],
            ["text_similarity", "0.161", "A moderate similarity score, showing the ASR and transcript are not identical but still aligned."],
            ["clip_duration_seconds", "30.0", "Human-readable duration of the clip."],
            ["word_timestamp_count", "86", "Number of word-level timestamps found by the aligner."],
            ["manual_review_required", "NO", "The pipeline considers this row usable without manual review."],
            ["split_group_id", "DM190", "Group identifier used to keep related rows together."],
            ["split_strategy", "witness_disjoint", "Shows the row was split by witness to reduce leakage."],
            ["quality_tier", "B", "Usable row for model work, but still not perfect enough to call A."],
            ["audio_present", "YES", "Audio was successfully found and decoded."],
            ["audio_rms", "0.031006", "Normalized audio energy level; a unitless loudness proxy."],
            ["silence_ratio", "0.684521", "About 68.5% of samples are silent or near-silent."],
            ["clipping_ratio", "0.0", "No clipping distortion in the clip."],
            ["sample_rate", "16000", "Audio was stored at 16 kHz."],
            ["audio_validation_status", "valid", "Audio passed the quality check."],
            ["face_detected", "UNKNOWN", "No reliable face-detection outcome yet."],
            ["face_visible_ratio", "blank", "No face-visibility metric is available yet."],
            ["shot_type", "UNKNOWN", "No camera-shot classification yet."],
            ["speaker_visible", "UNKNOWN", "No verified speaker-visibility flag yet."],
            ["video_quality_status", "VALID", "Video is technically usable."],
            ["emotion", "blank", "Emotion labels are not added yet."],
            ["credibility", "blank", "Credibility labels are not added yet."],
            ["split", "test", "Held out for evaluation rather than training."],
        ],
    )
    add_para(
        doc,
        "How to read this row in the project: the first group of fields tells you that the sample is real and traceable, the second group tells you that the transcript and media are aligned well enough to use, the third group tells you that the audio and video are technically valid, and the split fields tell you that the row is safely reserved for evaluation. In practical Phase 2 terms, this is the kind of row you would trust for a test-time multimodal sample, while the low-confidence review rows should stay out of the first training pass. In theory, this is the right shape of supervision for a multimodal courtroom model because it connects language, audio, and video to a grounded legal event instead of an invented label.",
    )
    add_para(
        doc,
        "One-line summary for this row: hear_7b2686bf1e5608e9_utt00002 is a grounded, high-confidence ICTR witness utterance with valid audio and video, held out in the test split, so it is best understood as a strong evaluation sample for Phase 2 multimodal alignment rather than a final labeled training target.",
    )

    add_section(
        doc,
        "Slide 18. Clean Implementation Checklist for What Remains",
        "Make it clear that the next work should improve quality, diversity, and validation before expanding more widely.",
        "The checklist slide is about restraint. The best next step is not just to collect more data. It is to improve the current bootstrap corpus first. That means reducing fallback alignments, growing case diversity, preserving hearing-disjoint or witness-disjoint splits, and improving the usefulness of the clips that are already exported. It also means keeping the Indian acquisition pack as the next stage, not as a distraction from fixing the current pipeline. The mentor should hear that you understand the difference between scaling the wrong thing and improving the right thing, including the choice of alignment backend.",
        [
            "Say that better quality is more useful than more hours if the utterance extraction is still noisy.",
            "Mention that split leakage must stay at zero.",
            "Explain that the Indian acquisition pack should be built only after the bootstrap corpus quality is stable.",
        ],
    )

    add_section(
        doc,
        "Slide 19. Realistic Indian Sources",
        "Explain what each Indian source family can actually contribute, and keep the scope conservative.",
        "This slide should sound practical, not aspirational. Supreme Court livestreams are valuable because they provide authentic Indian legal speech, but they are mainly arguments rather than witness testimony, so they are best for language and procedural adaptation. High Court livestreams have a similar value. District court material is the closest to real witness testimony, but public multimodal access is limited, so it should be used only where the records are actually available. Mock trials and judicial academies are useful because they often contain witness-style interactions and examination structure even if they are simulated. Legal education videos are useful for legal language only, not for witness emotion modeling. The point is to use each source family for what it realistically offers.",
        [
            "Use the phrase 'best use' rather than 'best source' to keep the claim honest.",
            "Say that some sources help with language adaptation, while others help with speaker-role structure.",
            "Make it clear that district court material is closest to testimony but is usually the hardest to obtain publicly.",
        ],
    )
    add_table(
        doc,
        ["Source family", "How to explain its best use"],
        [
            [
                "Supreme Court livestreams",
                "Best for Indian legal-language adaptation and procedural speech patterns.",
            ],
            [
                "High Court livestreams",
                "Best for additional Indian legal discourse variety and courtroom style.",
            ],
            [
                "District court material",
                "Best when you can find public multimodal testimony, because it is closest to real witness examination.",
            ],
            [
                "Mock trials and judicial academies",
                "Best for direct/cross-examination structure and speaker-role adaptation.",
            ],
            [
                "Legal education videos",
                "Best for legal-language adaptation only, not for witness emotion modeling.",
            ],
        ],
    )

    add_section(
        doc,
        "Slide 20. Why a Hybrid Corpus Is the Correct Design",
        "Explain why the dissertation needs both a tribunal bootstrap corpus and an Indian adaptation corpus.",
        "The hybrid-corpus slide is where you defend the research design. The tribunal corpus gives supervision, scale, and actual witness-testimony structure. The Indian corpus gives domain adaptation for accent, legal procedure, and local courtroom language. Mock trials help because they supply structure where public real testimony is scarce. If you say this clearly, the mentor will see that you are not overclaiming that any one source family solves the entire problem. You are instead building a two-stage methodology: prove the method on tribunal data, then adapt it to Indian sources.",
        [
            "Say that the tribunal corpus proves the method.",
            "Say that the Indian corpus proves transfer to the final research objective.",
            "Avoid implying that one source family can satisfy scale, witness metadata, and public availability all at once.",
        ],
    )

    add_section(
        doc,
        "Slide 21. Expected Outcomes",
        "State the outcome as a MELD-style utterance corpus plus a reusable Indian adaptation path, not as a completed model deployment.",
        "The expected outcome should be described as a practical dataset and methodology outcome. The aim is to create utterance-level rows with aligned transcript, audio, and video, not to keep collecting long hearing files. That gives the model a cleaner training unit and makes later emotion and credibility annotation feasible. The tribunal corpus will be the proof that the extraction pipeline works. The Indian corpus will be the final adaptation layer. This slide is also where you should say that the project is methodologically stronger when it uses smaller, grounded utterance samples than when it uses long noisy recordings.",
        [
            "Say that the main output is a clean utterance-level dataset, not just more hours of video.",
            "Mention that the dataset should support later emotion and credibility tasks.",
            "Keep the final Indian adaptation objective explicit.",
        ],
    )

    add_section(
        doc,
        "Slide 22. What I Want to Discuss With You",
        "Finish by asking for a scope and sequence check, not a blanket approval of all future data collection.",
        "The final slide should turn the discussion into a constructive mentor checkpoint. You want feedback on whether the tribunal bootstrap pipeline is clear enough as a proof of method, whether the Indian adaptation scope should remain explicitly part of Phase 2, whether the next effort should go into alignment quality or source expansion, and whether the literature reading order makes sense. This keeps the call focused on planning discipline rather than on collecting more data for its own sake.",
        [
            "Ask whether the tribunal bootstrap proof is convincing.",
            "Ask whether the Indian adaptation branch should remain in Phase 2 scope.",
            "Ask whether quality improvement should come before any larger source expansion.",
        ],
    )

    doc.add_heading("Quick Reference: Current Phase 2 Numbers", level=1)
    add_table(
        doc,
        ["Metric", "Current value"],
        [
            ["Verified cases represented", s["verified_cases"]],
            ["Hearings represented", s["hearings_represented"]],
            ["Utterances exported", s["utterances"]],
            ["High-confidence alignments", s["high_confidence"]],
            ["Medium-confidence alignments", s["medium_confidence"]],
            ["Low-confidence alignments", s["low_confidence"]],
            ["Train rows", s["train_rows"]],
            ["Dev rows", s["dev_rows"]],
            ["Test rows", s["test_rows"]],
            ["Split leakage violations", s["split_leakage"]],
            ["Distinct witnesses", s["distinct_witnesses"]],
            ["Audio-valid clips", s["audio_valid"]],
            ["Video-valid clips", s["video_valid"]],
        ],
    )
    add_para(
        doc,
        "Use these numbers only as a snapshot of the current validated run. Do not present them as the final dissertation scale. The main point is that the pipeline is real, reproducible, and already proving the method on tribunal data.",
    )

    doc.add_heading("Closing Practice Script", level=1)
    add_para(
        doc,
        "A clean way to close the meeting is: 'The project is now at the point where the verification-first tribunal pipeline works end to end. The next step is not to over-collect raw hours. The next step is to improve utterance quality, preserve leakage-safe splits, and then carry the same manifest logic into the Indian adaptation layer.'",
    )

    doc.add_heading("Reminder On Tone", level=1)
    add_bullets(
        doc,
        [
            "Be precise: verified does not mean final; it means grounded.",
            "Be conservative: Indian adaptation is a planned next layer, not a completed corpus.",
            "Be technical: the pipeline works because of specific scripts, manifests, and validation steps.",
            "Be honest about limits: the current dataset is small, but it already proves the method.",
        ],
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
