from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_Phase2_Dataset_Preparation_SOP.docx"
CANDIDATE_LEDGER = ROOT / "data" / "phase2" / "source_manifests" / "case_candidate_ledger.csv"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def add_code_map_table(doc: Document) -> None:
    add_table(
        doc,
        ["SOP step", "Script / module / function"],
        [
            [
                "Crawl and inspect source sites",
                "Candidate-ledger URLs in `data/phase2/source_manifests/case_candidate_ledger.csv`; search and crawl logic lives in `phase2/crawl_official_ucr_case_pages.py` and related run scripts.",
            ],
            [
                "Verify case metadata and reject placeholders",
                "`phase2/enrich_case_ledger_from_ucr_site.py` -> `resolve_case_verification()` and `validation_summary()`; wrapper: `phase2/run_enrich_case_ledger_from_ucr_site.sh`.",
            ],
            [
                "Promote only grounded rows into the inventory",
                "`phase2/build_ucr_case_inventory.py` -> `_eligible_rows()`, `_record_identity_key()`, and `resolve_case_verification()`; wrapper: `phase2/run_build_ucr_case_inventory.sh`.",
            ],
            [
                "Split inventory into media-specific manifests",
                "`phase2/split_ucr_inventory_by_media_type.py` and `phase2/run_split_ucr_inventory_by_media_type.sh`.",
            ],
            [
                "Filter TAP-bearing candidates from the expanded manifest",
                "`phase2/filter_tap_candidates_from_expanded_manifest.py` and `phase2/run_filter_tap_candidates_from_expanded_manifest.sh`.",
            ],
            [
                "Build hearing and witness manifests",
                "`phase2/build_hearing_manifest.py`, `phase2/build_witness_manifest.py`, `phase2/hearing_witness_manifest_utils.py`; wrapper: `phase2/run_build_hearing_witness_manifests.sh`.",
            ],
            [
                "Keep placeholder witness rows inert",
                "The SOP policy on `witness_harvest_manifest.csv` is enforced before corpus planning and download planning.",
            ],
            [
                "Build utterance-level LegalMELD samples",
                "`phase2/build_legalmeld_dataset.py` -> `segment_transcript()`, `_align_utterances()`, `_extract_segment()`, `_audio_metrics()`; wrapper: `phase2/run_build_legalmeld_dataset_validated.sh`.",
            ],
            [
                "Inspect the strict download index and validate unique MP4s",
                "`phase2/inspect_ucr_case_videos_strict_index.py` and `phase2/run_inspect_ucr_case_videos_strict_index.sh`; the fallback validator is `scripts/check_mp4_fallback.py`.",
            ],
            [
                "Select the next controlled TAP-bearing candidate",
                "`phase2/select_next_tap_verification_candidate.py` and `phase2/run_select_next_tap_verification_candidate.sh`.",
            ],
            [
                "Compare the expanded planning manifest against the verified inventory",
                "`phase2/compare_expanded_planning_to_verified_inventory.py` and `phase2/run_compare_expanded_planning_to_verified_inventory.sh`.",
            ],
            [
                "Discover official tribunal source-broadening candidates",
                "`phase2/discover_official_tribunal_source_broadening_candidates.py`, `phase2/run_discover_official_tribunal_source_broadening_candidates.sh`, and `phase2/run_refresh_tribunal_source_broadening.sh`.",
            ],
            [
                "Rank tribunal families for source broadening",
                "`phase2/rank_tribunal_source_broadening_candidates.py` and `phase2/run_rank_tribunal_source_broadening_candidates.sh`.",
            ],
            [
                "Build the broadened tribunal source manifest",
                "`phase2/build_broadened_tribunal_source_manifest.py` and `phase2/run_build_broadened_tribunal_source_manifest.sh`.",
            ],
            [
                "Append verified TAP-bearing case additions",
                "`data/phase2/source_manifests/verified_tap_case_additions.csv`, `phase2/build_broadened_tribunal_source_manifest.py`, and `phase2/run_build_broadened_tribunal_source_manifest.sh`.",
            ],
            [
                "Prepare for Indian adaptation",
                "The SOP describes the future Indian acquisition pack, but the reusable logic is the same candidate/verified/manifest pipeline already implemented for the tribunal bootstrap layer.",
            ],
        ],
    )


def add_pipeline_order_table(doc: Document) -> None:
    add_table(
        doc,
        ["Stage", "Produced by", "Used for", "Student-level meaning"],
        [
            [
                "1. case_candidate_ledger.csv",
                "Manual curation",
                "Starting candidate list",
                "The first shortlist of cases to inspect; it is not proof of usable media yet.",
            ],
            [
                "2. case_candidate_ledger_ucr_enriched.csv",
                "`phase2/enrich_case_ledger_from_ucr_site.py` via `run_enrich_case_ledger_from_ucr_site.sh`",
                "Verification layer",
                "The checked version of the shortlist; it shows what the official UCR records support.",
            ],
            [
                "3. verified_case_inventory.csv",
                "`phase2/build_ucr_case_inventory.py` via `run_build_ucr_case_inventory.sh`",
                "Canonical case inventory",
                "The first file that is trusted for corpus planning because it contains grounded records.",
            ],
            [
                "3b. verified_tap_case_additions.csv",
                "Manual verified additions plus `phase2/build_broadened_tribunal_source_manifest.py`",
                "Verified TAP-bearing additions",
                "A small reusable input that adds cases already confirmed to have public TAP recordings.",
            ],
            [
                "4. hearing_manifest.csv",
                "`phase2/build_hearing_manifest.py` via `run_build_hearing_witness_manifests.sh`",
                "Hearing/session grouping",
                "The hearing-level bridge between case records and utterance-level processing.",
            ],
            [
                "5. witness_harvest_manifest_resolved.csv",
                "`phase2/build_witness_manifest.py` via `run_build_hearing_witness_manifests.sh`",
                "Resolved witness rows",
                "The witness-level layer that stays safe for protected codes and later utterance selection.",
            ],
            [
                "6. legalmeld_metadata_validated.csv",
                "`phase2/build_legalmeld_dataset.py` via `run_build_legalmeld_dataset_validated.sh`",
                "Utterance-level master CSV",
                "The MELD-style file where each row is one aligned utterance with timing and modality fields.",
            ],
            [
                "7. alignment_review_sample.csv",
                "`phase2/build_legalmeld_dataset.py`",
                "QA sample",
                "A small review file that helps check whether alignment and extraction behaved correctly.",
            ],
            [
                "8. dataset_quality_summary.json",
                "`phase2/build_legalmeld_dataset.py`",
                "Validation summary",
                "The summary that tells you how many utterances, alignments, clips, and split rows were produced.",
            ],
            [
                "9. train.csv / dev.csv / test.csv",
                "`phase2/build_legalmeld_dataset.py`",
                "Final split files",
                "The actual train/dev/test outputs used for model work after validation is complete.",
            ],
            [
                "10. strict index inspection outputs",
                "`phase2/inspect_ucr_case_videos_strict_index.py` via `run_inspect_ucr_case_videos_strict_index.sh`",
                "Unique-file and row-source view",
                "The inspection layer that shows which rows reuse the same file and which source manifest entries they came from.",
            ],
        ],
    )


def load_candidate_sources(limit: int = 5) -> list[list[str]]:
    if not CANDIDATE_LEDGER.exists():
        return []
    rows: list[list[str]] = []
    seen: set[tuple[str, str]] = set()
    with CANDIDATE_LEDGER.open(newline="", encoding="utf-8-sig") as f:
        for row in csv.DictReader(f):
            source_url = (row.get("source_url") or "").strip()
            inventory_url = (row.get("inventory_search_url") or "").strip()
            if not source_url and not inventory_url:
                continue
            key = (source_url, inventory_url)
            if key in seen:
                continue
            seen.add(key)
            rows.append(
                [
                    row.get("tribunal", ""),
                    row.get("case_family", ""),
                    row.get("case_number", ""),
                    source_url,
                    inventory_url,
                ]
            )
            if len(rows) >= limit:
                break
    return rows


def build_doc() -> Document:
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LegalMemoCMT Phase 2 Dataset Preparation SOP for Indian Courtroom Testimony")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Current working SOP for the tribunal bootstrap ledger, verified inventory, grounded hearing/witness manifests, transcript-first discovery, balanced selection, and the planned Indian courtroom adaptation layer."
    )
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12.5)

    add_para(
        doc,
        "This document is the practical Phase 2 preparation guide for the scripts that are actually being used right now. It focuses on the curated case ledger, the corrected UCR enrichment pass, the verified inventory stage, the grounded hearing and witness manifest builders, and the tri-modal validation pass that checks real transcript/video records before any larger download step. It also records what must be ignored for the time being: the placeholder witness manifest, placeholder case rows, and any page-shell heuristics that are not backed by extracted record data. The tribunal corpus is the bootstrap layer; the Indian courtroom corpus is the adaptation target.",
    )

    doc.add_heading("Source Sites And Download Flow", level=1)
    add_para(
        doc,
        "Before the corpus can be trusted, the pipeline has to show where the cases come from and how records are found. The basic order is: crawl or inspect the official tribunal source, search by case number or witness code inside the UCR inventory, verify that the page returns actual case-specific records, and only then download or promote the record into the verified inventory.",
    )
    add_table(
        doc,
        ["Tribunal", "Case family", "Case number", "Source site link", "Inventory/search link"],
        load_candidate_sources(),
    )
    add_numbered(
        doc,
        [
            "Crawl or inspect the official tribunal page or document source.",
            "Search the UCR inventory or official case page using the case number, case family, or witness code.",
            "Confirm that the page contains actual case-specific records rather than only generic headings.",
            "Promote only grounded results into the verified inventory.",
            "Download only after verification has established real transcript or media assets.",
        ],
    )
    add_bullets(
        doc,
        [
            "Do not count static headings as proof of records.",
            "Do not count placeholder manifests as corpus evidence.",
            "Do not estimate video hours from page shells.",
        ],
    )
    add_para(
        doc,
        "The code behind that flow is split across the enrichment script, the inventory builder, the hearing and witness manifest builders, and the LegalMELD exporter. The SOP is intentionally written so the student can trace each stage back to a named script or module.",
    )
    add_para(
        doc,
        "A concrete example of this flow is `Nahimana et al. (ICTR-99-52)`. The official ICTR page shows this as a completed media case with a large witness record, and the UCR case downloader can enumerate downloadable video recordings for it. The student-level explanation is: probe one recording first, confirm that the MP4 is real, then let the existing enrichment, inventory, split, hearing, witness, and LegalMELD scripts carry the case forward. If the probe returns no recordings, the case stays transcript-only and should not be promoted into the tri-modal branch.",
    )

    doc.add_heading("Pipeline Order", level=1)
    add_para(
        doc,
        "This table shows which CSV each stage produces and what that CSV means in student language.",
    )
    add_pipeline_order_table(doc)

    doc.add_heading("Code Map", level=1)
    add_table(
        doc,
        ["SOP step", "Script / module / function"],
        [
            [
                "Crawl and inspect source sites",
                "Candidate-ledger URLs in `data/phase2/source_manifests/case_candidate_ledger.csv`; search and crawl logic lives in `phase2/crawl_official_ucr_case_pages.py` and related run scripts.",
            ],
            [
                "Verify case metadata and reject placeholders",
                "`phase2/enrich_case_ledger_from_ucr_site.py` -> `resolve_case_verification()` and `validation_summary()`; wrapper: `phase2/run_enrich_case_ledger_from_ucr_site.sh`.",
            ],
            [
                "Promote only grounded rows into the inventory",
                "`phase2/build_ucr_case_inventory.py` -> `_eligible_rows()`, `_record_identity_key()`, and `resolve_case_verification()`; wrapper: `phase2/run_build_ucr_case_inventory.sh`.",
            ],
            [
                "Split inventory into media-specific manifests",
                "`phase2/split_ucr_inventory_by_media_type.py` and `phase2/run_split_ucr_inventory_by_media_type.sh`.",
            ],
            [
                "Build hearing and witness manifests",
                "`phase2/build_hearing_manifest.py`, `phase2/build_witness_manifest.py`, `phase2/hearing_witness_manifest_utils.py`; wrapper: `phase2/run_build_hearing_witness_manifests.sh`.",
            ],
            [
                "Keep placeholder witness rows inert",
                "The SOP policy on `witness_harvest_manifest.csv` is enforced before corpus planning and download planning.",
            ],
            [
                "Build utterance-level LegalMELD samples",
                "`phase2/build_legalmeld_dataset.py` -> `segment_transcript()`, `_align_utterances()`, `_extract_segment()`, `_audio_metrics()`; wrapper: `phase2/run_build_legalmeld_dataset_validated.sh`.",
            ],
            [
                "Prepare for Indian adaptation",
                "The SOP describes the future Indian acquisition pack, but the reusable logic is the same candidate/verified/manifest pipeline already implemented for the tribunal bootstrap layer.",
            ],
        ],
    )

    doc.add_heading("1. What To Keep And What To Ignore", level=1)
    add_bullets(
        doc,
        [
            "Keep `data/phase2/source_manifests/case_candidate_ledger.csv` as the human-curated candidate list.",
            "Treat `data/phase2/source_manifests/case_candidate_ledger_ucr_enriched.csv` as a verification layer, not a download estimate source.",
            "Use `data/processed/phase2/verified_case_inventory.csv` as the canonical inventory output.",
            "Keep `data/phase2/source_manifests/witness_harvest_manifest.csv` as a placeholder until it contains real witness/hearing rows and resolved links.",
            "Do not use the current placeholder witness manifest to drive downloads or corpus estimates.",
            "Treat the tribunal bootstrap outputs as the source for proof of pipeline, not the final Indian adaptation corpus.",
        ],
    )
    add_para(
        doc,
        "The verified inventory is the file that should drive the next stage of corpus selection. The enriched ledger still matters for validation and traceability, but the old page-text flags are not to be trusted for actual media estimates. Once the tribunal bootstrap is stable, the same manifest logic can be mirrored for an Indian acquisition pack.",
    )

    doc.add_heading("2. Canonical File Order", level=1)
    add_table(
        doc,
        ["Order", "File", "Purpose"],
        [
            ["1", "case_candidate_ledger.csv", "Human-curated candidate list."],
            ["2", "case_candidate_ledger_ucr_enriched.csv", "Corrected UCR verification layer."],
            ["3", "verified_case_inventory.csv", "Actual case inventory built from verified records."],
            ["4", "hearing_manifest.csv", "Grounded hearing-level and transcript/video pairing manifest."],
            ["5", "witness_harvest_manifest_resolved.csv", "Resolved witness-level manifest built from hearing rows."],
            ["6", "hearing_manifest_validated.csv", "Pilot-validated hearing rows with media and transcript checks."],
            ["7", "witness_manifest_validated.csv", "Validated witness rows with utterance estimates."],
            ["8", "witness_harvest_manifest.csv", "Placeholder only; keep inert until fully replaced."],
            ["9", "utterance-level LegalMemoCMT dataset", "Final segmented training dataset."],
            [
                "10",
                "tribunal_sources_target_dataset_broadened.csv",
                "Widened tribunal source manifest after the planning shortlist is exhausted.",
            ],
            [
                "11",
                "tribunal_media_discovery.csv",
                "Discovery review that classifies missing tribunal families as video-bearing, transcript-only, or unresolved.",
            ],
            [
                "12",
                "tribunal_case_resolution_review.csv",
                "Targeted case-number resolution review for unresolved tribunal families.",
            ],
            [
                "13",
                "mict_hearings_candidate_manifest.csv",
                "Specific MICT hearing/proceeding candidates that replace the generic IRMCT Hearings label.",
            ],
            [
                "14",
                "tribunal_video_bearing_candidates.csv",
                "Compact shortlist of the missing-source families that actually look video-bearing.",
            ],
        ],
    )

    doc.add_heading("3. Scripts Actually Used So Far", level=1)
    add_table(
        doc,
        ["Script", "Role", "Inputs", "Outputs / Notes"],
        [
            [
                "phase2/run_deduplicate_case_ledger.sh",
                "Merge and deduplicate candidate ledgers.",
                "case_candidate_ledger.csv and any extra ledgers.",
                "Writes the deduped ledger for curation review.",
            ],
            [
                "phase2/run_enrich_case_ledger_from_ucr_site.sh",
                "API-backed verification of candidate cases.",
                "case_candidate_ledger.csv and any extra ledgers.",
                "Writes the corrected enriched ledger and validation report.",
            ],
            [
                "phase2/enrich_case_ledger_from_ucr_site.py",
                "Core UCR enrichment implementation.",
                "Ledger rows plus live UCR JSON APIs.",
                "Rejects placeholders, counts actual records, and validates controls.",
            ],
            [
                "phase2/run_build_ucr_case_inventory.sh",
                "Build verified inventory rows.",
                "case_candidate_ledger_ucr_enriched.csv.",
                "Writes verified_case_inventory.csv.",
            ],
            [
                "phase2/build_ucr_case_inventory.py",
                "Inventory construction logic.",
                "Verified enriched ledger.",
                "Keeps only verified cases with actual records.",
            ],
            [
                "phase2/run_split_ucr_inventory_by_media_type.sh",
                "Split the verified inventory by media type.",
                "verified_case_inventory.csv.",
                "Writes video and transcript-only manifests.",
            ],
            [
                "phase2/split_ucr_inventory_by_media_type.py",
                "Split logic.",
                "Verified inventory rows.",
                "Produces `ucr_video_candidate_manifest.csv` and `ucr_transcript_only_manifest.csv`.",
            ],
            [
                "phase2/run_build_hearing_witness_manifests.sh",
                "Build grounded hearing and witness manifests.",
                "verified_case_inventory.csv plus the split manifests.",
                "Writes `hearing_manifest.csv`, `witness_harvest_manifest_resolved.csv`, and the summary report.",
            ],
            [
                "phase2/build_hearing_manifest.py",
                "Construct hearing-level rows from actual UCR record metadata.",
                "Verified inventory and split manifests.",
                "Pairs transcript/video records only when the metadata supports a real hearing/session match.",
            ],
            [
                "phase2/build_witness_manifest.py",
                "Project hearing rows into a resolved witness manifest.",
                "hearing_manifest.csv.",
                "Creates witness-level rows without deanonymizing protected witnesses.",
            ],
            [
                "phase2/hearing_witness_manifest_utils.py",
                "Shared heuristics for hearing grouping, pairing, and witness extraction.",
                "Inventory rows and manifest rows.",
                "Centralizes the ground-truth filters and pairing helpers.",
            ],
            [
                "phase2/run_validate_trimodal_corpus.sh",
                "Run the hearing validation, witness resolution, and utterance-estimation stages.",
                "hearing_manifest.csv and the resolved witness manifest.",
                "Writes validated hearing and witness manifests plus `trimodal_validation_summary.json`.",
            ],
            [
                "phase2/validate_trimodal_hearings.py",
                "Probe the pilot tri-modal subset for real media and readable transcripts.",
                "hearing_manifest.csv.",
                "Adds media, transcript, and pilot-selection validation columns.",
            ],
            [
                "phase2/resolve_witnesses_from_transcripts.py",
                "Resolve witness identity from validated transcript text.",
                "hearing_manifest_validated.csv.",
                "Writes the validated witness manifest with witness type and eligibility flags.",
            ],
            [
                "phase2/estimate_utterance_counts.py",
                "Estimate utterance counts from validated transcripts.",
                "witness_manifest_validated.csv.",
                "Adds utterance counts and refreshes the tri-modal validation summary.",
            ],
            [
                "phase2/discover_witnesses_all_paired_hearings.py",
                "Transcript-first witness discovery across all paired hearings.",
                "Validated hearing rows and transcript files.",
                "Writes the grounded hearing/witness discovery manifest with actual witness candidates.",
            ],
            [
                "phase2/run_expand_phase2_planning_manifests.sh",
                "Build the tribunal planning manifest from the current source manifest.",
                "Tribunal source manifest plus witness placeholder manifest.",
                "Use `TRIBUNAL_SOURCES=data/phase2/source_manifests/tribunal_sources_target_dataset_broadened.csv` when the broadened source list should drive the rerun.",
            ],
            [
                "phase2/discover_tribunal_media_candidates.py",
                "Probe missing-source tribunal families for actual video-bearing or transcript-only evidence.",
                "Expanded planning missing sources plus the candidate ledger.",
                "Writes a media-discovery review file that says which families are video-bearing, transcript-only, or unresolved.",
            ],
            [
                "phase2/resolve_targeted_tribunal_case_numbers.py",
                "Attempt stronger case-number resolution for unresolved tribunal families.",
                "`tribunal_media_discovery.csv`.",
                "Writes a case-resolution review file for Akayesu, Ntakirutimana / Ntahobali, and IRMCT Hearings.",
            ],
            [
                "phase2/build_mict_hearings_candidate_manifest.py",
                "Build a specific MICT hearings candidate list from official tribunal case pages.",
                "Official MICT case pages and the unresolved IRMCT hearing placeholder.",
                "Writes a small MICT candidate manifest so the generic IRMCT label can be replaced by case-specific identifiers.",
            ],
            [
                "phase2/filter_tribunal_video_bearing_candidates.py",
                "Reduce the media-discovery review to only video-bearing rows.",
                "`tribunal_media_discovery.csv`.",
                "Writes a compact shortlist for the next download decision.",
            ],
            [
                "phase2/build_balanced_corpus_selection.py",
                "Selection scoring and cap enforcement.",
                "Transcript-first witness discovery rows.",
                "Chooses grounded rows while staying under the hour cap and preserving witness diversity.",
            ],
            [
                "phase2/run_build_text_only_diversity_supplement.sh",
                "Build the transcript-only diversity supplement.",
                "Validated transcript-only hearing rows.",
                "Writes a small text-only manifest that adds case diversity without adding video hours.",
            ],
            [
                "phase2/build_text_only_diversity_supplement.py",
                "Select the highest-signal transcript-only hearings.",
                "Validated transcript-only hearing rows and local transcript files.",
                "Produces a grounded text supplement from IT-04-81, IT-05-88, and IT-09-92.",
            ],
            [
                "phase2/run_build_legalmeld_dataset.sh",
                "Build the utterance-level LegalMELD dataset.",
                "Validated hearings with local transcript and video paths.",
                "Writes `legalmeld_metadata.csv`, `train.csv`, `dev.csv`, `test.csv`, and per-utterance clips/audio.",
            ],
            [
                "phase2/build_legalmeld_dataset.py",
                "Transcript parsing, alignment, and clip generation.",
                "Hearing-level transcript and media rows.",
                "Aligns speaker turns to word timestamps and exports MELD-style utterance samples.",
            ],
            [
                "phase2/trimodal_validation_utils.py",
                "Shared helpers for media probing, transcript extraction, witness identity, and utterance counting.",
                "Remote media URLs, transcript files, and manifest rows.",
                "Keeps the validation heuristics and extraction logic in one place.",
            ],
            [
                "phase2/run_phase2_dataset_pipeline.sh",
                "Current guarded preparation wrapper.",
                "Placeholder witness manifest and source manifests.",
                "Stops when the witness manifest is still placeholder-only.",
            ],
            [
                "Indian acquisition pack scripts",
                "Planned Indian adaptation layer.",
                "Supreme Court / High Court / mock-trial manifests.",
                "Not yet implemented as a completed acquisition pipeline; reserved for the Indian scope.",
            ],
            [
                "scripts/check_phase2_sources_ready.sh",
                "Source readiness check.",
                "Source manifests and tribunal indices.",
                "Reports file presence and row counts without requiring pandas.",
            ],
            [
                "scripts/check_phase2_dataset_ready.sh",
                "Dataset readiness check.",
                "Source manifests and derived artifacts.",
                "Reports dataset availability, splits, and language mix.",
            ],
            [
                "scripts/check_phase2_finetune_ready.sh",
                "Fine-tune readiness check.",
                "Tri-modal manifest and warm-start checkpoint.",
                "Confirms the training manifest before any run.",
            ],
            [
                "scripts/check_phase2_language_distribution.sh",
                "Language profile for the split manifest.",
                "Split manifest CSV.",
                "Reports English, Devanagari, and other-script shares.",
            ],
            [
                "scripts/check_phase2_video_integrity.sh",
                "Validate downloaded media files.",
                "Split or tri-modal manifest.",
                "Confirms the stored files are actual media.",
            ],
        ],
    )

    doc.add_heading("4. Current SOP", level=1)
    add_numbered(
        doc,
        [
            "Start from `case_candidate_ledger.csv` and keep that file as the curated source of candidates.",
            "Run `phase2/run_enrich_case_ledger_from_ucr_site.sh` to validate the candidate list against the live UCR APIs.",
            "Review `reports/phase2/ucr_enrichment_validation.json` and confirm the positive and negative controls passed.",
            "Build `verified_case_inventory.csv` with `phase2/run_build_ucr_case_inventory.sh`.",
            "Split the verified inventory with `phase2/run_split_ucr_inventory_by_media_type.sh`.",
            "Build `hearing_manifest.csv` and `witness_harvest_manifest_resolved.csv` with `phase2/run_build_hearing_witness_manifests.sh`.",
            "Run `phase2/run_validate_trimodal_corpus.sh` to validate the pilot tri-modal subset and refresh `hearing_manifest_validated.csv`, `witness_manifest_validated.csv`, and `reports/phase2/trimodal_validation_summary.json`.",
            "Run `phase2/run_build_balanced_corpus_selection.sh` to build transcript-first witness discovery rows and the balanced corpus selection manifest.",
            "Run `phase2/run_build_text_only_diversity_supplement.sh` to add a small transcript-only supplement from other cases without increasing video hours.",
            "Run `phase2/run_build_legalmeld_dataset.sh` to create the utterance-level MELD-style dataset with aligned clips and audio.",
            "Use the hearing, resolved witness, and selection manifests for later download planning, not the placeholder witness manifest.",
            "Run `scripts/check_phase2_dataset_ready.sh` and `scripts/check_phase2_finetune_ready.sh` only after the required derived files exist.",
            "Keep `witness_harvest_manifest.csv` inert until it contains real witness and hearing rows with resolved record links.",
            "For the dissertation's Indian scope, prepare a separate acquisition pack after the tribunal bootstrap path is stable.",
        ],
    )

    doc.add_heading("5. Current Discovery And Selection Snapshot", level=1)
    add_table(
        doc,
        ["Metric", "Current Value"],
        [
            ["Paired hearings scanned", "83"],
            ["Readable transcripts", "90"],
            ["Hearings with witness testimony", "84"],
            ["Distinct witnesses discovered", "48"],
            ["Selected hearings", "37"],
            ["Selected witnesses", "42"],
            ["Selected estimated hours", "29.92"],
            ["Estimated utterances", "1661"],
            ["Case families selected", "1"],
            ["Target hours status", "within_target"],
            ["Target witnesses status", "within_target"],
            ["Target utterances status", "below_target"],
        ],
    )
    add_para(
        doc,
        "This snapshot reflects the transcript-first discovery and balanced selection pass. It now stays under the 30-hour cap and within the witness-count window, but it is still below the utterance target, which is a signal that the current corpus selection is grounded but not yet large enough for the final training budget.",
    )

    doc.add_heading("6. Text-Only Diversity Supplement", level=1)
    add_table(
        doc,
        ["Metric", "Current Value"],
        [
            ["Rows selected", "5"],
            ["Case families selected", "3"],
            ["Case numbers selected", "3"],
            ["Estimated utterances", "2883"],
            ["Witness utterances", "861"],
            ["Counsel utterances", "1414"],
            ["Judge utterances", "608"],
        ],
    )
    add_para(
        doc,
        "This supplement is the key evidence that improving extraction and adding modest case diversity is likely to help more than simply adding extra hours of the same video-heavy case. Five transcript-only hearings from three additional cases already contribute 2,883 utterances, which is more new text signal than the current tri-modal selection contains, and it does so without increasing the video budget at all.",
    )

    doc.add_heading("7. LegalMELD Alignment Stage", level=1)
    add_table(
        doc,
        ["Metric", "Current Value"],
        [
            ["Hearings resolved", "1"],
            ["Utterances parsed", "10"],
            ["Matched alignments", "7"],
            ["Fuzzy alignments", "3"],
            ["Fallback alignments", "0"],
            ["Clips exported", "10"],
            ["Audio files exported", "10"],
            ["Case families represented", "1"],
        ],
    )
    add_para(
        doc,
        "This is the new Stage 3 artifact. It turns a hearing-level recording into a MELD-style utterance dataset, with one aligned audio clip and one aligned video clip per utterance, plus a master metadata CSV and train/dev/test split files. The pilot run currently proves the pipeline on the Karadzic judgment hearing: the transcript was segmented into 10 utterances, 7 aligned by direct ASR matching and 3 by fuzzy alignment, and the outputs were written under `data/processed/phase2/legalmeld/`. The alignment step is now configurable, so the same export can be rerun with the heuristic backend or a WhisperX-style backend without changing the schema.",
    )

    doc.add_heading("8. Guardrails", level=1)
    add_bullets(
        doc,
        [
            "The utterance builder should be run with `PYTHON_BIN=/usr/bin/python3` in this workspace so `faster_whisper` can be imported reliably.",
            "The alignment backend can be selected with `ALIGNMENT_BACKEND=auto`, `whisperx`, or `heuristic` depending on what you want to test.",
            "Invalid placeholder case numbers must be rejected before any UCR request is made.",
            "Do not infer `has_videos`, `has_transcripts`, `has_court_recordings`, or corpus size from the generic UCR page shell.",
            "Do not count generic navigation links or static headings as evidence of real records.",
            "Use only actual case-specific record rows, document IDs, TAP IDs, transcript records, and downloadable media URLs.",
            "Keep the placeholder witness manifest out of download and corpus-estimation steps until it is populated with real rows.",
        ],
    )

    doc.add_heading("9. What Is Not Yet Active", level=1)
    add_bullets(
        doc,
        [
            "The placeholder `witness_harvest_manifest.csv` remains inert and must not be used as corpus data.",
            "The validated hearing and witness manifests are pilot-stage checkpoints, not the final corpus.",
            "Witness resolution is still incomplete for most rows; `UNRESOLVED_WITNESS` is expected when the metadata does not expose a safe identifier.",
            "Audio extraction, utterance segmentation, and final dataset assembly remain downstream of the verified inventory and resolved hearing/witness manifests.",
        ],
    )

    doc.add_heading("10. Practical Handoff", level=1)
    add_para(
        doc,
        "The practical handoff after this SOP is simple: use the verified inventory to plan actual downloads, then create real witness or hearing rows, then segment utterances. Until that happens, the project should treat the enriched ledger as a validation layer and the verified inventory as the only reliable Phase 2 case-level source. After that, repeat the same manifest-and-alignment logic on Indian Supreme Court, High Court, and mock-trial material where synchronized media exists. If you need to explain the utterance builder, make it clear that the alignment backend is selectable and the Python binary is pinned only to make the current environment reproducible.",
    )

    add_para(
        doc,
        "If you need to explain the pipeline verbally, keep the story in this order: curated ledger, corrected enrichment, verified inventory, real witness/hearing manifest, tribunal bootstrap dataset, then Indian adaptation pack.",
    )

    doc.add_heading("11. Tribunal Expansion Before Indian Adaptation", level=1)
    add_para(
        doc,
        "This is the next practical step if the goal is to raise usable tribunal minutes before switching to the Indian adaptation track. The purpose is not to claim that every new tribunal case is ready for training. The purpose is to use the verified inventory, the tribunal planning manifest, and the existing download checks to identify new grounded cases, test one case at a time, and then widen the selection only when the evidence stays clean. That makes the corpus growth reproducible and safe to redo later if needed.",
    )
    add_para(
        doc,
        "In student language, this stage sits between the current tribunal bootstrap dataset and the future Indian corpus. It is the best place to increase utterance minutes because it reuses the same verified-case rules, the same hearing and witness manifest logic, and the same download validation checks that already proved the pipeline works.",
    )
    add_table(
        doc,
        ["Step", "Script / file", "What it does", "Why it matters"],
        [
            [
                "1",
                "`data/phase2/source_manifests/tribunal_sources_target_dataset.csv` and `data/processed/phase2/verified_case_inventory.csv`",
                "Compare planned tribunal targets with the cases already verified in the inventory.",
                "This shows which tribunal case families are genuinely new expansion candidates.",
            ],
            [
                "2",
                "`phase2/run_expand_phase2_planning_manifests.sh`",
                "Expands tribunal and witness planning manifests into a larger candidate inventory.",
                "This gives you a broader tribunal shortlist without pretending it is final corpus data.",
            ],
            [
                "3",
                "`phase2/download_ucr_case_video.py`",
                "Tests one exact tribunal case number and verifies that the downloaded MP4 is real media.",
                "This is the safest manual check before any broader case download or corpus promotion.",
            ],
            [
                "4",
                "`phase2/run_ucr_case_videos_strict.sh`",
                "Downloads only real video files for selected tribunal cases.",
                "This helps raise usable minutes without counting HTML pages, transcript-only rows, or fake media.",
            ],
            [
                "5",
                "`phase2/run_build_ucr_case_inventory.sh` and `phase2/run_split_ucr_inventory_by_media_type.sh`",
                "Rebuilds the verified inventory and splits it into video-bearing and transcript-only manifests.",
                "This refreshes the canonical inventory after any new tribunal case is promoted.",
            ],
            [
                "6",
                "`phase2/run_build_hearing_witness_manifests.sh` and `phase2/run_validate_trimodal_corpus.sh`",
                "Rebuilds hearing and witness manifests and checks tri-modal validity again.",
                "This confirms the new tribunal cases still support grounded hearing-level and witness-level rows.",
            ],
            [
                "7",
                "`phase2/run_build_legalmeld_dataset_validated.sh`",
                "Reruns the utterance-level dataset builder on the refreshed grounded set.",
                "This is the step that converts the expanded tribunal evidence into more usable utterance minutes if the rows are clean.",
            ],
        ],
    )
    add_bullets(
        doc,
        [
            "Keep the old tribunal bootstrap outputs as proof of method.",
            "Do not overwrite the logic that made the current verified inventory trustworthy.",
            "Only promote new tribunal cases when the records are grounded in actual case metadata.",
            "If a test download fails or returns HTML, stop at the case level and do not expand the batch.",
        ],
    )
    add_para(
        doc,
        "If you need to explain why this comes before Indian adaptation, the simple answer is that the tribunal branch is still the easiest place to increase corpus size in a controlled way. It already has the verified-inventory, manifest-building, download, and validation machinery in place, so you can improve utterance counts and case diversity without changing the project scope too early.",
    )

    doc.add_heading("12. Analysis of phase2_expanded_planning_manifest.csv", level=1)
    add_para(
        doc,
        "The file `data/processed/phase2/phase2_expanded_planning_manifest.csv` is the output of the expansion step, not a final corpus manifest. It tells you which tribunal and witness planning entries were expanded from the current source manifests so that the next download and verification pass can be controlled. The important idea is that this file is still a planning layer: it is useful for deciding what to inspect next, but it is not proof that the listed rows are already grounded media rows.",
    )
    add_para(
        doc,
        "The exact inspection workflow is: run `python3 phase2/check_ucr_case_resolution.py --source-csv data/processed/phase2/phase2_expanded_planning_manifest.csv`, then read the row-by-row console output. For each row, look for three signals. First, whether the case resolves at all in UCR. Second, whether the returned document mix contains `TAP` documents. Third, whether the case is already represented in `verified_case_inventory.csv` or is genuinely new. If a row prints `tap_docs: 0`, it is not a video-bearing expansion candidate. If a row prints `resolution: no candidate case number available`, it needs a better case-number mapping before you can inspect media. If a row prints `no case detail returned`, it should not be promoted until the case number is corrected or the row is replaced.",
    )
    add_bullets(
        doc,
        [
            "Use `tap_docs > 0` as the quick sign that the row is worth a download test.",
            "Use `total_docs` and `doc_types` to see whether the row is mostly transcript/judgment material or truly video-bearing.",
            "Use the verified inventory to separate already-covered rows from new expansion rows.",
            "Do not treat a resolved case with no TAP docs as evidence for tri-modal growth.",
        ],
    )
    add_table(
        doc,
        ["Metric", "Value", "Meaning"],
        [
            ["Total rows", "17", "The expanded planning set is still small enough to inspect by hand."],
            ["Tribunal rows", "11", "These are the tribunal-side expansion targets."],
            ["Witness rows", "6", "These are the witness-side expansion targets for the planning pass."],
            ["Tribunals represented", "3", "ICTY, ICTR, and IRMCT are all present in the planning output."],
            ["Already verified case families", "5", "Karadzic, Mladic, Popovic, Bagosora et al., and Perisic already exist in the verified inventory."],
            ["New expansion value", "Controlled", "The file helps identify which planned rows are still outside the verified inventory and therefore worth testing next."],
        ],
    )
    add_para(
        doc,
        "A student-level way to read the file is: the tribunal rows tell me which case-family candidates the expansion logic found, while the witness rows tell me which witness-side planning entries were carried forward for later inspection. The file should be read as a shortlist for action, not as a claim that all rows are ready for download or training.",
    )
    add_bullets(
        doc,
        [
            "The tribunal rows are the main expansion target because they can add usable minutes if the download and verification checks pass.",
            "The witness rows are still planning rows; they are useful for traceability, but they are not automatically evidence of corpus readiness.",
            "The overlap with the verified inventory matters because it helps separate already-covered cases from true expansion candidates.",
            "The file should be normalized against the verified inventory before any download decision is made.",
        ],
    )
    add_para(
        doc,
        "In practical terms, this CSV is the bridge between the current tribunal bootstrap and the next download test. If a row already exists in the verified inventory, it should be treated as coverage rather than new growth. If it does not exist there, it becomes a candidate for a one-case verification test. That is why the file belongs in the SOP: it explains how the next batch of tribunal work should be chosen without inflating the corpus estimate.",
    )

    doc.add_heading("13. Filter TAP Candidates from the Expanded Manifest", level=1)
    add_para(
        doc,
        "After the expanded planning manifest has been inspected, the next small step is to filter it down to the rows that actually return TAP documents in UCR. This produces `data/processed/phase2/tap_candidate_manifest.csv`. The purpose of this file is narrow: it is a shortlist of planning rows that are worth a real download test because the live UCR API returned video-bearing records. It is still not the verified inventory, but it is much closer to a case-level download queue than the raw expanded planning manifest.",
    )
    add_para(
        doc,
        "The exact command is `bash phase2/run_filter_tap_candidates_from_expanded_manifest.sh`. Internally, that wrapper runs `phase2/filter_tap_candidates_from_expanded_manifest.py`, which reads `phase2_expanded_planning_manifest.csv`, resolves each case against the live UCR API, keeps only rows with `tap_docs > 0`, and writes the filtered output plus a short summary JSON. The filter now includes a case-name hint for `Perisic Trial` as well, so the case can reach the UCR TAP lookup path instead of being skipped immediately. In student language, this means the script is checking whether a planning row actually has video recordings before you spend time trying to download or validate it.",
    )
    add_table(
        doc,
        ["Output", "What to look for", "How to interpret it"],
        [
            [
                "`tap_candidate_manifest.csv`",
                "Rows where `tap_docs > 0` and `doc_types` includes `TAP`.",
                "These are the planning rows that are worth a real download test.",
            ],
            [
                "`tap_candidate_manifest_summary.json`",
                "`rows_inspected`, `rows_kept`, and the list of case names kept.",
                "This tells you how many planning rows survived the TAP filter.",
            ],
            [
                "Console output from the script",
                "Resolved case number, document counts, and TAP doc examples.",
                "Use this to verify the row is grounded in actual UCR media metadata.",
            ],
        ],
    )
    add_bullets(
        doc,
        [
            "Keep rows where the live UCR response reports at least one TAP document.",
            "Treat rows with `tap_docs = 0` as transcript-only or non-video planning rows.",
            "Treat rows with no case-number mapping as unresolved planning rows, not download targets.",
            "Perisic Trial now has an explicit case hint, so it can be checked instead of bypassed at the hint stage.",
            "Use the filtered TAP manifest as the immediate input for one-case verification tests.",
        ],
    )
    add_para(
        doc,
        "For `Perisic Trial`, the live UCR probe now shows the exact reason it stays out of the TAP shortlist: `ByCaseDetail` resolves `IT-04-81` correctly, but `ByCaseDocsByLang` returns zero TAP documents for that case. The broader `ByMainCase` endpoint still shows appeal and judgment documents for `IT-04-81-A`, which confirms that the family exists in UCR but does not currently produce the kind of TAP-bearing courtroom recordings this shortlist is looking for. In other words, the case is resolvable, but it is not TAP-positive in the current planning filter.",
    )
    add_para(
        doc,
        "The correct interpretation for the project is to treat `Perisic Trial` as transcript-only for now. That means it can remain useful for transcript discovery, reading-flow analysis, and later text-only or transcript-first planning, but it should not be used as evidence that the tribunal branch has a new video-bearing TAP candidate. Leaving the TAP filter unchanged is the correct choice because the TAP shortlist is meant to stay video-grounded, and Perisic does not currently satisfy that criterion.",
    )
    add_para(
        doc,
        "For the current expansion set, this filter keeps the Karadzic and Bagosora planning rows that are video-bearing. That is the right outcome because the goal here is not to enlarge the manifest for its own sake. The goal is to isolate the rows that genuinely support more usable tribunal minutes before the project moves to the Indian adaptation stage.",
    )
    add_para(
        doc,
        "If you are choosing the first real download test from `tap_candidate_manifest.csv`, use the Karadzic row first: `ICTY_KARADZIC_01` / `IT-95-5/18`. It is the smallest clean video-bearing candidate, it resolves cleanly, and it has only six TAP docs, so it is safer for a first validation download than the much larger Bagosora set. After that test succeeds, the next expansion target is the Bagosora row because it has far more TAP material and therefore a bigger potential impact on usable minutes.",
    )

    doc.add_heading("14. Why the First Bagosora Attempt Skipped Video", level=1)
    add_para(
        doc,
        "The first generic Bagosora attempt did not fail because Bagosora had no video. It failed because the downloader's first-pass selection heuristic did not land on the right TAP row for the Bagosora planning entry. Earlier, the strict downloader used a different Bagosora case-number label than the one returned by the UCR API, so the planning filter and the downloader were not aligned. After normalizing both to `ICTR-98-41`, the case resolved correctly and the repository could see that Bagosora does in fact return a very large TAP set. The remaining issue was not case absence; it was the need to choose a specific tape/date row rather than the first TAP row returned by the generic heuristic.",
    )
    add_bullets(
        doc,
        [
            "Before the fix, the strict downloader and the planning filter did not agree on the Bagosora case-number label.",
            "After the fix, both paths resolve Bagosora as `ICTR-98-41`.",
            "The first strict pass still skipped Bagosora because it selected a non-video TAP row from the returned list.",
            "The live UCR API now confirms that Bagosora has 424 TAP documents, so the case is video-bearing and simply needs a better tape/date selector.",
        ],
    )

    doc.add_heading("15. Commanded Validation Path for Karadzic and Bagosora", level=1)
    add_para(
        doc,
        "This is the repeatable validation path you can show in a guidance call. First filter the expanded manifest into TAP-bearing candidates. Then use the Karadzic row as the first download test because it is the smallest clean video-bearing case. After that, use a specific Bagosora tape date so the downloader lands on a real video row instead of an arbitrary first match. Finally, use the fallback MP4 checker when `ffprobe` is not installed so the downloaded file can still be validated in a reproducible way.",
    )
    add_table(
        doc,
        ["Step", "Command", "Expected output", "Technical meaning"],
        [
            [
                "1",
                "`bash phase2/run_filter_tap_candidates_from_expanded_manifest.sh`",
                "`tap_candidate_manifest.csv` and `tap_candidate_manifest_summary.json`",
                "Builds the TAP-bearing shortlist from the expanded planning manifest.",
            ],
            [
                "2",
                "`python3 phase2/download_ucr_case_video.py --case-number IT-95-5/18 --verify`",
                "A Karadzic MP4 download with file-level validation; `ffprobe` is skipped when missing.",
                "Confirms the smallest clean video-bearing tribunal case works end to end.",
            ],
            [
                "3",
                "`python3 phase2/download_ucr_case_video.py --case-number ICTR-98-41 --date 01/06/2007 --index 1 --verify`",
                "A specific Bagosora TAP video is selected by date if the row matches.",
                "Shows that Bagosora is video-bearing and requires a precise tape filter.",
            ],
            [
                "4",
                "`python3 scripts/check_mp4_fallback.py <mp4-file>`",
                "PASS/FAIL output using `file` plus MP4 header checks when `ffprobe` is missing.",
                "Lets the repo validate MP4s even in a minimal environment.",
            ],
        ],
    )
    add_bullets(
        doc,
        [
            "Use Karadzic to prove the basic download path.",
            "Use Bagosora to prove the downloader can be steered to a specific TAP row.",
            "Use the fallback validator when the environment does not provide `ffprobe`.",
            "Do not interpret a skipped generic TAP selection as evidence that the case has no video.",
        ],
    )
    add_para(
        doc,
        "Technically, the outcome is that Karadzic proves the direct case-number-to-video path, while Bagosora proves that a case can be video-bearing but still require a more precise tape selection. That distinction is important because it prevents a false negative from being interpreted as corpus absence. It also gives you a cleaner explanation for the next SOP step: keep the planning filter, keep the verified case resolution, and use a fallback media check so the repository can still validate downloaded MP4s without `ffprobe`.",
    )

    doc.add_heading("16. Tribunal Expansion Output Summary", level=1)
    add_bullets(
        doc,
        [
            "The expanded planning manifest produced 17 rows, of which 6 survived the TAP filter.",
            "Karadzic is the first clean validation target and downloaded successfully as a real MP4.",
            "Bagosora is confirmed to be video-bearing, but the first strict pass skipped it because the generic selector did not land on the correct TAP row.",
            "The Bagosora case-number mapping was normalized to `ICTR-98-41`, which aligned the planning filter and the downloader.",
            "The new fallback MP4 checker now validates the downloaded Karadzic file even when `ffprobe` is absent.",
        ],
    )
    add_para(
        doc,
        "The practical meaning of this summary is simple: the tribunal branch is still alive as a usable bootstrap source, Karadzic gives you a clean proof-of-download example, Bagosora gives you a larger but more delicate expansion target, and the repository now has a fallback media-check path so the validation step does not depend on a missing system binary.",
    )

    doc.add_heading("17. Inspect the Strict Download Index by Unique File", level=1)
    add_para(
        doc,
        "Once the strict downloader has written `data/phase2/ucr_case_video_strict/index/ucr_case_videos_strict.csv`, the next student-friendly check is to inspect that file in two ways. First, collapse repeated rows into a unique-file view so you can see how many distinct MP4s were actually produced. Second, map each index row back to the source manifest so you can explain which planning rows pointed at the same file. Third, validate the unique MP4 files with the fallback media checker so the workflow still works even if `ffprobe` is missing from the machine.",
    )
    add_para(
        doc,
        "The exact command is `bash phase2/run_inspect_ucr_case_videos_strict_index.sh`. That wrapper runs `phase2/inspect_ucr_case_videos_strict_index.py`, which reads the strict index, writes a row-to-source mapping CSV, writes a unique-file CSV, and then calls the fallback MP4 validator on each unique file. In student terms, this step answers three questions at once: which rows are duplicates, which source rows created them, and whether the downloaded files are real MP4s.",
    )
    add_table(
        doc,
        ["Output", "What it shows", "How to read it"],
        [
            [
                "`ucr_case_videos_strict_row_source_map.csv`",
                "The original strict index rows plus source-manifest match fields.",
                "Use this to explain where each row came from and whether multiple rows point to the same recording.",
            ],
            [
                "`ucr_case_videos_strict_unique_files.csv`",
                "One row per distinct `local_video_path` or `resolved_video_url`.",
                "Use this to see the actual number of unique file paths rather than the number of repeated index rows.",
            ],
            [
                "`ucr_case_videos_strict_unique_recordings.csv`",
                "One row per distinct `resolved_video_url`.",
                "Use this to see how many recordings are truly distinct even when the same MP4 is stored in more than one folder.",
            ],
            [
                "`ucr_case_videos_strict_duplicate_file_groups.csv`",
                "One row per repeated file hash across unique paths.",
                "Use this to prove when two folder paths contain the exact same MP4 bytes.",
            ],
            [
                "`ucr_case_videos_strict_validation.csv`",
                "Fallback validation results for the unique MP4s.",
                "Use this to confirm the file is a real MP4 even when `ffprobe` is unavailable.",
            ],
            [
                "`ucr_case_videos_strict_inspection_summary.json`",
                "Counts for index rows, unique files, and validation results.",
                "Use this as the quick summary for a guidance call or checkpoint note.",
            ],
        ],
    )
    add_bullets(
        doc,
        [
            "If several rows share the same `local_video_path`, the same file path is being reused rather than redownloaded.",
            "If several rows share the same `resolved_video_url`, they are the same recording even if the file exists under more than one folder path.",
            "If two different local paths have the same SHA-256 hash, they are byte-identical copies of the same MP4.",
            "If the mapping points multiple rows to one source case and one tape, the index is showing planned duplication, not extra media.",
            "If the fallback validator says PASS, the file is a genuine MP4 even without `ffprobe`.",
            "This step is about explaining the evidence chain, not expanding the corpus further.",
        ],
    )
    add_para(
        doc,
        "The exact command sequence for this stage is: 1) download the strict manifest, 2) inspect the strict index, 3) validate the MP4 files, and 4) use the unique-recording view when you explain the result in a guidance call. In practice that means you first run `bash phase2/run_ucr_case_videos_strict.sh`, then `bash phase2/run_inspect_ucr_case_videos_strict_index.sh`, then check `data/phase2/ucr_case_video_strict/inspection/ucr_case_videos_strict_validation.csv`, and finally read `data/phase2/ucr_case_video_strict/inspection/ucr_case_videos_strict_unique_recordings.csv` to talk about actual recordings rather than repeated index rows.",
    )
    add_para(
        doc,
        "For the current strict index, this inspection step should show three unique local file paths, two distinct recordings by URL, and one duplicate-file hash group: the Karadzic MP4 exists under two folder paths but is byte-identical in both places, while the Bagosora MP4 exists under one folder path. The repeated rows in the index are still useful because they show which planning entries resolve to the same recording, but the unique-recording and duplicate-file views are the better answers when you are asked how many actual recordings or physical copies were produced.",
    )

    doc.add_heading("18. Select the Next Controlled TAP-bearing Candidate", level=1)
    add_para(
        doc,
        "Once the strict-index inspection is complete, the next implementation action is to ask a simple question: is there another TAP-bearing tribunal row in the shortlist that has not already been covered by the current strict downloads? The new selector script answers that question by comparing `tap_candidate_manifest.csv` against the source IDs already consumed by the strict download index. If a row remains, it is written into a one-row or small-row manifest for the next controlled verification download. If no row remains, that means the current TAP shortlist has been exhausted and the planning manifests need to be expanded again before more tribunal downloads are attempted.",
    )
    add_para(
        doc,
        "The exact command is `bash phase2/run_select_next_tap_verification_candidate.sh`. That wrapper runs `phase2/select_next_tap_verification_candidate.py`, which reads the TAP shortlist, removes source IDs already covered by the strict index, and writes `data/processed/phase2/next_tap_verification_manifest.csv` plus a short JSON summary. In student language, this is the gatekeeper step that tells you whether the next download test is ready or whether you need to widen the planning layer first.",
    )
    add_table(
        doc,
        ["Output", "What it shows", "How to use it"],
        [
            [
                "`next_tap_verification_manifest.csv`",
                "The next uncaptured TAP-bearing row or rows from the shortlist.",
                "Use this as the input to the next strict download test.",
            ],
            [
                "`next_tap_verification_summary.json`",
                "Counts of covered, remaining, and selected rows.",
                "Use this to check whether the shortlist still contains anything new.",
            ],
        ],
    )
    add_bullets(
        doc,
        [
            "If the output CSV is empty, there is no remaining TAP-bearing candidate in the current shortlist.",
            "If the output CSV has one row, run a strict download on that row only.",
            "Use a separate output root for the next candidate so the existing strict index stays readable.",
            "After the download, run the strict index inspection again and read the unique-recording view.",
        ],
    )
    add_para(
        doc,
        "The manual execution sequence is: 1) run the selector, 2) inspect the selected manifest, 3) if it is non-empty, run `bash phase2/run_ucr_case_videos_strict.sh --source-csv data/processed/phase2/next_tap_verification_manifest.csv --output-root data/phase2/ucr_case_video_next --index-csv data/phase2/ucr_case_video_next/index/ucr_case_videos_strict.csv --skip-existing --verify`, and 4) inspect the resulting index with `bash phase2/run_inspect_ucr_case_videos_strict_index.sh --index-csv data/phase2/ucr_case_video_next/index/ucr_case_videos_strict.csv --source-csv data/processed/phase2/next_tap_verification_manifest.csv --output-dir data/phase2/ucr_case_video_next/inspection`. If the selector returns nothing, the next action is not to force a download; it is to rerun the planning expansion and widen the shortlist.",
    )
    add_para(
        doc,
        "This is the correct next implementation step because it keeps the corpus growth controlled. You only move forward when the new candidate is grounded in the TAP shortlist, and you only keep moving if the download and validation path still works. That means the SOP now has a clear branch: use the current shortlist if it still contains a new row, or expand the planning manifests again if it does not.",
    )

    doc.add_heading("19. Compare the Expanded Planning Manifest to the Verified Inventory", level=1)
    add_para(
        doc,
        "If the TAP shortlist is already consumed, the next question is whether the broader expanded planning manifest contains tribunal rows that are still missing from the verified inventory. This comparison step answers that directly. It reads `phase2_expanded_planning_manifest.csv`, compares each row against `verified_case_inventory.csv`, and labels the row as already covered or still missing. In student language, this is the bridge between the planning layer and the canonical case inventory: it shows what is already safe to reuse and what still needs a broader source search.",
    )
    add_para(
        doc,
        "The exact command is `bash phase2/run_compare_expanded_planning_to_verified_inventory.sh`. That wrapper runs `phase2/compare_expanded_planning_to_verified_inventory.py`, which writes three outputs: a full comparison CSV, a missing-sources CSV, and a summary JSON. If you need to explain the purpose in a guidance call, you can say that this step is the evidence check for planning growth. It tells you whether the 17-row planning set is only repeating cases that are already verified, or whether it still contains genuinely new tribunal sources that can widen the corpus.",
    )
    add_table(
        doc,
        ["Output", "What it shows", "How to interpret it"],
        [
            [
                "`expanded_planning_vs_verified_inventory.csv`",
                "Every expanded planning row plus a comparison status.",
                "Use this to see whether each row is already covered by the verified inventory.",
            ],
            [
                "`expanded_planning_missing_sources.csv`",
                "Only the rows that are not matched by the verified inventory.",
                "Use this to identify the actual missing tribunal sources that justify broader expansion.",
            ],
            [
                "`expanded_planning_vs_verified_inventory_summary.json`",
                "Counts of covered rows, missing rows, and category totals.",
                "Use this as the quick decision file for the next planning action.",
            ],
        ],
    )
    add_bullets(
        doc,
        [
            "Rows marked `already_verified_by_case_number` are not new corpus growth.",
            "Rows marked `already_verified_by_case_name` are likely duplicates of already covered tribunal sources.",
            "Rows marked `new_expansion_source` are the ones that matter if you want to widen the tribunal set.",
            "Rows marked `unresolved_or_placeholder` should be cleaned before any further corpus decision is made.",
        ],
    )
    add_para(
        doc,
        "The manual execution sequence is: 1) run the comparison script, 2) inspect the summary JSON, 3) open the missing-sources CSV, and 4) decide whether the next action is to add new tribunal source manifests or to patch the current planning rules. If the missing-sources file is empty, the current planning layer is not broad enough to move the corpus forward, so the correct next step is to expand the source manifests rather than to keep re-running the same shortlist. If it is non-empty, then the missing rows become the evidence for the next controlled tribunal expansion batch.",
    )
    add_para(
        doc,
        "This step matters because it prevents false progress. Without it, the pipeline can keep re-generating the same 17 rows and the same six TAP candidates. With it, you can point to exactly which expanded rows are already covered and which ones are genuinely missing from the verified inventory. That is the information you need before creating more corpus, because it keeps the next download and manifest work grounded in the current state of the repo.",
    )

    doc.add_heading("20. Discover Official Tribunal Source-Broadening Candidates", level=1)
    add_para(
        doc,
        "This is the first step in the tribunal-expansion branch when the current TAP shortlist is exhausted or when you want to widen the source set without guessing. It reads the official ICTR, ICTY, and IRMCT case-list pages, compares the case families and case numbers against the current source manifest, and produces a review layer that separates already-covered cases from genuinely new tribunal families. In student language, this is the web-based discovery stage that tells you what the tribunal sites themselves say exists before you decide what to add to the pipeline.",
    )
    add_para(
        doc,
        "The exact command is `bash phase2/run_refresh_tribunal_source_broadening.sh`. That wrapper first runs `phase2/discover_official_tribunal_source_broadening_candidates.py`, which crawls the official case-list pages and probes the UCR case APIs for TAP and transcript evidence. It then runs `phase2/build_broadened_tribunal_source_manifest.py` through `phase2/run_build_broadened_tribunal_source_manifest.sh` with `--include-hold-for-link-validation`, so the broadened manifest can keep transcript-only families that still matter for source planning. The important point is that this step does not download media; it only refreshes the source-broadening review and the broadened planning input.",
    )
    add_table(
        doc,
        ["Output", "What it shows", "How to interpret it"],
        [
            [
                "`tribunal_source_broadening_review.csv`",
                "Official-site candidates after excluding already-covered case families and numbers.",
                "Use this as the decision layer for source broadening, not as download proof.",
            ],
            [
                "`tribunal_source_broadening_review_summary.json`",
                "Counts of broaden-now, hold, and manual-review rows from the official-site discovery pass.",
                "Use this to see how much of the new source set is ready, held, or unresolved.",
            ],
        ],
    )
    add_bullets(
        doc,
        [
            "A `broaden_now` row can be added to the source manifest immediately after review.",
            "A `hold_for_link_validation` row should stay transcript-only until a real hearing-level link is validated.",
            "A `manual_review` row should not be promoted until the case-number or hearing evidence is clarified.",
            "This step keeps Perisic transcript-only even if the official page lists a broader case family.",
        ],
    )
    add_para(
        doc,
        "The manual execution sequence is: 1) run the refresh wrapper, 2) inspect the review CSV and summary JSON, 3) confirm which rows are transcript-only versus TAP-bearing, and 4) only then move to the broadened-manifest build. If you need to explain why this exists, say that it prevents the planning layer from being driven by stale candidate-ledger assumptions when the live tribunal websites already show a more accurate case universe.",
    )
    add_para(
        doc,
        "In the current repo state, this refresh step is what keeps `Perisic Trial` documented as transcript-only. That matters because the site-level discovery can expose a tribunal family without proving it is video-bearing. The SOP needs to say that clearly, because source broadening is about planning coverage, not about claiming that corpus rows already exist.",
    )

    doc.add_heading("21. Rank Tribunal Families for Source Broadening", level=1)
    add_para(
        doc,
        "If the expanded planning comparison shows that the current planning shortlist is exhausted, the next question is which tribunal families should be added to the planning source set. This ranking step reads the human-curated candidate ledger and compares it against the current tribunal source manifest. In student language, it turns a long candidate list into a short decision list by separating families that are already represented, families that should be held for link validation, and families that are ready for a broadened source manifest.",
    )
    add_para(
        doc,
        "The exact command is `bash phase2/run_rank_tribunal_source_broadening_candidates.sh`. That wrapper runs `phase2/rank_tribunal_source_broadening_candidates.py`, which removes candidate-ledger rows already present in `tribunal_sources_target_dataset.csv` and ranks the rest by candidate priority, whether video is plausibly available, whether the notes point to a real UCR search path, and the estimated TAP coverage. The output is a review CSV plus a JSON summary. This step is not a download step; it is a source-manifest decision step.",
    )
    add_table(
        doc,
        ["Output", "What it shows", "How to interpret it"],
        [
            [
                "`tribunal_source_broadening_review.csv`",
                "Candidate-ledger families that are not already in the current tribunal source manifest.",
                "Use this to decide which tribunal family should be added next.",
            ],
            [
                "`tribunal_source_broadening_review_summary.json`",
                "Counts of broaden-now, hold, and manual-review rows.",
                "Use this as the short summary before building a broadened manifest.",
            ],
        ],
    )
    add_bullets(
        doc,
        [
            "A `broaden_now` row is the cleanest immediate addition.",
            "A `hold_for_link_validation` row stays out until the UCR link evidence is stronger.",
            "A `manual_review` row should not be promoted without a human decision.",
            "This step does not download media; it only changes the planning source set.",
            "Perisic stays transcript-only and should not be pushed into TAP filtering as a video-bearing case.",
        ],
    )
    add_para(
        doc,
        "For the current repo state, the broadening review points to `Perisic Trial` as the immediate new tribunal family to add at the planning level, but only as a transcript-oriented source. `Stanisic and Simatovic Trial` stays on hold because the link and case-number relationship needs more care, and the other reviewed families are either already represented in the source manifest or are not yet safe to promote. That is why the next step is to broaden the source manifest with the single `broaden_now` family first, while still keeping Perisic outside the TAP shortlist.",
    )
    add_para(
        doc,
        "After that manifest is widened, the planning comparison can grow even if the missing-source count stays the same. In the current run, adding `Perisic Trial` increased the number of expanded rows already covered by the verified inventory from 4 to 5. That is a good sign that the source layer is being broadened correctly, but it also shows why source broadening and inventory-missing analysis are separate decisions: one changes the planning input, while the other checks whether the planning rows are truly new or already verified.",
    )

    doc.add_heading("22. Build the Broadened Tribunal Source Manifest", level=1)
    add_para(
        doc,
        "This is the operational step that converts the ranking decision into a new planning input file. The broadened manifest keeps the original source rows and appends only the approved new family or families. In student language, this is where a careful manual decision becomes a reusable manifest that the planning pipeline can consume on the next run.",
    )
    add_para(
        doc,
        "The exact command is `bash phase2/run_build_broadened_tribunal_source_manifest.sh`. That wrapper runs `phase2/build_broadened_tribunal_source_manifest.py`, which reads the current tribunal source manifest and the broadening review CSV, adds the selected new family rows, and writes `data/phase2/source_manifests/tribunal_sources_target_dataset_broadened.csv`. After that, the next planning expansion should be run with `TRIBUNAL_SOURCES=data/phase2/source_manifests/tribunal_sources_target_dataset_broadened.csv bash phase2/run_expand_phase2_planning_manifests.sh`. The `TRIBUNAL_SOURCES` option is the important part because it tells the wrapper to use the broadened source manifest instead of the older planning input.",
    )
    add_table(
        doc,
        ["Output", "What it shows", "How to interpret it"],
        [
            [
                "`tribunal_sources_target_dataset_broadened.csv`",
                "The original tribunal source manifest plus the approved new family or families.",
                "Use this as the next planning input, not as download proof.",
            ],
            [
                "`phase2_expanded_planning_manifest.csv` after rerun",
                "The planning shortlist regenerated from the broadened source manifest.",
                "Use this to check whether the new family actually contributes new tribunal rows.",
            ],
        ],
    )
    add_bullets(
        doc,
        [
            "Do not add held rows just because they exist in the review file.",
            "Do not use the broadened manifest to claim corpus growth until the next comparison step shows new coverage.",
            "If the broadened manifest is unchanged, the manual decision step needs to be revisited.",
        ],
    )
    add_para(
        doc,
        "The manual step here is: inspect the review CSV, confirm that `Perisic Trial` is the family to add now, run the broadened-manifest builder, and then re-run the planning expansion with the broadened source file. In a guidance call, you can describe this as the point where the planning layer moves from a short bootstrap list to a slightly wider tribunal source set while still keeping every added family traceable back to the candidate ledger and the UCR verification logic.",
    )
    add_para(
        doc,
        "If you want to explain the result clearly, say that the broadened source manifest is now wider than the old one, and the rebuilt planning manifest contains 18 rows instead of 17. The comparison against the verified inventory now reports 5 already-covered rows and 13 genuinely new expansion-source rows. That means the repo has moved one step forward in planning breadth, but it has not yet turned the broader source manifest into new downloadable corpus volume. The next comparison and download work still have to be grounded row by row.",
    )

    doc.add_heading("23. Append Verified TAP-Bearing Cases", level=1)
    add_para(
        doc,
        "This step is for cases that have already been confirmed through the UCR APIs and the official tribunal pages as having public TAP recordings. It is different from the broadening review step because it is not asking whether a family might be promising; it is saying that the family already has verified video evidence and should therefore be included in the reusable source layer.",
    )
    add_para(
        doc,
        "The exact input file is `data/phase2/source_manifests/verified_tap_case_additions.csv`. That file currently holds `Nahimana et al. (ICTR-99-52)` and `Karemera et al. (ICTR-98-44)`, both of which were probed directly in UCR and returned public `TAP` documents. The exact command path is the same broadened-manifest builder, `bash phase2/run_build_broadened_tribunal_source_manifest.sh`, because the builder now merges the base manifest, the official broadening review, and the verified additions list into one reusable output file.",
    )
    add_table(
        doc,
        ["Output", "What it shows", "How to interpret it"],
        [
            [
                "`verified_tap_case_additions.csv`",
                "A tiny curated list of cases already verified to have public TAP recordings.",
                "Use this as a reusable seed list, not as a planning guess.",
            ],
            [
                "`tribunal_sources_target_dataset_broadened.csv` after rebuild",
                "The current base source manifest plus the reviewed broadening rows plus verified TAP-bearing additions.",
                "Use this as the planning source for the next expansion run.",
            ],
        ],
    )
    add_bullets(
        doc,
        [
            "Do not treat the additions file as a new download queue by itself; it is a source-manifest input.",
            "Do not remove the transcript-only broadening rows; they still matter for planning and traceability.",
            "Use this step to make the source manifest cover both older valid cases and the two newly confirmed TAP-bearing cases.",
        ],
    )
    add_para(
        doc,
        "For a student explanation, say that this is the point where you stop relying only on the original bootstrap cases and the transcript-only expansion rows, and you add two cases that are already proven to have public video recordings. That gives the pipeline a better chance of reaching the 120 to 150 hour target because the source set now includes both the older ground truth cases and the newly verified witnesses-heavy cases.",
    )

    doc.add_heading("24. Discover More Video-Bearing Tribunal Cases", level=1)
    add_para(
        doc,
        "After broadening the tribunal source manifest, the next useful check is to ask which of the remaining missing-source rows are actually video-bearing and which ones are transcript-only or unresolved. This discovery step reads `expanded_planning_missing_sources.csv` together with the candidate ledger, resolves each row against the live UCR APIs, and classifies it by media evidence. In student language, this is the stage where you test whether a missing tribunal family can really support new video work, or whether it should stay in a transcript-only or manual-search branch.",
    )
    add_para(
        doc,
        "The exact command is `bash phase2/run_discover_tribunal_media_candidates.sh`. That wrapper runs `phase2/discover_tribunal_media_candidates.py`, which tries to resolve each missing-source family to a real case number, pulls case-level documents from UCR, counts TAP documents, and records whether the family is video-bearing, transcript-only, or unresolved. The output is a media-discovery review CSV plus a summary JSON. This is not a TAP filter replacement; it is a separate discovery pass that helps you decide what should be added next to the planning layer.",
    )
    add_table(
        doc,
        ["Output", "What it shows", "How to interpret it"],
        [
            [
                "`tribunal_media_discovery.csv`",
                "Each missing-source family plus its resolved media status.",
                "Use this to see whether a family is video-bearing, transcript-only, or unresolved.",
            ],
            [
                "`tribunal_media_discovery_summary.json`",
                "Counts of video-bearing, transcript-only, no-document, and unresolved rows.",
                "Use this to decide whether more tribunal source broadening is worth doing.",
            ],
        ],
    )
    add_bullets(
        doc,
        [
            "Treat `video_bearing` as the only immediate signal that a family might support a future download test.",
            "Treat `transcript_only` as useful for text-first work, but not for TAP or video planning.",
            "Treat `unresolved` as a manual-search problem, not as corpus growth.",
            "Keep the TAP shortlist unchanged; this discovery pass only informs future planning.",
        ],
    )
    add_para(
        doc,
        "For the current repo state, the discovery run on the 13 missing-source rows found 4 video-bearing rows, 3 transcript-only rows, and 6 unresolved rows. The video-bearing families were `Karadzic` and `Bagosora et al.`; the transcript-only families were `Mladic` and `Popovic et al.`. The unresolved rows were the cases that still need better case-number resolution or a more specific tribunal path before they can be judged as video-bearing. This is important because it gives you a repeatable way to grow the tribunal bootstrap set without pretending that every case page is a valid video source.",
    )
    add_bullets(
        doc,
        [
            "Use `video_bearing` rows as the only immediate candidates for future download tests.",
            "Use `transcript_only` rows for text-first work, but keep them out of TAP planning.",
            "Use `unresolved` rows as manual-search items that need a better case-number path.",
        ],
    )

    doc.add_heading("23. Resolve Unresolved Tribunal Families with Better Case-Number Hints", level=1)
    add_para(
        doc,
        "Some missing-source families are not yet video-bearing or transcript-only because the first-pass case-number mapping is too weak. This step tries a stronger, family-specific case-number resolution before you decide whether to keep searching or to stop. In student language, this is the repair step for cases that the generic hint map could not resolve cleanly the first time.",
    )
    add_para(
        doc,
        "The exact command is `bash phase2/run_resolve_targeted_tribunal_case_numbers.sh`. That wrapper runs `phase2/resolve_targeted_tribunal_case_numbers.py`, which is focused on `Akayesu`, `Ntakirutimana / Ntahobali related cases`, and `IRMCT Hearings`. The resolver uses the stronger case-number hints that are already known from the public tribunal case lists: `ICTR-96-04` for Akayesu, `ICTR-96-17` for Ntakirutimana et al., and `ICTR-98-42` for the Nyiramasuhuko / Ntahobali Butare case family. If a row resolves, the script records whether UCR returns transcripts, TAP documents, or only appeal/judgment material. If it does not resolve, the row stays manual-search only.",
    )
    add_table(
        doc,
        ["Output", "What it shows", "How to interpret it"],
        [
            [
                "`tribunal_case_resolution_review.csv`",
                "The unresolved families plus stronger case-number candidates and UCR resolution results.",
                "Use this to see whether the first-pass unresolved rows can be repaired before any new corpus expansion.",
            ],
            [
                "`tribunal_case_resolution_review_summary.json`",
                "Counts of resolved and unresolved families plus the case numbers that worked.",
                "Use this as the quick decision file before any manual follow-up.",
            ],
        ],
    )
    add_bullets(
        doc,
        [
            "If the script resolves a case number, keep that case number for the next planning or media probe.",
            "If it only finds transcripts, keep the family out of TAP planning.",
            "If it still cannot resolve the family, treat it as a manual-search item instead of a download target.",
        ],
    )
    add_para(
        doc,
        "For the current repo state, the resolver returned three successful resolution rows across two real case numbers: `ICTR-96-04` for Akayesu and `ICTR-96-17` for Ntakirutimana et al. The remaining IRMCT Hearings rows stayed unresolved because they still do not map to a single stable case number. That is not a failure; it is the expected sign that the row needs a more precise tribunal identifier before it can be safely used.",
    )
    add_para(
        doc,
        "Because `IRMCT Hearings` is still a generic umbrella label, the repo now has a separate MICT hearings candidate manifest step. That step uses official MICT case pages to name specific case numbers such as `MICT-13-38` for Kabuga, `MICT-15-96` for Stanišić and Simatović, `MICT-12-29` for Ngirabatware, `MICT-12-17-R` for Ntakirutimana, and `MICT-17-111-R90` for Jojić and Radeta. This is a better way to explain the MICT branch in a guidance call because it shows specific tribunal identifiers instead of a generic placeholder.",
    )

    doc.add_heading("24. Build the MICT Hearings Candidate Manifest", level=1)
    add_para(
        doc,
        "This step takes the generic IRMCT hearing placeholder and turns it into a small, explicit MICT candidate list. The point is not to claim that every MICT matter is video-bearing; the point is to replace an umbrella label with concrete case numbers that can be inspected separately. In student terms, this is the bridge between a vague row and a real tribunal identifier.",
    )
    add_para(
        doc,
        "The exact command is `python3 phase2/build_mict_hearings_candidate_manifest.py` or the new wrapper `bash phase2/run_build_mict_hearings_candidate_manifest.sh`. That script writes a compact CSV of official MICT candidates and a JSON summary. It does not download media; it only creates a case-specific manifest that can be used in later planning, reading, or manual verification.",
    )
    add_table(
        doc,
        ["Output", "What it shows", "How to interpret it"],
        [
            [
                "`mict_hearings_candidate_manifest.csv`",
                "Specific MICT hearing/proceeding rows with case numbers and official evidence links.",
                "Use this when you need to explain the MICT branch as concrete tribunal cases instead of a generic placeholder.",
            ],
            [
                "`mict_hearings_candidate_manifest_summary.json`",
                "Counts of candidate rows plus the specific case numbers included.",
                "Use this as the quick check that the manifest contains the intended MICT cases.",
            ],
        ],
    )
    add_bullets(
        doc,
        [
            "This step is a planning and explanation aid, not a download or corpus-growth step.",
            "It keeps the generic `IRMCT Hearings` label out of any place where a case number is required.",
            "It is useful when a guidance call needs a concrete example of how unresolved tribunal families are turned into inspectable rows.",
        ],
    )

    doc.add_heading("25. Filter the Discovery Output to the Video-Bearing Shortlist", level=1)
    add_para(
        doc,
        "After the discovery pass, the next reusable step is to compress the review file into a smaller shortlist containing only the rows that are actually video-bearing. This keeps the next planning and download decision focused on the families that still have real media potential. In student language, this is the file you would open when you want to ask, 'Which missing tribunal families are still worth a video download test?'",
    )
    add_para(
        doc,
        "The exact command is `bash phase2/run_filter_tribunal_video_bearing_candidates.sh`. That wrapper runs `phase2/filter_tribunal_video_bearing_candidates.py`, which reads `tribunal_media_discovery.csv`, keeps only rows where `media_status = video_bearing`, and writes `data/processed/phase2/tribunal_video_bearing_candidates.csv` plus a small JSON summary. This is a filter step, not a new UCR probe, so it is fast and reusable after the discovery pass has already done the hard work.",
    )
    add_table(
        doc,
        ["Output", "What it shows", "How to interpret it"],
        [
            [
                "`tribunal_video_bearing_candidates.csv`",
                "Only the media-discovery rows classified as `video_bearing`.",
                "Use this as the compact shortlist for the next tribunal video planning decision.",
            ],
            [
                "`tribunal_video_bearing_candidates_summary.json`",
                "Counts of inspected rows and video-bearing rows.",
                "Use this as the quick status check after the discovery filter runs.",
            ],
        ],
    )
    add_bullets(
        doc,
        [
            "This shortlist is smaller than the full media-discovery file.",
            "It should contain only rows that already look video-bearing in UCR.",
            "Transcript-only and unresolved rows stay in the discovery review, not in this compact shortlist.",
        ],
    )
    add_para(
        doc,
        "For the current repo state, this filtered shortlist should contain the `Karadzic` and `Bagosora et al.` families only. That is the right behavior because the discovery run found those as the video-bearing missing-source families, while `Mladic` and `Popovic et al.` remained transcript-only and the others stayed unresolved. This final filter gives you a clean handoff point for any later download test or planning revision.",
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
