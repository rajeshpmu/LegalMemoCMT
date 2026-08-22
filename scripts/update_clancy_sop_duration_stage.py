from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOP = ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx"


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = "Menlo"
    r.font.size = Pt(8.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate(headers):
        t.cell(0, i).text = value
        for run in t.cell(0, i).paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(20, 48, 87)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def build() -> None:
    doc = Document(SOP)
    doc.add_page_break()
    doc.add_heading("Latest Controlled Duration-Subset Procedure", level=1)
    doc.add_paragraph(
        "This section records the next step after persistent rejection filtering. The purpose is to create two controlled Clancy candidate pools before applying any Phase 1 MELD model. The duration rule is a screening rule, not a semantic guarantee: a row still needs transcript, audio, video, alignment, speaker, and content review."
    )

    doc.add_heading("1. Starting Artifact and Current Baseline", level=2)
    doc.add_paragraph(
        "Use data/processed/phase2/clancy/clancy_turn_manifest_post_rejection.csv as the input. This file is preferred over clancy_turn_manifest_clipped.csv because it removes the persistent rejection IDs while preserving the original media files for audit. The current post-rejection population is 11,892 rows and approximately 25.62 candidate clip hours."
    )
    doc.add_paragraph("The two non-overlapping duration categories are:")
    for value in [
        "Pool A: 0.8 <= clip_duration_seconds < 20.0",
        "Pool B: 20.0 <= clip_duration_seconds <= 30.0",
        "Rows below 0.8 seconds or above 30 seconds remain outside the initial MELD-style training window and require review, splitting, or rejection.",
    ]:
        doc.add_paragraph(value, style="List Bullet")

    doc.add_heading("2. Inspect Duration Outliers First", level=2)
    doc.add_paragraph(
        "This step identifies the long and short rows before creating the two pools. The IQR report is a review aid. It must not automatically reject every statistical outlier because a long row may contain valid testimony that should be split."
    )
    add_code(doc, "PYTHON_BIN=$PWD/.venv/bin/python \\\nbash phase2/run_build_clancy_duration_outlier_report.sh")
    doc.add_paragraph("The command reads the post-rejection manifest and writes:")
    for value in [
        "data/processed/phase2/clancy/clancy_duration_outliers.csv",
        "reports/phase2/clancy_duration_outlier_summary.json",
    ]:
        doc.add_paragraph(value, style="List Bullet")
    add_table(doc, ["Reported measure", "Current result", "Student interpretation"], [
        ["Rows analysed", "11,892", "Current post-rejection candidate population."],
        ["IQR outlier rows", "924", "Rows requiring duration/content inspection."],
        ["Rows over 30 seconds", "443", "Do not automatically train as one MELD-style turn."],
        ["Rows over 60 seconds", "110", "High-priority split or manual-review group."],
        ["Rows over 300 seconds", "0", "No remaining five-minute turn candidates in this report."],
    ])

    doc.add_heading("3. Create the Two Duration Subsets", level=2)
    doc.add_paragraph(
        "The repository currently contains phase2/create_two_diration_subsets_clancy.py. The filename contains the historical spelling `diration`; retain that exact filename when following the current procedure unless it is later renamed deliberately with a compatibility note."
    )
    doc.add_paragraph(
        "The pasted attempt used `PYTHON_BIN=$PWD/.venv/bin/python phase2/create_two_diration_subsets_clancy.py`. That does not invoke Python. It sets an environment variable and asks the shell to execute the `.py` file directly, which explains the `Permission denied` message. `chmod +x` makes direct execution possible only if the file has a valid shebang; it is not necessary when the interpreter is called explicitly."
    )
    doc.add_paragraph("Use this corrected command:")
    add_code(doc, "PYTHON_BIN=$PWD/.venv/bin/python\n\"$PYTHON_BIN\" phase2/create_two_diration_subsets_clancy.py")
    doc.add_paragraph("The Python script:")
    for value in [
        "Reads clancy_turn_manifest_post_rejection.csv.",
        "Uses clip_duration_seconds as the numeric duration field.",
        "Writes every original column so source, transcript, split, and media traceability are retained.",
        "Writes Pool A to data/processed/phase2/clancy/clancy_turn_manifest_0_8_to_20.csv.",
        "Writes Pool B to data/processed/phase2/clancy/clancy_turn_manifest_20_to_30.csv.",
        "Prints row counts, minutes, and hours for each output.",
    ]:
        doc.add_paragraph(value, style="List Bullet")
    add_table(doc, ["Output", "Rows", "Minutes", "Hours", "Boundary"], [
        ["clancy_turn_manifest_0_8_to_20.csv", "9,073", "912.600", "15.2100", "0.8 inclusive, 20 exclusive"],
        ["clancy_turn_manifest_20_to_30.csv", "551", "220.891", "3.6815", "20 inclusive, 30 inclusive"],
    ])

    doc.add_heading("4. Validate Both Subsets", level=2)
    doc.add_paragraph(
        "Validation checks that the subset operation did not lose modality files or place a row in the wrong duration range. It is a structural check, not a claim that the transcript is semantically correct."
    )
    add_code(doc, "PYTHON_BIN=$PWD/.venv/bin/python\n\"$PYTHON_BIN\" phase2/validate_two_duration_subsets_clancy.py")
    add_table(doc, ["File", "Rows", "Missing video", "Missing audio", "Bad duration rows"], [
        ["0.8–<20 seconds", "9,073", "0", "0", "0"],
        ["20–30 seconds", "551", "0", "0", "0"],
    ])
    doc.add_paragraph(
        "A zero missing-media count means the manifest paths resolve on this machine. It does not prove that the audio and video content are synchronized with the text; that still requires manual sample inspection."
    )

    doc.add_heading("5. Generate Leakage-Aware Splits", level=2)
    doc.add_paragraph(
        "phase2/generate_dur_specific_split.sh is a convenience wrapper containing two separate calls to run_build_clancy_dataset_split.sh. The first call processes Pool A and the second processes Pool B. The script assigns complete YouTube source groups to train, dev, and test rather than randomly distributing individual rows."
    )
    add_code(doc, "bash phase2/generate_dur_specific_split.sh")
    add_table(doc, ["Pool", "Train", "Dev", "Test", "Groups", "Reason"], [
        ["0.8–<20 seconds", "6,245", "1,262", "1,566", "11", "Primary candidate pool with shorter turns."],
        ["20–30 seconds", "400", "76", "75", "11", "Secondary pool requiring stronger boundary review."],
    ])
    doc.add_paragraph(
        "The split summary reports `group_column=youtube_id`. This means all rows from a source video remain in one partition. The split is therefore safer than random row splitting, although it does not automatically prove speaker disjointness when speaker IDs are unavailable."
    )

    doc.add_heading("6. Manual Inspection Gate", level=2)
    doc.add_paragraph("Before using either pool for model-generated labels, inspect rows hearing/source by source. At minimum:")
    for value in [
        "Inspect 10 rows from Pool A and 10 rows from Pool B.",
        "Inspect examples from all 11 source videos.",
        "Include witness, prosecutor, defence, and judge turns where available.",
        "Check that the spoken words match utterance_text or turn_text.",
        "Check that the clip does not include a previous or following speaker turn.",
        "Reject or record breaks, lunch periods, news updates, courtroom transitions, and non-testimony material.",
        "Confirm that visual analysis is only claimed when the speaking person is visible.",
    ]:
        doc.add_paragraph(value, style="List Bullet")
    doc.add_paragraph(
        "If a new problem is confirmed, add the turn ID to clancy_turn_rejection_manifest.csv with a specific reason, then rerun the turn-manifest, post-rejection, duration-subset, validation, and split stages. Do not edit only the final training CSV because that breaks reproducibility."
    )

    doc.add_heading("7. Gate Before Phase 1 Model Use", level=2)
    doc.add_paragraph(
        "The Phase 1 MELD model should be used only after the two duration pools pass the structural and manual gates. Its output will be a weak basic-emotion suggestion, not courtroom-affect ground truth."
    )
    for value in [
        "Keep emotion_label for existing code compatibility.",
        "Store model-generated basic emotion separately with checkpoint, probability, and provenance fields.",
        "Do not automatically convert stress or confidence into the seven MELD emotions.",
        "Do not use the Phase 1 model to infer courtroom affect, deception, truthfulness, or credibility.",
        "Keep courtroom_affect UNKNOWN until a documented courtroom-specific annotation process is completed.",
    ]:
        doc.add_paragraph(value, style="List Bullet")
    doc.add_paragraph(
        "This gate prevents a common error: treating a valid duration and a Phase 1 model prediction as proof that a courtroom row has a correct legal-domain emotion label. The duration subsets are a controlled input population; they are not already gold-labelled training data."
    )

    doc.add_heading("8. Correct Manual Sequence", level=2)
    add_code(doc, "cd /Users/rajeshpmu/Desktop/LegalMemoCMT\n\nPYTHON_BIN=$PWD/.venv/bin/python bash phase2/run_build_clancy_duration_outlier_report.sh\n\nPYTHON_BIN=$PWD/.venv/bin/python\n\"$PYTHON_BIN\" phase2/create_two_diration_subsets_clancy.py\n\"$PYTHON_BIN\" phase2/validate_two_duration_subsets_clancy.py\n\nbash phase2/generate_dur_specific_split.sh")
    doc.add_paragraph(
        "Important shell correction: do not concatenate `PYTHON_BIN=...` and the interpreter without a separating newline or semicolon. The safe form is the two-line form shown below, where the variable is assigned first and then expanded as `\"$PYTHON_BIN\"`."
    )
    add_code(doc, "PYTHON_BIN=$PWD/.venv/bin/python\n\"$PYTHON_BIN\" phase2/create_two_diration_subsets_clancy.py\n\"$PYTHON_BIN\" phase2/validate_two_duration_subsets_clancy.py")

    doc.add_heading("9. Current Status Statement", level=2)
    doc.add_paragraph(
        "The Clancy corpus now has two duration-controlled candidate pools: 9,073 rows in the 0.8–<20 second pool and 551 rows in the 20–30 second pool. Both pools passed the current path and boundary validation with zero missing video paths, zero missing audio paths, and zero bad-duration rows. The next action is manual content review and provenance-controlled basic-emotion pseudo-labelling; courtroom-affect annotation remains a separate future stage."
    )

    doc.save(SOP)
    print(f"Updated {SOP}")


if __name__ == "__main__":
    build()
