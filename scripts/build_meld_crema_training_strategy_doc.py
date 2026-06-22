from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_MELD_vs_CREMA_D_Training_Strategy_Deep_Dive.docx"


def style_document(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def add_section(doc: Document, title: str, paragraphs: list[str]) -> None:
    doc.add_heading(title, level=1)
    for para in paragraphs:
        add_para(doc, para)


def build_doc() -> Document:
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run("MELD vs CREMA-D Training Strategy Deep Dive")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = 1
    run = subtitle.add_run(
        "Why MELD remains the primary paper-aligned training target, when CREMA-D is useful, and how imbalance-aware training should be applied in the current LegalMemoCMT pipeline."
    )
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)

    add_para(
        doc,
        "This document answers a practical question that appears immediately once the dataset statistics are inspected: MELD is visibly dominated by Neutral emotion, while CREMA-D appears cleaner and more balanced. That observation is correct, but it does not automatically imply that CREMA-D should replace MELD for training. The right answer depends on the research objective, the target benchmark, the type of signal the model is meant to learn, and the evaluation protocol the project is trying to reproduce.",
    )
    add_para(
        doc,
        "The short version is simple. If the goal is to reproduce the MemoCMT paper and study conversational multimodal emotion recognition, MELD should remain the main training target. If the goal is to establish a cleaner public baseline or to pretrain audiovisual representations before moving to a harder conversational dataset, CREMA-D is useful. The project should treat CREMA-D as a secondary benchmark, not a replacement for MELD.",
    )

    add_section(
        doc,
        "1. Executive Position",
        [
            "MELD should remain the main benchmark because it matches the research problem the project is actually trying to solve: multimodal emotion recognition in dialogue. In MELD, each utterance is part of a conversation, and the model must interpret emotional state in the presence of surrounding turns, speaker changes, and conversational context.",
            "CREMA-D should be retained as a secondary benchmark because it is still scientifically valuable. It is easier to reproduce, has cleaner labels, and often produces more stable optimization behavior. Those properties make it useful for validating that the model learns robust audiovisual cues, but they do not make it an appropriate substitute for MELD when the core claim is about dialogue-level emotion recognition.",
            "The imbalance in MELD should be addressed through the learning objective and the evaluation metrics, not by changing the benchmark to a fundamentally different dataset. Weighted loss, focal loss, and macro-F1 are the main tools for that purpose.",
        ],
    )

    add_section(
        doc,
        "2. Why Dataset Choice Changes the Problem",
        [
            "A dataset is not just a collection of samples. It defines the inductive bias of the model, the type of error the model is likely to make, and the kind of result that can be meaningfully reported. CREMA-D and MELD are both multimodal emotion datasets, but they encode different scientific questions.",
            "MELD is conversational and multiparty. Each utterance is embedded in dialogue history, turn-taking, speaker context, and pragmatic flow. This means that emotion recognition is not just a matter of detecting facial expression or vocal tone; it also depends on how the utterance relates to what was said before and who said it.",
            "CREMA-D is closer to an utterance-level audiovisual recognition problem. The labels are cleaner and the class distribution is easier, but the conversational complexity is much lower. A system that performs well on CREMA-D may still fail to capture dialogic emotion dynamics in MELD.",
            "That difference matters because the MemoCMT paper is not simply about classifying isolated emotional utterances. It is about multimodal interaction, and MELD is the benchmark that actually exercises that interaction in context.",
        ],
    )

    add_table(
        doc,
        ["Property", "MELD", "CREMA-D", "Implication"],
        [
            ["Structure", "Multi-party dialogue", "Single-utterance / acted clips", "MELD tests context reasoning; CREMA-D tests cleaner audiovisual pattern learning"],
            ["Label distribution", "Highly Neutral-dominated", "More balanced", "MELD needs imbalance-aware training; CREMA-D is easier to optimize"],
            ["Context", "Dialogue turns, speaker ID, turn order", "Minimal conversational context", "MELD better matches conversation-aware fusion"],
            ["Evaluation difficulty", "Harder", "Easier", "CREMA-D can inflate apparent model quality"],
            ["Reproducibility role", "Primary paper-aligned benchmark", "Secondary public benchmark", "Do not replace MELD if paper fidelity matters"],
        ],
    )

    add_section(
        doc,
        "3. Neutral Dominance in MELD Is a Training Signal, Not a Reason to Abandon MELD",
        [
            "The fact that Neutral dominates MELD is not a defect in the dataset; it is a characteristic of the target distribution. Conversational emotion recognition is often dominated by neutral or low-arousal statements because many turns in dialogue are informational, transitional, or structurally necessary rather than strongly emotional.",
            "This means that the model must learn to discriminate subtle emotional deviations from a neutral conversational baseline. That is a harder and more realistic problem than training on a dataset with more evenly distributed acted labels.",
            "If the project were to replace MELD with CREMA-D simply because MELD is imbalanced, the model would be optimized for the wrong operating conditions. It would learn from a cleaner distribution but would no longer be validated on the conversational setting that the paper and the repo are designed around.",
            "The correct interpretation is that Neutral imbalance informs the training strategy. It motivates class weighting, focal loss, macro-F1 reporting, and cross-validation discipline. It does not motivate changing the benchmark itself.",
        ],
    )

    add_section(
        doc,
        "4. Why MELD Remains the Main Paper-Aligned Result",
        [
            "The repository is aligned around the MemoCMT paper, and the paper’s MELD case study is specifically about text and speech fusion in a conversational setting. The paper compares fusion pooling options on MELD and reports that MIN gives the best test result among the listed variants. That is a benchmark-specific conclusion, not a generic one.",
            "Because the base paper uses MELD as an additional case study, the closest implementation path in this project is to keep MELD as the primary reporting dataset and to reproduce the paper’s style of text-plus-speech fusion with a paper-aligned evaluation workflow.",
            "MELD also has a better conceptual match to the model design. A bidirectional cross-attention module is most meaningful when there is a plausible interaction between text and audio that may evolve across dialogue turns. That is exactly what MELD provides.",
            "If CREMA-D were used as the main benchmark, the project would drift away from the paper’s scenario. The resulting evaluation could still be useful, but it would no longer be the main paper-aligned result.",
        ],
    )

    add_section(
        doc,
        "5. Where CREMA-D Still Helps",
        [
            "CREMA-D is useful for two reasons. First, it provides a cleaner sanity-check benchmark where the model can be tested on audiovisual emotion recognition without the added difficulty of dialogue context. Second, it can be used as a secondary benchmark to verify that a pretrained text/audio architecture is not overfitting to MELD-specific conversational patterns.",
            "A useful way to think about CREMA-D is as a control experiment. If the model performs reasonably on CREMA-D, then the implementation likely learned something general about multimodal emotion cues. If it fails on CREMA-D, then the model may not have even learned reliable audiovisual emotion representations, which would be a separate problem from dialogue complexity.",
            "However, a strong CREMA-D result cannot replace a MELD result. The model may do very well on acted data and still struggle on conversational affect. This is why CREMA-D should sit beside MELD in the evaluation plan, not instead of it.",
        ],
    )

    add_section(
        doc,
        "6. What Weighted Loss Changes",
        [
            "Weighted loss corrects the fact that frequent classes dominate the gradient. In plain cross-entropy, a large class such as Neutral can consume much of the optimization pressure because the model sees it so often. The model then learns a safe but weak strategy: predict Neutral too often and collect a superficially acceptable accuracy score.",
            "When class weights are introduced, rare classes receive larger penalties when misclassified. This rebalances the gradient signal so that anger, sadness, joy, surprise, disgust, and fear are not ignored. The model is still free to predict Neutral often when it is correct, but it is no longer rewarded for collapsing minority emotions into the dominant class.",
            "Weighted cross-entropy is a natural first response when class imbalance is known and the label set is stable. It is easy to interpret, easy to implement, and less sensitive to the hyperparameter tuning burden than more complex imbalance remedies.",
            "In the current codebase, weighted loss is the default recommendation for MELD because it matches the observed imbalance and does not alter the benchmark definition.",
        ],
    )

    add_section(
        doc,
        "7. When Focal Loss Is Better",
        [
            "Focal loss is useful when the model quickly becomes confident on easy examples and then stops paying attention to hard examples. In imbalanced emotion recognition, that can happen when Neutral and a few frequent classes dominate the training signal.",
            "Focal loss down-weights easy examples and concentrates the gradient on misclassified or uncertain examples. That makes it a better fit when the failure mode is not only class frequency but also the model’s tendency to over-fit obvious neutral cases and neglect subtle emotional classes.",
            "The tradeoff is that focal loss introduces another hyperparameter, usually gamma, which needs tuning. Too little focusing and it behaves like weighted cross-entropy; too much focusing and training can become noisy or unstable.",
            "For this project, focal loss is best treated as an ablation or an alternate training regime, while weighted cross-entropy remains the simpler default. The useful teaching point is that both losses address imbalance, but they do so in slightly different ways.",
        ],
    )

    add_section(
        doc,
        "8. Why Macro-F1 Matters More Than Accuracy Alone",
        [
            "Accuracy is easy to understand, but it can be misleading under class imbalance. A model that does very well on Neutral and weakly on minority emotions may still report a respectable accuracy because Neutral occupies a large fraction of the dataset.",
            "Macro-F1 gives each class equal weight in the averaging process. That means a model cannot hide poor performance on rare classes behind a strong Neutral score. If the minority emotions are not being learned, macro-F1 will reveal that weakness.",
            "This is the right metric for the MELD problem because the scientific question is not simply whether the model can exploit class frequency. The question is whether it can distinguish a broad range of emotional states in a conversational environment.",
            "Weighted F1 and accuracy are still useful, but they should be read together with macro-F1. In practice, the gap between weighted and macro metrics is often the clearest sign that the model is over-relying on the dominant class.",
        ],
    )

    add_section(
        doc,
        "9. A Practical Training Strategy",
        [
            "The most defensible strategy is to train on MELD as the primary benchmark, use a paper-aligned text-plus-audio setup, and choose a loss function that explicitly addresses imbalance. The training loop should report both overall and class-balanced metrics, and the evaluation should be separated cleanly from training so that validation decisions do not leak into test reporting.",
            "A strong implementation strategy has four layers: data preparation, imbalance-aware training, fold-based validation, and final held-out evaluation. First, the dataset is made manifest-driven so the sample mapping is explicit. Second, the loss function is adjusted to prevent Neutral collapse. Third, the MELD train/dev dialogues are split into folds for more robust estimates. Fourth, the official MELD test split is reserved for the final score.",
            "This strategy is superior to simply switching to CREMA-D because it preserves the benchmark’s difficulty and keeps the project aligned with the paper. It also makes error analysis more meaningful, since the model’s behavior on rare MELD emotions is exactly the part that needs scrutiny.",
            "A practical implementation rule is: use MELD for the main experiments, report macro-F1 prominently, and then use CREMA-D to confirm that the system also learns robust audiovisual affect cues in a cleaner setting.",
        ],
    )

    add_table(
        doc,
        ["Training choice", "Benefit", "Risk"],
        [
            ["Train on MELD", "Matches the paper and the conversational setting", "Harder optimization because of Neutral dominance"],
            ["Weighted cross-entropy", "Rebalances frequent and rare classes", "May still favor easy classes if the model is too confident"],
            ["Focal loss", "Forces attention on difficult examples", "Needs tuning and can be unstable if overused"],
            ["Use macro-F1", "Reveals minority-class behavior", "Can look pessimistic if the model is already strong on dominant classes"],
            ["Use CREMA-D as secondary benchmark", "Validates audiovisual learning on a cleaner dataset", "Does not replace conversational generalization"],
        ],
    )

    add_section(
        doc,
        "10. When Pretraining on CREMA-D Makes Sense",
        [
            "A more nuanced option is pretraining on CREMA-D and fine-tuning on MELD. This is a reasonable strategy if the goal is to help the model learn speech and audiovisual affect patterns before exposing it to the harder conversational setting.",
            "That approach can be beneficial when the model is small, the optimization is unstable, or the pretrained backbones need a stronger emotion-specific adaptation signal. CREMA-D can function as a cleaner source of supervised emotion learning before the model is asked to reason over MELD dialogue context.",
            "Even in that scenario, MELD remains the benchmark that matters most. CREMA-D is then a stepping stone, not the endpoint. The final evaluation should still be on MELD, because that is the benchmark that corresponds to the paper-aligned research question.",
            "If pretraining is used, the implementation should keep the two stages explicit in the documentation so that readers do not confuse an auxiliary optimization step with the main benchmark result.",
        ],
    )

    add_section(
        doc,
        "11. Why the Current Paper-Aligned Pipeline Uses MELD 5-Fold CV",
        [
            "The purpose of cross-validation is to reduce variance in the estimate of model quality. In a dataset like MELD, where there are many dialogues but also speaker and dialogue-specific patterns, a single train/dev split can overstate or understate performance depending on which dialogues end up in validation.",
            "Five folds are a practical compromise. They provide multiple validation views of the same benchmark without making the process prohibitively expensive. They also make it easier to compare pooling methods and loss functions under more than one split configuration.",
            "The paper itself discusses cross-validation as part of its evaluation design. Using five folds therefore improves alignment with the base methodology while also making the experiment more robust than a one-off split.",
            "In this repo, the new CV workflow is therefore the main paper-aligned path for MELD. It keeps the test split reserved, but it uses the train/dev portion more rigorously to stabilize model selection.",
        ],
    )

    add_section(
        doc,
        "12. Decision Rule for This Project",
        [
            "If the question is 'Which dataset should be used for the main paper-aligned result?', the answer is MELD.",
            "If the question is 'How should MELD’s Neutral imbalance be handled?', the answer is weighted loss, focal loss, and macro-F1 reporting.",
            "If the question is 'Should CREMA-D be thrown away because MELD is harder?', the answer is no. CREMA-D is valuable, but as a secondary benchmark.",
            "If the question is 'Should the project replace MELD with CREMA-D for main training?', the answer is also no. That would change the task and weaken paper alignment.",
            "If the question is 'Can CREMA-D be used before MELD?', the answer is yes, as a pretraining or diagnostic step, provided the final benchmark remains MELD.",
        ],
    )

    add_section(
        doc,
        "13. Recommended Implementation Policy",
        [
            "Use MELD for the main paper-aligned result. This preserves the conversational benchmark and keeps the project faithful to the MemoCMT framing.",
            "Use weighted loss by default and focal loss as a controlled ablation. This addresses the Neutral-heavy class distribution without changing the benchmark.",
            "Use macro-F1 as a primary reporting metric for model selection and interpretation. Accuracy alone is not enough under imbalance.",
            "Use CREMA-D as a secondary benchmark to validate general audiovisual emotion learning. It is complementary, not substitutive.",
            "If pretraining is introduced, document it as a separate stage and still report the final MELD result as the main outcome.",
        ],
    )

    doc.add_heading("14. Final Answer", level=1)
    add_para(
        doc,
        "No, MELD should not be replaced by CREMA-D just because MELD is dominated by Neutral. Neutral dominance is exactly why imbalance-aware training and better metrics are required. MELD remains the correct main training benchmark because it matches the paper’s conversational multimodal setting. CREMA-D should remain in the project as a useful secondary benchmark, or as a possible pretraining source, but not as the replacement for MELD.",
    )
    add_para(
        doc,
        "In other words: train on MELD for the main paper-aligned result, use weighted loss or focal loss together with macro-F1 to address imbalance, and keep CREMA-D as a complementary benchmark that helps verify whether the system learns robust audiovisual emotion cues outside the dialogue setting.",
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
