from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_Guidance_Call_Comparison_Report.docx"
INDEX = ROOT / "implementation_docments" / "LegalMemoCMT_Student_Document_Index.docx"

MELD_SUMMARY = ROOT / "results" / "paper_aligned_meld_cv" / "cmt_min" / "summary.json"
MELD_FOLD2_METRICS = ROOT / "results" / "paper_aligned_meld_cv" / "cmt_min" / "fold_2" / "metrics.json"
CREMA_METRICS = ROOT / "results" / "paper_aligned_crema_d" / "cmt_min" / "metrics.json"

BASE_PAPER_MELD = {
    "Acc": 64.18,
    "F1": 62.52,
    "Prec": 63.82,
    "Rec": 64.18,
}


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def build_doc() -> Document:
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LegalMemoCMT Guidance Call Comparison Report")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "A student-level comparison of the current system results against the MemoCMT base paper, with discussion points, gaps, and next-step questions for a mentor call."
    )
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)

    meld_summary = read_json(MELD_SUMMARY)["metrics"]
    meld_fold2 = read_json(MELD_FOLD2_METRICS)
    crema = read_json(CREMA_METRICS)

    add_para(
        doc,
        "Purpose: this document is designed for a guidance call. It summarizes what has been implemented so far, how the current results compare with the MemoCMT base paper, where the current system is close to the paper, where it is far away, and what reasonable next improvements should be considered. It also ends with concrete questions that can be asked to the mentor so the next stage is well guided rather than guessed.",
    )
    add_para(
        doc,
        "Important note: the MELD comparison is not an exact apples-to-apples reproduction of the paper because the current repository uses a 5-fold MELD cross-validation workflow while the paper reports its MELD case-study values in a different evaluation framing. Even so, the comparison is still informative because it shows whether the model is moving in the same direction as the paper and whether the current implementation is in the right performance range.",
    )
    add_para(
        doc,
        "Small metric-definition note: in the current codebase, weighted accuracy is reported as standard sample accuracy, while unweighted accuracy is the mean per-class accuracy over the classes that appear in the split. That is the right quantity to carry into the speech-emotion cross-validation summary, where the report should be read using the base-paper style terminology of W-Acc and UW-Acc.",
    )

    doc.add_heading("1. What Has Been Done So Far", level=1)
    add_bullets(
        doc,
        [
            "The MELD paper-aligned path was implemented with pretrained BERT for text and HuBERT for audio.",
            "The fusion block was changed to a bidirectional cross-attention CMT design, and MIN pooling is used for the paper-aligned MELD result.",
            "A raw MELD manifest and MELD 5-fold cross-validation workflow were added so the main MELD result could be evaluated more robustly.",
            "A MELD fold-level error analysis was completed, including confusion matrices and predicted-vs-actual tables.",
            "A CREMA-D paper-aligned run was added as a secondary benchmark and its analysis was completed.",
            "The project documentation was expanded so a student can read the workflow, the results, and the interpretation in a structured way.",
        ],
    )

    doc.add_heading("1.1 What the new benchmark scripts do", level=2)
    add_para(
        doc,
        "The new benchmark layer does not change the core model code. It changes how the existing implementation is organized for the new experimental goal: CREMA-D is now treated as the primary speech-emotion benchmark, and MELD is kept as the primary conversational benchmark. The scripts are therefore wrappers around the same pretrained text-plus-audio model, but they separate the datasets and the evaluation style so the reports are easier to interpret.",
    )
    add_para(
        doc,
        "There are two levels of script logic. The first level is the speech-emotion track, which is paper-style and centered on CREMA-D. The second level is the conversational track, which remains paper-aligned and centered on MELD. Both levels use the same underlying model family, but they differ in how the data is split and how the output should be read.",
    )
    add_bullets(
        doc,
        [
            "scripts/build_crema_d_cv_folds.py creates speaker-independent CREMA-D folds from the manifest.",
            "scripts/run_primary_speech_emotion_crema_d_cv.sh trains and evaluates CREMA-D in 5 CV folds.",
            "scripts/run_primary_speech_emotion_crema_d_cv_analysis.sh exports and analyzes one CREMA-D fold’s predictions.",
            "scripts/aggregate_crema_d_cv_metrics.py turns the CREMA-D fold metrics into a summary JSON and markdown report.",
            "scripts/run_primary_conversational_meld_cv.sh keeps the MELD 5-fold dialogue-based workflow.",
            "scripts/run_primary_conversational_meld_analysis.sh analyzes a representative MELD fold, usually fold 2.",
            "scripts/run_primary_benchmark_suite.sh runs both tracks one after the other.",
        ],
    )
    add_para(
        doc,
        "For a student, the easiest way to think about the new scripts is this: one group prepares and evaluates CREMA-D as a speech-emotion benchmark, and the other group prepares and evaluates MELD as a conversational benchmark. The scripts are not inventing a new model. They are organizing the same model into two benchmark stories that match the project’s goals more cleanly.",
    )

    doc.add_heading("1.2 Expected outputs from the new scripts", level=2)
    add_para(
        doc,
        "Each script writes files that make the run inspectable after it finishes. This matters because training alone is not enough to understand model behavior. You need the manifests, metrics, predictions, and confusion summaries to see whether the model is learning the right patterns or collapsing into one dominant class.",
    )
    add_table(
        doc,
        ["Script", "Primary outputs", "What the outputs are for"],
        [
            [
                "build_crema_d_cv_folds.py",
                "crema_d_fold_*_train.csv, crema_d_fold_*_val.csv, crema_d_cv_summary.json",
                "Creates speaker-independent CREMA-D folds so training can be done in 5 cross-validation rounds.",
            ],
            [
                "run_primary_speech_emotion_crema_d_cv.sh",
                "results/primary_speech_emotion/crema_d_cv/cmt_min/fold_*/metrics.json and predictions_val.csv",
                "Runs training and validation for each CREMA-D fold.",
            ],
            [
                "run_primary_speech_emotion_crema_d_cv_analysis.sh",
                "analysis_val/predictions_test.csv or predictions_val.csv, confusion_matrix.csv, top_confusions.csv, predicted_vs_actual_first20.csv, predicted_vs_actual_first20.md",
                "Shows how the CREMA-D model is making mistakes on one fold.",
            ],
            [
                "aggregate_crema_d_cv_metrics.py",
                "summary.json and summary.md",
                "Summarizes the CREMA-D fold metrics into one readable report.",
            ],
            [
                "run_primary_conversational_meld_cv.sh",
                "results/primary_conversational/meld_cv/cmt_min/fold_*/metrics.json, predictions_test.csv, summary.json, summary.md",
                "Runs MELD 5-fold cross-validation and aggregates the fold results.",
            ],
            [
                "run_primary_conversational_meld_analysis.sh",
                "analysis_test/confusion_matrix.csv, top_confusions.csv, predicted_vs_actual_first20.csv, predicted_vs_actual_first20.md",
                "Explains one representative MELD fold in detail.",
            ],
        ],
    )
    add_para(
        doc,
        "The most important output directories are separate by design. CREMA-D primary speech-emotion results go into results/primary_speech_emotion, while MELD conversational results go into results/primary_conversational. That separation makes it easier to present the results cleanly in a report or guidance call.",
    )

    doc.add_heading("1.3 What metrics these scripts evaluate", level=2)
    add_para(
        doc,
        "The metrics are chosen to fit the type of benchmark. CREMA-D, as the speech-emotion benchmark, is summarized with W-Acc and UW-Acc style language in the cross-validation summary. MELD, as the conversational benchmark, is summarized using accuracy, weighted F1, macro F1, and fold-level error analysis.",
    )
    add_bullets(
        doc,
        [
            "W-Acc in this codebase is the standard sample accuracy quantity reported using the base-paper naming style.",
            "UW-Acc is the mean per-class accuracy over the observed classes in the split, so it is closer to balanced accuracy than to plain accuracy.",
            "Accuracy measures how many predictions are correct overall, but it can hide class imbalance.",
            "Weighted F1 gives more weight to common classes, so it reflects the result on the dataset distribution.",
            "Macro F1 treats every class equally, which makes it the clearest metric for minority-class behavior.",
            "Confusion matrices show where the model is confusing one class with another, which is often more informative than a single score.",
        ],
    )
    add_para(
        doc,
        "For CREMA-D, the most important thing to read is whether W-Acc and UW-Acc move above chance and whether the fold summary shows stable behavior across the five folds. For MELD, the most important thing to read is whether the model stays close to the paper on the main accuracy range while still struggling less on the minority classes than before. In both cases, the confusion matrix and the first-20 predicted-vs-actual table tell you how the numbers arise in practice.",
    )

    doc.add_heading("2. The Core Comparison at a Glance", level=1)
    add_table(
        doc,
        ["Benchmark", "Your result", "Base paper / reference", "Interpretation"],
        [
            [
                "MELD",
                f"CV mean Acc {meld_summary['accuracy']['mean']*100:.2f}%, weighted F1 {meld_summary['weighted_f1']['mean']*100:.2f}%",
                "Paper CMT + MIN test: Acc 64.18%, F1 62.52%, Prec 63.82%, Rec 64.18%",
                "Close to the paper and clearly in the right direction.",
            ],
            [
                "CREMA-D",
                f"Acc {crema['accuracy']*100:.2f}%, macro F1 {crema['macro_f1']*100:.2f}%",
                "No direct CREMA-D result in the base paper",
                "Very weak; the model is close to chance and needs more work.",
            ],
        ],
    )

    doc.add_heading("3. MELD Compared with the Base Paper", level=1)
    add_para(
        doc,
        f"The base paper reports MELD CMT + MIN test results of Acc {BASE_PAPER_MELD['Acc']:.2f}%, F1 {BASE_PAPER_MELD['F1']:.2f}%, Prec {BASE_PAPER_MELD['Prec']:.2f}%, and Rec {BASE_PAPER_MELD['Rec']:.2f}%. The current paper-aligned MELD CV summary has mean accuracy {meld_summary['accuracy']['mean']*100:.2f}% and mean weighted F1 {meld_summary['weighted_f1']['mean']*100:.2f}%. That places the system only about {BASE_PAPER_MELD['Acc'] - meld_summary['accuracy']['mean']*100:.2f} percentage points below the paper on the accuracy number, which is a good sign.",
    )
    add_para(
        doc,
        f"The best MELD fold reached accuracy {meld_fold2['accuracy']:.4f} and weighted F1 {meld_fold2['weighted_f1']:.4f}. That best-fold result is very close to the paper’s MELD CMT + MIN test accuracy of {BASE_PAPER_MELD['Acc']:.2f}%. This does not mean the system is an exact reproduction, but it does mean the implementation is moving in the same general direction and the main architecture choice is sensible.",
    )
    add_bullets(
        doc,
        [
            "The MELD result is good, not perfect.",
            "It is close enough to say the implementation is in the right range.",
            "The gap is small enough that further training and implementation refinements could plausibly close it.",
            "The fold-level variation shows the model is reasonably stable, but still not fully robust on minority emotions.",
        ],
    )

    doc.add_heading("4. Why the MELD Result Is Not an Exact Match Yet", level=1)
    add_bullets(
        doc,
        [
            "The evaluation protocol is not exactly the same as the paper’s reporting setup, so the numeric comparison should be treated as close rather than identical.",
            "The current workflow is a 5-fold CV system, which gives a more robust estimate but is not the same as a single fixed test split result.",
            "The fold analysis still shows a clear bias toward Neutral and related classes, which means the model has not fully solved the class imbalance problem.",
            "The macro-F1 is notably lower than the weighted metrics, showing that the minority classes still need improvement.",
        ],
    )

    doc.add_heading("5. What the MELD Error Analysis Indicates", level=1)
    add_para(
        doc,
        "The Fold 2 analysis is the best window into the model’s actual behavior. The confusion matrix shows that the model often collapses minority emotions into Neutral or nearby emotions such as Joy, Anger, and Surprise. That is the kind of failure expected when a conversational dataset is imbalanced and the model is not yet fully separating emotion classes.",
    )
    add_para(
        doc,
        "The important point is that the MELD errors are structured, not random. The model is not completely broken. It is learning something useful about affect, but it still prefers the dominant class too often. That is why the weighted metrics look reasonably strong while macro-F1 remains lower.",
    )
    add_bullets(
        doc,
        [
            "This is the right direction because the result is far above a random baseline and close to the paper on the main accuracy number.",
            "This is still not ideal because minority-class recognition is not strong enough.",
            "The next improvement target is better class balance handling, not a completely different benchmark.",
        ],
    )

    doc.add_heading("6. CREMA-D Compared with the Base Paper Context", level=1)
    add_para(
        doc,
        "The base MemoCMT paper does not provide a CREMA-D result, so there is no direct paper number to compare against. That means CREMA-D must be judged on its own internal quality and on whether it behaves like a useful secondary benchmark for this project.",
    )
    add_para(
        doc,
        f"The current CREMA-D paper-aligned run is weak: accuracy {crema['accuracy']*100:.2f}%, weighted F1 {crema['weighted_f1']*100:.2f}%, and macro F1 {crema['macro_f1']*100:.2f}% across {crema['num_samples']} test samples. Since random chance on six classes is about 16.67%, the result is only barely above chance and should be treated as very poor model performance rather than as a useful baseline.",
    )
    add_para(
        doc,
        "The CREMA-D confusion analysis makes the failure mode obvious: the model collapses heavily toward fear. In the first twenty examples, every prediction is fear. That means the current CREMA-D setup is not learning class separation in any meaningful way.",
    )

    doc.add_heading("7. What the CREMA-D Result Indicates", level=1)
    add_bullets(
        doc,
        [
            "The pipeline itself works end-to-end, because training, checkpointing, prediction export, and analysis all completed.",
            "The learned representation is not yet adequate for CREMA-D, because the model collapses to one class too strongly.",
            "This suggests a training problem or dataset fit problem, not just a logging or script problem.",
            "The result is scientifically useful because it shows that success on MELD does not automatically transfer to CREMA-D.",
        ],
    )

    doc.add_heading("8. How the Two Results Compare Overall", level=1)
    add_table(
        doc,
        ["Question", "MELD answer", "CREMA-D answer"],
        [
            ["Is the system in the right direction?", "Yes", "No, not yet"],
            ["Is it close to the paper?", "Fairly close on the main MELD accuracy number", "No direct paper comparison exists"],
            ["Does the pipeline run correctly?", "Yes", "Yes"],
            ["Is class balance handled well?", "Partially, but not fully", "No"],
            ["Is the result ready for a strong claim?", "Cautiously yes for MELD, with caveats", "No"],
        ],
    )

    add_para(
        doc,
        "The combined reading is: the project is technically on the right track for MELD and technically functioning for CREMA-D, but CREMA-D has not yet become a meaningful predictive model. That is a better and more honest summary than calling both results equally successful.",
    )

    doc.add_heading("9. Recommended Order of Execution From Here", level=1)
    add_para(
        doc,
        "If you want to start from the new benchmark split, the cleanest order is to begin with the CREMA-D speaker-independent cross-validation workflow and then return to MELD as the conversational benchmark. That keeps the speech-emotion side paper-style and the conversational side paper-aligned, without changing the underlying implementation.",
    )
    add_numbered(
        doc,
        [
            "Run the CREMA-D 5-fold CV workflow with scripts/run_primary_speech_emotion_crema_d_cv.sh.",
            "Inspect the CREMA-D CV summary in results/primary_speech_emotion/crema_d_cv/cmt_min/summary.json and summary.md.",
            "Run the CREMA-D CV analysis wrapper on a representative fold, usually fold 2 unless another fold is more representative.",
            "Read the CREMA-D confusion matrix and first-20 predicted-vs-actual table to see whether the model is still collapsing into one class or starting to separate classes better.",
            "Then run the MELD conversational CV workflow with scripts/run_primary_conversational_meld_cv.sh.",
            "Inspect the MELD summary and fold-level outputs, especially the fold that gives the clearest error picture.",
            "Use the guidance-call comparison report to decide whether the next improvement should target imbalance handling, training schedule, or a CREMA-D-specific setup issue.",
        ],
    )

    doc.add_heading("10. What Should Be Improved Next", level=1)
    add_numbered(
        doc,
        [
            "Improve class imbalance handling on MELD further, especially minority emotions such as disgust and fear.",
            "Investigate whether focal loss or stronger class weighting improves the MELD macro-F1 without hurting the overall accuracy too much.",
            "Check whether the bidirectional CMT fusion block can be refined further, for example by adjusting sequence projection, pooling, or fine-tuning behavior.",
            "For CREMA-D, revisit the training setup because the current run is too close to a single-class collapse.",
            "Consider whether CREMA-D needs a different learning-rate schedule, more epochs, or a more conservative fine-tuning regime.",
            "Consider pretraining on CREMA-D and fine-tuning on MELD only if the mentor agrees that it helps the final objective.",
            "Keep the documentation updated so future runs remain interpretable rather than becoming a collection of unrelated experiments.",
        ],
    )

    doc.add_heading("11. Questions to Ask the Mentor", level=1)
    add_para(
        doc,
        "These are the questions that seem most useful to ask in the guidance call because they directly affect the next step and help avoid unnecessary work.",
    )
    add_bullets(
        doc,
        [
            "Should MELD remain the only primary benchmark for the paper-aligned story, with CREMA-D kept strictly as a secondary benchmark?",
            "Is the current MELD result close enough to the paper to treat as a successful reproduction direction, or should I also run a stricter paper-exact MELD configuration before writing the final summary?",
            "Should I prioritize improving MELD minority-class performance with weighted loss / focal loss before trying to rescue CREMA-D?",
            "For CREMA-D, is the current collapse to fear a sign that I should revise the training schedule, the fine-tuning strategy, or the model architecture first?",
            "Would the mentor prefer one unified report emphasizing MELD only, or a comparative report that explicitly shows MELD as strong and CREMA-D as weak?",
            "Is it worthwhile to attempt CREMA-D pretraining followed by MELD fine-tuning, or would that add complexity without enough benefit?",
            "Should I report accuracy and weighted F1 as the main numbers, or should macro-F1 be elevated more strongly in the final writeup because of the imbalance issue?",
        ],
    )

    doc.add_heading("12. Suggested Way to Present This on the Call", level=1)
    add_para(
        doc,
        "A clean way to speak about the work is to separate it into three claims: what has been implemented, what the current numbers say, and what still needs improvement. That keeps the discussion factual and avoids overstating the result.",
    )
    add_bullets(
        doc,
        [
            "Implemented: paper-aligned MELD with pretrained text/audio, bidirectional CMT, MIN pooling, 5-fold CV, and full error analysis.",
            "Observed: MELD is close to the MemoCMT paper on the main accuracy range; CREMA-D is currently very weak and near chance.",
            "Next: improve imbalance handling, decide whether to keep or adjust CREMA-D, and ask the mentor which direction is most valuable for the final project.",
        ],
    )

    doc.add_heading("13. Bottom-Line Judgment", level=1)
    add_para(
        doc,
        "The MELD result is good and in the right direction. It is reasonably close to the base paper on the main metric and supports the paper-aligned architecture choice. The CREMA-D result is very bad in its current form and should be treated as a diagnostic failure rather than a success. The project is therefore in a mixed state: strong enough to justify the current MELD direction, but still needing meaningful improvement before CREMA-D can be considered useful.",
    )
    add_para(
        doc,
        "That is the most honest summary to take into a guidance call. It shows progress, it shows the gap, and it creates clear next-step questions instead of pretending the current system is already finished.",
    )

    return doc


def update_index_entry() -> None:
    path = ROOT / "scripts" / "build_student_document_index.py"
    text = path.read_text(encoding="utf-8")
    if "LegalMemoCMT_Guidance_Call_Comparison_Report.docx" in text:
        return
    marker = '    (\n        "LegalMemoCMT_MELD_CV_Paper_Aligned_Report_From_Submission_First.docx",\n'
    insert = '''    (\n        "LegalMemoCMT_Guidance_Call_Comparison_Report.docx",\n        "Guidance-call comparison document that summarizes the current MELD and CREMA-D results against the MemoCMT base paper and lists the next-step mentor questions.",\n        "Use this before a mentor call when you need a concise but deep comparison of what has been done, how it matches the paper, and what should improve next.",\n    ),\n'''
    if marker not in text:
        raise SystemExit("Could not find insertion point in build_student_document_index.py")
    text = text.replace(marker, insert + marker, 1)
    path.write_text(text, encoding="utf-8")


def main() -> None:
    doc = build_doc()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    update_index_entry()
    print(OUTPUT)


if __name__ == "__main__":
    main()
