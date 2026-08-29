"""Add a canonical 30-slide reading sequence without deleting existing manual notes."""
from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "implementation_docments/LegalMemoCMT_Phase2_Second_Review_Call_Speaking_Doc.docx"
MARKER = "CANONICAL_30_SLIDE_SPEAKING_ORDER_V1"

SLIDES = [
("LegalMemoCMT", "I will introduce LegalMemoCMT as a multimodal courtroom testimony project. The important scope is not ordinary video classification; each example must retain text, audio, video, speaker role, and provenance. I will also state that the current Clancy branch is the primary benchmark while the tribunal branch remains a secondary bootstrap and comparison corpus."),
("Objective and Novelty of Phase 2", "Phase 2 converts long courtroom recordings into auditable utterance or turn-level examples. The novelty is the combination of courtroom-specific witness selection, synchronized modalities, speaker-role evidence, and conservative affect annotation rather than treating every emotional word as the speaker's emotion."),
("Title and Abstract for Phase 2", "I will explain that the title describes the research direction, while the abstract states the actual implementation boundary. At this stage the pipeline demonstrates corpus construction and machine-assisted annotation; it does not claim human-gold emotion or credibility labels."),
("Algorithms and Techniques: What To Read", "This slide connects each engineering component to the literature I need to understand. I will explain that forced alignment, diarization, multimodal fusion, weak supervision, and leakage-aware evaluation solve different problems and should not be presented as one interchangeable algorithm."),
("Second Guidance Call: Corpus Expansion", "The corpus is organized incrementally: source discovery, download verification, transcript and media alignment, role mapping, and quality-controlled selection. Expansion is only useful when it increases grounded witness evidence, not merely the number of URLs or raw hours."),
("LegalMemoCMT Courtroom Corpus", "I will distinguish the corpus branches by their research role. Clancy supplies a primary, witness-rich courtroom benchmark; ICTY/ICTR/IRMCT supply a tribunal bootstrap with real testimony; Tupac/Keffe D and Indian-SIM are planned or supplementary branches with different risks and purposes."),
("Source Sites and Download Flow", "The source URL is retained before downloading so every local file can be traced back to its public origin. The downloader verifies media, records status, and avoids treating a failed or generic page as evidence. For YouTube sources, download success and subtitle availability are separate checks."),
("IRMCT / ICTY EDA", "This EDA is a secondary tribunal baseline, not evidence that the Indian corpus already exists. I will report cases, hearings, durations, and modality availability separately because a transcript-only hearing cannot contribute video hours."),
("Lindsay Clancy EDA", "The original EDA describes the Clancy source inventory and the first controlled analysis batch. I will emphasize that the 200 rows are a pilot selected from a larger witness-visible pool, not the final corpus size."),
("Clancy EDA Layer 1: Full Source-Shared Utterance Corpus", "The 77,442 rows are subtitle-derived source rows across 11 videos. They show extraction coverage, but shared video and subtitle paths mean the rows are not independent media examples. The summed 29.904 hours therefore must not be reported as unique usable training hours."),
("Clancy EDA Layer 2: Speaker-Cluster Processing", "The diarization stage supplies time intervals and anonymous local speaker clusters. Pyannote can indicate which voice is active, but it cannot know that a cluster is a witness or lawyer; that legal-role decision requires mapping and review. Overlapping segment-to-turn mappings are evidence that the boundaries need cautious interpretation."),
("Clancy EDA Layer 3: Witness-Visible Speaking Corpus", "This is the first layer that is appropriate for witness-only selection. The 2,229 usable rows passed the current witness-speaking and exclusion rules, while 350 rows remain review/outlier material. I will not add the review duration to usable corpus hours until those rows are inspected."),
("Clancy Pilot Scope and Final Corpus", "The pilot is intentionally limited to 200 rows so feature paths, model outputs, scope rules, and acceptance gates can be inspected quickly. Once the rules are stable, the same pipeline will process the complete eligible witness pool and recompute final counts rather than extrapolating from 200."),
("Clancy Corpus Pipeline: From Sources to the 200-Row Pilot", "This diagram shows the dependency order: source evidence comes first, then turn construction, diarization, role mapping, witness filtering, and quality checks. The pilot is downstream of those controls, so it is not a shortcut around them. This ordering is what makes the pilot reproducible and auditable."),
("What Is Done in the First 200-Row Pilot?", "Each pilot row is checked for provenance and then receives independent evidence: ViT video features, Phase 1 basic-emotion prediction, audio-SER measurements, and transcript-scope cues. The acceptance gate separates machine-assisted SILVER candidates from unresolved rows requiring review. DeBERTa scope inference is a planned comparison stage, not silently applied gold labeling."),
("Clancy Duration Window and Emotion EDA", "The duration windows make the data closer to MELD-style samples and reduce long breaks or mixed segments. I will explain that duration is a usability control, not an emotion label. Outliers are retained in review manifests so the decision is traceable rather than silently deleting evidence."),
("Tupac / Keffe D: Current Status and EDA Boundary", "This branch is treated as a possible diversity extension, not as already validated witness testimony. Before inclusion, each source needs role, testimony relevance, media, transcript, and duration checks. I will not combine its planned rows with Clancy totals until those checks pass."),
("Indian-SIM: Planned Adaptation Branch", "Indian-SIM is a controlled adaptation source for Indian courtroom procedure and speech. It is useful for adaptation and evaluation, but simulated testimony must be labeled as simulated. This prevents the project from overstating that the bootstrap tribunal data is already an Indian real-court corpus."),
("Cross-Corpus Duration and Outlier Comparison", "The comparison uses the same duration and quality definitions across branches. I will separate source duration, candidate duration, and validated multimodal duration. This prevents a larger raw corpus from appearing better merely because it contains more unfiltered material."),
("What I Want to Confirm in Guidance", "I will ask the mentor to confirm whether witness-visible speaking should remain the primary inclusion criterion and whether the proposed label taxonomy is appropriate. I will also ask whether the pilot gate is conservative enough before scaling to the full Clancy pool."),
("Phase 2 Review: Gated Annotation EDA", "This EDA reports how many rows pass the machine-assisted gate and how many remain unresolved. SILVER means eligible for controlled downstream use with provenance, not human gold. The unresolved count is useful because it identifies where courtroom domain shift affects the MELD-trained model."),
("Phase 2 Hugging Face Models", "Each model has a bounded purpose. ViT creates visual features, Pyannote creates speech segments and clusters, SpeechBrain and dimensional audio models provide acoustic evidence, and DeBERTa is intended for text scope hypotheses. None of these models independently determines courtroom truth, credibility, or deception."),
("Example: SILVER Auto-Adjudication", "A SILVER example passes only when confidence and affect evidence satisfy the configured gate and no critical conflict is present. I will preserve the original Phase 1 output and record the machine-assisted decision separately. This makes the row useful for calibration without disguising it as a human annotation."),
("Example: Weak / Unresolved Candidate", "A weak example remains unresolved when confidence is low, modalities disagree, or distress is unsupported. The correct engineering response is to route it to review, not force a categorical label. These rows are valuable active-learning candidates because they expose failure modes."),
("How the Gated CSV Was Produced", "The CSV is produced by joining preserved Phase 1 predictions with audio, transcript-scope, and courtroom-affect evidence. The gate then writes status, tier, and conflict fields while retaining the inputs. I will explain the CSV as an auditable decision record, not as an unexplained final-label table."),
("Slide 17 Explained: Why the SILVER Row Passed", "I will walk through the row-level evidence and show which thresholds were met. The important point is that the row passed a reproducible machine rule, while the final research label still depends on the annotation policy and any human validation sample."),
("Slide 18 Explained: Why the Row Stayed WEAK", "This row demonstrates why semantic negativity, moderate excitement, or a single SER output cannot establish the witness's emotion. The unresolved status records uncertainty and protects the training set from a confident but unsupported label."),
("Acceptance Gate: Two Critical-Failure Checks", "The gate checks for conflicts such as a low-confidence Phase 1 non-neutral prediction being converted without supporting evidence, and an affect class such as distress being proposed without corroboration. A critical failure forces unresolved handling even if another score is high."),
("Critical Conflicts Versus Ordinary Weak Rows", "A critical conflict is a specific safety failure, while an ordinary weak row may simply lack enough confidence. Both can require review, but the distinction helps prioritize manual work and lets me report why a row was withheld."),
("Decision and Next Action", "The immediate next action is to inspect the gated and unresolved rows, validate a stratified sample, and run the DeBERTa scope pilot when its environment is available. Only after those checks should I generate the full Clancy annotation manifest and prepare a training split."),
]

def main():
    doc = Document(DOCX)
    if any(MARKER in p.text for p in doc.paragraphs):
        print("Canonical speaking order already present")
        return
    backup = DOCX.with_name(DOCX.name + ".before_canonical_30_slide_order.docx")
    if not backup.exists(): shutil.copy2(DOCX, backup)
    anchor = next((p for p in doc.paragraphs if p.text.strip() == "Slide-by-slide speaking guidance"), doc.paragraphs[0])
    created = []
    h = doc.add_heading("Canonical Current Deck Order: Slides 1-30", level=1); created.append(h)
    marker = doc.add_paragraph(MARKER); marker.runs[0].font.size = Pt(8); created.append(marker)
    intro = doc.add_paragraph("Use this section as the primary reading sequence for the current 30-slide PPTX. The earlier notes remain below and are intentionally retained so previous manual explanations are not lost."); created.append(intro)
    for num, (title, note) in enumerate(SLIDES, 1):
        h = doc.add_heading(f"Slide {num}: {title}", level=2); created.append(h)
        p = doc.add_paragraph(note); created.append(p)
    body = doc._body._body
    anchor_index = list(body).index(anchor._p)
    for offset, item in enumerate(created):
        body.insert(anchor_index + offset, item._p)
    doc.save(DOCX)
    print("Added canonical 30-slide speaking order to", DOCX)

if __name__ == "__main__": main()
