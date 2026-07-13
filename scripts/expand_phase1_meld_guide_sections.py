from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "implementation_docments" / "LegalMemoCMT_Phase1_ESA_Prep_Guide.docx"


def insert_paragraph_before(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def remove_section(paragraphs, heading_text, next_heading_prefix):
    start = next((i for i, p in enumerate(paragraphs) if p.text.strip() == heading_text), None)
    if start is None:
        return False
    end = next(
        (i for i in range(start + 1, len(paragraphs)) if paragraphs[i].text.strip().startswith(next_heading_prefix)),
        len(paragraphs),
    )
    for i in range(start, end):
        el = paragraphs[i]._element
        el.getparent().remove(el)
    return True


def update_table(doc: Document):
    table = doc.tables[0]
    row7 = table.rows[7]
    row7.cells[1].text = (
        "Explain MELD as a conversational benchmark and use Table 10 to show why the dataset statistics matter: shorter turns, neutral-heavy labels, multi-party context, and dialogue-heavy structure."
    )
    row8 = table.rows[8]
    row8.cells[1].text = (
        "Explain the contextual analysis: inter-speaker influence, emotion shifts, contextual distance, and why these factors matter for dialogue-aware evaluation in Phase 1."
    )


def update_docx():
    doc = Document(str(DOCX_PATH))
    paras = list(doc.paragraphs)

    remove_section(paras, "2.2 Slide 7 MELD Dataset EDA & Table 10", "2.3 Slide 8 MELD Contextual Analysis")
    paras = list(doc.paragraphs)
    remove_section(paras, "2.3 Slide 8 MELD Contextual Analysis", "2.4 Literature Review Slides")
    paras = list(doc.paragraphs)

    anchor = next((p for p in doc.paragraphs if p.text.strip().startswith("2.4 Literature Review Slides")), None)
    if anchor is None:
        raise RuntimeError("Could not find the literature review anchor section.")

    # Slide 7 section.
    h7 = insert_paragraph_before(anchor, style="Heading 2")
    h7.add_run("2.2 Slide 7 MELD Dataset EDA & Table 10")
    slide7_sections = [
        "MELD is not just a list of emotion labels. It is a multi-party, multimodal conversation dataset, which means every utterance appears inside a social and temporal context. That context is part of the label-making process, because the meaning of a turn is often shaped by what was said before and by who said it.",
        "The statistics in Table 10 are useful because they summarize the dataset in a way that immediately explains the modeling challenge. MELD has short utterances, a much stronger neutral class, and dialogue characteristics that are different from older datasets such as IEMOCAP. The short average utterance length means local language cues are compact, while the neutral skew means that a model can appear strong on accuracy without actually learning minority emotions well.",
        "Table 10 also helps explain why the project should care about conversation-level behavior rather than only utterance-level prediction. When a dataset has many turns per dialogue and multi-party interactions, the same emotion label may depend on the surrounding context and on speaker identity. That is why the project later uses fold-safe grouping, weighted F1, and confusion-matrix analysis instead of relying on accuracy alone.",
        "For the ESA explanation, a good way to phrase the point is that MELD provides a realistic benchmark for conversational emotion analysis: it is small enough to analyze deeply, but complex enough to expose the weakness of simple single-utterance classifiers. This is exactly the kind of benchmark that lets Phase 1 prove the implementation and expose the error pattern before any legal-domain adaptation is claimed.",
        "The practical consequence is that the project must preserve the dialogue structure in the data pipeline. If the utterances are treated as independent rows without respecting conversational grouping, then the model can leak context across folds or misrepresent the real difficulty of the benchmark. So the dataset EDA is not a side note; it is the reason the evaluation strategy is designed the way it is.",
    ]
    for t in slide7_sections:
        p = insert_paragraph_before(anchor)
        r = p.add_run(t)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11.3)
        p.alignment = 3

    # Slide 8 section.
    h8 = insert_paragraph_before(anchor, style="Heading 2")
    h8.add_run("2.3 Slide 8 MELD Contextual Analysis")
    slide8_sections = [
        "Inter-speaker influence means that the emotion of one utterance is often affected by a different speaker’s prior response. In a multi-party conversation, this is important because a speaker may react to a previous statement, interrupt another speaker, or shift the tone of the conversation. The MELD paper shows that these cross-speaker dependencies are common, which is why conversational models should not ignore speaker interaction.",
        "Emotion shifts mean that the same speaker can move from one emotion to another as the conversation develops. This is technically important because the model is not just learning a single static emotion profile for a person; it is learning how emotion changes across the dialogue timeline. If the model misses a shift, it may still predict the correct class for one utterance but fail to explain why that label appeared at that moment.",
        "Contextual distance refers to how far the useful supporting utterance is from the target utterance in the conversation. Nearby turns usually matter most, but distant history can still influence the prediction. This is why the conversation should be interpreted as a sequence rather than as isolated points. It also explains why a bidirectional or attention-based model can be useful: the right context may appear before or after the target utterance.",
        "For Phase 1, these three ideas change how the results should be read. A model that only looks good on aggregate accuracy may still be weak at handling shifts, speaker interactions, and long-range context. That is why the project emphasizes weighted metrics, confusion matrices, and per-sample error analysis, so the examiner can see whether the model is learning dialogue structure or only taking advantage of the dominant label distribution.",
        "The broader takeaway for the ESA is that MELD is not merely a benchmark with many labels. It is a benchmark that requires context-aware reasoning. That makes it a suitable place to validate the paper-aligned baseline and later justify why legal-testimony work will need even richer metadata and explanation support.",
    ]
    for t in slide8_sections:
        p = insert_paragraph_before(anchor)
        r = p.add_run(t)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11.3)
        p.alignment = 3

    update_table(doc)

    doc.save(str(DOCX_PATH))


def main():
    update_docx()
    print(f"Updated {DOCX_PATH}")


if __name__ == "__main__":
    main()
