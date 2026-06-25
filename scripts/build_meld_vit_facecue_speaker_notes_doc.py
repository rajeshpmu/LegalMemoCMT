from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "implementation_docments" / "LegalMemoCMT_MELD_Facial_Cues_ViT_Speaker_Notes.docx"
FIG_DIR = ROOT / "implementation_docments" / "figures" / "meld_vit_facecue_notes"
FIG_DIR.mkdir(parents=True, exist_ok=True)
SAMPLED_FRAMES_PNG = FIG_DIR / "sampled_rgb_frames.png"
SAMPLED_FRAMES_SVG = FIG_DIR / "sampled_rgb_frames.svg"
OUTPUTS_PNG = FIG_DIR / "pipeline_outputs.png"
OUTPUTS_SVG = FIG_DIR / "pipeline_outputs.svg"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)


def add_para(doc: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.italic = italic
    r.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc: Document, code: str) -> None:
    for line in code.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(9.2)


def render_mermaid(code: str, svg_path: Path, png_path: Path) -> None:
    mmd_path = png_path.with_suffix(".mmd")
    mmd_path.write_text(code, encoding="utf-8")
    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(mmd_path), "-o", str(svg_path), "-b", "white"],
        check=True,
        cwd=str(ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(mmd_path), "-o", str(png_path), "-b", "white"],
        check=True,
        cwd=str(ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )


def build_doc() -> Document:
    render_mermaid(
        """sequenceDiagram
  participant A as Raw MELD utterance clip
  participant B as cv2.VideoCapture
  participant C as np.linspace frame indices
  participant D as cv2.resize to 224x224
  participant E as BGR -> RGB conversion
  participant F as ViT image processor
  participant G as ViT CLS embeddings
  participant H as Saved .npy file
  A->>B: open the clip
  B->>C: count frames and choose evenly spaced indices
  C->>D: sample the selected frames
  D->>E: standardize the image size and color order
  E->>F: convert images to tensors
  F->>G: extract facial embeddings
  G->>H: save the embedding array for reuse
""",
        SAMPLED_FRAMES_SVG,
        SAMPLED_FRAMES_PNG,
    )
    render_mermaid(
        """sequenceDiagram
  participant A as ViT .npy embeddings
  participant B as fold CSVs
  participant C as training script
  participant D as best_model.pt
  participant E as evaluation step
  participant F as metrics.json
  participant G as predictions_test.csv
  participant H as confusion matrix / error analysis
  A->>C: load visual features
  B->>C: map each sample to train or val
  C->>D: save the best validation checkpoint
  D->>E: evaluate on held-out data
  E->>F: write summary metrics
  E->>G: write per-sample predictions
  G->>H: build the error analysis tables
""",
        OUTPUTS_SVG,
        OUTPUTS_PNG,
    )

    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    run = title.add_run("LegalMemoCMT MELD Facial Cues ViT Speaker Notes")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(
        "Detailed slide-by-slide speaker notes for the 14-slide presentation, including the guidance-call continuation slides and the facial-cue implementation section."
    )
    run.italic = True
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"

    doc.add_paragraph(
        "Purpose: this document is designed to be read while presenting the full facial-cue slide deck. It explains the slides in student-friendly language, shows what the code is doing technically, and points to the exact script locations that support each stage of the pipeline. The first three slides continue the second guidance-call story, and the remaining slides explain the ViT facial-cue implementation in detail."
    )
    doc.add_paragraph(
        "Primary code references used throughout the notes: scripts/build_meld_vit_facecue_manifest.py, src/data/preprocessing.py, and the facial-cue training entrypoint scripts/run_meld_vit_facecue_fold2.sh / scripts/run_meld_vit_facecue_fold4.sh."
    )

    doc.add_heading("1. Slide 1 - Project Review and Guidance Call Continuation", level=1)
    add_para(
        doc,
        "Open by explaining that these three slides continue the second guidance-call update. The purpose is to remind the audience where Phase 1 stands, what the strongest Fold 2 result is, and why the next work should move from loss tuning into the video-enhanced stage.",
    )
    add_bullets(
        doc,
        [
            "The project is still a multimodal emotion-recognition framework for Indian courtroom testimony.",
            "Phase 1 remains the working codebase for the pretrained/paper-aligned baseline.",
            "Phase 2 is the legal-domain adaptation stage, where the video branch becomes more important.",
            "The opening slide should make the research scope clear: emotional cues only, not guilt, deception, or intent.",
        ],
    )
    add_para(
        doc,
        "Student takeaway: this is the project framing slide. It tells the audience what the thesis is about and keeps the scope legally and scientifically narrow.",
    )

    doc.add_heading("2. Slide 2 - Fold 2 Three-Way Training Comparison", level=1)
    add_para(
        doc,
        "This slide compares the same Fold 2 split under three training strategies. The important lesson is not just which number is highest. The lesson is that the training objective and checkpoint strategy change the result, even when the architecture is the same.",
    )
    add_bullets(
        doc,
        [
            "Weighted cross-entropy is the strongest overall Fold 2 baseline.",
            "Focal loss from scratch is the weakest result because the objective alone is too disruptive here.",
            "Warm-start focal is a partial recovery because it begins from a better learned boundary.",
            "The comparison shows that optimization history matters, not just the final loss choice.",
        ],
    )
    add_para(
        doc,
        "When presenting this slide, explain that the comparison is same fold, same architecture, but different training path. That makes it a controlled experiment about the learning objective rather than about model redesign.",
    )

    doc.add_heading("3. Slide 3 - Summary So Far", level=1)
    add_para(
        doc,
        "This slide explains the current project direction after the second guidance call. The evidence says that the baseline is strong, the focal-loss experiments were diagnostic, and the next stage should move toward the video-enhanced LegalMemoCMT path.",
    )
    add_bullets(
        doc,
        [
            "Keep the weighted-CE MELD Phase 1 result as the backbone baseline.",
            "Treat the focal-loss runs as diagnostics that did not beat the baseline.",
            "Proceed to HuBERT + BERT with Cross-Modal Transformer fusion plus ViT facial cues for the next stage.",
            "Keep the text-audio path as a non-visual comparison so the video stage can be measured cleanly.",
            "Frame the novelty around legal-domain adaptation and multimodal video integration, not a small MELD gain.",
        ],
    )
    add_para(
        doc,
        "Student takeaway: the slide tells you where the project is now and what the next research branch should be.",
    )

    doc.add_heading("3.1 Slide 3.1 - Gated Fusion Follow-Up", level=1)
    add_para(
        doc,
        "This section gives the spoken explanation for the gated-fusion step. The point is not to claim that video should dominate the prediction. The point is to let the model learn when the face-crop branch is useful and when it should be down-weighted in favor of text and audio.",
    )
    add_bullets(
        doc,
        [
            "Explain that the gate is a learned weighting mechanism over modality representations.",
            "Say that the gate lets the model trust video only when the face-crop evidence is helpful.",
            "Emphasize that this is safer than forcing a video loss too early because the video-only run was weak.",
            "State that the goal is selective contribution, not unconditional dominance of the visual branch.",
        ],
    )
    add_para(
        doc,
        "Student explanation: the gate is like a switchboard that decides how much text, audio, and video should matter for a given utterance. If the face crop is informative, the gate raises its influence. If it is noisy or weak, the gate lowers its influence and lets the text/audio backbone stay in control.",
    )
    add_code(
        doc,
        """text_repr + audio_repr + video_repr
        -> gated fusion
        -> weighted modality mix
        -> classifier
        -> emotion label""",
    )
    add_para(
        doc,
        "Recommended execution order for the gated follow-up: first run scripts/run_meld_vit_facecrop_gated_fold2.sh, then scripts/analyze_meld_vit_facecrop_gated_fold2.sh, then scripts/run_meld_vit_facecrop_gated_fold4.sh, and finally scripts/analyze_meld_vit_facecrop_gated_fold4.sh. Use scripts/run_meld_vit_facecrop_gated_suite.sh only if you want the same sequence wrapped into a single command after the control behavior is already understood.",
    )
    add_bullets(
        doc,
        [
            "Fold 2 comes first because it is the most useful debugging fold.",
            "Fold 4 comes second because it checks whether the gate generalizes to another balanced-class anchor.",
            "The analysis scripts must follow each training run so the confusion matrix is available before interpretation.",
        ],
    )
    add_para(
        doc,
        "Confusion-matrix interpretation for the gated Fold 2 run: the model is still neutral-biased, but it is less collapsed than the video-only control. Neutral remains the dominant prediction, which means the imbalance problem is still present. At the same time, surprise, fear, and joy are being recognized better than in the video-only run, which suggests that the gate is actually letting some useful facial signal through.",
    )
    add_bullets(
        doc,
        [
            "Neutral is still the most common prediction, so the model has not fully escaped the dataset bias.",
            "Surprise, fear, and joy have meaningful diagonal counts, showing the gate preserved some visual signal.",
            "Sadness, anger, and disgust still mix with neutral and with nearby emotions, which means the boundary is still soft.",
            "The pattern is better than video-only, but not yet strong enough to claim that gated fusion solved the class imbalance problem.",
        ],
    )
    add_para(
        doc,
        "Student-level way to say it: the gate improves the situation by making the video branch useful, but the model still prefers neutral when it is uncertain. So the confusion matrix says the gate is helpful, not perfect. That is the right message for a second-review discussion: the branch is promising, but it still needs careful tuning or a stronger visual-fusion strategy before it can be treated as final.",
    )

    doc.add_heading("3.2 Slide 3.2 - Gated Fusion + Video Auxiliary Loss", level=1)
    add_para(
        doc,
        "This section explains the next experiment after gated fusion. The model keeps the face-crop ViT branch and the gated fusion block, but now it also adds a small auxiliary loss on the video branch. The purpose is to give the video encoder a direct training signal so it does not remain too passive behind the text/audio backbone.",
    )
    add_bullets(
        doc,
        [
            "Say that the main classifier still uses the gated multimodal fusion output.",
            "Explain that the auxiliary video head makes the video branch solve emotion classification on its own.",
            "Tell the audience that the auxiliary loss is small, lambda = 0.1, so it supports the video branch without dominating the overall objective.",
            "Mention that the learning rate is reduced to 2e-5 because this is a refinement step, not a from-scratch training run.",
            "State that the run uses at most 8 epochs with early stopping patience = 2, monitored on validation weighted F1.",
            "State that the best checkpoint should be chosen by validation weighted F1 because that is the safer class-aware selection rule for MELD.",
            "Mention explicitly that this is still a warm-start experiment from the paper-aligned MELD checkpoint.",
        ],
    )
    add_para(
        doc,
        "Student explanation: this experiment is like saying to the video branch, 'You are not only helping the fused model, you must also learn to predict emotion directly.' The gate decides when video should matter in the multimodal classifier, and the auxiliary loss makes the video branch stronger on its own. Together, they test whether the face-crop cue can become more useful without forcing it to dominate the whole system.",
    )
    add_code(
        doc,
        """main loss = weighted CE on gated fused output
video loss = weighted CE on video-only head
total loss = main loss + 0.1 * video loss""",
    )
    add_para(
        doc,
        "Recommended manual run order for this next experiment: first make sure the face-crop control folds are already built, then run scripts/run_meld_vit_facecrop_gated_video_aux_fold2.sh, then scripts/analyze_meld_vit_facecrop_gated_video_aux_fold2.sh, then scripts/run_meld_vit_facecrop_gated_video_aux_fold4.sh, and finally scripts/analyze_meld_vit_facecrop_gated_video_aux_fold4.sh. Use scripts/run_meld_vit_facecrop_gated_video_aux_suite.sh only if you want the same flow in one convenience command. In the talk, state that the run is capped at 8 epochs, stops early after 2 non-improving epochs on validation weighted F1, and keeps the best checkpoint by validation weighted F1 rather than by accuracy.",
    )
    add_bullets(
        doc,
        [
            "Fold 2 remains the best debugging fold for early inspection.",
            "Fold 4 checks whether the same improvement generalizes in a more balanced setting.",
            "The analysis steps matter because the confusion matrix is what tells you whether the auxiliary loss actually improves the visual branch.",
        ],
    )

    doc.add_heading("4. Slide 4 - LegalMemoCMT MELD Facial Cues ViT Implementation Plan", level=1)
    add_para(
        doc,
        "Start the presentation by telling the audience that this is the next stage after the weighted-CE MELD Phase 1 baseline. The message is not that the baseline failed; the message is that the baseline is strong enough to support a careful visual extension.",
    )
    add_bullets(
        doc,
        [
            "The goal is to add facial cues as a supported extension, not as a replacement for the current text+audio backbone.",
            "The visual branch is intentionally framed as a controlled experiment so that any gain can be attributed to the added facial information.",
            "The audience should understand the pipeline order before looking at any result numbers: raw clip -> frames -> ViT -> .npy -> warm-start training -> validation trend.",
        ],
    )
    add_para(
        doc,
        "Teaching note: emphasize that this is a backbone-plus-extension story. The current model already learned a usable MELD boundary with text and audio, so facial cues are added to see whether the visual evidence helps the classifier make better emotion decisions.",
    )

    doc.add_heading("5. Slide 5 - What the ViT Step Means in Practice", level=1)
    add_para(
        doc,
        "This slide explains the central idea of the facial-cue pipeline. The ViT step is where raw video is converted into a learned representation. In other words, the model is no longer reading a raw clip; it is reading visual embeddings derived from sampled frames.",
    )
    add_bullets(
        doc,
        [
            "The script first resolves the raw MELD clip for the utterance.",
            "It then samples a fixed number of full frames from that clip.",
            "Those frames are sent through a pretrained ViT model from Hugging Face.",
            "The result is a compact embedding that can be reused later during training.",
        ],
    )
    add_para(
        doc,
        "Explain that ViT is not installed by the script. The script downloads or reuses the pretrained checkpoint through transformers. The code path that matters is scripts/build_meld_vit_facecue_manifest.py:101-102 for loading the processor and model, and scripts/build_meld_vit_facecue_manifest.py:44-47 for running the frames through ViT.",
    )
    add_para(
        doc,
        "Why 224 x 224? Because that is the standard input size used by the pretrained ViT model family and it is a sensible transfer-learning default. Matching the expected size keeps the visual branch compatible with the pretrained checkpoint and avoids making the first facial-cue experiment unnecessarily heavy.",
    )
    add_bullets(
        doc,
        [
            "224 x 224 matches the common pretrained ViT image size.",
            "It preserves enough detail for facial-expression cues without making the pipeline too slow.",
            "A larger size would cost more memory and time; a smaller size would lose facial detail.",
            "For a first facial-cue baseline, 224 is the best balance between quality and efficiency.",
        ],
    )
    add_code(
        doc,
        """image_processor = AutoImageProcessor.from_pretrained(args.vit_model)
vit_model = AutoModel.from_pretrained(args.vit_model).to(device)
outputs = vit_model(**inputs)
hidden = outputs.last_hidden_state[:, 0, :]""",
    )
    add_para(
        doc,
        "Student takeaway: the CLS token embedding is the visual summary of each frame. The saved .npy file contains those embeddings so the model does not have to recompute ViT every time it trains.",
    )

    doc.add_heading("5.0.1 What ViT-Base Means", level=3)
    add_para(
        doc,
        "ViT-Base is the standard medium-sized Vision Transformer variant used as the visual backbone in this project. It is not trained from the MELD dataset itself. Instead, it is a pretrained image model that gives the pipeline a strong starting point for learning visual patterns in frames.",
    )
    add_bullets(
        doc,
        [
            "ViT means Vision Transformer, so the model reads image patches the way a language model reads tokens.",
            "Base means the model is the standard middle-sized version, large enough to be useful but smaller than the heavy ViT-Large variants.",
            "patch16 means the image is split into 16 x 16 pixel patches before the transformer reads it.",
            "224 means the model expects 224 x 224 input images, which matches the resize step in the pipeline.",
            "in21k means the model was pretrained on ImageNet-21k, a very large generic image dataset, before being used here.",
        ],
    )
    add_para(
        doc,
        "Student intuition: ViT-Base is like a pretrained visual reader that already knows general image structure. Your pipeline does not teach it from scratch. It simply adapts that reader to the utterance-level facial-expression task by feeding it sampled MELD frames and extracting embeddings.",
    )
    add_para(
        doc,
        "Why the embedding size becomes 768: that number comes from the ViT-Base architecture, not from the dataset. The model configuration defines a 768-dimensional hidden state, and the CLS token embedding inherits that size. So the dataset supplies the frames, but the model architecture decides how wide each learned feature vector will be.",
    )
    add_code(
        doc,
        """model_name = "google/vit-base-patch16-224-in21k"
processor = AutoImageProcessor.from_pretrained(model_name)
vit_model = AutoModel.from_pretrained(model_name)
cls_embedding = outputs.last_hidden_state[:, 0, :]  # shape: [num_frames, 768]""",
    )

    doc.add_heading("5.1 Sampled RGB Frames Explained", level=2)
    add_para(
        doc,
        "This is the part that students usually find confusing. The raw clip is not passed into ViT as one video object. Instead, the code opens the clip with OpenCV, counts how many frames the clip has, chooses evenly spaced frame indices with np.linspace, and then keeps only those selected frames.",
    )
    add_bullets(
        doc,
        [
            "The frames are sampled evenly across the entire utterance clip.",
            "OpenCV reads frames in BGR order, so the code converts them to RGB before saving.",
            "Each frame is resized to 224 x 224 so every image matches the ViT input expectation.",
            "The sampled frames are normalized to float values between 0 and 1 before inference.",
            "The final result is a small set of representative still images, not the full raw video.",
        ],
    )
    add_para(
        doc,
        "Why this matters: if you give ViT the full video, the model would have to understand temporal structure directly, which is not what this script is doing. Here, ViT is being used as an image encoder. The video is therefore reduced to a sequence of representative RGB images, and each image becomes one embedding vector.",
    )
    add_para(
        doc,
        "Student intuition: think of the utterance as a strip of film. The script cuts out a small number of still frames from that strip, converts them into clean RGB pictures, and asks ViT to describe each picture in feature-space terms. Those feature vectors are the facial-cue evidence.",
    )
    add_picture = doc.add_picture
    add_picture(str(SAMPLED_FRAMES_PNG), width=Inches(6.9))
    add_code(
        doc,
        """cap = cv2.VideoCapture(str(video_path))
total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT)) or 1
indices = np.linspace(0, total_frames - 1, cfg.num_frames).astype(int)
frame = cv2.resize(frame, (cfg.frame_size, cfg.frame_size))
frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
frames.append(frame.astype(np.float32) / 255.0)""",
    )

    doc.add_heading("5.2 Default ViT and Training Settings to Mention", level=2)
    add_para(
        doc,
        "When presenting the ViT step, it helps to explain why the experiment uses these defaults. The values are chosen to match the pretrained backbone, keep the pipeline reproducible, and protect the current MELD checkpoint during warm-start training.",
    )
    add_bullets(
        doc,
        [
            "ViT checkpoint default: google/vit-base-patch16-224-in21k.",
            "Frame size default: 224 x 224, which matches the common pretrained ViT input size.",
            "Frame count default: 16 sampled frames per utterance.",
            "Audio defaults carried over from the pipeline: 16 kHz sample rate and 10-second cap.",
            "Warm-start defaults: 5 epochs, learning rate 5e-5, freeze-backbone-epochs 1, batch size 4.",
            "Baseline checkpoint default: results/paper_aligned_meld_cv/cmt_min.",
            "Device default: cpu unless the script is overridden with another device.",
        ],
    )
    add_para(
        doc,
        "Student explanation: the defaults are not random. They are there to make the first visual experiment stable, comparable, and not too expensive to run. The pretrained ViT expects 224 x 224 style inputs, and the small warm-start learning rate protects the already strong text/audio backbone.",
    )
    add_code(
        doc,
        """--vit-model google/vit-base-patch16-224-in21k
--frame-size 224
--num-frames 16
--sample-rate 16000
--max-audio-seconds 10.0
--batch-size 8
--epochs 5
--freeze-backbone-epochs 1
--lr 5e-5
--batch-size 4
--checkpoint results/paper_aligned_meld_cv/cmt_min""",
    )

    doc.add_heading("5.3 What the output files mean", level=2)
    add_para(
        doc,
        "This is the part of the pipeline that students usually care about after the embeddings are created. The output files are the visible evidence that the experiment actually ran, saved a checkpoint, evaluated the model, and produced a confusion analysis.",
    )
    add_bullets(
        doc,
        [
            "The .npy files are the saved ViT facial embeddings. They are the reusable visual features for each utterance.",
            "The fold CSVs tell the trainer which utterances belong to train and validation for each fold.",
            "best_model.pt is the checkpoint chosen by validation performance. It is the version used for the final test evaluation.",
            "metrics.json is the compact numerical summary of the run. It usually includes accuracy, weighted accuracy, unweighted accuracy, macro F1, and weighted F1.",
            "predictions_test.csv lists each utterance with its predicted label, which is what lets you inspect mistakes one by one.",
            "confusion_matrix.csv shows which classes were confused with which other classes. This is the most direct error-analysis output.",
        ],
    )
    add_para(
        doc,
        "Student interpretation: the output files are not separate random by-products. They form a chain. First the embeddings are saved, then the model is trained, then the best checkpoint is kept, then the checkpoint is evaluated, then metrics and predictions are written, and finally the confusion matrix explains the error pattern.",
    )
    doc.add_picture(str(OUTPUTS_PNG), width=Inches(6.9))
    add_code(
        doc,
        """data/processed/MELD_VIT_FACECUE/.../*.npy -> visual embeddings
fold CSVs -> train/val split instructions
best_model.pt -> strongest checkpoint
metrics.json -> summary scores
predictions_test.csv -> per-sample predictions
confusion_matrix.csv -> class confusion structure""",
    )

    doc.add_heading("5.4 Video-Only Control Execution Order", level=2)
    add_para(
        doc,
        "If you want to verify that the cropped ViT facial embeddings are useful on their own, present the control workflow in this order. This is a signal-check path before the tri-modal result is discussed.",
    )
    add_bullets(
        doc,
        [
            "Run scripts/run_meld_vit_facecrop_prepare.sh once so the face-crop manifest and control folds exist.",
            "Run scripts/run_meld_vit_facecrop_video_only_control_fold2.sh to train Fold 2 with only the cropped face embeddings.",
            "Run scripts/analyze_meld_vit_facecrop_video_only_control_fold2.sh to inspect the Fold 2 predictions and confusion matrix.",
            "Run scripts/run_meld_vit_facecrop_video_only_control_fold4.sh to repeat the same check for Fold 4.",
            "Run scripts/analyze_meld_vit_facecrop_video_only_control_fold4.sh to inspect the Fold 4 predictions and confusion matrix.",
            "Optionally run scripts/run_meld_vit_facecrop_video_only_control_suite.sh if you want a single wrapper for both folds.",
        ],
    )
    add_para(
        doc,
        "Student interpretation: this order tells the story that the face crop branch is not being trusted blindly. It is first tested alone, then later compared with the text+audio+video setting, so the audience can see whether the visual branch has independent emotion signal.",
    )

    doc.add_heading("6. Slide 6 - Full Frames vs Face Crop for Courtroom Testimony", level=1)
    add_para(
        doc,
        "This slide is about design choice, not just implementation. For courtroom testimony, the face is usually the most direct source of emotional evidence because it captures eyes, brows, mouth shape, hesitation, and stress. That is why face cropping is the better long-term facial-cue design.",
    )
    add_bullets(
        doc,
        [
            "Face crops keep the model focused on the expressive region of the speaker.",
            "Full frames are still a reasonable first baseline because they are robust and simple.",
            "Courtroom clips may have distance, occlusion, or side angles, so full frames can be safer at the beginning.",
            "The best scientific practice is to compare both and keep the version that actually helps MELD performance.",
        ],
    )
    add_para(
        doc,
        "The current code uses full sampled frames, not face crops. That is why the slide says face crop is better conceptually, but full-frame ViT is the current first implementation. This distinction is important for a student: the project is not claiming face-only ViT yet; it is testing a practical first version and keeping the face-crop idea as the next refinement.",
    )
    add_para(
        doc,
        "If asked why full frames were chosen first, the safest answer is that they preserve more visual context and avoid the failure mode of missing or badly cropped faces. If face detection later proves reliable, you can move toward a true facial-region pipeline.",
    )

    doc.add_heading("7. Slide 7 - Why Warm-Start from the Existing Checkpoint", level=1)
    add_para(
        doc,
        "This slide combines two ideas: why the baseline was chosen and why the new experiment should warm-start from it. The weighted-CE MELD Phase 1 checkpoint is the backbone because it is the strongest current conversational result. Starting from it makes the facial-cue experiment more controlled and more defensible.",
    )
    add_bullets(
        doc,
        [
            "Warm-starting preserves the useful text and audio knowledge already learned in Phase 1.",
            "It isolates the impact of the new visual branch instead of retraining the whole model from scratch.",
            "A small learning rate prevents the new branch from destroying the old decision boundary.",
            "The experiment becomes: does the face add value on top of an already-good MELD backbone?",
        ],
    )
    add_para(
        doc,
        "Code-level interpretation: the training script loads the old MELD checkpoint, rebuilds the model with video enabled, and then fine-tunes the combined system. That is the standard transfer-learning logic. You are not throwing away what the model already knows; you are adding another signal on top of it.",
    )
    add_code(
        doc,
        """# Conceptual warm-start flow
load weighted-CE checkpoint
keep text/audio backbone stable at first
enable the video branch
train with a smaller learning rate
keep the best validation checkpoint""",
    )

    doc.add_heading("8. Slide 8 - Utterance-to-Embedding Flow With Code Lines", level=1)
    add_para(
        doc,
        "This is the most technical slide in the deck. Explain it slowly and in order. For one utterance, the pipeline first finds the raw clip, then samples frames, then runs ViT, then saves a .npy file, and finally writes a manifest row that points to that file.",
    )
    add_bullets(
        doc,
        [
            "find_meld_clip(...) locates the raw clip for the utterance.",
            "sample_video_frames(...) selects evenly spaced frames and resizes them to 224 x 224.",
            "AutoImageProcessor converts the frames into ViT-ready tensors.",
            "AutoModel runs the pretrained ViT and returns token embeddings.",
            "outputs.last_hidden_state[:, 0, :] extracts the CLS embedding per frame.",
            "np.save(...) writes the final embedding array to disk as a .npy file.",
        ],
    )
    add_para(
        doc,
        "The exact code locations are scripts/build_meld_vit_facecue_manifest.py:123-135 for clip resolution and saving, scripts/build_meld_vit_facecue_manifest.py:101-102 for loading ViT, scripts/build_meld_vit_facecue_manifest.py:44-47 for inference, and src/data/preprocessing.py:41-75 for frame sampling.",
    )
    add_code(
        doc,
        """clip_path = find_meld_clip(meld_root / "raw", split, dialogue_id, utterance_id)
frames = sample_video_frames(video_path, cfg)
inputs = image_processor(images=list(batch), return_tensors="pt")
outputs = vit_model(**inputs)
hidden = outputs.last_hidden_state[:, 0, :]
np.save(video_feat_path, vit_embeddings)""",
    )
    add_para(
        doc,
        "Use the table in the slide to teach the data flow from left to right: annotation row, raw clip, sampled frame tensor, CLS embedding array, .npy file, and then manifest row. The student should leave this slide knowing that the manifest is a map from original video to saved features, not just a CSV for bookkeeping.",
    )
    add_para(
        doc,
        "The output column is the key to understanding what the pipeline produces at each step. It does not mean one single final answer. Instead, it shows the artifact created after each processing stage. In student terms, every input row is transformed into a usable visual feature that the later training code can load directly without repeating the expensive video processing work.",
    )
    add_bullets(
        doc,
        [
            "After the raw clip is read, the output is a sampled frame sequence.",
            "After the sampled frames are passed through ViT, the output is a CLS embedding tensor.",
            "After the embedding is saved, the output is a .npy file on disk.",
            "After the manifest is written, the output is a CSV row that points to that .npy file and keeps the dataset organized for training.",
        ],
    )
    add_para(
        doc,
        "The practical lesson is that the output column shows how the pipeline converts video into reusable data. This matters because the model training step later reads the .npy file and the manifest instead of decoding the raw video every time. That makes the experiment faster, cleaner, and easier to reproduce.",
    )
    add_para(
        doc,
        "If the slide table includes example columns, explain them as metadata rather than model inputs. The student should understand that these columns identify the sample, show where the raw media lives, and record where the processed feature was written.",
    )
    add_bullets(
        doc,
        [
            "utterance_id or sample_id identifies which dialogue turn the row belongs to.",
            "dialogue_id groups several utterances from the same conversation so splits can stay speaker- and dialogue-safe.",
            "label stores the emotion class that the model is trying to predict.",
            "raw clip path points to the original video file before any sampling or resizing.",
            "output embedding path points to the saved .npy file produced by ViT.",
            "split or fold columns tell the trainer whether the row is used for training, validation, or evaluation.",
        ],
    )
    add_para(
        doc,
        "A concrete example makes this easier to remember. A row might look like utterance_id = dia125_utt3, dialogue_id = dia125, label = surprise, split = train, raw clip path = data/MELD/raw/train/dia125_utt3.mp4, and output embedding path = data/processed/MELD_VIT_FACECUE/train/video/dia125_utt3.npy.",
    )
    add_para(
        doc,
        "That example means the row is telling you: this utterance came from dialogue 125, the emotion label is surprise, the original clip is stored in the raw MELD video folder, and the processed ViT feature was saved as a .npy file so training can reuse it later.",
    )
    add_para(
        doc,
        "A good way to say it in the presentation is that the example columns are the index card for the sample, while the output column is the final file created by the feature-extraction step. Together they let the student track one utterance from the original MELD annotation to the stored facial embedding.",
    )
    add_para(
        doc,
        "If the example output shows shapes such as (16, 224, 224, 3) and (16, 768) float32, explain them carefully. These shapes are the clearest summary of what the pipeline is doing numerically.",
    )
    add_bullets(
        doc,
        [
            "(16, 224, 224, 3) means the sample contains 16 frames, each resized to 224 by 224 pixels, with 3 color channels for RGB.",
            "The 16 is the number of sampled time steps, so the model sees the utterance as a short sequence of visual snapshots.",
            "The 224 x 224 size matches the standard ViT input resolution used by the pretrained backbone.",
            "The 3 channels mean red, green, and blue, which is the normal color format for image models.",
            "(16, 768) float32 means ViT produced 16 embedding vectors, one for each frame, and each vector has 768 features.",
            "float32 means the numbers are stored as 32-bit floating-point values, which is the usual numeric format for neural-network tensors.",
        ],
    )
    add_para(
        doc,
        "For a similar mental picture, you can imagine one row with sampled frame values shaped like (16, 224, 224, 3) and one saved feature file holding embeddings shaped like (16, 768) float32. So the first shape describes the image data before ViT, and the second shape describes the learned feature data after ViT. That is the key conceptual jump in the slide: pixels go in, embeddings come out.",
    )
    add_para(
        doc,
        "When you inspect one saved .npy file directly, the output is a compact numeric summary of the facial-cue representation. The file path tells you which utterance was processed, and the array values tell you that the model has converted the sampled frames into learned features rather than leaving them as raw image pixels.",
    )
    add_bullets(
        doc,
        [
            "shape: (16, 768) means one utterance became 16 frame embeddings, each with 768 features.",
            "dtype: float32 means the feature values are stored in standard neural-network precision.",
            "size: 12288 means the file contains 16 x 768 total numbers.",
            "finite_values: 12288/12288 means the file is numerically healthy and contains no NaN or infinity values.",
            "min / max show the numeric range of the learned embedding values.",
            "mean / std describe the center and spread of the embedding distribution.",
        ],
    )
    add_para(
        doc,
        "The first_frame_first_8_values and second_frame_first_8_values lines are only small previews of the embedding vector. They are not class labels and they are not directly human-interpretable emotions. They are just example coordinates in the learned feature space, showing that each frame becomes a different vector after ViT processing.",
    )
    add_para(
        doc,
        "A good way to explain this to a student is: the .npy file is the packed visual memory of one utterance. Instead of keeping the whole video around, the pipeline stores the important learned representation so the model can train faster and reproducibly from the same visual features.",
    )

    doc.add_heading("9. Slide 9 - Epoch Trend So Far", level=1)
    add_para(
        doc,
        "This slide explains how to read the training curve. The main point is that the warm-start facial-cue run does not keep improving forever. The best checkpoint appears early, around epoch 2, and later epochs begin to overfit.",
    )
    add_bullets(
        doc,
        [
            "Epoch 1 shows that training is stable and the warm-start is working.",
            "Epoch 2 gives the best validation accuracy so far.",
            "Epoch 3 and epoch 4 improve training accuracy but hurt validation.",
            "The correct interpretation is that the best checkpoint is likely the early validation peak, not the last epoch.",
        ],
    )
    add_para(
        doc,
        "This is a very important student lesson: training accuracy can keep rising while validation accuracy gets worse. That is overfitting. In this run, the model learns enough by epoch 2 to be useful, but later epochs start memorizing the training fold instead of improving generalization.",
    )
    add_code(
        doc,
        """epoch=1 | train_loss=1.4195 | train_acc=0.5975 | val_loss=1.6160 | val_acc=0.5414
epoch=2 | train_loss=1.5937 | train_acc=0.5576 | val_loss=1.6063 | val_acc=0.6010
epoch=3 | train_loss=1.3378 | train_acc=0.6492 | val_loss=1.7225 | val_acc=0.5893
epoch=4 | train_loss=1.1202 | train_acc=0.7261 | val_loss=1.9397 | val_acc=0.5399""",
    )
    doc.add_heading("9.1 Validation Curve Story", level=2)
    add_para(
        doc,
        "If you show the validation-curve chart, explain it as the same information as the table but in visual form. The chart helps the audience see the turning point faster: training keeps getting better, but validation peaks early and then weakens.",
    )
    add_bullets(
        doc,
        [
            "Epoch 1 is the warm-start ramp-up phase.",
            "Epoch 2 is the best validation point.",
            "Epochs 3 and 4 continue to fit the training fold more strongly.",
            "The visual gap between training and validation is the overfitting signal.",
        ],
    )
    add_para(
        doc,
        "Student takeaway: a validation curve is not just a pretty graph. It tells you when to stop training, which checkpoint to keep, and whether Phase 1 is already good enough to freeze before moving to Phase 2.",
    )
    add_para(
        doc,
        "Conclude the deck by saying that the next step is to keep the best validation checkpoint, compare the facial-cue run against the weighted-CE baseline, and then decide whether to refine the visual branch further or move on to the next modeling stage.",
    )
    add_para(
        doc,
        "If you want a short closing line for the live explanation: the current run suggests that facial cues are learnable, but the model should be evaluated using early stopping because the validation curve peaks before the end of training.",
        italic=True,
    )

    doc.add_heading("10. Slide 10 - Current Run Comparison: Where the Project Stands Now", level=1)
    add_para(
        doc,
        "Use this slide to compare the four MELD Fold 2 runs as a single story. The weighted-CE baseline is still the strongest stable model, focal loss from scratch is the weakest, warm-start focal is a controlled recovery attempt, and warm-start + ViT is the best improvement branch so far.",
    )
    add_bullets(
        doc,
        [
            "Weighted CE is the reference point because it gives the best overall balance of accuracy and F1.",
            "Focal loss from scratch is too unstable in this setting and should not be treated as a winning baseline.",
            "Warm-start focal is better than focal from scratch because it reuses the better MELD solution.",
            "Warm-start + ViT is the most informative improvement run because it includes the visual branch while still preserving the learned baseline.",
        ],
    )
    add_para(
        doc,
        "The teaching point is that this slide is not about one magic number. It is about ranking the experiments and showing that the current project status is: the baseline is strong, the improvement path is promising, and the visual branch is the most useful bridge to the next plan.",
    )

    doc.add_heading("10.1 What the comparison metrics mean", level=2)
    add_para(
        doc,
        "When you explain the comparison slide, do not treat the three scores as interchangeable. Each one answers a different question about the model.",
    )
    add_bullets(
        doc,
        [
            "Accuracy tells you the overall fraction of correct predictions. It is easy to read, but it can hide class imbalance.",
            "Weighted F1 gives more influence to classes that appear more often. It is useful when you want a single summary score that still respects the dataset frequency pattern.",
            "Macro F1 gives each class the same importance, even rare classes. It is the best score for checking whether the model is really helping minority emotions and not only the frequent ones.",
            "Best validation trend or best checkpoint means the epoch where validation performance was strongest. That is the model version you usually keep, even if later epochs have lower training loss.",
        ],
    )
    add_para(
        doc,
        "Student interpretation: if accuracy goes up but macro F1 stays weak, the model may be doing well on common classes while still missing rare ones. If weighted F1 is high but macro F1 is lower, that usually means the model is still biased toward frequent emotions. The best checkpoint is therefore chosen from validation, not from the final training loss alone.",
    )
    add_code(
        doc,
        """Accuracy      = overall correct predictions / all predictions
Weighted F1    = F1 score averaged with class-frequency weights
Macro F1       = unweighted average of per-class F1 scores
Best checkpoint = epoch with the strongest validation result""",
    )

    doc.add_heading("11. Slide 11 - Current Status: What the Experiments Say", level=1)
    add_para(
        doc,
        "This slide converts the comparison into a plain technical conclusion. The baseline has already solved the main Phase 1 problem well enough, so the remaining work is not more baseline squeezing. Instead, the project is now about controlled extension and legal-domain adaptation.",
    )
    add_bullets(
        doc,
        [
            "The weighted-CE result remains the best anchor for MELD Fold 2.",
            "The focal-only run shows that a loss change alone does not automatically improve the model.",
            "Warm-starting helps because it transfers the stronger learned boundary instead of restarting the whole optimization.",
            "The ViT branch is useful because it adds the facial-cue signal that the later courtroom work needs.",
        ],
    )
    add_para(
        doc,
        "For a student audience, the important explanation is that the project has moved from 'Can we reproduce a good baseline?' to 'What is the best next extension on top of that baseline?' That is a very different research stage.",
    )

    doc.add_heading("12. Slide 12 - Next Plan: Where the Remaining Work Should Go", level=1)
    add_para(
        doc,
        "This slide explains the next implementation step. The remaining effort should not be spent on repeating loss-function tuning. It should be spent on carrying the best current checkpoint into the courtroom-testimony novelty work.",
    )
    add_bullets(
        doc,
        [
            "Freeze the weighted-CE checkpoint as the stable reference.",
            "Keep the warm-start + ViT branch as the best improvement candidate.",
            "Shift the implementation effort to the legal-domain novelty pipeline.",
            "Use the same evaluation discipline so the new phase can be compared fairly against the current backbone.",
        ],
    )
    add_para(
        doc,
        "This is the point to say: the project is no longer just about benchmark chasing. The technical value now comes from extending the learned backbone into the courtroom-testimony use case.",
    )

    doc.add_heading("13. Slide 13 - Project Direction: Final Recommendation for the Current Stage", level=1)
    add_para(
        doc,
        "This final slide should sound like a clean recommendation rather than a transition gate. The current stage is strong enough to freeze the best baseline, keep the improvement branch as evidence of what helped, and move the remaining implementation effort into the legal-domain novelty work.",
    )
    add_bullets(
        doc,
        [
            "The current status is that Phase 1 has a strong baseline and a measured improvement path.",
            "The next plan is to use the best current checkpoint as the base for the courtroom-testimony research direction.",
            "The thesis novelty should focus on legal-domain adaptation, not on claiming a small benchmark gain as the main contribution.",
            "The deck should leave the reviewer with a clear message: the model is ready to be extended, not endlessly tuned.",
        ],
    )
    add_para(
        doc,
        "If you want one sentence to end the talk, say that the present results are strong enough to freeze the Phase 1 backbone and move the remaining engineering effort into the courtroom-testimony novelty implementation.",
        italic=True,
    )

    doc.add_heading("14. Slide 14 - Thesis Summary: What to do in the next two weeks", level=1)
    add_para(
        doc,
        "This final slide should answer the practical question the reviewer will ask: should you still test face crops before leaving Phase 1? The answer is yes, but only as one controlled ablation, not as a new long experiment track.",
    )
    add_bullets(
        doc,
        [
            "Use face crops because they focus ViT on the speaker's face instead of the whole frame.",
            "Keep the same MELD splits and the current best checkpoint so the comparison stays fair.",
            "Change only the visual input path: full-frame sampling becomes face detection + crop + resize + ViT.",
            "If face crops improve the result, freeze that setting as the final MELD visual design.",
            "If face crops do not help, keep the full-frame path and move to the courtroom-testimony implementation without further delay.",
        ],
    )
    add_para(
        doc,
        "Technical story to tell: the next two weeks should be spent on one clean visual ablation that answers whether face-focused ViT features are better than whole-frame ViT features on MELD. That is a stronger thesis move than endlessly tuning the same full-frame experiment.",
    )
    add_code(
        doc,
        """clip -> face detector -> crop speaker face -> resize 224x224 -> pretrained ViT -> .npy face embedding -> compare against full-frame baseline""",
    )

    doc.add_heading("15. Slide 15 - Third Guidance Call: Combined Phase 1 Design", level=1)
    add_para(
        doc,
        "This slide is the bridge into the third guidance call. It should explain that the project has now accumulated several Phase 1 enhancements on top of the original weighted-CE conversational baseline: full-frame ViT facial cues, face-crop ViT, gated fusion, and a video auxiliary-loss branch whose Fold 4 result is now complete and ready for comparison.",
    )
    add_bullets(
        doc,
        [
            "The backbone is still the paper-aligned MELD Phase 1 checkpoint.",
            "The visual branch has moved from full frames to face crops because the speaker face is the most relevant signal for testimony-style modeling.",
            "Gated fusion is used so the model can learn when to trust the visual branch instead of forcing a fixed contribution.",
            "The aux-loss branch now has a completed Fold 4 result, so it should be presented as a completed comparison point rather than a WIP item.",
        ],
    )
    add_para(
        doc,
        "Student explanation: this is not one new model. It is a controlled Phase 1 design that layers the strongest baseline, the facial-cue branch, and the gating mechanism in a way that keeps the comparison fair and interpretable.",
    )

    doc.add_heading("16. Slide 16 - Experimental Results, Confusion Matrix, and Error Analysis", level=1)
    add_para(
        doc,
        "This slide should present the combined Phase 1 evidence in a careful way. The baseline metrics, the facial-cue runs, the gated-fusion run, and the auxiliary-loss Fold 4 result can all be discussed as completed evidence, but the aux-loss row should still be interpreted as an incremental refinement rather than a new strongest baseline.",
    )
    add_bullets(
        doc,
        [
            "Show accuracy, weighted F1, and macro F1 together because each metric highlights a different aspect of performance.",
            "Use the confusion matrix to show where the model is still mixing similar emotions.",
            "In the error analysis, explain whether the main mistakes are due to class imbalance, neutral-class dominance, or visual branch under-weighting.",
            "Mark the aux-loss numbers as completed comparison evidence, but not as the final strongest claim.",
        ],
    )
    add_para(
        doc,
        "The important technical point is that error analysis should not only report the final scores. It should explain which emotions the model confuses and whether the new visual branch is changing the error pattern in a useful way.",
    )

    doc.add_heading("17. Slide 17 - Performance Evaluation and Model Comparison", level=1)
    add_para(
        doc,
        "This slide should compare the major Phase 1 variants in one place. The purpose is to show how the model evolved rather than to claim that every new branch beat the baseline.",
    )
    add_bullets(
        doc,
        [
            "Compare the paper-aligned weighted-CE baseline with the focal-loss-only attempt.",
            "Compare the warm-start focal run against the from-scratch focal run to show why warm-starting is the more stable training path.",
            "Compare the full-frame ViT, face-crop ViT, and gated-fusion variants to show how each architectural choice affects the visual contribution.",
            "Keep the auxiliary-loss row in the table and describe it as a completed refinement with modest gains relative to earlier facial-cue runs.",
        ],
    )
    add_para(
        doc,
        "Student-level takeaway: the goal of the table is not simply to sort models by accuracy. It is to explain which design choice changed the behavior, why it changed, and whether the change is worth carrying into the next research stage.",
    )

    doc.add_heading("18. Slide 18 - Journal Paper Draft", level=1)
    add_para(
        doc,
        "This slide should turn the Phase 1 work into a paper narrative. The key is to frame the contribution as an adaptation pipeline for legal-domain testimony modeling, not as a small benchmark improvement on MELD.",
    )
    add_bullets(
        doc,
        [
            "Abstract: pretrained multimodal baseline plus facial-cue extensions for courtroom-style testimony modeling.",
            "Introduction: explain why legal-domain testimony needs stronger speaker-aware cues than generic emotion recognition.",
            "Method: describe the weighted-CE backbone, facial-cue extraction, warm-starting, gated fusion, and the auxiliary-loss branch.",
            "Experiments: report the completed results, including the aux-loss Fold 4 comparison.",
            "Discussion: explain that the Phase 1 work builds the bridge to the next-stage novelty implementation.",
        ],
    )
    add_para(
        doc,
        "If you speak to the reviewer as a student researcher, the paper message is: 'We reproduced a strong multimodal backbone, added facial-cue experiments to understand visual contribution, and prepared the model for the legal-domain research stage.'",
    )

    doc.add_heading("Appendix - Overall Phase 1 Architecture Slide", level=1)
    add_para(
        doc,
        "This speaker-note section supports the added architecture slide in the deck. The point of the slide is to let the reviewer see the entire Phase 1 pipeline in one picture: raw utterance in, modality preprocessing in the middle, fusion and classification at the end, and the auxiliary-loss head included as part of the completed comparison set.",
    )
    add_bullets(
        doc,
        [
            "Tell the audience that the architecture is a complete multimodal pipeline, not a single isolated model.",
            "Explain that the video branch can be full-frame or face-crop depending on the experiment being discussed.",
            "Point out that the aux-loss head is included for design completeness and now has a completed Fold 4 result for reporting.",
            "Use the slide to connect the implementation details with the later metrics and confusion-matrix discussion.",
        ],
    )
    add_para(
        doc,
        "A student-friendly way to say it: the whole system takes the raw case utterance, encodes each modality with its own pretrained encoder, and then combines the results so the classifier can make one final emotion decision. That is the core of the Phase 1 architecture.",
    )

    doc.add_heading("Appendix - Training Parameters by Stage", level=1)
    add_para(
        doc,
        "This note block explains the hyperparameter choices across the Phase 1 experiments. The main teaching point is that each stage is a controlled variant of the same backbone, so the audience can see what changed and why.",
    )
    add_bullets(
        doc,
        [
            "Baseline stage: weighted cross entropy, best validation accuracy selection, and the original paper-aligned checkpoint.",
            "Full-frame ViT stage: 5 epochs, learning rate 5e-5, one warm-up freeze epoch, batch size 4.",
            "Face-crop stage: warm-start from the baseline, keep the same fold logic, and use a small learning rate to protect the original backbone.",
            "Gated-fusion stage: warm-start again, but allow the model to decide how much to trust the video branch.",
            "Aux-loss stage: warm-start, learning rate 2e-5, max 8 epochs, early stopping patience 2, and best checkpoint by validation weighted F1.",
        ],
    )
    add_para(
        doc,
        "The simplest student explanation is: earlier stages establish the backbone, visual stages test whether facial cues help, gated fusion tests whether the model can trust video selectively, and the auxiliary-loss stage is a final refinement whose Fold 4 result can now be discussed as part of the completed model comparison.",
    )

    return doc


def main() -> None:
    doc = build_doc()
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
