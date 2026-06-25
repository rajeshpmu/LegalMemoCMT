from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_Phase2_Dataset_Preparation_Deep_Dive.docx"
PIPELINE_PNG = ROOT / "phase2" / "assets" / "phase2_pipeline.png"
FINE_TUNE_PNG = ROOT / "phase2" / "assets" / "phase2_finetune_path.png"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def build_doc() -> Document:
    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LegalMemoCMT Phase 2 Dataset Preparation Deep Dive")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "A student-oriented explanation of how the courtroom-testimony dataset is being prepared, why these sources were chosen, and how the scripts work together."
    )
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12.5)

    add_para(
        doc,
        "This document is intentionally careful about claims. The dataset-preparation pipeline is a real contribution to the project because it makes the legal data path reproducible, traceable, and modular. That said, I should not call the dataset itself a novelty unless the final curated corpus is large enough, consistently validated, and materially different from a standard download-and-convert workflow. The safest phrasing is that the Phase 2 dataset preparation is a methodological novelty and a project contribution.",
    )
    add_para(
        doc,
        "From a student point of view, the important idea is that Phase 2 is not only about training a model. It is about building a usable courtroom-testimony dataset pipeline from public records, with enough structure that the later LegalMemoCMT fine-tuning step can trust the data.",
    )

    doc.add_heading("1. What Phase 2 Dataset Preparation Is Trying To Achieve", level=1)
    add_bullets(
        doc,
        [
            "Collect public courtroom and tribunal records in a reproducible way.",
            "Separate source planning from witness-level harvesting.",
            "Resolve transcript and video links from public pages or source URLs.",
            "Download or materialize raw media only when it is actually available.",
            "Turn the raw material into a structured dataset that LegalMemoCMT can read later.",
            "Add weak labels and status indicators so the data can be reviewed instead of blindly trusted.",
        ],
    )
    add_para(
        doc,
        "The point is not to pretend that public legal records already arrive as a perfect emotion dataset. They do not. The point is to create a disciplined preparation pipeline that turns heterogeneous public sources into something that can be used safely for research and later fine-tuning.",
    )

    doc.add_heading("2. Why These Sources Were Chosen", level=1)
    add_table(
        doc,
        ["Source", "Why it was chosen", "How it helps the project"],
        [
            ["IRMCT / ICTR / ICTY public judicial records", "They are the best match for courtroom and testimony-style analysis.", "They provide real judicial testimony structure, adversarial questioning, and emotionally relevant speaking patterns."],
            ["Supreme Court oral argument transcripts", "They are public, official, and easy to justify as legal-domain text.", "They help legal-language adaptation and transcript normalization even when audio/video is not available."],
            ["Witness-harvest manifest", "It gives the project a witness-level plan instead of a vague corpus-level target.", "It lets the pipeline work at the level of witnesses, cases, durations, and utterance counts."],
            ["Tribunal source target dataset", "It provides a high-level acquisition map across tribunal families.", "It tells the project what source families and target volumes should be collected."],
        ],
    )
    add_para(
        doc,
        "As a student, I would think of the sources in two layers. The tribunal source CSV says where to look in general. The witness harvest manifest says which witness-level records to actually harvest first. That separation is important because it keeps the project organized and makes incremental updates possible.",
    )

    doc.add_heading("3. The Dataset Design Philosophy", level=1)
    add_bullets(
        doc,
        [
            "Traceability first: every row should remain linked back to a public source record.",
            "Incremental updates: new witnesses should be addable without rebuilding the whole dataset.",
            "Modality-aware curation: keep track of whether transcript, audio, and video are really available.",
            "Weak supervision before overclaiming: assign provisional labels transparently and mark uncertainty.",
            "Legal safety: preserve protected witness identifiers and never infer guilt, innocence, or deception.",
        ],
    )
    add_para(
        doc,
        "This is where the novelty argument becomes careful and useful. If the project shows a clean, reproducible, public-record-to-manifest pipeline with explicit validation, then the preparation method itself is a contribution. I would describe that as a novelty in dataset preparation methodology, not as a claim that the data source alone is a new benchmark.",
    )

    doc.add_heading("4. The Main Scripts and What Each One Does", level=1)
    add_table(
        doc,
        ["Script", "Purpose", "Why a student should care"],
        [
            ["phase2/dataset_builder.py", "Core module with validation, resolution, download, segmentation, labels, and dashboard functions.", "This is the main implementation. The rest of the pipeline calls into it."],
            ["phase2/run_phase2_dataset_pipeline.sh", "Runs the end-to-end data preparation sequence.", "This is the easiest way to execute the pipeline step by step."],
            ["phase2/run_phase2_full.sh", "Chains dataset prep, fine-tuning, and evaluation.", "This is the single-command version for the whole Phase 2 flow."],
            ["phase2/run_phase2_finetune.sh", "Fine-tunes from the best MELD checkpoint.", "This is the adaptation stage after the dataset is ready."],
            ["phase2/evaluate_phase2_checkpoint.sh", "Evaluates the saved Phase 2 checkpoint.", "This produces the final metrics file."],
            ["scripts/check_manifest_assets.py", "Checks whether manifest-referenced files already exist on disk.", "This answers the practical question: do we still need the raw data?"],
            ["scripts/check_phase1_meld_ready.sh", "Checks multiple MELD manifests for readiness.", "This helps the RunPod workflow before training starts."],
            ["scripts/check_meld_ready.sh", "Checks one manifest for file presence.", "This is the lightweight version of the readiness check."],
        ],
    )

    doc.add_heading("5. Script-by-Script Technical Explanation", level=1)

    doc.add_heading("5.1 phase2/dataset_builder.py", level=2)
    add_para(
        doc,
        "This is the main Phase 2 implementation file. It is not just one function. It is the place where the full data lifecycle is defined: load the source CSVs, validate them, resolve links, download or materialize files, segment transcripts, build the final dataset, generate weak labels, and write the dashboard. In other words, the module is the operational center of Phase 2.",
    )
    add_bullets(
        doc,
        [
            "load_tribunal_sources() validates the acquisition plan CSV.",
            "load_manifest() validates the witness-level manifest and assigns stable manifest IDs.",
            "resolve_transcript_links() and resolve_video_links() search or normalize public URLs.",
            "download_transcript(), download_video(), and extract_audio() materialize the media into local folders.",
            "segment_transcript() turns transcript text into utterance-level rows.",
            "build_final_dataset() creates the final LegalMemoCMT-ready CSV.",
            "generate_weak_labels() exports separate label files for review.",
            "build_dashboard() writes the dataset status HTML report.",
        ],
    )

    doc.add_heading("5.2 phase2/run_phase2_dataset_pipeline.sh", level=2)
    add_para(
        doc,
        "This wrapper is the safest way to run the Phase 2 dataset preparation stages in order. It is intentionally not a black box. A student can read it and understand the sequence: validate source CSVs, resolve records, materialize files, build the dataset, generate weak labels, and write the dashboard.",
    )
    add_bullets(
        doc,
        [
            "It is CPU-bound because it mostly handles parsing, resolving, and file creation.",
            "It does not train a model.",
            "It should be run before any Phase 2 checkpoint fine-tuning.",
            "It is the best wrapper when you want to check the dataset state step by step.",
        ],
    )

    doc.add_heading("5.3 phase2/run_phase2_full.sh", level=2)
    add_para(
        doc,
        "This is the convenience wrapper that chains Phase 2 dataset preparation, fine-tuning, and evaluation. It is useful when the dataset is already stable and you want a single command. It should not replace understanding the individual stages, because debugging is easier when you know which step failed.",
    )

    doc.add_heading("5.4 phase2/run_phase2_finetune.sh", level=2)
    add_para(
        doc,
        "This script fine-tunes LegalMemoCMT from the best MELD checkpoint. The idea is that MELD has already taught the model how to do multimodal emotion recognition, so Phase 2 does not start from zero. Instead, it adapts that representation to courtroom testimony and legal-domain text.",
    )
    add_bullets(
        doc,
        [
            "It prefers CUDA automatically if a GPU is visible.",
            "It uses the paper encoder mode and text+audio as the core path.",
            "It loads the MELD checkpoint and resets the classifier head if needed.",
            "It uses a smaller learning rate because this is adaptation, not pretraining.",
        ],
    )

    doc.add_heading("5.5 phase2/evaluate_phase2_checkpoint.sh", level=2)
    add_para(
        doc,
        "This script evaluates the saved Phase 2 checkpoint and writes the output metrics. It is separate from training because the final report should be based on a saved checkpoint, not on whatever happened to be the last batch during training.",
    )

    doc.add_heading("5.6 scripts/check_manifest_assets.py", level=2)
    add_para(
        doc,
        "This checker is the practical answer to the question, 'Do I still need raw data on the pod?' It scans the manifest and confirms whether the referenced files already exist. If everything is already present, you may not need the raw archive for that particular run. If files are missing, then the raw data or the derived features are still required.",
    )

    doc.add_heading("5.7 scripts/check_phase1_meld_ready.sh and scripts/check_meld_ready.sh", level=2)
    add_para(
        doc,
        "These wrappers are readiness tools for the MELD workflow. They are not data builders. Their job is to quickly tell the student whether the manifest references are in place, whether the environment is healthy, and whether the pod is ready to move on to training.",
    )

    doc.add_heading("5.8 Why the wrappers matter", level=2)
    add_para(
        doc,
        "As a student, I should not think of the wrappers as unnecessary shell files. They are there to protect the workflow from accidental order mistakes. Data prep, training, and evaluation each have different responsibilities, and the wrappers make those responsibilities visible.",
    )

    doc.add_heading("6. How the Dataset Is Created", level=1)
    add_numbered(
        doc,
        [
            "Start with the two source CSVs stored under data/phase2/source_manifests/.",
            "Validate that the source rows are structurally correct and have usable URLs or URL placeholders.",
            "Resolve transcript and video links into concrete downloadable locations or direct file references.",
            "Materialize the records into data/raw/transcripts/, data/raw/videos/, and data/raw/audio/ when media exists.",
            "Build the final dataset CSV with utterance-level rows and modality paths.",
            "Generate weak labels for emotion, question type, and credibility status.",
            "Write the dashboard so collection coverage can be reviewed.",
        ],
    )
    add_para(
        doc,
        "This flow matters because it mirrors how a real research dataset should be handled. First you prove the source exists. Then you locate the data. Then you store it locally. Then you convert it into model-ready rows. Then you label or weakly label it. Then you review it before training.",
    )

    doc.add_heading("7. How the Dataset Helps LegalMemoCMT", level=1)
    add_bullets(
        doc,
        [
            "It gives the model courtroom-style text rather than generic conversational text.",
            "It creates a bridge from public judicial records to emotion modeling.",
            "It allows a legal-domain warm start before Phase 2 modeling.",
            "It keeps the modality paths explicit so later experiments can compare text-only, audio-only, and multimodal settings.",
            "It supports explanation because every row stays traceable to a source and a weak-label rule.",
        ],
    )
    add_para(
        doc,
        "The biggest contribution here is not simply that the model has more data. It is that the data is structured in a way that fits LegalMemoCMT and can be defended in a thesis or mentor discussion. That is the real value of the preparation pipeline.",
    )

    doc.add_heading("8. How to Interpret the Novelty Carefully", level=1)
    add_para(
        doc,
        "If the thesis says the dataset preparation is a novelty, the claim should be specific. The novelty is in the legally grounded, manifest-driven, weakly supervised, incremental preparation pipeline. It is not safe to claim that public tribunal records by themselves are novel. The novelty is in the way the project turns them into a reproducible research dataset for LegalMemoCMT.",
    )
    add_bullets(
        doc,
        [
            "Good claim: this project introduces a structured Phase 2 data-preparation pipeline for courtroom testimony.",
            "Good claim: the pipeline uses public judicial records with traceable manifests and weak labels.",
            "Too strong: the dataset is a new benchmark unless the final corpus is truly large and validated enough to justify that wording.",
            "Too strong: the dataset preparation alone proves model superiority.",
        ],
    )

    doc.add_heading("9. Pipeline Diagram", level=1)
    if PIPELINE_PNG.exists():
        doc.add_picture(str(PIPELINE_PNG), width=Inches(6.3))
        cap = doc.add_paragraph("Figure 1. Phase 2 data preparation and warm-start pipeline.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if FINE_TUNE_PNG.exists():
        doc.add_picture(str(FINE_TUNE_PNG), width=Inches(6.1))
        cap = doc.add_paragraph("Figure 2. Warm-start path from the best MELD checkpoint.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("10. What A Student Should Remember", level=1)
    add_bullets(
        doc,
        [
            "The source manifests are the starting point.",
            "The dataset builder is the core implementation.",
            "The wrappers exist to preserve order and reproducibility.",
            "The raw data is not needed if the manifest references already exist and are fully materialized.",
            "The Phase 2 dataset preparation is a methodological contribution and may be described as a novelty only with that caution.",
            "The final training benefit is that LegalMemoCMT can adapt to courtroom testimony instead of only benchmark emotion corpora.",
        ],
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()

