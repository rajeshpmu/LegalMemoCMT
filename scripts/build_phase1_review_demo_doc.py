#!/usr/bin/env python3
from __future__ import annotations

import pandas as pd
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DEMO_DIR = ROOT / "results" / "phase1_review_demo" / "fold2_baseline"
OUTPUT = ROOT / "implementation_docments" / "Phase1_Review_Demo_Guide.docx"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def load_examples() -> pd.DataFrame:
    path = DEMO_DIR / "demo_examples.csv"
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)


def load_metrics() -> pd.DataFrame:
    path = DEMO_DIR / "metrics_summary.csv"
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)


def build_doc() -> Document:
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Phase 1 Review Demo Guide")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Student-level step-by-step explanation of the live demo, demo scripts, example selection, and expected reviewer discussion")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)

    add_para(
        doc,
        "Purpose: this guide explains how to present the Phase 1 demo in a way that a student can follow and explain clearly during the final review. The demo uses saved MELD prediction outputs, so it shows real model behavior without requiring a rerun of training. The reviewer can see example videos, predicted emotions, metrics, and confusion patterns in one flow.",
    )

    doc.add_heading("1. What the Demo Is Trying to Prove", level=1)
    add_bullets(
        doc,
        [
            "The model can accept test examples and produce emotion predictions.",
            "The predictions can be compared against the ground-truth labels.",
            "The full test-set metrics summarize performance beyond one example.",
            "The confusion matrix and top confusions show the model’s actual failure pattern.",
            "The demo shows that Phase 1 is not just trained, but also inspectable and explainable.",
        ],
    )
    add_para(
        doc,
        "In student terms, the demo proves that the pipeline is operational from input to output: test sample in, emotion prediction out, metrics summary available, and error analysis available."
    )

    doc.add_heading("2. Scripts Used for the Demo", level=1)
    add_para(
        doc,
        "The demo is intentionally lightweight. It does not train a new model. It reads already completed MELD outputs and packages them into a reviewer-friendly bundle.",
    )
    add_numbered(
        doc,
        [
            "Run scripts/run_phase1_review_demo_bundle.sh to build the demo bundle from existing MELD Fold 2 outputs.",
            "The wrapper calls scripts/build_phase1_review_demo_bundle.py, which merges the saved predictions with the manifest and picks example cases.",
            "The script writes a metrics summary, a demo example table, and copies the confusion matrix and top confusions into a single output folder.",
        ],
    )
    add_para(
        doc,
        "The important point is that the demo is based on already saved evaluation artifacts. That makes the review reproducible and avoids changing the model right before the presentation."
    )

    doc.add_heading("2.1 How to Run One Video at a Time", level=1)
    add_para(
        doc,
        "If you want a live-style demo, use the single-video runner instead of the bundle script. This script reads one MELD sample_id, runs the saved Phase 1 checkpoint on that example, and prints the final emotion prediction plus the top-3 classes.",
    )
    add_numbered(
        doc,
        [
            "Choose a sample_id from the demo bundle, for example test_dia185_utt4.",
            "Run bash scripts/run_phase1_single_video_demo.sh <sample_id>.",
            "The script loads the saved MELD Fold 2 checkpoint.",
            "It reads the matching manifest row for that sample.",
            "It prints the ground truth emotion, the predicted emotion, whether the result is correct, and the top-3 predicted classes.",
        ],
    )
    add_bullets(
        doc,
        [
            "Use this when you want to speak through one sample interactively.",
            "Use the bundle script when you want to prepare the example list and the summary tables first.",
            "The single-video runner is manifest-backed, so it uses the stored feature paths from the pipeline instead of training a new model.",
        ],
    )
    add_para(
        doc,
        "Student explanation: this script is like a live microscope. It lets the reviewer see one utterance go through the already trained Phase 1 model and returns the emotion decision immediately."
    )
    add_para(
        doc,
        "The output also prints a 'Missing keys during load' message for the extra gated-fusion and auxiliary-loss layers. That is expected when the checkpoint comes from the earlier paper-aligned baseline, because the newer model class contains additional modules that were not part of the original checkpoint. The important point is that the core baseline weights still load and the prediction still runs."
    )

    doc.add_heading("2.1.1 Worked Single-Video Examples", level=2)
    add_para(
        doc,
        "The table below uses the attached single-video demo outputs. These examples are useful because they show both correct and incorrect decisions in a very small live explanation.",
    )
    add_table(
        doc,
        ["Sample ID", "Truth", "Prediction", "Correct?", "Top-1 confidence", "Student interpretation"],
        [
            ["test_dia143_utt2", "neutral", "neutral", "Yes", "0.3237", "The model is confident enough to keep the neutral label."],
            ["test_dia244_utt14", "neutral", "neutral", "Yes", "0.2579", "This is another neutral example where the model stays correct."],
            ["test_dia153_utt5", "neutral", "neutral", "Yes", "0.3419", "The model again uses neutral as the top class, but not with extreme confidence."],
            ["test_dia22_utt5", "neutral", "surprise", "No", "0.3033", "This shows the main failure mode: a neutral utterance can shift to a nearby emotion like surprise."],
        ],
    )
    add_bullets(
        doc,
        [
            "The first three examples show that the model can recognize neutral correctly.",
            "The last example shows that the model is not perfect and still confuses neutral with an emotionally close class.",
            "The top-3 probabilities show that the model is often unsure between neutral and nearby emotions, which is exactly what the confusion matrix also shows.",
        ],
    )
    add_para(
        doc,
        "Student takeaway: the live single-video demo is not just about showing one label. It is about showing how the model behaves under uncertainty and how its confidence shifts when it sees emotionally close examples."
    )

    doc.add_heading("2.1.2 Raw MP4-to-Prediction Pipeline", level=2)
    add_para(
        doc,
        "This is the true live-video version of the demo. Instead of starting from a saved .npy feature file, the pipeline starts from a raw .mp4 clip, extracts visual features first, and then feeds those features into the same trained Phase 1 model. That is the cleanest way to show a reviewer the full input-to-output story.",
    )
    add_numbered(
        doc,
        [
            "Read one raw MELD .mp4 clip directly from disk.",
            "Sample fixed frames from the clip.",
            "Run the sampled frames through a pretrained ViT.",
            "Save the resulting embeddings as a cached .npy file.",
            "Load the corresponding text/audio metadata from the manifest.",
            "Run the already trained Phase 1 checkpoint on the extracted features.",
            "Print the final prediction, confidence, and top-3 classes.",
        ],
    )
    add_para(
        doc,
        "The command for this mode is:"
    )
    add_para(
        doc,
        "bash scripts/run_phase1_raw_mp4_demo.sh <sample_id> <raw_video.mp4>",
        bold=True,
    )
    add_bullets(
        doc,
        [
            "Use vision-mode facecrop to crop the face before ViT when you want a courtroom-style face-focused demo.",
            "Use vision-mode fullframe when you want to show the broader scene-processing branch.",
            "The raw mp4 wrapper caches the extracted .npy file so the same clip does not need to be reprocessed every time.",
        ],
    )
    add_para(
        doc,
        "Student explanation: this script is the bridge between the original video evidence and the final emotion prediction. It makes the preprocessing visible, which is useful when the reviewer wants to understand exactly what the model sees before it makes a decision."
    )
    add_para(
        doc,
        "The two attached raw-mp4 runs are good worked examples because one is a correct prediction and the other is a near-miss. That lets you explain both success and failure from the same live pipeline.",
    )
    add_table(
        doc,
        ["Sample ID", "Truth", "Prediction", "Correct?", "Confidence", "Top-3 interpretation"],
        [
            ["test_dia279_utt9", "neutral", "neutral", "Yes", "0.2722", "Neutral is the top class, with anger and joy as nearby alternatives."],
            ["test_dia278_utt5", "surprise", "anger", "No", "0.2123", "A near-miss: anger wins by a small margin, but surprise remains in the top-3."],
        ],
    )
    add_bullets(
        doc,
        [
            "The first example shows the raw-video pipeline can give the right emotion label end to end.",
            "The second example shows that the model still confuses emotionally close classes when the signal is ambiguous.",
            "These two cases are useful together because they show the reviewer the model is learning, but not perfectly.",
        ],
    )
    add_para(
        doc,
        "Student takeaway: the raw-mp4 demo is strongest when you use one clear correct case and one clear near-miss case. That combination tells the story of the model much better than a long list of outputs."
    )
    doc.add_heading("2.1.3 Viva / Examiner Quick Checks", level=2)
    add_para(
        doc,
        "If an examiner asks why a prediction is correct or wrong, the best answer is to walk through the code in the same order as the pipeline. That shows you understand where the output came from, not just the final label.",
    )
    add_numbered(
        doc,
        [
            "Check the manifest row first so you know the sample_id, transcript, audio path, and video path.",
            "Inspect src/data/preprocessing.py to see how the raw clip becomes sampled frames and then ViT embeddings.",
            "If the demo used facecrop, inspect scripts/build_meld_vit_facecrop_manifest.py to confirm the face-cropping step.",
            "If the demo used fullframe, inspect scripts/build_meld_vit_facecue_manifest.py to confirm the full-frame ViT path.",
            "Inspect scripts/predict_phase1_raw_mp4_demo.py to see how the cached .npy file and manifest metadata are loaded into the trained Phase 1 model.",
            "Compare ground_truth_label and predicted_label, then use the top-3 probabilities to explain why the model chose that class.",
        ],
    )
    add_bullets(
        doc,
        [
            "If the correct class is top-1, explain that the model found the strongest emotion signal for that clip.",
            "If the correct class is in the top-3 but not top-1, explain that the model saw the right emotion but another class was slightly stronger.",
            "If the correct class is not in the top-3, explain that the model probably fell back to a nearby class or neutral because the signal was ambiguous.",
        ],
    )
    add_para(
        doc,
        "Important clarification: the top-3 values are not accuracy, weighted F1, unweighted F1, or macro F1. They are softmax probabilities for one sample. They show how the model distributes its belief across classes for that single video. Accuracy and F1 are dataset-level metrics computed after many samples are predicted."
    )
    add_para(
        doc,
        "Why softmax is used: the model first produces raw scores called logits. Those logits are not easy to read directly because they can be any positive or negative numbers. Softmax converts them into probability-like values between 0 and 1, and the values across all emotion classes add up to 1. That makes the output human-readable during a demo."
    )
    add_bullets(
        doc,
        [
            "Logits are internal model scores before normalization.",
            "Softmax converts those scores into a distribution over the emotion classes.",
            "The highest softmax value becomes the predicted class.",
            "The top-3 values show the model's closest alternatives for that one sample.",
            "This is helpful because a reviewer can see whether the model was confident or only slightly leaning toward one label.",
        ],
    )
    add_para(
        doc,
        "Student-level explanation: softmax is not the same as accuracy or F1. It is the model's per-sample belief distribution. In the demo, softmax helps you say not only 'the model predicted surprise' but also 'surprise was 21%, neutral was 18%, and anger was 16%, so the model was uncertain between nearby emotions.'"
    )
    add_para(
        doc,
        "Student wording: 'The top-3 values tell me the model's confidence for this one video. Accuracy and F1 tell me how the model behaved on the whole test set.' That is the clean distinction to remember in a viva."
    )
    add_table(
        doc,
        ["Possible viva question", "Short answer", "Code location to inspect"],
        [
            ["Why is this prediction correct or wrong?", "Trace the sample from manifest row to preprocessing to ViT features to model output.", "src/data/preprocessing.py, scripts/predict_phase1_raw_mp4_demo.py"],
            ["Why use softmax?", "It turns raw logits into probability-like values that are easy to read for one sample.", "scripts/predict_phase1_raw_mp4_demo.py"],
            ["What does top-3 mean?", "The three highest softmax probabilities for that single video.", "scripts/predict_phase1_raw_mp4_demo.py"],
            ["Why is neutral common?", "The dataset is imbalanced and the model often falls back to neutral when uncertain.", "results/paper_aligned_meld_cv/cmt_min/fold_2/analysis_test/confusion_matrix.csv"],
            ["How do I change facecrop vs fullframe?", "Edit the preprocessing branch that extracts ViT features before inference.", "scripts/build_meld_vit_facecrop_manifest.py, scripts/build_meld_vit_facecue_manifest.py"],
        ],
    )
    add_para(
        doc,
        "Student note: this table is useful in a viva because it tells you both what to say and where the code is. If the examiner wants a deeper answer, you can open the listed script and explain the exact step that produced the prediction."
    )
    add_table(
        doc,
        ["Training / pipeline question", "Short answer", "Code location to inspect"],
        [
            ["How is the model trained?", "The train script loads a manifest, builds batches, runs epochs, computes loss, and saves the best checkpoint.", "src/train/train.py"],
            ["What changes in gated fusion?", "The gate learns how much to trust the video branch relative to text and audio.", "src/models/model.py and scripts/run_meld_vit_facecrop_gated_fold2.sh"],
            ["What does aux-loss do?", "It adds an extra video-related objective so the model can learn a useful visual signal during training.", "src/models/model.py and scripts/run_meld_vit_facecrop_gated_video_aux_fold2.sh"],
            ["How are metrics computed?", "Prediction CSVs are converted into accuracy, weighted F1, macro F1, and confusion tables.", "src/train/evaluate.py, scripts/analyze_predictions.py"],
            ["What does the batch raw-mp4 demo do?", "It runs the single-video raw pipeline on multiple clips and stores one log per sample.", "scripts/run_phase1_raw_mp4_demo_batch.sh"],
        ],
    )
    add_para(
        doc,
        "Student note: the second table is the one to use when the examiner asks about the learning mechanism, the fusion mechanism, or the evaluation logic. The first table is for prediction-level questions, while this one is for training-level questions."
    )
    add_table(
        doc,
        ["Comparison question", "Short answer", "Code / result to inspect"],
        [
            ["Fold 2 vs Fold 4?", "Fold 2 is the strongest stable anchor; Fold 4 is useful to check whether the behavior stays similar on another split.", "results/paper_aligned_meld_cv/cmt_min/fold_2/metrics.json and fold_4/metrics.json"],
            ["Weighted-CE vs warm-start focal?", "Weighted-CE is still the stronger baseline; warm-start focal is a diagnostic run that did not surpass it.", "results/paper_aligned_meld_cv/cmt_min/fold_2/metrics.json and results/improvement/warmstart_focal/meld_selected/cmt_min/fold_2/metrics.json"],
            ["Facecrop vs fullframe?", "Facecrop focuses on the speaker's face and is better when the goal is facial-cue analysis; fullframe keeps more scene context.", "scripts/build_meld_vit_facecrop_manifest.py, scripts/build_meld_vit_facecue_manifest.py"],
            ["Why Phase 1 stops here?", "Phase 1 has a strong, reproducible backbone and clear error patterns, so the next research gain should come from Phase 2 rather than endless tuning.", "Phase 1 demo bundle, confusion matrices, and the Phase 2 planning docs"],
            ["Why move to Phase 2?", "The remaining errors are systematic and suggest that legal-domain adaptation and courtroom-specific modeling are the next meaningful step.", "Phase 1 conclusion slides and the Phase 2 roadmap"],
        ],
    )
    add_para(
        doc,
        "Student note: this third table is the one to use when the examiner compares experimental stages or asks why you are stopping Phase 1 instead of continuing to tune it. It keeps the comparison short, but you can always open the referenced result files if the examiner wants the exact numbers."
    )
    add_bullets(
        doc,
        [
            "If you want to change frame sampling, edit src/data/preprocessing.py in sample_video_frames.",
            "If you want to change the face crop logic, edit scripts/build_meld_vit_facecrop_manifest.py.",
            "If you want to change the full-frame ViT extraction, edit scripts/build_meld_vit_facecue_manifest.py.",
            "If you want to change how the raw mp4 demo behaves, edit scripts/predict_phase1_raw_mp4_demo.py.",
            "If you want to change the actual training method, edit src/train/train.py and the run scripts that point to it.",
        ],
    )
    add_para(
        doc,
        "This separation is important: preprocessing changes the input representation, training changes how the model learns from that representation, and the demo wrapper only connects the two for live explanation."
    )

    doc.add_heading("2.1.4 Batch Raw MP4 Demo", level=2)
    add_para(
        doc,
        "If the reviewer wants to see several raw clips one after another, use the batch runner. The batch version reads a CSV list of sample_id and video_path pairs and runs the raw-mp4 demo once per row. This is useful when you want to prepare a small demo set in advance and then walk through the clips in a controlled order.",
    )
    add_numbered(
        doc,
        [
            "Create a CSV file with columns sample_id and video_path.",
            "List the test clips in the order you want to present them.",
            "Run bash scripts/run_phase1_raw_mp4_demo_batch.sh <pairs.csv>.",
            "The batch runner calls the single-video script for each row.",
            "Each clip gets its own log file.",
            "A summary.csv file is written at the end so you can review which clips succeeded and which failed.",
        ],
    )
    add_para(
        doc,
        "The batch script is helpful because it turns the raw-video pipeline into a repeatable mini-benchmark for the review. It is still the same model, but now the user can show multiple examples without manually typing each command."
    )
    add_bullets(
        doc,
        [
            "The batch summary is saved in results/phase1_review_demo/raw_mp4_batch/summary.csv by default.",
            "Per-sample logs are saved in the logs/ subfolder.",
            "This lets the reviewer inspect the exact output of each video after the demo.",
        ],
    )
    add_para(
        doc,
        "Student explanation: the batch runner is just a convenience layer over the single-video runner. It does not change the model. It only automates the sequence of demos and keeps the outputs organized."
    )

    doc.add_heading("2.2 Which Trained Model Stage the Demo Uses", level=1)
    add_para(
        doc,
        "The demo uses the saved paper-aligned MELD Fold 2 checkpoint from Phase 1. That checkpoint corresponds to the trained HuBERT + BERT + cross-modal transformer fusion model with MIN pooling, which is the strongest stable conversational baseline currently available in the project.",
    )
    add_bullets(
        doc,
        [
            "Backbone: pretrained text and audio encoders with cross-modal fusion.",
            "Training style: paper-aligned MELD 5-fold cross-validation.",
            "Checkpoint used for the demo bundle: the Fold 2 best_model.pt from the paper-aligned MELD run.",
            "Why this one: it is stable, already analyzed, and strong enough to present as the Phase 1 baseline.",
        ],
    )
    add_para(
        doc,
        "From a student point of view, this is the model stage that represents the project after the main Phase 1 implementation work was completed. It is the reference model for both the demo and the planned Phase 2 comparison."
    )

    doc.add_heading("2.3 How the Model Was Trained", level=1)
    add_numbered(
        doc,
        [
            "The MELD dataset was divided into five dialogue-level folds.",
            "For each fold, the model was trained on the fold-specific train split and validated on the fold-specific validation split.",
            "The pretrained text and audio encoders were fused with the cross-modal transformer block.",
            "The best checkpoint for each fold was selected using validation performance.",
            "The selected checkpoint was then evaluated on the held-out MELD test set.",
            "Predictions were exported to CSV and analyzed to produce the confusion matrix and top confusions.",
        ],
    )
    add_para(
        doc,
        "This means the demo is based on a proper trained-and-evaluated research model, not a toy example or an unverified checkpoint."
    )

    doc.add_heading("2.4 Training and Fine-Tuning Results That Justify the Baseline", level=1)
    add_table(
        doc,
        ["Metric", "Fold 2 baseline value"],
        [
            ["Accuracy", "0.6375"],
            ["Weighted F1", "0.6254"],
            ["Macro F1", "0.4430"],
            ["Unweighted Accuracy", "0.4369"],
            ["Weighted Accuracy", "0.6375"],
        ],
    )
    add_bullets(
        doc,
        [
            "The baseline is strong on the headline weighted metrics.",
            "Macro F1 is lower than weighted F1, which means minority classes are still harder to separate.",
            "The model is learning real emotion structure rather than random label patterns.",
            "That makes it a good Phase 1 anchor because it is stable, reproducible, and already interpreted.",
        ],
    )
    add_para(
        doc,
        "The short comparison below shows where the other Phase 1 experiments sit relative to this baseline. It is useful for the review because it explains why the weighted-CE Fold 2 checkpoint is still the reference model even though later video experiments add useful information.",
    )
    add_table(
        doc,
        ["Stage", "Accuracy", "Weighted F1", "Macro F1", "What it means"],
        [
            ["Weighted-CE baseline", "0.6375", "0.6254", "0.4430", "Strongest stable conversational anchor."],
            ["Warm-start focal", "0.4992", "0.5043", "0.3224", "Loss change alone did not beat the baseline."],
            ["Face-crop ViT", "0.3452", "0.3104", "0.1352", "Face crop alone is too weak as a standalone branch."],
            ["Gated fusion", "0.6222", "0.6109", "0.4191", "Video becomes useful when the model learns to trust it selectively."],
            ["Aux-loss", "0.6054", "0.6022", "0.4351", "Useful refinement, but still below the weighted-CE baseline."],
        ],
    )
    add_para(
        doc,
        "The key interpretation is that the model is good enough to serve as a scientific reference point. That is exactly what a baseline should be: not perfect, but trustworthy."
    )

    doc.add_heading("2.5 Why This Is the Baseline for Phase 1 and Planned Phase 2", level=1)
    add_bullets(
        doc,
        [
            "It is the best stable conversational result currently available in the project.",
            "It has already been trained, validated, exported, and analyzed.",
            "It gives a clean reference for all later facial-cue and courtroom-adaptation experiments.",
            "It shows the model already learns useful multimodal emotion structure, which is important before adding a new domain.",
            "Phase 2 should build on a known-good backbone instead of starting from an unstable model.",
        ],
    )
    add_para(
        doc,
        "For the planned Phase 2 courtroom testimony work, this baseline matters because it gives a controlled starting point. If the next stage adds facial cues, gated fusion, or courtroom-domain adaptation, the reviewer can compare every new result against this checkpoint instead of against a weaker or changing model."
    )

    doc.add_heading("3. Command to Build the Demo Bundle", level=1)
    add_para(doc, "The default command is:")
    add_para(
        doc,
        "bash scripts/run_phase1_review_demo_bundle.sh",
        bold=True,
    )
    add_para(
        doc,
        "The default source files are the MELD Fold 2 paper-aligned outputs, but you can override the paths with environment variables if you want to show a different run.",
    )
    add_bullets(
        doc,
        [
            "MANIFEST defaults to data/manifests/meld_train.csv in the current bundle script.",
            "PREDICTIONS_CSV defaults to results/paper_aligned_meld_cv/cmt_min/fold_2/predictions_test.csv.",
            "METRICS_JSON defaults to results/paper_aligned_meld_cv/cmt_min/fold_2/metrics.json.",
            "ANALYSIS_DIR defaults to results/paper_aligned_meld_cv/cmt_min/fold_2/analysis_test.",
        ],
    )

    doc.add_heading("4. What the Demo Bundle Contains", level=1)
    add_bullets(
        doc,
        [
            "demo_examples.csv and demo_examples.md: the selected 3 to 5 review videos with their labels and predictions.",
            "metrics_summary.csv and metrics_summary.md: the headline evaluation numbers.",
            "confusion_matrix.csv: the full class-by-class confusion matrix.",
            "top_confusions.csv: the strongest error pairs.",
            "README.md: the suggested live demo sequence.",
        ],
    )
    add_para(
        doc,
        "This is the folder you can open during the review. It gives you everything needed for a controlled live explanation."
    )

    doc.add_heading("5. Which Test Videos Are Chosen and Why", level=1)
    examples = load_examples()
    if not examples.empty:
        add_para(
            doc,
            "The bundle chooses examples from the saved MELD predictions using a simple teacher-friendly logic: it starts with one correct, high-confidence example, then tries to add neutral-heavy mistakes, and then adds emotionally close confusions if they are available.",
        )
        rows = []
        for _, row in examples.iterrows():
            rows.append(
                [
                    str(row.get("sample_id", "")),
                    str(row.get("ground_truth", "")),
                    str(row.get("prediction", "")),
                    str(row.get("correct_or_wrong", "")),
                    f"{float(row.get('confidence', 0.0)):.4f}" if pd.notna(row.get("confidence", None)) else "",
                ]
            )
        add_table(doc, ["Sample ID", "Ground Truth", "Prediction", "Result", "Confidence"], rows)
        add_para(
            doc,
            "The selected examples are not random. They are chosen to tell a story: first a success case, then the common failure mode, and then the difficult emotionally similar cases. That helps the reviewer understand the model behavior quickly.",
        )
    else:
        add_para(doc, "The demo examples file was not found yet, so the actual example list is not available inside this document at generation time.")

    doc.add_heading("6. Why Those Videos Were Chosen", level=1)
    add_numbered(
        doc,
        [
            "One correct example shows that the model can make a valid prediction, not only mistakes.",
            "A neutral-heavy error example shows the most common bias in MELD emotion recognition.",
            "A neutral-versus-joy or sadness-versus-neutral example shows that the model struggles with close emotional neighbors.",
            "A fear-versus-surprise example shows that the model is still sensitive to emotionally similar classes.",
        ],
    )
    add_para(
        doc,
        "In review language, these cases are useful because they demonstrate both success and limitation. The reviewer should see that the system is learning signal, but that class imbalance and emotionally close categories still matter."
    )

    doc.add_heading("7. How to Present One Example Live", level=1)
    add_numbered(
        doc,
        [
            "Show the video name.",
            "Read the ground-truth label.",
            "Read the model prediction.",
            "State whether it was correct.",
            "Mention the confidence value if useful.",
            "Point to the transcript and the stored path if the reviewer asks where the input came from.",
        ],
    )
    add_para(
        doc,
        "A simple speaking pattern is: 'This utterance was predicted as neutral, but the true label is anger. That tells us the model is still using neutral as a fallback class when it is uncertain.'"
    )

    doc.add_heading("8. How to Explain the Metrics Summary", level=1)
    metrics = load_metrics()
    if not metrics.empty:
        rows = []
        for _, row in metrics.iterrows():
            value = row.get("value", "")
            if isinstance(value, (float, int)):
                value = f"{float(value):.4f}"
            rows.append([str(row.get("metric", "")), str(value)])
        add_table(doc, ["Metric", "Value"], rows)
    add_para(
        doc,
        "The metrics summary should be read as the headline scorecard for the whole test set. Accuracy shows overall correctness, weighted F1 reflects the majority-class-sensitive balance, macro F1 shows how well the model treats every class equally, and unweighted accuracy helps explain imbalance behavior.",
    )
    add_bullets(
        doc,
        [
            "Accuracy: overall percentage of correct predictions.",
            "Weighted F1: good for comparing performance when common classes dominate.",
            "Macro F1: class-balanced score that penalizes poor minority-class performance.",
            "Unweighted accuracy: another imbalance-sensitive summary that helps explain class bias.",
        ],
    )

    doc.add_heading("9. How to Explain the Error Analysis Panel", level=1)
    add_bullets(
        doc,
        [
            "Predicted-vs-actual table: shows the first few sample-level examples.",
            "Confusion matrix: shows the full class-by-class error structure.",
            "Top confusions: shows the largest off-diagonal mistakes in a compact form.",
        ],
    )
    add_para(
        doc,
        "This is where you show the reviewer that the model’s errors are systematic rather than random. For MELD, the important pattern is usually neutral being used too often as the fallback prediction."
    )

    doc.add_heading("10. Suggested Live Demo Flow", level=1)
    add_numbered(
        doc,
        [
            "Start with one correct example.",
            "Move to one neutral-heavy mistake.",
            "Then show one emotionally close confusion.",
            "Present the metrics summary.",
            "Finish with the confusion matrix and top confusions.",
        ],
    )
    add_para(
        doc,
        "This ordering works well because it tells a story from concrete example to overall statistics. The reviewer first sees that the model works, then sees where it fails, and finally sees that the failure pattern is measurable and understandable."
    )

    doc.add_heading("11. What You Should Say During the Demo", level=1)
    add_bullets(
        doc,
        [
            "This is the held-out MELD test output, not the training set.",
            "The model predicts emotion labels for each video-utterance example.",
            "The summary metrics describe the overall test behavior.",
            "The confusion matrix shows the main systematic mistakes.",
            "The result is good enough to demonstrate Phase 1 completion, but still leaves room for Phase 2 improvement.",
        ],
    )

    doc.add_heading("13. One-Minute Closing Answer", level=1)
    add_para(
        doc,
        "If you need to close the review quickly, say this: 'Phase 1 is now a complete and reproducible multimodal emotion-recognition baseline. I can demonstrate it on saved test examples or directly from raw mp4 clips, and the outputs include per-video predictions, confidence, top-3 probabilities, metrics, and confusion analysis. The model is not perfect, but its errors are systematic, which tells me the backbone is stable enough. That is why I am stopping Phase 1 refinement here and moving the next effort into Phase 2 courtroom-testimony adaptation.'",
    )
    add_bullets(
        doc,
        [
            "Phase 1 is complete enough to demonstrate end to end.",
            "The demo works from saved features and from raw mp4 input.",
            "The reviewer can inspect both correct and incorrect predictions.",
            "The remaining errors are systematic, not random.",
            "The next meaningful step is Phase 2 domain adaptation.",
        ],
    )

    doc.add_heading("12. Short Student Summary", level=1)
    add_para(
        doc,
        "The demo is a proof-of-work presentation. It shows that the Phase 1 pipeline can take test videos, produce emotion predictions, summarize the overall metrics, and reveal the main error patterns. The selected examples are chosen on purpose so the reviewer sees one success, one neutral-heavy failure, and one emotionally close confusion. That is the right way to explain both progress and limitation in a final review."
    )

    doc.save(OUTPUT)
    return doc


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
