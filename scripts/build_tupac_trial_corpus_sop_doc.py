from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_Tupac_Trial_Corpus_Student_SOP.docx"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(11.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
            styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.color.rgb = RGBColor(55, 55, 55)
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(11.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(3)


def add_code(doc: Document, text: str) -> None:
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9.2)
    doc.add_paragraph()


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def build_doc() -> Document:
    doc = Document()
    configure(doc)
    add_page_number(doc.sections[0].footer.paragraphs[0])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LegalMemoCMT\nTupac Trial Mini-Corpus SOP")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Student-level, manual Phase 2 workflow for downloading, processing, checking, and documenting testimony videos")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11.5)
    add_para(doc, "Version: 1.0 | Project: LegalMemoCMT | Source list: data/tupac_trial_urls.txt", italic=True)

    add_para(doc, "Purpose", bold=True)
    add_para(doc, "This SOP explains how to build a small Tupac trial testimony corpus using the existing LegalMemoCMT Phase 2 scripts. It is written so a student can run each command manually, inspect the files created at every stage, understand common errors, and explain why this case is useful for the project.")
    add_para(doc, "The workflow downloads YouTube video, English subtitles where available, and mono 16 kHz WAV audio. It then creates utterance rows, speaker-independent train/dev/test groups, weak labels for exploratory work, merged testimony turns, and review clips. It does not make legal findings and it does not infer that an emotional display proves truthfulness, deception, guilt, or innocence.")

    doc.add_heading("1. Why This Mini-Corpus Helps LegalMemoCMT", level=1)
    add_para(doc, "LegalMemoCMT is studying observable emotion and interaction patterns in legal speech. A useful adaptation corpus should therefore contain more than one speaking style. The proposed Tupac witness set is valuable because it supplies contrast: different levels of affect, challenge, medical explanation, and police procedure. That contrast can help the model learn courtroom interaction patterns instead of memorising one person’s voice or one emotional style.")
    add_para(doc, "Working project rationale:", bold=True)
    add_para(doc, "For LegalMemoCMT, Mob James + Reggie Wright full testimony + Lisa Gavin + Dean O'Kelley should provide a particularly useful mini-corpus: hostile or high-affect testimony, a challenged witness, professional medical testimony, and procedural police testimony. This diversity is more useful than collecting many witnesses who all speak in similar roles and emotional conditions.")
    add_table(doc, ["Proposed witness group", "Analytical value", "What the student should look for"], [
        ["Mob James", "Hostile/high-affect or confrontational interaction", "Raised intensity, interruptions, emphatic wording, visible frustration, rapid changes in engagement."],
        ["Reggie Wright full testimony", "Extended challenged-witness sequence", "Consistency across a long examination, responses to corrections, pauses, defensiveness, composure under questioning."],
        ["Lisa Gavin", "Professional medical testimony", "Technical explanation, controlled delivery, precision, calm affect, explanation under examination."],
        ["Dean O'Kelley", "Procedural police testimony", "Chronology, evidence handling, formal narration, restrained delivery, answers to procedural questions."],
    ])
    add_para(doc, "These are research categories, not legal conclusions. The labels must be checked against the actual video, title, transcript, and source context. If a URL does not contain the expected witness, keep the URL in the source audit but correct the annotation rather than forcing it into the proposed category.")

    doc.add_heading("2. What You Need Before Starting", level=1)
    add_bullets(doc, [
        "A terminal opened in the LegalMemoCMT repository root.",
        "Python 3. The repository prefers `.venv/bin/python` when that environment exists.",
        "`yt-dlp` available on PATH for YouTube downloads.",
        "`ffmpeg` available on PATH for audio extraction.",
        "A browser cookie source if YouTube requires authentication. The wrappers default to Chrome; change it if your browser is different.",
        "Enough storage for video, WAV audio, subtitles, and derived clips. WAV files are large because they are uncompressed.",
    ])
    add_code(doc, """cd /Users/rajeshpmu/Desktop/LegalMemoCMT
command -v python3
command -v yt-dlp
command -v ffmpeg
python3 --version
yt-dlp --version
ffmpeg -version | head -n 1""")
    add_para(doc, "If `yt-dlp` or `ffmpeg` is missing, stop and install the missing dependency before continuing. Do not replace the downloader with a browser screen recording: the pipeline needs stable files, timestamps, and a reproducible manifest.")

    doc.add_heading("3. Understand the Repository Paths", level=1)
    add_table(doc, ["Path", "Meaning"], [
        ["data/tupac_trial_urls.txt", "Input list: one YouTube URL per line. The current file contains eight URLs."],
        ["data/phase2/tupac/corpus/raw/", "Downloaded MP4, VTT subtitle, and WAV files, organised by title."],
        ["data/processed/phase2/tupac/", "CSV manifests and train/dev/test files produced by processing."],
        ["reports/phase2/tupac/", "JSON summaries and review CSVs."],
        ["phase2/build_clancy_corpus.py", "Generic downloader/extractor. The name is historical; the Tupac paths are supplied through command-line arguments."],
        ["phase2/build_clancy_utterance_manifest.py", "Parses VTT subtitle cues into utterance rows."],
        ["phase2/build_clancy_dataset_split.py", "Creates source-grouped train/dev/test partitions."],
        ["phase2/build_clancy_turn_manifest.py", "Merges adjacent utterances into longer turns."],
    ])
    add_para(doc, "Important naming note: several wrappers retain `clancy` in their filename because they were first written for the Clancy corpus. This SOP overrides their input and output paths. Every command below is Tupac-specific because it sets `OUTPUT_ROOT`, `RAW_ROOT`, `INPUT_CSV`, and related variables explicitly.")

    doc.add_heading("4. Define Tupac Variables Once", level=1)
    add_para(doc, "Run the following block in one terminal session. The variables remain available for later commands in that same session. If you close the terminal, run the block again.")
    add_code(doc, """cd /Users/rajeshpmu/Desktop/LegalMemoCMT

export TUPAC_ROOT="$PWD/data/phase2/tupac"
export TUPAC_RAW="$TUPAC_ROOT/corpus/raw"
export TUPAC_PROC="$PWD/data/processed/phase2/tupac"
export TUPAC_REPORTS="$PWD/reports/phase2/tupac"
export TUPAC_URLS="$PWD/data/tupac_trial_urls.txt"
export TUPAC_COOKIES_FROM_BROWSER="chrome"

mkdir -p "$TUPAC_RAW" "$TUPAC_PROC" "$TUPAC_REPORTS"
wc -l "$TUPAC_URLS"
sed -n '1,20p' "$TUPAC_URLS"
""")
    add_para(doc, "The expected URL count is currently eight. The command `wc -l` is a simple audit: if it prints a different count, investigate before downloading. Blank lines and comments are ignored by the Python downloader, but one URL should still be kept per line.")

    doc.add_heading("5. Optional Smoke Test", level=1)
    add_para(doc, "A smoke test processes only the first URL. Use it to confirm cookies, YouTube access, format selection, subtitle access, and ffmpeg extraction before downloading all sources.")
    add_code(doc, """COOKIES_FROM_BROWSER="$TUPAC_COOKIES_FROM_BROWSER" \\
URLS_FILE="$TUPAC_URLS" \\
OUTPUT_ROOT="$TUPAC_ROOT/corpus" \\
MANIFEST_CSV="$TUPAC_PROC/tupac_corpus_manifest_smoke.csv" \\
SUMMARY_JSON="$TUPAC_REPORTS/tupac_corpus_summary_smoke.json" \\
bash phase2/run_build_clancy_corpus.sh --limit 1 --skip-existing --verify""")
    add_para(doc, "Inspect the terminal output. A successful row should report a video, audio, and subtitle status. Subtitles may be missing even when the media download succeeds; that is a data-quality issue to record, not an automatic reason to claim the URL failed.")
    add_code(doc, """find "$TUPAC_ROOT/corpus/raw" -maxdepth 3 -type f | sort
sed -n '1,12p' "$TUPAC_PROC/tupac_corpus_manifest_smoke.csv"
cat "$TUPAC_REPORTS/tupac_corpus_summary_smoke.json"
""")

    doc.add_heading("6. Download the Complete Source Media", level=1)
    add_para(doc, "After the smoke test, run the complete eight-URL download. `--skip-existing` makes the command safe to rerun: existing files are reused instead of downloaded again. `--verify` runs file and ffprobe checks on available media.")
    add_code(doc, """COOKIES_FROM_BROWSER="$TUPAC_COOKIES_FROM_BROWSER" \\
URLS_FILE="$TUPAC_URLS" \\
OUTPUT_ROOT="$TUPAC_ROOT/corpus" \\
MANIFEST_CSV="$TUPAC_PROC/tupac_corpus_manifest.csv" \\
SUMMARY_JSON="$TUPAC_REPORTS/tupac_corpus_summary.json" \\
bash phase2/run_build_clancy_corpus.sh --skip-existing --verify""")
    add_para(doc, "The downloader first probes metadata, then downloads MP4 video, downloads English auto-subtitles in VTT format, and extracts mono 16 kHz PCM WAV audio. It writes one manifest row per input URL. The manifest is the audit trail connecting a URL to its local files and statuses.")
    add_code(doc, """sed -n '1,20p' "$TUPAC_PROC/tupac_corpus_manifest.csv"
cat "$TUPAC_REPORTS/tupac_corpus_summary.json"
find "$TUPAC_RAW" -type f | sort""")
    add_bullets(doc, [
        "`video_status=downloaded` or `exists` means an MP4 was found or created.",
        "`audio_status=downloaded` or `exists` means ffmpeg produced the WAV file.",
        "`subtitle_status=downloaded` or `exists` means a VTT file was found.",
        "`failed` or `missing` requires inspection of `notes` and the terminal error.",
        "`media_verified=YES` is meaningful only when verification was requested and both video and audio passed the script's checks.",
    ])

    doc.add_heading("7. Assign and Audit the Witness Categories", level=1)
    add_para(doc, "The URL list is a download list, not a verified witness annotation. Open the downloaded title information and review each source. Record which URL corresponds to Mob James, Reggie Wright, Lisa Gavin, Dean O'Kelley, or another speaker. Do not infer identity from a filename alone.")
    add_table(doc, ["Audit field", "Student action"], [
        ["source_url", "Copy the exact URL from the corpus manifest."],
        ["youtube_id", "Use the stable ID generated by yt-dlp."],
        ["title", "Check whether the title supports the expected witness and testimony segment."],
        ["witness_name", "Write the verified witness name, or `unknown` if not established."],
        ["witness_group", "Use one of `hostile_high_affect`, `challenged_witness`, `medical_professional`, `procedural_police`, or `other`."],
        ["verification_note", "Record how the identity was checked: title, transcript, courtroom caption, or manual video review."],
    ])
    add_para(doc, "Keep this annotation separate from the generated corpus manifest unless you are deliberately extending the schema. The existing utterance builder only requires `youtube_id`, `source_url`, `title`, `category`, and `priority`; missing category and priority values are allowed but will remain blank. For analysis, a separate witness annotation CSV is safer because rerunning the downloader will not overwrite your manual notes.")

    doc.add_heading("8. Build the Utterance Manifest", level=1)
    add_para(doc, "This stage reads VTT subtitle cues and converts them into rows with start time, end time, duration, text, and paths to the corresponding media. A subtitle cue is not necessarily a complete human sentence or a speaker turn; it is a timestamped transcript unit. That distinction matters later when interpreting emotion.")
    add_code(doc, """RAW_ROOT="$TUPAC_RAW" \\
SOURCE_MANIFEST="$TUPAC_PROC/tupac_corpus_manifest.csv" \\
OUTPUT_CSV="$TUPAC_PROC/tupac_utterance_manifest.csv" \\
SUMMARY_JSON="$TUPAC_REPORTS/tupac_utterance_summary.json" \\
bash phase2/run_build_clancy_utterance_manifest.sh""")
    add_code(doc, """sed -n '1,8p' "$TUPAC_PROC/tupac_utterance_manifest.csv"
cat "$TUPAC_REPORTS/tupac_utterance_summary.json"
find "$TUPAC_RAW" -name '*.vtt' -print | sort""")
    add_para(doc, "If the utterance CSV has zero rows, the most likely cause is missing VTT subtitles. Check the corpus manifest's `subtitle_path`, `subtitle_status`, and `notes`. A successful video download alone is not enough for this subtitle-based utterance stage.")

    doc.add_heading("9. Create Source-Grouped Train, Dev, and Test Splits", level=1)
    add_para(doc, "The split script groups rows by `youtube_id` so utterances from one video do not appear in both training and evaluation. This is essential: randomly splitting individual cues would let the model see nearly identical neighbouring speech from the same source in both sets, producing an over-optimistic result.")
    add_code(doc, """INPUT_CSV="$TUPAC_PROC/tupac_utterance_manifest.csv" \\
OUTPUT_CSV="$TUPAC_PROC/tupac_dataset_manifest.csv" \\
TRAIN_CSV="$TUPAC_PROC/train.csv" \\
DEV_CSV="$TUPAC_PROC/dev.csv" \\
TEST_CSV="$TUPAC_PROC/test.csv" \\
SUMMARY_JSON="$TUPAC_REPORTS/tupac_dataset_split_summary.json" \\
bash phase2/run_build_clancy_dataset_split.sh""")
    add_code(doc, """for f in "$TUPAC_PROC/tupac_dataset_manifest.csv" "$TUPAC_PROC/train.csv" "$TUPAC_PROC/dev.csv" "$TUPAC_PROC/test.csv"; do
  echo "--- $f"
  wc -l "$f"
  sed -n '1,3p' "$f"
done
cat "$TUPAC_REPORTS/tupac_dataset_split_summary.json"
""")
    add_para(doc, "With only a small number of videos, a three-way split may be unstable or impossible to interpret. Treat this split as a pipeline check and exploratory adaptation set, not as a statistically strong benchmark. Report the number of source videos in each split, not only the number of utterance rows.")

    doc.add_heading("10. Validate the Dataset Manifest", level=1)
    add_code(doc, """INPUT_CSV="$TUPAC_PROC/tupac_dataset_manifest.csv" \\
SUMMARY_JSON="$TUPAC_REPORTS/tupac_dataset_validation_summary.json" \\
bash phase2/run_build_clancy_dataset_validation.sh""")
    add_para(doc, "Read the validation output rather than assuming that a zero exit code makes the data scientifically correct. Check for empty text, missing audio, missing video, invalid time intervals, duplicate IDs, and rows marked unusable. Keep the validation JSON with the corpus for later reporting.")

    doc.add_heading("11. Add Weak Labels for Exploration", level=1)
    add_para(doc, "Weak labels are heuristic labels generated from transcript keywords. They are useful for inspecting coverage and producing a first review list, but they are not human ground truth and should not be described as expert annotation. They can be especially misleading in legal speech because a word such as ‘angry’ may be quoted, denied, or used by a lawyer rather than expressing the witness's emotion.")
    add_code(doc, """INPUT_CSV="$TUPAC_PROC/tupac_dataset_manifest.csv" \\
LABELS_CSV="$TUPAC_PROC/tupac_weak_labels.csv" \\
TRAINING_CSV="$TUPAC_PROC/tupac_training_manifest_weak.csv" \\
SUMMARY_JSON="$TUPAC_REPORTS/tupac_weak_label_summary.json" \\
REVIEW_CSV="$TUPAC_REPORTS/tupac_weak_label_review.csv" \\
bash phase2/run_build_clancy_weak_labels.sh""")
    add_code(doc, """sed -n '1,8p' "$TUPAC_PROC/tupac_training_manifest_weak.csv"
sed -n '1,12p' "$TUPAC_REPORTS/tupac_weak_label_review.csv"
cat "$TUPAC_REPORTS/tupac_weak_label_summary.json"
""")

    doc.add_heading("12. Check Training Readiness", level=1)
    add_code(doc, """INPUT_CSV="$TUPAC_PROC/tupac_training_manifest_weak.csv" \\
SUMMARY_JSON="$TUPAC_REPORTS/tupac_training_readiness_summary.json" \\
bash phase2/run_build_clancy_training_readiness.sh""")
    add_para(doc, "This check is a gate for the next stage. If it reports missing files or too few usable rows, fix or document the problem before creating clips. Do not silently remove failures from the CSV; the failed rows are part of the provenance record.")

    doc.add_heading("13. Build Merged Testimony Turns", level=1)
    add_para(doc, "Subtitle cues are short and may split one answer across several rows. The turn builder merges compatible adjacent cues into longer turn records. A turn is more useful for courtroom interaction analysis because it provides more local context than a single subtitle fragment.")
    add_code(doc, """INPUT_CSV="$TUPAC_PROC/tupac_utterance_manifest.csv" \\
OUTPUT_CSV="$TUPAC_PROC/tupac_turn_manifest.csv" \\
SUMMARY_JSON="$TUPAC_REPORTS/tupac_turn_summary.json" \\
ISSUES_CSV="$TUPAC_REPORTS/tupac_turn_issues.csv" \\
REJECTION_CSV="$TUPAC_PROC/tupac_turn_rejection_manifest.csv" \\
bash phase2/run_build_clancy_turn_manifest.sh""")
    add_code(doc, """sed -n '1,6p' "$TUPAC_PROC/tupac_turn_manifest.csv"
cat "$TUPAC_REPORTS/tupac_turn_summary.json"
sed -n '1,12p' "$TUPAC_REPORTS/tupac_turn_issues.csv"
""")
    add_para(doc, "The rejection CSV is optional. If you create it after manual review, keep one row per rejected turn with a reason. The turn builder can apply those exclusions on a rerun, making the review decision reproducible.")

    doc.add_heading("14. Create Turn and Utterance Clips", level=1)
    add_para(doc, "Clips are derived files for listening and visual review. They should never replace the original MP4, WAV, VTT, or manifests. Build turn clips first because they are the most natural unit for testimony analysis.")
    add_code(doc, """INPUT_CSV="$TUPAC_PROC/tupac_turn_manifest.csv" \\
OUTPUT_ROOT="$TUPAC_ROOT/turn_clips" \\
bash phase2/run_build_clancy_turn_clips.sh --skip-existing""")
    add_code(doc, """INPUT_CSV="$TUPAC_PROC/tupac_training_manifest_weak.csv" \\
OUTPUT_ROOT="$TUPAC_ROOT/utterance_clips" \\
bash phase2/run_build_clancy_utterance_clips.sh --skip-existing""")
    add_code(doc, """find "$TUPAC_ROOT/turn_clips" -type f | sort | sed -n '1,20p'
find "$TUPAC_ROOT/utterance_clips" -type f | sort | sed -n '1,20p'""")
    add_para(doc, "Listen to a sample from every witness group and from every split. Check that the clip starts and ends at sensible times, that the audio is audible, and that the visible speaker matches the transcript context. This is where timestamp problems become obvious.")

    doc.add_heading("15. Make a Long-Turn Review List", level=1)
    add_code(doc, """INPUT_CSV="$TUPAC_PROC/tupac_turn_manifest.csv" \\
OUTPUT_CSV="$TUPAC_REPORTS/tupac_long_turn_review.csv" \\
bash phase2/run_build_clancy_long_turn_review.sh""")
    add_para(doc, "Long turns are useful for checking extended testimony, but they also have a higher risk of containing multiple ideas, interruptions, or subtitle alignment errors. Review the longest rows manually and note whether they should remain one turn or be split for analysis.")

    doc.add_heading("16. Manual Quality-Control Checklist", level=1)
    add_numbered(doc, [
        "Confirm the URL count and preserve the original URL file unchanged.",
        "Confirm each downloaded row has a stable YouTube ID and source URL.",
        "Open at least one MP4, WAV, and VTT file from each proposed witness group.",
        "Check that video duration, audio duration, and subtitle timestamps are plausible.",
        "Record missing subtitles separately from failed video downloads.",
        "Verify that train, dev, and test do not share a YouTube ID.",
        "Review weak labels as suggestions only; do not treat them as ground truth.",
        "Listen to clips from high-affect, challenged, medical, and procedural examples.",
        "Keep a short annotation log explaining identity, role, source quality, and exclusions.",
        "Save the summary JSON files with the corpus so another student can reproduce the run.",
    ])

    doc.add_heading("17. Common Problems and Safe Responses", level=1)
    add_table(doc, ["Problem", "Likely cause", "Response"], [
        ["YouTube access error", "Cookies, rate limit, age restriction, or changed site behaviour", "Check the browser name, rerun one URL, and inspect the yt-dlp error. Do not delete the manifest."],
        ["Video exists but subtitle is missing", "No English auto-caption or captions disabled", "Keep the source in the audit, mark it subtitle-missing, and do not expect the VTT utterance stage to use it."],
        ["No utterance rows", "No VTT files found under the raw root", "Check `subtitle_path` and the `RAW_ROOT` override first."],
        ["Audio extraction fails", "Corrupt media, unsupported format, or ffmpeg issue", "Run ffprobe on the specific MP4 and preserve the failure note."],
        ["Split has too few groups", "Mini-corpus has few source videos", "Use the split for pipeline validation and report the limitation; do not claim robust generalisation."],
        ["Wrong files appear in output", "A Clancy default was not overridden", "Stop, inspect the command variables, and use the Tupac paths shown in this SOP."],
    ])

    doc.add_heading("18. Expected Output Inventory", level=1)
    add_table(doc, ["Artifact", "Why it matters"], [
        ["tupac_corpus_manifest.csv", "URL-to-media provenance and download status."],
        ["tupac_corpus_summary.json", "Counts of successes, failures, subtitles, audio, and duration."],
        ["tupac_utterance_manifest.csv", "Timestamped subtitle cue rows."],
        ["tupac_dataset_manifest.csv, train.csv, dev.csv, test.csv", "Source-grouped split manifests."],
        ["tupac_weak_labels.csv and tupac_training_manifest_weak.csv", "Exploratory heuristic labels and training-ready row view."],
        ["tupac_turn_manifest.csv", "Merged testimony turns for courtroom-context analysis."],
        ["tupac_turn_issues.csv and tupac_long_turn_review.csv", "Quality-control queues."],
        ["turn_clips/ and utterance_clips/", "Reviewable derived media clips."],
    ])

    doc.add_heading("19. How to Explain the Project Contribution", level=1)
    add_para(doc, "In a report or supervision meeting, describe this corpus as a controlled domain-adaptation and stress-test resource. Its contribution is not its size. Its contribution is the deliberate diversity of witness roles and interaction conditions.")
    add_bullets(doc, [
        "It gives LegalMemoCMT courtroom speech that is less homogeneous than a single witness collection.",
        "It tests whether learned audio, video, and text features remain useful when speech is formal, adversarial, technical, or procedural.",
        "It creates a natural comparison between high-affect and restrained testimony without claiming that either style reveals credibility.",
        "It supports turn-level analysis because extended testimony contains changes over time, responses to challenge, and transitions between explanation and defensiveness.",
        "It provides a small, manually inspectable corpus that can expose alignment and generalisation problems before a larger legal-domain expansion.",
    ])
    add_para(doc, "A careful limitation statement is equally important: YouTube captions may be noisy, speaker labels may require manual verification, the source selection is not a random sample of legal testimony, and the small number of videos cannot establish broad performance claims. The mini-corpus is best presented as a targeted adaptation and qualitative validation set.")

    doc.add_heading("20. Recommended Run Order", level=1)
    add_numbered(doc, [
        "Check dependencies and inspect `data/tupac_trial_urls.txt`.",
        "Export the Tupac paths and create the output directories.",
        "Run the one-URL smoke test with `--verify`.",
        "Run the complete downloader with `--skip-existing --verify`.",
        "Audit the corpus manifest and summary JSON.",
        "Build the VTT utterance manifest.",
        "Create source-grouped train/dev/test files.",
        "Run dataset validation.",
        "Build weak labels and inspect the review CSV.",
        "Run the training-readiness check.",
        "Build merged turns and inspect turn issues.",
        "Create turn and utterance clips.",
        "Review clips across the four proposed witness groups and record annotations.",
    ])
    add_para(doc, "The core principle is to preserve the chain from source URL to downloaded media to subtitle cue to split row to turn clip. If a student can follow that chain for one example from each witness group, the corpus is understandable, auditable, and useful for the next LegalMemoCMT experiment.")

    return doc


if __name__ == "__main__":
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_doc().save(OUTPUT)
    print(f"Wrote {OUTPUT}")
