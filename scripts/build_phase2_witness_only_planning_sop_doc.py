from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_Phase2_Witness_Only_Planning_SOP.docx"
VALIDATED_INPUT = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated" / "legalmeld_metadata_validated.csv"
WITNESS_OUTPUT_DIR = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated" / "witness_only_rows"
DISCOVERY_INPUT = ROOT / "data" / "processed" / "phase2" / "paired_hearing_witness_discovery.csv"
CONTROLLED_VALIDATION_INPUT = WITNESS_OUTPUT_DIR / "witness_controlled_validation_subset.csv"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def count_rows(path: Path) -> int:
    import csv

    if not path.exists():
        return 0
    with path.open(newline="", encoding="utf-8-sig") as f:
        return sum(1 for _ in csv.DictReader(f))


def sample_hearing_plan_rows(limit: int = 8) -> list[list[str]]:
    import csv

    path = WITNESS_OUTPUT_DIR / "witness_hearing_plan.csv"
    if not path.exists():
        return []
    rows: list[list[str]] = []
    with path.open(newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(
                [
                    row.get("hearing_id", ""),
                    row.get("case_number", ""),
                    row.get("hearing_date", ""),
                    row.get("witness_name_or_code", ""),
                    row.get("discovery_row_count", ""),
                ]
            )
            if len(rows) >= limit:
                break
    return rows


def sample_controlled_validation_rows(limit: int = 10) -> list[list[str]]:
    import csv

    path = CONTROLLED_VALIDATION_INPUT
    if not path.exists():
        return []
    rows: list[list[str]] = []
    with path.open(newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(
                [
                    row.get("validation_track", ""),
                    row.get("hearing_id", ""),
                    row.get("hearing_date", ""),
                    row.get("witness_name_or_code", ""),
                    row.get("controlled_validation_score", ""),
                ]
            )
            if len(rows) >= limit:
                break
    return rows


def load_controlled_validation_details() -> list[dict[str, str]]:
    import csv

    path = CONTROLLED_VALIDATION_INPUT
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def build_doc() -> Document:
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LegalMemoCMT Phase 2 Witness-Only Planning SOP")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Student SOP for rerunning Phase 2 with only witness utterance rows while keeping the tribunal bootstrap pipeline grounded and reproducible."
    )
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12.5)

    add_para(
        doc,
        "This SOP is for the witness-only planning layer of Phase 2. The goal is not to rebuild the full tribunal dataset again. The goal is to isolate the witness utterances from the validated LegalMELD export, inspect them clearly, and use them as the controlled subset for later witness-centric analysis. That keeps the pipeline aligned with the project focus and avoids mixing in judge or prosecutor turns when the next step needs testimony only.",
    )
    add_para(
        doc,
        "The plan now has two levels. The first level is the immediately usable witness-only utterance subset from the validated LegalMELD export. The second level is a broader hearing-level discovery plan built from the witness discovery layer, which shows additional witness-bearing hearings that are applicable but not yet converted into utterance rows. That broader plan keeps the scope honest: you can talk about more hearings without pretending they are already in the final training set.",
    )

    doc.add_heading("1. What This SOP Keeps And What It Ignores", level=1)
    add_bullets(
        doc,
        [
            "Keep the validated utterance-level export as the source of truth.",
            "Keep only rows where `speaker_role = Witness` for this planning path.",
            "Keep the training-use buckets so usable, review, and reject rows remain visible.",
            "Ignore judge, prosecutor, and defence rows for the witness-only planning set.",
            "Do not treat witness-only filtering as new data collection; it is a selection step on already validated rows.",
        ],
    )

    doc.add_heading("2. Why Witness-Only Filtering Is Useful", level=1)
    add_bullets(
        doc,
        [
            "It matches the project’s testimony focus more closely than a mixed-role dataset.",
            "It reduces noise from argument-heavy or procedural speech.",
            "It makes manual review easier because the rows already correspond to testimony turns.",
            "It gives a cleaner base for later emotion, credibility, and witness-behavior analysis.",
            "It keeps the corpus planning step honest because the selection is traceable back to one field: speaker_role.",
        ],
    )

    doc.add_heading("3. Input And Output Files", level=1)
    add_table(
        doc,
        ["File", "Role in the witness-only plan", "What to inspect"],
        [
            [
                "data/processed/phase2/legalmeld_validated/legalmeld_metadata_validated.csv",
                "Main validated export",
                "Check `speaker_role`, `quality_tier`, `alignment_confidence`, `split`, and `manual_review_required`.",
            ],
            [
                "data/processed/phase2/legalmeld_validated/witness_only_rows/witness_hearing_plan.csv",
                "Broader hearing-level witness discovery plan",
                "Check `hearing_id`, `witness_name_or_code`, and the discovery row count per hearing.",
            ],
            [
                "data/processed/phase2/legalmeld_validated/witness_only_rows/witness_rows.csv",
                "Witness-only master subset",
                "Confirm that every row is a witness utterance and that the row count looks reasonable.",
            ],
            [
                "data/processed/phase2/legalmeld_validated/witness_only_rows/witness_rows_usable.csv",
                "High-trust subset",
                "These are the witness rows that are the most practical for downstream use.",
            ],
            [
                "data/processed/phase2/legalmeld_validated/witness_only_rows/witness_rows_review.csv",
                "Manual-review subset",
                "These rows are witness turns but still need inspection before training.",
            ],
            [
                "data/processed/phase2/legalmeld_validated/witness_only_rows/witness_rows_reject.csv",
                "Rejected subset",
                "These witness rows should not be used for model training yet.",
            ],
            [
                "data/processed/phase2/legalmeld_validated/witness_only_rows/witness_rows_summary.json",
                "Summary report",
                "Check row counts, distinct witnesses, hearing coverage, and split distribution.",
            ],
        ],
    )

    doc.add_heading("4. Script And Wrapper", level=1)
    add_table(
        doc,
        ["Artifact", "What it does", "Why it matters"],
        [
            [
                "phase2/filter_legalmeld_rows_by_witness_role.py",
                "Filters the validated export down to witness rows and writes the witness-only buckets.",
                "This is the real implementation of the witness-only selection step.",
            ],
            [
                "phase2/run_filter_legalmeld_rows_by_witness_role.sh",
                "Shell wrapper for rerunning the witness-only filter with a single command.",
                "This is the easiest script to show during a guidance call.",
            ],
            [
                "phase2/filter_legalmeld_rows_by_use.py",
                "General-quality classifier reused by the witness-only filter.",
                "This keeps the witness-only plan aligned with the existing quality logic.",
            ],
        ],
    )

    doc.add_heading("5. Exact Run Order", level=1)
    add_numbered(
        doc,
        [
            "Open the validated export and confirm that the rows you want to keep are actually witness rows.",
            "Run `bash phase2/run_filter_legalmeld_rows_by_witness_role.sh`.",
            "Inspect `witness_rows.csv` to confirm that all rows have `speaker_role = Witness`.",
            "Inspect `witness_rows_usable.csv`, `witness_rows_review.csv`, and `witness_rows_reject.csv` to see how the witness rows are distributed by quality.",
            "Inspect `witness_rows_summary.json` to confirm the row counts, distinct witness count, and hearing coverage.",
            "Only after that, use the witness-only subset for the next planning or analysis step.",
        ],
    )

    doc.add_heading("6. How To Explain The Filtering Logic To A Guide", level=1)
    add_bullets(
        doc,
        [
            "The filter is not inventing new data.",
            "It is selecting only the utterances that are already marked as witness speech.",
            "The hearing plan is not claiming those discoveries are final utterance rows; it is a planning view built from the discovery manifest.",
            "The quality buckets are still preserved so the rows can be split into usable, review, and reject groups.",
            "The result is a smaller but more relevant subset for witness-focused work.",
            "The pipeline stays reproducible because the same input CSV and the same script always give the same witness-only selection.",
        ],
    )

    doc.add_heading("7. Student-Level Interpretation", level=1)
    add_para(
        doc,
        "At a student level, the simplest way to understand this stage is: first the pipeline builds a mixed courtroom dataset, then the witness-only filter removes everything except testimony utterances. That is useful because the project’s final emotion-analysis use case is more meaningful on witness speech than on lawyer argument or judge management speech. The witness-only plan therefore acts like a lens that narrows the broader tribunal dataset into the part that matters most for later modeling.",
    )
    add_para(
        doc,
        "This does not mean the other roles are unimportant. It only means they are outside this specific planning path. If the project later needs courtroom-procedure analysis, the mixed-role rows can still be reused from the validated export.",
    )
    add_para(
        doc,
        "From the broader discovery layer, the plan currently extends to many more witness-bearing hearings in ICTR-98-41, not just the four validated utterance hearings. The key point is that these extra hearings live in the planning layer, where they help you explain the wider corpus direction, but they are still separate from the immediately usable utterance subset.",
    )

    doc.add_heading("8. Current Scope Snapshot", level=1)
    add_table(
        doc,
        ["Layer", "Current count", "Meaning"],
        [
            ["Validated witness utterance rows", str(count_rows(WITNESS_OUTPUT_DIR / "witness_rows.csv")), "Rows that can already be reviewed as witness utterances."],
            ["Broader hearing discovery rows", str(count_rows(DISCOVERY_INPUT)), "All hearing-level discovery rows before witness-only grouping."],
            ["Broader hearing plan rows", str(count_rows(WITNESS_OUTPUT_DIR / "witness_hearing_plan.csv")), "Grouped hearings that contain actual witness testimony in the discovery layer."],
        ],
    )
    if sample_hearing_plan_rows():
        doc.add_paragraph("Example hearings in the broader discovery plan:")
        add_table(
            doc,
            ["Hearing ID", "Case number", "Date", "Witness name/code", "Discovery rows"],
            sample_hearing_plan_rows(),
        )

    doc.add_heading("9. Controlled Validation Pass", level=1)
    add_para(
        doc,
        "This is the small promoted subset that sits between the full hearing discovery plan and any future download or validation work. I selected it manually from the broader hearing plan after checking the hearing-by-hearing evidence. The purpose is not to claim that these hearings are already fully ready for training; the purpose is to create a tighter validation group that can be inspected more carefully before any scale-up.",
    )
    add_para(
        doc,
        "In simple terms, the first four rows are the hearings already validated in the witness utterance subset. The remaining six rows are the hearings I manually promoted from the broader discovery plan because they show strong testimony signals and good evidence coverage. That makes this a controlled review set, not a final corpus claim.",
    )
    if sample_controlled_validation_rows():
        add_table(
            doc,
            ["Track", "Hearing ID", "Date", "Witness name/code", "Score"],
            sample_controlled_validation_rows(),
        )
    add_bullets(
        doc,
        [
            "Manual note: this step is intentionally conservative and traceable.",
            "Manual note: the promoted hearings are still part of a planning subset, not a download-ready corpus.",
            "Manual note: the score is only a prioritization aid, not a truth label.",
        ],
    )

    doc.add_heading("10. What This Means", level=1)
    add_para(
        doc,
        "What this means is that Phase 2 now has a small, explainable progression inside the witness-only path. You can show the fully validated witness utterance subset first, then show a broader hearing discovery plan, and then show a manually promoted controlled validation subset. That sequence demonstrates disciplined corpus growth: first validation, then broader discovery, then cautious promotion. It is a better research story than pretending all witness-bearing hearings are equally ready at the same time.",
    )
    add_para(
        doc,
        "For the project, this is useful because it keeps the dataset grounded. For the mentor, it is easy to explain: 'I inspected the hearing plan hearing by hearing, kept the four already-validated witness hearings as anchors, then manually promoted six additional hearings into a controlled validation set. I am not using this as final training data yet; I am using it to inspect whether the witness-only path remains clean as I expand the plan.'"
    )

    doc.add_heading("11. Hearing-by-Hearing Walkthrough", level=1)
    add_para(
        doc,
        "This section is the clearest way to talk through the controlled validation subset in a guidance call. You can present each hearing as one evidence-backed planning decision. The key student-level message is that every row here is still tied to the same tribunal family, but the reason for keeping it differs: some rows are anchors because they already produced usable witness utterances, and some rows are promoted because they looked strong enough in the broader hearing discovery plan to justify a closer look.",
    )
    add_para(
        doc,
        "A good way to explain the table is: anchor hearings are already proven at the utterance level, while promoted hearings are candidates from the broader discovery layer that now deserve controlled inspection. That is a healthy research workflow because it separates proof from priority.",
    )
    details = load_controlled_validation_details()
    if details:
        add_table(
            doc,
            ["Track", "Hearing ID", "Date", "Witness", "Why this hearing is here"],
            [
                [
                    row.get("validation_track", ""),
                    row.get("hearing_id", ""),
                    row.get("hearing_date", ""),
                    row.get("witness_name_or_code", ""),
                    row.get("validation_reason", ""),
                ]
                for row in details
            ],
        )
        for row in details:
            track = row.get("validation_track", "")
            hearing_id = row.get("hearing_id", "")
            witness = row.get("witness_name_or_code", "")
            reason = row.get("validation_reason", "")
            note = row.get("manual_note", "")
            score = row.get("controlled_validation_score", "")
            add_para(
                doc,
                (
                    f"Hearing {hearing_id} ({row.get('hearing_date', '')}) is a {track} row with witness label {witness}. "
                    f"The score shown in the planning subset is {score}. In student terms, this means the hearing already looks structurally strong enough to matter in the corpus story. "
                    f"The reason recorded in the subset is: {reason}. Manual note: {note}. "
                    "When you explain it aloud, say whether it is an anchor or promoted hearing, then state that the hearing remains in a controlled validation group rather than a final training release."
                ),
            )

    doc.add_heading("12. Compact Comparison Table", level=1)
    add_para(
        doc,
        "This compact table is the quickest version to show on a slide or read aloud in a guidance call. It separates the four anchor hearings, which are already validated, from the six manually promoted hearings, which are only in the controlled validation pass. The important student-level message is that the anchor rows prove the pipeline, while the promoted rows test whether the broader hearing discovery plan is still clean enough to expand carefully.",
    )
    add_table(
        doc,
        ["Group", "Hearing ID", "Date", "Witness", "Reason selected"],
        [
            ["Anchor", "hear_2ef83c852251d65c", "21/06/2004", "ANTIPAS NYANJWA", "Already validated in the utterance-level witness subset."],
            ["Anchor", "hear_7b2686bf1e5608e9", "03/05/2005", "DM190 | EMMANUEL NERETSE", "Already validated in the utterance-level witness subset."],
            ["Anchor", "hear_8280237faba9d96f", "16/09/2004", "FILIP REYNTJENS", "Already validated in the utterance-level witness subset."],
            ["Anchor", "hear_f5d485391c1d04cc", "19/01/2004", "ROMÉO DALLAIRE", "Already validated in the utterance-level witness subset."],
            ["Promoted", "hear_dea210cdb4c728e0", "03/02/2004", "BRENT BEARDSLEY | Exanination-in-chief by Mr. White", "Manually promoted from the broader hearing discovery plan for controlled validation."],
            ["Promoted", "hear_8a80539e19e2df44", "05/07/2005", "DK120 | MATHIEU NGIRUMPATSE", "Manually promoted from the broader hearing discovery plan for controlled validation."],
            ["Promoted", "hear_da77e01d6076db4b", "18/06/2003", "OMAR SERUSHAGO", "Manually promoted from the broader hearing discovery plan for controlled validation."],
            ["Promoted", "hear_b3fafce46640756e", "21/09/2006", "ALOYS NTABAKUZE", "Manually promoted from the broader hearing discovery plan for controlled validation."],
            ["Promoted", "hear_23dd766aab93f457", "26/01/2004", "ROMÉO DALLAIRE", "Manually promoted from the broader hearing discovery plan for controlled validation."],
            ["Promoted", "hear_ebd43a70140ec251", "27/01/2004", "ROMÉO DALLAIRE", "Manually promoted from the broader hearing discovery plan for controlled validation."],
        ],
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
