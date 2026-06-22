from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path("/Users/rajeshpmu/Desktop/LegalMemoCMT")
DOC_PATH = ROOT / "implementation_docments" / "code" / "LegalMemoCMT_Phase1_Code_Explanation.docx"


def add_heading_text(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def main() -> None:
    doc = Document(str(DOC_PATH))
    if any("BidirectionalCrossAttentionCMT" in p.text for p in doc.paragraphs):
        print("Fusion explainer already present.")
        return

    doc.add_heading("6. BidirectionalCrossAttentionCMT and CrossModalFusion", level=2)
    doc.add_paragraph(
        "The repository uses two different fusion blocks. CrossModalFusion is the older three-modality Transformer fusion block used in the legacy path. BidirectionalCrossAttentionCMT is the custom paper-style fusion block used for the pretrained text and audio path. Both are custom modules written in this project; neither is a pretrained backbone."
    )

    add_table(
        doc,
        ["Module", "Input structure", "Core operation", "Main use in the project"],
        [
            [
                "CrossModalFusion",
                "Three modality vectors: text, audio, video",
                "Stacks modality tokens, adds learned modality embeddings, applies a Transformer encoder, then pools the three-token sequence.",
                "Legacy baseline and video-enabled hybrid branch.",
            ],
            [
                "BidirectionalCrossAttentionCMT",
                "Text token sequence and audio token sequence",
                "Lets text attend to audio and audio attend to text using two MultiheadAttention blocks, applies residual connections and layer norm, concatenates the updated sequences, then pools with a mask-aware rule.",
                "Paper-style pretrained path for the MemoCMT-like implementation.",
            ],
        ],
    )

    doc.add_paragraph(
        "BidirectionalCrossAttentionCMT is the closer implementation of the MemoCMT-style idea because it preserves token sequences longer and explicitly learns interaction in both directions. Text does not just sit beside audio; text queries audio and audio queries text."
    )
    doc.add_paragraph(
        "The module is implemented with standard PyTorch primitives. It creates two nn.MultiheadAttention layers, one for text-to-audio and one for audio-to-text. Each branch uses the other modality as key and value, while the current modality acts as the query. That is the key mechanism that allows the model to learn cross-modal dependency rather than simple concatenation."
    )

    add_heading_text(doc, "6.1 Step-by-Step Implementation", level=3)
    add_bullets(
        doc,
        [
            "The text encoder produces a sequence of contextual token embeddings.",
            "The audio encoder produces a sequence of acoustic token embeddings.",
            "Attention masks mark which positions are valid and which positions are padding.",
            "Text queries audio through text_to_audio and receives an updated text context.",
            "Audio queries text through audio_to_text and receives an updated audio context.",
            "Residual connections and LayerNorm stabilize the updated sequences.",
            "The two updated sequences are concatenated into one fused sequence.",
            "The pooled output is selected with cls, mean, max, or min and then passed to the classifier.",
        ],
    )

    add_heading_text(doc, "6.2 What Masking Does Here", level=3)
    doc.add_paragraph(
        "Masking is essential because text and audio sequences are padded to a common length in batching. The module converts the masks to boolean key_padding_mask values for MultiheadAttention and uses the same masks again in the pooling step. That way, padded tokens do not contribute to the learned attention pattern or the final pooled representation."
    )
    doc.add_paragraph(
        "For mean pooling, the code sums only valid positions and divides by the number of valid tokens. For max pooling, invalid positions are filled with negative infinity before taking the maximum. For min pooling, invalid positions are filled with positive infinity before taking the minimum. CLS simply takes the first fused token, which is treated as the summary position in this implementation."
    )

    add_heading_text(doc, "6.3 Why This Is Different from CrossModalFusion", level=3)
    add_bullets(
        doc,
        [
            "CrossModalFusion works at the modality-vector level: text, audio, and video are first reduced to one vector each, then stacked as three tokens.",
            "BidirectionalCrossAttentionCMT works at the token-sequence level: text and audio stay as sequences long enough for one modality to attend to the other.",
            "CrossModalFusion is simpler and is useful for the legacy branch and for hybrid video-enabled experiments.",
            "BidirectionalCrossAttentionCMT is more faithful to the MemoCMT-style design because it explicitly models interaction between text and speech before pooling.",
        ],
    )

    doc.add_paragraph(
        "A useful way to explain the design in a thesis is to say that CrossModalFusion summarizes modality evidence after each modality has already been compressed, while BidirectionalCrossAttentionCMT performs evidence exchange before the final summarization step."
    )

    add_heading_text(doc, "6.4 Implementation Summary", level=3)
    doc.add_paragraph(
        "In short, BidirectionalCrossAttentionCMT is a custom cross-modal fusion layer that uses bidirectional attention, residual updates, masking, and configurable pooling to build a paper-style text-audio representation. It is the project’s main pretrained fusion block for the MemoCMT-like path."
    )

    doc.save(str(DOC_PATH))
    print("Updated phase1 code explanation with fusion-module explainer.")


if __name__ == "__main__":
    main()
