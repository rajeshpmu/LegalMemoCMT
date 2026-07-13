#!/usr/bin/env python3
from __future__ import annotations

import json
import shutil
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "implementation_docments" / "phase1_project_report_pes"
FIG_DIR = OUT_DIR / "figures"
DOCX_OUT = OUT_DIR / "LegalMemoCMT_Phase1_Project_Report.docx"
TEX_OUT = OUT_DIR / "LegalMemoCMT_Phase1_Project_Report.tex"
PDF_OUT = OUT_DIR / "LegalMemoCMT_Phase1_Project_Report.pdf"
IEEE_DOCX_OUT = ROOT / "implementation_docments" / "LegalMemoCMT_Phase1_IEEE_Style_Project_Report.docx"
IEEE_TEX_OUT = ROOT / "implementation_docments" / "LegalMemoCMT_Phase1_IEEE_Style_Project_Report.tex"
IEEE_PDF_OUT = ROOT / "implementation_docments" / "LegalMemoCMT_Phase1_IEEE_Style_Project_Report.pdf"
LOGO_SRC = Path("/Users/rajeshpmu/Downloads/PES_Mtech_Project_report_Sample_Latex/PES_Mtech_Project_report_latex_format/pes_logo.png")
LOGO_DST = OUT_DIR / "pes_logo.png"

ABSTRACT_TEXT = (
    "This Phase 1 project establishes LegalMemoCMT as a cross-modal emotional cue analysis framework "
    "and validates the core implementation on MELD before extending it toward Indian courtroom testimony "
    "use cases. The project reproduces a paper-aligned conversational baseline with pretrained BERT for "
    "text, HuBERT for audio, cross-modal transformer fusion, and MIN pooling, then adds a ViT-based "
    "facial-cue path from raw mp4 clips to cached embeddings. The title reflects the broader multilingual "
    "courtroom direction, while Phase 1 focuses on a reproducible benchmark result and a transparent visual "
    "support path. Across five MELD folds, the baseline achieves mean accuracy of 0.6247, mean weighted F1 "
    "of 0.6195, and mean macro F1 of 0.4395. The facial-cue branch supports boundary-case analysis, error "
    "inspection, and future courtroom-testimony adaptation, but it does not replace the conversational "
    "baseline as the main Phase 1 result."
)

CONTENTS_ITEMS = [
    "Certificate",
    "Declaration",
    "Acknowledgement",
    "Abstract",
    "Contents",
    "1. Introduction",
    "2. Related Work",
    "3. Materials and Methods",
    "4. Results",
    "5. Discussion",
    "6. Limitations",
    "7. Conclusion",
    "8. Project Objectives and Scope",
    "9. Dataset Organization, Label Space, and Fold Rules",
    "10. Baseline Architecture and Modality Roles",
    "11. Training and Evaluation Protocol",
    "12. ViT Facial-Cue Extension and Video Caching",
    "13. Comparative Results and Error Analysis",
    "14. Reproducibility, Scripts, and Outputs",
    "15. Limitations and Phase 2 Direction",
    "16. Background",
    "17. Problem Statement",
    "18. Literature Survey",
    "19. System Requirements Specification",
    "20. Proposed Methodology",
    "21. Implementation Details",
    "22. Intermediate Results and Discussion",
    "23. Conclusions and Future Work",
    "24. Comparative Metric Summary",
    "Appendix A. Script List and Command Sequence",
    "References",
]


def add_word_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_sep)

    placeholder = paragraph.add_run("Right-click and update this field in Word to refresh the table of contents.")
    placeholder.italic = True

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def ensure_dirs() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if LOGO_SRC.exists():
        shutil.copy2(LOGO_SRC, LOGO_DST)


def append_front_matter_docx(doc: Document) -> None:
    for title, body in [
        ("Certificate", "This is to certify that the Phase 1 project entitled LegalMemoCMT: An Explainable Multilingual Multimodal Emotional Cue Analysis Framework for Indian Courtroom Testimony Using Cross-Modal Transformers is a bona fide work carried out under the guidance of Ramesh Prakash Guledgudd in partial fulfilment of the requirements for the award of the degree of Master of Technology in AI and Machine Learning."),
        ("Declaration", "I hereby declare that the Phase 1 project work presented in this report is carried out as part of the LegalMemoCMT project under the prescribed academic guidance and that the submitted content reflects the implementation, analysis, and evaluation carried out in the repository."),
        ("Acknowledgement", "I acknowledge the guidance, review support, and academic feedback received during the Phase 1 project. The work also benefited from the available benchmark datasets, pretrained encoders, and the project scripts that made the implementation traceable from raw input to final metrics."),
    ]:
        doc.add_heading(title, level=1)
        doc.add_paragraph(body)
        doc.add_page_break()


def render_mermaid(code: str, base_name: str) -> tuple[Path, Path]:
    mmd = FIG_DIR / f"{base_name}.mmd"
    svg = FIG_DIR / f"{base_name}.svg"
    png = FIG_DIR / f"{base_name}.png"
    mmd.write_text(code, encoding="utf-8")
    import subprocess

    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(mmd), "-o", str(svg), "-b", "white"],
        check=True,
        cwd=str(ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(mmd), "-o", str(png), "-b", "white"],
        check=True,
        cwd=str(ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    return svg, png


def build_figures() -> dict[str, Path]:
    figures: dict[str, Path] = {}
    _, figures["pipeline_png"] = render_mermaid(
        """flowchart LR
  A["MELD raw clips and annotations"] --> B["Build raw and fold manifests"]
  B --> C["Train paper-aligned CMT + MIN"]
  C --> D["Evaluate held-out MELD test folds"]
  D --> E["Export metrics, predictions, confusion matrix"]
  E --> F["Interpret imbalance and class confusion"]
""",
        "phase1_replication_pipeline",
    )
    _, figures["vit_png"] = render_mermaid(
        """flowchart LR
  A["Raw .mp4 utterance"] --> B["Sample RGB frames"]
  B --> C["Face-crop or full-frame preprocessing"]
  C --> D["Pretrained ViT"]
  D --> E["Cached .npy visual embeddings"]
  E --> F["Warm-start multimodal Phase 1 model"]
  F --> G["Prediction, confidence, top-3 probabilities"]
""",
        "vit_support_pipeline",
    )
    # Reuse an existing confusion matrix image from the repo.
    fold2_conf_src = ROOT / "results" / "paper_aligned_meld_cv" / "cmt_min" / "fold_2" / "analysis_test" / "confusion_matrix.png"
    fold2_conf_dst = FIG_DIR / "confusion_matrix.png"
    if fold2_conf_src.exists():
        shutil.copy2(fold2_conf_src, fold2_conf_dst)
    figures["fold2_conf"] = fold2_conf_dst
    return figures


def metric_chart() -> Path:
    out = FIG_DIR / "phase1_metric_comparison.png"
    labels = ["Paper-aligned\nMELD mean", "Face-crop\ngated Fold 2", "Face-crop\ngated Fold 4", "Face-crop+\naux Fold 2", "Face-crop+\naux Fold 4"]
    accuracy = [0.6247, 0.6222, 0.5973, 0.6054, 0.5992]
    macro = [0.4395, 0.4191, 0.4156, 0.4351, 0.4330]
    weighted = [0.6195, 0.6109, 0.5968, 0.6022, 0.6056]
    x = range(len(labels))
    fig, ax = plt.subplots(figsize=(11, 4.7))
    ax.bar([i - 0.23 for i in x], accuracy, width=0.22, label="Accuracy", color="#2F5D8A")
    ax.bar(list(x), weighted, width=0.22, label="Weighted F1", color="#4E8D5D")
    ax.bar([i + 0.23 for i in x], macro, width=0.22, label="Macro F1", color="#D28E39")
    ax.set_ylim(0, 0.75)
    ax.set_ylabel("Score")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels)
    ax.set_title("Phase 1 paper-aligned baseline and ViT-support comparison")
    ax.legend(frameon=False, ncol=3, loc="upper center", bbox_to_anchor=(0.5, 1.12))
    ax.grid(axis="y", linestyle="--", alpha=0.25)
    fig.tight_layout()
    fig.savefig(out, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return out


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def make_docx(figs: dict[str, Path]) -> None:
    doc = Document()
    settings = doc.settings
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.element.append(update_fields)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_DST.exists():
        p.add_run().add_picture(str(LOGO_DST), width=Inches(0.7))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Master of Technology in AI and Machine Learning")
    r.bold = True
    r.font.size = Pt(15)
    r.font.name = "Times New Roman"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UE20CS971 Project Phase-1")
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "Times New Roman"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LegalMemoCMT: An Explainable Multilingual Multimodal Emotional Cue Analysis Framework for Indian Courtroom Testimony Using Cross-Modal Transformers")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "Times New Roman"

    def center_line(text: str, size: int = 12, bold: bool = False, italic: bool = False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        return p

    center_line("Submitted by:", 13, bold=True)
    center_line("Rajesh Upadhyaya", 13, bold=True)
    center_line("PES2PGE24DS200", 12, bold=True)
    doc.add_paragraph()
    center_line("Under the guidance of:", 13, bold=True)
    center_line("Ramesh Prakash Guledgudd", 13, bold=True)
    doc.add_paragraph()
    center_line("PES University, Bengaluru", 13, bold=True)

    append_front_matter_docx(doc)
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(ABSTRACT_TEXT)
    doc.add_page_break()
    doc.add_heading("Contents", level=1)
    toc_p = doc.add_paragraph()
    add_word_toc(toc_p)
    doc.add_page_break()
    sections = [
        ("1. Introduction", [
            "Emotion recognition in conversations is difficult because each utterance depends on dialogue context, speaker history, and a mix of textual, acoustic, and visual signals. MELD is a useful benchmark for this setting because it contains multi-party dialogue, seven emotion classes, and both emotion and sentiment annotations. The project therefore focuses on reproducing a stable paper-aligned baseline first, then extending the system with ViT-based facial cues so the raw video pathway can be inspected in the same framework.",
            "The main aim of Phase 1 is implementation fidelity and interpretability. The conversational baseline must be strong enough to serve as the reference model, while the visual extension should be understandable as a support path that can help explain success cases, near-misses, and confident errors. This balance is important because a multimodal emotion model can appear stronger on accuracy while still failing to separate minority emotions or emotionally close classes. The adaptation keeps the transformer-based formulation close to the source design rather than replacing it with a different architecture, which makes the results easier to interpret against the original paper [1,18].",
            "In practical terms, Phase 1 is a controlled replication exercise. Each script, manifest, checkpoint, and evaluation file is treated as part of one traceable pipeline. That means the report must explain not only what the model predicts, but also how the data enters the system, how the representation is built, and how the final metric is selected. This is why the document retains both the benchmark result and the explanation of the supporting visual branch.",
        ]),
        ("2. Related Work", [
            "MELD established a widely used multimodal conversational benchmark for emotion recognition in dialogue [2]. DialogueRNN and DialogueGCN showed that speaker context and conversational structure improve emotion prediction [10,11]. BERT and HuBERT provide pretrained language and speech representations that transfer well to downstream tasks [4,5]. Vision Transformer offers a strong visual backbone for frame-based facial feature extraction [6]. The cross-modal fusion logic used in the project follows the original Transformer attention design [18]. Focal loss and weighted cross-entropy remain standard choices for imbalance-aware learning [9].",
            "The project also aligns with more recent multimodal emotion-recognition work that shows the same general pattern: adding modalities can help, but only when the model is able to trust the correct modality at the right time. This is the motivation for keeping the support path modular and for measuring weighted F1 and macro F1 together rather than relying only on accuracy.",
        ]),
        ("3. Materials and Methods", [
            "The project uses MELD as the primary benchmark. The data is organized through manifests so that each utterance retains its sample identifier, split, label, transcript, and media path. Fold generation is dialogue-aware, which prevents train-validation leakage across utterances from the same dialogue. This setup makes the reported results reproducible and easier to trace from raw data to final evaluation.",
            "The paper-aligned baseline uses BERT for the text stream, HuBERT for the audio stream, a cross-modal transformer to model their interaction, and MIN pooling for the final utterance representation. The ViT support path starts from a raw mp4 clip, samples RGB frames, applies face-crop or full-frame preprocessing, encodes the frames with a pretrained ViT, and stores the resulting features as a cached .npy file. The cached embeddings are then loaded into the multimodal model for inference and analysis.",
            "The methods section is intentionally data-centric because the value of the report depends on whether a reviewer can recreate the same sample flow from the raw dataset to the final prediction files. The manifest layer is therefore treated as a first-class artifact rather than a convenience file.",
        ]),
        ("4. Results", [
            "The paper-aligned MELD baseline is the strongest result in the project. Five-fold evaluation gives a mean accuracy of 0.6247, a mean weighted F1 of 0.6195, and a mean macro F1 of 0.4395. Fold 2 reaches 0.6375 accuracy and 0.6254 weighted F1, while Fold 4 reaches 0.6199 accuracy and 0.6194 weighted F1. The gap between weighted F1 and macro F1 shows that class imbalance remains a central challenge.",
            "The ViT support runs are informative but mixed. Face-crop gated Fold 2 reaches 0.6222 accuracy and 0.6109 weighted F1, while Fold 4 reaches 0.5973 accuracy and 0.5968 weighted F1. Adding auxiliary loss gives 0.6054 accuracy and 0.6022 weighted F1 on Fold 2, and 0.5992 accuracy and 0.6056 weighted F1 on Fold 4. These numbers show that the facial path can help some clips, but it does not consistently surpass the stronger conversational baseline.",
            "The results are therefore read as a hierarchy of evidence. The baseline establishes what the paper-aligned implementation can do. The facial-cue runs show where the visual branch contributes information. The auxiliary-loss run shows how much the visual path can be encouraged to learn without changing the model architecture itself.",
        ]),
        ("5. Discussion", [
            "The error patterns explain why accuracy alone is not sufficient. The Fold 2 confusion matrix shows a neutral-heavy structure, with confusion between neutral and nearby emotions such as anger, joy, and surprise. That pattern is common in conversational emotion recognition, where the model often sees short utterances and incomplete visual context. The facial-cue runs can correct some neutral boundary cases, but they also remain confidently wrong on ambiguous clips, which means confidence is not a substitute for correctness.",
            "The project therefore treats the paper-aligned baseline as the main Phase 1 result and the ViT path as a supported extension. This keeps the report faithful to what the code currently demonstrates: a reproducible conversational model, a visual support branch, and a transparent analysis of where each one helps or fails.",
            "This discussion is also useful because it explains the behavior of the model in the reviewer’s terms: a good metric is not the same thing as a good model on every class. A stable weighted score can hide the fact that the neutral class dominates the decision boundary and that minority emotions are still difficult to separate reliably.",
        ]),
        ("6. Limitations", [
            "The facial-cue branch is supplementary and does not consistently outperform the conversational baseline.",
            "MELD remains class imbalanced, so macro F1 stays substantially lower than weighted F1.",
            "Some clips are still confidently misclassified, especially when the transcript and visual evidence point in different directions.",
            "CREMA-D is supported as a secondary benchmark, but the project’s main replication story is MELD.",
            "The report also does not claim that the ViT path is universally better. Instead, it shows where the path is useful, where it is neutral, and where it introduces new ambiguity. That framing is closer to the actual behavior of the experiments.",
        ]),
        ("7. Conclusion", [
            "Phase 1 now contains a reproducible paper-aligned MELD baseline and a supporting ViT facial-cue path. The baseline is strong enough to serve as the reference model, and the visual branch provides a practical route for future courtroom-style testimony adaptation. The main contribution of Phase 1 is therefore a complete, inspectable, and benchmarked multimodal pipeline rather than a claim of universal improvement from adding video.",
            "The final takeaway for the report is that the project has reached the point where the implementation can be audited, discussed, and extended. The baseline result is stable enough to anchor later experimentation, and the visual branch gives a concrete direction for the next research step.",
        ]),
        ("8. Project Objectives and Scope", [
            "The scope of the project in Phase 1 is deliberately narrow in one sense and broad in another. It is narrow because the main benchmark claim is tied to the paper-aligned MELD conversational setup, where the baseline architecture, the training path, and the evaluation strategy are kept close to the reference implementation. It is broad because the repository also carries the surrounding infrastructure needed to inspect manifests, raw media, cached embeddings, fold-level runs, error analysis, and demo behavior.",
            "The practical objective is to show that the full benchmark loop works end to end. That means raw dialogue data must be converted into stable manifests, the model must train without leakage between dialogue groups, the test results must be exported in a form that can be reviewed later, and the outcome must remain explainable through confusion matrices and per-sample predictions. This is why the project report emphasizes traceability from source data to final metrics.",
            "A second objective is support for later courtroom-testimony work. The Phase 1 project does not attempt to solve that later problem directly, but it creates a visual pathway that can be reused and stress-tested. The same input-output logic that maps a conversational utterance to an emotion label can later be reinterpreted for testimony-style clips, where facial cues and spoken content both matter.",
        ]),
        ("9. Dataset Organization, Label Space, and Fold Rules", [
            "MELD contains seven emotion classes and an imbalanced label distribution, with neutral appearing far more often than the minority classes. That imbalance is not a minor implementation detail; it directly changes what the model learns. If the training objective only optimizes the most frequent class, the model can appear strong on accuracy while still failing the classes that matter for nuanced emotion separation.",
            "The repository therefore organizes the dataset through explicit manifests rather than ad hoc file traversal. Each row links a sample identifier to its transcript, label, processed audio path, processed video path, and split assignment. This is important because it makes the data pipeline inspectable. If a result looks suspicious, the manifest can be used to trace whether the clip came from train, validation, or test, and whether any preprocessing artifact changed the signal.",
            "Fold construction is dialogue-aware. This means utterances from the same dialogue are kept together rather than scattered randomly across training and validation partitions. That choice matters because nearby utterances often share speaker identity, scene context, and even emotional trajectory. A random split would make the evaluation look artificially better by leaking conversational context into the held-out side.",
            "The same grouping logic is retained for the ViT support path, so the visual embeddings remain aligned with the same sample identifiers that the text and audio path use. This makes the visual branch a true support mechanism rather than a separate dataset pipeline with different sample ordering.",
        ]),
        ("10. Baseline Architecture and Modality Roles", [
            "The baseline model treats text as the strongest conversational channel and audio as the complementary paralinguistic channel. BERT converts each utterance transcript into a contextual token representation, while HuBERT converts the audio waveform into a sequence of speech features. The two streams are then fused through cross-modal attention so the classifier does not simply concatenate unrelated embeddings.",
            "Cross-modal transformer fusion is useful because emotion is rarely expressed in a single modality in a clean way. A short utterance may look neutral in text but carry frustration in tone, or vice versa. The fusion block gives the model a chance to reweight each stream according to the context of the current clip. The design is especially relevant in MELD, where short utterances and speaker turn changes are common.",
            "MIN pooling is used at the utterance aggregation stage. Rather than averaging every token or every frame in a uniform way, MIN pooling keeps a compact summary of the strongest aligned signals produced by the fused representation. In practice, this helps the model preserve the most informative parts of the utterance while still keeping the classifier lightweight enough to train on the available hardware.",
            "The classification head then maps the fused representation to the seven emotion classes. The output is not just a class label; it is also a confidence distribution over the labels. That distribution is useful for understanding whether the model is decisively correct, barely correct, or confidently wrong.",
        ]),
        ("11. Training and Evaluation Protocol", [
            "Training is organized around fold-by-fold checkpoints so that the reported result is not tied to a single lucky split. Each fold uses the training partition to learn parameters and the validation partition to decide which checkpoint is best. The project keeps that checkpointing logic explicit because the best epoch is not necessarily the last epoch. A later epoch can fit the training data more closely while performing worse on held-out examples.",
            "The evaluation loop exports accuracy, weighted F1, macro F1, unweighted accuracy, predictions, confusion matrices, and a top-confusion summary. Accuracy gives the simplest summary of the percentage of correct predictions. Weighted F1 compensates for class frequency by giving more weight to common labels. Macro F1 treats every class equally, which makes it much more sensitive to minority-class behavior. Unweighted accuracy is useful because it exposes how well the model performs when the class imbalance is not allowed to dominate the score.",
            "The result files are not just bookkeeping. They are the basis for error analysis. If a fold performs well on weighted F1 but poorly on macro F1, the model is likely over-reliant on frequent classes. If the confusion matrix is neutral-heavy, then the model has probably learned a generic safe prediction behavior rather than a robust emotion separator.",
            "This protocol is also why the project report emphasizes fold 2 and fold 4. Those folds are useful inspection points because they expose whether the behavior is stable or whether it changes in a way that depends on the dialogue partition.",
        ]),
        ("12. ViT Facial-Cue Extension and Video Caching", [
            "The ViT extension begins with raw mp4 input rather than with an already prepared feature tensor. The script samples RGB frames from the clip, optionally applies face-crop preprocessing, and sends the resulting images through a pretrained Vision Transformer. This converts a variable-length video into a compact fixed-shape representation that can be cached and reused.",
            "The cached .npy file is important because it separates feature generation from inference. Once the visual embeddings are saved, the model can reload them without repeating the full video decoding process every time. That makes experiments faster, easier to reproduce, and easier to inspect. It also means that later demo runs can focus on inference and explanation instead of repeating the expensive image encoding stage.",
            "Face-crop preprocessing narrows the signal to the speaker region. For emotion recognition and testimony analysis, that is often more informative than full-frame context because facial expression is usually the most direct visual cue for affect. Full-frame preprocessing still has value, especially when surrounding context carries meaningful evidence, but it also brings more background variation and more opportunity for noise.",
            "The ViT branch is not a replacement for the conversational model. It is a support path that reveals whether visual evidence changes the prediction boundary, whether it helps resolve neutral-heavy examples, and whether it still fails when the spoken content and the facial signal disagree.",
        ]),
        ("13. Comparative Results and Error Analysis", [
            "The baseline remains the strongest overall result in the current project. The paper-aligned MELD run still sets the main reference point, and the facial-cue runs should be read as supporting experiments that probe what happens when the visual path is introduced. The most useful comparison is not a single win/loss statement, but the pattern across accuracy, weighted F1, macro F1, and unweighted accuracy.",
            "The folded results show a consistent pattern. The face-crop gated run improves some clips, but it does not reliably exceed the conversational baseline. The auxiliary-loss variant helps the visual branch learn a little more structure, yet the final metrics still remain mixed. That suggests the visual features are adding signal, but not enough signal to overwhelm the strong text-plus-audio baseline.",
            "The confusion analysis shows why this happens. Neutral is the dominant class, and the model often confuses neutral with joy, anger, surprise, or fear. That is a classic imbalance pattern. The model is learning something useful because it often places the correct label among the top predictions, but it still has difficulty separating emotionally close classes when the evidence is sparse or contradictory.",
            "A confident wrong prediction is especially important to inspect. If the output probability is high but incorrect, then the issue is not simply uncertainty. It may indicate a strong but misplaced bias toward one class, or a preprocessing path that emphasized the wrong portion of the clip. For that reason, the confidence score must always be read together with the confusion matrix and the predicted-vs-actual table.",
        ]),
        ("14. Reproducibility, Scripts, and Outputs", [
            "The project is reproducible because the important actions are separated into named scripts. One set of scripts builds the MELD fold manifests. Another set trains the baseline or the ViT support variants. A third set analyzes predictions, exports confusion matrices, and summarizes errors. The demo scripts then read the saved checkpoints and run inference on a single raw mp4 clip or on a small set of curated clips.",
            "This separation is deliberate. It lets the project in Phase 1 keep the same model weights while changing only one ingredient at a time. That is the only reliable way to understand whether the change in outcome came from the loss function, the visual input, the fusion mechanism, or some accidental difference in the execution path.",
            "The output artifacts are also part of the deliverable. Metrics.json files capture the final numbers. Prediction CSVs keep the sample-level output. Confusion matrices show where the model fails in aggregate. Top-confusion tables identify the most frequent error pairs. These are the files that make the project auditable during review.",
        ]),
        ("15. Limitations and Phase 2 Direction", [
            "The current Phase 1 implementation is strong enough to support review and explanation, but it is not the endpoint. The video branch still needs more work if the goal is to make courtroom-style facial reasoning central rather than supplementary. The model also remains sensitive to neutral-heavy data behavior, which limits how far accuracy alone can be pushed without changing the input strategy or the fusion strategy.",
            "For that reason, the project uses Phase 1 to establish the baseline and to identify where visual support is informative. The next direction is to push the visual path further, especially where the clips contain meaningful face information and where the current model still confuses emotionally close labels. The important point is that the remaining work is now targeted rather than exploratory.",
            "This makes the current report suitable as a project milestone document. It records what is implemented, what is measured, what remains weak, and what the next step must focus on.",
        ]),
        ("16. Background", [
            "The broader use case behind LegalMemoCMT is courtroom testimony analysis, where spoken content, delivery style, and facial cues can all affect the perceived emotion or intent of a witness or speaker. In that setting, a system needs to do more than output a label. It needs to preserve traceability so that the evidence behind a prediction can be reviewed later.",
            "MELD is not a courtroom dataset, but it is a practical Phase 1 benchmark because it provides multi-party conversational context, overlapping emotions, and visible class imbalance. These properties make it a good test bed for the technical parts of the pipeline that will eventually be reused in testimony analysis.",
            "The background problem is therefore one of transferability. The project must first prove that the conversational baseline and the visual feature path can be implemented in a reproducible way on a benchmark dataset. Only then does it become meaningful to adapt the same design to courtroom-style material, where the audio and video signals are often noisier and the emotional boundaries may be even more subtle.",
        ]),
        ("17. Problem Statement", [
            "The core problem in Phase 1 is to build an explainable multimodal pipeline that can replicate the paper-aligned conversational result and also accept a visual branch derived from raw video. The pipeline must remain traceable from raw data to output metrics, and it must support inspection of errors rather than only headline accuracy.",
            "The technical challenge is that the target labels are imbalanced, the conversational context is short, and the video path may contain either strong face cues or distracting background information. That means the model can learn a dominant safe class, such as neutral, while still failing to separate emotionally close labels. The project therefore needs an architecture, a training objective, and an evaluation strategy that make these failure modes visible.",
            "A second part of the problem is reproducibility. The same sample must map to the same transcript, audio features, visual features, and label across training and analysis. If the data path is not stable, the model cannot be audited and the observed behavior cannot be trusted as a benchmark result.",
        ]),
        ("18. Literature Survey", [
            "The base paper MemoCMT motivates the main baseline design by showing that cross-modal feature fusion can be used to combine text and audio representations through an attention-based transformer block [1]. This is the most direct architectural reference for the project in Phase 1.",
            "MELD provides the benchmark environment and also illustrates why conversational emotion recognition is hard: many clips are short, the label distribution is skewed, and the same emotional category may appear in very different dialogue settings [2]. CREMA-D complements this by showing how multimodal emotion signals also appear in actor-based speech corpora [3].",
            "BERT and HuBERT provide the pretrained text and audio encoders used in the baseline [4,5]. Vision Transformer is the visual backbone used to convert frame samples into embeddings [6]. The general attention design comes from Transformer-based sequence modeling [18], while weighted cross-entropy and focal loss provide standard tools for imbalance-aware optimization [9].",
            "The broader emotion-recognition literature also shows that context-aware models such as DialogueRNN and DialogueGCN are useful in conversational settings [10,11]. Surveys on multimodal emotion recognition repeatedly note that multimodal gains are not automatic and that error analysis is often more informative than aggregate accuracy alone [13-17]. That observation is consistent with the project results, where the visual path helps some examples but does not universally beat the stronger conversational baseline.",
        ]),
        ("19. System Requirements Specification", [
            "Hardware requirements are kept modest enough for local development and review, but the pipeline benefits from a CUDA-capable GPU when the visual or multimodal runs are trained. CPU execution is supported for debugging and analysis, although the data-loading stage can become the bottleneck because video decoding and audio preprocessing take time.",
            "Software requirements include Python, PyTorch, Transformers, NumPy, Pandas, scikit-learn, python-docx, report generation tools, ffmpeg for media handling, and the model checkpoints or pretrained weights required by the text, audio, and ViT components. The report generation scripts also require the ability to render Mermaid figures and to build PDF and Word outputs.",
            "Functional requirements include manifest creation, fold generation, model training, evaluation, per-sample prediction export, confusion-matrix generation, raw-mp4 demo inference, and repeatable checkpoint selection. Non-functional requirements include reproducibility, traceability, low leakage across folds, and interpretable error analysis.",
            "The main assumptions are that the manifests correctly map each utterance to one label, the raw media paths are valid, and the fold assignment keeps dialogue groups intact. The main constraint is that class imbalance remains present even after cross-validation, so metric selection must include macro F1 and unweighted accuracy rather than only accuracy.",
            "The system requirements section also makes clear which external pieces are assumed by the workflow. If a dependency such as ffmpeg, a Hugging Face model, or a cached video path is missing, the pipeline can fail before the model itself is even evaluated. That is why the report treats environment setup as part of the technical system rather than as a side note.",
        ]),
        ("20. Proposed Methodology", [
            "The proposed methodology uses a staged pipeline. First, the data is organized into manifests and dialogue-safe folds. Second, the paper-aligned baseline is trained with text and audio encoders plus cross-modal transformer fusion. Third, the visual branch is added through raw mp4 sampling and ViT embeddings. Fourth, the outputs are evaluated through metrics, confusion matrices, and top-confusion summaries.",
            "The architecture is intentionally modular. Text, audio, and video can be studied independently or in combination, which makes it easier to isolate the effect of a design change. The same modularity also helps during review because a single change in the input path does not force a complete rewrite of the rest of the pipeline.",
            "The workflow is therefore not only a training recipe but also an explanation recipe. Each stage produces artifacts that can be checked separately: a manifest shows what data was used, a checkpoint shows what model was selected, a metrics file shows how well it performed, and a confusion matrix shows where it failed.",
        ]),
        ("21. Implementation Details", [
            "Implementation starts from the repository scripts. Separate scripts build manifests, train the paper-aligned MELD folds, train the face-crop gated runs, train the gated plus auxiliary-loss runs, run analysis, and execute demo predictions on raw mp4 clips. This separation keeps the codebase inspectable and makes it easier to explain what each run is doing.",
            "The data layer converts raw inputs into reproducible intermediate artifacts such as manifests, cached audio/video features, and .npy files. The model layer loads pretrained text and audio encoders and optionally a ViT-based visual branch. The evaluation layer exports per-sample CSVs, confusion matrices, and metrics. The documentation layer turns these artifacts into reportable outputs.",
            "The key implementation details that matter for correctness are sample-id alignment, consistent split assignment, checkpoint loading, and the use of the same metric definitions across all runs. Without those checks, a result may appear better or worse simply because the data or the metric changed, not because the model truly changed.",
        ]),
        ("22. Intermediate Results and Discussion", [
            "The baseline results remain the strongest and most stable. The MELD paper-aligned run provides the main reference values, and the ViT runs are interpreted against that baseline rather than in isolation. The comparison shows that the visual branch is informative but still not dominant over the conversational features.",
            "The confusion matrix behavior is especially important. Neutral-heavy errors remain the major pattern, and several emotionally close classes are still mixed up. This pattern is visible in the fold-level outputs, where accuracy can look acceptable while macro F1 stays noticeably lower. That gap shows that the model is still biased toward easier or more frequent classes.",
            "The face-crop gated and auxiliary-loss runs add nuance. They show that the video pathway can change the prediction for certain clips, including some near-miss neutral cases, but they also show that the same pathway can produce confident wrong outputs when the visual evidence is ambiguous or when the facial signal is weak. The result is improvement in observability more than universal score improvement.",
        ]),
        ("23. Conclusions and Future Work", [
            "The Phase 1 project now contains a reproducible conversational baseline and a supporting ViT facial-cue path, both of which can be traced from raw data through model inference to evaluation output. That is the main technical achievement of the current stage.",
            "Future work should focus on making the visual branch more central, especially for testimony-style clips where the facial signal may carry more predictive value than in MELD. The next work also needs to refine the fusion strategy, because simply attaching more modalities does not guarantee that the model will use them effectively.",
            "A further direction is to study more domain-relevant courtroom data and to compare whether face crops, full frames, and gated fusion behave differently when the speaker is intentionally framed for testimony analysis rather than conversational emotion recognition.",
        ]),
        ("24. Comparative Metric Summary", [
            "The metric summary places the baseline and the visual-support branches side by side so the result can be read as a sequence of controlled experiments rather than as isolated scores.",
        ]),
        ("Appendix A. Script List and Command Sequence", [
            "The appendix lists the scripts used to reproduce the Phase 1 project workflow from training to analysis and demo. The exact sequence depends on whether the paper-aligned baseline, the ViT face-crop support path, or the gated plus auxiliary-loss variant is being reviewed.",
            "1. Build or verify the paper-aligned MELD folds: bash scripts/run_paper_aligned_meld_cv.sh",
            "2. Analyze Fold 2 predictions and confusion matrix: bash scripts/analyze_meld_vit_facecue_fold2.sh",
            "3. Analyze Fold 4 predictions and confusion matrix: bash scripts/analyze_meld_vit_facecue_fold4.sh",
            "4. Train the face-crop gated variant: bash scripts/run_meld_vit_facecrop_gated_fold2.sh and bash scripts/run_meld_vit_facecrop_gated_fold4.sh",
            "5. Train the face-crop gated plus auxiliary-loss variant: bash scripts/run_meld_vit_facecrop_gated_video_aux_fold2.sh and bash scripts/run_meld_vit_facecrop_gated_video_aux_fold4.sh",
            "6. Run the paper-aligned raw-mp4 demo: bash scripts/run_demo_paper_aligned_raw_mp4.sh <sample_id> <video_path>",
            "7. Run the gated plus auxiliary-loss raw-mp4 demo: bash scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh <sample_id> <video_path>",
            "8. Review the per-sample predictions, metrics.json, top-confusion table, and confusion matrix generated by the analysis scripts before presenting the results.",
        ]),
    ]

    for heading, paras in sections:
        doc.add_heading(heading, level=1)
        for para in paras:
            p = doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph("Table 1. Comparative Metric Summary.")
    metric_table = doc.add_table(rows=1, cols=5)
    metric_table.style = "Table Grid"
    headers = ["Setting", "Accuracy", "Weighted F1", "Macro F1", "Unweighted Acc."]
    for cell, text in zip(metric_table.rows[0].cells, headers):
        cell.text = text
    metric_rows = [
        ["Paper-aligned MELD mean", "0.6247", "0.6195", "0.4395", "0.4417"],
        ["Face-crop gated Fold 2", "0.6222", "0.6109", "0.4191", "0.4173"],
        ["Face-crop gated Fold 4", "0.5973", "0.5968", "0.4156", "0.4254"],
        ["Face-crop gated + aux Fold 2", "0.6054", "0.6022", "0.4351", "0.4492"],
        ["Face-crop gated + aux Fold 4", "0.5992", "0.6056", "0.4330", "0.4638"],
    ]
    for row in metric_rows:
        cells = metric_table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text
    doc.add_paragraph("The table above keeps the baseline and the visual-support branches side by side so the result can be read as a progression of experiments rather than as isolated checkpoints.").paragraph_format.space_after = Pt(6)

    doc.add_heading("4.1 Key Figures", level=2)
    for caption, img in [
        ("Figure 1. Paper-aligned Phase 1 replication workflow.", figs["pipeline_png"]),
        ("Figure 2. Raw-mp4 ViT facial-cue support path.", figs["vit_png"]),
        ("Figure 3. Phase 1 metric comparison across the baseline and ViT-supported runs.", FIG_DIR / "phase1_metric_comparison.png"),
        ("Figure 4. Fold 2 confusion matrix for the paper-aligned MELD baseline.", figs["fold2_conf"]),
    ]:
        if img.exists():
            doc.add_paragraph().add_run().add_picture(str(img), width=Inches(5.8))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True

    doc.add_heading("References", level=1)
    refs = [
        "1. Khan, M., Tran, P.-N., Pham, N. T., El Saddik, A., & Othmani, A. (2025). MemoCMT: multimodal emotion recognition using cross-modal transformer-based feature fusion. Scientific Reports, 15(1), 5473. https://doi.org/10.1038/s41598-025-89202-x",
        "2. Poria, S., Hazarika, D., Majumder, N., Naik, G., Cambria, E., & Mihalcea, R. (2019). MELD: A Multimodal Multi-Party Dataset for Emotion Recognition in Conversations. ACL. https://aclanthology.org/P19-1050/",
        "3. Cao, H., Cooper, D., Keutmann, M., Gur, R., Nenkova, A., & Verma, R. (2014). CREMA-D: Crowd-sourced Emotional Multimodal Actors Dataset. IEEE Transactions on Affective Computing. https://pmc.ncbi.nlm.nih.gov/articles/PMC4313618/",
        "4. Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. https://arxiv.org/abs/1810.04805",
        "5. Hsu, W.-N., Bolte, B., Tsai, Y.-H. H., Lakhotia, K., Salakhutdinov, R., & Mohamed, A. (2021). HuBERT: Self-Supervised Speech Representation Learning by Masked Prediction of Hidden Units. https://arxiv.org/abs/2106.07447",
        "6. Dosovitskiy, A., Beyer, L., Kolesnikov, A., et al. (2021). An Image is Worth 16x16 Words: Transformers for Image Recognition at Scale. https://arxiv.org/abs/2010.11929",
        "7. Schneider, S., Baevski, A., Collobert, R., & Auli, M. (2019). wav2vec: Unsupervised Pre-training for Speech Recognition. https://arxiv.org/abs/1904.05862",
        "8. Baevski, A., Zhou, H., Mohamed, A., & Auli, M. (2020). wav2vec 2.0: A Framework for Self-Supervised Learning of Speech Representations. https://arxiv.org/abs/2006.11477",
        "9. Lin, T.-Y., Goyal, P., Girshick, R., He, K., & Dollár, P. (2018). Focal Loss for Dense Object Detection. https://arxiv.org/abs/1708.02002",
        "10. Majumder, N., Poria, S., Hazarika, D., Mihalcea, R., Gelbukh, A., & Cambria, E. (2019). DialogueRNN: An Attentive RNN for Emotion Detection in Conversations. https://arxiv.org/abs/1811.00405",
        "11. Ghosal, D., Majumder, N., Poria, S., Chhaya, N., & Gelbukh, A. (2020). DialogueGCN: A Graph Convolutional Neural Network for Emotion Recognition in Conversation. https://aclanthology.org/D19-1015/",
        "12. Ghosh, S., Ramaneswaran, S., Tyagi, U., et al. (2022). M-MELD: A Multilingual Multi-Party Dataset for Emotion Recognition in Conversations. https://arxiv.org/abs/2203.16799",
        "13. Wu, C., Cai, Y., Liu, Y., Zhu, P., Xue, Y., Gong, Z., Hirschberg, J., & Ma, B. (2025). Multimodal Emotion Recognition in Conversations: A Survey of Methods, Trends, Challenges and Prospects. https://arxiv.org/abs/2505.20511",
        "14. Shou, Y., Meng, T., Ai, W., Yin, N., & Li, K. (2023). A Comprehensive Survey on Multi-modal Conversational Emotion Recognition with Deep Learning. https://arxiv.org/abs/2312.05735",
        "15. Zhang, X., et al. (2024). Emotion Recognition in Conversations: A Survey Focusing on Multimodal Approaches. Electronics, 12(22), 4714. https://www.mdpi.com/2079-9292/12/22/4714",
        "16. TelME: Teacher-leading Multimodal Fusion Network for Emotion Recognition in Conversation. https://aclanthology.org/2024.naacl-long.5/",
        "17. Multimodal Emotion Recognition Calibration in Conversations. 2024. https://www.atailab.cn/seminar2024Fall/pdf/2024_MM_Multimodal%20Emotion%20Recognition%20Calibration%20in%20Conversations.pdf",
        "18. Vaswani, A., Shazeer, N., Parmar, N., et al. (2017). Attention Is All You Need. NeurIPS 2017. https://arxiv.org/abs/1706.03762",
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    doc.save(DOCX_OUT)


def tex_escape(text: str) -> str:
    repl = {
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
        "\\": r"\textbackslash{}",
    }
    for k, v in repl.items():
        text = text.replace(k, v)
    return text


def make_tex(figs: dict[str, Path]) -> None:
    lines: list[str] = []
    add = lines.append
    add(r"\documentclass[12pt]{report}")
    add(r"\usepackage[a4paper,margin=1in]{geometry}")
    add(r"\usepackage{graphicx}")
    add(r"\usepackage{hyperref}")
    add(r"\usepackage{booktabs}")
    add(r"\usepackage{longtable}")
    add(r"\usepackage{array}")
    add(r"\usepackage{float}")
    add(r"\usepackage{caption}")
    add(r"\usepackage{fancyhdr}")
    add(r"\pagestyle{fancy}")
    add(r"\fancyhf{}")
    add(r"\rhead{LegalMemoCMT Phase 1}")
    add(r"\lhead{Project Report}")
    add(r"\cfoot{\thepage}")
    add(r"\begin{document}")
    add(r"\begin{titlepage}")
    add(r"\centering")
    add(rf"\includegraphics[width=2.5cm]{{{LOGO_DST.name}}}\\[0.7cm]")
    add(r"\textbf{\Large Master of Technology in AI and Machine Learning}\\[0.5cm]")
    add(r"\textbf{\Large UE20CS971 Project Phase-1}\\[1cm]")
    add(r"\textbf{\LARGE LegalMemoCMT: An Explainable Multilingual Multimodal Emotional Cue Analysis Framework for Indian Courtroom Testimony Using Cross-Modal Transformers}\\[1cm]")
    add(r"Submitted by:\\[0.2cm]")
    add(r"\textbf{Rajesh Upadhyaya}\\")
    add(r"\textbf{PES2PGE24DS200}\\[0.8cm]")
    add(r"Under the guidance of:\\[0.2cm]")
    add(r"\textbf{Ramesh Prakash Guledgudd}\\[1cm]")
    add(r"PES University, Bengaluru")
    add(r"\end{titlepage}")
    add(r"\section*{Certificate}")
    add(r"This is to certify that the Phase 1 project entitled LegalMemoCMT: An Explainable Multilingual Multimodal Emotional Cue Analysis Framework for Indian Courtroom Testimony Using Cross-Modal Transformers is a bona fide work carried out under the guidance of Ramesh Prakash Guledgudd in partial fulfilment of the requirements for the award of the degree of Master of Technology in AI and Machine Learning.")
    add(r"\newpage")
    add(r"\section*{Declaration}")
    add(r"I hereby declare that the Phase 1 project work presented in this report is carried out as part of the LegalMemoCMT project under the prescribed academic guidance and that the submitted content reflects the implementation, analysis, and evaluation carried out in the repository.")
    add(r"\newpage")
    add(r"\section*{Acknowledgement}")
    add(r"I acknowledge the guidance, review support, and academic feedback received during the Phase 1 project. The work also benefited from the available benchmark datasets, pretrained encoders, and the project scripts that made the implementation traceable from raw input to final metrics.")
    add(r"\newpage")
    add(r"\section*{Abstract}")
    add(tex_escape(ABSTRACT_TEXT) + "\n\n")
    add(r"\newpage")
    add(r"\tableofcontents")
    add(r"\newpage")
    add(r"\section{Introduction}")
    intro_paras = [
        "Emotion recognition in conversations is difficult because each utterance depends on dialogue context, speaker history, and a mix of textual, acoustic, and visual signals. MELD is a useful benchmark for this setting because it contains multi-party dialogue, seven emotion classes, and both emotion and sentiment annotations. The project therefore focuses on reproducing a stable paper-aligned baseline first, then extending the system with ViT-based facial cues so the raw video pathway can be inspected in the same framework.",
        "The main aim of Phase 1 is implementation fidelity and interpretability. The conversational baseline must be strong enough to serve as the reference model, while the visual extension should be understandable as a support path that can help explain success cases, near-misses, and confident errors. The adaptation keeps the transformer-based formulation close to the source design rather than replacing it with a different architecture, which makes the results easier to interpret against the original paper [1,18].",
    ]
    for p in intro_paras:
        add(tex_escape(p) + "\n\n")
    add(r"\section{Related Work}")
    rel = [
        "MELD established a widely used multimodal conversational benchmark for emotion recognition in dialogue [2]. DialogueRNN and DialogueGCN showed that speaker context and conversational structure improve emotion prediction [10,11]. BERT and HuBERT provide pretrained language and speech representations that transfer well to downstream tasks [4,5]. Vision Transformer offers a strong visual backbone for frame-based facial feature extraction [6]. The cross-modal fusion logic used in the project follows the original Transformer attention design [18]. Focal loss and weighted cross-entropy remain standard choices for imbalance-aware learning [9].",
    ]
    for p in rel:
        add(tex_escape(p) + "\n\n")
    add(r"\section{Materials and Methods}")
    add(r"\section{Dataset and Manifests}")
    add(r"The project uses MELD as the primary benchmark. The data is organized through manifests so that each utterance retains its sample identifier, split, label, transcript, and media path. Fold generation is dialogue-aware, which prevents train-validation leakage across utterances from the same dialogue.")
    add(r"\section{Baseline Model}")
    add(r"The paper-aligned baseline uses BERT for the text stream, HuBERT for the audio stream, a cross-modal transformer to model their interaction, and MIN pooling for the final utterance representation.")
    add(r"\section{ViT Facial-Cue Support Path}")
    add(r"The facial-cue extension starts from a raw mp4 clip, samples RGB frames, applies face-crop or full-frame preprocessing, encodes the frames with a pretrained ViT, and stores the resulting features as a cached .npy file.")
    add(r"\section{Evaluation Protocol}")
    add(r"Evaluation uses accuracy, weighted F1, macro F1, and unweighted accuracy. Accuracy and weighted F1 tend to look stronger than macro F1 under class imbalance, so macro F1 is essential for understanding minority-class behavior.")
    add(r"\section{Results}")
    add(r"The paper-aligned MELD baseline is the strongest result in the project. Five-fold evaluation gives a mean accuracy of 0.6247, a mean weighted F1 of 0.6195, and a mean macro F1 of 0.4395. Fold 2 reaches 0.6375 accuracy and 0.6254 weighted F1, while Fold 4 reaches 0.6199 accuracy and 0.6194 weighted F1. The gap between weighted F1 and macro F1 shows that class imbalance remains a central challenge.")
    add(r"The ViT support runs are informative but mixed. Face-crop gated Fold 2 reaches 0.6222 accuracy and 0.6109 weighted F1, while Fold 4 reaches 0.5973 accuracy and 0.5968 weighted F1. Adding auxiliary loss gives 0.6054 accuracy and 0.6022 weighted F1 on Fold 2, and 0.5992 accuracy and 0.6056 weighted F1 on Fold 4.")
    add(r"\begin{figure}[H]\centering")
    add(rf"\includegraphics[width=0.85\textwidth]{{figures/{figs['pipeline_png'].name}}}")
    add(r"\caption{Paper-aligned Phase 1 replication workflow.}")
    add(r"\end{figure}")
    add(r"\begin{figure}[H]\centering")
    add(rf"\includegraphics[width=0.85\textwidth]{{figures/{figs['vit_png'].name}}}")
    add(r"\caption{Raw-mp4 ViT facial-cue support path.}")
    add(r"\end{figure}")
    add(r"\begin{figure}[H]\centering")
    add(rf"\includegraphics[width=0.9\textwidth]{{figures/{(FIG_DIR / 'phase1_metric_comparison.png').name}}}")
    add(r"\caption{Phase 1 metric comparison across the baseline and ViT-supported runs.}")
    add(r"\end{figure}")
    add(r"\begin{figure}[H]\centering")
    add(rf"\includegraphics[width=0.72\textwidth]{{figures/{figs['fold2_conf'].name}}}")
    add(r"\caption{Fold 2 confusion matrix for the paper-aligned MELD baseline.}")
    add(r"\end{figure}")
    add(r"\section{Discussion}")
    add(r"The error patterns explain why accuracy alone is not sufficient. The Fold 2 confusion matrix shows a neutral-heavy structure, with confusion between neutral and nearby emotions such as anger, joy, and surprise. The facial-cue runs can correct some neutral boundary cases, but they also remain confidently wrong on ambiguous clips, which means confidence is not a substitute for correctness.")
    add(r"\section{Limitations}")
    add(r"The facial-cue branch is supplementary and does not consistently outperform the conversational baseline. MELD remains class imbalanced, so macro F1 stays substantially lower than weighted F1. Some clips are still confidently misclassified, especially when the transcript and visual evidence point in different directions.")
    add(r"\section{Conclusion}")
    add(r"Phase 1 now contains a reproducible paper-aligned MELD baseline and a supporting ViT facial-cue path. The baseline is strong enough to serve as the reference model, and the visual branch provides a practical route for future courtroom-style testimony adaptation.")
    add(r"\section{Project Objectives and Scope}")
    add(r"The scope of the project in Phase 1 is deliberately narrow in one sense and broad in another. It is narrow because the main benchmark claim is tied to the paper-aligned MELD conversational setup, where the baseline architecture, the training path, and the evaluation strategy are kept close to the reference implementation. It is broad because the repository also carries the surrounding infrastructure needed to inspect manifests, raw media, cached embeddings, fold-level runs, error analysis, and demo behavior.")
    add(r"The practical objective is to show that the full benchmark loop works end to end. That means raw dialogue data must be converted into stable manifests, the model must train without leakage between dialogue groups, the test results must be exported in a form that can be reviewed later, and the outcome must remain explainable through confusion matrices and per-sample predictions. This is why the project report emphasizes traceability from source data to final metrics.")
    add(r"The same system also creates a support path for later courtroom-testimony work. The Phase 1 project does not attempt to solve that later problem directly, but it creates a visual pathway that can be reused and stress-tested. The same input-output logic that maps a conversational utterance to an emotion label can later be reinterpreted for testimony-style clips, where facial cues and spoken content both matter.")
    add(r"\section{Dataset Organization, Label Space, and Fold Rules}")
    add(r"MELD contains seven emotion classes and an imbalanced label distribution, with neutral appearing far more often than the minority classes. That imbalance is not a minor implementation detail; it directly changes what the model learns. If the training objective only optimizes the most frequent class, the model can appear strong on accuracy while still failing the classes that matter for nuanced emotion separation.")
    add(r"The repository therefore organizes the dataset through explicit manifests rather than ad hoc file traversal. Each row links a sample identifier to its transcript, label, processed audio path, processed video path, and split assignment. This is important because it makes the data pipeline inspectable. If a result looks suspicious, the manifest can be used to trace whether the clip came from train, validation, or test, and whether any preprocessing artifact changed the signal.")
    add(r"Fold construction is dialogue-aware. This means utterances from the same dialogue are kept together rather than scattered randomly across training and validation partitions. That choice matters because nearby utterances often share speaker identity, scene context, and even emotional trajectory. A random split would make the evaluation look artificially better by leaking conversational context into the held-out side.")
    add(r"The same grouping logic is retained for the ViT support path, so the visual embeddings remain aligned with the same sample identifiers that the text and audio path use. This makes the visual branch a true support mechanism rather than a separate dataset pipeline with different sample ordering.")
    add(r"\section{Baseline Architecture and Modality Roles}")
    add(r"The baseline model treats text as the strongest conversational channel and audio as the complementary paralinguistic channel. BERT converts each utterance transcript into a contextual token representation, while HuBERT converts the audio waveform into a sequence of speech features. The two streams are then fused through cross-modal attention so the classifier does not simply concatenate unrelated embeddings.")
    add(r"Cross-modal transformer fusion is useful because emotion is rarely expressed in a single modality in a clean way. A short utterance may look neutral in text but carry frustration in tone, or vice versa. The fusion block gives the model a chance to reweight each stream according to the context of the current clip. The design is especially relevant in MELD, where short utterances and speaker turn changes are common.")
    add(r"MIN pooling is used at the utterance aggregation stage. Rather than averaging every token or every frame in a uniform way, MIN pooling keeps a compact summary of the strongest aligned signals produced by the fused representation. In practice, this helps the model preserve the most informative parts of the utterance while still keeping the classifier lightweight enough to train on the available hardware.")
    add(r"The classification head then maps the fused representation to the seven emotion classes. The output is not just a class label; it is also a confidence distribution over the labels. That distribution is useful for understanding whether the model is decisively correct, barely correct, or confidently wrong.")
    add(r"\section{Training and Evaluation Protocol}")
    add(r"Training is organized around fold-by-fold checkpoints so that the reported result is not tied to a single lucky split. Each fold uses the training partition to learn parameters and the validation partition to decide which checkpoint is best. The project keeps that checkpointing logic explicit because the best epoch is not necessarily the last epoch. A later epoch can fit the training data more closely while performing worse on held-out examples.")
    add(r"The evaluation loop exports accuracy, weighted F1, macro F1, unweighted accuracy, predictions, confusion matrices, and a top-confusion summary. Accuracy gives the simplest summary of the percentage of correct predictions. Weighted F1 compensates for class frequency by giving more weight to common labels. Macro F1 treats every class equally, which makes it much more sensitive to minority-class behavior. Unweighted accuracy is useful because it exposes how well the model performs when the class imbalance is not allowed to dominate the score.")
    add(r"The result files are not just bookkeeping. They are the basis for error analysis. If a fold performs well on weighted F1 but poorly on macro F1, the model is likely over-reliant on frequent classes. If the confusion matrix is neutral-heavy, then the model has probably learned a generic safe prediction behavior rather than a robust emotion separator.")
    add(r"This protocol is also why the project report emphasizes fold 2 and fold 4. Those folds are useful inspection points because they expose whether the behavior is stable or whether it changes in a way that depends on the dialogue partition.")
    add(r"\section{ViT Facial-Cue Extension and Video Caching}")
    add(r"The ViT extension begins with raw mp4 input rather than with an already prepared feature tensor. The script samples RGB frames from the clip, optionally applies face-crop preprocessing, and sends the resulting images through a pretrained Vision Transformer. This converts a variable-length video into a compact fixed-shape representation that can be cached and reused.")
    add(r"The cached .npy file is important because it separates feature generation from inference. Once the visual embeddings are saved, the model can reload them without repeating the full video decoding process every time. That makes experiments faster, easier to reproduce, and easier to inspect. It also means that later demo runs can focus on inference and explanation instead of repeating the expensive image encoding stage.")
    add(r"Face-crop preprocessing narrows the signal to the speaker region. For emotion recognition and testimony analysis, that is often more informative than full-frame context because facial expression is usually the most direct visual cue for affect. Full-frame preprocessing still has value, especially when surrounding context carries meaningful evidence, but it also brings more background variation and more opportunity for noise.")
    add(r"The ViT branch is not a replacement for the conversational model. It is a support path that reveals whether visual evidence changes the prediction boundary, whether it helps resolve neutral-heavy examples, and whether it still fails when the spoken content and the facial signal disagree.")
    add(r"\section{Comparative Results and Error Analysis}")
    add(r"The baseline remains the strongest overall result in the current project. The paper-aligned MELD run still sets the main reference point, and the facial-cue runs should be read as supporting experiments that probe what happens when the visual path is introduced. The most useful comparison is not a single win/loss statement, but the pattern across accuracy, weighted F1, macro F1, and unweighted accuracy.")
    add(r"The folded results show a consistent pattern. The face-crop gated run improves some clips, but it does not reliably exceed the conversational baseline. The auxiliary-loss variant helps the visual branch learn a little more structure, yet the final metrics still remain mixed. That suggests the visual features are adding signal, but not enough signal to overwhelm the strong text-plus-audio baseline.")
    add(r"The confusion analysis shows why this happens. Neutral is the dominant class, and the model often confuses neutral with joy, anger, surprise, or fear. That is a classic imbalance pattern. The model is learning something useful because it often places the correct label among the top predictions, but it still has difficulty separating emotionally close classes when the evidence is sparse or contradictory.")
    add(r"A confident wrong prediction is especially important to inspect. If the output probability is high but incorrect, then the issue is not simply uncertainty. It may indicate a strong but misplaced bias toward one class, or a preprocessing path that emphasized the wrong portion of the clip. For that reason, the confidence score must always be read together with the confusion matrix and the predicted-vs-actual table.")
    add(r"\section{Reproducibility, Scripts, and Outputs}")
    add(r"The project is reproducible because the important actions are separated into named scripts. One set of scripts builds the MELD fold manifests. Another set trains the baseline or the ViT support variants. A third set analyzes predictions, exports confusion matrices, and summarizes errors. The demo scripts then read the saved checkpoints and run inference on a single raw mp4 clip or on a small set of curated clips.")
    add(r"This separation is deliberate. It lets the project in Phase 1 keep the same model weights while changing only one ingredient at a time. That is the only reliable way to understand whether the change in outcome came from the loss function, the visual input, the fusion mechanism, or some accidental difference in the execution path.")
    add(r"The output artifacts are also part of the deliverable. Metrics.json files capture the final numbers. Prediction CSVs keep the sample-level output. Confusion matrices show where the model fails in aggregate. Top-confusion tables identify the most frequent error pairs. These are the files that make the project auditable during review.")
    add(r"\section{Limitations and Phase 2 Direction}")
    add(r"The current Phase 1 implementation is strong enough to support review and explanation, but it is not the endpoint. The video branch still needs more work if the goal is to make courtroom-style facial reasoning central rather than supplementary. The model also remains sensitive to neutral-heavy data behavior, which limits how far accuracy alone can be pushed without changing the input strategy or the fusion strategy.")
    add(r"For that reason, the project uses Phase 1 to establish the baseline and to identify where visual support is informative. The next direction is to push the visual path further, especially where the clips contain meaningful face information and where the current model still confuses emotionally close labels. The important point is that the remaining work is now targeted rather than exploratory.")
    add(r"This makes the current report suitable as a project milestone document. It records what is implemented, what is measured, what remains weak, and what the next step must focus on.")
    add(r"\section{Background}")
    add(r"The broader use case behind LegalMemoCMT is courtroom testimony analysis, where spoken content, delivery style, and facial cues can all affect the perceived emotion or intent of a witness or speaker. In that setting, a system needs to do more than output a label. It needs to preserve traceability so that the evidence behind a prediction can be reviewed later.")
    add(r"MELD is not a courtroom dataset, but it is a practical Phase 1 benchmark because it provides multi-party conversational context, overlapping emotions, and visible class imbalance. These properties make it a good test bed for the technical parts of the pipeline that will eventually be reused in testimony analysis.")
    add(r"The background problem is therefore one of transferability. The project must first prove that the conversational baseline and the visual feature path can be implemented in a reproducible way on a benchmark dataset. Only then does it become meaningful to adapt the same design to courtroom-style material, where the audio and video signals are often noisier and the emotional boundaries may be even more subtle.")
    add(r"\section{Problem Statement}")
    add(r"The core problem in Phase 1 is to build an explainable multimodal pipeline that can replicate the paper-aligned conversational result and also accept a visual branch derived from raw video. The pipeline must remain traceable from raw data to output metrics, and it must support inspection of errors rather than only headline accuracy.")
    add(r"The technical challenge is that the target labels are imbalanced, the conversational context is short, and the video path may contain either strong face cues or distracting background information. That means the model can learn a dominant safe class, such as neutral, while still failing to separate emotionally close labels. The project therefore needs an architecture, a training objective, and an evaluation strategy that make these failure modes visible.")
    add(r"A second part of the problem is reproducibility. The same sample must map to the same transcript, audio features, visual features, and label across training and analysis. If the data path is not stable, the model cannot be audited and the observed behavior cannot be trusted as a benchmark result.")
    add(r"\section{Literature Survey}")
    add(r"The base paper MemoCMT motivates the main baseline design by showing that cross-modal feature fusion can be used to combine text and audio representations through an attention-based transformer block [1]. This is the most direct architectural reference for the project in Phase 1.")
    add(r"MELD provides the benchmark environment and also illustrates why conversational emotion recognition is hard: many clips are short, the label distribution is skewed, and the same emotional category may appear in very different dialogue settings [2]. CREMA-D complements this by showing how multimodal emotion signals also appear in actor-based speech corpora [3].")
    add(r"BERT and HuBERT provide the pretrained text and audio encoders used in the baseline [4,5]. Vision Transformer is the visual backbone used to convert frame samples into embeddings [6]. The general attention design comes from Transformer-based sequence modeling [18], while weighted cross-entropy and focal loss provide standard tools for imbalance-aware optimization [9].")
    add(r"The broader emotion-recognition literature also shows that context-aware models such as DialogueRNN and DialogueGCN are useful in conversational settings [10,11]. Surveys on multimodal emotion recognition repeatedly note that multimodal gains are not automatic and that error analysis is often more informative than aggregate accuracy alone [13-17]. That observation is consistent with the project results, where the visual path helps some examples but does not universally beat the stronger conversational baseline.")
    add(r"\section{System Requirements Specification}")
    add(r"The system requirements for the project in Phase 1 separate the work into data handling, model execution, evaluation, and report generation. At the hardware level, a CUDA-capable GPU reduces the training time for the visual and multimodal runs, but the pipeline can also be inspected on CPU for small demo cases. At the software level, the implementation relies on Python, PyTorch, Hugging Face Transformers, NumPy, Pandas, scikit-learn, python-docx, ffmpeg, and the report-generation dependencies used to build the PDF and LaTeX outputs.")
    add(r"The functional requirements are simple to state but important to satisfy: the code must create manifests, maintain fold-safe splits, train the baseline and the ViT-supported variants, export per-sample predictions, produce confusion matrices, and support a raw-mp4 demo path. The non-functional requirements are traceability, reproducibility, and consistent sample-id alignment from preprocessing through inference.")
    add(r"\section{Proposed Methodology}")
    add(r"The proposed methodology follows a staged experimental design. First, the MELD fold manifests are built and validated. Second, the paper-aligned text-plus-audio baseline is trained with cross-modal fusion and MIN pooling. Third, the raw video clips are converted into sampled RGB frames and then into cached ViT embeddings. Fourth, the support path is attached through face-crop or gated variants, and then the effect of auxiliary loss is tested as a separate improvement step.")
    add(r"This methodology makes every change inspectable. If a score changes, the report can attribute the change to the loss function, the visual input strategy, or the fusion mechanism. That is important because the current project goal is not only to obtain a label but also to understand why the label changed for a specific clip.")
    add(r"\section{Implementation Details}")
    add(r"The implementation is organized as a set of scripts and modules. The data layer converts raw MELD information into manifests and cached feature files. The model layer uses the shared \texttt{src/train/train.py}, \texttt{src/models/model.py}, \texttt{src/data/preprocessing.py}, \texttt{src/train/evaluate.py}, and the supporting dataset utilities to keep the training and inference paths consistent. The analysis layer exports predictions, reports, confusion matrices, and top-confusion tables, while the demo layer runs a single mp4 through the same feature pipeline that the batch scripts use.")
    add(r"The most important implementation detail is that the same sample identifier must resolve to the same transcript, audio feature, and video feature across train, evaluation, and demo. That alignment is what makes the outputs auditable. Without it, the model could appear to predict an emotion from one clip while actually consuming another clip’s cached features.")
    add(r"\section{Intermediate Results and Discussion}")
    add(r"The intermediate results show that the paper-aligned baseline remains the best single reference point, which is consistent with the five-fold metrics and the fold-level confusion analysis. The visual branches do add information, but the gains are uneven. Some clips become easier to interpret when the face-crop path is active, while other clips remain stubbornly neutral-heavy or become confidently wrong.")
    add(r"The discussion therefore focuses on the error profile rather than on one accuracy number. Weighted F1 and macro F1 give a more realistic picture of behavior under imbalance, and the confusion matrices show that the model still struggles with emotionally close classes such as neutral, joy, anger, fear, and surprise. This is exactly the kind of behavior that has to be explained in review because it is where the model is most likely to generalize poorly.")
    add(r"\section{Conclusions and Future Work}")
    add(r"Phase 1 is complete enough to show a full and reproducible benchmark pipeline, but the report also makes clear that the facial-cue branch is still a support path rather than the dominant source of improvement. The strongest output remains the paper-aligned conversational baseline, and the visual experiments are valuable because they identify where facial information helps and where it does not.")
    add(r"Future work should focus on pushing the visual pathway toward the courtroom-testimony setting, especially by studying how face crops, gated fusion, auxiliary supervision, and additional domain-specific data affect the final decision boundary. The current Phase 1 results provide the reference line for that next stage.")
    add(r"\section{Comparative Metric Summary}")
    add(r"\begin{center}")
    add(r"\begin{tabular}{|p{4.2cm}|c|c|c|c|}")
    add(r"\hline")
    add(r"\textbf{Setting} & \textbf{Accuracy} & \textbf{Weighted F1} & \textbf{Macro F1} & \textbf{Unweighted Acc.} \\ \hline")
    for row in [
        ("Paper-aligned MELD mean", "0.6247", "0.6195", "0.4395", "0.4417"),
        ("Face-crop gated Fold 2", "0.6222", "0.6109", "0.4191", "0.4173"),
        ("Face-crop gated Fold 4", "0.5973", "0.5968", "0.4156", "0.4254"),
        ("Face-crop gated + aux Fold 2", "0.6054", "0.6022", "0.4351", "0.4492"),
        ("Face-crop gated + aux Fold 4", "0.5992", "0.6056", "0.4330", "0.4638"),
    ]:
        add(rf"{tex_escape(row[0])} & {row[1]} & {row[2]} & {row[3]} & {row[4]} \\ \hline")
    add(r"\end{tabular}")
    add(r"\end{center}")
    add(r"The table compares the main baseline with the visual-support branches so the result can be read as a sequence of controlled experiments rather than as isolated scores.")
    add(r"\section{Appendix A. Script List and Command Sequence}")
    add(r"The appendix lists the scripts used to reproduce the Phase 1 project workflow from training to analysis and demo. The exact sequence depends on whether the paper-aligned baseline, the ViT face-crop support path, or the gated plus auxiliary-loss variant is being reviewed.")
    add(r"1. Build or verify the paper-aligned MELD folds: bash scripts/run_paper_aligned_meld_cv.sh")
    add(r"2. Analyze Fold 2 predictions and confusion matrix: bash scripts/analyze_meld_vit_facecue_fold2.sh")
    add(r"3. Analyze Fold 4 predictions and confusion matrix: bash scripts/analyze_meld_vit_facecue_fold4.sh")
    add(r"4. Train the face-crop gated variant: bash scripts/run_meld_vit_facecrop_gated_fold2.sh and bash scripts/run_meld_vit_facecrop_gated_fold4.sh")
    add(r"5. Train the face-crop gated plus auxiliary-loss variant: bash scripts/run_meld_vit_facecrop_gated_video_aux_fold2.sh and bash scripts/run_meld_vit_facecrop_gated_video_aux_fold4.sh")
    add(r"6. Run the paper-aligned raw-mp4 demo: bash scripts/run_demo_paper_aligned_raw_mp4.sh <sample_id> <video_path>")
    add(r"7. Run the gated plus auxiliary-loss raw-mp4 demo: bash scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh <sample_id> <video_path>")
    add(r"8. Review the per-sample predictions, metrics.json, top-confusion table, and confusion matrix generated by the analysis scripts before presenting the results.")
    add(r"\begin{thebibliography}{99}")
    refs = [
        r"\bibitem{ref1} Khan, M., Tran, P.-N., Pham, N. T., El Saddik, A., and Othmani, A., ``MemoCMT: multimodal emotion recognition using cross-modal transformer-based feature fusion,'' \textit{Scientific Reports}, vol. 15, no. 1, 2025. \url{https://doi.org/10.1038/s41598-025-89202-x}.",
        r"\bibitem{ref2} Poria, S., Hazarika, D., Majumder, N., Naik, G., Cambria, E., and Mihalcea, R., ``MELD: A Multimodal Multi-Party Dataset for Emotion Recognition in Conversations,'' ACL, 2019. \url{https://aclanthology.org/P19-1050/}.",
        r"\bibitem{ref3} Cao, H., Cooper, D., Keutmann, M., Gur, R., Nenkova, A., and Verma, R., ``CREMA-D: Crowd-sourced Emotional Multimodal Actors Dataset,'' \textit{IEEE Transactions on Affective Computing}, 2014. \url{https://pmc.ncbi.nlm.nih.gov/articles/PMC4313618/}.",
        r"\bibitem{ref4} Devlin, J., Chang, M.-W., Lee, K., and Toutanova, K., ``BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding,'' 2019. \url{https://arxiv.org/abs/1810.04805}.",
        r"\bibitem{ref5} Hsu, W.-N., Bolte, B., Tsai, Y.-H. H., Lakhotia, K., Salakhutdinov, R., and Mohamed, A., ``HuBERT: Self-Supervised Speech Representation Learning by Masked Prediction of Hidden Units,'' 2021. \url{https://arxiv.org/abs/2106.07447}.",
        r"\bibitem{ref6} Dosovitskiy, A., Beyer, L., Kolesnikov, A., et al., ``An Image is Worth 16x16 Words: Transformers for Image Recognition at Scale,'' 2021. \url{https://arxiv.org/abs/2010.11929}.",
        r"\bibitem{ref7} Schneider, S., Baevski, A., Collobert, R., and Auli, M., ``wav2vec: Unsupervised Pre-training for Speech Recognition,'' 2019. \url{https://arxiv.org/abs/1904.05862}.",
        r"\bibitem{ref8} Baevski, A., Zhou, H., Mohamed, A., and Auli, M., ``wav2vec 2.0: A Framework for Self-Supervised Learning of Speech Representations,'' 2020. \url{https://arxiv.org/abs/2006.11477}.",
        r"\bibitem{ref9} Lin, T.-Y., Goyal, P., Girshick, R., He, K., and Doll\'ar, P., ``Focal Loss for Dense Object Detection,'' 2018. \url{https://arxiv.org/abs/1708.02002}.",
        r"\bibitem{ref10} Majumder, N., Poria, S., Hazarika, D., Mihalcea, R., Gelbukh, A., and Cambria, E., ``DialogueRNN: An Attentive RNN for Emotion Detection in Conversations,'' 2019. \url{https://arxiv.org/abs/1811.00405}.",
        r"\bibitem{ref11} Ghosal, D., Majumder, N., Poria, S., Chhaya, N., and Gelbukh, A., ``DialogueGCN: A Graph Convolutional Neural Network for Emotion Recognition in Conversation,'' 2020. \url{https://aclanthology.org/D19-1015/}.",
        r"\bibitem{ref12} Ghosh, S., Ramaneswaran, S., Tyagi, U., et al., ``M-MELD: A Multilingual Multi-Party Dataset for Emotion Recognition in Conversations,'' 2022. \url{https://arxiv.org/abs/2203.16799}.",
        r"\bibitem{ref13} Wu, C., Cai, Y., Liu, Y., Zhu, P., Xue, Y., Gong, Z., Hirschberg, J., and Ma, B., ``Multimodal Emotion Recognition in Conversations: A Survey of Methods, Trends, Challenges and Prospects,'' 2025. \url{https://arxiv.org/abs/2505.20511}.",
        r"\bibitem{ref14} Shou, Y., Meng, T., Ai, W., Yin, N., and Li, K., ``A Comprehensive Survey on Multi-modal Conversational Emotion Recognition with Deep Learning,'' 2023. \url{https://arxiv.org/abs/2312.05735}.",
        r"\bibitem{ref15} Zhang, X., et al., ``Emotion Recognition in Conversations: A Survey Focusing on Multimodal Approaches,'' \textit{Electronics}, 2024. \url{https://www.mdpi.com/2079-9292/12/22/4714}.",
        r"\bibitem{ref16} ``TelME: Teacher-leading Multimodal Fusion Network for Emotion Recognition in Conversation,'' 2024. \url{https://aclanthology.org/2024.naacl-long.5/}.",
        r"\bibitem{ref17} ``Multimodal Emotion Recognition Calibration in Conversations,'' 2024. \url{https://www.atailab.cn/seminar2024Fall/pdf/2024_MM_Multimodal%20Emotion%20Recognition%20Calibration%20in%20Conversations.pdf}.",
        r"\bibitem{ref18} Vaswani, A., Shazeer, N., Parmar, N., et al., ``Attention Is All You Need,'' NeurIPS 2017. \url{https://arxiv.org/abs/1706.03762}.",
    ]
    lines.extend(refs)
    add(r"\end{thebibliography}")
    add(r"\end{document}")
    TEX_OUT.write_text("\n".join(lines), encoding="utf-8")


def make_pdf(figs: dict[str, Path]) -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCenter",
        parent=styles["Title"],
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        spaceAfter=10,
    )
    subtitle_style = ParagraphStyle(
        "SubCenter",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        fontName="Helvetica",
        fontSize=11,
        leading=14,
        spaceAfter=8,
    )
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=14, leading=18, spaceBefore=10, spaceAfter=6)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12, leading=15, spaceBefore=8, spaceAfter=4)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.6, leading=13, alignment=TA_JUSTIFY)
    small = ParagraphStyle("Small", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.3, leading=10, alignment=TA_CENTER)

    def p(text: str, style=body):
        return Paragraph(text.replace("&", "&amp;"), style)

    story: list = []
    if LOGO_DST.exists():
        story.append(Image(str(LOGO_DST), width=2.2 * cm, height=2.2 * cm))
        story.append(Spacer(1, 0.3 * cm))
    story.extend([
        p("Master of Technology in AI and Machine Learning", title_style),
        p("UE20CS971 Project Phase-1", subtitle_style),
        p("LegalMemoCMT: An Explainable Multilingual Multimodal Emotional Cue Analysis Framework for Indian Courtroom Testimony Using Cross-Modal Transformers", ParagraphStyle("Title2", parent=title_style, fontSize=20, leading=24)),
        Spacer(1, 0.2 * cm),
        p("Submitted by:", subtitle_style),
        p("Rajesh Upadhyaya", subtitle_style),
        p("PES2PGE24DS200", subtitle_style),
        Spacer(1, 0.15 * cm),
        p("Under the guidance of:", subtitle_style),
        p("Ramesh Prakash Guledgudd", subtitle_style),
        p("PES University, Bengaluru", subtitle_style),
        p("Certificate", h1),
        p("This is to certify that the Phase 1 project entitled LegalMemoCMT: An Explainable Multilingual Multimodal Emotional Cue Analysis Framework for Indian Courtroom Testimony Using Cross-Modal Transformers is a bona fide work carried out under the guidance of Ramesh Prakash Guledgudd in partial fulfilment of the requirements for the award of the degree of Master of Technology in AI and Machine Learning."),
        PageBreak(),
        p("Declaration", h1),
        p("I hereby declare that the Phase 1 project work presented in this report is carried out as part of the LegalMemoCMT project under the prescribed academic guidance and that the submitted content reflects the implementation, analysis, and evaluation carried out in the repository."),
        PageBreak(),
        p("Acknowledgement", h1),
        p("I acknowledge the guidance, review support, and academic feedback received during the Phase 1 project. The work also benefited from the available benchmark datasets, pretrained encoders, and the project scripts that made the implementation traceable from raw input to final metrics."),
        PageBreak(),
        p("Abstract", h1),
        p(ABSTRACT_TEXT),
        PageBreak(),
        p("Contents", h1),
    ])
    for item in CONTENTS_ITEMS:
        story.append(p(item))
    story.append(Spacer(1, 0.15 * cm))
    story.append(PageBreak())

    def scaled_image_flowable(img_path: Path, width_cm: float):
        with PILImage.open(img_path) as im:
            w, h = im.size
        width = width_cm * cm
        height = width * h / w
        return Image(str(img_path), width=width, height=height)

    content = [
        ("1. Introduction", [
            "Emotion recognition in conversations is difficult because each utterance depends on dialogue context, speaker history, and a mix of textual, acoustic, and visual signals. MELD is a useful benchmark for this setting because it contains multi-party dialogue, seven emotion classes, and both emotion and sentiment annotations. The project therefore focuses on reproducing a stable paper-aligned baseline first, then extending the system with ViT-based facial cues so the raw video pathway can be inspected in the same framework.",
            "The main aim of Phase 1 is implementation fidelity and interpretability. The conversational baseline must be strong enough to serve as the reference model, while the visual extension should be understandable as a support path that can help explain success cases, near-misses, and confident errors.",
        ]),
        ("2. Related Work", [
            "MELD established a widely used multimodal conversational benchmark for emotion recognition in dialogue [2]. DialogueRNN and DialogueGCN showed that speaker context and conversational structure improve emotion prediction [10,11]. BERT and HuBERT provide pretrained language and speech representations that transfer well to downstream tasks [4,5]. Vision Transformer offers a strong visual backbone for frame-based facial feature extraction [6]. Focal loss and weighted cross-entropy remain standard choices for imbalance-aware learning [9].",
        ]),
        ("3. Materials and Methods", [
            "The project uses MELD as the primary benchmark. The data is organized through manifests so that each utterance retains its sample identifier, split, label, transcript, and media path. Fold generation is dialogue-aware, which prevents train-validation leakage across utterances from the same dialogue.",
            "The paper-aligned baseline uses BERT for the text stream, HuBERT for the audio stream, a cross-modal transformer to model their interaction, and MIN pooling for the final utterance representation. The ViT support path starts from a raw mp4 clip, samples RGB frames, applies face-crop or full-frame preprocessing, encodes the frames with a pretrained ViT, and stores the resulting features as a cached .npy file.",
        ]),
        ("4. Results", [
            "The paper-aligned MELD baseline is the strongest result in the project. Five-fold evaluation gives a mean accuracy of 0.6247, a mean weighted F1 of 0.6195, and a mean macro F1 of 0.4395. Fold 2 reaches 0.6375 accuracy and 0.6254 weighted F1, while Fold 4 reaches 0.6199 accuracy and 0.6194 weighted F1. The gap between weighted F1 and macro F1 shows that class imbalance remains a central challenge.",
            "The ViT support runs are informative but mixed. Face-crop gated Fold 2 reaches 0.6222 accuracy and 0.6109 weighted F1, while Fold 4 reaches 0.5973 accuracy and 0.5968 weighted F1. Adding auxiliary loss gives 0.6054 accuracy and 0.6022 weighted F1 on Fold 2, and 0.5992 accuracy and 0.6056 weighted F1 on Fold 4.",
        ]),
        ("5. Discussion", [
            "The error patterns explain why accuracy alone is not sufficient. The Fold 2 confusion matrix shows a neutral-heavy structure, with confusion between neutral and nearby emotions such as anger, joy, and surprise. The facial-cue runs can correct some neutral boundary cases, but they also remain confidently wrong on ambiguous clips, which means confidence is not a substitute for correctness.",
            "The project therefore treats the paper-aligned baseline as the main Phase 1 result and the ViT path as a supported extension. This keeps the report faithful to what the code currently demonstrates: a reproducible conversational model, a visual support branch, and a transparent analysis of where each one helps or fails.",
        ]),
        ("6. Limitations", [
            "The facial-cue branch is supplementary and does not consistently outperform the conversational baseline.",
            "MELD remains class imbalanced, so macro F1 stays substantially lower than weighted F1.",
            "Some clips are still confidently misclassified, especially when the transcript and visual evidence point in different directions.",
            "CREMA-D is supported as a secondary benchmark, but the project’s main replication story is MELD.",
        ]),
        ("7. Conclusion", [
            "Phase 1 now contains a reproducible paper-aligned MELD baseline and a supporting ViT facial-cue path. The baseline is strong enough to serve as the reference model, and the visual branch provides a practical route for future courtroom-style testimony adaptation. The main contribution of Phase 1 is therefore a complete, inspectable, and benchmarked multimodal pipeline rather than a claim of universal improvement from adding video.",
        ]),
        ("8. Project Objectives and Scope", [
            "The scope of the project in Phase 1 is deliberately narrow in one sense and broad in another. It is narrow because the main benchmark claim is tied to the paper-aligned MELD conversational setup, where the baseline architecture, the training path, and the evaluation strategy are kept close to the reference implementation. It is broad because the repository also carries the surrounding infrastructure needed to inspect manifests, raw media, cached embeddings, fold-level runs, error analysis, and demo behavior.",
            "The practical objective is to show that the full benchmark loop works end to end. That means raw dialogue data must be converted into stable manifests, the model must train without leakage between dialogue groups, the test results must be exported in a form that can be reviewed later, and the outcome must remain explainable through confusion matrices and per-sample predictions. This is why the project report emphasizes traceability from source data to final metrics.",
            "A second objective is support for later courtroom-testimony work. The Phase 1 project does not attempt to solve that later problem directly, but it creates a visual pathway that can be reused and stress-tested. The same input-output logic that maps a conversational utterance to an emotion label can later be reinterpreted for testimony-style clips, where facial cues and spoken content both matter.",
        ]),
        ("9. Dataset Organization, Label Space, and Fold Rules", [
            "MELD contains seven emotion classes and an imbalanced label distribution, with neutral appearing far more often than the minority classes. That imbalance is not a minor implementation detail; it directly changes what the model learns. If the training objective only optimizes the most frequent class, the model can appear strong on accuracy while still failing the classes that matter for nuanced emotion separation.",
            "The repository therefore organizes the dataset through explicit manifests rather than ad hoc file traversal. Each row links a sample identifier to its transcript, label, processed audio path, processed video path, and split assignment. This is important because it makes the data pipeline inspectable. If a result looks suspicious, the manifest can be used to trace whether the clip came from train, validation, or test, and whether any preprocessing artifact changed the signal.",
            "Fold construction is dialogue-aware. This means utterances from the same dialogue are kept together rather than scattered randomly across training and validation partitions. That choice matters because nearby utterances often share speaker identity, scene context, and even emotional trajectory. A random split would make the evaluation look artificially better by leaking conversational context into the held-out side.",
            "The same grouping logic is retained for the ViT support path, so the visual embeddings remain aligned with the same sample identifiers that the text and audio path use. This makes the visual branch a true support mechanism rather than a separate dataset pipeline with different sample ordering.",
        ]),
        ("10. Baseline Architecture and Modality Roles", [
            "The baseline model treats text as the strongest conversational channel and audio as the complementary paralinguistic channel. BERT converts each utterance transcript into a contextual token representation, while HuBERT converts the audio waveform into a sequence of speech features. The two streams are then fused through cross-modal attention so the classifier does not simply concatenate unrelated embeddings.",
            "Cross-modal transformer fusion is useful because emotion is rarely expressed in a single modality in a clean way. A short utterance may look neutral in text but carry frustration in tone, or vice versa. The fusion block gives the model a chance to reweight each stream according to the context of the current clip. The design is especially relevant in MELD, where short utterances and speaker turn changes are common.",
            "MIN pooling is used at the utterance aggregation stage. Rather than averaging every token or every frame in a uniform way, MIN pooling keeps a compact summary of the strongest aligned signals produced by the fused representation. In practice, this helps the model preserve the most informative parts of the utterance while still keeping the classifier lightweight enough to train on the available hardware.",
            "The classification head then maps the fused representation to the seven emotion classes. The output is not just a class label; it is also a confidence distribution over the labels. That distribution is useful for understanding whether the model is decisively correct, barely correct, or confidently wrong.",
        ]),
        ("11. Training and Evaluation Protocol", [
            "Training is organized around fold-by-fold checkpoints so that the reported result is not tied to a single lucky split. Each fold uses the training partition to learn parameters and the validation partition to decide which checkpoint is best. The project keeps that checkpointing logic explicit because the best epoch is not necessarily the last epoch. A later epoch can fit the training data more closely while performing worse on held-out examples.",
            "The evaluation loop exports accuracy, weighted F1, macro F1, unweighted accuracy, predictions, confusion matrices, and a top-confusion summary. Accuracy gives the simplest summary of the percentage of correct predictions. Weighted F1 compensates for class frequency by giving more weight to common labels. Macro F1 treats every class equally, which makes it much more sensitive to minority-class behavior. Unweighted accuracy is useful because it exposes how well the model performs when the class imbalance is not allowed to dominate the score.",
            "The result files are not just bookkeeping. They are the basis for error analysis. If a fold performs well on weighted F1 but poorly on macro F1, the model is likely over-reliant on frequent classes. If the confusion matrix is neutral-heavy, then the model has probably learned a generic safe prediction behavior rather than a robust emotion separator.",
            "This protocol is also why the project report emphasizes fold 2 and fold 4. Those folds are useful inspection points because they expose whether the behavior is stable or whether it changes in a way that depends on the dialogue partition.",
        ]),
        ("12. ViT Facial-Cue Extension and Video Caching", [
            "The ViT extension begins with raw mp4 input rather than with an already prepared feature tensor. The script samples RGB frames from the clip, optionally applies face-crop preprocessing, and sends the resulting images through a pretrained Vision Transformer. This converts a variable-length video into a compact fixed-shape representation that can be cached and reused.",
            "The cached .npy file is important because it separates feature generation from inference. Once the visual embeddings are saved, the model can reload them without repeating the full video decoding process every time. That makes experiments faster, easier to reproduce, and easier to inspect. It also means that later demo runs can focus on inference and explanation instead of repeating the expensive image encoding stage.",
            "Face-crop preprocessing narrows the signal to the speaker region. For emotion recognition and testimony analysis, that is often more informative than full-frame context because facial expression is usually the most direct visual cue for affect. Full-frame preprocessing still has value, especially when surrounding context carries meaningful evidence, but it also brings more background variation and more opportunity for noise.",
            "The ViT branch is not a replacement for the conversational model. It is a support path that reveals whether visual evidence changes the prediction boundary, whether it helps resolve neutral-heavy examples, and whether it still fails when the spoken content and the facial signal disagree.",
        ]),
        ("13. Comparative Results and Error Analysis", [
            "The baseline remains the strongest overall result in the current project. The paper-aligned MELD run still sets the main reference point, and the facial-cue runs should be read as supporting experiments that probe what happens when the visual path is introduced. The most useful comparison is not a single win/loss statement, but the pattern across accuracy, weighted F1, macro F1, and unweighted accuracy.",
            "The folded results show a consistent pattern. The face-crop gated run improves some clips, but it does not reliably exceed the conversational baseline. The auxiliary-loss variant helps the visual branch learn a little more structure, yet the final metrics still remain mixed. That suggests the visual features are adding signal, but not enough signal to overwhelm the strong text-plus-audio baseline.",
            "The confusion analysis shows why this happens. Neutral is the dominant class, and the model often confuses neutral with joy, anger, surprise, or fear. That is a classic imbalance pattern. The model is learning something useful because it often places the correct label among the top predictions, but it still has difficulty separating emotionally close classes when the evidence is sparse or contradictory.",
            "A confident wrong prediction is especially important to inspect. If the output probability is high but incorrect, then the issue is not simply uncertainty. It may indicate a strong but misplaced bias toward one class, or a preprocessing path that emphasized the wrong portion of the clip. For that reason, the confidence score must always be read together with the confusion matrix and the predicted-vs-actual table.",
        ]),
        ("14. Reproducibility, Scripts, and Outputs", [
            "The project is reproducible because the important actions are separated into named scripts. One set of scripts builds the MELD fold manifests. Another set trains the baseline or the ViT support variants. A third set analyzes predictions, exports confusion matrices, and summarizes errors. The demo scripts then read the saved checkpoints and run inference on a single raw mp4 clip or on a small set of curated clips.",
            "This separation is deliberate. It lets the project in Phase 1 keep the same model weights while changing only one ingredient at a time. That is the only reliable way to understand whether the change in outcome came from the loss function, the visual input, the fusion mechanism, or some accidental difference in the execution path.",
            "The output artifacts are also part of the deliverable. Metrics.json files capture the final numbers. Prediction CSVs keep the sample-level output. Confusion matrices show where the model fails in aggregate. Top-confusion tables identify the most frequent error pairs. These are the files that make the project auditable during review.",
        ]),
        ("15. Limitations and Phase 2 Direction", [
            "The current Phase 1 implementation is strong enough to support review and explanation, but it is not the endpoint. The video branch still needs more work if the goal is to make courtroom-style facial reasoning central rather than supplementary. The model also remains sensitive to neutral-heavy data behavior, which limits how far accuracy alone can be pushed without changing the input strategy or the fusion strategy.",
            "For that reason, the project uses Phase 1 to establish the baseline and to identify where visual support is informative. The next direction is to push the visual path further, especially where the clips contain meaningful face information and where the current model still confuses emotionally close labels. The important point is that the remaining work is now targeted rather than exploratory.",
            "This makes the current report suitable as a project milestone document. It records what is implemented, what is measured, what remains weak, and what the next step must focus on.",
        ]),
        ("16. Background", [
            "The broader use case behind LegalMemoCMT is courtroom testimony analysis, where spoken content, delivery style, and facial cues can all affect the perceived emotion or intent of a witness or speaker. In that setting, a system needs to do more than output a label. It needs to preserve traceability so that the evidence behind a prediction can be reviewed later.",
            "MELD is not a courtroom dataset, but it is a practical Phase 1 benchmark because it provides multi-party conversational context, overlapping emotions, and visible class imbalance. These properties make it a good test bed for the technical parts of the pipeline that will eventually be reused in testimony analysis.",
            "The background problem is therefore one of transferability. The project must first prove that the conversational baseline and the visual feature path can be implemented in a reproducible way on a benchmark dataset. Only then does it become meaningful to adapt the same design to courtroom-style material, where the audio and video signals are often noisier and the emotional boundaries may be even more subtle.",
        ]),
        ("17. Problem Statement", [
            "The core problem in Phase 1 is to build an explainable multimodal pipeline that can replicate the paper-aligned conversational result and also accept a visual branch derived from raw video. The pipeline must remain traceable from raw data to output metrics, and it must support inspection of errors rather than only headline accuracy.",
            "The technical challenge is that the target labels are imbalanced, the conversational context is short, and the video path may contain either strong face cues or distracting background information. That means the model can learn a dominant safe class, such as neutral, while still failing to separate emotionally close labels. The project therefore needs an architecture, a training objective, and an evaluation strategy that make these failure modes visible.",
            "A second part of the problem is reproducibility. The same sample must map to the same transcript, audio features, visual features, and label across training and analysis. If the data path is not stable, the model cannot be audited and the observed behavior cannot be trusted as a benchmark result.",
        ]),
        ("18. Literature Survey", [
            "The base paper MemoCMT motivates the main baseline design by showing that cross-modal feature fusion can be used to combine text and audio representations through an attention-based transformer block [1]. This is the most direct architectural reference for the project in Phase 1.",
            "MELD provides the benchmark environment and also illustrates why conversational emotion recognition is hard: many clips are short, the label distribution is skewed, and the same emotional category may appear in very different dialogue settings [2]. CREMA-D complements this by showing how multimodal emotion signals also appear in actor-based speech corpora [3].",
            "BERT and HuBERT provide the pretrained text and audio encoders used in the baseline [4,5]. Vision Transformer is the visual backbone used to convert frame samples into embeddings [6]. The general attention design comes from Transformer-based sequence modeling [18], while weighted cross-entropy and focal loss provide standard tools for imbalance-aware optimization [9].",
            "The broader emotion-recognition literature also shows that context-aware models such as DialogueRNN and DialogueGCN are useful in conversational settings [10,11]. Surveys on multimodal emotion recognition repeatedly note that multimodal gains are not automatic and that error analysis is often more informative than aggregate accuracy alone [13-17]. That observation is consistent with the project results, where the visual path helps some examples but does not universally beat the stronger conversational baseline.",
        ]),
        ("19. System Requirements Specification", [
            "The system requirements for the project in Phase 1 separate the work into data handling, model execution, evaluation, and report generation. At the hardware level, a CUDA-capable GPU reduces the training time for the visual and multimodal runs, but the pipeline can also be inspected on CPU for small demo cases. At the software level, the implementation relies on Python, PyTorch, Hugging Face Transformers, NumPy, Pandas, scikit-learn, python-docx, ffmpeg, and the report-generation dependencies used to build the PDF and LaTeX outputs.",
            "The functional requirements are simple to state but important to satisfy: the code must create manifests, maintain fold-safe splits, train the baseline and the ViT-supported variants, export per-sample predictions, produce confusion matrices, and support a raw-mp4 demo path. The non-functional requirements are traceability, reproducibility, and consistent sample-id alignment from preprocessing through inference.",
            "The main assumptions are that the manifests correctly map each utterance to one label, the raw media paths are valid, and the fold assignment keeps dialogue groups intact. The main constraint is that class imbalance remains present even after cross-validation, so metric selection must include macro F1 and unweighted accuracy rather than only accuracy.",
            "The system requirements section also makes clear which external pieces are assumed by the workflow. If a dependency such as ffmpeg, a Hugging Face model, or a cached video path is missing, the pipeline can fail before the model itself is even evaluated. That is why the report treats environment setup as part of the technical system rather than as a side note.",
        ]),
        ("20. Proposed Methodology", [
            "The proposed methodology follows a staged experimental design. First, the MELD fold manifests are built and validated. Second, the paper-aligned text-plus-audio baseline is trained with cross-modal fusion and MIN pooling. Third, the raw video clips are converted into sampled RGB frames and then into cached ViT embeddings. Fourth, the support path is attached through face-crop or gated variants, and then the effect of auxiliary loss is tested as a separate improvement step.",
            "This methodology makes every change inspectable. If a score changes, the report can attribute the change to the loss function, the visual input strategy, or the fusion mechanism. That is important because the current project goal is not only to obtain a label but also to understand why the label changed for a specific clip.",
            "The architecture is intentionally modular. Text, audio, and video can be studied independently or in combination, which makes it easier to isolate the effect of a design change. The same modularity also helps during review because a single change in the input path does not force a complete rewrite of the rest of the pipeline.",
            "The workflow is therefore not only a training recipe but also an explanation recipe. Each stage produces artifacts that can be checked separately: a manifest shows what data was used, a checkpoint shows what model was selected, a metrics file shows how well it performed, and a confusion matrix shows where it failed.",
        ]),
        ("21. Implementation Details", [
            "The implementation is organized as a set of scripts and modules. The data layer converts raw MELD information into manifests and cached feature files. The model layer uses the shared src/train/train.py, src/models/model.py, src/data/preprocessing.py, src/train/evaluate.py, and the supporting dataset utilities to keep the training and inference paths consistent. The analysis layer exports predictions, reports, confusion matrices, and top-confusion tables, while the demo layer runs a single mp4 through the same feature pipeline that the batch scripts use.",
            "The most important implementation detail is that the same sample identifier must resolve to the same transcript, audio feature, and video feature across train, evaluation, and demo. That alignment is what makes the outputs auditable. Without it, the model could appear to predict an emotion from one clip while actually consuming another clip’s cached features.",
        ]),
        ("22. Intermediate Results and Discussion", [
            "The intermediate results show that the paper-aligned baseline remains the best single reference point, which is consistent with the five-fold metrics and the fold-level confusion analysis. The visual branches do add information, but the gains are uneven. Some clips become easier to interpret when the face-crop path is active, while other clips remain stubbornly neutral-heavy or become confidently wrong.",
            "The discussion therefore focuses on the error profile rather than on one accuracy number. Weighted F1 and macro F1 give a more realistic picture of behavior under imbalance, and the confusion matrices show that the model still struggles with emotionally close classes such as neutral, joy, anger, fear, and surprise. This is exactly the kind of behavior that has to be explained in review because it is where the model is most likely to generalize poorly.",
            "The face-crop gated and auxiliary-loss runs add nuance. They show that the video pathway can change the prediction for certain clips, including some near-miss neutral cases, but they also show that the same pathway can produce confident wrong outputs when the visual evidence is ambiguous or when the facial signal is weak. The result is improvement in observability more than universal score improvement.",
        ]),
        ("23. Conclusions and Future Work", [
            "Phase 1 is complete enough to show a full and reproducible benchmark pipeline, but the report also makes clear that the facial-cue branch is still a support path rather than the dominant source of improvement. The strongest output remains the paper-aligned conversational baseline, and the visual experiments are valuable because they identify where facial information helps and where it does not.",
            "Future work should focus on pushing the visual pathway toward the courtroom-testimony setting, especially by studying how face crops, gated fusion, auxiliary supervision, and additional domain-specific data affect the final decision boundary. The current Phase 1 results provide the reference line for that next stage.",
        ]),
        ("Appendix A. Script List and Command Sequence", [
            "The appendix lists the scripts used to reproduce the Phase 1 project workflow from training to analysis and demo. The exact sequence depends on whether the paper-aligned baseline, the ViT face-crop support path, or the gated plus auxiliary-loss variant is being reviewed.",
            "1. Build or verify the paper-aligned MELD folds: bash scripts/run_paper_aligned_meld_cv.sh",
            "2. Analyze Fold 2 predictions and confusion matrix: bash scripts/analyze_meld_vit_facecue_fold2.sh",
            "3. Analyze Fold 4 predictions and confusion matrix: bash scripts/analyze_meld_vit_facecue_fold4.sh",
            "4. Train the face-crop gated variant: bash scripts/run_meld_vit_facecrop_gated_fold2.sh and bash scripts/run_meld_vit_facecrop_gated_fold4.sh",
            "5. Train the face-crop gated plus auxiliary-loss variant: bash scripts/run_meld_vit_facecrop_gated_video_aux_fold2.sh and bash scripts/run_meld_vit_facecrop_gated_video_aux_fold4.sh",
            "6. Run the paper-aligned raw-mp4 demo: bash scripts/run_demo_paper_aligned_raw_mp4.sh <sample_id> <video_path>",
            "7. Run the gated plus auxiliary-loss raw-mp4 demo: bash scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh <sample_id> <video_path>",
            "8. Review the per-sample predictions, metrics.json, top-confusion table, and confusion matrix generated by the analysis scripts before presenting the results.",
        ]),
    ]

    for head, paras in content:
        story.append(p(head, h1))
        for para in paras:
            story.append(p(para))
        story.append(Spacer(1, 0.15 * cm))

    metric_data = [
        ["Setting", "Accuracy", "Weighted F1", "Macro F1", "Unweighted Acc."],
        ["Paper-aligned MELD mean", "0.6247", "0.6195", "0.4395", "0.4417"],
        ["Face-crop gated Fold 2", "0.6222", "0.6109", "0.4191", "0.4173"],
        ["Face-crop gated Fold 4", "0.5973", "0.5968", "0.4156", "0.4254"],
        ["Face-crop gated + aux Fold 2", "0.6054", "0.6022", "0.4351", "0.4492"],
        ["Face-crop gated + aux Fold 4", "0.5992", "0.6056", "0.4330", "0.4638"],
    ]
    metric_table = Table(metric_data, colWidths=[5.0 * cm, 2.1 * cm, 2.2 * cm, 1.8 * cm, 2.5 * cm])
    metric_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E8F5")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("ALIGN", (1, 1), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
    ]))
    story.append(p("Comparative Metric Summary", h1))
    story.append(p("Table 1. Comparative Metric Summary.", small))
    story.append(metric_table)
    story.append(Spacer(1, 0.25 * cm))
    story.append(p("The table compares the main baseline with the visual-support branches so the result can be read as a sequence of controlled experiments rather than as isolated scores.", body))
    story.append(Spacer(1, 0.2 * cm))

    def img_with_cap(img_path: Path, cap_text: str, width_cm: float = 15.2):
        if img_path.exists():
            story.append(scaled_image_flowable(img_path, width_cm))
            story.append(p(cap_text, small))
            story.append(Spacer(1, 0.2 * cm))

    img_with_cap(figs["pipeline_png"], "Figure 1. Paper-aligned Phase 1 replication workflow.")
    img_with_cap(figs["vit_png"], "Figure 2. Raw-mp4 ViT facial-cue support path.")
    img_with_cap(FIG_DIR / "phase1_metric_comparison.png", "Figure 3. Phase 1 metric comparison across the baseline and ViT-supported runs.")
    img_with_cap(figs["fold2_conf"], "Figure 4. Fold 2 confusion matrix for the paper-aligned MELD baseline.", width_cm=13.8)

    story.append(p("References", h1))
    refs = [
        "1. Khan, M., Tran, P.-N., Pham, N. T., El Saddik, A., and Othmani, A. MemoCMT: multimodal emotion recognition using cross-modal transformer-based feature fusion. Scientific Reports, 15(1), 5473. https://doi.org/10.1038/s41598-025-89202-x",
        "2. Poria, S., Hazarika, D., Majumder, N., Naik, G., Cambria, E., and Mihalcea, R. MELD: A Multimodal Multi-Party Dataset for Emotion Recognition in Conversations. ACL 2019. https://aclanthology.org/P19-1050/",
        "3. Cao, H., Cooper, D., Keutmann, M., Gur, R., Nenkova, A., and Verma, R. CREMA-D: Crowd-sourced Emotional Multimodal Actors Dataset. IEEE Transactions on Affective Computing, 2014. https://pmc.ncbi.nlm.nih.gov/articles/PMC4313618/",
        "4. Devlin, J., Chang, M.-W., Lee, K., and Toutanova, K. BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. https://arxiv.org/abs/1810.04805",
        "5. Hsu, W.-N., Bolte, B., Tsai, Y.-H. H., Lakhotia, K., Salakhutdinov, R., and Mohamed, A. HuBERT: Self-Supervised Speech Representation Learning by Masked Prediction of Hidden Units. https://arxiv.org/abs/2106.07447",
        "6. Dosovitskiy, A., Beyer, L., Kolesnikov, A., et al. An Image is Worth 16x16 Words: Transformers for Image Recognition at Scale. https://arxiv.org/abs/2010.11929",
        "7. Schneider, S., Baevski, A., Collobert, R., and Auli, M. wav2vec: Unsupervised Pre-training for Speech Recognition. https://arxiv.org/abs/1904.05862",
        "8. Baevski, A., Zhou, H., Mohamed, A., and Auli, M. wav2vec 2.0: A Framework for Self-Supervised Learning of Speech Representations. https://arxiv.org/abs/2006.11477",
        "9. Lin, T.-Y., Goyal, P., Girshick, R., He, K., and Dollár, P. Focal Loss for Dense Object Detection. https://arxiv.org/abs/1708.02002",
        "10. Majumder, N., Poria, S., Hazarika, D., Mihalcea, R., Gelbukh, A., and Cambria, E. DialogueRNN: An Attentive RNN for Emotion Detection in Conversations. https://arxiv.org/abs/1811.00405",
        "11. Ghosal, D., Majumder, N., Poria, S., Chhaya, N., and Gelbukh, A. DialogueGCN: A Graph Convolutional Neural Network for Emotion Recognition in Conversation. https://aclanthology.org/D19-1015/",
        "12. Ghosh, S., Ramaneswaran, S., Tyagi, U., et al. M-MELD: A Multilingual Multi-Party Dataset for Emotion Recognition in Conversations. https://arxiv.org/abs/2203.16799",
        "13. Wu, C., Cai, Y., Liu, Y., Zhu, P., Xue, Y., Gong, Z., Hirschberg, J., and Ma, B. Multimodal Emotion Recognition in Conversations: A Survey of Methods, Trends, Challenges and Prospects. https://arxiv.org/abs/2505.20511",
        "14. Shou, Y., Meng, T., Ai, W., Yin, N., and Li, K. A Comprehensive Survey on Multi-modal Conversational Emotion Recognition with Deep Learning. https://arxiv.org/abs/2312.05735",
        "15. Zhang, X., et al. Emotion Recognition in Conversations: A Survey Focusing on Multimodal Approaches. Electronics, 12(22), 4714. https://www.mdpi.com/2079-9292/12/22/4714",
        "16. TelME: Teacher-leading Multimodal Fusion Network for Emotion Recognition in Conversation. https://aclanthology.org/2024.naacl-long.5/",
        "17. Multimodal Emotion Recognition Calibration in Conversations. 2024. https://www.atailab.cn/seminar2024Fall/pdf/2024_MM_Multimodal%20Emotion%20Recognition%20Calibration%20in%20Conversations.pdf",
        "18. Vaswani, A., Shazeer, N., Parmar, N., et al. Attention Is All You Need. https://arxiv.org/abs/1706.03762",
    ]
    for ref in refs:
        story.append(p(ref))

    doc = SimpleDocTemplate(str(PDF_OUT), pagesize=A4, rightMargin=1.6 * cm, leftMargin=1.6 * cm, topMargin=1.6 * cm, bottomMargin=1.6 * cm)

    def on_page(canvas, _doc):
        canvas.setFont("Helvetica", 9)
        canvas.drawRightString(A4[0] - 1.6 * cm, 1.0 * cm, f"{canvas.getPageNumber()}")

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)


def main() -> None:
    ensure_dirs()
    figs = build_figures()
    metric_chart()
    make_docx(figs)
    make_tex(figs)
    make_pdf(figs)
    print(f"Wrote {DOCX_OUT}")
    print(f"Wrote {TEX_OUT}")
    print(f"Wrote {PDF_OUT}")


if __name__ == "__main__":
    main()
