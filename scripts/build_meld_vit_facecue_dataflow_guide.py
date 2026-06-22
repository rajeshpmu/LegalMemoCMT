from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "implementation_docments" / "LegalMemoCMT_MELD_Facial_Cues_ViT_Dataflow_Student_Guide.docx"
FIG_DIR = ROOT / "implementation_docments" / "figures" / "meld_vit_facecue_dataflow"
FIG_DIR.mkdir(parents=True, exist_ok=True)

PIPELINE_SVG = FIG_DIR / "meld_vit_facecue_dataflow.svg"
PIPELINE_PNG = FIG_DIR / "meld_vit_facecue_dataflow.png"


def render_mermaid(code: str, svg_path: Path, png_path: Path) -> None:
    mmd_path = png_path.with_suffix(".mmd")
    mmd_path.write_text(code, encoding="utf-8")
    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(mmd_path), "-o", str(svg_path), "-b", "white"],
        check=True,
        cwd=str(ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(mmd_path), "-o", str(png_path), "-b", "white"],
        check=True,
        cwd=str(ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc: Document, code: str) -> None:
    for line in code.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(9.2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build_doc() -> Document:
    render_mermaid(
        """flowchart LR
  A["MELD raw annotations"] --> B["Find raw MP4 clip"]
  B --> C["Sample full video frames"]
  C --> D["Run pretrained ViT"]
  D --> E["Save facial embedding as .npy"]
  E --> F["Write meld_vit_facecue.csv"]
  F --> G["Split into 5 CV folds"]
  G --> H["Warm-start fold training"]
  H --> I["Validation + test evaluation"]
  I --> J["metrics.json + predictions_test.csv + confusion matrix"]
""",
        PIPELINE_SVG,
        PIPELINE_PNG,
    )

    doc = Document()
    configure(doc)

    t = doc.add_paragraph()
    r = t.add_run("LegalMemoCMT MELD Facial Cues ViT Dataflow Student Guide")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "Times New Roman"

    s = doc.add_paragraph()
    r = s.add_run(
        "A student-focused explanation of how the facial-cue implementation goes from raw MELD data to manifests, .npy embeddings, validation, and saved results."
    )
    r.italic = True
    r.font.size = Pt(13)
    r.font.name = "Times New Roman"

    add_para(
        doc,
        "This guide is meant to show the full implementation path in one place. It explains which datasets are used, what each manifest means, where the .npy files come from, how the fold splits work, and how training/evaluation produces the final result files.",
    )

    doc.add_heading("1. What datasets are used", level=1)
    add_para(
        doc,
        "The ViT facial-cue pipeline is built on the MELD dataset. The script reads the MELD annotation CSV files, finds the matching raw clips, samples frames, and converts the clips into ViT embeddings. The same MELD data is then split into cross-validation folds for training and validation.",
    )
    add_bullets(
        doc,
        [
            "Raw MELD annotations come from data/MELD/annotations/train_sent_emo.csv, dev_sent_emo.csv, and test_sent_emo.csv.",
            "Raw video clips come from data/MELD/raw/MELD.Raw/<split>/.../.mp4.",
            "The facial-cue features are written as .npy files under data/processed/MELD_VIT_FACECUE/.",
            "The main manifest is data/manifests/meld_vit_facecue.csv.",
            "The 5-fold CV manifests are under data/manifests/meld_vit_facecue_cv/.",
        ],
    )
    add_para(
        doc,
        "For students, the key idea is that MELD is not loaded as one giant tensor. It is turned into a manifest so that every utterance has a clear row containing the text, label, raw clip reference, and generated visual feature path.",
    )

    doc.add_heading("2. Raw data to facial-cue manifest", level=1)
    add_para(
        doc,
        "The manifest builder starts from the MELD annotation CSV files and converts them into a feature-oriented CSV. Each row keeps the original utterance information but adds a path to the extracted facial embedding .npy file.",
    )
    add_bullets(
        doc,
        [
            "build_meld_vit_facecue_manifest.py reads the MELD annotation files.",
            "find_meld_clip(...) resolves the raw video clip for each utterance.",
            "sample_video_frames(...) samples fixed-length frames from the clip.",
            "A pretrained ViT converts those frames into embeddings.",
            "The embeddings are saved as .npy files and the manifest points to them.",
        ],
    )
    add_para(
        doc,
        "The facial-cue manifest columns are sample_id, split, label, video_path, audio_path, and transcript. In this project, video_path points to the saved .npy facial embedding file, while audio_path still points to the raw MP4 clip used to find the utterance and keep the modality traceable.",
    )
    add_code(
        doc,
        """sample_id, split, label, video_path, audio_path, transcript
train_dia0_utt0, train, 0, data/processed/MELD_VIT_FACECUE/train/video/train_dia0_utt0.npy, data/MELD/raw/MELD.Raw/train/train_splits/dia0_utt0.mp4, also I was the point person...
train_dia0_utt1, train, 0, data/processed/MELD_VIT_FACECUE/train/video/train_dia0_utt1.npy, data/MELD/raw/MELD.Raw/train/train_splits/dia0_utt1.mp4, You must’ve had your hands full.""",
    )

    doc.add_heading("3. What the .npy files are", level=1)
    add_para(
        doc,
        "The .npy files store the facial embeddings produced by ViT. They are not raw images. They are numerical arrays, usually one embedding vector per sampled frame, that summarize the visual content of the utterance.",
    )
    add_bullets(
        doc,
        [
            "Each .npy file is a saved NumPy array.",
            "The array contains ViT CLS embeddings for the sampled frames.",
            "This lets training reuse the visual features without running ViT again every epoch.",
            "The stored representation is compact, fast to load, and reproducible.",
        ],
    )
    add_para(
        doc,
        "A student should understand the difference between raw video and a .npy feature file: the raw video is the original evidence, while the .npy file is the model-ready summary of that evidence.",
    )

    doc.add_heading("4. How the fold CSVs are created", level=1)
    add_para(
        doc,
        "The fold builder converts the raw facial-cue manifest into five cross-validation folds. It groups utterances by dialogue so that the same dialogue does not appear in both train and validation sets for the same fold.",
    )
    add_bullets(
        doc,
        [
            "The input is data/manifests/meld_vit_facecue.csv.",
            "The output is data/manifests/meld_vit_facecue_cv/meld_fold_0_train.csv through meld_fold_4_val.csv.",
            "The script also writes meld_fold_assignments.csv and meld_cv_summary.json.",
            "Fold 2 and Fold 4 are the selected folds for the current improvement experiments.",
        ],
    )
    add_code(
        doc,
        """python3 scripts/build_meld_cv_folds.py \\
  --manifest data/manifests/meld_vit_facecue.csv \\
  --output-dir data/manifests/meld_vit_facecue_cv \\
  --base-splits train,dev \\
  --num-folds 5""",
    )
    add_para(
        doc,
        "The purpose of the fold split is to create a fair validation setup. The model trains on four folds and validates on the held-out fold, then the next fold becomes the validation fold in turn.",
    )

    doc.add_heading("5. How the training script uses the manifests", level=1)
    add_para(
        doc,
        "The training script does not read the raw MELD annotation CSV directly during the fold runs. It reads the fold CSVs produced by the fold builder. Each row already contains the label, text, raw clip link, and the .npy feature path for the video branch.",
    )
    add_bullets(
        doc,
        [
            "scripts/run_meld_vit_facecue_fold2.sh calls scripts/resume_meld_vit_facecue_fold.py --fold 2.",
            "scripts/run_meld_vit_facecue_fold4.sh calls the same script with --fold 4.",
            "The script loads the weighted-CE checkpoint from results/paper_aligned_meld_cv/cmt_min/fold_<fold>/best_model.pt.",
            "It trains with text, audio, and video modalities enabled, then saves the best checkpoint and test predictions.",
        ],
    )
    add_para(
        doc,
        "Student explanation: the fold CSVs are the training and validation instructions. The model code simply follows those instructions and learns from the rows listed there.",
    )
    add_code(
        doc,
        """train_manifest = data/manifests/meld_vit_facecue_cv/meld_fold_2_train.csv
val_manifest   = data/manifests/meld_vit_facecue_cv/meld_fold_2_val.csv
raw_manifest   = data/manifests/meld_vit_facecue.csv
baseline_ckpt  = results/paper_aligned_meld_cv/cmt_min/fold_2/best_model.pt""",
    )

    doc.add_heading("6. What the training outputs are", level=1)
    add_para(
        doc,
        "After training, the fold output directory contains the model checkpoint and the metrics/predictions used for reporting. The evaluation step then turns the saved checkpoint into test results and analysis files.",
    )
    add_table(
        doc,
        ["File", "Meaning", "Why it matters"],
        [
            ["best_model.pt", "Best checkpoint saved during validation", "This is the model version you keep"],
            ["metrics.json", "Accuracy, weighted accuracy, unweighted accuracy, macro F1, weighted F1", "This is the main numerical summary"],
            ["predictions_test.csv", "Per-sample predictions on the test split", "This is used for confusion-matrix analysis"],
            ["analysis_test/confusion_matrix.csv", "Class-by-class confusion counts", "Shows which emotions are confused"],
        ],
    )
    add_para(
        doc,
        "For the student story, the important distinction is training versus evaluation: training decides which checkpoint is best, and evaluation uses that checkpoint on the test split to produce the final score and error analysis.",
    )

    doc.add_heading("7. How to read the results", level=1)
    add_para(
        doc,
        "The final result files are not the end of the pipeline; they are the evidence that lets you judge whether the facial-cue branch helped. The metrics tell you how well the model performed, while the confusion matrix tells you what kinds of mistakes it made.",
    )
    add_bullets(
        doc,
        [
            "Accuracy tells the overall fraction of correct predictions.",
            "Weighted F1 gives more influence to frequent classes.",
            "Macro F1 gives each class equal importance, which is important for imbalance analysis.",
            "The best validation checkpoint is the model version kept for test evaluation.",
        ],
    )
    add_para(
        doc,
        "If you want to describe the implementation simply: raw MELD clips become sampled frames, sampled frames become ViT embeddings, embeddings become .npy files, .npy files are listed in the manifest, the manifest is split into folds, and the folds are used for training, validation, and final test analysis.",
    )

    doc.add_heading("8. End-to-end summary for students", level=1)
    add_para(
        doc,
        "The full pipeline is: raw MELD data -> annotation CSV -> facial-cue manifest -> .npy feature files -> fold CSVs -> warm-start training -> validation checkpoint selection -> test evaluation -> metrics and confusion matrix. That is the complete student-level implementation story.",
        italic=True,
    )
    doc.add_picture(str(PIPELINE_PNG), width=Inches(6.7))

    return doc


def main() -> None:
    doc = build_doc()
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
