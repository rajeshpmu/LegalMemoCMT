from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path("/Users/rajeshpmu/Desktop/LegalMemoCMT")
OUT_PATH = ROOT / "implementation_docments/Phase1_Raw_MP4_Single_Clip_Live_Demo_SOP.docx"


def style_run(run, *, size=11, bold=False, color="000000"):
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text, *, size=11, bold=False, color="000000", align=None, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    style_run(r, size=14 if level == 1 else 12, bold=True, color="122F55")
    return p


def build_doc():
    doc = Document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Phase 1 Raw MP4 Single-Clip Live Demo SOP")
    style_run(r, size=18, bold=True, color="122F55")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Short live-demo guide for one clip at a time, without the reviewer bundle")
    style_run(r, size=11, color="5E5E5E")

    add_para(
        doc,
        "Purpose: this short SOP is for a simple live demo on Runpod where you process one raw .mp4 clip at a time, explain the output, and move on. It does not use the reviewer bundle. That makes the review easier to follow because the examiner sees input, prediction, confidence, and explanation for one clip before the next clip is shown.",
    )

    add_heading(doc, "1. What You Need Ready")
    for item in [
        "A raw MELD .mp4 clip from the approved demo list.",
        "The correct checkpoint path for the model you want to show.",
        "Runpod or another CUDA-capable environment if you want the faster path.",
        "The demo scripts and the raw clips on the same machine so the command can run directly.",
    ]:
        add_para(doc, f"• {item}")

    add_heading(doc, "2. Single-Clip Live Workflow")
    add_para(doc, "Use the same pattern for every clip:")
    add_para(doc, "1) Pick one clip.")
    add_para(doc, "2) Run the baseline command or the gated + aux command on that same clip.")
    add_para(doc, "3) Read the ground truth, predicted label, confidence, and top-3 probabilities.")
    add_para(doc, "4) Explain whether it is correct, a near-miss, or a confident wrong case.")
    add_para(doc, "5) Move to the next clip only after finishing the explanation for the current one.")

    add_heading(doc, "3. Exact Commands")
    add_para(doc, "Baseline paper-aligned run:")
    add_para(
        doc,
        "DEVICE=cuda bash scripts/run_demo_paper_aligned_raw_mp4.sh test_dia279_utt9 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia279_utt9.mp4",
        size=10,
        color="404040",
    )
    add_para(doc, "Gated + auxiliary-loss run:")
    add_para(
        doc,
        "DEVICE=cuda CHECKPOINT=results/facial_cues/meld_vit_facecrop_gated_video_aux/fold_4/best_model.pt bash scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh test_dia279_utt9 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia279_utt9.mp4",
        size=10,
        color="404040",
    )
    add_para(doc, "Change only the sample_id and mp4 path when you move to another approved clip. Keep the same clip for the baseline and gated runs so the comparison is fair.")

    add_heading(doc, "4. Approved One-Clip Demo Set")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Clip"
    hdr[1].text = "Category"
    hdr[2].text = "Why show it"
    hdr[3].text = "Live explanation"
    rows = [
        ("test_dia279_utt9", "Correct neutral / near-miss baseline", "Shows the baseline can be close but still wrong, while the later gated model can correct it.", "Use it to explain boundary behavior and confidence."),
        ("test_dia4_utt6", "Correct non-neutral", "Shows the model is not only predicting neutral.", "Use it as a clean success case."),
        ("test_dia278_utt5", "Confident wrong", "Shows a strong wrong prediction with high softmax confidence.", "Use it to explain modal bias, label ambiguity, and sampling limits."),
        ("test_dia153_utt5", "Hard neutral failure", "Shows that neutral can still fail badly in the baseline and can remain difficult later.", "Use it to show that improvement is selective, not universal."),
        ("test_dia244_utt14", "Backup neutral success", "Useful as a spare neutral clip if you want a second clean success case.", "Use it only if you need an extra safe example."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text

    add_heading(doc, "5. What to Say for Every Clip")
    for item in [
        "Ground truth: the real MELD label from the manifest or output printout.",
        "Predicted label: the class with the highest softmax probability.",
        "Confidence: the top-1 softmax score, which means certainty, not correctness.",
        "Top-3: the three highest softmax probabilities for that sample. They show how close the alternatives were.",
    ]:
        add_para(doc, f"• {item}")

    add_heading(doc, "6. How to Explain Wrong Predictions")
    for item in [
        "Modal bias: the transcript or audio may push the model toward the wrong class.",
        "Label ambiguity: the emotion can be genuinely hard to separate in MELD context.",
        "Frame sampling: the sampled frames may miss the exact facial peak.",
        "Fusion behavior: one strong modality can overrule the others.",
    ]:
        add_para(doc, f"• {item}")
    add_para(doc, "If the reviewer asks where to inspect the issue in code, mention src/data/preprocessing.py for frame sampling and video feature creation, src/models/model.py for fusion, src/train/train.py for checkpoint selection and loss, and src/train/evaluate.py for the confusion-matrix output.")

    add_heading(doc, "7. What Not to Use")
    for item in [
        "Do not use the reviewer bundle if the goal is a live one-clip explanation.",
        "Do not batch too many clips together in one command if you want a clean viva flow.",
        "Do not claim that the gated model is always better; explain improvement and failure honestly.",
    ]:
        add_para(doc, f"• {item}")

    add_heading(doc, "8. Short Closing Line")
    add_para(
        doc,
        "A good closing sentence is: Phase 1 is now a reproducible raw-video demo pipeline. I can show one clip at a time, explain the output technically, and describe why a prediction is correct or wrong using the transcript, the sampled frames, the fusion behavior, and the confusion analysis.",
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build_doc()
