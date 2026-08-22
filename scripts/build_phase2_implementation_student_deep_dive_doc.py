from __future__ import annotations

import csv
import json
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_Phase2_Implementation_Student_Deep_Dive.docx"
FIG_DIR = ROOT / "implementation_docments" / "phase2_student_deep_dive_figures"
DIAGRAM_MMD = FIG_DIR / "phase2_indian_scope_architecture.mmd"
DIAGRAM_SVG = FIG_DIR / "phase2_indian_scope_architecture.svg"
DIAGRAM_PNG = FIG_DIR / "phase2_indian_scope_architecture.png"
INDIA_DIAGRAM_MMD = FIG_DIR / "phase2_indian_acquisition_pack.mmd"
INDIA_DIAGRAM_SVG = FIG_DIR / "phase2_indian_acquisition_pack.svg"
INDIA_DIAGRAM_PNG = FIG_DIR / "phase2_indian_acquisition_pack.png"
CANDIDATE_LEDGER = ROOT / "data" / "phase2" / "source_manifests" / "case_candidate_ledger.csv"
VALIDATED_ROOT = ROOT / "data" / "processed" / "phase2" / "legalmeld_validated"
SUMMARY_JSON = VALIDATED_ROOT / "dataset_quality_summary.json"
LEGACY_SUMMARY_JSON = VALIDATED_ROOT / "legalmeld_dataset_summary.json"
VERIFIED_INVENTORY = ROOT / "data" / "processed" / "phase2" / "verified_case_inventory.csv"
HEARING_MANIFEST = ROOT / "data" / "processed" / "phase2" / "hearing_manifest.csv"
WITNESS_MANIFEST = ROOT / "data" / "phase2" / "source_manifests" / "witness_harvest_manifest_resolved.csv"
LEGALMELD_METADATA = VALIDATED_ROOT / "legalmeld_metadata_validated.csv"
TRAIN_CSV = VALIDATED_ROOT / "train.csv"
DEV_CSV = VALIDATED_ROOT / "dev.csv"
TEST_CSV = VALIDATED_ROOT / "test.csv"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)


def add_pipeline_order_table(doc: Document) -> None:
    add_table(
        doc,
        ["Stage", "Produced by", "Used for", "Student-level meaning"],
        [
            [
                "1. case_candidate_ledger.csv",
                "Manual curation",
                "Starting candidate list",
                "The first shortlist of cases to inspect; it is not proof of usable media yet.",
            ],
            [
                "2. case_candidate_ledger_ucr_enriched.csv",
                "`phase2/enrich_case_ledger_from_ucr_site.py` via `phase2/run_enrich_case_ledger_from_ucr_site.sh`",
                "Verification layer",
                "Checks whether the UCR page really resolves to case-specific records before anything is promoted.",
            ],
            [
                "3. verified_case_inventory.csv",
                "`phase2/build_ucr_case_inventory.py` via `phase2/run_build_ucr_case_inventory.sh`",
                "Canonical case inventory",
                "Keeps only cases with grounded metadata and removes placeholder-style false positives.",
            ],
            [
                "4. hearing_manifest.csv",
                "`phase2/build_hearing_manifest.py` via `phase2/run_build_hearing_witness_manifests.sh`",
                "Hearing-level grouping",
                "Turns verified cases into grounded hearing rows with transcript and video pairing evidence.",
            ],
            [
                "5. witness_harvest_manifest_resolved.csv",
                "`phase2/build_witness_manifest.py` via `phase2/run_build_hearing_witness_manifests.sh`",
                "Resolved witness layer",
                "Extracts witness or protected-code information without treating placeholder rows as corpus data.",
            ],
            [
                "6. legalmeld_metadata_validated.csv",
                "`phase2/build_legalmeld_dataset.py` via `phase2/run_build_legalmeld_dataset_validated.sh`",
                "Utterance-level export",
                "Creates the MELD-style row for each aligned utterance with audio, video, and transcript metadata.",
            ],
            [
                "7. alignment_review_sample.csv",
                "`phase2/build_legalmeld_dataset.py`",
                "Manual inspection sample",
                "Shows which utterances need review so the pipeline can be improved before scaling.",
            ],
            [
                "8. dataset_quality_summary.json",
                "`phase2/build_legalmeld_dataset.py`",
                "Quality summary",
                "Summarizes how many rows passed, how many were rejected, and how strong the alignment was.",
            ],
            [
                "9. train.csv / dev.csv / test.csv",
                "`phase2/build_legalmeld_dataset.py`",
                "Final split files",
                "Creates leakage-aware train, development, and test partitions for model work.",
            ],
        ],
    )


def render_mermaid(code: str) -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    DIAGRAM_MMD.write_text(code.strip() + "\n")
    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(DIAGRAM_MMD), "-o", str(DIAGRAM_SVG), "-b", "white"],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.PIPE,
    )
    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(DIAGRAM_MMD), "-o", str(DIAGRAM_PNG), "-b", "white"],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.PIPE,
    )


def render_india_mermaid(code: str) -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    INDIA_DIAGRAM_MMD.write_text(code.strip() + "\n")
    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(INDIA_DIAGRAM_MMD), "-o", str(INDIA_DIAGRAM_SVG), "-b", "white"],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.PIPE,
    )
    subprocess.run(
        ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(INDIA_DIAGRAM_MMD), "-o", str(INDIA_DIAGRAM_PNG), "-b", "white"],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.PIPE,
    )


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(6.9))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10.5)


def read_json(path: Path) -> dict[str, object]:
    if not path.exists():
        return {}
    return json.loads(path.read_text())


def count_rows(path: Path) -> int:
    if not path.exists():
        return 0
    with path.open(newline="", encoding="utf-8-sig") as f:
        return sum(1 for _ in csv.DictReader(f))


def fmt_int(value: object) -> str:
    try:
        return f"{int(value):,}"
    except Exception:
        return "0"


def fmt_float(value: object, digits: int = 3) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return f"{0.0:.{digits}f}"


def load_candidate_sources(limit: int = 5) -> list[list[str]]:
    if not CANDIDATE_LEDGER.exists():
        return []
    rows: list[list[str]] = []
    seen: set[tuple[str, str]] = set()
    with CANDIDATE_LEDGER.open(newline="", encoding="utf-8-sig") as f:
        for row in csv.DictReader(f):
            source_url = (row.get("source_url") or "").strip()
            inventory_url = (row.get("inventory_search_url") or "").strip()
            if not source_url and not inventory_url:
                continue
            key = (source_url, inventory_url)
            if key in seen:
                continue
            seen.add(key)
            rows.append(
                [
                    row.get("tribunal", ""),
                    row.get("case_family", ""),
                    row.get("case_number", ""),
                    source_url,
                    inventory_url,
                ]
            )
            if len(rows) >= limit:
                break
    return rows


def snapshot() -> dict[str, str]:
    data = read_json(SUMMARY_JSON)
    legacy = read_json(LEGACY_SUMMARY_JSON)
    out = {
        "verified_cases": fmt_int(data.get("cases_represented", 0)),
        "hearings_represented": fmt_int(data.get("hearings_represented", 0)),
        "utterances": fmt_int(data.get("total_utterances", data.get("utterances_exported", 0))),
        "total_clip_hours": fmt_float(data.get("total_clip_hours", 0.0)),
        "high_confidence": fmt_int(data.get("high_confidence_alignments", 0)),
        "medium_confidence": fmt_int(data.get("alignment_confidence_counts", {}).get("MEDIUM", 0) if isinstance(data.get("alignment_confidence_counts"), dict) else 0),
        "low_confidence": fmt_int(data.get("alignment_confidence_counts", {}).get("LOW", 0) if isinstance(data.get("alignment_confidence_counts"), dict) else 0),
        "quality_b": fmt_int(data.get("quality_tier_counts", {}).get("B", 0) if isinstance(data.get("quality_tier_counts"), dict) else 0),
        "quality_reject": fmt_int(data.get("quality_tier_counts", {}).get("REJECT", 0) if isinstance(data.get("quality_tier_counts"), dict) else 0),
        "train_rows": fmt_int(count_rows(TRAIN_CSV)),
        "dev_rows": fmt_int(count_rows(DEV_CSV)),
        "test_rows": fmt_int(count_rows(TEST_CSV)),
        "metadata_rows": fmt_int(count_rows(LEGALMELD_METADATA)),
        "inventory_rows": fmt_int(count_rows(VERIFIED_INVENTORY)),
        "hearing_manifest_rows": fmt_int(count_rows(HEARING_MANIFEST)),
        "witness_manifest_rows": fmt_int(count_rows(WITNESS_MANIFEST)),
        "legacy_summary_rows": fmt_int(legacy.get("total_utterances", 0)),
        "split_leakage": fmt_int(data.get("split_leakage_violations", 0)),
        "distinct_witnesses": fmt_int(data.get("distinct_witnesses", 0)),
        "witness_visible": fmt_int(data.get("witness_visible_count", 0)),
        "audio_valid": fmt_int(data.get("audio_valid_count", 0)),
        "video_valid": fmt_int(data.get("video_valid_count", 0)),
        "manual_review": fmt_int(data.get("manual_review_count", 0)),
        "cases_counted": fmt_int(data.get("cases_represented", 0)),
        "hearings_resolved": fmt_int(data.get("hearings_resolved", 0)),
    }
    return out


def build_doc() -> Document:
    s = snapshot()
    doc = Document()
    configure(doc)
    render_mermaid(
        """
flowchart LR
    A[Phase 1 multimodal baseline] --> B[Tribunal bootstrap corpus]
    B --> C[Corrected UCR enrichment]
    C --> D[verified_case_inventory.csv]
    D --> E[Hearing manifest builder]
    D --> F[Witness manifest builder]
    E --> G[Transcript-first hearing selection]
    F --> G
    G --> H[Transcript parser and speaker-turn segmentation]
    H --> I[Forced alignment with ASR word timestamps]
    I --> J[Utterance timestamps]
    J --> K[ffmpeg clip extraction]
    K --> L[Audio validation]
    K --> M[Video validation]
    L --> N[legalmeld_metadata_validated.csv]
    M --> N
    N --> O[Group-based train/dev/test split]
    O --> P[Tribunal bootstrap ready]

    P --> Q[Indian corpus acquisition pack]
    Q --> R1[Supreme Court livestreams]
    Q --> R2[High Court livestreams]
    Q --> R3[District courts]
    Q --> R4[Mock trials and judicial academies]
    R1 --> S[Indian candidate ledger]
    R2 --> S
    R3 --> S
    R4 --> S
    S --> T[Indian verification layer]
    T --> U[Indian hearing manifest]
    U --> V[Indian witness manifest]
    V --> W[Indian alignment manifest]
    W --> X[Indian courtroom adaptation corpus]

    C --> C1[Negative controls]
    C1 --> C2[Placeholder rejection]
    D --> D1[Positive controls]
    D1 --> D2[Verified real case records]
""",
    )

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LegalMemoCMT Phase 2: Multilingual Multimodal Emotion Analysis for Indian Courtroom Testimony")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "A technically grounded but student-friendly report on the tribunal bootstrap pipeline, the Indian courtroom adaptation strategy, and the current readiness for fine-tuning."
    )
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12.5)

    add_para(
        doc,
        "This document is written as a bridge from Phase 1 into Phase 2. Phase 1 established the multimodal backbone. Phase 2 now uses that backbone to build a tribunal bootstrap corpus and then adapt the same pipeline toward Indian courtroom testimony. The report explains why Phase 2 is a methodological novelty, how the proposed system works, what has already been implemented in code, and what still needs to be completed before the Indian scope is ready for serious fine-tuning.",
    )

    doc.add_heading("Source Sites And Download Flow", level=1)
    add_para(
        doc,
        "The candidate ledger starts from official tribunal sources and official UCR inventory pages. In this project, crawling means visiting the public source pages or case pages, searching means checking the case number or case family inside the UCR inventory, and downloading means pulling the actual transcript or video asset only after the record is verified.",
    )
    add_table(
        doc,
        ["Tribunal", "Case family", "Case number", "Source site link", "Inventory/search link"],
        load_candidate_sources(),
    )
    add_numbered(
        doc,
        [
            "Start from the official source page or case-page link in the candidate ledger.",
            "Search the UCR inventory by case number, case family, or witness code.",
            "Verify that the case resolves to actual record metadata, not a generic page shell.",
            "Promote only grounded records into the verified inventory and then into the manifest builders.",
            "Download media only after the row is verified and the record points to real transcript or video assets.",
        ],
    )
    add_bullets(
        doc,
        [
            "Generic headings are not proof of records.",
            "Placeholder manifests are not corpus evidence.",
            "Page shells must never be used as download estimates.",
        ],
    )
    add_para(
        doc,
        "The implementation is split across a small set of scripts and modules. That is deliberate: each step has a narrow responsibility, and the docs now point to the exact code that performs the work.",
    )

    doc.add_heading("Pipeline Order", level=1)
    add_para(
        doc,
        "This table shows the exact file order the student should understand first. The idea is simple: the ledger starts the workflow, the verified inventory proves which cases are real, the manifests organize hearings and witnesses, and the LegalMELD exporter produces the utterance-level dataset.",
    )
    add_pipeline_order_table(doc)

    doc.add_heading("Code Map", level=1)
    add_table(
        doc,
        ["Implementation step", "Script / module / function"],
        [
            [
                "Reject placeholders and verify case records",
                "`phase2/enrich_case_ledger_from_ucr_site.py` -> `resolve_case_verification()` and `validation_summary()`; wrapper: `phase2/run_enrich_case_ledger_from_ucr_site.sh`.",
            ],
            [
                "Promote verified cases into the inventory",
                "`phase2/build_ucr_case_inventory.py` -> `_eligible_rows()`, `_record_identity_key()`, and `resolve_case_verification()`; wrapper: `phase2/run_build_ucr_case_inventory.sh`.",
            ],
            [
                "Build hearing rows from grounded records",
                "`phase2/build_hearing_manifest.py` with shared helpers from `phase2/hearing_witness_manifest_utils.py`; wrapper: `phase2/run_build_hearing_witness_manifests.sh`.",
            ],
            [
                "Build witness rows without deanonymizing protected codes",
                "`phase2/build_witness_manifest.py` plus `extract_witness_identity()` from `phase2/hearing_witness_manifest_utils.py`.",
            ],
            [
                "Align transcript utterances and generate clips",
                "`phase2/build_legalmeld_dataset.py` -> `segment_transcript()`, `_align_utterances()`, `_extract_segment()`, `_audio_metrics()`, `_resolve_hearing_rows()`; wrapper: `phase2/run_build_legalmeld_dataset_validated.sh`.",
            ],
            [
                "Validate the tri-modal subset and split safely",
                "`phase2/validate_trimodal_hearings.py`, `phase2/resolve_witnesses_from_transcripts.py`, `phase2/estimate_utterance_counts.py`, and `phase2/common.py` utilities such as `ensure_dir()` and `write_csv()`.",
            ],
            [
                "Bootstrap on tribunal data and reuse the same logic for India",
                "Current implementation covers the tribunal bootstrap layer; the Indian acquisition pack is the planned reuse layer for Supreme Court, High Court, and mock-trial sources.",
            ],
        ],
    )

    doc.add_heading("Quick Status Snapshot", level=1)
    add_table(
        doc,
        ["Metric", "Current validated value"],
        [
            ["Verified cases", s["verified_cases"]],
            ["Hearings represented", s["hearings_represented"]],
            ["Utterances exported", s["utterances"]],
            ["Total clip hours", s["total_clip_hours"]],
            ["High-confidence alignments", s["high_confidence"]],
            ["Medium-confidence alignments", s["medium_confidence"]],
            ["Low-confidence alignments", s["low_confidence"]],
            ["Quality tier B", s["quality_b"]],
            ["Quality tier REJECT", s["quality_reject"]],
            ["Train / Dev / Test rows", f'{s["train_rows"]} / {s["dev_rows"]} / {s["test_rows"]}'],
            ["Split leakage violations", s["split_leakage"]],
            ["Distinct witnesses", s["distinct_witnesses"]],
            ["Witness-visible clips", s["witness_visible"]],
            ["Audio-valid clips", s["audio_valid"]],
        ],
    )

    doc.add_heading("1. Phase 1 Review", level=1)
    add_para(
        doc,
        "Phase 1 should be understood as the multimodal baseline stage. It established that LegalMemoCMT can combine text, audio, and video evidence instead of relying on one modality only. In the earlier workflow, the project focused on benchmark-style multimodal emotion recognition, where a hearing or clip could be treated as a single training example. That is useful, but it is still coarse.",
    )
    add_para(
        doc,
        "The main limitation of a Phase 1-style setup is that it does not yet reflect how courtroom data is actually spoken. A full hearing can last a long time, many speakers appear in the same recording, and only some moments are useful for model training. If the model sees only long recordings, it learns from noisy context, not from clean utterance-level samples. Phase 2 is the response to that limitation.",
    )
    add_bullets(
        doc,
        [
            "Phase 1 proves the project can work in a multimodal setting.",
            "Phase 1 also shows the need for finer-grained units than whole hearings or long clips.",
            "Phase 2 shifts the focus from benchmark-style inputs to grounded legal utterances.",
            "Phase 2 also adds stronger corpus validation and traceability than a simple download pipeline.",
        ],
    )

    doc.add_heading("2. Objective and Novelty of Phase 2", level=1)
    add_para(
        doc,
        "The objective of Phase 2 is to support LegalMemoCMT's original Indian courtroom testimony goal by building a MELD-style utterance corpus and a reproducible adaptation pipeline. In MELD, each sample is one utterance with aligned text, audio, and video. Phase 2 builds the legal-domain equivalent of that idea: one courtroom utterance, one transcript span, one audio segment, and one video segment. Because public Indian courtroom testimony is hard to obtain at scale, the tribunal corpus is used as a supervised bootstrap set, and the same manifest-and-alignment design is then reused for Indian sources.",
    )
    add_para(
        doc,
        "The novelty is not that the model architecture is new. The novelty is the data construction method and the adaptation strategy. Phase 2 adds a grounded pipeline that rejects placeholder cases, verifies actual UCR record metadata, builds hearing and witness manifests only from real records, and then exports utterance-level training samples only when transcript alignment and media validation succeed. At the same time, it creates a reusable preparation pattern that can later be applied to Indian sources such as Supreme Court livestreams, High Court archives, district-court material, and mock-trial material.",
    )
    add_bullets(
        doc,
        [
            "Novelty 1: grounded UCR enrichment instead of page-shell heuristics.",
            "Novelty 2: verified inventory and manifest generation based on actual record metadata.",
            "Novelty 3: transcript-first utterance segmentation with alignment to media.",
            "Novelty 4: quality-tiered tri-modal export with leakage-aware splits.",
            "Novelty 5: legal-domain dataset preparation that is reproducible and rerunnable.",
            "Novelty 6: a practical bootstrap path from international tribunal testimony to Indian courtroom adaptation.",
        ],
    )

    doc.add_heading("3. Title and Abstract for Phase 2", level=1)
    add_para(doc, "Proposed title:", bold=True)
    add_para(
        doc,
        "LegalMemoCMT Phase 2: Multilingual Multimodal Emotion Analysis for Indian Courtroom Testimony",
    )
    add_para(doc, "Proposed abstract:", bold=True)
    add_para(
        doc,
        "Phase 2 extends LegalMemoCMT from a general multimodal baseline into a legal-domain dataset construction and adaptation pipeline for Indian courtroom testimony. The immediate bootstrap corpus is built from public tribunal hearings, because those are the best available large-scale multimodal testimony records. The pipeline starts from a curated case ledger, verifies UCR records using actual case-specific metadata, builds hearing and witness manifests from grounded evidence, and then parses transcripts into speaker turns that can be aligned to audio/video timestamps. Each exported sample is validated for alignment confidence, audio quality, video availability, and split leakage risk before it is written to the final dataset. The resulting corpus is designed for courtroom emotion recognition, credibility analysis, and later multi-task fine-tuning. Phase 2 is therefore both a dataset engineering contribution and the bridge between Phase 1 multimodal modeling and the Indian courtroom adaptation objective.",
    )

    doc.add_heading("4. Proposed System for Phase 2", level=1)
    add_para(
        doc,
        "The proposed system is a pipeline, not a single model. It starts with curated source rows and ends with utterance-level training clips. The important design choice is that every stage preserves traceability back to a tribunal case number, a hearing record, and eventually an Indian source manifest row.",
    )
    add_para(
        doc,
        "From a project-design point of view, the system has two layers. The first layer is the tribunal bootstrap layer, which produces the grounded utterance corpus. The second layer is the Indian adaptation layer, which reuses the same manifest-and-alignment logic to process Indian legal speech where transcript, audio, and video are available. That second layer is important because the original research problem is Indian courtroom testimony, not only international tribunal testimony.",
    )
    add_table(
        doc,
        ["Stage", "What happens", "Why it matters"],
        [
            ["Curated ledger", "Keep the human-curated candidate list.", "This is the starting point and the source of truth for what should be checked."],
            ["UCR correction", "Reject placeholders and validate real records.", "This prevents false positives from generic page shells."],
            ["Verified inventory", "Keep only records with grounded case metadata.", "This becomes the canonical case-level source."],
            ["Hearing manifest", "Group records into real hearing/session units.", "This creates the bridge from case rows to hearing rows."],
            ["Witness manifest", "Resolve witness or code-level rows where possible.", "This lets later steps reason about who spoke."],
            ["Transcript-first alignment", "Segment, align, and export utterances.", "This is the MELD-style dataset building step."],
            ["Validation and splits", "Check audio/video quality and avoid leakage.", "This protects model training from noisy or duplicated samples."],
        ],
    )

    doc.add_heading("5. Proposed Mermaid Architecture Diagram", level=1)
    add_para(
        doc,
        "The diagram below shows the intended flow. It is rendered from Mermaid into SVG and PNG so the document contains the actual architecture figure instead of only source text.",
    )
    add_figure(doc, DIAGRAM_PNG, "Figure: Phase 2 tribunal bootstrap and Indian courtroom adaptation architecture")
    add_para(
        doc,
        "In simple terms, the novelty is that the pipeline does not jump directly from legal video to training data. It inserts validation, pairing, alignment, and quality control in between, and it keeps a separate path open for later Indian adaptation.",
    )

    doc.add_heading("6. Technical Explanation of the Proposed System", level=1)
    add_para(
        doc,
        "A student can think of the system as a sequence of filters. The first filter says, 'Is this even a real case number?' The second filter says, 'Does the UCR record actually contain case-specific evidence?' The third filter says, 'Can I build a real hearing from this evidence?' The fourth filter says, 'Can I align what was spoken with the audio and video timeline?' Only after all four filters pass do we get a training sample. After that, the same logic can be reused for Indian sources once a manifest, transcript, and media path are available.",
    )
    add_bullets(
        doc,
        [
            "The ledger stage handles candidate selection.",
            "The inventory stage handles record verification.",
            "The hearing stage handles grouping and pairing.",
            "The witness stage handles who-said-what metadata.",
            "The utterance stage handles segmentation and timing.",
            "The final stage handles export, validation, and splitting.",
            "The adaptation stage reuses the same logic for Indian sources.",
        ],
    )
    add_para(
        doc,
        "This is important because legal hearings are not naturally organized as machine-learning samples. Humans hear a discussion, but models need structured rows. Phase 2 creates that structure while preserving the legal provenance of every row.",
    )

    doc.add_heading("7. Algorithms and Techniques Used", level=1)
    add_bullets(
        doc,
        [
            "Placeholder and negative-control rejection. The enrichment layer rejects non-specific identifiers such as TO_BE_FILLED, TBD, and empty case labels before making a request.",
            "Record grouping and deduplication. Hearing rows are built by grouping case records that share the same date, title, session, and record identity.",
            "Witness identity extraction. Record titles and metadata are scanned to recover public witness names or protected codes without deanonymizing them.",
            "Transcript parsing. Speaker labels and Q/A patterns are turned into utterance objects before alignment starts.",
            "ASR-based alignment. The builder now supports a selectable alignment backend: `auto`, `whisperx`, or `heuristic`. In `auto` mode it prefers WhisperX-style alignment when that dependency is available; otherwise it falls back to the local `faster_whisper` plus fuzzy token matching path. This keeps the alignment pipeline reproducible while making it clear that the backend choice is an implementation detail, not the main methodological novelty.",
            "Sequence matching. `difflib.SequenceMatcher` is used for fuzzy alignment when exact matching fails.",
            "Fallback alignment. When the transcript cannot be matched confidently, proportional fallback timing is used and the row is marked low confidence or rejected.",
            "Media extraction. `ffmpeg` creates per-utterance MP4 and WAV clips from the resolved source media.",
            "Audio validation. RMS, silence ratio, clipping ratio, and sample rate checks reject corrupt or silent audio.",
            "Group-based splitting. Train/dev/test separation is done by hearing or related groups, not random row shuffling, to avoid leakage.",
        ],
    )
    add_para(
        doc,
        "These techniques are simple enough to explain at student level, but they are strong enough to support a reproducible legal dataset pipeline. The main idea is to reduce uncertainty step by step instead of trusting one heuristic too early. The alignment backend is intentionally configurable so that a local development run can stay on the heuristic path while a more alignment-focused run can use WhisperX-style timestamps when available.",
    )

    doc.add_heading("7.1 What To Read For Each Technique", level=2)
    add_table(
        doc,
        ["Pipeline part", "What to read first", "Why this is the right reading order"],
        [
            [
                "Utterance-level dataset design",
                "MELD: A Multimodal Multi-Party Dataset for Emotion Recognition in Conversations",
                "This is the closest benchmark for the MELD-style row format the project is trying to build.",
            ],
            [
                "ASR and word timestamps",
                "Whisper: Robust Speech Recognition via Large-Scale Weak Supervision, plus WhisperX-style alignment when available",
                "This gives the student a clear model of robust transcript-first speech recognition and noisy-audio handling, while WhisperX explains how the same audio can be turned into more precise word-level timing when that backend is installed.",
            ],
            [
                "Forced alignment",
                "NeMo Forced Aligner documentation and the 2023 NeMo Forced Aligner paper; then the Montreal Forced Aligner paper as a comparison point",
                "These explain how text and audio are turned into word- and segment-level timestamps for clip generation.",
            ],
            [
                "Text encoder for transcripts",
                "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding",
                "This is the standard reference for contextual text encoding in the transcript branch.",
            ],
            [
                "Visual encoder for facial cues",
                "An Image is Worth 16x16 Words: Transformers for Image Recognition at Scale",
                "This gives the transformer-based visual background for frame or face encoding.",
            ],
            [
                "Imbalanced labels and focal training",
                "Focal Loss for Dense Object Detection",
                "This is the classic reading for why imbalance-aware training helps when some classes dominate.",
            ],
            [
                "Media extraction and validation",
                "ffmpeg documentation and basic audio-signal checks",
                "These are implementation tools rather than research papers, so they are best read as system references.",
            ],
            [
                "Sequence matching and cleanup",
                "Python `difflib` documentation and transcript-alignment practice notes",
                "This explains the fuzzy-matching layer used when exact token matching is not enough.",
            ],
        ],
    )
    add_para(
        doc,
        "For the guidance call, the key point is not that every line of code comes from a paper. The key point is that the important modeling ideas are grounded in a small set of well-known references, while the file handling and validation logic are project-specific engineering choices.",
    )

    doc.add_heading("8. Expected Outcomes", level=1)
    add_para(
        doc,
        "If Phase 2 is completed as intended, the output will not be a long-video corpus. It will be a MELD-style dataset where each row is a single utterance and each sample has synchronized text, audio, and video. That is the right unit for fine-tuning a multimodal model. For the Indian scope, this means a tribunal bootstrap dataset plus a smaller but carefully curated Indian adaptation dataset.",
    )
    add_bullets(
        doc,
        [
            "A grounded legal-domain utterance dataset with traceable provenance.",
            "Better training signal than whole-hearing or whole-video samples.",
            "More useful case diversity than simply collecting extra hours from the same hearing family.",
            "A clean path for adding emotion, credibility, and speaker-role labels later.",
            "A corpus that can support courtroom analysis, not only generic multimodal classification.",
            "A separate Indian adaptation corpus for testing transfer beyond tribunal testimony.",
        ],
    )
    add_para(
        doc,
        "For the project, the biggest practical benefit is that the model gets many short aligned examples instead of a few very long ones. That usually helps learning because each training item is more focused and less noisy.",
    )

    doc.add_heading("9. What Is Implemented So Far", level=1)
    add_para(
        doc,
        "The current Phase 2 implementation is already real and reproducible. It includes corrected UCR inventory logic, hearing and witness manifest builders, an utterance-level LegalMELD exporter, and a validated rerun script that writes outputs under the validated dataset directory. What is implemented now is the tribunal bootstrap layer. The Indian acquisition and adaptation manifests are still part of the planned scope rather than the completed code path.",
    )
    add_table(
        doc,
        ["Implemented item", "Code-level description", "Current state"],
        [
            [
                "UCR case inventory correction and validation",
                "The enrichment and inventory path rejects placeholders before request time, validates actual record metadata, and keeps only cases that can be grounded in real record data. The verified inventory becomes the canonical case-level source.",
                f"{VERIFIED_INVENTORY.relative_to(ROOT)} exists and is the current verified inventory output.",
            ],
            [
                "Hearing manifest builder",
                "`phase2/build_hearing_manifest.py` groups verified inventory rows into hearing-level records, selects preferred transcript and video records, records pairing status, confidence, and witness identity, and emits `hearing_manifest.csv`.",
                f"{HEARING_MANIFEST.relative_to(ROOT)} is the hearing-level output path.",
            ],
            [
                "Witness manifest builder",
                "`phase2/build_witness_manifest.py` projects hearing rows into a resolved witness manifest, preserving protected codes and avoiding deanonymization.",
                f"{WITNESS_MANIFEST.relative_to(ROOT)} is the resolved witness manifest output.",
            ],
            [
                "LegalMELD utterance stage",
                "`phase2/build_legalmeld_dataset.py` parses transcript text into utterances, aligns them with ASR word timestamps, extracts clip segments with ffmpeg, validates audio/video quality, and writes train/dev/test splits. The current code path also allows the student to switch between the heuristic backend and a WhisperX-style backend without changing the export schema.",
                f"{LEGALMELD_METADATA.relative_to(ROOT)} plus clip folders and split CSVs are already generated.",
            ],
            [
                "Validated rerun path",
                "`phase2/run_build_legalmeld_dataset_validated.sh` runs the LegalMELD builder in pilot mode using the validated hearing manifest, media manifest, and selection manifest. It keeps the rerun path reproducible and preserves the backend choice through an environment variable.",
                "The rerun shell script is present and wired to the validated output root; in this workspace the reliable invocation is `PYTHON_BIN=/usr/bin/python3` because that interpreter can import `faster_whisper`.",
            ],
            [
                "Validated outputs",
                "The current outputs include the validated metadata CSV, an alignment review sample, a quality summary, and train/dev/test split files under the validated directory.",
                f"{VALIDATED_ROOT.relative_to(ROOT)} contains the current validated artifacts.",
            ],
            [
                "Indian adaptation pack",
                "The proposed Indian manifests are defined conceptually, but they are not yet implemented as a completed acquisition pipeline.",
                "This is the next scope item for the dissertation if the project proceeds beyond the bootstrap corpus.",
            ],
        ],
    )
    add_para(
        doc,
        "At code level, the LegalMELD builder uses `WhisperModel` or a WhisperX-style backend when available, `SequenceMatcher` for fuzzy matching, and `ffmpeg` for clip extraction. It also computes audio metrics such as RMS, silence ratio, clipping ratio, and sample rate, which means it does not blindly trust the media file just because the file exists. In this workspace the validated rerun path is executed with `PYTHON_BIN=/usr/bin/python3` so the alignment code can import the required speech module cleanly.",
    )
    add_para(doc, "Current validated snapshot from the latest rerun:", bold=True)
    add_bullets(
        doc,
        [
            f"{s['hearings_resolved']} hearings resolved.",
            f"{s['utterances']} utterances exported.",
            f"{s['verified_cases']} cases represented.",
            f"{s['distinct_witnesses']} distinct witnesses represented.",
            f"{s['train_rows']} train rows, {s['dev_rows']} dev rows, and {s['test_rows']} test rows.",
            f"{s['high_confidence']} high-confidence alignments, {s['medium_confidence']} medium-confidence alignments, and {s['low_confidence']} low-confidence alignments.",
            f"{s['quality_b']} Tier B rows and {s['quality_reject']} rejected rows.",
            f"{s['split_leakage']} split leakage violations.",
            f"{s['witness_visible']} witness-visible clips.",
            f"{s['total_clip_hours']} total clip hours.",
        ],
    )

    doc.add_heading("10. Clean Implementation Checklist for What Remains", level=1)
    add_numbered(
        doc,
        [
            "Expand the validated corpus beyond the current small pilot set of hearings and cases.",
            "Improve witness-visible clip recovery so the dataset contains real courtroom speaking-face evidence, not only audio-valid clips.",
            "Increase the number of Tier A rows by improving alignment confidence and visual quality checks.",
            "Grow case diversity so the dataset is not concentrated in only one or two case families.",
            "Keep group-based splitting strict so the same hearing or witness does not leak across splits.",
            "Add or refine manual review samples for the low-confidence and rejected alignment cases.",
            "Continue treating the placeholder witness harvest manifest as non-corpus data until it is fully replaced by resolved witness rows.",
            "Create the final ready-check script or checklist that defines the exact threshold for fine-tuning approval.",
            "Leave emotion and credibility annotation for the later stage, after the utterance dataset itself is stable.",
            "Build the Indian corpus acquisition pack and ingest a modest Indian evaluation/adaptation set.",
            "Validate the Indian sources on the same utterance-level pipeline so the method demonstrates transfer beyond tribunal testimony.",
        ],
    )
    add_para(
        doc,
        "The most important remaining work is not to collect more raw hours for their own sake. It is to improve utterance quality, witness diversity, and grounded pairing so that the samples are actually useful for learning.",
    )

    doc.add_heading("11. One-Week Task List", level=1)
    add_para(
        doc,
        "This is the concrete next-step plan for the next seven days. The focus is on improving the quality of the existing bootstrap corpus before expanding to the Indian acquisition layer.",
    )
    add_numbered(
        doc,
        [
            "Day 1: Inspect the current alignment outputs and identify the weak rows. Run no new corpus expansion yet. Open `data/processed/phase2/legalmeld_validated/alignment_review_sample.csv`, `data/processed/phase2/legalmeld_validated/dataset_quality_summary.json`, and `data/processed/phase2/legalmeld_validated/legalmeld_metadata_validated.csv`. Read the alignment logic in `phase2/build_legalmeld_dataset.py`, especially `segment_transcript()`, `_align_utterances()`, `_extract_segment()`, and `_audio_metrics()`.",
            "Day 2: Fix transcript parsing or alignment issues before changing the corpus. Inspect `phase2/build_legalmeld_dataset.py` and the helper modules that drive hearing resolution. If the issue is hearing grouping, inspect `phase2/build_hearing_manifest.py` and `phase2/hearing_witness_manifest_utils.py`. If the issue is witness identity resolution, inspect `phase2/build_witness_manifest.py` and `phase2/hearing_witness_manifest_utils.py`.",
            "Day 3: Rerun the validated pilot using the same wrapper. Run `bash phase2/run_build_legalmeld_dataset_validated.sh`. If you need to vary the pilot size, use the script’s supported options such as `--max-hearings`, `--max-utterances`, or `--include-hearing-ids` through the shell wrapper or direct Python call.",
            "Day 4: Inspect the rerun outputs and compare them to the previous batch. Open `data/processed/phase2/legalmeld_validated/legalmeld_metadata_validated.csv`, `data/processed/phase2/legalmeld_validated/train.csv`, `data/processed/phase2/legalmeld_validated/dev.csv`, `data/processed/phase2/legalmeld_validated/test.csv`, `data/processed/phase2/legalmeld_validated/alignment_review_sample.csv`, and `data/processed/phase2/legalmeld_validated/dataset_quality_summary.json`.",
            "Day 5: Check whether the hearing and witness layers still agree with the utterance layer. Inspect `data/processed/phase2/hearing_manifest.csv`, `data/phase2/source_manifests/witness_harvest_manifest_resolved.csv`, and the verified inventory `data/processed/phase2/verified_case_inventory.csv` to confirm the corpus is still grounded in real case metadata.",
            "Day 6: Decide whether the next iteration should widen within the tribunal bootstrap set or stay focused on fixing quality. If the quality has improved, choose one or two additional grounded hearing IDs and rerun `phase2/build_legalmeld_dataset.py` via the validated wrapper with `--include-hearing-ids` so you can test a slightly broader but still controlled batch.",
            "Day 7: Write a short status summary for the guidance call or supervisor check-in. Summarize the number of utterances exported, the high- and medium-confidence alignment counts, the number of rejected rows, the split leakage status, and whether any witness-visible clips improved. Use the validated JSON and CSV files as the evidence base.",
        ],
    )
    add_bullets(
        doc,
        [
            "Do not add the Indian acquisition pack yet unless the bootstrap corpus quality has improved.",
            "Do not expand raw hours before the utterance-level extraction is cleaner.",
            "Do not let the placeholder witness manifest drive any new estimates.",
        ],
    )
    add_para(
        doc,
        "The exact point of this week is to improve the current implementation where it already exists, not to jump ahead to a larger corpus. If the quality gets better, then the Indian adaptation layer becomes the next stage. If the quality does not get better, then the correct move is another focused iteration on parsing, alignment, witness/manifest traceability, and the choice of alignment backend.",
    )

    doc.add_heading("12. Indian Courtroom Adaptation Strategy", level=1)
    add_para(
        doc,
        "The Phase 2 tribunal corpus is technically valuable, but it is not an Indian courtroom testimony corpus. If the original dissertation objective is LegalMemoCMT for Indian courtroom testimony, then the final methodology should include an Indian adaptation and evaluation corpus. The strongest argument is that the tribunal data is the supervised bootstrap set, while the Indian data is the domain-adaptation target.",
    )
    add_para(
        doc,
        "This is a realistic design because there is currently no large public Indian courtroom dataset that provides synchronized video, audio, transcript, and witness metadata at MELD scale. That means the project cannot rely on a single perfect Indian source. Instead, it should use a hybrid corpus strategy.",
    )
    add_bullets(
        doc,
        [
            "Tribunal hearings provide the large supervised multimodal testimony base.",
            "Indian sources provide domain adaptation for accent, procedure, and legal discourse.",
            "The same utterance-level pipeline can be reused once media and transcripts are available.",
            "The final dissertation becomes stronger because it demonstrates both supervised learning and domain transfer.",
        ],
    )

    doc.add_heading("12.1 Realistic Indian Sources", level=2)
    add_table(
        doc,
        ["Source family", "Strength", "Limitation", "Best use"],
        [
            [
                "Supreme Court of India livestreams",
                "Official, authentic Indian courtroom speech with good audio/video quality.",
                "Mostly arguments by advocates and judges; witness examination is rare.",
                "Legal-language adaptation, judge speech, advocate speech, procedure-aware modeling.",
            ],
            [
                "High Court livestreams",
                "Indian legal proceedings with real courtroom style and official context.",
                "Usually argument-heavy, with limited witness testimony.",
                "Additional adaptation data and legal discourse variety.",
            ],
            [
                "District Court material",
                "Closest source to real witness testimony.",
                "Public video and transcript access is usually limited.",
                "Only use when public recordings and transcript evidence are available.",
            ],
            [
                "Mock trials from law schools and judicial academies",
                "Often contain direct and cross-examination with visible speakers.",
                "Simulated rather than real proceedings.",
                "Very useful for witness-role structure, examination phases, and utterance alignment practice.",
            ],
            [
                "National Judicial Academy and State Judicial Academy material",
                "Can include witness-examination demonstrations and advocacy training.",
                "Usually small-scale and non-uniform.",
                "Supplementary evidence for speaker-role and examination-type learning.",
            ],
            [
                "Parliament TV, Law Commission, and legal education videos",
                "Useful legal speech material.",
                "Not witness testimony.",
                "Text and legal-language adaptation only.",
            ],
        ],
    )

    doc.add_heading("12.2 Why a Hybrid Corpus Is the Correct Design", level=2)
    add_para(
        doc,
        "A hybrid corpus is the most defensible design because no single Indian source currently solves every requirement. Supreme Court and High Court livestreams are excellent for authentic Indian legal speech, but they are not rich witness-testimony corpora. District courts are closer to witness testimony, but public multimodal access is limited. Mock trials and judicial-academy recordings are not real cases, but they often give the exact utterance-level structure needed for alignment, role classification, and cross-examination modeling.",
    )
    add_bullets(
        doc,
        [
            "Tribunal corpora solve scale and witness-testimony supervision.",
            "Indian corpora solve domain shift and local procedural adaptation.",
            "Mock trials help where public real testimony is not available.",
            "The combined result is much stronger than using only one source family.",
        ],
    )

    doc.add_heading("12.3 Proposed Indian Corpus Acquisition Pack", level=2)
    add_para(
        doc,
        "If the project continues into the Indian adaptation phase, the next practical step is to create a separate acquisition pack. That pack should mirror the tribunal pipeline but keep the Indian sources separate from the bootstrap corpus so that the provenance remains clear.",
    )
    add_code_block(
        doc,
        """
indian_case_candidate_ledger.csv
indian_video_sources.csv
indian_mock_trial_manifest.csv
indian_supreme_court_manifest.csv
indian_high_court_manifest.csv
indian_alignment_manifest.csv
""",
    )
    add_para(
        doc,
        "These files would not replace the tribunal pipeline. They would sit beside it and feed the same manifest, validation, alignment, and utterance-export logic that Phase 2 already uses. That is why the current tribunal work is still valuable even though the final research target is Indian courtroom testimony.",
    )

    render_india_mermaid(
        """
flowchart LR
    A[Indian acquisition pack] --> B[Supreme Court livestreams]
    A --> C[High Court livestreams]
    A --> D[District court material]
    A --> E[Mock trials and judicial academies]
    B --> F[Indian candidate ledger]
    C --> F
    D --> F
    E --> F
    F --> G[Indian verification layer]
    G --> H[Indian verified inventory]
    H --> I[Indian hearing manifest]
    H --> J[Indian witness manifest]
    I --> K[Transcript-first segmentation]
    J --> K
    K --> L[Forced alignment and clip generation]
    L --> M[Indian LegalMELD-style utterance corpus]
"""
    )

    doc.add_heading("12.4 Project Judgment for the Indian Objective", level=2)
    add_bullets(
        doc,
        [
            "Yes, the project can be demonstrated with ICTY / ICTR / IRMCT data.",
            "No, that alone does not fully satisfy the Indian courtroom testimony objective.",
            "A modest Indian adaptation corpus of even 5 to 10 hours would materially strengthen the dissertation.",
            "The tribunal corpus should therefore be presented as bootstrap data, not as the final domain target.",
        ],
    )
    add_para(
        doc,
        "This framing is academically stronger because it avoids overstating the scope. It says clearly that the pipeline has been proven on a real multimodal legal corpus, and it also says clearly that the final research contribution is the adaptation of that pipeline to Indian legal speech.",
    )

    doc.add_heading("12.5 Indian Acquisition Pack Flow", level=2)
    add_figure(doc, INDIA_DIAGRAM_PNG, "Figure: Indian courtroom acquisition pack and adaptation flow")

    doc.add_heading("13. Ready vs Not Ready Judgment for Fine-Tuning", level=1)
    add_table(
        doc,
        ["Status", "Judgment"],
        [
            [
                "Ready",
                "The implementation is ready as a reproducible pilot pipeline, a documentation target, and a method-validation artifact. The inventory, manifests, LegalMELD exporter, and validated rerun path all exist.",
            ],
            [
                "Not ready",
                "The corpus is not yet ready for final fine-tuning at the target Indian scope because it is still small, covers only two cases and four hearings in the latest validated run, has no witness-visible clips, and is dominated by Tier B rows. It is also still a tribunal bootstrap corpus rather than the full Indian adaptation corpus.",
            ],
            [
                "Final judgment",
                "Ready for controlled experimentation and iterative dataset growth, but not yet ready for the final Indian courtroom fine-tuning milestone.",
            ],
        ],
    )
    add_para(
        doc,
        "So the precise answer is: Phase 2 is implementation-ready as a tribunal bootstrap pipeline, but not yet fine-tuning-ready as the full Indian courtroom corpus. That is the correct technical judgment based on the current validated outputs and the current scope of the dissertation.",
    )

    doc.add_heading("Appendix: Validation Numbers Used in This Report", level=1)
    add_bullets(
        doc,
        [
            f"Validated metadata rows: {s['metadata_rows']}.",
            f"Validated inventory rows: {s['inventory_rows']}.",
            f"Hearing manifest rows: {s['hearing_manifest_rows']}.",
            f"Witness manifest rows: {s['witness_manifest_rows']}.",
            f"Audio-valid clips: {s['audio_valid']}.",
            f"Video-valid clips: {s['video_valid']}.",
            f"Manual review rows: {s['manual_review']}.",
        ],
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
