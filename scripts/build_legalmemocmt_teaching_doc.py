from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


OUTPUT = Path("implementation_docments/LegalMemoCMT_Multimodal_Learning_Teaching_Draft.docx")


def set_document_defaults(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            styles[style_name].font.name = "Times New Roman"

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LegalMemoCMT Multimodal Learning Teaching Draft")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Speech, Video, Text, Fusion, and Literature Survey for Student-Level Instruction"
    )
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)

    doc.add_paragraph()
    meta = [
        "Purpose: to explain how multimodal emotion recognition works from first principles and how the base MemoCMT line of work maps into the LegalMemoCMT project.",
        "Scope: speech emotion recognition, video-based emotion recognition, text emotion recognition, fusion learning, article-by-article literature review, and legal-AI limitations.",
        "Audience: a student who needs a teaching document that can bridge theory, implementation, and the project-specific architecture.",
        "Framing: the system should support interpretation of emotional patterns, not legal decision-making.",
    ]
    for line in meta:
        doc.add_paragraph(line)

    doc.add_paragraph()


def add_objectives(doc: Document) -> None:
    doc.add_heading("Clear Objectives", level=1)
    objectives = [
        "Explain how speech, video, and text are processed by machine learning systems for emotion recognition, starting from raw signals and ending at classification.",
        "Teach unimodal learning before introducing fusion, so the student understands why multimodal systems outperform isolated channels in many settings.",
        "Use the selected literature to show how the field evolved from canonical correlation and Bayesian fusion toward transformers, graph reasoning, and cross-modal attention.",
        "Connect the theory to LegalMemoCMT, with special attention to how the project can support explanatory analysis of emotional patterns without turning those patterns into legal conclusions.",
        "Discuss fairness, bias, explainability, dataset limitations, and the risks of emotional inference in judicial settings and adjacent high-stakes domains.",
    ]
    for item in objectives:
        doc.add_paragraph(item, style="List Bullet")


def add_paragraphs(doc: Document, paragraphs):
    for text in paragraphs:
        doc.add_paragraph(text)


def add_subsection(doc: Document, title: str, paragraphs):
    doc.add_heading(title, level=2)
    add_paragraphs(doc, paragraphs)


def add_table(doc: Document, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def add_section_one(doc: Document) -> None:
    doc.add_heading("1. Core Teaching Path", level=1)
    add_paragraphs(
        doc,
        [
            "The right way to teach multimodal emotion recognition is to move from one modality at a time to the combined system. If a student jumps straight to fusion, the model looks like magic. If the student first understands speech emotion recognition, then visual emotion recognition, then text emotion recognition, the design logic becomes visible.",
            "The central idea is that emotion is not a single observable quantity. It is inferred from cues distributed across speech prosody, facial expression, body posture, lexical content, temporal context, speaker identity, and the interaction between them. A robust system learns representations for each cue, aligns them, and then decides how much weight each cue should receive in the final prediction.",
            "For LegalMemoCMT, the pedagogical aim is not only to classify emotion labels. It is to show how each modality contributes an interpretable pattern, what each branch learns, and where the model can fail. That is the bridge between a research prototype and a responsible high-level project in a sensitive domain.",
        ],
    )

    add_subsection(
        doc,
        "1.1 How Speech Emotion Recognition Works",
        [
            "Speech emotion recognition starts from the observation that the same words can carry different emotional meanings depending on tone, pace, pauses, pitch movement, energy, spectral shape, and voice quality. A person can say a neutral sentence in a frustrated tone, and the acoustic signal will contain cues that the transcript alone will not reveal. That is why speech is often treated as a distinct source of affective information rather than as a side note to language.",
            "The speech pipeline usually begins with audio collection and pre-processing. The waveform is normalized, resampled, clipped or padded to a fixed duration, and optionally denoised. Traditional systems extract handcrafted features such as MFCCs, pitch, energy, jitter, shimmer, formants, speaking rate, and delta features. Deep learning systems often learn a representation directly from the waveform or from a time-frequency representation such as a mel spectrogram.",
            "A modern speech model usually contains a front-end encoder and a classifier. The encoder can be a CNN, a recurrent model, an attention-based model, or a pretrained transformer such as HuBERT, wav2vec 2.0, or XLS-R. The encoder compresses the raw sequence into a compact latent vector while keeping useful temporal structure. The classifier then maps that latent vector to emotion categories or to continuous affective dimensions such as valence and arousal.",
            "A student should understand the difference between local and global cues. Local cues include short bursts of emphasis, pitch spikes, and pauses. Global cues include the overall speaking style across a whole utterance or dialogue turn. Recurrent and convolutional models are good at local and medium-range structure, while transformers are stronger when the model must reason over longer dependencies. That is why transformer-based speech systems have become common in recent literature.",
            "Evaluation is not just about accuracy. Speech emotion recognition in conversation is usually reported with macro F1, weighted F1, unweighted accuracy, per-class recall, and confusion analysis because class imbalance is common. This matters in teaching: if the student only sees accuracy, they may miss the fact that a system can perform well on frequent emotions and poorly on minority classes such as fear or disgust.",
            "The main teaching message for speech is that the model must convert a continuous acoustic stream into a representation that preserves emotional timing. Emotion is rarely located in a single frame. It emerges from spans of sound and from how the speaker moves across those spans. That is why pooling strategy, context window length, and attention design matter as much as the backbone itself.",
        ],
    )

    add_subsection(
        doc,
        "1.2 How Video Emotion Recognition Works",
        [
            "Video emotion recognition covers the visible part of affective expression: facial muscles, head motion, gaze direction, hand motion, body posture, and co-speech gestures. In practice, video is one of the most informative modalities because it captures nonverbal cues that are often withheld or flattened in text and speech. A video branch is usually asked to detect how the face and body evolve over time, not just what one frame looks like.",
            "The simplest video pipeline begins with frame extraction, face detection, alignment, and temporal aggregation. In a stronger pipeline, the system can also track body keypoints, facial landmarks, action units, and gesture trajectories. Deep models then process either raw frames or precomputed visual features. CNNs learn appearance cues, while 3D CNNs, temporal convolutional networks, and vision transformers learn motion and sequence structure.",
            "A key concept for students is that emotion in video is not just classification from a still image. Movement matters. A fleeting smile, a restrained jaw, an averted gaze, or a tense shoulder can change the inferred label. That means the model needs both spatial and temporal reasoning. If it only sees a single frame, it can confuse a posed expression with a genuine emotional state or miss the transition from neutral to angry.",
            "Video emotion recognition also has a serious data challenge. Labels are often noisy because human annotators disagree, and many datasets are acted rather than spontaneous. Acted datasets are useful for learning a clean signal, but spontaneous datasets are better for real-world generalization. Students should be taught that better benchmark accuracy does not automatically mean better real-world reliability.",
            "For LegalMemoCMT, the video branch should be interpreted as a complementary evidence stream, not as a source of truth. Visual emotion cues can support a broader pattern analysis, but they are vulnerable to occlusion, camera angle, lighting, cultural style, and individual differences. This is one reason the final system should present emotional patterns with uncertainty and explanatory context rather than direct judgments.",
        ],
    )

    add_subsection(
        doc,
        "1.3 How Text Emotion Recognition Works",
        [
            "Text emotion recognition extracts affect from words, syntax, discourse context, and pragmatic cues. The text branch is not only about sentiment polarity. It can learn that a phrase is sarcastic, dismissive, anxious, apologetic, or supportive depending on the surrounding context. In conversation, the same utterance may change meaning when preceded by a question, a correction, or a contradiction.",
            "A standard text pipeline includes normalization, tokenization, embedding, and contextual encoding. Early systems used bag-of-words or static embeddings, but modern systems rely on pretrained language models such as BERT, RoBERTa, DeBERTa, and task-specific dialogue encoders. These models learn that word meaning changes across context and that emotion often depends on discourse structure rather than isolated tokens.",
            "Students should learn that text emotion models work at several layers. The lexical layer sees emotionally loaded words. The syntactic layer sees negation, intensifiers, and dependency structure. The discourse layer sees whether the utterance is a response, a correction, a continuation, or a contradiction. The conversational layer sees speaker history and prior turns. Dialogue-based emotion recognition is difficult precisely because the emotion label for a turn may be inherited from or modified by the previous turns.",
            "Transformer encoders dominate current text emotion recognition because self-attention can relate any token to any other token in the sequence. That is useful for long-distance dependencies such as negation scope, sarcasm, contrastive clauses, or emotion shifting across a dialogue. However, transformers also bring costs: they may overfit on dataset-specific style and can be less transparent than simpler models.",
            "For LegalMemoCMT, the text branch can be taught as the semantic anchor. It tells us what was said, which is essential for interpreting the other modalities. But the textual signal should not be treated as emotionally complete. In legal or quasi-legal settings, a transcript can be misleading if it is detached from tone, pacing, and speaker demeanor. Therefore the text branch should be explained as one lens among several, not as a substitute for context.",
        ],
    )

    add_subsection(
        doc,
        "1.4 How Fusion Recognition Works",
        [
            "Fusion is the step that turns isolated modality models into a multimodal system. The core question is not whether to combine modalities, but when and how. Early fusion concatenates inputs or features before classification. Late fusion combines individual predictions. Intermediate fusion uses a joint hidden space. Hybrid and cross-modal fusion models do both, often with attention or gating.",
            "A student should be taught that fusion is a representation learning problem. The model must align information across modalities that are sampled differently, shaped differently, and sometimes contradictory. Speech may express anger, the text may remain polite, and the video may show discomfort. A good fusion module must learn whether the disagreement means noise, irony, masking, or a true mixed state.",
            "Classic fusion methods include canonical correlation analysis, multiset correlation methods, and Bayesian decision fusion. Their value is conceptual clarity: they show the student that multimodal learning is partly about finding shared structure across modalities. Modern systems push this idea further with cross-attention, modality-specific gating, graph reasoning, and memory-based fusion. These models are better at sequential conversation, but they are also more complex to train and interpret.",
            "In practice, fusion design controls both performance and explainability. If the fusion layer is too simple, the model behaves like a concatenation baseline and misses cross-modal interactions. If the fusion layer is too complex, the system may become difficult to analyze and fragile under missing or noisy modalities. The best teaching approach is to compare simple concatenation, attention-based fusion, graph-based fusion, and cross-modal transformer fusion side by side.",
            "For LegalMemoCMT, the strongest lesson is that fusion should be interpreted as evidence integration. Each modality contributes partial evidence, and the fusion model estimates how strongly to trust each source in a given context. That framing supports a responsible system design: the output can be presented as an interpretable emotional pattern, with confidence and modality contribution information, rather than as a deterministic claim about intent or truth.",
        ],
    )


def add_section_two(doc: Document) -> None:
    doc.add_heading("2. Literature Survey Matrix", level=1)
    add_paragraphs(
        doc,
        [
            "The papers below are best treated as a staged reading list. The first paper gives a clear trimodal baseline with an explicit fusion method. The second paper provides a methodical review of speech emotion recognition in conversational settings. The third paper shows how transformer and graph fusion can be combined in dialogue. The fourth paper is the conceptual anchor for the project because it uses cross-modal transformers. The fifth paper serves as a survey backbone for tasks, datasets, and evaluation. The sixth paper is a simpler multimodal baseline that is useful for teaching and comparison.",
            "For each article, the teaching objective is the same: identify the modalities, identify the representation method, identify the fusion mechanism, identify the dataset, identify the evaluation logic, and identify the limit of the method. That pattern makes it much easier for a student to compare papers without getting lost in implementation detail.",
        ],
    )

    headers = ["Paper", "Year", "Main modalities", "Main technique", "Why it matters for LegalMemoCMT"]
    rows = [
        [
            "Yan et al., Multimodal Emotion Recognition Based on Facial Expressions, Speech, and Body Gestures",
            "2024",
            "Face, speech, body gesture",
            "SLSMKCCA and SSLSMKCCA fusion",
            "Best trimodal baseline for feature fusion and classic multimodal alignment.",
        ],
        [
            "Alhussein et al., Speech Emotion Recognition in Conversations Using Artificial Intelligence",
            "2025",
            "Speech-centered conversation, with broader review scope",
            "Systematic review and meta-analysis",
            "Provides the review backbone for datasets, metrics, bias, and method trends.",
        ],
        [
            "Jin et al., Multimodal Emotion Recognition in Conversations Using Transformer and Graph Neural Networks",
            "2025",
            "Text, audio, visual",
            "Transformer fusion plus directed graph reasoning",
            "Shows how to combine global attention and local conversational structure.",
        ],
        [
            "Khan et al., MemoCMT: Multimodal Emotion Recognition Using Cross-Modal Transformer-Based Feature Fusion",
            "2025",
            "Audio and text",
            "Cross-modal transformer fusion",
            "Main conceptual reference for the project architecture and replication logic.",
        ],
        [
            "Hu et al., Recent Trends of Multimodal Affective Computing: A Survey from NLP Perspective",
            "2024",
            "Text-dominant multimodal tasks",
            "Survey of task families, datasets, metrics, and approaches",
            "Useful for framing the NLP side of affective computing and comparative evaluation.",
        ],
        [
            "Kraack, A Multimodal Emotion Recognition System: Integrating Facial Expressions, Body Movement, Speech, and Spoken Language",
            "2024 preprint",
            "Face, body, speech, text",
            "Integrated practical multimodal system",
            "Useful as an implementation-oriented system-level example when a fuller multimodal stack is needed.",
        ],
    ]
    add_table(doc, headers, rows)


def add_paper_section(doc: Document, title: str, overview, method, strengths, limitations, relation):
    doc.add_heading(title, level=2)
    doc.add_heading("Problem and Scope", level=3)
    add_paragraphs(doc, overview)
    doc.add_heading("Method and Teaching Points", level=3)
    add_paragraphs(doc, method)
    doc.add_heading("Strengths", level=3)
    for x in strengths:
        doc.add_paragraph(x, style="List Bullet")
    doc.add_heading("Limitations and What to Teach", level=3)
    for x in limitations:
        doc.add_paragraph(x, style="List Bullet")
    doc.add_heading("Relation to LegalMemoCMT", level=3)
    add_paragraphs(doc, relation)


def add_section_three(doc: Document) -> None:
    doc.add_heading("3. Paper-by-Paper Teaching Notes", level=1)

    add_paper_section(
        doc,
        "3.1 Yan et al. (2024): Multimodal Emotion Recognition Based on Facial Expressions, Speech, and Body Gestures",
        [
            "This paper is an implementation-oriented trimodal baseline. It is valuable because it deals with the exact kind of cross-modal complementarity that students need to understand before they encounter more abstract transformer systems. The paper studies facial expressions, speech, and body gestures together, which makes it a natural classroom example of how multimodal affective information is distributed across visual and acoustic channels.",
            "The study uses the GEMEP and Polish databases, which matters because it demonstrates that benchmark choice shapes the type of emotion signal available to the model. The work is not framed as a generic deep learning benchmark. Instead, it carefully organizes the pipeline around feature extraction, feature fusion, and multimodal classification. That makes it easy to teach as a sequence rather than as a black box.",
        ],
        [
            "The main technical contribution is the use of supervised least squares multiset kernel canonical correlation analysis, or SLSMKCCA, and its sparse variant, SSLSMKCCA. The student should understand this as a disciplined way of forcing different modality feature spaces to agree in a shared latent structure. In simple terms, the method tries to find correlated projections of face, speech, and body-gesture features while respecting supervision from emotion labels.",
            "The sparse variant adds L1 regularization so the model can drop irrelevant dimensions. That is pedagogically useful because it shows the student that multimodal learning is not only about adding more information. It is also about removing noise and reducing redundancy. The paper therefore teaches an important lesson: more modalities do not automatically mean better performance unless the fusion step can filter them intelligently.",
            "The solving strategies, alternating least squares and augmented Lagrangian multiplier methods, are also worth teaching. They show that multimodal fusion is often an optimization problem before it is a neural-network problem. This helps students understand the mathematical roots of the field and prevents them from assuming that all multimodal systems are transformer-based.",
            "The paper compares monomodal, bimodal, and trimodal recognition and reports that the trimodal setup performs best on average. That comparison is important because it offers a clear teaching point: gains come not just from adding modalities, but from learning a better relationship between modalities. The paper also shows why body gesture is worth including even though it is harder to annotate and extract reliably than face or speech.",
        ],
        [
            "The paper provides a clean example of explicit multimodal feature fusion.",
            "It connects classic correlation-based learning to emotion recognition in a way students can follow.",
            "It demonstrates how dataset choice and modality choice interact.",
            "It is a strong baseline for comparing against attention-based or graph-based fusion.",
        ],
        [
            "The article is not primarily a conversation model, so it does not capture dialogue history in the way LegalMemoCMT must. It is also more classical than the newer transformer literature, which means students should not treat it as the final state of the art. The value lies in its clarity, not in its recency.",
            "The paper is strongest as a conceptual bridge between handcrafted multimodal fusion and modern neural fusion. Teach the student to ask what information is preserved, what information is discarded, and whether the fusion layer is supervised enough to align the modalities in a human-interpretable way.",
        ],
        [
            "For LegalMemoCMT, this paper is useful because it shows how to think about visual evidence when adding a video branch. If the project extends beyond speech and transcript into face and body cues, the student can use this paper as a control reference for classical trimodal fusion and then compare it with the transformer-based project design.",
        ],
    )

    add_paper_section(
        doc,
        "3.2 Alhussein et al. (2025): Speech Emotion Recognition in Conversations Using Artificial Intelligence",
        [
            "This paper is a systematic review and meta-analysis of AI-based speech emotion recognition in conversation. It is not a single model paper, and that is precisely why it is valuable for teaching. A student often learns more from a structured review of the field than from one more benchmark result because the review tells them which model families are stable, which datasets are dominant, and where the literature is biased.",
            "The review adopts a speech-centered perspective while still acknowledging the importance of multimodality in the broader ERC field. It examines studies from January 2010 through March 2023 and organizes the literature around methods, datasets, evaluation outcomes, and methodological quality. The core lesson is that the field has matured from handcrafted features and classical ML into deep learning and transformer-informed modeling.",
        ],
        [
            "The study follows a PROSPERO-registered protocol and the PRISMA-DTA reporting structure. Students should be taught to value that rigor because it prevents a review from becoming a subjective list of papers. The work identifies studies that report quantitative performance, extracts study attributes such as feature extraction method and database, and then synthesizes the landscape of speech emotion recognition in conversations.",
            "The review distinguishes between handcrafted feature pipelines and deep learning pipelines. That division is useful in the classroom because it explains a key historical transition in affective computing. Handcrafted systems typically depend on engineered acoustic descriptors, while deep models learn representations directly from data. The review also highlights transformer-based and attention-based trends, which helps students understand why context modeling has become central.",
            "Another important teaching point is methodological quality. A review is not only about ranking models by score. It is also about examining class imbalance, dataset size, annotation noise, task definition, and reproducibility. For a student working on LegalMemoCMT, this is essential because high apparent performance can hide fragile assumptions about the data or the evaluation setup.",
            "The review is especially helpful for explaining why dialogue emotion recognition is hard. Emotion changes across turns, speaker roles matter, and the local meaning of a speech segment depends on prior context. The review therefore sets up the need for memory, context windows, and speaker-aware architectures.",
        ],
        [
            "It gives a field-level map rather than a single model snapshot.",
            "It uses systematic review logic, which is useful for teaching research methods.",
            "It emphasizes datasets, evaluation metrics, and bias in conversational emotion recognition.",
            "It helps explain the shift from handcrafted speech features to transformer-era methods.",
        ],
        [
            "The review is speech-centered, so it does not cover visual emotion signals in depth. That is not a weakness for its stated aim, but students should not confuse it with a full multimodal survey. It also does not replace a model paper when the goal is implementation; it provides the landscape in which implementation choices should be interpreted.",
            "The student should be taught that reviews are useful for justification, but not enough for architecture design. The review can tell us what the field values, but not exactly how to build the next system. That is why it should be read alongside MemoCMT and the graph-fusion paper.",
        ],
        [
            "For LegalMemoCMT, this review is the evidence base for the speech branch and for the broader methodological discussion. It helps justify why speech emotion recognition cannot be treated as a solved problem and why conversational context, bias, and reproducibility must be discussed in the thesis or project report.",
        ],
    )

    add_paper_section(
        doc,
        "3.3 Jin et al. (2025): Multimodal Emotion Recognition in Conversations Using Transformer and Graph Neural Networks",
        [
            "This paper is a modern conversational multimodal system that combines transformer fusion and graph reasoning. It is directly relevant because it addresses the exact challenge that many students face when they move from single-utterance emotion recognition to dialogue-level emotion recognition: the model must reason over both long-range conversational dependencies and local speaker-conditioned shifts.",
            "The paper introduces a model called MTG-ERC. The architecture is designed to capture global and local emotional information simultaneously. That matters because conversational emotion is not only about the current utterance. It is also about how the utterance sits in the dialogue history, how one speaker responds to another, and how emotional state changes over time.",
        ],
        [
            "The first branch is a multi-level Transformer fusion module. It uses multi-head self-attention and cross-modal attention to model intra-modal and inter-modal interactions. For teaching, this branch should be explained as the global reasoning component. It can discover dependencies across distant utterances and across modalities, which is difficult for a plain recurrent model.",
            "The second branch is a directed multi-relational graph fusion module. This graph encodes local structure with multiple relation types, including directed conversational edges and speaker-aware links. Students should understand that the graph is used to preserve local conversational structure that a pure attention model may smooth away. In other words, the graph branch protects short-range context and speaker-dependent transitions.",
            "The model then combines the outputs of the two branches through gated fusion. This is pedagogically important because it shows a hybrid design pattern: instead of forcing one mechanism to do everything, the system divides labor between global attention and local graph reasoning. That is a pattern students can reuse when they design their own systems.",
            "The paper reports improvements of about one absolute point on IEMOCAP and MELD relative to strong baselines. That size of improvement is realistic in modern multimodal research, and it is a good lesson for students: once the field is mature, performance gains often come from architectural refinements, better alignment, and better use of context rather than from dramatic new leaps.",
        ],
        [
            "It teaches a hybrid design that combines transformer fusion and graph reasoning.",
            "It is explicitly conversation-aware and speaker-aware.",
            "It shows how global and local reasoning can be separated and then fused.",
            "It uses public benchmarks that are familiar in multimodal ERC research.",
        ],
        [
            "The model is more complex than the earlier canonical-correlation baseline, so students may need help mapping the graph relations to actual dialogue events. It is also more abstract than a pure speech or text system, which means it can be harder to explain line by line without visual diagrams.",
            "A student should be warned not to think of the graph as just another layer. The graph encodes assumptions about conversation structure, so the choice of edge types and relation types is part of the scientific claim. That point is easy to miss when reading advanced papers.",
        ],
        [
            "For LegalMemoCMT, this paper is the best bridge to the project's multimodal dialogue reasoning story. If the project includes conversational history and multiple channels, MTG-ERC gives a strong reference for how to combine attention-based global fusion with graph-based local structure. It is especially useful for teaching how a model can stay interpretable by separating the two reasoning pathways.",
        ],
    )

    add_paper_section(
        doc,
        "3.4 Khan et al. (2025): MemoCMT: Multimodal Emotion Recognition Using Cross-Modal Transformers",
        [
            "MemoCMT is the conceptual center of the project. The paper uses audio and text, which makes it a speech-plus-language model rather than a full visual multimodal system. That distinction matters. The paper does not merely concatenate features; it uses a cross-modal transformer mechanism to integrate audio and text representations in a way that captures local and global dependencies.",
            "The paper is important for teaching because it is both practical and methodologically clean. It uses pretrained HuBERT for audio and BERT for text, then passes those representations through a cross-modal transformer. This is a strong example of how modern systems separate feature extraction from fusion. The student can therefore learn the architecture as a pipeline rather than as a monolith.",
        ],
        [
            "MemoCMT explains why cross-modal transformers are attractive for emotional understanding. The audio branch captures prosody, timbre, and temporal affective patterns. The text branch captures semantic content and discourse meaning. The cross-modal transformer lets these streams interact so that the model can look for agreement, contradiction, and complementarity between what is said and how it is said.",
            "The paper also compares aggregation choices such as CLS, mean, max, and min. This is a useful teaching hook because it shows that pooling is not a trivial detail. Different aggregation methods can change how the model summarizes a sequence of fused tokens. In the paper, min aggregation performs best in the reported setup. Students should learn to treat such details as design choices, not afterthoughts.",
            "Another teaching point is interpretability. The paper acknowledges that deep emotion systems can be black boxes and explicitly discusses the value of layer-by-layer training and representation analysis. Even if the system remains a deep network, the student can still learn to inspect t-SNE patterns, ablation studies, and performance under different fusion settings.",
            "The model reports strong performance on IEMOCAP and ESD and includes a MELD case study. That combination is ideal for teaching because it shows both benchmark evaluation and a conversational use case. It also illustrates the project logic that the same architecture can be stressed under multiple corpora.",
        ],
        [
            "It is the main architecture reference for LegalMemoCMT.",
            "It cleanly separates speech encoding, text encoding, and cross-modal fusion.",
            "It uses pretrained encoders that are easy to explain to students.",
            "It compares multiple pooling strategies, which supports a controlled ablation discussion.",
        ],
        [
            "MemoCMT is still not a full face-plus-body system, so the project should not present it as such. Students need to know that adding video is an extension, not a direct feature of the original paper. The paper is also focused on speech emotion recognition and multimodal fusion of audio and text, so the visual branch must be introduced as a project-level enhancement.",
            "The best way to teach MemoCMT is to emphasize the role of cross-modal alignment. If the student understands why HuBERT and BERT are used, and why a transformer-based fusion block sits between them and the classifier, the rest of the architecture becomes much easier to reason about.",
        ],
        [
            "For LegalMemoCMT, this is the principal reference. The project should be framed as a broader multimodal teaching and exploration layer that uses the MemoCMT design philosophy, while extending the explanation to video and to legal-AI concerns. In a thesis or project report, MemoCMT is the anchor point from which the architecture discussion should start.",
        ],
    )

    add_paper_section(
        doc,
        "3.5 Hu et al. (2024): Recent Trends of Multimodal Affective Computing: A Survey from NLP Perspective",
        [
            "This survey is valuable because it frames multimodal affective computing from an NLP perspective and organizes the space around task families rather than around a single dataset or model. The paper covers multimodal sentiment analysis, multimodal emotion recognition in conversation, multimodal aspect-based sentiment analysis, and multimodal multi-label emotion recognition.",
            "That task-based organization is especially useful for teaching students because it helps them see that multimodal affective computing is not one problem. It is a family of related tasks with different labels, different evaluation standards, and different modeling assumptions. A student who understands that distinction will have a much stronger foundation for reading the rest of the literature.",
        ],
        [
            "The survey provides a unified view of task formulations, benchmark datasets, evaluation protocols, representative methods, and future challenges. It also briefly connects the NLP-centered view to facial, acoustic, and physiological modalities. This wide framing is useful because it stops the student from assuming that multimodal emotion recognition begins and ends with speech and video only.",
            "For LegalMemoCMT, the survey helps define the scope of the thesis or report. It explains how the project sits at the intersection of multimodal sentiment analysis, emotion recognition in conversation, and affective computing more broadly. It also helps justify why datasets, metrics, and context modeling should be discussed before model architecture.",
            "The survey's value is partly organizational. It gives the student a taxonomy of the field and a place to put each method they encounter. That reduces confusion when reading papers that use different terms for similar ideas, such as multimodal affective computing, conversational emotion recognition, and multimodal sentiment analysis.",
        ],
        [
            "It organizes the field around tasks, datasets, and evaluation protocols.",
            "It gives a broad view that is especially useful for literature review writing.",
            "It helps students distinguish between related but different affective computing problems.",
            "It broadens the discussion beyond audio-text fusion into wider multimodal affective research.",
        ],
        [
            "The paper is a survey, so it does not provide implementation detail or ablation logic. Students should not expect a build guide from it. Its strength is synthesis, not model design. It also has a stronger NLP orientation than some vision-first affective computing surveys, so it should be used alongside a more system-level multimodal source.",
        ],
        [
            "For LegalMemoCMT, the survey provides the conceptual umbrella. It is useful for the introductory chapter, the taxonomy chapter, and the dataset/metric chapter. It helps show why the project is not just an isolated model implementation, but part of a larger research field that has its own tasks, constraints, and evaluation habits.",
        ],
    )

    add_paper_section(
        doc,
        "3.6 Kraack (2024 preprint): A Multimodal Emotion Recognition System Integrating Facial Expressions, Body Movement, Speech, and Spoken Language",
        [
            "This preprint is useful as an implementation-oriented system example because it explicitly combines facial expressions, body movement, speech, and spoken language. It is not the main conceptual anchor for the project, but it is useful to show how a more complete multimodal stack can be built when the goal is practical human-state interpretation.",
            "The key teaching value is that the paper is clearly system-based. It supports the idea that emotion recognition can be used as a data-driven support tool rather than a subjective human-only interpretation process. That makes it a good example when discussing clinical support, evaluation support, or other human-facing contexts where bias and fatigue can distort judgement.",
        ],
        [
            "The work is useful because it shows the practicality of combining facial expressions, speech, spoken language, and body movement. Students often think of multimodal systems as purely academic. This paper helps show the opposite: the same design logic can support real-world interaction settings where subtle nonverbal cues are important.",
            "The preprint also helps explain why transcript-only interpretation is insufficient in emotionally sensitive contexts. If a spoken message is interpreted without nonverbal context, the result can be incomplete or misleading. That lesson is directly relevant to LegalMemoCMT because legal-AI contexts are extremely sensitive to overconfident reading of human behavior.",
        ],
        [
            "It gives an applied, system-level multimodal example.",
            "It supports discussion of body movement as a genuine affective cue.",
            "It makes the case for multi-source evidence rather than single-channel interpretation.",
        ],
        [
            "Because it is a preprint, it should be presented as a developing source rather than as a settled benchmark. It also should not be used to overstate what emotion recognition can reliably infer in sensitive settings. Students should be reminded that more modalities can increase coverage, but they do not eliminate ambiguity.",
        ],
        [
            "For LegalMemoCMT, this preprint can be used as a supplementary example when the student needs to see a broader pipeline than the base MemoCMT paper provides. It is not a replacement for the main references, but it is useful for explaining how a system can integrate face, body, speech, and language into one analysis stack.",
        ],
    )


def add_gap_reading_plan(doc: Document) -> None:
    doc.add_heading("4. Reading Plan to Close the Gaps", level=1)
    add_paragraphs(
        doc,
        [
            "This section turns the gap analysis into a concrete study plan. The goal is not to read everything at once. The goal is to read the papers that directly answer the limitations of the six core references and to use them as supporting literature for LegalMemoCMT Phase 2.",
            "The reading order below is organized by problem: missing modalities, conversational reasoning, adaptive fusion, explainability, bias, and legal-domain caution.",
        ],
    )

    doc.add_heading("4.1 To Extend the Trimodal Baseline", level=2)
    add_paragraphs(
        doc,
        [
            "These papers help close the gaps in Yan et al. (2024), where the main limitations are classical fusion, acted datasets, and weak support for missing or noisy modalities.",
        ],
    )
    for text in [
        "Dynamic Modality and View Selection for Multimodal Emotion Recognition with Missing Modalities (2024, arXiv): https://arxiv.org/abs/2404.12251. Read this for dynamic modality selection when one stream is missing or degraded.",
        "Multimodal Prompt Learning with Missing Modalities for Sentiment Analysis and Emotion Recognition (2024, arXiv): https://arxiv.org/abs/2407.05374. Read this for prompt-based missing-modality generation and compact adaptation.",
        "Leveraging Retrieval Augment Approach for Multimodal Emotion Recognition Under Missing Modalities (2024, arXiv): https://arxiv.org/abs/2410.02804. Read this for retrieval-assisted recovery when one modality is absent.",
        "Enhancing Emotion Recognition in Incomplete Data: A Novel Cross-Modal Alignment, Reconstruction, and Refinement Framework (2024, arXiv): https://arxiv.org/abs/2407.09029. Read this for reconstruction and refinement under incomplete inputs.",
        "EMERSK -- Explainable Multimodal Emotion Recognition with Situational Knowledge (2023, arXiv): https://arxiv.org/abs/2306.08657. Read this for modular multimodal explanation and situational context.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("4.2 To Strengthen the Speech Review and Conversational Scope", level=2)
    add_paragraphs(
        doc,
        [
            "These papers support Alhussein et al. (2025) by adding current conversational fusion and a newer survey of multimodal emotion recognition in conversations.",
        ],
    )
    for text in [
        "Multimodal Emotion Recognition in Conversations: A Survey of Methods, Trends, Challenges and Prospects (2025, arXiv): https://arxiv.org/abs/2505.20511. Read this as the newest high-level survey for conversational multimodal emotion recognition.",
        "Multi-modal emotion recognition in conversation based on prompt learning with text-audio fusion features (2025, Scientific Reports): https://www.nature.com/articles/s41598-025-89758-8. Read this for prompt-based text-audio fusion in conversational ERC.",
        "Cross-modal gated feature enhancement for multimodal emotion recognition in conversations (2025, Scientific Reports): https://www.nature.com/articles/s41598-025-11989-6. Read this for gated fusion and robust cross-modal interaction in ERC.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("4.3 To Upgrade Transformer and Graph Fusion", level=2)
    add_paragraphs(
        doc,
        [
            "These papers help close the gap in Jin et al. (2025), where the model is strong but still benchmark-centric and not yet adapted to legal conversation structure.",
        ],
    )
    for text in [
        "Adaptive multimodal transformer based on exchanging for multimodal sentiment analysis (2025, Scientific Reports): https://www.nature.com/articles/s41598-025-11848-4. Read this for adaptive exchange fusion and modality interaction efficiency.",
        "Hierarchical cross-modal attention and dual audio pathways for enhanced multimodal sentiment analysis (2025, Scientific Reports): https://www.nature.com/articles/s41598-025-09000-3. Read this for hierarchical attention and multi-path audio design that can inspire more robust fusion.",
        "Cross-modal gated feature enhancement for multimodal emotion recognition in conversations (2025, Scientific Reports): https://www.nature.com/articles/s41598-025-11989-6. Read this again as a practical comparison point for gated fusion versus graph-heavy fusion.",
        "Multi-modal emotion recognition in conversation based on prompt learning with text-audio fusion features (2025, Scientific Reports): https://www.nature.com/articles/s41598-025-89758-8. Read this for prompt-guided cross-fusion in ERC pipelines.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("4.4 To Adapt MemoCMT for Phase 2 and Missing Modalities", level=2)
    add_paragraphs(
        doc,
        [
            "These papers are the most important follow-ups for the Phase 1 to Phase 2 transition, because MemoCMT is the base paper and the legal project must become more robust to partial evidence.",
        ],
    )
    for text in [
        "Multimodal Prompt Learning with Missing Modalities for Sentiment Analysis and Emotion Recognition (2024, arXiv): https://arxiv.org/abs/2407.05374. Read this for parameter-efficient missing-modality handling.",
        "Dynamic Modality and View Selection for Multimodal Emotion Recognition with Missing Modalities (2024, arXiv): https://arxiv.org/abs/2404.12251. Read this for dynamic modality routing in the face of incomplete evidence.",
        "Leveraging Retrieval Augment Approach for Multimodal Emotion Recognition Under Missing Modalities (2024, arXiv): https://arxiv.org/abs/2410.02804. Read this for retrieval-augmented robustness when inputs are missing.",
        "Enhancing Emotion Recognition in Incomplete Data: A Novel Cross-Modal Alignment, Reconstruction, and Refinement Framework (2024, arXiv): https://arxiv.org/abs/2407.09029. Read this for reconstruction and refinement of missing modality embeddings.",
        "Adaptive multimodal transformer based on exchanging for multimodal sentiment analysis (2025, Scientific Reports): https://www.nature.com/articles/s41598-025-11848-4. Read this for a stronger adaptive fusion model than a fixed pooling baseline.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("4.5 To Cover Fairness, Bias, Explainability, and Legal Caution", level=2)
    add_paragraphs(
        doc,
        [
            "These papers are essential because LegalMemoCMT is not a general benchmark-only system. It is a legal-domain support project, so interpretability and bias controls are part of the core research scope.",
        ],
    )
    for text in [
        "Emo-bias: A Large Scale Evaluation of Social Bias on Speech Emotion Recognition (2024, arXiv): https://arxiv.org/abs/2406.05065. Read this to understand gender bias, dataset bias, and model bias in SER.",
        "A Review on Explainability in Multimodal Deep Neural Nets (2021, arXiv): https://arxiv.org/abs/2105.07878. Read this for the general explainability toolkit in multimodal deep models.",
        "EMERSK -- Explainable Multimodal Emotion Recognition with Situational Knowledge (2023, arXiv): https://arxiv.org/abs/2306.08657. Read this for modular explanation generation in multimodal emotion recognition.",
        "Interpretable Multimodal Emotion Recognition using Facial Features and Physiological Signals (2023, arXiv): https://arxiv.org/abs/2306.02845. Read this for modality contribution analysis and interpretability methods.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    add_paragraphs(
        doc,
        [
            "The practical reading strategy is: start with the core six papers, then read the missing-modality papers, then read the conversation surveys and fusion papers, then finish with explainability and bias papers. That order tracks the actual risk profile of the project and keeps Phase 2 grounded in methods that can survive partial evidence and legal scrutiny.",
        ],
    )


def add_section_four(doc: Document) -> None:
    doc.add_heading("5. Fairness, Bias, Explainability, and Legal Limits", level=1)
    add_paragraphs(
        doc,
        [
            "A legal-AI project must be much more careful than a standard benchmark project. Emotion inference is inherently uncertain, culturally dependent, and context sensitive. A facial expression can signal stress, concentration, masking, fatigue, social politeness, neurodiversity, or a learned professional demeanor. The same is true for voice tone and language choice. Because of that, the model output should never be treated as a direct statement of truth, intent, honesty, credibility, or guilt.",
            "Bias enters through data selection, label construction, and model design. If the training corpus overrepresents a narrow cultural group, the system may encode that group as the implicit default. If annotators are asked to map nuanced expressions into coarse emotion labels, the labels can collapse important diversity. If one modality is missing or noisy more often for certain groups, the model may learn performance differences that look technical but are actually social inequities.",
            "Explainability should therefore be built into the project from the beginning. The system should expose which modality contributed most to a prediction, what the attention or fusion mechanism emphasized, and what uncertainty exists around the decision. The student should be taught to distinguish interpretability from truth. A readable explanation of model behavior does not mean the model is correct. It means the model is inspectable.",
            "The legal limit is simple: emotional inference may support interpretation of a communication pattern, but it should not be used to replace legal reasoning or human judgment. In judicial settings, emotion recognition can become dangerous if it is treated as evidence of intent, deception, or reliability without validation. The project should explicitly state that its outputs are descriptive and analytical, not adjudicative.",
        ],
    )
    bullets = [
        "Prefer calibrated probabilities or uncertainty bands instead of single hard claims when presenting emotion predictions.",
        "Document dataset composition, class balance, annotation source, and the likely cultural scope of the training data.",
        "Use ablations and modality contribution analysis to show what the model is actually using.",
        "Avoid language that suggests the system can read a person's mind, legal culpability, or moral character.",
        "Treat disagreement between modalities as a meaningful research signal, not as proof that one modality is false.",
        "Include a clear ethical note that the system is a support tool for interpretation, not a decision engine for law or discipline.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")


def add_section_five(doc: Document) -> None:
    doc.add_heading("6. How This Maps to LegalMemoCMT", level=1)
    add_paragraphs(
        doc,
        [
            "The best way to position LegalMemoCMT is as a teaching and analysis framework that connects speech, video, and text in a multimodal pipeline while staying faithful to the MemoCMT design philosophy. The base paper gives the audio-text cross-modal transformer core. The broader project should explain how a visual branch can be added conceptually and how the resulting system can be used to interpret patterns rather than to issue normative judgments.",
            "In a project report, the sequence should be: first explain the unimodal encoders, then explain the cross-modal alignment mechanism, then explain the conversational or temporal structure, then explain the output analysis and uncertainty presentation. This order matters because it mirrors how students actually learn the subject.",
            "For implementation, the student should think in modules. A speech module extracts paralinguistic cues. A video module extracts facial and body motion cues. A text module extracts semantic and discourse cues. A fusion module aligns these streams. A reporting layer then shows the predicted emotion, the confidence, and the modalities most responsible for the result. That reporting layer is what makes the project more useful for human interpretation.",
            "The project should also distinguish between two levels of explanation. The first level is model explanation: what the network did. The second level is domain explanation: what the result might mean in a real conversation. The second level must remain cautious, especially in legal contexts, because emotion is not equivalent to intent, credibility, or liability.",
        ],
    )
    bullets = [
        "Use MemoCMT as the architectural anchor for speech and text fusion.",
        "Use the trimodal and graph-fusion papers to justify visual and conversational extensions.",
        "Use the survey papers to justify dataset and metric selection.",
        "Use the fairness section to constrain the system's legal-language framing.",
        "Keep the final deliverable educational: it should teach how the model works, not promise that the model can make legal decisions.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")


def add_phase2_section(doc: Document) -> None:
    doc.add_heading("7. Phase 2 Gap Analysis and Research Directions", level=1)
    add_paragraphs(
        doc,
        [
            "Phase 2 should not be framed as a simple performance upgrade. It should be framed as a controlled research extension that answers the limitations left open by the current multimodal emotion recognition literature and converts them into a legal-domain analysis system. The key shift is from benchmark emotion classification to explainable, uncertainty-aware, legally cautious emotional pattern analysis.",
            "The main research question for Phase 2 is: how can a multimodal model estimate emotional patterns in legal conversations, documents, or hearings while remaining robust to missing modalities, speaker variability, domain shift, and fairness concerns? That question sits directly on top of the gaps identified in the literature.",
        ],
    )

    doc.add_heading("7.1 Cross-Paper Gaps That Matter Most", level=2)
    gap_bullets = [
        "Domain shift: public datasets are acted, curated, or conversationally narrow, while legal data are noisy, procedural, and high-stakes.",
        "Missing modalities: many systems assume that audio, text, and video are all available and clean at the same time.",
        "Weak long-context reasoning: many models still struggle with long hearings, long dialogue histories, or turn-by-turn emotional drift.",
        "Limited explainability: attention maps alone are not sufficient for legal use because they do not explain why the model trusted one modality over another.",
        "Bias and fairness risks: emotion labels can be culturally biased, speaker-biased, accent-biased, and recording-quality-biased.",
        "Calibration problems: a model can be accurate on average but still badly overconfident on uncertain or out-of-domain legal data.",
        "Incomplete legal framing: the literature rarely states clearly that emotion recognition should not be used to infer credibility, guilt, or legal intent.",
    ]
    for item in gap_bullets:
        doc.add_paragraph(item, style="List Bullet")

    add_paragraphs(
        doc,
        [
            "The practical implication is that Phase 2 needs a research design that is more than a classifier. It should behave like an evidential analysis system: it should ingest multiple modalities, detect when evidence is incomplete, explain which evidence stream mattered, and refuse to overstate the result when the context is weak.",
        ],
    )

    def add_recent_advances_subsection(title: str, gaps, advances):
        doc.add_heading(title, level=2)
        doc.add_heading("Gap Profile", level=3)
        for g in gaps:
            doc.add_paragraph(g, style="List Bullet")
        doc.add_heading("Recent Advances Now Allow", level=3)
        for a in advances:
            doc.add_paragraph(a["heading"], style="List Bullet")
            add_paragraphs(doc, a["paras"])

    add_recent_advances_subsection(
        "7.2 Yan et al. 2024: Trimodal Face, Speech, and Body Gestures",
        [
            "The paper relies on relatively small and acted datasets, which limits robustness in real conversational environments.",
            "Its fusion mechanism is mathematically strong but not designed for modern large-scale, missing-modality, or legal-domain settings.",
            "The work does not model dialogue structure, speaker role, or long temporal context.",
        ],
        [
            {
                "heading": "Multimodal pretraining and self-supervised encoders now make richer visual and acoustic features feasible.",
                "paras": [
                    "In 2024-era trimodal work, the biggest weakness is not the idea of fusion but the quality of the modality representations. Modern vision and speech encoders can now learn stronger embeddings from raw or lightly processed inputs. For Phase 2, this means the face branch can be upgraded from handcrafted or shallow visual descriptors to pretrained facial-video backbones, while the speech branch can use robust self-supervised speech representations and the body branch can use pose or landmark embeddings.",
                    "This is relevant for legal settings because law-related recordings are often imperfect: low frame rate, noisy audio, partial face visibility, and inconsistent lighting. A stronger pretrained front-end reduces the risk that the legal system will depend on brittle handcrafted signals.",
                ],
            },
            {
                "heading": "Missing-modality training now allows the system to work even when one channel is absent.",
                "paras": [
                    "The original trimodal pipeline assumes that all three channels are present and usable. That assumption is weak for legal evidence. In a deposition, one may only have audio and transcript. In a hearing, the video may be partial or obstructed. Modern training with modality dropout, partial fusion, or learned modality gating lets the model function under such realistic constraints.",
                    "This is not a minor engineering trick. It changes the legal usability of the model. A system that fails whenever the face stream is missing cannot support real-world evidence analysis. A Phase 2 design should therefore explicitly train on partial-modality conditions and report what confidence it has under each input scenario.",
                ],
            },
            {
                "heading": "Modern cross-modal attention can replace or augment classical correlation fusion.",
                "paras": [
                    "SLSMKCCA and SSLSMKCCA are elegant, but current fusion methods can represent non-linear interactions more flexibly. Cross-modal transformers, gated cross-attention, and multimodal adapters can learn when body motion strengthens or contradicts speech and facial cues. This is especially important in legal contexts where disguise, restraint, and social performance are common.",
                    "The practical advantage is not just accuracy. It is the ability to inspect cross-modal attention patterns and use them as part of an explanation layer. A legal user can be shown that the model relied more on vocal strain and posture than on a neutral transcript, which is far more informative than a raw label.",
                ],
            },
            {
                "heading": "The body-gesture gap can now be addressed with pose estimation and multimodal video modeling.",
                "paras": [
                    "One of the historical gaps in trimodal emotion recognition is the difficulty of reliably extracting body cues. That gap is narrower now because modern pose estimation, dense tracking, and video transformers can produce more stable motion representations. The body branch can therefore move from fragile low-level signals to structured temporal keypoints or motion embeddings.",
                    "For Phase 2, this matters because legal conversation often includes stress behaviors that appear in posture, hand motion, or self-adaptors. The model should not be asked to make claims about intent from these cues, but they can be useful for pattern analysis if the system clearly marks uncertainty and context dependence.",
                ],
            },
            {
                "heading": "Phase 2 Research Use",
                "paras": [
                    "Use Yan et al. as the rationale for a legal trimodal extension: face, speech, and body motion should be treated as complementary evidence streams. The Phase 2 model should not assume complete inputs, should not depend only on classical fusion, and should always expose which modalities contributed to the prediction.",
                ],
            },
        ],
    )

    add_recent_advances_subsection(
        "7.3 Alhussein et al. 2025: Systematic Review and Meta-Analysis of Speech Emotion Recognition",
        [
            "The review exposes weak reporting standards and high heterogeneity in the field.",
            "It highlights that many systems still rely on datasets and protocols that are not comparable.",
            "It shows that speech emotion recognition remains sensitive to feature choices, annotation schemes, and conversational context.",
        ],
        [
            {
                "heading": "Foundation speech encoders now reduce dependence on fragile handcrafted feature sets.",
                "paras": [
                    "The review notes that deep features and hybrid approaches are increasingly favored. That trend is stronger today because pretrained speech models learn richer representations from large-scale unlabeled audio. For Phase 2, this means the speech branch can be built on top of a robust speech encoder instead of depending entirely on MFCCs or manually engineered prosody statistics.",
                    "In a legal setting, this is important because speech is often noisy, interrupted, and domain-specific. A pretrained encoder can better survive recording artifacts, accents, and variable speaking rates, while still preserving emotional timing and speaker style. However, these encoders should be fine-tuned carefully so they do not simply memorize benchmark regularities.",
                ],
            },
            {
                "heading": "Transformer-era conversational models now make speaker and turn context tractable.",
                "paras": [
                    "A major limitation in earlier speech emotion work is that it often treats utterances as isolated examples. Current transformer-based conversational models can keep track of turn history, speaker identity, and contextual dependencies much more effectively. This means Phase 2 can model not only the current utterance but the buildup before it and the reaction after it.",
                    "That is directly relevant to legal dialogues, where a statement may be emotionally neutral in isolation but highly charged in context. Question-answer sequences, objections, interruptions, and repeated prompts all change how speech should be interpreted. Context-aware speech modeling is therefore a necessity, not an optional enhancement.",
                ],
            },
            {
                "heading": "Better evaluation practice now supports fairness, calibration, and subgroup analysis.",
                "paras": [
                    "The review emphasizes the weaknesses in reporting and methodological transparency. Recent AI practice gives us better tools to address this: calibration curves, confidence intervals, subgroup performance by speaker group or recording condition, and error analysis under noise. These tools matter because legal-AI systems cannot afford to hide their failure modes.",
                    "Phase 2 should use these tools to show whether the speech branch is systematically less reliable for certain accents, genders, or speech styles. The goal is not to enforce demographic prediction. The goal is to detect whether the model is learning unfair shortcuts that would be problematic in a legal environment.",
                ],
            },
            {
                "heading": "Hybrid feature strategies now let you combine interpretable acoustic cues with deep features.",
                "paras": [
                    "The review suggests that handcrafted expert knowledge still has value. Recent systems can combine deep embeddings with explicit prosodic cues, giving both performance and interpretability. For Phase 2, this is valuable because a legal user may want to know whether the model reacted to speech rate, pitch rise, energy spikes, or pauses.",
                    "A hybrid design also helps with auditability. If the deep model flags an emotionally charged segment, the system can highlight the corresponding acoustic cues. This is not proof of emotion, but it is a more inspectable path than a black-box score alone.",
                ],
            },
            {
                "heading": "Phase 2 Research Use",
                "paras": [
                    "Use the review to justify a context-aware legal speech model with pretrained encoders, explicit calibration, and subgroup analysis. The speech branch should be able to say not only what emotion pattern it found, but also how stable that pattern is under noise and contextual variation.",
                ],
            },
        ],
    )

    add_recent_advances_subsection(
        "7.4 Jin et al. 2025: Transformer and Graph Fusion for Conversational Multimodal Emotion Recognition",
        [
            "The model is strong, but it is still benchmark-centric rather than legal-domain-specific.",
            "It improves context modeling, but it does not explicitly solve legal discourse structure.",
            "Its graph design assumes a conversation graph, not a procedurally structured legal interaction graph.",
        ],
        [
            {
                "heading": "Graph transformers now allow heterogeneous legal conversation structure.",
                "paras": [
                    "The paper shows that combining transformer reasoning with graph reasoning is effective. Since then, graph transformers and heterogeneous graph models have become even more useful. For Phase 2, this means the conversation graph can include speaker role, turn type, question, answer, interruption, pause, and procedural metadata as explicit relation types.",
                    "In legal settings, this is important because the same statement can carry different emotional meaning depending on whether it is an answer, a challenge, a clarification, or an objection. A heterogeneous legal graph lets the model preserve those distinctions instead of flattening everything into generic utterances.",
                ],
            },
            {
                "heading": "Long-context attention can complement graph locality.",
                "paras": [
                    "The paper already separates global transformer reasoning from local graph reasoning. Recent advances in long-context modeling make the global side stronger. For Phase 2, this means the model can look across longer transcripts or hearings while still using graph links to preserve local speaker transitions and dialogue structure.",
                    "This is particularly relevant for legal cases, where a key emotional pattern may emerge only after many turns. A model that only sees a narrow window may misread cumulative stress, hesitation, or escalation.",
                ],
            },
            {
                "heading": "Cross-modal and graph-based explanations are now more usable.",
                "paras": [
                    "The hybrid structure is naturally explainable if used properly. The transformer branch can show broad cross-modal influence, while the graph branch can show which conversational edges were important. Together, they enable a more defensible narrative than raw attention alone.",
                    "For Phase 2, this supports a legal-facing explanation layer. The system can say that the emotional pattern was driven by a repeated question sequence, a delayed response, and a visible voice-energy shift, instead of just outputting 'anxious'.",
                ],
            },
            {
                "heading": "Statistical rigor and variance reporting are now easier to incorporate.",
                "paras": [
                    "The paper's modest performance gains are a reminder that legal research should not overclaim. Modern ML practice encourages multi-run reporting, significance tests, and variance analysis. Phase 2 should follow that standard because legal applications need reliability, not leaderboard chasing.",
                    "If a model's performance varies sharply across folds or speakers, that instability must be visible in the report. That is a core requirement for any system that could influence human interpretation in a legal context.",
                ],
            },
            {
                "heading": "Phase 2 Research Use",
                "paras": [
                    "Use this paper to justify a legal-conversation graph with procedural roles and long-range transformer context. The legal version should encode question-answer structure, interruptions, and speaker roles explicitly, then fuse them with multimodal evidence streams for cautious emotional pattern analysis.",
                ],
            },
        ],
    )

    add_recent_advances_subsection(
        "7.5 MemoCMT 2025: Cross-Modal Transformer Fusion",
        [
            "The paper already notes high compute cost and dependence on complete input modalities.",
            "It is limited in real-time and constrained-device settings.",
            "It does not include video, and its explainability remains partial.",
        ],
        [
            {
                "heading": "Efficient multimodal adaptation now reduces compute overhead.",
                "paras": [
                    "Transformer fusion used to imply a heavy end-to-end model. Current efficient training and adaptation techniques make it easier to extend a pretrained core with lightweight modality adapters, low-rank modules, or parameter-efficient fine-tuning. For Phase 2, this means the legal system does not have to retrain a huge model from scratch just to add a new modality or domain.",
                    "This matters because legal data are usually scarce and sensitive. A parameter-efficient approach allows the project to start from a strong pretrained backbone and adapt to legal-style data with lower risk of overfitting and lower engineering cost.",
                ],
            },
            {
                "heading": "Missing-modality robustness is now a central research area.",
                "paras": [
                    "MemoCMT notes that the original system expects full modality inputs. That limitation is now being addressed through modality dropout, learned imputation, mixture-of-experts routing, and fallback fusion. For Phase 2, this is crucial because legal data often have partial evidence: transcript only, audio with weak video, or video without reliable speech transcripts.",
                    "A robust legal system should therefore train under different modality availability patterns and report separate performance for full, partial, and degraded input conditions. This makes the tool more realistic and more honest about its limitations.",
                ],
            },
            {
                "heading": "Long-context and hierarchical transformers now make hearing-length data more tractable.",
                "paras": [
                    "The MemoCMT architecture is effective, but it is still bounded by sequence length and compute. Recent long-context and hierarchical attention methods can segment long hearings into chunks, summarize each chunk, and then reason across the summaries. That is much closer to how legal conversations actually unfold.",
                    "For Phase 2, this is especially important if the goal is to analyze emotional trends across a hearing or interview rather than a single utterance. The model should be able to report local spikes and global trajectories separately.",
                ],
            },
            {
                "heading": "Explainability now has stronger tooling than simple attention inspection.",
                "paras": [
                    "MemoCMT already discusses representation analysis, but modern explanation methods go further. Token and modality attribution, counterfactual perturbation, and uncertainty estimation can show whether the prediction changed because of audio tone, text content, or cross-modal conflict. This is much more useful for legal analysis than a single opaque class score.",
                    "In Phase 2, the explanation layer should allow the user to inspect the model's reasoning at the level of modality contribution and confidence, while still avoiding any claim that the model has discovered legal truth.",
                ],
            },
            {
                "heading": "Phase 2 Research Use",
                "paras": [
                    "Use MemoCMT as the audio-text core and extend it with video, missing-modality training, and long-context handling. The legal version should be lighter, more robust, and more explicit about uncertainty than the benchmark version.",
                ],
            },
        ],
    )

    add_recent_advances_subsection(
        "7.6 Hu et al. 2024: Multimodal Affective Computing Survey",
        [
            "The survey is broad, but it is not implementation-oriented.",
            "Its task taxonomy is useful, yet the legal domain is not central.",
            "It does not itself provide an operational framework for explainability or fairness.",
        ],
        [
            {
                "heading": "Unified task taxonomies now make legal task definition easier.",
                "paras": [
                    "A key contribution of the survey is that it organizes the field into task families. That structure is now even more useful because legal Phase 2 needs a narrow, defensible task definition. Instead of 'detect emotion', the system can be defined as estimating emotional pattern shifts, tension, uncertainty, or escalation in context.",
                    "This task narrowing is important because legal applications should not over-claim. A precise task definition reduces the chance that the model output will be interpreted as a legal judgment rather than an analytical signal.",
                ],
            },
            {
                "heading": "Modern benchmarking now supports richer multi-metric evaluation.",
                "paras": [
                    "The survey emphasizes datasets and metrics. Current practice extends this by encouraging metric bundles: accuracy, macro F1, weighted F1, calibration, and subgroup performance. For Phase 2, this allows the project to report not only how often the model is right, but how stable and fair it is across evidence conditions.",
                    "That matters for legal deployment because a single headline number is inadequate. A legal support system must be evaluated under uncertainty and across user groups, modalities, and dialogue types.",
                ],
            },
            {
                "heading": "Retrieval and instruction-based explanation layers are now practical.",
                "paras": [
                    "The survey helps students see the field's broader structure. Today, a multimodal system can be paired with an explanation layer that retrieves supporting examples, summarizes evidence, and generates a plain-language rationale for the model's output. For Phase 2, this makes the system easier to inspect by non-technical users.",
                    "The explanation layer should remain descriptive. It can say which cues were detected and how they relate to the emotion pattern, but it must not convert affective inference into legal characterization.",
                ],
            },
            {
                "heading": "Phase 2 Research Use",
                "paras": [
                    "Use the survey to justify the task definition, metrics, and chapter structure of Phase 2. The legal system should be positioned as a constrained multimodal affective analysis task, not a general-purpose legal reasoning engine.",
                ],
            },
        ],
    )

    add_recent_advances_subsection(
        "7.7 Kraack 2024 Preprint: Full Multimodal Emotion Recognition System",
        [
            "The preprint is useful as a system-level example, but it is less mature than the main benchmark papers.",
            "It does not define a legal-domain protocol or legal safety framework.",
            "Its value is architectural breadth rather than methodological completeness.",
        ],
        [
            {
                "heading": "System integration is now easier through modular multimodal pipelines.",
                "paras": [
                    "The broad trimodal and quad-modal structure of the preprint becomes easier to build now because modern tooling supports modular encoders and fusion adapters. Phase 2 can therefore integrate facial, body, speech, and language streams without forcing all branches to be trained identically.",
                    "This modularity matters in legal settings because different cases will have different data quality. The system should be able to use whatever evidence is present and degrade gracefully when a branch is unavailable.",
                ],
            },
            {
                "heading": "Body-movement modeling is now more precise and less fragile.",
                "paras": [
                    "The body-movement gap is smaller today because pose estimation, keypoint tracking, and motion transformers can provide better structural cues than earlier systems. For Phase 2, this means gestures, posture shifts, and self-adaptors can be represented in a way that is more stable and easier to analyze.",
                    "In legal settings, the model should still be cautious. Body motion should be treated as one interpretive cue among many, not as proof of deceptive behavior or intent.",
                ],
            },
            {
                "heading": "Human-facing explanation layers are more important than ever.",
                "paras": [
                    "The preprint's integrated perspective supports a human-facing workflow. Phase 2 can add a presentation layer that shows the emotional trajectory across modalities and time, which is especially useful when the user needs to review an interaction rather than issue a verdict.",
                    "This is important because legal-domain users need reviewability. A good system should let them see which modality was strong, which was weak, and where the modalities conflicted.",
                ],
            },
            {
                "heading": "Phase 2 Research Use",
                "paras": [
                    "Use this preprint as a supplementary example when extending the architecture beyond MemoCMT. It supports the case for a more complete multimodal stack, but the legal thesis should still anchor itself in the more methodologically rigorous papers above.",
                ],
            },
        ],
    )

    doc.add_heading("7.8 Practical Phase 2 Direction", level=2)
    add_paragraphs(
        doc,
        [
            "The strongest Phase 2 design is a legal-aware multimodal analysis pipeline with the following properties: modality-specific encoders, context-aware fusion, missing-modality robustness, speaker-aware temporal modeling, calibrated outputs, and explanation tooling. The model should be used to study emotional patterns in legal interaction, not to infer legal conclusions.",
            "If the project is implemented well, Phase 2 can answer a meaningful research question: how much emotional structure can be recovered from legal communication when the system is explicitly designed to be robust, cautious, and interpretable? That is a real AI/ML research problem, and it is more defensible than claiming that the model can judge people.",
        ],
    )


def add_reference_table_section(doc: Document) -> None:
    doc.add_heading("8. Reference Table for Reading Plan", level=1)
    add_paragraphs(
        doc,
        [
            "This table collects the additional reading articles and journals that were selected to fill the gaps identified in the core paper-by-paper teaching notes. The links are provided directly so the reading list can be used immediately during thesis preparation.",
        ],
    )

    headers = ["Gap area", "Article", "Link", "Why to read"]
    rows = [
        [
            "Missing modalities",
            "Dynamic Modality and View Selection for Multimodal Emotion Recognition with Missing Modalities",
            "https://arxiv.org/abs/2404.12251",
            "Learn dynamic modality selection when one stream is missing or degraded.",
        ],
        [
            "Missing modalities",
            "Multimodal Prompt Learning with Missing Modalities for Sentiment Analysis and Emotion Recognition",
            "https://arxiv.org/abs/2407.05374",
            "See prompt-based missing-modality generation and compact adaptation.",
        ],
        [
            "Missing modalities",
            "Leveraging Retrieval Augment Approach for Multimodal Emotion Recognition Under Missing Modalities",
            "https://arxiv.org/abs/2410.02804",
            "Study retrieval-assisted recovery when one modality is absent.",
        ],
        [
            "Missing modalities",
            "Enhancing Emotion Recognition in Incomplete Data: A Novel Cross-Modal Alignment, Reconstruction, and Refinement Framework",
            "https://arxiv.org/abs/2407.09029",
            "Read for reconstruction and refinement under incomplete inputs.",
        ],
        [
            "Explainability",
            "EMERSK -- Explainable Multimodal Emotion Recognition with Situational Knowledge",
            "https://arxiv.org/abs/2306.08657",
            "Use for modular multimodal explanation and situational context.",
        ],
        [
            "Conversational MER",
            "Multimodal Emotion Recognition in Conversations: A Survey of Methods, Trends, Challenges and Prospects",
            "https://aclanthology.org/2025.findings-emnlp.332/",
            "Use as the newest high-level survey for MERC.",
        ],
        [
            "Conversational MER",
            "Multi-modal emotion recognition in conversation based on prompt learning with text-audio fusion features",
            "https://www.nature.com/articles/s41598-025-89758-8",
            "Read for prompt-based text-audio fusion in conversational ERC.",
        ],
        [
            "Conversational MER",
            "Cross-modal gated feature enhancement for multimodal emotion recognition in conversations",
            "https://www.nature.com/articles/s41598-025-11989-6",
            "Learn gated fusion and cross-modal interaction in ERC.",
        ],
        [
            "Fusion / adaptation",
            "Adaptive multimodal transformer based on exchanging for multimodal sentiment analysis",
            "https://www.nature.com/articles/s41598-025-11848-4",
            "Study adaptive exchange fusion beyond fixed pooling baselines.",
        ],
        [
            "Fusion / adaptation",
            "Hierarchical cross-modal attention and dual audio pathways for enhanced multimodal sentiment analysis",
            "https://www.nature.com/articles/s41598-025-09000-3",
            "See hierarchical attention and stronger audio processing design.",
        ],
        [
            "Bias / fairness",
            "Emo-bias: A Large Scale Evaluation of Social Bias on Speech Emotion Recognition",
            "https://arxiv.org/abs/2406.05065",
            "Understand gender bias, dataset bias, and model bias in SER.",
        ],
        [
            "Explainability",
            "A Review on Explainability in Multimodal Deep Neural Nets",
            "https://arxiv.org/abs/2105.07878",
            "Get the general explainability toolkit for multimodal deep models.",
        ],
        [
            "Explainability",
            "Interpretable Multimodal Emotion Recognition using Facial Features and Physiological Signals",
            "https://arxiv.org/abs/2306.02845",
            "Study modality contribution analysis and interpretability methods.",
        ],
    ]
    add_table(doc, headers, rows)


def add_phase1_reference_table_section(doc: Document) -> None:
    doc.add_heading("9. Phase 1 Main Reference Table", level=1)
    add_paragraphs(
        doc,
        [
            "This table lists the main reference articles that directly support Phase 1 implementation, architecture, and evaluation. These are the papers that should stay closest to the current codebase and benchmark framing.",
        ],
    )

    headers = ["Phase 1 role", "Article", "Link", "Why it matters for Phase 1"]
    rows = [
        [
            "Base paper (primary reference)",
            "MemoCMT: multimodal emotion recognition using cross-modal transformer-based feature fusion",
            "https://www.nature.com/articles/s41598-025-89202-x",
            "This is the main Phase 1 architecture reference and the direct source for the CMT design.",
        ],
        [
            "Trimodal baseline",
            "Multimodal Emotion Recognition Based on Facial Expressions, Speech, and Body Gestures",
            "https://www.mdpi.com/2079-9292/13/18/3756",
            "This is the strongest classical trimodal baseline for face, speech, and body-gesture fusion.",
        ],
        [
            "Speech and conversation",
            "Speech emotion recognition in conversations using artificial intelligence: a systematic review and meta-analysis",
            "https://link.springer.com/article/10.1007/s10462-025-11197-8",
            "This survey supports the speech branch, metrics, and conversational evaluation framing.",
        ],
        [
            "Conversational multimodal fusion",
            "Multimodal Emotion Recognition in Conversations Using Transformer and Graph Neural Networks",
            "https://www.mdpi.com/2076-3417/15/22/11971",
            "This is a strong reference for dialogue structure, transformer fusion, and graph reasoning.",
        ],
        [
            "Survey backbone",
            "Recent Trends of Multimodal Affective Computing: A Survey from NLP Perspective",
            "https://arxiv.org/abs/2409.07388",
            "This survey helps define the task taxonomy, datasets, and metric framing for the thesis narrative.",
        ],
        [
            "Broader multimodal system",
            "A Multimodal Emotion Recognition System: Integrating Facial Expressions, Body Movement, Speech, and Spoken Language",
            "https://arxiv.org/abs/2412.17907",
            "This gives a broader applied multimodal example beyond the MemoCMT core.",
        ],
    ]
    add_table(doc, headers, rows)


def add_references(doc: Document) -> None:
    doc.add_heading("10. Source Links and Reference Notes", level=1)
    refs = [
        "Yan et al. (2024), Multimodal Emotion Recognition Based on Facial Expressions, Speech, and Body Gestures: https://www.mdpi.com/2079-9292/13/18/3756",
        "Alhussein et al. (2025), Speech emotion recognition in conversations using artificial intelligence: a systematic review and meta-analysis: https://link.springer.com/article/10.1007/s10462-025-11197-8",
        "Jin et al. (2025), Multimodal Emotion Recognition in Conversations Using Transformer and Graph Neural Networks: https://www.mdpi.com/2076-3417/15/22/11971",
        "Khan et al. (2025), MemoCMT: multimodal emotion recognition using cross-modal transformer-based feature fusion: https://www.nature.com/articles/s41598-025-89202-x.pdf",
        "Hu et al. (2024), Recent Trends of Multimodal Affective Computing: A Survey from NLP Perspective: https://arxiv.org/abs/2409.07388",
        "Kraack (2024 preprint), A Multimodal Emotion Recognition System: Integrating Facial Expressions, Body Movement, Speech, and Spoken Language: https://arxiv.org/abs/2412.17907",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Bullet")


def main():
    doc = Document()
    set_document_defaults(doc)
    add_title_page(doc)
    add_objectives(doc)
    add_section_one(doc)
    add_section_two(doc)
    add_section_three(doc)
    add_gap_reading_plan(doc)
    add_section_four(doc)
    add_section_five(doc)
    add_phase2_section(doc)
    add_reference_table_section(doc)
    add_phase1_reference_table_section(doc)
    add_references(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
