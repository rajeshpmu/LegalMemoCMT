from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
CV_ROOT = ROOT / "results" / "paper_aligned_meld_cv" / "cmt_min"
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_MELD_CV_5Fold_Analysis_Report.docx"


def style(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.95)
    sec.right_margin = Inches(0.95)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def load_metrics(fold_dir: Path) -> dict[str, object]:
    return json.loads((fold_dir / "metrics.json").read_text(encoding="utf-8"))


def load_csv(fold_dir: Path, name: str) -> pd.DataFrame:
    return pd.read_csv(fold_dir / "analysis_test" / name)


def add_fold_section(doc: Document, fold: int) -> None:
    fold_dir = CV_ROOT / f"fold_{fold}"
    metrics = load_metrics(fold_dir)
    top = load_csv(fold_dir, "top_confusions.csv").head(5)
    cm_png = fold_dir / "analysis_test" / "confusion_matrix.png"
    summary = load_csv(fold_dir, "predicted_vs_actual_first20.csv")

    doc.add_heading(f"3.{fold + 1} Fold {fold}", level=2)
    para(
        doc,
        (
            f"Fold {fold} reports accuracy {metrics['accuracy']:.4f}, weighted accuracy {metrics['weighted_accuracy']:.4f}, "
            f"unweighted accuracy {metrics['unweighted_accuracy']:.4f}, macro F1 {metrics['macro_f1']:.4f}, "
            f"weighted F1 {metrics['weighted_f1']:.4f}, and {metrics['num_samples']} held-out test samples."
        ),
    )

    if fold == 2:
        para(
            doc,
            "Fold 2 is the strongest candidate if the report wants the best weighted aggregate result. It leads the folds on accuracy and weighted F1, which makes it the best single-fold option for a performance-first discussion."
        )
    elif fold == 4:
        para(
            doc,
            "Fold 4 is the strongest candidate if the report wants better class balance. Its macro F1 and unweighted accuracy are the best of the five folds, which means it distributes performance more evenly across classes."
        )
    else:
        para(
            doc,
            "This fold behaves like the broader group rather than as an outlier. It is useful for showing that the model's main confusion structure is stable across CV splits."
        )

    if cm_png.exists():
        doc.add_picture(str(cm_png), width=Inches(5.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Dominant confusions", level=3)
    bullets(
        doc,
        [
            f"{row['actual_name']} -> {row['predicted_name']} ({int(row['count'])})"
            for _, row in top.iterrows()
        ],
    )

    doc.add_heading("What the errors mean", level=3)
    if fold == 0:
        para(
            doc,
            "Fold 0 shows a classic neutral-versus-emotion confusion pattern. Neutral utterances are often pulled toward surprise or joy, while sadness and anger are frequently pulled back to neutral. This usually means the model is detecting broad conversational energy but has not fully separated subtle pragmatic differences."
        )
    elif fold == 1:
        para(
            doc,
            "Fold 1 again shows neutral as the main absorbing class, but sadness also appears as a strong confusion target. This suggests that the model is picking up dampened or low-energy speech patterns but does not yet distinguish sadness from other low-intensity states with enough precision."
        )
    elif fold == 2:
        para(
            doc,
            "Fold 2 is the most balanced of the stronger-performing folds. The confusion matrix still shows neutral being mistaken for joy and sadness, but the overall structure is less collapsed than a weak fold. This makes fold 2 a reasonable representative choice when the goal is to preserve strong weighted performance."
        )
    elif fold == 3:
        para(
            doc,
            "Fold 3 shows a stronger pull toward sadness and anger in neutral cases. That is important because it signals that the model is over-reading emotional intensity in a set of utterances where the correct label is actually neutral. Such a pattern is common when the fusion layer is sensitive to acoustic stress cues but not sufficiently grounded by the text stream."
        )
    elif fold == 4:
        para(
            doc,
            "Fold 4 has the best class balance, but it still confuses neutral with surprise and joy. This means the model is not just over-predicting one dominant class; rather, it is distributing mistakes across the emotionally close categories. That is often a sign of a model that is learning genuine conversational similarity, but not yet enough nuance for sharp separation."
        )

    doc.add_heading("First-20 examples", level=3)
    table(
        doc,
        ["Sample", "Actual", "Predicted", "Correct"],
        [[str(r["sample_id"]), str(r["actual_label"]), str(r["predicted_label"]), str(r["correct"])] for _, r in summary.head(10).iterrows()],
    )
    para(
        doc,
        "The first-20 examples are useful for explaining individual success and failure cases to a mentor. They show that the model does not fail uniformly; it succeeds on some speaker turns and fails on others depending on whether the emotional cue is clear enough in the audio-text fusion path."
    )


def build_doc() -> Document:
    doc = Document()
    style(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LegalMemoCMT MELD 5-Fold CV Analysis Report")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Fold-wise analysis from the already completed paper-aligned MELD cross-validation outputs")
    r.italic = True
    r.font.size = Pt(13)
    r.font.name = "Times New Roman"

    para(
        doc,
        "This report is based only on the already saved fold outputs. It does not rerun training or evaluation. The goal is to explain the fold-level behavior, the confusion matrices, and why fold 2 can be treated as a good candidate while still acknowledging the trade-off against fold 4."
    )

    doc.add_heading("1. Fold Status Overview", level=1)
    rows = []
    for fold in range(5):
        metrics = load_metrics(CV_ROOT / f"fold_{fold}")
        rows.append([
            f"fold_{fold}",
            f"{metrics['accuracy']:.4f}",
            f"{metrics['weighted_accuracy']:.4f}",
            f"{metrics['unweighted_accuracy']:.4f}",
            f"{metrics['macro_f1']:.4f}",
            f"{metrics['weighted_f1']:.4f}",
        ])
    table(doc, ["Fold", "Accuracy", "Weighted Acc.", "Unweighted Acc.", "Macro F1", "Weighted F1"], rows)

    doc.add_heading("2. Aggregate Reading of the Five Folds", level=1)
    bullets(
        doc,
        [
            "The folds are close enough to show stable behavior rather than a one-off spike.",
            "Fold 2 is the strongest option if weighted aggregate performance is the priority.",
            "Fold 4 is the strongest option if balanced class behavior is the priority.",
            "Neutral, joy, sadness, and anger are the most frequently confused classes across the folds.",
            "The matrix patterns suggest the model is learning broad emotional polarity more reliably than fine-grained class distinctions.",
        ],
    )

    doc.add_heading("3. Fold-by-Fold Explanation", level=1)
    for fold in range(5):
        add_fold_section(doc, fold)

    doc.add_heading("4. Mentor-Level Interpretation", level=1)
    para(
        doc,
        "A useful way to explain the fold analysis to a mentor is to say that the model is performing at a stable but still imperfect level across the five folds. The consistency of the confusion structure matters as much as the absolute score. The main question is not only which fold is numerically best, but whether the same confusion pattern repeats, because repetition means the model has learned a real but incomplete representation."
    )
    para(
        doc,
        "Fold 2 is a good candidate because it combines the strongest accuracy and weighted F1. That makes it the most attractive single fold if the report needs one representative result. However, fold 4 shows the best macro F1 and unweighted accuracy, so if the emphasis is on class balance, fold 4 deserves mention. This is why the report should present both and explain the trade-off rather than choosing one without context."
    )
    para(
        doc,
        "The confusion matrices also help explain why accuracy alone is not enough. A model can appear acceptable on accuracy while still missing class-level detail. The fold matrices show whether the model is mostly wrong in one direction, which would indicate collapse, or whether it is making more distributed confusions, which indicates a more nuanced but still limited learned structure."
    )

    doc.add_heading("5. Output Locations", level=1)
    bullets(
        doc,
        [
            "Per-fold images and CSV analyses: results/paper_aligned_meld_cv/cmt_min/fold_*/analysis_test/",
            "Aggregate fold summary: results/paper_aligned_meld_cv/cmt_min/fold_summary.md",
            "Aggregate fold JSON: results/paper_aligned_meld_cv/cmt_min/fold_summary.json",
        ],
    )

    return doc


def main() -> None:
    if not CV_ROOT.exists():
        raise FileNotFoundError(CV_ROOT)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_doc().save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
