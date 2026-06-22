from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
RUN_ROOT = ROOT / "results" / "paper_aligned_crema_d" / "cmt_min"
OUTPUT = ROOT / "implementation_docments" / "LegalMemoCMT_CREMA_D_Analysis_Report.docx"


def style(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.95)
    sec.right_margin = Inches(0.95)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def load_metrics() -> dict[str, object]:
    return json.loads((RUN_ROOT / "metrics.json").read_text(encoding="utf-8"))


def load_csv(name: str) -> pd.DataFrame:
    if name == "confusion_matrix.csv":
        return pd.read_csv(RUN_ROOT / "analysis_test" / name, index_col=0)
    return pd.read_csv(RUN_ROOT / "analysis_test" / name)


def build_doc() -> Document:
    doc = Document()
    style(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LegalMemoCMT CREMA-D Analysis Report")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Analysis of the saved CREMA-D held-out output")
    r.italic = True
    r.font.size = Pt(13)
    r.font.name = "Times New Roman"

    metrics = load_metrics()
    cm = load_csv("confusion_matrix.csv")
    top = load_csv("top_confusions.csv").head(5)
    first20 = load_csv("predicted_vs_actual_first20.csv")
    cm_png = RUN_ROOT / "analysis_test" / "confusion_matrix.png"

    para(
        doc,
        "This report summarizes the currently available CREMA-D held-out analysis for the paper-aligned CMT + MIN setup. It is intentionally limited to the saved output already present in the workspace and does not claim a full 5-fold CREMA-D evaluation."
    )
    para(
        doc,
        "The purpose of the report is to make the confusion pattern easy to explain to a mentor: the model is not simply mediocre overall; it is visibly collapsing toward a narrow set of predicted classes, which the confusion matrix shows very clearly."
    )

    doc.add_heading("1. Metrics", level=1)
    table(doc, ["Metric", "Value"], [[k, f"{v:.6f}" if isinstance(v, float) else str(v)] for k, v in metrics.items()])

    doc.add_heading("2. What the numbers mean", level=1)
    bullets(
        doc,
        [
            "Accuracy is low enough to indicate weak separation on this run.",
            "Macro F1 is very low, which means some classes are effectively not being learned.",
            "Weighted F1 is also low, so this is not just a minority-class issue; overall prediction quality is poor.",
            "The confusion structure suggests the model is strongly biased toward predicting fear.",
        ],
    )

    doc.add_heading("3. Confusion Matrix", level=1)
    if cm_png.exists():
        doc.add_picture(str(cm_png), width=Inches(5.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    rows = []
    for idx, row in cm.iterrows():
        rows.append([idx] + [str(int(v)) for v in row.tolist()])
    table(doc, ["Actual \\ Pred"] + list(cm.columns), rows)

    doc.add_heading("4. Top Confusions", level=1)
    rows = []
    for _, row in top.iterrows():
        rows.append([f"{int(row['actual_label'])}:{row['actual_name']}", f"{int(row['predicted_label'])}:{row['predicted_name']}", str(int(row['count']))])
    table(doc, ["Actual", "Predicted", "Count"], rows)

    doc.add_heading("5. First-20 Examples", level=1)
    rows = []
    for _, row in first20.head(10).iterrows():
        rows.append([str(row["sample_id"]), str(row["actual_label"]), str(row["predicted_label"]), str(row["correct"])])
    table(doc, ["Sample", "Actual", "Predicted", "Correct"], rows)

    doc.add_heading("6. Error Analysis for Students", level=1)
    para(
        doc,
        "The CREMA-D matrix is best explained as a near-collapse into fear as the dominant prediction. Many actual anger, disgust, happy, sad, and neutral samples are mapped into fear, which means the model is detecting some generic emotional intensity but not preserving class identity."
    )
    para(
        doc,
        "For a mentor explanation, the key point is that the error is not random. Random error would spread across many wrong classes. Here the errors are concentrated, which means the model has learned a narrow decision rule that over-triggers one class. That usually indicates a mismatch between the dataset, the learned representation, and the current training setup."
    )
    para(
        doc,
        "This is useful diagnostically. It tells us that the current CREMA-D run should not be presented as a strong result. Instead, it should be presented as a worked example of how confusion analysis reveals a failure mode that aggregate metrics alone can hide."
    )
    para(
        doc,
        "The first-20 sample table is also important for teaching because it shows concrete examples of the same pattern. A student can point to one sample and say exactly what the model predicted, whether it was correct, and why that prediction fits the confusion matrix."
    )

    doc.add_heading("7. What can be claimed", level=1)
    bullets(
        doc,
        [
            "The held-out CREMA-D predictions were analyzed successfully.",
            "The confusion matrix and top-confusion structure were generated from the saved output.",
            "The report should not claim a full 5-fold CREMA-D benchmark because that is not what exists in the current outputs.",
        ],
    )

    doc.add_heading("8. Output Locations", level=1)
    bullets(
        doc,
        [
            "Saved predictions: results/paper_aligned_crema_d/cmt_min/predictions_test.csv",
            "Image confusion matrix: results/paper_aligned_crema_d/cmt_min/analysis_test/confusion_matrix.png",
            "Top confusions: results/paper_aligned_crema_d/cmt_min/analysis_test/top_confusions.csv",
        ],
    )

    return doc


def main() -> None:
    if not RUN_ROOT.exists():
        raise FileNotFoundError(RUN_ROOT)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_doc().save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
