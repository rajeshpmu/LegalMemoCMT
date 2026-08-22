from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOP = ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx"


def code(document: Document, value: str) -> None:
    p = document.add_paragraph(style="No Spacing")
    r = p.add_run(value)
    r.font.name = "Courier New"
    r.font.size = Pt(8)


document = Document(SOP)
document.add_section(WD_SECTION.NEW_PAGE)
document.add_heading("pyannote Authentication API Correction", level=1)
document.add_paragraph(
    "The isolated environment uses pyannote.audio 3.4.0. In this version, "
    "`Pipeline.from_pretrained` accepts the Hugging Face credential as `use_auth_token`, not `token`. "
    "The preflight initially failed with `TypeError: unexpected keyword argument 'token'`. Both the "
    "preflight checker and diarization worker now try the newer `token` keyword and fall back to "
    "`use_auth_token` when the installed pyannote version rejects it."
)
code(document, '''export HF_TOKEN="<your Hugging Face token>"
./.venv-diarization/bin/python phase2/check_clancy_diarization_prerequisites.py \\
  --model pyannote/speaker-diarization-3.1 \\
  --load-model''')
document.add_paragraph(
    "The local compatibility checks now pass with torch 2.2.2, torchaudio 2.2.2, "
    "AudioMetaData available, and pyannote.audio 3.4.0. The remaining model-access result depends on "
    "the supplied token and accepted Hugging Face model terms."
)
document.save(SOP)
print(f"Updated {SOP}")
