from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOP = ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx"


def code(document: Document, value: str) -> None:
    paragraph = document.add_paragraph(style="No Spacing")
    run = paragraph.add_run(value)
    run.font.name = "Courier New"
    run.font.size = Pt(8)


def table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = document.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, value in zip(t.rows[0].cells, headers):
        cell.text = value
    for row in rows:
        cells = t.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)


document = Document(SOP)
document.add_section(WD_SECTION.NEW_PAGE)
document.add_heading("Clancy Speaker-Role Resolution Stage", level=1)
document.add_paragraph(
    "Purpose. The Phase 1 emotion checkpoint predicts an emotion class; it does not know whether the "
    "current speaker is a witness, prosecutor, defence lawyer, judge, or another participant. This stage "
    "adds speaker and witness evidence before weak labels are used for a witness-focused corpus. It is "
    "deliberately three-layered: subtitle evidence, source-level speaker diarization, and manual role "
    "mapping. Automatic outputs remain UNKNOWN until a person verifies the cluster role."
)

document.add_heading("Fields Added", level=2)
table(document, ["Field", "Purpose", "Safe values"], [
    ["witness_in_segment", "Whether the testimony context contains or involves a witness.", "YES, NO, UNKNOWN"],
    ["witness_speaking_status", "Whether the witness is speaking rather than listening or being addressed.", "SPEAKING, LISTENING_OR_ADDRESSED, UNKNOWN"],
    ["speaker_role", "Role of the person producing the audio turn.", "Witness, Prosecutor, Defence, Judge, Other, UNKNOWN"],
    ["speaker_role_source", "How the role was assigned.", "manual_verified_diarization, subtitle_heuristic, unresolved"],
    ["speaker_role_confidence", "Confidence in the role assignment, not emotion accuracy.", "HIGH, MEDIUM, LOW"],
    ["speaker_cluster_id", "Anonymous voice cluster produced by diarization.", "For example SPEAKER_00 or UNKNOWN"],
    ["visual_target_role", "Role of the person shown by the camera.", "Witness, Prosecutor, Defence, Judge, Other, UNKNOWN"],
    ["visual_speaker_match", "Whether the visible target matches the speaking person.", "YES, NO, UNKNOWN"],
])
document.add_paragraph(
    "A row can have `witness_in_segment=YES` while `speaker_role=Prosecutor`; that means a witness is "
    "being questioned or discussed, not that the witness is speaking. For the primary witness-emotion "
    "pool, the preferred condition is `speaker_role=Witness` and `witness_speaking_status=SPEAKING`."
)

document.add_heading("Layer 1: Subtitle Evidence", level=2)
document.add_paragraph(
    "`phase2/build_clancy_subtitle_evidence.py` parses WebVTT timestamps and overlapping caption text. "
    "It preserves the relationship between a manifest row and its subtitle file, records caption counts, "
    "question-pattern evidence, witness-related terms, and probable testimony evidence. The parser does "
    "not promote a row to a named legal role merely because a caption contains the word witness."
)
code(document, '''PYTHON_BIN=$PWD/.venv/bin/python \\
bash phase2/run_build_clancy_subtitle_evidence.sh \\
  --input-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200.csv \\
  --output-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_subtitle_evidence.csv''')
document.add_paragraph(
    "The subtitle pilot produced 200 rows. All 200 remained `speaker_role=UNKNOWN` and "
    "`witness_speaking_status=UNKNOWN`; 153 rows contained a question-pattern signal. This is an "
    "expected conservative result because the downloaded Clancy VTT files contain `>>` caption-change "
    "markers and text, but not reliable names such as WITNESS or PROSECUTOR for each cue."
)

document.add_heading("Layer 2: Source-Level Speaker Diarization", level=2)
document.add_paragraph(
    "Diarization separates voices into anonymous clusters. It must run on the complete source audio for "
    "each YouTube recording, not independently on every short utterance clip. Running it per clip would "
    "reassign cluster numbers repeatedly and make the clusters impossible to compare across a hearing. "
    "The script uses the Hugging Face pyannote model `pyannote/speaker-diarization-3.1`, then overlaps each "
    "manifest turn with the source-level diarization segments. It assigns the cluster with the greatest "
    "time overlap. Diarization identifies different voices, but it does not know their legal roles."
)
document.add_paragraph(
    "Before this layer, install the optional dependency in the environment used for inference and accept "
    "the model terms on Hugging Face. A Hugging Face access token is required. Do not place the token in a "
    "CSV, script, DOCX, or committed repository file."
)
code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python
"$PYTHON_BIN" -m pip install "pyannote.audio>=3.3,<4"
export HF_TOKEN="<your Hugging Face token>"''')
document.add_paragraph("Run one source first:", style="List Bullet")
code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python HF_TOKEN="$HF_TOKEN" \\
bash phase2/run_clancy_diarization.sh \\
  --input-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_subtitle_evidence.csv \\
  --output-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_diarized.csv \\
  --segments-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_diarization_segments_200.csv \\
  --max-sources 1''')
document.add_paragraph("After the one-source check, diarize all source recordings represented by the input:", style="List Bullet")
code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python HF_TOKEN="$HF_TOKEN" \\
bash phase2/run_clancy_diarization.sh \\
  --input-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_subtitle_evidence.csv \\
  --output-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_diarized.csv \\
  --segments-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_diarization_segments_200.csv''')

document.add_heading("Layer 3: Manual Cluster-Role Mapping", level=2)
document.add_paragraph(
    "`phase2/create_clancy_speaker_role_review.py` creates one row per source recording and anonymous "
    "speaker cluster. The review sheet contains example utterance IDs, audio paths, and text so the "
    "student can listen to representative samples and assign the legal role. The reviewer must use "
    "`UNKNOWN` when the cluster cannot be identified. A diarization label such as SPEAKER_00 is not itself "
    "a witness label."
)
code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python \\
"$PYTHON_BIN" phase2/create_clancy_speaker_role_review.py \\
  --input-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_diarized.csv \\
  --output-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_speaker_role_review_200.csv \\
  --examples-per-cluster 3''')
document.add_paragraph(
    "Open the review CSV and complete `role_label`, `role_confidence`, `witness_in_segment`, "
    "`witness_speaking_status`, `visual_target_role`, `visual_speaker_match`, and `review_notes`. "
    "The permitted role values are Witness, Prosecutor, Defence, Judge, Other, and UNKNOWN. The "
    "reviewer should listen to the sample audio and inspect the corresponding video before assigning "
    "visual speaker match. This manual edit is an evidence-review step, not a model prediction."
)
document.add_heading("Apply the Verified Role Map", level=2)
code(document, '''PYTHON_BIN=/opt/anaconda3/bin/python \\
"$PYTHON_BIN" phase2/apply_clancy_speaker_role_map.py \\
  --input-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_diarized.csv \\
  --role-map-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_speaker_role_review_200.csv \\
  --output-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_role_verified.csv''')

document.add_heading("Role-Aware Weak-Label Routing", level=2)
table(document, ["Verified condition", "Phase 1 route", "Use"], [
    ["Witness + SPEAKING + visual match YES", "text,audio,video", "Primary witness multimodal weak-label candidate"],
    ["Prosecutor, Defence, Judge, or Other speaking", "text,audio", "Non-witness interactional or auxiliary corpus"],
    ["Witness in context but lawyer speaking", "text,audio", "Do not assign witness facial emotion"],
    ["Speaker or visual target UNKNOWN", "text,audio or ABSTAIN", "Review-required row; do not promote automatically"],
])
document.add_paragraph(
    "Use the video-enabled face-crop checkpoint only when the video feature file has shape `(16, 768)` "
    "and the visual target is plausibly the speaking person. Use the text-audio checkpoint when the "
    "camera target is another participant or visual alignment is unresolved. This prevents the model from "
    "using a lawyer's face to label a witness's speech."
)

document.add_heading("Validation and Guardrails", level=2)
for item in [
    "Every cluster role must retain source recording, cluster ID, sample audio, and sample text traceability.",
    "Do not convert the words witness, doctor, or testimony in a caption directly into speaker_role=Witness.",
    "Do not infer a speaker's legal role from face visibility alone.",
    "Do not use diarization cluster IDs as personal identities without manual verification.",
    "Rows with speaker_role=UNKNOWN remain review or abstain rows for witness-emotion training.",
    "Basic emotion remains separate from courtroom affect; no deception, credibility, or truthfulness label is generated.",
    "The primary witness pool should require speaker_role=Witness and witness_speaking_status=SPEAKING.",
]:
    document.add_paragraph(item, style="List Bullet")

document.add_heading("Current Implementation Status", level=2)
document.add_paragraph(
    "Completed: subtitle parser, diarization worker, cluster review-sheet builder, role-map application "
    "script, wrappers, and the 200-row subtitle-evidence pilot. Pending: install pyannote.audio, configure "
    "a Hugging Face token, run source-level diarization, manually complete the cluster-role review sheet, "
    "apply the map, and only then route Phase 1 weak-label inference by role and visual match."
)

document.save(SOP)
print(f"Updated {SOP}")
