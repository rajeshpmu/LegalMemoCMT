#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "Phase1_Final_Demo_List_Student_Guide.docx"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_code(doc: Document, text: str) -> None:
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(10.2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Phase 1 Final Demo List")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Student-level safe demo list with verified clips and commands for the baseline and gated + aux raw-mp4 demos")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12.5)

    add_para(
        doc,
        "Purpose: this short guide lists five verified demo clips and the exact commands to run them against the paper-aligned baseline and the gated + auxiliary-loss checkpoint. It is written for a student who needs a safe final demo list, not a speculative one.",
    )

    doc.add_heading("1. Final Demo Clips", level=1)
    add_table(
        doc,
        ["Order", "sample_id", "Raw mp4", "Why it is useful"],
        [
            ["1", "test_dia279_utt9", "data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia279_utt9.mp4", "Correct neutral in gated + aux; baseline is a near-boundary error."],
            ["2", "test_dia4_utt6", "data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia4_utt6.mp4", "Verified correct non-neutral disgust example in both checkpoints."],
            ["3", "test_dia278_utt5", "data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia278_utt5.mp4", "Confident wrong example: both checkpoints predict joy for surprise."],
            ["4", "test_dia153_utt5", "data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia153_utt5.mp4", "Hard neutral failure: baseline predicts fear, gated + aux predicts disgust."],
            ["5", "test_dia1_utt1", "data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia1_utt1.mp4", "Backup correct non-neutral joy example."],
        ],
    )

    doc.add_heading("2. Safe Demo Categories", level=1)
    add_table(
        doc,
        ["Category", "Clip", "Verified behavior", "How to present it"],
        [
            ["Correct neutral", "test_dia279_utt9", "Gated + aux correct; baseline wrong", "Use it to show the improvement story."],
            ["Correct non-neutral", "test_dia4_utt6", "Disgust -> disgust in both checkpoints", "Use it to show the model can classify a non-neutral emotion correctly."],
            ["Near-miss", "test_dia279_utt9", "Baseline is close to the boundary between anger and neutral", "Use it to explain uncertainty and top-2 closeness."],
            ["Confident wrong", "test_dia278_utt5", "Surprise -> joy with high confidence", "Use it to explain modal bias and label ambiguity."],
            ["Hard neutral failure", "test_dia153_utt5", "Neutral -> fear/disgust", "Use it to show that neutral can still be difficult."],
        ],
    )

    doc.add_heading("3. Commands to Run", level=1)
    add_table(
        doc,
        ["sample_id", "Baseline command", "Gated + aux command", "Expected demo role"],
        [
            [
                "test_dia279_utt9",
                "MODALITIES=text,audio bash scripts/run_demo_paper_aligned_raw_mp4.sh test_dia279_utt9 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia279_utt9.mp4",
                "DEVICE=cuda CHECKPOINT=results/facial_cues/meld_vit_facecrop_gated_video_aux/fold_4/best_model.pt bash scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh test_dia279_utt9 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia279_utt9.mp4",
                "Baseline near-miss, gated + aux correct neutral.",
            ],
            [
                "test_dia4_utt6",
                "MODALITIES=text,audio bash scripts/run_demo_paper_aligned_raw_mp4.sh test_dia4_utt6 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia4_utt6.mp4",
                "DEVICE=cuda CHECKPOINT=results/facial_cues/meld_vit_facecrop_gated_video_aux/fold_4/best_model.pt bash scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh test_dia4_utt6 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia4_utt6.mp4",
                "Correct non-neutral disgust in both checkpoints.",
            ],
            [
                "test_dia278_utt5",
                "MODALITIES=text,audio bash scripts/run_demo_paper_aligned_raw_mp4.sh test_dia278_utt5 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia278_utt5.mp4",
                "DEVICE=cuda CHECKPOINT=results/facial_cues/meld_vit_facecrop_gated_video_aux/fold_4/best_model.pt bash scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh test_dia278_utt5 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia278_utt5.mp4",
                "Confident wrong surprise -> joy.",
            ],
            [
                "test_dia153_utt5",
                "MODALITIES=text,audio bash scripts/run_demo_paper_aligned_raw_mp4.sh test_dia153_utt5 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia153_utt5.mp4",
                "DEVICE=cuda CHECKPOINT=results/facial_cues/meld_vit_facecrop_gated_video_aux/fold_4/best_model.pt bash scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh test_dia153_utt5 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia153_utt5.mp4",
                "Hard neutral failure.",
            ],
        ],
    )

    doc.add_heading("4. What to Say", level=1)
    add_bullets = [
        "The baseline and gated + aux checkpoints see the same clip, so the comparison is fair.",
        "The baseline can be uncertain or wrong on neutral clips.",
        "The gated + aux model can correct some boundary cases, but not all of them.",
        "A high confidence score means the model is certain, not necessarily correct.",
        "The final demo should include both success and failure cases so the reviewer sees the real behavior."
    ]
    for item in add_bullets:
        doc.add_paragraph(item, style="List Bullet")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
