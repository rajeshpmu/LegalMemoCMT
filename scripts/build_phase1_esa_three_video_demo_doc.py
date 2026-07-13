#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "Phase1_ESA_Three_Video_Demo_Guide.docx"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Phase 1 ESA Review: Three-Video Demo Guide")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Student-level guide for the ESA review demo using three raw MELD test clips and current outputs")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12.5)

    add_para(
        doc,
        "Purpose: this guide gives you a compact but technically defensible demo set for the ESA review. It explains which three clips to show, how to run the baseline and the gated + auxiliary-loss model, what the outputs mean, and how to explain correct and incorrect predictions without overstating the results.",
    )

    doc.add_heading("1. Why These Three Clips Were Chosen", level=1)
    add_bullets(
        doc,
        [
            "One correct neutral clip: to show the model can still classify the common neutral class correctly.",
            "One correct non-neutral clip: to show the model is not only guessing neutral.",
            "One confidently wrong clip: to show the failure pattern clearly and honestly.",
            "Together, the three clips let you show success, limitation, and confidence behavior in a short live demo.",
        ],
    )
    add_table(
        doc,
        ["Category", "sample_id", "Truth", "Why it is useful"],
        [
            ["Correct neutral", "test_dia279_utt9", "neutral", "Gated + aux corrects a baseline near-miss."],
            ["Correct non-neutral", "test_dia4_utt6", "disgust", "Both models are correct, so it is a safe non-neutral success case."],
            ["Confident wrong", "test_dia278_utt5", "surprise", "Both models predict joy; gated + aux is especially confident."],
        ],
    )

    doc.add_heading("2. Exact Commands to Run", level=1)
    add_para(
        doc,
        "Use the same raw .mp4 clip with the baseline command and the gated + aux command so the reviewer can compare the two model stages on identical input.",
    )
    add_table(
        doc,
        ["sample_id", "Baseline command", "Gated + aux command"],
        [
            [
                "test_dia4_utt6",
                "MODALITIES=text,audio bash scripts/run_demo_paper_aligned_raw_mp4.sh test_dia4_utt6 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia4_utt6.mp4",
                "DEVICE=cuda CHECKPOINT=results/facial_cues/meld_vit_facecrop_gated_video_aux/fold_4/best_model.pt bash scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh test_dia4_utt6 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia4_utt6.mp4",
            ],
            [
                "test_dia279_utt9",
                "MODALITIES=text,audio bash scripts/run_demo_paper_aligned_raw_mp4.sh test_dia279_utt9 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia279_utt9.mp4",
                "DEVICE=cuda CHECKPOINT=results/facial_cues/meld_vit_facecrop_gated_video_aux/fold_4/best_model.pt bash scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh test_dia279_utt9 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia279_utt9.mp4",
            ],
            [
                "test_dia278_utt5",
                "MODALITIES=text,audio bash scripts/run_demo_paper_aligned_raw_mp4.sh test_dia278_utt5 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia278_utt5.mp4",
                "DEVICE=cuda CHECKPOINT=results/facial_cues/meld_vit_facecrop_gated_video_aux/fold_4/best_model.pt bash scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh test_dia278_utt5 data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia278_utt5.mp4",
            ],
        ],
    )
    add_para(
        doc,
        "Student explanation: the baseline demo uses the stable paper-aligned checkpoint and is run in text+audio mode, because that checkpoint is the conversational reference. The gated + aux demo uses the later facial-cue checkpoint and the same raw clip, so the difference comes from the model stage rather than from a different input video.",
    )

    doc.add_heading("3. Current Outputs to Present", level=1)
    add_table(
        doc,
        ["sample_id", "Model stage", "Ground truth", "Prediction", "Confidence", "Correct?", "What to say"],
        [
            ["test_dia4_utt6", "Baseline", "disgust", "disgust", "0.5126", "Yes", "A non-neutral case that both models get right."],
            ["test_dia4_utt6", "Gated + aux", "disgust", "disgust", "0.9670", "Yes", "The later model is much more confident on this clip."],
            ["test_dia279_utt9", "Baseline", "neutral", "anger", "0.4422", "No", "A near-boundary neutral error; anger is close to neutral."],
            ["test_dia279_utt9", "Gated + aux", "neutral", "neutral", "0.9428", "Yes", "The gated + aux model corrects the neutral case strongly."],
            ["test_dia278_utt5", "Baseline", "surprise", "joy", "0.4809", "No", "A near-miss with surprise and joy; the prediction is wrong but not random."],
            ["test_dia278_utt5", "Gated + aux", "surprise", "joy", "0.9435", "No", "A confident wrong prediction; the model is very certain but still incorrect."],
        ],
    )

    doc.add_heading("4. How to Explain the Outputs", level=1)
    add_bullets(
        doc,
        [
            "Ground truth is the real MELD label from the manifest.",
            "Predicted label is the top class after softmax over emotion logits.",
            "Confidence is the top-1 softmax probability, not the accuracy of the model.",
            "Top-3 classes help show whether the model is unsure or whether it has strongly committed to the wrong class.",
        ],
    )
    add_para(
        doc,
        "Student explanation: the baseline and gated + aux models should be compared on the same sample_id. That makes the demo fair. The neutral example shows an actual correction, the disgust example shows a stable success, and the surprise example shows that the model can still be very confident and still wrong.",
    )
    add_numbered(
        doc,
        [
            "Show the clip name and the ground-truth label first.",
            "Run the baseline command and read the printed prediction.",
            "Run the gated + aux command on the same clip.",
            "Compare the confidence and the top-3 values.",
            "Explain whether the clip is a success, near-miss, or confident wrong prediction.",
        ],
    )

    doc.add_heading("5. Technical Interpretation for the ESA Review", level=1)
    add_table(
        doc,
        ["Clip", "Technical point", "Reviewer-ready explanation"],
        [
            ["test_dia4_utt6", "Both models are correct.", "This is a safe success example for the demo because it shows the pipeline can recognize a non-neutral emotion."],
            ["test_dia279_utt9", "Baseline near-miss, gated + aux correct.", "This is the best improvement example because the later model fixes the baseline’s neutral confusion."],
            ["test_dia278_utt5", "Both models are wrong, gated + aux is confident.", "This is the honest failure case: the model’s certainty is high, but the label is still wrong."],
        ],
    )
    add_para(
        doc,
        "Do not overstate the gated + aux model as always better. The correct review language is that the later model improves some cases, especially neutral boundary cases, but it still has failure modes and can be confidently wrong on ambiguous clips.",
    )
    add_para(
        doc,
        "If the reviewer asks why that happens, explain that the result depends on how the transcript, sampled frames, and learned fusion weights interact. In other words, the model is making a multimodal decision, not looking at the label directly.",
    )

    doc.add_heading("6. Suggested Demo Order", level=1)
    add_numbered(
        doc,
        [
            "Start with test_dia4_utt6 so the reviewer sees one clean success case.",
            "Move to test_dia279_utt9 to show how the gated + aux model fixes a neutral baseline error.",
            "End with test_dia278_utt5 to show a confident wrong prediction and explain the model’s remaining limitation.",
        ],
    )
    add_para(
        doc,
        "This order is useful because it starts with confidence, moves to improvement, and ends with a limitation. That gives a balanced and technically honest presentation.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
