#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "implementation_docments" / "Phase1_Raw_MP4_Demo_Differences_Student_Guide.docx"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_code(doc: Document, text: str) -> None:
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(10.2)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Phase 1 Raw MP4 Demo Differences")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Student-level explanation of why the same clip can be wrong in the baseline and correct in the gated + aux model")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12.5)

    add_para(
        doc,
        "Purpose: this document explains the difference between the paper-aligned baseline raw-mp4 demo and the later gated + auxiliary-loss raw-mp4 demo. It is written so a student can explain why the same clip may produce different predictions under different checkpoints, and where in the code that difference comes from.",
    )

    doc.add_heading("1. What Is Being Compared", level=1)
    add_bullets(
        doc,
        [
            "Same raw mp4 clip.",
            "Same MELD sample_id.",
            "Same raw-video preprocessing path.",
            "Different trained checkpoint.",
            "Different fusion behavior in the model.",
            "Potentially different final prediction and confidence.",
        ],
    )
    add_para(
        doc,
        "The important scientific idea is that we are not changing the input clip. We are changing the model checkpoint. That means any change in prediction can be discussed as a checkpoint-level behavior change rather than a data-level change.",
    )

    doc.add_heading("2. The Two Demo Paths", level=1)
    add_table(
        doc,
        ["Path", "Wrapper", "Checkpoint style", "What it is good for"],
        [
            ["Baseline raw-mp4 demo", "scripts/run_demo_paper_aligned_raw_mp4.sh", "Paper-aligned baseline checkpoint", "Showing the stable reference behavior."],
            ["Gated + aux raw-mp4 demo", "scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh", "Face-crop ViT + gated fusion + auxiliary loss", "Showing the later facial-cue refinement."],
        ],
    )
    add_para(
        doc,
        "The raw mp4 extraction script is the same idea in both cases: it samples frames, extracts features, caches the `.npy` file, and passes that representation into the loaded checkpoint. The checkpoint decides how those features are interpreted.",
    )

    doc.add_heading("3. The Exact Commands", level=1)
    add_code(
        doc,
        r"""# Baseline paper-aligned demo
DEVICE=cuda \
bash scripts/run_demo_paper_aligned_raw_mp4.sh \
  test_dia279_utt9 \
  data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia279_utt9.mp4

# Gated + aux demo on the same clip
DEVICE=cuda \
CHECKPOINT=results/facial_cues/meld_vit_facecrop_gated_video_aux/fold_4/best_model.pt \
bash scripts/run_demo_meld_vit_facecrop_gated_video_aux_raw_mp4.sh \
  test_dia279_utt9 \
  data/MELD/raw/MELD.Raw/test/output_repeated_splits_test/dia279_utt9.mp4""",
    )
    add_para(
        doc,
        "Student note: the exact same clip is used in both runs so the reviewer can see a true apples-to-apples comparison.",
    )

    doc.add_heading("4. Verified Same-Clip Result", level=1)
    add_table(
        doc,
        ["Checkpoint", "Ground truth", "Prediction", "Confidence", "Top-3 meaning"],
        [
            ["Paper-aligned baseline", "neutral", "anger", "0.4422", "The model nearly split between anger and neutral, then tipped slightly toward anger."],
            ["Gated + aux", "neutral", "neutral", "0.9428", "The model strongly concentrated on neutral and moved away from the wrong classes."],
        ],
    )
    add_para(
        doc,
        "This is the key comparison result you should use in discussion: the baseline is not random, but it is uncertain and slightly biased toward the wrong class. The gated + aux checkpoint makes the same sample much more stable and correct.",
    )

    doc.add_heading("4.1 Same Clip, Still Wrong in Both Checkpoints", level=2)
    add_para(
        doc,
        "A different clip, `test_dia278_utt5`, shows a different kind of lesson. Here, both checkpoints are wrong, but the confidence changes a lot. That means the later model does not automatically fix every hard sample. It can also become more confident in the wrong answer if the evidence is ambiguous or biased toward the wrong class.",
    )
    add_table(
        doc,
        ["Checkpoint", "Ground truth", "Prediction", "Confidence", "What it tells you"],
        [
            ["Paper-aligned baseline", "surprise", "joy", "0.4809", "The model leans toward joy, but the score gap is still modest."],
            ["Gated + aux", "surprise", "joy", "0.9435", "The model becomes much more certain about the same wrong answer."],
        ],
    )
    add_bullets(
        doc,
        [
            "The transcript sounds positive or playful, so the text branch can bias the model toward joy.",
            "The sampled face frames may not capture the peak surprise expression strongly enough.",
            "The gated + aux model may strengthen the dominant signal instead of correcting it if the input evidence is already misleading.",
            "This is why you should not say that a later checkpoint always improves every clip.",
        ],
    )
    add_para(
        doc,
        "Student explanation: this second example is important because it shows that model improvement is selective. The gated + aux checkpoint improves some clips, like the neutral example above, but it can still fail on harder or more ambiguous clips. In fact, it can become more confident in the wrong class if the transcript and sampled frames are already leaning in that direction.",
    )
    add_para(
        doc,
        "The technical lesson is that the same model family can show two behaviors: correction on some boundary cases and stronger overcommitment on some ambiguous cases. That is why you should present both a success case and a failure case in the review.",
    )

    doc.add_heading("4.2 Same Clip, Hard Neutral Failure", level=2)
    add_para(
        doc,
        "The clip `test_dia153_utt5` gives a third behavior pattern. Here the ground truth is neutral, but both checkpoints fail and choose a non-neutral class. The baseline predicts fear, while the gated + aux checkpoint predicts disgust. This is important because it shows that the improved model does not guarantee a better outcome on every neutral clip.",
    )
    add_table(
        doc,
        ["Checkpoint", "Ground truth", "Prediction", "Confidence", "What it tells you"],
        [
            ["Paper-aligned baseline", "neutral", "fear", "0.6990", "The model is fairly sure the clip is not neutral and leans toward fear."],
            ["Gated + aux", "neutral", "disgust", "0.7620", "The later model is also wrong and becomes even more certain, but in a different non-neutral direction."],
        ],
    )
    add_bullets(
        doc,
        [
            "Neutral is often a subtle class, so it can be easy for the model to drift toward a stronger emotion label.",
            "If the clip contains small tension or a slight facial reaction, the model may not keep it in the neutral bucket.",
            "The text or audio may also sound more emotionally loaded than the neutral label suggests.",
            "This example shows that the gated + aux model improves some clips, but it can still overcommit on hard neutral cases.",
        ],
    )
    add_para(
        doc,
        "Student explanation: this case is useful because it stops you from overselling the improvement. The gated + aux checkpoint is better on the verified neutral example `test_dia279_utt9`, but on `test_dia153_utt5` it still fails and becomes more confident in a non-neutral label. So the improvement is real, but it is selective rather than universal.",
    )
    add_para(
        doc,
        "A good viva sentence is: 'The improved checkpoint can correct some neutral examples, but hard neutral clips can still be pushed into fear or disgust if the visual or linguistic evidence is slightly non-neutral. That is why I treat Phase 1 as improving, but not fully solved.'",
    )

    doc.add_heading("5. How to Read the Difference", level=1)
    add_numbered(
        doc,
        [
            "First, notice that the input clip never changed.",
            "Second, notice that the baseline decision was close to the boundary between anger and neutral.",
            "Third, notice that the gated + aux decision became much more decisive.",
            "Fourth, interpret the confidence as model certainty, not truth.",
            "Fifth, connect the change to gated fusion and auxiliary supervision, not to the clip itself.",
        ],
    )
    add_bullets(
        doc,
        [
            "The baseline likely relied on a weaker combination of modalities for this clip.",
            "The gated model can learn when to trust the video branch more carefully.",
            "The auxiliary loss makes the video branch learn emotion information directly, not just passively through fusion.",
            "That can improve the final fused representation, especially for clips where neutral cues are subtle but present.",
        ],
    )

    doc.add_heading("6. Student-Level Technical Explanation", level=1)
    add_para(
        doc,
        "In the model code, the final decision is not made by one branch alone. The text, audio, and video branches produce embeddings; then the fusion module combines them; then the classifier outputs logits. If a checkpoint was trained with a different fusion strategy, the same input features can produce a different label.",
    )
    add_table(
        doc,
        ["Code area", "What it controls", "Why it matters here"],
        [
            ["src/data/preprocessing.py", "How frames and audio are sampled.", "If sampling misses the peak moment, the feature can be weaker."],
            ["scripts/predict_phase1_raw_mp4_demo.py", "How raw mp4 becomes cached `.npy` features.", "This is the common input path for both demos."],
            ["src/models/model.py", "How embeddings are fused and classified.", "This is where the checkpoint family differs."],
            ["src/train/train.py", "How the model was trained and which loss was used.", "This affects whether the model learns to rely on video or ignore it."],
        ],
    )
    add_para(
        doc,
        "A student explanation should sound like this: 'The input is fixed, but the baseline checkpoint and the gated + aux checkpoint learn different decision boundaries. The gated + aux model has a stronger incentive to use the video branch and can therefore correct a neutral clip that the baseline pushed toward anger.'",
    )

    doc.add_heading("7. Why the Baseline Can Fail on the Same Clip", level=1)
    add_bullets(
        doc,
        [
            "The baseline may give too much weight to a misleading modality.",
            "The sampled frames may miss the strongest neutral cue.",
            "The model may see a short emotional transition and overreact to it.",
            "The learned boundary between neutral and anger may still be too sharp or poorly calibrated.",
        ],
    )
    add_para(
        doc,
        "This is a good example of why confidence alone should not be treated as correctness. A model can be moderately confident and still be wrong if the learned boundary is slightly off.",
    )

    doc.add_heading("8. Why the Gated + Aux Model Helps", level=1)
    add_bullets(
        doc,
        [
            "Gated fusion learns when the model should trust video more or less.",
            "The auxiliary loss forces the video branch to learn emotion-discriminative information directly.",
            "Together, these changes can stabilize the final fused representation.",
            "That is why the same clip can move from an uncertain anger prediction to a confident neutral prediction.",
        ],
    )
    add_para(
        doc,
        "Student takeaway: the gated + aux model does not magically solve every error, but it gives the video branch more influence and more supervision. That is enough to flip some near-boundary clips in the right direction.",
    )

    doc.add_heading("9. What This Means for Phase 1", level=1)
    add_para(
        doc,
        "This comparison supports the claim that Phase 1 is evolving from a stable baseline into a more video-aware system. The baseline remains the reference point, but the gated + aux model is the stronger bridge to later facial-cue work because it shows that the visual branch can change the outcome on the same clip.",
    )
    add_para(
        doc,
        "For a viva, the key sentence is: 'The same input video produces different predictions because the checkpoints have learned different fusion strategies and different levels of reliance on the video branch.'",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
