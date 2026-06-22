from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
BASE_DOC_PATH = ROOT / "implementation_docments" / "First_Review_LegalMemoCMT.docx"
DOC_PATH = ROOT / "implementation_docments" / "Second_Review_LegalMemoCMT.docx"
MELD_SUMMARY = ROOT / "results" / "paper_aligned_meld_cv" / "cmt_min" / "summary.json"
MELD_FOLD2 = ROOT / "results" / "paper_aligned_meld_cv" / "cmt_min" / "fold_2" / "metrics.json"
CREMA_METRICS = ROOT / "results" / "paper_aligned_crema_d" / "cmt_min" / "metrics.json"
FOLD2_FOCAL = ROOT / "results" / "improvement" / "class_balanced_focal" / "meld_selected" / "cmt_min" / "fold_2" / "metrics.json"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def add_section_title(doc: Document, title: str, level: int = 1) -> None:
    doc.add_heading(title, level=level)


def build_doc() -> Document:
    doc = Document(BASE_DOC_PATH)
    configure(doc)

    meld_summary = read_json(MELD_SUMMARY)["metrics"]
    meld_fold2 = read_json(MELD_FOLD2)
    crema = read_json(CREMA_METRICS)
    fold2_focal = read_json(FOLD2_FOCAL)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    add_section_title(doc, "12. Post-First-Review Technical Progress", level=1)
    add_para(
        doc,
        "This section is a continuation of the earlier review document. It summarizes the technical progress made after the first review, using the MELD CV report and the guidance comparison report as the main evidence base. The aim is to show what has changed in the codebase, what has been validated experimentally, and what still remains before the project can be described as fully adapted or fully reproduced.",
    )
    add_bullets(
        doc,
        [
            "The pretrained/paper-aligned branch is now operational and uses BERT for text and HuBERT for audio.",
            "The fusion block has been upgraded to a bidirectional cross-attention CMT path for the paper-aligned text+audio case.",
            "MELD has been run in a 5-fold CV workflow with MIN pooling as the main paper-aligned conversational result.",
            "CREMA-D has been added as the speech-emotion benchmark in the new benchmark split.",
            "The project now has new benchmark scripts, analysis wrappers, and student-facing documents that explain the result chain end to end.",
        ],
    )

    add_section_title(doc, "13. What the New Benchmark Scripts Do", level=1)
    add_para(
        doc,
        "The second review should be read together with two existing analysis documents: LegalMemoCMT_MELD_CV_5Fold_Analysis_Report.docx and LegalMemoCMT_CREMA_D_Analysis_Report.docx. Those reports are the result basis for the discussion here. The new benchmark script layer does not change the model implementation. It changes how the same model is organized for reporting. One path is the speech-emotion track, where CREMA-D is treated as the primary benchmark and is evaluated with speaker-independent cross-validation. The other path is the conversational track, where MELD remains the dialogue-level benchmark and is evaluated with the paper-aligned 5-fold CV workflow.",
    )
    add_table(
        doc,
        ["Script", "What it does", "Expected output"],
        [
            ["scripts/run_primary_speech_emotion_crema_d_cv.sh", "Builds speaker-independent CREMA-D folds, trains 5 folds, and aggregates W-Acc / UW-Acc-style results.", "crema_d_cv summary JSON/MD, fold metrics, fold predictions"],
            ["scripts/run_primary_speech_emotion_crema_d_cv_analysis.sh", "Exports and analyzes one CREMA-D fold's predictions.", "confusion matrix, top confusions, first-20 table, report.txt"],
            ["scripts/run_primary_conversational_meld_cv.sh", "Builds MELD raw manifests, creates 5 dialogue folds, trains and evaluates each fold.", "fold metrics, fold predictions, summary JSON/MD"],
            ["scripts/run_primary_conversational_meld_analysis.sh", "Analyzes one representative MELD fold.", "confusion matrix, top confusions, first-20 table, report.txt"],
            ["scripts/run_primary_benchmark_suite.sh", "Runs the CREMA-D track and the MELD track in sequence.", "Both benchmark result trees"],
        ],
    )
    add_para(
        doc,
        "The important student lesson is that the scripts are wrappers around the same model family, but they enforce a different narrative: CREMA-D is now the main speech-emotion benchmark, while MELD is the main conversational benchmark. That separation makes the project easier to discuss because the benchmark choice matches the task type instead of mixing several benchmark stories together.",
    )

    add_section_title(doc, "14. Metric Meaning in the Current Implementation", level=1)
    add_para(
        doc,
        "The project now uses two metric styles depending on the benchmark, and the discussion should be interpreted through the completed analysis reports. For MELD, the main numbers are accuracy, weighted F1, macro F1, and confusion-based error analysis as written in the MELD 5-fold CV analysis report. For CREMA-D, the summary is framed as W-Acc and UW-Acc so that the report follows the base-paper speech-emotion style more closely, while the CREMA-D analysis report remains the source for the current held-out failure pattern.",
    )
    add_bullets(
        doc,
        [
            "W-Acc is the sample-accuracy style summary used for the speech-emotion CV report.",
            "UW-Acc is the mean per-class accuracy over the classes present in the split.",
            "Accuracy is useful but can hide imbalance, so it should never be read alone.",
            "Macro F1 is the clearest indicator of whether minority classes are being learned.",
            "Weighted F1 reflects how the model behaves on the actual class distribution.",
            "Confusion matrices are needed to understand which emotion classes are being merged by the classifier.",
        ],
    )
    add_para(
        doc,
        "For a student, the easiest way to remember the difference is: W-Acc tells you how much overall sample-level correctness you have, while UW-Acc tells you whether the model is behaving fairly across classes. That is why the paper-style speech-emotion track needs both numbers, not just one. The CREMA-D analysis report uses the current held-out run as a diagnostic example, while the MELD 5-fold report uses fold-level stability and confusion structure to explain the conversational result.",
    )

    add_section_title(doc, "15. MELD Paper-Aligned Result After the First Review", level=1)
    add_para(
        doc,
        f"The MELD 5-fold CV summary shows mean accuracy {meld_summary['accuracy']['mean']:.4f}, mean weighted F1 {meld_summary['weighted_f1']['mean']:.4f}, and mean macro F1 {meld_summary['macro_f1']['mean']:.4f}. The best fold, Fold 2, reached accuracy {meld_fold2['accuracy']:.4f}, weighted F1 {meld_fold2['weighted_f1']:.4f}, and macro F1 {meld_fold2['macro_f1']:.4f}. Compared with the base paper’s MELD CMT + MIN test accuracy of 64.18%, the current result is close enough to be considered in the right direction, but not close enough to be described as a perfect reproduction.",
    )
    add_para(
        doc,
        "The Fold 2 error analysis shows a clear neutral-collapse pattern. Neutral, Joy, Anger, and Surprise are mixed repeatedly, while minority emotions remain difficult. That means the architecture is useful and working, but the class balance problem is still visible. The result should therefore be described as good and directionally correct, not as finished.",
    )
    add_bullets(
        doc,
        [
            "The MELD result is much closer to the base paper than a random or broken model would be.",
            "The main remaining weakness is minority-class separation, not the absence of a working pipeline.",
            "Fold-level variation is moderate, which suggests the model is usable but still not fully stable.",
        ],
    )

    add_section_title(doc, "15.1 Fold Status Overview", level=2)
    add_para(
        doc,
        "The fold status overview is the cleanest way to explain the MELD 5-fold report to a student or mentor because it makes the trade-off between weighted aggregate performance and balanced class behavior visible in one place. Fold 2 is the strongest fold if the priority is accuracy and weighted F1, while Fold 4 is the strongest fold if the priority is balanced per-class behavior.",
    )
    add_table(
        doc,
        ["Fold", "Acc / W-Acc", "UW-Acc", "Macro F1", "Weighted F1", "Reading"],
        [
            ["0", "0.6257", "0.4322", "0.4334", "0.6184", "Stable, neutral-heavy"],
            ["1", "0.6238", "0.4425", "0.4395", "0.6221", "Stable, sadness + neutral"],
            ["2", "0.6375", "0.4369", "0.4430", "0.6254", "Best weighted aggregate"],
            ["3", "0.6165", "0.4295", "0.4209", "0.6122", "Lowest overall fold"],
            ["4", "0.6199", "0.4672", "0.4606", "0.6194", "Best balanced-class fold"],
        ],
    )
    add_para(
        doc,
        "This table is important for interpretation. Fold 2 has the highest accuracy and weighted F1, so it is the best candidate when the aim is to compare against the paper’s MELD CMT + MIN result. Fold 4 has the best unweighted accuracy and macro F1, which means it is the strongest fold when the goal is to check whether minority classes are being handled more fairly. The existence of this trade-off is normal in imbalanced emotion data: one fold can look better under overall metrics while another looks better under class-balanced metrics.",
    )
    add_para(
        doc,
        "For student-level explanation, the key point is that the MELD folds are not random duplicates. They show consistent behavior with small but meaningful variation. That variation is useful because it tells us the model is not collapsing completely, but it is also not solving the class imbalance problem. The fold overview therefore supports a careful claim: the implementation is working and reasonably strong, but the result is still incomplete relative to the ideal paper-style target.",
    )

    add_section_title(doc, "15.2 Fold 4 Three-Way Training Comparison", level=2)
    add_para(
        doc,
        "Fold 4 is the balanced-class anchor in the MELD report, so it is the right place to check whether the alternative loss improves minority behavior. The student-level interpretation is that this fold is not chosen because it has the highest overall accuracy; it is chosen because it gives the clearest view of whether the model behaves more fairly across the emotion classes.",
    )
    add_table(
        doc,
        ["Fold 4 run", "Script used", "Training strategy", "Accuracy / W-Acc", "UW-Acc", "Macro F1", "Weighted F1"],
        [
            [
                "Weighted-CE baseline",
                "scripts/run_paper_aligned_meld_cv.sh",
                "Fresh paper-aligned run with weighted cross-entropy",
                f"{read_json(ROOT / 'results' / 'paper_aligned_meld_cv' / 'cmt_min' / 'fold_4' / 'metrics.json')['accuracy']:.4f}",
                f"{read_json(ROOT / 'results' / 'paper_aligned_meld_cv' / 'cmt_min' / 'fold_4' / 'metrics.json')['unweighted_accuracy']:.4f}",
                f"{read_json(ROOT / 'results' / 'paper_aligned_meld_cv' / 'cmt_min' / 'fold_4' / 'metrics.json')['macro_f1']:.4f}",
                f"{read_json(ROOT / 'results' / 'paper_aligned_meld_cv' / 'cmt_min' / 'fold_4' / 'metrics.json')['weighted_f1']:.4f}",
            ],
            [
                "Warm-start focal",
                "scripts/run_improvement_class_balanced_focal_meld_selected.sh",
                "Warm-start focal-loss run on the same Fold 4 split",
                f"{read_json(ROOT / 'results' / 'improvement' / 'warmstart_focal' / 'meld_selected' / 'cmt_min' / 'fold_4' / 'metrics.json')['accuracy']:.4f}",
                f"{read_json(ROOT / 'results' / 'improvement' / 'warmstart_focal' / 'meld_selected' / 'cmt_min' / 'fold_4' / 'metrics.json')['unweighted_accuracy']:.4f}",
                f"{read_json(ROOT / 'results' / 'improvement' / 'warmstart_focal' / 'meld_selected' / 'cmt_min' / 'fold_4' / 'metrics.json')['macro_f1']:.4f}",
                f"{read_json(ROOT / 'results' / 'improvement' / 'warmstart_focal' / 'meld_selected' / 'cmt_min' / 'fold_4' / 'metrics.json')['weighted_f1']:.4f}",
            ],
        ],
    )
    add_para(
        doc,
        "Fold 4 is a useful companion to Fold 2 because it shows whether the model is handling minority classes more fairly. In the current outputs, the weighted-CE baseline still beats the warm-start focal run. That tells us the alternative loss did not improve the balanced-class behavior enough to replace the baseline.",
    )

    add_section_title(doc, "15.3 Fold 2 Focal-Loss Comparison", level=2)
    add_para(
        doc,
        f"The new focal-loss improvement run for Fold 2 produced accuracy {fold2_focal['accuracy']:.4f}, weighted F1 {fold2_focal['weighted_f1']:.4f}, macro F1 {fold2_focal['macro_f1']:.4f}, and unweighted accuracy {fold2_focal['unweighted_accuracy']:.4f}. Compared with the earlier weighted-CE Fold 2 result of accuracy {meld_fold2['accuracy']:.4f}, weighted F1 {meld_fold2['weighted_f1']:.4f}, macro F1 {meld_fold2['macro_f1']:.4f}, and unweighted accuracy {meld_fold2['unweighted_accuracy']:.4f}, the focal-loss version is much worse.",
    )
    add_table(
        doc,
        ["Fold 2 run", "Accuracy / W-Acc", "UW-Acc", "Macro F1", "Weighted F1"],
        [
            ["Weighted-CE baseline", f"{meld_fold2['accuracy']:.4f}", f"{meld_fold2['unweighted_accuracy']:.4f}", f"{meld_fold2['macro_f1']:.4f}", f"{meld_fold2['weighted_f1']:.4f}"],
            ["Focal-loss improvement", f"{fold2_focal['accuracy']:.4f}", f"{fold2_focal['unweighted_accuracy']:.4f}", f"{fold2_focal['macro_f1']:.4f}", f"{fold2_focal['weighted_f1']:.4f}"],
        ],
    )
    add_para(
        doc,
        "The Fold 2 focal-loss confusion pattern is important because it explains why the result is bad in a way that a student can understand. Instead of improving minority classes, the focal-loss run collapsed into a single dominant predicted label. In the saved predictions, almost every actual class was mapped to the same predicted class, which means focal loss in this configuration did not create a richer decision boundary. It created a narrower one.",
    )
    add_bullets(
        doc,
        [
            "The earlier weighted-CE Fold 2 run was imperfect but still meaningful.",
            "The focal-loss Fold 2 run became much more biased and much less useful.",
            "This suggests that focal loss, as currently configured, is too aggressive for this setup or needs a different hyperparameter/balancing strategy.",
            "The failure is not in the architecture alone; it is in how the loss reshaped the gradient during training.",
        ],
    )
    add_para(
        doc,
        "From a student point of view, the lesson is that a theoretically better-sounding loss is not automatically better in practice. Focal loss is designed to focus on hard examples, but if the class distribution, optimizer, or fine-tuning schedule does not suit it, the model may still collapse. That is why the next step should not be to assume focal loss always wins. The next step should be to study why it failed here and whether a different imbalance-aware strategy or a gentler setting is needed.",
    )

    add_section_title(doc, "15.4 Fold 2 Three-Way Training Comparison", level=2)
    add_para(
        doc,
        "This section compares the three Fold 2 training/evaluation paths that were used in the project. They all keep the same MELD Fold 2 split and the same paper-aligned text+audio architecture, but they differ in how the objective is optimized and whether the run starts fresh or resumes from the weighted-CE checkpoint. That makes the comparison useful for isolating the effect of the training objective and the continuation strategy.",
    )
    add_para(
        doc,
        "From a student point of view, this is an ablation-style comparison. The dataset split does not change, the architecture does not change, and the paper-style text/audio path does not change. What changes is the training objective and the starting point of optimization. That is why this section is technically important: it shows whether the loss itself is helping or harming the final decision boundary.",
    )
    add_table(
        doc,
        ["Fold 2 run", "Script used", "Training strategy", "Accuracy / W-Acc", "UW-Acc", "Macro F1", "Weighted F1"],
        [
            [
                "Weighted-CE baseline",
                "scripts/run_paper_aligned_meld_cv.sh",
                "Fresh paper-aligned run with weighted cross-entropy",
                f"{meld_fold2['accuracy']:.4f}",
                f"{meld_fold2['unweighted_accuracy']:.4f}",
                f"{meld_fold2['macro_f1']:.4f}",
                f"{meld_fold2['weighted_f1']:.4f}",
            ],
            [
                "Focal from scratch",
                "scripts/run_improvement_class_balanced_focal_meld_selected.sh",
                "Fresh run with focal loss and the same fold split",
                f"{read_json(ROOT / 'results' / 'improvement' / 'warmstart_focal' / 'meld_selected' / 'cmt_min' / 'fold_2' / 'metrics.json')['accuracy']:.4f}",
                f"{read_json(ROOT / 'results' / 'improvement' / 'warmstart_focal' / 'meld_selected' / 'cmt_min' / 'fold_2' / 'metrics.json')['unweighted_accuracy']:.4f}",
                f"{read_json(ROOT / 'results' / 'improvement' / 'warmstart_focal' / 'meld_selected' / 'cmt_min' / 'fold_2' / 'metrics.json')['macro_f1']:.4f}",
                f"{read_json(ROOT / 'results' / 'improvement' / 'warmstart_focal' / 'meld_selected' / 'cmt_min' / 'fold_2' / 'metrics.json')['weighted_f1']:.4f}",
            ],
            [
                "Resume warm-start, 5 epochs",
                "scripts/run_resume_warmstart_focal_meld_fold2.sh --epochs 5",
                "Load the weighted-CE checkpoint and continue with focal loss",
                f"{read_json(ROOT / 'results' / 'improvement' / 'warmresume_focal' / 'meld_fold_2' / 'metrics.json')['accuracy']:.4f}",
                f"{read_json(ROOT / 'results' / 'improvement' / 'warmresume_focal' / 'meld_fold_2' / 'metrics.json')['unweighted_accuracy']:.4f}",
                f"{read_json(ROOT / 'results' / 'improvement' / 'warmresume_focal' / 'meld_fold_2' / 'metrics.json')['macro_f1']:.4f}",
                f"{read_json(ROOT / 'results' / 'improvement' / 'warmresume_focal' / 'meld_fold_2' / 'metrics.json')['weighted_f1']:.4f}",
            ],
        ],
    )
    add_para(
        doc,
        "The table shows why the three runs are not equivalent even though they all use the same fold and the same architecture. The weighted-CE baseline is the strongest overall result and is the closest to the paper-style MELD number. The focal-from-scratch run performs much worse, which indicates that focal loss alone, when used from initialization, was too aggressive for this dataset and model combination. The resume warm-start run is a partial recovery: it improves over focal-from-scratch, but it still does not reach the original weighted-CE baseline.",
    )
    add_para(
        doc,
        "The key technical lesson is that the learning path matters. Weighted cross-entropy gives the model a stable first solution. Focal loss from scratch changes the gradient shape too sharply and pushes the model into a poor decision boundary. Warm-start focal is better because it begins from the already sensible weighted-CE boundary, but it still cannot fully recover the baseline once the focal objective starts reshaping the classifier. This is why the section is not just about numbers; it is about how optimization history affects model behavior.",
    )
    add_bullets(
        doc,
        [
            "Weighted-CE baseline: best overall, strongest comparison point, and the cleanest paper-aligned reference.",
            "Focal from scratch: shows that changing the loss alone can destabilize the decision boundary if the run starts from random task-specific layers.",
            "Resume warm-start, 5 epochs: shows that continuing from the weighted-CE checkpoint is safer than starting from scratch, but still not enough to beat the baseline.",
        ],
    )
    add_para(
        doc,
        "For a student, the important technical lesson is that training strategy matters as much as architecture. A loss function is not a magic improvement by itself. It interacts with initialization, checkpoint history, learning rate, and how long training is allowed to continue. In this project, those interactions are visible directly in the Fold 2 metrics and confusion patterns.",
    )
    add_para(
        doc,
        "When you explain this in the review, the safe wording is: the baseline weighted-CE Fold 2 run is the best current MELD anchor; the focal-loss experiments are diagnostics showing that the alternative objective does not improve the model on this dataset; and the warm-start version is the strongest of the focal runs, but still not strong enough to replace the baseline. That lets you justify stopping loss optimization and moving on to the video stage.",
    )

    add_section_title(doc, "15.5 Phase 1 Result Summary for the Thesis Narrative", level=2)
    add_para(
        doc,
        "For the thesis narrative, the weighted-CE Phase 1 result is the best version of the current text+audio baseline and the most defensible stopping point before adding video. The 5-fold MELD CV summary gives mean accuracy 0.6247, mean weighted F1 0.6195, and mean macro F1 0.4395. Fold 2 is the strongest single-fold anchor, with accuracy 0.6375, weighted F1 0.6254, and macro F1 0.4430.",
    )
    add_table(
        doc,
        ["Phase 1 result", "Accuracy / W-Acc", "UW-Acc", "Macro F1", "Weighted F1", "Use in thesis"],
        [
            ["MELD 5-fold mean", f"{meld_summary['accuracy']['mean']:.4f}", f"{meld_summary['unweighted_accuracy']['mean']:.4f}", f"{meld_summary['macro_f1']['mean']:.4f}", f"{meld_summary['weighted_f1']['mean']:.4f}", "Primary baseline summary"],
            ["MELD Fold 2 anchor", f"{meld_fold2['accuracy']:.4f}", f"{meld_fold2['unweighted_accuracy']:.4f}", f"{meld_fold2['macro_f1']:.4f}", f"{meld_fold2['weighted_f1']:.4f}", "Best single-fold anchor"],
        ],
    )
    add_bullets(
        doc,
        [
            "The weighted-CE baseline is better than the focal-loss alternatives currently tested on Fold 2.",
            "The thesis novelty is therefore not a small MELD gain; it is the adaptation to legal testimony and the addition of video frames.",
            "A defensible summary line is: HuBERT + BERT + Cross-Entropy produced a competitive MELD baseline and serves as the backbone for LegalMemoCMT before video is added.",
        ],
    )
    add_para(
        doc,
        "If you want a compact version for the second review slide or abstract-style discussion, you can say: Phase 1 baseline, HuBERT + BERT + weighted cross-entropy, achieved approximately 62.5% mean accuracy with 61.9% weighted F1 and 44.0% macro F1 on MELD 5-fold CV. That is strong enough to justify moving forward to the video-enhanced stage without further loss-function optimization unless you want a secondary ablation.",
    )

    add_section_title(doc, "16. CREMA-D as the Primary Speech-Emotion Benchmark", level=1)
    add_para(
        doc,
        f"CREMA-D is now the project’s primary speech-emotion benchmark in the new split. The latest finished single-split paper-aligned run produced accuracy {crema['accuracy']:.4f}, weighted F1 {crema['weighted_f1']:.4f}, and macro F1 {crema['macro_f1']:.4f} on {crema['num_samples']} test samples. That result is very weak in absolute terms, and the confusion analysis shows a strong collapse toward a single class. The new CREMA-D CV workflow was created to make the speech-emotion side more paper-like and more robust than that single-split result.",
    )
    add_para(
        doc,
        "The important point is not that CREMA-D is already good. The important point is that the project now has a proper speech-emotion track with its own cross-validation structure, summary metrics, and analysis scripts. That is the correct place to improve if the goal is to move toward a stronger speech-emotion result while keeping MELD as the conversational benchmark.",
    )
    add_bullets(
        doc,
        [
            "The old single-split CREMA-D result is weak and near chance.",
            "The new CREMA-D CV scripts are intended to give a more paper-style speech-emotion evaluation.",
            "The CV output should be judged using W-Acc and UW-Acc together with fold consistency.",
            "If the CV runs still collapse, the next step is tuning or revisiting the CREMA-D training setup.",
        ],
    )
    add_para(
        doc,
        "W-Acc and UW-Acc are the two metric views that matter most on the speech-emotion side. In this codebase, W-Acc is the sample-level accuracy across all predictions, while UW-Acc is the mean of the per-class accuracies. That makes W-Acc easy to read as the overall score, but UW-Acc more sensitive to class imbalance.",
    )
    add_para(
        doc,
        "Example: if 8 of 10 test predictions are correct, then W-Acc = 0.80. But if the per-class accuracies are 1.0, 0.5, and 0.0, then UW-Acc = (1.0 + 0.5 + 0.0) / 3 = 0.50. The example shows why a model can look good overall while still failing on one or more classes.",
    )

    add_section_title(doc, "17. What Remains to Reach a Stronger Final State", level=1)
    add_numbered(
        doc,
        [
            "Run the CREMA-D CV workflow cleanly to completion and inspect its summary and fold-level behavior.",
            "Keep MELD as the conversational benchmark and decide whether the paper-aligned or paper-exact MELD branch should be treated as the final authority for the writeup.",
            "Improve class balance handling where needed, especially on MELD minority emotions and on the CREMA-D collapse pattern if it persists.",
            "Keep the reporting language consistent so the speech-emotion benchmark is described as W-Acc / UW-Acc style and the conversational benchmark is described with accuracy and F1 style metrics.",
            "Use the guidance documents to make the final benchmark choice explicit rather than leaving it implicit.",
        ],
    )
    add_para(
        doc,
        "In thesis terms, the project is no longer missing the core machinery. What remains is the final selection of benchmark emphasis, the last tuning steps on the weak benchmark, and the decision about how strict the reproduction target must be.",
    )

    add_section_title(doc, "18. Second Review Questions and Talking Points", level=1)
    add_bullets(
        doc,
        [
            "Is the new CREMA-D speaker-independent CV track the right way to represent the paper-style speech-emotion side, or should I pursue a different speech-emotion benchmark strategy?",
            "Should the completed MELD paper-aligned CV result remain the main conversational result for now, with the paper-exact MELD template treated as a future protocol refinement unless I run it to completion as well?",
            "Should the final thesis emphasize the paper-aligned MELD track and use CREMA-D only as a secondary speech-emotion benchmark, or should the new benchmark split be described as equally central?",
            "What is the best next improvement target: better imbalance handling, a stronger tuning schedule, or a stronger speech-emotion benchmark result?",
            "Is the current paper-adaptation completeness high enough to proceed with documentation and presentation refinement, or is one more protocol-level refinement needed first?",
        ],
    )

    add_section_title(doc, "19. Second Review Bottom Line", level=1)
    add_para(
        doc,
        "The project is now substantially implemented and clearly beyond the first-review stage. The MELD paper-aligned result is close to the base paper’s conversational case study, the new benchmark split is in place, and the codebase supports both conversational and speech-emotion evaluation paths. The remaining gap is not a missing system, but a combination of benchmark finalization, stronger speech-emotion performance, and the final decision about whether the project should be presented as paper-aligned or paper-exact in the thesis narrative.",
    )

    return doc


def main() -> None:
    doc = build_doc()
    doc.save(DOC_PATH)
    print(DOC_PATH)


if __name__ == "__main__":
    main()
