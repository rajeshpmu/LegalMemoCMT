from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor


ROOT = Path("/Users/rajeshpmu/Desktop/LegalMemoCMT")
DOCX_PATH = ROOT / "implementation_docments/LegalMemoCMT_Phase1_ESA_Reading_Script.docx"


def style_runs(paragraph, *, bold=False, size=11, color="000000"):
    for run in paragraph.runs:
        run.font.name = "Aptos"
        run.font.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)


def insert_before(anchor_paragraph, text, *, bold=False, size=11, color="000000"):
    new_p = OxmlElement("w:p")
    anchor_paragraph._p.addprevious(new_p)
    paragraph = Paragraph(new_p, anchor_paragraph._parent)
    paragraph.text = text
    style_runs(paragraph, bold=bold, size=size, color=color)
    return paragraph


def find_heading_indices(doc):
    headings = {}
    for idx, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("Slide ") and ":" in t:
            num_str = t.split(":", 1)[0].split()[1]
            if num_str.isdigit():
                headings[int(num_str)] = idx
    return headings


def append_block(doc, slide_num, lines):
    headings = find_heading_indices(doc)
    if slide_num not in headings:
        return
    start_idx = headings[slide_num]
    next_indices = [idx for n, idx in headings.items() if idx > start_idx]
    end_anchor = doc.paragraphs[min(next_indices)] if next_indices else None
    if end_anchor is None:
        end_anchor = doc.add_paragraph("")
    # avoid duplicate insertions if the marker already exists in this slide section
    first_marker = lines[0]
    section_text = []
    for idx in range(start_idx + 1, min(next_indices) if next_indices else len(doc.paragraphs)):
        section_text.append(doc.paragraphs[idx].text)
    if any(first_marker in t for t in section_text):
        return
    for text in reversed(lines):
        insert_before(end_anchor, text, size=11)


def main():
    doc = Document(str(DOCX_PATH))

    additions = {
        4: [
            "Expanded detail: Phase 1 should be framed as the verification step of a larger research pipeline. The point is not to claim the final courtroom system yet, but to show that the multimodal baseline is stable enough to support later legal-domain adaptation.",
            "Expanded detail: the ViT path matters because it converts raw video into cached facial embeddings. Once those embeddings exist, the same clip can be reused without recomputing frame extraction every time, which makes the pipeline measurable and reproducible.",
        ],
        7: [
            "Expanded detail: Table 10 is not only a descriptive statistic table. It explains why MELD is a hard benchmark: the turns are short, the neutral class dominates, and the model has to infer emotion from limited lexical evidence inside a dialogue.",
            "Expanded detail: because the dataset is conversation-based, the training and validation folds must be dialogue-safe. If utterances from the same dialogue leak into both splits, the benchmark becomes easier than it should be and the evaluation no longer reflects the real task difficulty.",
        ],
        8: [
            "Expanded detail: inter-speaker influence means the model may need to understand who spoke just before the current utterance and how that turn changes the emotional interpretation of the response.",
            "Expanded detail: emotion shift is important because one speaker can move from neutral to angry or from calm to sad across a dialogue. The current pipeline can partially react to this through pretrained encoders and fusion, but it does not yet keep explicit speaker memory.",
            "Expanded detail: the code path to mention is scripts/build_meld_cv_folds.py for dialogue-safe splitting, src/models/model.py for fusion behavior, src/train/train.py for imbalance-aware loss selection, and src/train/evaluate.py for confusion-matrix analysis.",
        ],
        11: [
            "Expanded detail: weighted cross entropy gives extra learning pressure to minority classes by scaling the loss contribution of rare labels. That is why it is the stable baseline choice when neutral dominates the dataset.",
            "Expanded detail: focal loss is a more aggressive imbalance-aware objective. It reduces the importance of easy examples so the model spends more gradient signal on difficult or ambiguous examples. This can help minority classes, but it can also destabilize training if the data or warm start is not well aligned.",
            "Expanded detail: the key evaluation point is that a better loss is not proven by accuracy alone. You must look at weighted F1, macro F1, and the confusion matrix to see whether minority-class behavior really improved.",
        ],
        12: [
            "Expanded detail: the design constraints are the things that make the benchmark trustworthy. Class imbalance means the project cannot be defended by accuracy alone, and the sample-to-cache alignment means every input and output must refer to the same utterance.",
            "Expanded detail: the hidden dependency is that the system relies on pretrained backbones, ffmpeg, and cached media assets. If any of these change, the reproduction story becomes weaker because the same command may no longer produce the same artifact path or the same feature tensor.",
        ],
        13: [
            "Expanded detail: the design approach is intentionally controlled. The baseline is tested first, the ViT branch is attached second, and the gating or auxiliary-loss variants are tested later so you can isolate the effect of each modification.",
            "Expanded detail: the reason this matters at MTech level is attribution. If the architecture is changed too many times at once, you lose the ability to say whether a metric change came from the visual branch, the fusion rule, the loss, or the data split.",
        ],
        14: [
            "Expanded detail: novelty here is not just adding another branch. It is the combination of a reproducible multimodal baseline and a courtroom-oriented adaptation story, with the pipeline still traceable from source files to predictions.",
            "Expanded detail: portability and maintainability are important because the same code has to run in more than one place. In practice, the pipeline should work on local development hardware and on CUDA-based cloud execution without changing the scientific meaning of the result.",
        ],
        16: [
            "Expanded detail: interoperability means that MP4, WAV, CSV manifests, cached .npy features, and checkpoints must stay synchronized. The project is only meaningful if the same utterance identity survives each stage of the pipeline.",
            "Expanded detail: reliability is demonstrated when the fold metrics, confusion matrix, and prediction CSV agree with the trained checkpoint. That is what makes the implementation reviewable rather than merely executable.",
        ],
        17: [
            "Expanded detail: the methodology is a sequence of decisions. Train, pick the best validation checkpoint, run evaluation, inspect the confusion matrix, and only then decide whether another experiment is justified.",
            "Expanded detail: this is why the next step is not chosen by intuition. It is chosen by error structure. If the confusion matrix still shows strong neutral bias or emotionally close confusions, that determines whether gating, auxiliary loss, or future context modeling is the right next move.",
        ],
        18: [
            "Expanded detail: the system architecture starts with preprocessing because raw media is too large and too inconsistent to feed directly into the model. Preprocessing normalizes the data into reusable artifacts that can be tracked in manifests.",
            "Expanded detail: once frame sampling, face cropping, audio extraction, and text normalization are complete, the same sample can travel through training, evaluation, and demo code without any ambiguity about what was actually used.",
        ],
        19: [
            "Expanded detail: BERT is responsible for contextual semantics in text, HuBERT is responsible for audio representation, and ViT is responsible for facial cues. These are different information sources, so the model needs a fusion stage rather than one shared encoder.",
            "Expanded detail: cross-modal transformer blocks allow one modality to influence another through attention, while gated fusion adds a learned control mechanism that can decide how much each modality should contribute before classification.",
        ],
        20: [
            "Expanded detail: Python and PyTorch form the machine learning core because they make the encoders, losses, and training loops easy to express. Hugging Face transformers provide the pretrained BERT and HuBERT model loading logic, while OpenCV and ffmpeg handle the media side of the pipeline.",
            "Expanded detail: numpy, pandas, and scikit-learn are used for manifest processing, metric computation, and evaluation reports. These libraries are not the novelty themselves, but they are what make the experiment reproducible and analyzable.",
        ],
        21: [
            "Expanded detail: saying the project is 'in progress' is not enough. You should explain that the baseline is complete, the facial-cue branch is implemented, and the current model behavior has already been analyzed through metrics and confusion outputs.",
            "Expanded detail: that makes the work a research checkpoint rather than a partially finished prototype. The next iteration should be guided by the observed failure modes, not by a vague desire to make the number bigger.",
        ],
        23: [
            "Expanded detail: the top-3 output values are softmax probabilities, so they represent the model’s confidence distribution over emotion classes. They are useful because they let you see whether the model was uncertain, nearly correct, or confidently wrong.",
            "Expanded detail: if a prediction is wrong, the right explanation is usually not 'the model failed' but a more technical reason such as label ambiguity, neutral dominance, speaker context, or the visual branch being overruled by stronger text/audio evidence.",
        ],
        25: [
            "Expanded detail: explicit inter-speaker modeling would add speaker identity or speaker-state memory so the model can learn who said what before the current turn. Emotion-shift modeling would try to capture how emotions change across turns instead of predicting each utterance in isolation.",
            "Expanded detail: long-range dialogue context would extend the model beyond local neighboring turns, while courtroom-role metadata would let Phase 2 use legal structure such as witness, judge, or attorney roles. Those ideas are valuable, but they increase complexity, so they are better treated as the next technical layer after Phase 1 is locked.",
            "Expanded detail: if you have a little time left in Phase 1, prototype only a small context ablation on MELD. If not, keep it as the bridge into Phase 2 so the current benchmark remains stable and defensible.",
        ],
        26: [
            "Expanded detail: this slide should be read as the proof slide for the demo. Each row uses the same raw clip for both checkpoints, so the comparison is controlled: the only thing changing is the checkpoint and its learned fusion behavior.",
            "Expanded detail: test_dia279_utt9 is the clearest improvement because the baseline is almost tied between anger and neutral, but the gated+aux checkpoint pushes neutral far above the rest. That means the later model learned a much cleaner decision boundary for a short, weakly emotional utterance.",
            "Expanded detail: test_dia4_utt6 is a stable success case. Both checkpoints are correct, but the gated+aux model separates disgust from anger much more strongly, which shows sharper class separation rather than just a lucky top-1 pick.",
            "Expanded detail: test_dia278_utt5 is the strongest reminder not to overclaim. The transcript sounds playful or positive, so both checkpoints lean toward joy, but the true label is surprise. The later model becomes more confident in the wrong class, so this is a real failure, not a hidden win.",
            "Expanded detail: test_dia153_utt5 and test_dia244_utt14 are hard neutral failures. Their wording contains fear-like or angry surface cues, so both checkpoints move away from neutral. The gated+aux model does not fix that issue; it only changes which non-neutral class wins.",
            "Expanded detail: the actual conclusion is selective improvement. Gated+aux helps on the boundary neutral clip and keeps one non-neutral success strong, but it does not solve all ambiguous samples. That is the honest MTech-level reading of the results.",
        ],
    }

    # Insert from the end so paragraph indexes remain stable.
    for slide_num in sorted(additions.keys(), reverse=True):
        append_block(doc, slide_num, additions[slide_num])

    doc.save(str(DOCX_PATH))
    print(f"Updated {DOCX_PATH}")


if __name__ == "__main__":
    main()
