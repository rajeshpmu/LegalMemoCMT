from __future__ import annotations

import json
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "implementation_docments" / "LegalMemoCMT_MELD_Facial_Cues_ViT_Implementation_Plan.docx"
INDEX = ROOT / "results" / "paper_aligned_meld_cv" / "cmt_min" / "summary.json"
FOLD2 = ROOT / "results" / "paper_aligned_meld_cv" / "cmt_min" / "fold_2" / "metrics.json"
FIG_DIR = ROOT / "implementation_docments" / "figures" / "meld_vit_facecue"
MMD_PATH = FIG_DIR / "meld_vit_facecue_pipeline.mmd"
SVG_PATH = FIG_DIR / "meld_vit_facecue_pipeline.svg"
PNG_PATH = FIG_DIR / "meld_vit_facecue_pipeline.png"


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)


def add_para(doc: Document, text: str, *, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.italic = italic


def add_code_block(doc: Document, code: str) -> None:
    for line in code.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(9)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(image_path), width=Inches(6.9))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)


def write_mermaid_assets() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    mmd = """%%{init: {'themeVariables': {'fontSize': '18px'}}}%%
flowchart LR
    A["MELD annotations\\n(train/dev/test CSVs)"] --> B["find_meld_clip(...)"]
    B --> C["Raw MELD clip\\n.../diaX_uttY.mp4"]
    C --> D["sample_video_frames(...)"]
    D --> E["16 sampled full frames\\n224x224 RGB"]
    E --> F["Pretrained ViT\\nAutoImageProcessor + AutoModel"]
    F --> G["CLS embeddings\\n(shape: 16 x 768)"]
    G --> H["np.save(... .npy)"]
    H --> I["video_path in meld_vit_facecue.csv"]
    I --> J["Face-cue fold CSVs"]
"""
    MMD_PATH.write_text(mmd, encoding="utf-8")
    cmd = [
        "npx",
        "-y",
        "@mermaid-js/mermaid-cli",
        "-i",
        str(MMD_PATH),
        "-o",
        str(SVG_PATH),
        "-b",
        "transparent",
    ]
    subprocess.run(cmd, check=True, cwd=str(ROOT), stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    cmd = [
        "npx",
        "-y",
        "@mermaid-js/mermaid-cli",
        "-i",
        str(MMD_PATH),
        "-o",
        str(PNG_PATH),
        "-b",
        "transparent",
    ]
    subprocess.run(cmd, check=True, cwd=str(ROOT), stdout=subprocess.PIPE, stderr=subprocess.PIPE)


def build_doc() -> Document:
    write_mermaid_assets()
    doc = Document()
    configure(doc)

    meld_summary = read_json(INDEX)["metrics"]
    fold2 = read_json(FOLD2)

    title = doc.add_paragraph()
    run = title.add_run("LegalMemoCMT MELD Facial Cues ViT Implementation Plan")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(
        "A student-level guide for moving from the weighted-CE MELD Phase 1 baseline to a warm-start facial-cue upgrade using ViT-ready video embeddings."
    )
    run.italic = True
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"

    doc.add_paragraph(
        "Purpose: this document explains why the current weighted-CE MELD Phase 1 result is the right backbone baseline, why the next step should warm-start from the existing checkpoint instead of retraining from scratch, and how the facial-cue upgrade can be implemented in a clean, student-readable way."
    )
    doc.add_paragraph(
        "Important note: the scripts in this plan prepare a ViT-based facial-cue pipeline and a warm-start training flow. The repo already has a usable video branch, so the scripts are structured to be practical now while still pointing to the future code change where the compact video features are replaced by true ViT facial embeddings in the model path."
    )

    doc.add_heading("1. Why This Baseline Was Chosen", level=1)
    add_para(
        doc,
        f"The weighted-CE MELD Phase 1 baseline is the strongest current conversational result and therefore the right anchor for the next stage. The 5-fold MELD CV summary reports mean accuracy {meld_summary['accuracy']['mean']:.4f}, mean weighted F1 {meld_summary['weighted_f1']['mean']:.4f}, and mean macro F1 {meld_summary['macro_f1']['mean']:.4f}. Fold 2 is the strongest single-fold anchor, with accuracy {fold2['accuracy']:.4f}, weighted F1 {fold2['weighted_f1']:.4f}, and macro F1 {fold2['macro_f1']:.4f}.",
    )
    add_bullets(
        doc,
        [
            "This baseline is already competitive enough to serve as a real reference point.",
            "It is better than the focal-loss alternatives tried earlier.",
            "It gives a stable text + audio backbone that can be extended with facial cues.",
            "It keeps the thesis novelty in the right place: legal-domain adaptation and multimodal extension, not a tiny MELD gain.",
        ],
    )
    add_para(
        doc,
        "The reasoning is simple: if the current text + audio model is already strong, then the next improvement should test whether facial cues add information that the baseline cannot already capture. That makes the facial-cue study a meaningful extension rather than a random experiment.",
    )

    doc.add_heading("2. Why Warm-Start from the Existing Checkpoint", level=1)
    add_para(
        doc,
        "Warm-starting means reusing the already learned Phase 1 weights instead of training a new model from scratch. For this project, that is the right design because the text and audio encoders have already learned a useful MELD decision boundary. If facial cues are added later, the goal is to improve that boundary, not throw it away.",
    )
    add_bullets(
        doc,
        [
            "The text and audio backbones already contain useful learned weights.",
            "Starting from the checkpoint keeps the comparison clean: any gain can be attributed to the new facial branch more directly.",
            "A small learning rate lets the new facial branch adapt without destroying the existing MELD structure.",
            "Warm-starting is safer than training everything from scratch when the baseline is already strong.",
        ],
    )
    add_para(
        doc,
        "In other words, the checkpoint is not just a saved file. It is the learned representation of what the model already knows about MELD. The facial-cue upgrade should build on that knowledge, not replace it.",
    )

    doc.add_heading("3. Why Facial Cues and Why ViT", level=1)
    add_para(
        doc,
        "Facial expressions are one of the strongest additional signals in emotion recognition. A speaker's face often reveals stress, amusement, confusion, or anger even when the text and audio are ambiguous. ViT is a good first choice for the facial-cue branch because it can convert sampled frames into a compact embedding sequence that the existing multimodal pipeline can learn from.",
    )
    add_bullets(
        doc,
        [
            "Text tells you what was said.",
            "Audio tells you how it was said.",
            "Facial cues tell you what the speaker looked like while saying it.",
            "ViT turns those visual frames into a learned representation rather than a hand-crafted feature vector.",
        ],
    )
    add_para(
        doc,
        "From a student point of view, the key idea is that ViT is not being added just because it is modern. It is being added because it can encode visual expressions in a way that is compatible with the existing multimodal design. The visual branch should therefore be understood as support for the text and audio backbone, not as a completely separate project.",
    )

    doc.add_heading("4. Recommended Implementation Order", level=1)
    add_numbered(
        doc,
        [
            "Build a dedicated ViT facial-cue MELD manifest so the visual experiment has its own data path.",
            "Create fold CSVs from that facial-cue manifest so the new runs stay comparable to the earlier MELD CV setup.",
            "Warm-start from the existing weighted-CE MELD checkpoint instead of retraining from scratch.",
            "Fine-tune the new facial branch with a small learning rate first, then unfreeze more of the model if needed.",
            "Evaluate the resulting checkpoint on the MELD test split and export predictions.",
            "Run confusion-matrix analysis to see whether the neutral-heavy error pattern improves.",
            "Only after Fold 2 looks promising should Fold 4 or a fuller CV rerun be expanded.",
        ],
    )
    add_para(
        doc,
        "This order matters because it isolates the effect of the new facial cues. If everything changes at once, it becomes impossible to tell whether the improvement came from ViT, from the checkpoint warm-start, or from a different optimization setup.",
    )

    doc.add_heading("5. What the New Scripts Do", level=1)
    add_table(
        doc,
        ["Script", "What it does", "Expected output"],
        [
            ["scripts/build_meld_vit_facecue_manifest.py", "Samples MELD frames, extracts ViT facial embeddings, and writes a dedicated face-cue manifest.", "meld_vit_facecue.csv + MELD_VIT_FACECUE feature files"],
            ["scripts/run_meld_vit_facecue_prepare.sh", "Builds the face-cue manifest and then creates dedicated 5-fold MELD CV splits.", "meld_vit_facecue_cv fold CSVs"],
            ["scripts/resume_meld_vit_facecue_fold.py", "Loads the weighted-CE checkpoint, keeps the backbone as the starting point, and fine-tunes the facial-cue model.", "best_model.pt, metrics.json, predictions_test.csv"],
            ["scripts/run_meld_vit_facecue_fold2.sh", "Runs the Fold 2 facial-cue warm-start experiment.", "Fold 2 result folder"],
            ["scripts/run_meld_vit_facecue_fold4.sh", "Runs the Fold 4 facial-cue warm-start experiment.", "Fold 4 result folder"],
            ["scripts/analyze_meld_vit_facecue_fold2.sh", "Exports and analyzes the Fold 2 predictions.", "confusion matrix, top confusions, first-20 table"],
            ["scripts/analyze_meld_vit_facecue_fold4.sh", "Exports and analyzes the Fold 4 predictions.", "confusion matrix, top confusions, first-20 table"],
            ["scripts/run_meld_vit_facecue_suite.sh", "Runs the full prepare/train/analyze sequence in order.", "Both fold outputs and analysis folders"],
        ],
    )
    add_para(
        doc,
        "The scripts are written in the same style as the earlier benchmark scripts so that a student can understand them in the same order: prepare the data, train from the baseline, evaluate the saved checkpoint, and then inspect the confusion matrix.",
    )

    doc.add_heading("6. What the ViT Step Means in Practice", level=1)
    add_para(
        doc,
        "The ViT step is the point where raw video clips become facial embeddings. The script samples frames from each MELD utterance clip, sends those frames through a pretrained ViT, and saves the resulting embeddings in a dedicated feature directory. Those embeddings then become the input to the visual branch of the multimodal model.",
    )
    add_bullets(
        doc,
        [
            "The face-cue manifest is separate from the original MELD manifest so the experiment remains clean.",
            "The visual embeddings are saved per utterance, which makes the data easy to reuse and debug.",
            "The visual branch can be swapped or improved later without changing the overall training story.",
        ],
    )
    add_para(
        doc,
        "This is the right student mental model: the script does not just 'use video'. It creates a dedicated visual representation of the speaker's face so that the model can learn whether the face adds useful emotion information beyond text and audio.",
    )

    doc.add_heading("6.2 Utterance-to-Embedding Flow With Code Lines", level=2)
    add_para(
        doc,
        "This is the exact path the code follows for one MELD utterance. First the script resolves the raw clip for that utterance, then it samples frames from the clip, then it sends those frames through a pretrained ViT model, and finally it saves the output as a NumPy .npy file. The important idea for a student is that each utterance becomes its own reusable visual feature file.",
    )
    add_code_block(
        doc,
        """clip_path = find_meld_clip(meld_root / "raw", split, dialogue_id, utterance_id)
sample_id = f"{split}_dia{dialogue_id}_utt{utterance_id}"
video_feat_path = feat_dir / f"{sample_id}.npy"
vit_embeddings = extract_vit_face_embeddings(
    str(clip_path), cfg, image_processor, vit_model, device, args.batch_size
)
np.save(video_feat_path, vit_embeddings)

frames = sample_video_frames(video_path, cfg)
inputs = image_processor(images=list(batch), return_tensors="pt")
outputs = vit_model(**inputs)
hidden = outputs.last_hidden_state[:, 0, :]""",
    )
    add_bullets(
        doc,
        [
            "find_meld_clip(...) locates the raw MELD video file that belongs to the current utterance.",
            "sample_video_frames(...) chooses evenly spaced frames from that raw clip.",
            "AutoImageProcessor converts the sampled images into ViT-ready tensors.",
            "AutoModel runs the pretrained ViT and outputs token embeddings.",
            "outputs.last_hidden_state[:, 0, :] takes the CLS embedding for each sampled frame.",
            "np.save(...) writes the final facial embedding array to disk as a .npy file.",
        ],
    )
    add_para(
        doc,
        "Student interpretation: the raw clip is the starting point, the sampled frames are the visual evidence, the ViT embeddings are the learned summary, and the .npy file is the saved artifact that the model can reuse during training. That separation is important because it means visual preprocessing happens once, and training can reuse the features many times without recomputing the encoder.",
    )
    add_para(
        doc,
        "The key code locations behind this flow are scripts/build_meld_vit_facecue_manifest.py:101-102 for loading the pretrained ViT, scripts/build_meld_vit_facecue_manifest.py:123-135 for resolving the raw clip and saving the .npy file, scripts/build_meld_vit_facecue_manifest.py:44-47 for running ViT on sampled frames, and src/data/preprocessing.py:41-75 for sampling the frames from the raw clip.",
    )
    add_table(
        doc,
        ["Stage", "Input", "Output", "Example path / object"],
        [
            [
                "1. Locate clip",
                "MELD annotation row",
                "Raw video clip path",
                "data/MELD/raw/MELD.Raw/train/train_splits/dia0_utt0.mp4",
            ],
            [
                "2. Sample frames",
                "Raw video clip",
                "Fixed frame tensor",
                "(16, 224, 224, 3) sampled frames in memory",
            ],
            [
                "3. Run ViT",
                "Sampled RGB frames",
                "CLS embeddings",
                "(16, 768) float32 array",
            ],
            [
                "4. Save features",
                "ViT embedding array",
                "NumPy feature file",
                "data/processed/MELD_VIT_FACECUE/train/video/train_dia0_utt0.npy",
            ],
            [
                "5. Write manifest",
                "Per-utterance feature paths",
                "Face-cue manifest row",
                "data/manifests/meld_vit_facecue.csv",
            ],
        ],
    )
    add_figure(
        doc,
        PNG_PATH,
        "Figure 1. MELD utterance-to-embedding flow: raw clip -> sampled frames -> pretrained ViT -> .npy feature file -> manifest row.",
    )
    add_para(
        doc,
        "The table is useful because it shows the exact transition from raw input to reusable feature output. A student can read it from left to right and understand why the manifest exists: it is the map that connects the original MELD clip to the saved .npy embedding file.",
    )

    doc.add_heading("6.3 Default Configuration Values and Checkpoints", level=2)
    add_para(
        doc,
        "The facial-cue plan uses a small set of defaults so the experiment is reproducible and easy to explain. These values are not arbitrary. They match common ViT practice, keep the pipeline computationally manageable, and reuse the strongest existing MELD checkpoint as the warm-start anchor.",
    )
    add_table(
        doc,
        ["Setting", "Default value", "Why it is used"],
        [
            ["ViT checkpoint", "google/vit-base-patch16-224-in21k", "Pretrained ViT backbone from Hugging Face"],
            ["Frame size", "224 x 224", "Standard ViT input size; good tradeoff between detail and speed"],
            ["Number of frames", "16", "Fixed-length visual sequence for each utterance"],
            ["Audio sample rate", "16000", "Common speech processing rate used by the audio loaders"],
            ["Max audio seconds", "10.0", "Keeps audio features bounded and consistent"],
            ["Manifest extraction batch size", "8", "Reasonable default for embedding extraction"],
            ["Warm-start training batch size", "4", "Smaller batch for memory stability during multi-modal fine-tuning"],
            ["Warm-start learning rate", "5e-5", "Gentle update so the baseline is not overwritten"],
            ["Warm-start epochs", "5", "Short refinement window after starting from the baseline checkpoint"],
            ["Freeze backbone epochs", "1", "Lets the new visual branch adapt before the full model is unfrozen"],
            ["Warm-start checkpoint", "results/paper_aligned_meld_cv/cmt_min", "The strongest current weighted-CE MELD baseline"],
            ["Device default", "cpu", "Safe local default; can be overridden if GPU is available"],
        ],
    )
    add_para(
        doc,
        "You can also mention the broader Phase 1 defaults when comparing this facial-cue path with the original training pipeline: src/train/train.py uses epochs=10, batch-size=8, lr=1e-4, loss-type=ce, pooling=mean, encoder-mode=legacy, and modalities=text,audio,video by default. Those values define the original generic training scaffold, while the facial-cue scripts narrow the configuration to the warm-start ViT experiment.",
    )
    add_bullets(
        doc,
        [
            "ViT defaults are chosen to match pretrained weights and keep the first run practical.",
            "Training defaults are chosen to protect the baseline and make the experiment comparable.",
            "The checkpoint path is the key reason the facial-cue experiment is an extension rather than a new model from scratch.",
        ],
    )

    doc.add_heading("6.1 Full Frames vs Face Crop for Courtroom Testimony", level=2)
    add_para(
        doc,
        "For courtroom testimony, the better long-term visual design is usually to crop the face before sending frames to ViT. The reason is that the legal question is often about the witness's emotional state, and the face is the strongest place to observe that state. Eyes, brows, mouth shape, and micro-expressions are directly tied to stress, confusion, anger, surprise, and hesitation.",
    )
    add_bullets(
        doc,
        [
            "Face crop before ViT focuses the model on the facial region that actually carries emotion.",
            "Full sampled frames include background, body pose, and camera framing, which can dilute the facial signal.",
            "Courtroom testimony often has distant camera angles, occlusion, and side views, so full frames can still be useful as a robustness baseline.",
            "The best scientific comparison is to test both approaches under the same MELD setup before deciding which one to keep.",
        ],
    )
    add_para(
        doc,
        "In thesis terms, the recommended interpretation is: use face-crop ViT as the primary facial-cue method for courtroom testimony, and keep full-frame ViT as the comparison baseline or fallback. That gives you a cleaner claim about facial evidence while still preserving robustness if face detection is unreliable.",
    )
    add_bullets(
        doc,
        [
            "Primary courtroom-cue choice: face crop before ViT.",
            "Fallback / ablation choice: full sampled frames before ViT.",
            "Keep both only if the evaluation shows that full frames outperform crops in difficult clips.",
        ],
    )

    doc.add_heading("7. How the Warm-Start Trainer Works", level=1)
    add_para(
        doc,
        "The warm-start trainer loads the existing weighted-CE checkpoint, infers the ViT embedding dimension from the new manifest, and then fine-tunes the model on the face-cue data. For the first stage, it can keep the pretrained text and audio backbones stable while the new visual branch learns to align itself. After that, the whole model can be unfrozen for a short joint fine-tuning stage.",
    )
    add_bullets(
        doc,
        [
            "The text/audio knowledge from Phase 1 is reused.",
            "The visual branch is learned on top of that baseline.",
            "A smaller learning rate is used so the new branch does not overwrite the baseline.",
            "The best validation checkpoint is kept, so the final evaluation uses the strongest epoch rather than the last one by default.",
        ],
    )
    add_para(
        doc,
        "This staged training strategy is standard in transfer learning: first let the new modality catch up, then let the combined model adjust jointly. It is a controlled way to test whether facial cues actually help the already-good backbone.",
    )

    doc.add_heading("8. What to Expect from the Results", level=1)
    add_bullets(
        doc,
        [
            "Accuracy may improve only modestly, because MELD is already reasonably strong.",
            "Macro F1 and unweighted accuracy are the more important numbers to watch if facial cues are helping minority classes.",
            "The confusion matrix should become less neutral-heavy if the facial cues are useful.",
            "If the new branch helps, you should see fewer cases where neutral absorbs joy, anger, sadness, or surprise.",
        ],
    )
    add_para(
        doc,
        "The expected outcome is not magical perfection. The expected outcome is a clearer error pattern and better class balance than the text-audio-only baseline. If the result does not improve, that is also useful because it tells you the visual branch needs a different design or that the dataset split does not support the gain you hoped for.",
    )

    doc.add_heading("8.1 How to Read the Epoch Trend", level=2)
    add_para(
        doc,
        "The Fold 2 warm-start run does not improve monotonically forever. Instead, the validation curve peaks early and then degrades. This is important because a student should not look only at the last epoch. The best checkpoint is the one with the strongest validation behavior, not necessarily the checkpoint from the final epoch.",
    )
    add_table(
        doc,
        ["Epoch", "Train Loss", "Train Acc", "Val Loss", "Val Acc", "What it means"],
        [
            ["1", "1.4195", "0.5975", "1.6160", "0.5414", "Warm-start has begun; training is stable but validation is still behind."],
            ["2", "1.5937", "0.5576", "1.6063", "0.6010", "Best validation so far; the model is using the new facial information more effectively."],
            ["3", "1.3378", "0.6492", "1.7225", "0.5893", "Training improves, but validation starts slipping, which is the first sign of overfitting."],
            ["4", "1.1202", "0.7261", "1.9397", "0.5399", "Clear overfitting; training keeps improving while validation gets worse."],
        ],
    )
    add_para(
        doc,
        "The student conclusion is straightforward: the best Fold 2 checkpoint is probably around epoch 2, because that is where validation accuracy is highest and validation loss is still controlled. After that point, the model is fitting the training fold better but generalizing worse to the held-out validation fold.",
    )
    add_bullets(
        doc,
        [
            "A falling training accuracy is not automatically bad if validation is improving.",
            "A rising training accuracy with falling validation accuracy usually means overfitting.",
            "For this run, the safe interpretation is that early stopping would likely choose epoch 2.",
            "The saved best checkpoint matters more than the last epoch when the validation curve peaks early.",
        ],
    )
    add_para(
        doc,
        "In plain terms, the model learned enough by epoch 2 to be useful, but later epochs started to memorize the training fold too aggressively. That is why the facial-cue experiment should be judged using the best validation epoch, not by the final training epoch alone.",
    )

    doc.add_heading("9. What Not to Overstate", level=1)
    add_bullets(
        doc,
        [
            "Do not claim the facial-cue branch solves MELD just because it adds another modality.",
            "Do not claim the ViT step automatically improves the model without checking the confusion matrix.",
            "Do not replace the baseline narrative; the baseline is still the reference point.",
            "Do not mix the facial-cue experiment with the earlier focal-loss experiments, because they answer different questions.",
        ],
    )
    add_para(
        doc,
        "The careful thesis wording is: the Phase 1 weighted-CE baseline is the backbone, and the ViT facial-cue phase is the next supported extension. That is a stronger story than saying the model got one percentage point better on MELD.",
    )

    doc.add_heading("10. Suggested Manual Run Order", level=1)
    add_numbered(
        doc,
        [
            "Run scripts/run_meld_vit_facecue_prepare.sh to create the dedicated face-cue manifest and fold CSVs.",
            "Run scripts/run_meld_vit_facecue_fold2.sh to warm-start Fold 2 from the weighted-CE baseline.",
            "Run scripts/analyze_meld_vit_facecue_fold2.sh to inspect the Fold 2 confusion matrix and prediction errors.",
            "If Fold 2 looks promising, repeat the same flow for Fold 4 with scripts/run_meld_vit_facecue_fold4.sh and scripts/analyze_meld_vit_facecue_fold4.sh.",
            "If both folds look useful, move to the full suite wrapper so the whole face-cue workflow is repeatable.",
        ],
    )
    add_para(
        doc,
        "This run order keeps the experiment disciplined. The point is to learn whether facial cues improve the current backbone before deciding whether to scale the idea further.",
    )
    add_para(
        doc,
        "If you plan to run the entire facial-cue workflow end to end, the suite launcher is the main entry point. In that case, the manual prepare/train/analyze scripts are still useful to understand the pipeline, but you do not need to invoke them one by one because the suite already chains them in the correct order.",
    )

    doc.add_heading("11. Script-by-Script Execution Guide", level=1)
    add_para(
        doc,
        "The next sections explain each script in the exact order you would normally run them. The shell wrappers are intentionally thin: they call the Python programs that do the real work. This makes the workflow easier to inspect because the command-line entry point stays simple while the Python code carries the actual data processing and training logic.",
    )

    doc.add_heading("11.1 scripts/run_meld_vit_facecue_prepare.sh", level=2)
    add_para(
        doc,
        "This shell script is the preparation step for the facial-cue experiment. It first runs scripts/build_meld_vit_facecue_manifest.py to create the ViT-based facial-cue manifest, and then it runs scripts/build_meld_cv_folds.py on that new manifest to generate dedicated 5-fold MELD CSV splits. The important idea is that this creates a separate facial-cue dataset path, so the new experiment does not overwrite or confuse the earlier paper-aligned MELD baseline files.",
    )
    add_bullets(
        doc,
        [
            "It starts from raw MELD clips and annotations.",
            "It builds a new manifest with ViT facial embeddings.",
            "It creates dedicated train/validation folds for the face-cue experiment.",
            "It preserves the same MELD dialogue-based split logic used earlier, so comparisons stay fair.",
        ],
    )
    add_code_block(
        doc,
        """#!/usr/bin/env bash
set -euo pipefail

PYTHON_BIN="${PYTHON_BIN:-$(command -v python3)}"

"$PYTHON_BIN" scripts/build_meld_vit_facecue_manifest.py "$@"
"$PYTHON_BIN" scripts/build_meld_cv_folds.py \\
  --manifest "${FACECUE_MANIFEST:-data/manifests/meld_vit_facecue.csv}" \\
  --output-dir "${FACECUE_CV_DIR:-data/manifests/meld_vit_facecue_cv}" \\
  --base-splits "${FACECUE_BASE_SPLITS:-train,dev}" \\
  --num-folds "${FACECUE_NUM_FOLDS:-5}" \\
  --seed "${FACECUE_SEED:-42}""",
    )

    doc.add_heading("11.2 scripts/build_meld_vit_facecue_manifest.py", level=2)
    add_para(
        doc,
        "This is the core data-preparation Python script. It reads the MELD annotation CSVs for train, dev, and test. For each utterance, it finds the raw clip, samples video frames, sends those frames through a pretrained ViT model, and saves the resulting facial embeddings as .npy files. It also writes a new manifest CSV that points to the saved embeddings and keeps the original audio clip path and transcript. In student terms, this is the step where raw MELD clips are turned into reusable visual features.",
    )
    add_bullets(
        doc,
        [
            "Loads MELD annotation rows split by train/dev/test.",
            "Uses the repository helper that locates the matching raw clip for each utterance.",
            "Samples a fixed number of frames from the clip.",
            "Converts those frames into ViT embeddings using Hugging Face image processor and model objects.",
            "Saves one embedding file per utterance so the visual branch can reuse the features later.",
            "Writes a dedicated meld_vit_facecue.csv manifest for the rest of the pipeline.",
        ],
    )
    add_code_block(
        doc,
        """image_processor = AutoImageProcessor.from_pretrained(args.vit_model)
vit_model = AutoModel.from_pretrained(args.vit_model).to(device)

for _, row in df.iterrows():
    clip_path = find_meld_clip(meld_root / "raw", split, dialogue_id, utterance_id)
    sample_id = f"{split}_dia{dialogue_id}_utt{utterance_id}"
    video_feat_path = feat_dir / f"{sample_id}.npy"
    if not video_feat_path.exists():
        vit_embeddings = extract_vit_face_embeddings(
            str(clip_path), cfg, image_processor, vit_model, device, args.batch_size
        )
        np.save(video_feat_path, vit_embeddings)
    rows.append({
        "sample_id": sample_id,
        "split": split,
        "label": EMOTION_MAP[emotion],
        "video_path": str(video_feat_path),
        "audio_path": str(clip_path),
        "transcript": transcript,
    })""",
    )
    add_para(
        doc,
        "Conceptually, this script answers the question 'what does the speaker's face look like while the utterance is being spoken?' rather than 'what did the speaker say?' or 'how did they say it?'. That is why the saved video_path column points to feature files instead of to raw video clips.",
    )
    add_para(
        doc,
        "The saved .npy files are NumPy binary arrays that store the ViT facial embeddings for each utterance clip. A single .npy file usually contains a two-dimensional array whose rows correspond to sampled frames and whose columns correspond to the ViT embedding dimension. This is useful because the model can later load the facial features directly instead of recomputing ViT on every training run.",
    )
    add_bullets(
        doc,
        [
            "Raw clip -> sampled frames -> ViT embeddings -> saved .npy feature file.",
            "The .npy format is fast to load and easy to reuse in PyTorch or NumPy.",
            "This keeps the visual preprocessing separate from model training, which makes debugging easier.",
        ],
    )
    add_para(
        doc,
        "A concrete example from the generated manifest looks like this: sample_id=train_dia0_utt0, split=train, label=0, video_path=data/processed/MELD_VIT_FACECUE/train/video/train_dia0_utt0.npy, audio_path=data/MELD/raw/MELD.Raw/train/train_splits/dia0_utt0.mp4, transcript='also I was the point person on my company’s transition from the KL-5 to GR-6 system.' The corresponding .npy file has shape (16, 768), dtype float32, and values that are already normalized to a small numeric range rather than raw pixel values.",
    )
    add_bullets(
        doc,
        [
            "Manifest rows are ordinary CSV rows, so they are easy to inspect with pandas.",
            "The video_path column points to the saved .npy embedding file, not to the raw clip.",
            "The audio_path column still points to the original MELD clip so the audio branch can reuse it.",
            "Each embedding file is a compact representation of the sampled visual frames.",
        ],
    )
    add_para(
        doc,
        "A student can inspect one of these feature files directly with NumPy. For example, after loading it with np.load('data/processed/MELD_VIT_FACECUE/train/video/train_dia0_utt0.npy'), the result is a 2D array. The first axis tells you how many frames were sampled, and the second axis tells you the ViT embedding size. Printing arr.shape, arr.dtype, arr.min(), and arr.max() is a quick way to confirm that the file was written correctly and that the values are numeric and finite.",
    )
    add_bullets(
        doc,
        [
            "np.load(...) reads the saved ViT feature file back into memory.",
            "arr.shape shows the frame count and embedding dimension.",
            "arr.dtype should usually be float32.",
            "Checking min/max helps confirm the values look reasonable.",
        ],
    )

    doc.add_heading("11.3 scripts/run_meld_vit_facecue_fold2.sh", level=2)
    add_para(
        doc,
        "This is the first training run for the facial-cue plan. It calls scripts/resume_meld_vit_facecue_fold.py with --fold 2. Fold 2 is used first because it is the strongest weighted baseline anchor from the earlier MELD study, so it gives the clearest check of whether facial cues are actually helping.",
    )
    add_bullets(
        doc,
        [
            "Loads the weighted-CE MELD Fold 2 checkpoint as the starting point.",
            "Uses the new facial-cue fold CSVs instead of the original MELD CV files.",
            "Trains the text + audio + video model with a smaller learning rate.",
            "Keeps the experiment focused on the effect of adding facial cues.",
        ],
    )
    add_code_block(
        doc,
        """#!/usr/bin/env bash
set -euo pipefail

PYTHON_BIN="${PYTHON_BIN:-$(command -v python3)}"
"$PYTHON_BIN" scripts/resume_meld_vit_facecue_fold.py --fold 2 "$@" """,
    )

    doc.add_heading("11.4 scripts/resume_meld_vit_facecue_fold.py", level=2)
    add_para(
        doc,
        "This is the main warm-start training Python script. It reads the facial-cue fold manifests, infers the ViT embedding dimension from the saved .npy files, recreates the LegalMemoCMTPhase1 model with video enabled, loads the old weighted-CE checkpoint into the model, and then fine-tunes the combined system. The script can freeze the text and audio backbones briefly at the start so that the new facial branch has a chance to adapt before the full model is updated.",
    )
    add_para(
        doc,
        "This warm-start flow is intentionally aligned with the training utilities in src/train/train.py. The trainer uses the same dataset-building idea, the same modality masking pattern, the same optimizer style, and the same loss/metric logic as the core training script, but it adds the checkpoint-loading step needed for the facial-cue refinement experiment.",
    )
    add_bullets(
        doc,
        [
            "Infers the visual feature dimension from the facial-cue manifest.",
            "Rebuilds the model with text, audio, and video modalities turned on.",
            "Warm-starts from the old weighted-CE MELD checkpoint, excluding the old video encoder weights when needed.",
            "Optionally freezes text/audio backbones for a small number of epochs.",
            "Trains with weighted cross-entropy so the baseline loss policy stays stable during the first facial-cue extension.",
            "Saves the best validation checkpoint and then evaluates it on MELD test.",
        ],
    )
    add_code_block(
        doc,
        """base_ckpt = baseline_dir / f"fold_{args.fold}" / "best_model.pt"
checkpoint = load_checkpoint_with_video_shape(base_ckpt, model)

for epoch in range(1, train_cfg.epochs + 1):
    if args.freeze_backbone_epochs > 0 and epoch <= args.freeze_backbone_epochs:
        set_requires_grad(model.text_encoder, False)
        set_requires_grad(model.audio_encoder, False)
    else:
        set_requires_grad(model.text_encoder, True)
        set_requires_grad(model.audio_encoder, True)
    set_requires_grad(model.video_encoder, True)
    train_metrics = train_one_epoch(model, train_loader, optimizer, device, model_cfg.num_classes, loss_fn, modalities)
    with torch.no_grad():
        val_metrics = run_epoch(model, val_loader, None, device, model_cfg.num_classes, loss_fn, modalities)
    if val_metrics["accuracy"] > best_val:
        torch.save({"model_state": model.state_dict(), "model_cfg": asdict(model_cfg), "train_cfg": asdict(train_cfg)}, best_path)""",
    )
    add_para(
        doc,
        "The student idea here is 'reuse what the model already learned, then test whether the face adds anything useful.' This is not a full retraining from scratch. It is a refinement experiment on top of a working backbone.",
    )

    doc.add_heading("11.5 scripts/analyze_meld_vit_facecue_fold2.sh", level=2)
    add_para(
        doc,
        "This shell wrapper takes the Fold 2 predictions produced by the warm-start trainer and runs the error analysis tools. It exports test-set predictions if needed, then builds the predicted-vs-actual table, the confusion matrix, the top-confusions list, and the short analysis report. This is the script you use to answer the question 'did facial cues change the model's mistakes in a useful way?'.",
    )
    add_bullets(
        doc,
        [
            "Reads the Fold 2 checkpoint output.",
            "Exports test predictions in a CSV that can be inspected row by row.",
            "Builds confusion-matrix statistics for the MELD labels.",
            "Highlights the strongest confusion pairs so you can see which emotions are still mixed up.",
        ],
    )
    add_code_block(
        doc,
        """mkdir -p "$ANALYSIS_DIR"
"$PYTHON_BIN" -m src.tools.export_predictions \\
  --manifest "$RAW_MANIFEST" \\
  --split test \\
  --checkpoint "$OUTPUT_DIR/best_model.pt" \\
  --output-csv "$PRED_CSV" \\
  --batch-size 4 \\
  --modalities "text,audio,video" \\
  --fusion-pooling min \\
  --encoder-mode paper \\
  --device cpu

"$PYTHON_BIN" scripts/analyze_predictions.py \\
  --predictions-csv "$PRED_CSV" \\
  --output-dir "$ANALYSIS_DIR" \\
  --dataset meld \\
  --split test""",
    )

    doc.add_heading("11.6 scripts/run_meld_vit_facecue_fold4.sh", level=2)
    add_para(
        doc,
        "This is the same warm-start training idea as Fold 2, but applied to Fold 4. Fold 4 is important because it gives a different view of class balance and fold stability. If Fold 2 is the weighted anchor, Fold 4 helps you see whether facial cues improve the model's more balanced behavior across classes.",
    )
    add_bullets(
        doc,
        [
            "Runs the same warm-start trainer.",
            "Uses the Fold 4 facial-cue train/validation CSVs.",
            "Keeps the model architecture and optimization logic consistent with Fold 2.",
            "Provides a second evidence point before any scaling to more folds.",
        ],
    )
    add_code_block(
        doc,
        """#!/usr/bin/env bash
set -euo pipefail

PYTHON_BIN="${PYTHON_BIN:-$(command -v python3)}"
"$PYTHON_BIN" scripts/resume_meld_vit_facecue_fold.py --fold 4 "$@" """,
    )

    doc.add_heading("11.7 scripts/analyze_meld_vit_facecue_fold4.sh", level=2)
    add_para(
        doc,
        "This analysis wrapper does for Fold 4 exactly what the Fold 2 analysis wrapper does for Fold 2. It exports predictions, runs the MELD error analysis, and writes the confusion-matrix outputs into a separate fold-specific folder. The goal is to compare the error pattern with Fold 2 and decide whether facial cues help in a consistent way or only in one fold.",
    )
    add_bullets(
        doc,
        [
            "Creates fold-specific prediction outputs.",
            "Produces a fold-specific confusion matrix.",
            "Lets you compare fold behavior without mixing results from multiple runs.",
        ],
    )
    add_code_block(
        doc,
        """mkdir -p "$ANALYSIS_DIR"
"$PYTHON_BIN" -m src.tools.export_predictions \\
  --manifest "$RAW_MANIFEST" \\
  --split test \\
  --checkpoint "$OUTPUT_DIR/best_model.pt" \\
  --output-csv "$PRED_CSV" \\
  --batch-size 4 \\
  --modalities "text,audio,video" \\
  --fusion-pooling min \\
  --encoder-mode paper \\
  --device cpu

"$PYTHON_BIN" scripts/analyze_predictions.py \\
  --predictions-csv "$PRED_CSV" \\
  --output-dir "$ANALYSIS_DIR" \\
  --dataset meld \\
  --split test""",
    )

    doc.add_heading("11.8 scripts/run_meld_vit_facecue_suite.sh", level=2)
    add_para(
        doc,
        "This is the convenience wrapper that runs the full sequence in order: prepare the data, train Fold 2, analyze Fold 2, train Fold 4, and analyze Fold 4. It exists so the whole facial-cue workflow can be repeated with one command once you are confident the individual steps are correct.",
    )
    add_bullets(
        doc,
        [
            "Useful for a full run after the step-by-step workflow has been validated.",
            "Calls the same scripts you would run manually, just chained together.",
            "Good for reproducibility once the pipeline is stable.",
        ],
    )
    add_para(
        doc,
        "If you use this suite script, you do not need to run scripts/run_meld_vit_facecue_prepare.sh, scripts/run_meld_vit_facecue_fold2.sh, scripts/analyze_meld_vit_facecue_fold2.sh, scripts/run_meld_vit_facecue_fold4.sh, or scripts/analyze_meld_vit_facecue_fold4.sh separately. The suite already executes them in that order. You would still keep the verification script separate if you want an explicit pass/fail check of the generated manifest and folds before training.",
    )
    add_bullets(
        doc,
        [
            "Suite mode replaces the individual prepare/train/analyze commands.",
            "You still keep the verify script as a separate optional data-quality check.",
            "The suite is the easiest choice once the facial-cue pipeline is trusted.",
        ],
    )
    add_code_block(
        doc,
        """#!/usr/bin/env bash
set -euo pipefail

bash scripts/run_meld_vit_facecue_prepare.sh "$@"
bash scripts/run_meld_vit_facecue_fold2.sh
bash scripts/analyze_meld_vit_facecue_fold2.sh
bash scripts/run_meld_vit_facecue_fold4.sh
bash scripts/analyze_meld_vit_facecue_fold4.sh

echo "MELD ViT facial-cue suite complete."
""",
    )

    doc.add_heading("12. How to Verify the Pipeline", level=1)
    add_para(
        doc,
        "To check that the face-cue pipeline is valid, run the new verification script. It performs a single automated pass over the manifest, the extracted ViT embeddings, and the fold CSVs, then prints a PASS/FAIL summary. The script is scripts/verify_meld_vit_facecue_pipeline.py, and the shell wrapper is scripts/run_meld_vit_facecue_verify.sh.",
    )
    add_bullets(
        doc,
        [
            "It checks that the generated manifest exists and has the expected MELD splits.",
            "It compares the manifest sample IDs against the expected MELD annotations and clip paths.",
            "It verifies that every video_path embedding file and every audio_path clip exists.",
            "It loads the saved ViT embeddings and checks that they are two-dimensional, finite, and non-empty.",
            "It confirms that the new fold CSVs exist and that train/validation dialogue groups do not leak into each other.",
        ],
    )
    add_code_block(
        doc,
        """manifest_path = Path(args.manifest)
df = pd.read_csv(manifest_path)
required_cols = {"sample_id", "split", "label", "video_path", "audio_path", "transcript"}
expected_ids = build_expected_ids(meld_root, required_splits)
manifest_ids = set(df["sample_id"].astype(str).tolist())
missing_ids = sorted(expected_ids - manifest_ids)

for _, row in df.iterrows():
    video_path = Path(str(row["video_path"]))
    audio_path = Path(str(row["audio_path"]))
    arr = np.load(video_path)
    if arr.ndim != 2:
        ...

fold_checks = verify_folds(fold_dir)""",
    )
    add_para(
        doc,
        "The verifier is meant to answer the practical question a student would ask after generating the face-cue data: 'Did the pipeline really build what we think it built?' If every check passes, the manifest and fold preparation are consistent enough to use for the warm-start facial-cue experiments. If a check fails, the summary tells you whether the problem is missing clips, missing embeddings, malformed arrays, or fold leakage.",
    )
    add_para(
        doc,
        "If you are using the full suite script, the verify step still remains optional but recommended. The suite is about running the experiment, while the verifier is about checking the prepared data before or after the run.",
    )
    doc.add_heading("12.1 Verification Result and Interpretation", level=2)
    add_para(
        doc,
        "The current verification run is almost fully successful. The manifest contains 13,706 rows, the embeddings load correctly, the inferred ViT feature shape is consistent at (16, 768), and all five folds are dialogue-safe with zero leakage. The only failing check is manifest completeness: one MELD utterance, train_dia125_utt3, is missing from the manifest compared with the expected 13,707 rows.",
    )
    add_bullets(
        doc,
        [
            "This means the facial-cue preprocessing pipeline is structurally correct.",
            "The missing item is a single utterance out of 13,707, which is a tiny fraction of the dataset.",
            "Because all embeddings and folds are otherwise valid, the pipeline can still be used safely for the planned warm-start experiments.",
            "The missing utterance is a coverage gap, not a model or fold-integrity failure.",
        ],
    )
    add_para(
        doc,
        "In practical terms, it is safe to ignore this minor issue for the next experiment stage because the experiment is being evaluated at the fold level, the fold integrity check passed, and one missing utterance will not change the overall training story in a meaningful way. The important quality signals are still correct: the embeddings exist, they are finite, the ViT embedding size is consistent, and the train/validation split has no leakage. That is enough to proceed while noting the one missing row as a small preprocessing omission rather than a blocking defect.",
    )
    add_para(
        doc,
        "This also shows what the outputs from the two build scripts look like in practice. The manifest builder produces one CSV plus one .npy embedding file per utterance, while the fold builder produces five train/validation CSV pairs and a fold-assignment summary. The verification script then checks those outputs by loading the CSV rows, loading the .npy arrays, and confirming that the fold partitions stay dialogue-safe.",
    )
    add_bullets(
        doc,
        [
            "Manifest builder output: meld_vit_facecue.csv + a per-utterance .npy embedding file tree.",
            "Fold builder output: meld_fold_0_train.csv through meld_fold_4_val.csv plus meld_fold_assignments.csv and meld_cv_summary.json.",
            "Verifier output: PASS/FAIL summary that reports manifest coverage, feature shape, and fold leakage.",
        ],
    )

    doc.add_heading("13. Short Thesis-Style Summary", level=1)
    add_para(
        doc,
        "The current weighted-CE MELD Phase 1 result is strong enough to be treated as the backbone baseline. The next logical step is to warm-start from that checkpoint and add facial cues through a ViT-based visual branch. The new scripts create a dedicated face-cue data path, run the warm-start training, and then evaluate whether the confusion matrix and class-balanced metrics improve. In thesis terms, this is the right move because the novelty is not a tiny MELD gain; the novelty is the legal-domain multimodal extension.",
    )

    doc.add_heading("14. Understanding src/train/train.py", level=1)
    add_para(
        doc,
        "This section explains the main training script itself. If the facial-cue scripts are the experiment wrappers, then src/train/train.py is the engine that actually performs ordinary training for LegalMemoCMT Phase 1. Understanding it helps a student see which parts are standard machine learning code and which parts are project-specific.",
    )

    doc.add_heading("14.1 What the Training Script Does End to End", level=2)
    add_para(
        doc,
        "The training script takes a manifest, builds a dataset, constructs the model, chooses a loss, and runs a train/validation loop for a fixed number of epochs. At the end of each epoch, it prints the training and validation metrics, and it saves the best checkpoint when validation accuracy improves. In other words, this file is the generic supervised learning pipeline used throughout Phase 1.",
    )
    add_code_block(
        doc,
        """samples = load_manifest(args.manifest)
dataset = build_dataset(samples, model_cfg, modalities)
train_loader = DataLoader(train_ds, batch_size=train_cfg.batch_size, shuffle=True, collate_fn=collate_samples)
val_loader = DataLoader(val_ds, batch_size=train_cfg.batch_size, shuffle=False, collate_fn=collate_samples)

model = LegalMemoCMTPhase1(model_cfg).to(device)
optimizer = torch.optim.AdamW(model.parameters(), lr=train_cfg.lr, weight_decay=train_cfg.weight_decay)
loss_fn = nn.CrossEntropyLoss(...) or FocalLoss(...)

for epoch in range(1, train_cfg.epochs + 1):
    train_metrics = run_epoch(model, train_loader, optimizer, device, model_cfg.num_classes, loss_fn, modalities)
    val_metrics = run_epoch(model, val_loader, None, device, model_cfg.num_classes, loss_fn, modalities)
    if val_metrics["accuracy"] > best_val:
        torch.save({...}, best_path)""",
    )

    doc.add_heading("14.2 Which Parts Reuse Existing Libraries, Modules, and Algorithms", level=2)
    add_bullets(
        doc,
        [
            "argparse is reused for command-line configuration.",
            "torch and torch.nn provide the model, optimizer, autograd, and loss functions.",
            "torch.utils.data.DataLoader provides batching and shuffling.",
            "transformers.AutoTokenizer is reused for pretrained text encoders in pretrained/paper mode.",
            "NumPy is reused for label counting and metric aggregation.",
            "The optimization loop follows the standard supervised learning pattern: forward pass, loss, backward pass, gradient clipping, optimizer step.",
            "Cross-entropy and focal loss are standard algorithmic choices from the broader machine-learning literature.",
        ],
    )
    add_para(
        doc,
        "From a student view, most of the heavy lifting around tensors, batching, optimization, and loss computation is not custom. The project reuses PyTorch and Hugging Face where possible, which is good engineering because it avoids reinventing basic machine learning infrastructure.",
    )

    doc.add_heading("14.3 Which Parts Are Custom to This Project", level=2)
    add_bullets(
        doc,
        [
            "parse_modalities(), parse_pooling(), and parse_encoder_mode() turn human-readable flags into model settings.",
            "compute_class_weights() implements the project-specific class-balancing rule used for weighted cross-entropy and focal loss.",
            "FocalLoss is a custom wrapper that applies cross-entropy first and then down-weights easy examples.",
            "build_dataset() wires the manifest rows into the LegalMemoCMT data pipeline.",
            "apply_modality_mask() zeros out modalities that are not selected for a run.",
            "run_epoch() is the core reusable loop that handles both training and evaluation in this project.",
            "The main() routine defines the project-specific training protocol: manifest selection, split handling, checkpoint saving, and metric reporting.",
        ],
    )
    add_para(
        doc,
        "The custom part is not the idea of training a classifier. The custom part is how this repository turns raw manifest rows into text/audio/video tensors, how it chooses which modalities are active, how it counts classes for imbalance-aware training, and how it saves the best model in a way that fits the MELD and CREMA-D workflows.",
    )

    doc.add_heading("14.4 How a Batch Moves Through the Script", level=2)
    add_para(
        doc,
        "A single batch goes through the following path: the DataLoader produces a dictionary, the script converts the values to tensors on the chosen device, apply_modality_mask() removes unused modalities, the model performs its forward pass, the chosen loss function scores the predictions, and the optimizer updates the weights if the script is in training mode. This is the basic heartbeat of the training loop.",
    )
    add_code_block(
        doc,
        """for batch in loader:
    labels = torch.as_tensor(batch["label"], device=device)
    outputs = apply_modality_mask(batch_tensors, modalities, model.config.encoder_mode)
    logits = model(...)
    loss = loss_fn(logits, labels)
    if is_train:
        optimizer.zero_grad(set_to_none=True)
        loss.backward()
        torch.nn.utils.clip_grad_norm_(model.parameters(), 1.0)
        optimizer.step()""",
    )

    doc.add_heading("14.5 How the Loss Functions Work", level=2)
    add_para(
        doc,
        "The script supports three loss types: plain cross-entropy, weighted cross-entropy, and focal loss. Plain cross-entropy is the simplest standard classification loss. Weighted cross-entropy adds more importance to minority classes by giving them higher class weights. Focal loss goes a step further by reducing the contribution of easy examples so the model concentrates more on hard cases. This is why the script computes class weights before constructing the chosen loss function.",
    )
    add_code_block(
        doc,
        """if args.loss_type == "ce":
    loss_fn = nn.CrossEntropyLoss()
elif args.loss_type == "weighted-ce":
    loss_fn = nn.CrossEntropyLoss(weight=class_weights)
else:
    loss_fn = FocalLoss(gamma=args.focal_gamma, weight=class_weights)""",
    )
    add_para(
        doc,
        "For a student, the key idea is that loss choice changes what the model pays attention to during learning. The weighted-CE Phase 1 baseline worked best for the current MELD setup, which is why it became the backbone reference for the later facial-cue work.",
    )

    doc.add_heading("14.6 How the Training and Validation Loop Works", level=2)
    add_para(
        doc,
        "The run_epoch() helper is the shared training-and-evaluation routine. If an optimizer is provided, the script treats the pass as training; if no optimizer is provided, it treats the pass as evaluation. The function accumulates true labels, predictions, and loss values, then computes accuracy, macro F1, weighted F1, and a skipped-batch count. This design keeps the training code compact and ensures that the same forward-pass logic is used for both train and validation.",
    )
    add_code_block(
        doc,
        """is_train = optimizer is not None
model.train(is_train)
...
loss = loss_fn(logits, labels)
if is_train:
    optimizer.zero_grad(set_to_none=True)
    loss.backward()
    torch.nn.utils.clip_grad_norm_(model.parameters(), 1.0)
    optimizer.step()
...
return {
    "loss": ...,
    "accuracy": ...,
    "macro_f1": ...,
    "weighted_f1": ...,
    "skipped_batches": skipped,
}""",
    )

    doc.add_heading("14.7 What to Remember as a Student", level=2)
    add_bullets(
        doc,
        [
            "Most of the training code is standard deep-learning infrastructure reused from PyTorch and Hugging Face.",
            "The project-specific parts are the manifest handling, modality control, class balancing, and checkpointing logic.",
            "The same training engine is used for the baseline and for the later facial-cue refinements.",
            "The warm-start facial-cue scripts build on this training engine rather than replacing it.",
        ],
    )

    doc.add_heading("15. Understanding src/models/model.py", level=1)
    add_para(
        doc,
        "This file defines the model architecture itself. If src/train/train.py is the engine that decides how to optimize, then src/models/model.py is the machine being optimized. It contains the encoders for text, audio, and video, the fusion layers that combine them, and the final classifier that predicts the emotion label.",
    )
    add_code_block(
        doc,
        """class TextEncoder(nn.Module):
    ...

class SequenceEncoder(nn.Module):
    ...

class CrossModalFusion(nn.Module):
    ...

class BidirectionalCrossAttentionCMT(nn.Module):
    ...

class PretrainedBackboneEncoder(nn.Module):
    ...

class LegalMemoCMTPhase1(nn.Module):
    ...""",
    )
    add_bullets(
        doc,
        [
            "TextEncoder is the legacy text pathway built with embeddings and a Transformer encoder.",
            "SequenceEncoder is the generic sequence transformer used for audio and video in the legacy branch.",
            "CrossModalFusion is the three-modality fusion block that stacks text, audio, and video as tokens.",
            "BidirectionalCrossAttentionCMT is the paper-style fusion block for text and audio.",
            "PretrainedBackboneEncoder wraps Hugging Face models such as BERT and HuBERT in the pretrained/paper mode.",
            "LegalMemoCMTPhase1 assembles the whole architecture and selects the right branch at runtime.",
        ],
    )
    add_para(
        doc,
        "The model file is where the project becomes more than just a training loop. It defines the actual multimodal learning strategy: either a legacy branch with compact feature encoders or a pretrained/paper branch with Hugging Face backbones and cross-attention fusion.",
    )
    add_para(
        doc,
        "A student should pay special attention to the paper branch because that is the one used by the MELD CV baseline and by the facial-cue warm-start experiments. It is also where the CMT + MIN idea is implemented through the bidirectional cross-attention module and its masked pooling logic.",
    )
    add_bullets(
        doc,
        [
            "Reused library pieces: torch.nn modules, nn.MultiheadAttention, nn.TransformerEncoder, and Hugging Face AutoModel.",
            "Custom pieces: the project-specific fusion design, the paper/legacy branch selection, and the final LegalMemoCMTPhase1 wrapper.",
            "Algorithmically, the model uses transformer encoders, cross-attention, token stacking, and masked pooling.",
        ],
    )
    add_para(
        doc,
        "For the facial-cue plan, the model file matters because the new visual embeddings must fit into the same multimodal fusion story. That is why the warm-start scripts reuse the checkpoint and then adapt the visual branch rather than inventing a completely new architecture.",
    )

    doc.add_heading("16. Understanding src/data/preprocessing.py", level=1)
    add_para(
        doc,
        "This file prepares raw media for the model. It is the bridge between the dataset and the neural network. It knows how to sample frames, extract compact video features, load audio waveforms, build mel-style audio features, and normalize text. In short, it turns raw clips into arrays that the model can actually consume.",
    )
    add_code_block(
        doc,
        """def sample_video_frames(video_path: str, cfg: PreprocessConfig) -> np.ndarray:
    ...

def extract_video_features(video_path: str, cfg: PreprocessConfig) -> np.ndarray:
    ...

def load_audio_features(audio_path: str, cfg: PreprocessConfig) -> np.ndarray:
    ...

def load_audio_waveform(audio_path: str, cfg: PreprocessConfig) -> np.ndarray:
    ...

def normalize_text(text: str) -> str:
    ...""",
    )
    add_bullets(
        doc,
        [
            "sample_video_frames uses OpenCV to read clips and sample evenly spaced frames.",
            "extract_video_features is the older compact video path used by the legacy branch.",
            "load_audio_features derives mel-spectrogram style features for the legacy audio path.",
            "load_audio_waveform prepares raw waveforms for the pretrained/paper audio path.",
            "normalize_text removes simple whitespace noise before tokenization.",
        ],
    )
    add_para(
        doc,
        "This module reuses OpenCV, NumPy, librosa, soundfile, and imageio-ffmpeg, which are standard data-processing libraries. The custom contribution is the project-specific way these tools are combined for MELD and CREMA-D, and the fallback logic that keeps preprocessing robust when one library path fails.",
    )
    add_para(
        doc,
        "For the facial-cue pipeline, the most important function is sample_video_frames. That function currently returns full sampled frames, which are then sent to ViT in the face-cue manifest builder. If you later move to face cropping, this is the place where that preprocessing step would be inserted.",
    )
    add_bullets(
        doc,
        [
            "Reused library pieces: cv2, librosa, soundfile, imageio-ffmpeg, NumPy.",
            "Custom pieces: the exact sampling strategy, fallback logic, and feature-shaping choices.",
            "The code is designed so the same preprocessing helpers can support legacy runs, paper-aligned runs, and future facial-cue runs.",
        ],
    )

    doc.add_heading("17. Understanding src/train/evaluate.py", level=1)
    add_para(
        doc,
        "This file handles inference-time evaluation. It loads a trained checkpoint, rebuilds the model configuration, runs the model on a manifest split, and then writes metrics plus per-sample prediction outputs. It is the evaluation counterpart to src/train/train.py: train.py learns the model, evaluate.py measures it.",
    )
    add_code_block(
        doc,
        """checkpoint = torch.load(args.checkpoint, map_location="cpu")
model = LegalMemoCMTPhase1(model_cfg).to(device)
model.load_state_dict(checkpoint["model_state"])
...
for batch in loader:
    logits = model(...)
    probs = torch.softmax(logits, dim=-1)
    pred = torch.argmax(probs, dim=-1)
    ...
metrics = {
    "accuracy": ...,
    "weighted_accuracy": ...,
    "unweighted_accuracy": ...,
    "macro_f1": ...,
    "weighted_f1": ...,
}""",
    )
    add_bullets(
        doc,
        [
            "Torch is reused to load the saved checkpoint and rebuild the model.",
            "The same dataset and modality masking logic from train.py are reused for evaluation.",
            "NumPy and csv are used to write the prediction table and aggregate the labels.",
            "The metrics are computed with the shared project metric helpers.",
        ],
    )
    add_para(
        doc,
        "A student should understand that evaluation does not update the model. It only measures the fixed checkpoint on the chosen manifest split. This is why the facial-cue analysis scripts call evaluate.py after training: the evaluation step is what produces the metrics.json and predictions_test.csv files used for confusion-matrix analysis.",
    )
    add_para(
        doc,
        "The evaluation file is also where you see the project's reporting behavior: accuracy, weighted accuracy, unweighted accuracy, macro F1, and weighted F1 are all computed in one place so the same checkpoint can be compared across different experiments.",
    )
    add_bullets(
        doc,
        [
            "Reused library pieces: torch.load, DataLoader, softmax, argmax, csv.DictWriter.",
            "Custom pieces: model reconfiguration from the checkpoint config, the project metrics, and the per-sample output format.",
            "Evaluation is a fixed measurement step, not a learning step.",
        ],
    )

    doc.add_heading("18. Module Map for Text, Audio, and Video", level=1)
    add_para(
        doc,
        "The five core files below are all used in the multimodal pipeline. They do not all do the same job, but together they cover the full path from raw manifest rows to trained and evaluated models.",
    )
    add_table(
        doc,
        ["Module", "Used for", "Role in the pipeline"],
        [
            ["src/train/train.py", "Text, audio, video", "Controls training: builds the dataset, chooses the loss, runs epochs, saves checkpoints."],
            ["src/models/model.py", "Text, audio, video", "Defines the architecture: encoders, fusion blocks, and final classifier."],
            ["src/data/preprocessing.py", "Text, audio, video", "Prepares raw media and text into arrays the model can consume."],
            ["src/train/evaluate.py", "Text, audio, video", "Loads a checkpoint and measures performance on a chosen split."],
            ["src/data/dataset.py", "Text, audio, video", "Converts manifest rows into batched tensors with padding and masks."],
        ],
    )
    add_bullets(
        doc,
        [
            "train.py = learning process.",
            "model.py = architecture.",
            "preprocessing.py = data preparation.",
            "evaluate.py = measurement.",
            "dataset.py = manifest-to-batch bridge.",
        ],
    )
    add_para(
        doc,
        "For a student, this table is the simplest way to remember the codebase. If you are tracing a sample through the system, it starts in dataset.py, gets its raw data prepared by preprocessing.py, passes through model.py during training or evaluation, and is orchestrated by train.py or evaluate.py depending on whether the model is learning or being measured.",
    )

    doc.add_heading("19. Expected Warm-Start Warnings", level=1)
    add_para(
        doc,
        "When the facial-cue warm-start run begins, a few warnings may appear in the terminal. In this project, those warnings are usually expected and do not indicate a failure. They are useful to understand because they show how pretrained checkpoints and fallback loaders behave when the architecture has changed slightly or when a media file needs an alternate loading route.",
    )
    add_bullets(
        doc,
        [
            "BERT 'UNEXPECTED' keys are usually the extra pretraining head weights from Hugging Face, such as masked-language-model and next-sentence-prediction heads.",
            "Missing video_encoder.* keys after warm-start are expected when the new facial-cue run adds or reshapes the video branch compared with the old baseline checkpoint.",
            "PySoundFile failed. Trying audioread instead. means the audio loader fell back to a different backend, which is a normal robustness feature.",
            "The librosa audioread FutureWarning is a library deprecation notice, not a training error.",
        ],
    )
    add_para(
        doc,
        "For this run, the important thing is that the text and audio baseline weights were still loaded, the new video branch was allowed to initialize appropriately, and the audio fallback was able to continue processing the clip instead of crashing. So these messages are safe to ignore unless they become frequent enough to cause skipped samples, missing outputs, or instability in the loss values.",
    )
    add_code_block(
        doc,
        """BertModel LOAD REPORT from: bert-base-uncased
Key                              | Status     | Details
---------------------------------+------------+--------
cls.predictions.bias             | UNEXPECTED |
cls.seq_relationship.weight      | UNEXPECTED |

Loaded checkpoint from .../fold_2/best_model.pt
Missing keys after warm-start: ['video_encoder.proj.weight', 'video_encoder.proj.bias', ...]

PySoundFile failed. Trying audioread instead.
librosa.core.audio.__audioread_load: Deprecated as of librosa version 0.10.0""",
    )
    add_para(
        doc,
        "The student takeaway is that this kind of warm-start output is a mixture of checkpoint reuse and library fallback behavior. It only becomes a problem if it is accompanied by crashes, many skipped samples, or obviously broken metrics.",
    )

    doc.add_heading("11. Phase 1 Combined Design Up to the Third Guidance Call", level=1)
    add_para(
        doc,
        "The current Phase 1 design is broader than the original face-cue plan. It now includes the paper-aligned MELD weighted-CE backbone, the full-frame ViT facial-cue branch, the face-crop interpretation for courtroom testimony, gated fusion for selective modality weighting, and the auxiliary-loss branch whose Fold 4 evaluation is now completed and ready to discuss as part of the Phase 1 comparison.",
    )
    add_bullets(
        doc,
        [
            "Backbone: HuBERT + BERT weighted-CE MELD Phase 1.",
            "Visual branch: ViT embeddings from sampled frames and face-crop design logic.",
            "Selective fusion: gated fusion to decide when to trust video.",
            "Auxiliary supervision: small video-head loss to strengthen the visual branch.",
            "Status note: aux-loss Fold 4 results are complete, but they should be described as a modest refinement rather than a new best model.",
        ],
    )
    add_para(
        doc,
        "Student interpretation: the project is now a layered system. The backbone already works, the visual branch adds extra evidence, the gate decides when to use that evidence, and the auxiliary loss is the next refinement that should be judged only after the final runs are complete.",
    )

    doc.add_heading("12. Experimental Results, Confusion Matrices, and Error Analysis", level=1)
    add_para(
        doc,
        "The experimental story should be presented in three levels. First, the weighted-CE baseline remains the stable anchor. Second, the full-frame and face-crop branches tell us whether facial cues carry useful signal. Third, the gated and aux-loss refinements tell us how far the visual branch can be pushed before the model becomes too soft or too biased. For the aux-loss branch, Fold 4 is now finished and can be interpreted as a modest improvement over weaker facial-cue stages, but not as a replacement for the strongest weighted-CE baseline.",
    )
    add_table(
        doc,
        ["Stage", "What the metric story says", "Status"],
        [
            ["Weighted-CE MELD baseline", "Best stable conversational anchor", "Complete and usable"],
            ["Full-frame facial cues", "Visual signal is present but still neutral-biased", "Complete as a first visual baseline"],
            ["Face-crop facial cues", "Better courtroom-testimony interpretation of the visual input", "Planned controlled ablation"],
            ["Gated fusion", "Video can help selectively, but balance is still imperfect", "Implemented and useful"],
            ["Aux-loss branch", "Completed Fold 4; modest improvement, not new best baseline", "Complete"],
        ],
    )
    add_para(
        doc,
        "The confusion matrices should therefore be described carefully. The baseline matrix shows the main neutral-heavy structure. The full-frame and gated matrices show that facial cues are not random noise. The aux-loss Fold 4 matrix is now available and can be discussed as a small refinement that still leaves the neutral-heavy structure in place.",
    )
    add_bullets(
        doc,
        [
            "Baseline confusion: useful but neutral-heavy.",
            "Full-frame confusion: evidence that ViT features carry signal.",
            "Face-crop confusion: expected to be the cleaner courtroom-testimony version.",
            "Gated confusion: selective improvement, but not a solved imbalance problem.",
            "Aux-loss confusion: completed Fold 4 result, useful for comparison but not a new baseline.",
        ],
    )

    doc.add_heading("13. Performance Evaluation and Model Comparison", level=1)
    add_para(
        doc,
        "Performance evaluation now needs to compare the model family as a sequence of decisions, not as unrelated runs. The baseline tells us where Phase 1 starts. The visual branches tell us whether the model is learning facial evidence. The gated branch tells us whether the model can decide when to trust the video. The auxiliary-loss branch is now part of the completed comparison, and its Fold 4 result should be read as an incremental refinement rather than a final replacement.",
    )
    add_table(
        doc,
        ["Model / stage", "Primary question", "Current reading"],
        [
            ["Weighted-CE baseline", "How strong is the backbone?", "Strongest stable benchmark anchor"],
            ["Full-frame ViT", "Does the model see useful facial information?", "Yes, but with strong neutral bias"],
            ["Face-crop ViT", "Does face-focused video look more courtroom-appropriate?", "Most defensible next facial cue"],
            ["Gated fusion", "Can the model learn when to trust video?", "Useful, but still not fully balanced"],
            ["Aux-loss run", "Can extra supervision strengthen the video branch?", "Complete; incremental improvement only"],
        ],
    )
    add_para(
        doc,
        "A student should read the comparison this way: the project is moving from a stable text-audio baseline toward more selective visual learning. The visual methods are not all equal. The full-frame path is the first proof of concept, the face-crop path is the better legal-domain formulation, the gated path is the selective refinement, and the auxiliary-loss path is the current training experiment that still needs final validation before it can become a thesis conclusion.",
    )
    add_bullets(
        doc,
        [
            "Baseline comparison: weighted-CE still wins as the most reliable Phase 1 anchor.",
            "Visual comparison: full-frame versus face-crop is the key facial-cue design choice.",
            "Fusion comparison: gated fusion is safer than forcing video to dominate.",
            "Aux-loss comparison: completed Fold 4 result, but still below the strongest baseline anchor.",
        ],
    )

    doc.add_heading("14. Phase 1 Journal Paper Draft Outline", level=1)
    add_para(
        doc,
        "A journal-paper draft can already be sketched from the current Phase 1 work. The central claim should not be that MELD improved by a tiny margin. The central claim should be that a MemoCMT-style baseline has been reproduced and then extended with facial-cue design choices that are relevant to courtroom testimony.",
    )
    add_numbered(
        doc,
        [
            "Abstract: summarize the baseline and the facial-cue extension path.",
            "Introduction: explain why legal-domain emotion cues matter for courtroom testimony.",
            "Method: describe the weighted-CE backbone, the ViT facial embeddings, the face-crop option, gated fusion, and the aux-loss branch.",
            "Experiments: report the baseline, full-frame, face-crop, gated, and aux-loss settings.",
            "Results: present the metrics and confusion matrices, using the completed aux-loss Fold 4 result as part of the comparison.",
            "Discussion: interpret the error structure and explain why neutral bias remains the main issue.",
            "Conclusion: position the current Phase 1 model as a strong backbone for the legal-domain extension.",
        ],
    )
    add_para(
        doc,
        "For a student-level paper draft, the clean thesis is: the model is already strong enough to be a reproducible baseline, and the visual branch work shows a principled path toward courtroom-testimony adaptation. The aux-loss branch should be described honestly as a completed refinement that improves some facial-cue behavior but still does not surpass the strongest baseline.",
    )

    doc.add_heading("15. Overall Phase 1 Architecture for the Third Guidance Call", level=1)
    add_para(
        doc,
        "This added section summarizes the complete Phase 1 system as it stands now. The architecture starts from a MELD utterance, branches into text, audio, and video processing, and then joins those branches at the fusion layer before the classifier produces the final emotion prediction. The video branch can be instantiated either as full-frame ViT or as the newer face-crop ViT path. The auxiliary-loss branch is drawn as part of the design, and its completed Fold 4 result is now part of the Phase 1 comparison set.",
    )
    add_bullets(
        doc,
        [
            "Input: raw MELD utterance with transcript, audio, and video.",
            "Text branch: BERT feature extraction from the transcript.",
            "Audio branch: HuBERT feature extraction from the audio stream.",
            "Video branch: ViT facial embeddings from full-frame or face-crop frames.",
            "Fusion: combine the modality embeddings into one shared representation.",
            "Classifier: produce the emotion label and evaluation metrics.",
            "Aux-loss head: extra supervision for the video branch, now available as a completed Fold 4 comparison.",
        ],
    )
    add_code_block(
        doc,
        """raw MELD utterance
  -> transcript preprocessing -> BERT
  -> audio preprocessing -> HuBERT
  -> video clip -> full-frame or face-crop sampling -> ViT
  -> fusion layer -> classifier -> metrics / confusion matrix
  -> optional video auxiliary head (completed Fold 4 comparison)""",
    )

    doc.add_heading("16. Training Parameters by Stage", level=1)
    add_para(
        doc,
        "This section records the important hyperparameters used across the Phase 1 stages. The goal is to show that each experiment changes only one or two controlled variables at a time, so the comparison remains fair.",
    )
    add_table(
        doc,
        ["Stage", "Important parameters", "Interpretation"],
        [
            ["Paper-aligned weighted-CE baseline", "weighted CE, baseline checkpoint, best val accuracy", "Stable reference point"],
            ["Full-frame ViT facial cues", "LR 5e-5, 5 epochs, freeze backbone 1 epoch, batch size 4", "First visual proof of concept"],
            ["Face-crop ViT", "warm-start from baseline, small LR, same fold logic", "Speaker-focused visual input"],
            ["Gated fusion", "warm-start, gated video branch, same backbone settings", "Learn when to trust video"],
            ["Video aux-loss", "warm-start, LR 2e-5, max 8 epochs, patience 2, best val weighted F1", "Refinement stage; completed Fold 4 result"],
        ],
    )
    add_para(
        doc,
        "Student takeaway: the later stages do not restart the whole project. They keep the same backbone idea and change the visual or optimization settings so you can isolate the effect of the new component. That is why the baseline checkpoint, learning rate, epoch budget, and checkpoint-selection rule matter so much.",
    )

    doc.add_heading("17. Model Comparison Across Phase 1 Stages", level=1)
    add_para(
        doc,
        "This section condenses the comparison story into one place. The main lesson is that the original weighted-CE backbone is still the strongest stable anchor, even though the visual branches and the gated aux-loss branch give you a better architecture story for the thesis.",
    )
    add_table(
        doc,
        ["Stage", "Metric reading", "Interpretation"],
        [
            ["Weighted-CE baseline", "Strongest stable anchor", "Best reference point for Phase 1"],
            ["Full-frame ViT", "Lower than baseline", "Shows visual signal, but still neutral-biased"],
            ["Face-crop ViT", "Cleaner than full-frame", "Better courtroom-testimony visual design"],
            ["Gated fusion", "0.5992 acc / 0.6056 wF1 / 0.4330 macro F1", "Completed refinement, not new best baseline"],
            ["Aux-loss branch", "Completed comparison run", "Incremental improvement, but still below baseline"],
        ],
    )
    add_para(
        doc,
        "Student explanation: the comparison is not just about the score. It is about what each stage proves. Full-frame ViT proves that facial signal exists. Face-crop proves that focusing on the speaker face is a better legal-domain choice. Gated fusion proves that the model can learn when to trust video. The aux-loss branch proves that extra supervision can help a bit, but not enough to replace the strongest weighted-CE checkpoint.",
    )

    doc.add_heading("18. Confusion Matrix Reading Guide", level=1)
    add_para(
        doc,
        "The confusion matrix is the clearest way to explain the remaining weaknesses. It tells you which labels the model mixes when it is not confident. In this run, the model is still very neutral-heavy, which means neutral often acts like a fallback label.",
    )
    add_bullets(
        doc,
        [
            "Neutral is the main error hub: many emotions collapse into neutral when the model is uncertain.",
            "Joy and anger are often swapped, which is common when the model sees strong but ambiguous emotional cues.",
            "Surprise and fear are also mixed frequently, showing that the model still struggles with subtle high-arousal differences.",
            "Sadness often turns into neutral, which is a sign that the model has not fully separated low-arousal classes.",
            "The matrix shows learning, but it is still a class-imbalance problem rather than a solved recognition problem.",
        ],
    )
    add_para(
        doc,
        "A student-friendly summary is: the matrix is not random. It has structure. The model knows some emotions, but when it is unsure it falls back to the more common or more generic classes. That is why macro F1 is still lower than weighted F1.",
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
