from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "implementation_docments" / "LegalMemoCMT_Phase2_Dataset_Preparation_and_Finetuning_Blueprint.docx"
ASSET_DIR = ROOT / "phase2" / "assets"


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def box(ax, x, y, w, h, text, facecolor="#f3f6fb", edgecolor="#244c74", fontsize=10):
    patch = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.02,rounding_size=0.02",
        linewidth=1.4,
        facecolor=facecolor,
        edgecolor=edgecolor,
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fontsize, family="DejaVu Sans")


def arrow(ax, start, end):
    ax.add_patch(
        FancyArrowPatch(
            start,
            end,
            arrowstyle="->",
            mutation_scale=12,
            linewidth=1.2,
            color="#244c74",
        )
    )


def render_pipeline_assets() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    fig, ax = plt.subplots(figsize=(16, 6.2))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 6.2)
    ax.axis("off")
    box(ax, 0.3, 4.1, 3.0, 1.2, "Public sources\nSCOTUS oral arguments\nIRMCT / ICTR / ICTY archives")
    box(ax, 4.0, 4.1, 2.8, 1.2, "Download and normalize\nPDF / HTML / TXT\nmetadata + transcripts")
    box(ax, 7.3, 4.1, 2.8, 1.2, "Segment and align\ncase, speaker, time window")
    box(ax, 10.6, 4.1, 2.8, 1.2, "Weak supervision\nlexical rules + confidence\nmanual review if needed")
    box(ax, 13.9, 4.1, 1.7, 1.2, "Phase 2\nmanifest")
    box(ax, 2.0, 1.0, 3.2, 1.3, "Text-only legal corpus\nlanguage adaptation\nSCOTUS transcripts")
    box(ax, 6.0, 1.0, 3.4, 1.3, "Multimodal courtroom corpus\ntranscript + audio/video\nwhen records provide it")
    box(ax, 10.0, 1.0, 3.2, 1.3, "Warm-start from best MELD fold\nload encoders + fusion\nreset classifier head")
    box(ax, 13.7, 1.0, 1.9, 1.3, "Fine-tune\n+ evaluate")
    arrow(ax, (3.3, 4.7), (4.0, 4.7))
    arrow(ax, (6.8, 4.7), (7.3, 4.7))
    arrow(ax, (10.1, 4.7), (10.6, 4.7))
    arrow(ax, (13.4, 4.7), (13.9, 4.7))
    arrow(ax, (14.75, 4.1), (14.75, 2.3))
    arrow(ax, (5.2, 2.3), (10.0, 2.3))
    arrow(ax, (13.2, 2.3), (13.7, 2.3))
    fig.suptitle("Phase 2 data preparation and warm-start fine-tuning pipeline", fontsize=15, family="DejaVu Sans")
    fig.tight_layout()
    fig.savefig(ASSET_DIR / "phase2_pipeline.png", dpi=220, bbox_inches="tight")
    fig.savefig(ASSET_DIR / "phase2_pipeline.svg", bbox_inches="tight")
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(15, 4.8))
    ax.set_xlim(0, 15)
    ax.set_ylim(0, 4.8)
    ax.axis("off")
    box(ax, 0.5, 1.8, 2.3, 1.1, "Best MELD fold\ncheckpoint")
    box(ax, 3.4, 1.8, 2.5, 1.1, "Load backbone\nweights")
    box(ax, 6.5, 1.8, 2.7, 1.1, "Discard or replace\nclassifier head")
    box(ax, 9.9, 1.8, 2.6, 1.1, "Train on Phase 2\nmanifest")
    box(ax, 13.1, 1.8, 1.3, 1.1, "Test\nanalysis")
    arrow(ax, (2.8, 2.35), (3.4, 2.35))
    arrow(ax, (5.9, 2.35), (6.5, 2.35))
    arrow(ax, (9.2, 2.35), (9.9, 2.35))
    arrow(ax, (12.5, 2.35), (13.1, 2.35))
    fig.suptitle("Phase 2 checkpoint warm-start path", fontsize=15, family="DejaVu Sans")
    fig.tight_layout()
    fig.savefig(ASSET_DIR / "phase2_finetune_path.png", dpi=220, bbox_inches="tight")
    fig.savefig(ASSET_DIR / "phase2_finetune_path.svg", bbox_inches="tight")
    plt.close(fig)

    (ASSET_DIR / "phase2_pipeline.mmd").write_text(
        """flowchart LR
    A[Public sources\\nSCOTUS oral arguments\\nIRMCT / ICTR / ICTY archives] --> B[Download and normalize\\nPDF / HTML / TXT]
    B --> C[Segment and align\\ncase, speaker, time window]
    C --> D[Weak supervision\\nlexical rules + confidence]
    D --> E[Phase 2 manifest]
    A --> F[Text-only legal corpus\\nSCOTUS transcripts]
    D --> G[Multimodal courtroom corpus\\ntranscript + audio/video]
    G --> H[Warm-start from best MELD fold\\nload encoders + fusion]
    H --> I[Fine-tune + evaluate]
""",
        encoding="utf-8",
    )
    (ASSET_DIR / "phase2_finetune_path.mmd").write_text(
        """flowchart LR
    A[Best MELD fold checkpoint] --> B[Load backbone weights]
    B --> C[Discard or replace classifier head]
    C --> D[Train on Phase 2 manifest]
    D --> E[Test analysis]
""",
        encoding="utf-8",
    )


def build_document() -> Document:
    render_pipeline_assets()
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("LegalMemoCMT Phase 2 Dataset Preparation and Fine-Tuning Blueprint")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(
        "Student-focused technical guide for building the courtroom-testimony adaptation pipeline, from public data collection to weak supervision and MELD warm-start fine-tuning."
    )
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12.5)

    add_para(
        doc,
        "This document explains what Phase 2 is trying to accomplish, how the public legal sources should be organized, how the manifest should be structured, how weak supervision can be used when labels are not available, and how the best MELD checkpoint can be used as the starting point for adaptation. The emphasis is on a practical student-level plan rather than a claim that the final courtroom dataset is already complete.",
    )
    add_para(
        doc,
        "The key design choice is to separate two different kinds of legal data: text-only legal language corpora, which are useful for language adaptation and weak-label design, and multimodal courtroom or testimony records, which are the better fit for the current emotion-analysis model when audio or video is available. That separation keeps the project technically honest and easier to debug.",
    )

    doc.add_heading("1. What Phase 2 Is Trying To Do", level=1)
    add_bullets(
        doc,
        [
            "Adapt the Phase 1 LegalMemoCMT model to courtroom and judicial-record settings.",
            "Keep the output limited to observable emotion cues, not guilt, innocence, deception, or legal judgment.",
            "Use public transcripts and public judicial records to build a legally relevant dataset pipeline.",
            "Create a data path that is realistic even when some records are text-only and other records provide audio or video.",
            "Warm-start from the best MELD fold checkpoint instead of training a new model from scratch.",
        ],
    )

    doc.add_heading("2. Why The Eyewitness Incongruence Paper Matters", level=1)
    add_para(
        doc,
        "The eyewitness incongruence paper is useful as a structuring reference, not as the main emotion dataset. Its value is methodological: it shows how to organize testimony, pair evidence, reason over context windows, and think about inconsistency signals. That is helpful for Phase 2 because courtroom testimony is also context-sensitive and often requires segment-level interpretation rather than whole-document labeling.",
    )
    add_bullets(
        doc,
        [
            "Use it to think about segmentation and witness-turn structure.",
            "Use it to think about how context affects interpretation.",
            "Use it to motivate weak supervision when gold labels are scarce.",
            "Do not treat it as the final emotion benchmark for this project.",
        ],
    )

    doc.add_heading("3. Recommended Data Strategy", level=1)
    add_table(
        doc,
        ["Source", "Primary role", "Why it is useful"],
        [
            ["Supreme Court oral argument transcripts", "Text corpus for legal-language adaptation", "Strong public source, stable legal language, useful for transcript normalization and weak supervision design."],
            ["IRMCT / ICTR / ICTY public judicial records", "Main courtroom/testimony source", "Closer to real judicial evidence and more suitable for emotional pattern analysis in proceedings."],
            ["Eyewitness incongruence paper", "Structuring reference", "Useful for how to organize testimony and context, but not the main data source."],
        ],
    )
    add_para(
        doc,
        "A practical Phase 2 path is therefore two-stage: first, collect and normalize the legal text and public tribunal records; second, build a labeled training manifest only from the records that actually support the intended multimodal training mode. That is why the Phase 2 scripts separate source download, weak labeling, and training manifest creation.",
    )

    doc.add_heading("4. Manifest Design", level=1)
    add_para(
        doc,
        "A manifest is the bridge between raw public documents and the training loop. It should record where each sample came from, what modality files exist, what the model should read, what label was assigned, and how confident that label is. The manifest is the file you can inspect when you need to explain exactly what a row means.",
    )
    add_table(
        doc,
        ["Column", "Suggested content", "Required?"],
        [
            ["sample_id", "Unique string for each record or segment", "Yes"],
            ["case_id", "Case or proceeding identifier used for split grouping", "Yes"],
            ["court", "SCOTUS, IRMCT, ICTR, ICTY, or related archive name", "Yes"],
            ["source_type", "oral_argument_transcript, judicial_record, testimony, etc.", "Yes"],
            ["language", "en, hi, ta, te, kn, or mixed", "Recommended"],
            ["split", "train, dev, test", "Yes for training"],
            ["transcript", "Normalized text of the segment", "Yes"],
            ["audio_path", "Path to audio file or waveform", "When available"],
            ["video_path", "Path to video clip or frame sequence", "When available"],
            ["label", "Integer emotion label", "Yes for supervised rows"],
            ["label_text", "Human-readable label name", "Recommended"],
            ["confidence", "Weak-label confidence from 0 to 1", "Recommended"],
            ["weak_rule", "Which rule or heuristic assigned the label", "Recommended"],
            ["notes", "Curation notes or caveats", "Recommended"],
        ],
    )
    add_para(
        doc,
        "The current Phase 1 training code expects at least sample_id, split, label, video_path, audio_path, and transcript. The Phase 2 manifest can contain more columns than that, but the training subset should still preserve the columns needed by the loader. That is why the Phase 2 builder produces separate labeled multimodal and text-only corpora.",
    )

    doc.add_heading("5. Suggested Split Strategy", level=1)
    add_para(
        doc,
        "The split strategy should be case-aware, not just row-aware. If two rows come from the same case or the same witness, they should not leak across train and test. That would make the metrics look better than they really are because the model would see too much of the same context during training.",
    )
    add_bullets(
        doc,
        [
            "Split by case_id first, then by segment or row only inside each case.",
            "If the dataset is small, prefer 70/15/15 or a case-level cross-validation scheme.",
            "Keep the dev set stable so checkpoint selection is reproducible.",
            "If transcripts are multilingual, try to distribute languages across splits instead of clustering one language in a single split.",
            "For public legal records, keep source_type balanced where possible.",
        ],
    )
    add_para(
        doc,
        "This is especially important in courtroom data because emotional tone, legal topic, and speaker identity are all strongly correlated within the same case. A case-level split keeps the benchmark closer to a real generalization test.",
    )

    doc.add_heading("6. Weak Supervision Rules", level=1)
    add_para(
        doc,
        "Weak supervision is the practical bridge between public records and an emotion model when gold labels do not exist. The idea is to use transparent heuristics that are not perfect but are explainable. Each rule contributes a tentative label and a confidence score.",
    )
    add_numbered(
        doc,
        [
            "Lexical cues: words such as scared, anxious, tense, angry, stressed, worried, and their multilingual or transliterated variants raise the score for the corresponding emotion.",
            "Uncertainty cues: phrases such as I don't recall, maybe, perhaps, not sure, unclear, or cannot remember push the sample toward uncertainty or anxiety.",
            "Stress cues: repeated references to pressure, exhaustion, inability, or difficulty push the sample toward stress.",
            "Anger cues: objections, hostile wording, or direct expressions of frustration push the sample toward anger.",
            "Fear cues: explicit fear or alarm wording pushes the sample toward fear.",
            "Fallback behavior: if the confidence is low, mark the row as uncertain rather than pretending the label is reliable.",
        ],
    )
    add_para(
        doc,
        "This is not the final scientific truth. It is a controlled approximation that lets the student build a training pipeline, inspect the behavior, and later replace the weak labels with stronger annotation if needed.",
    )

    doc.add_heading("7. How The Public Data Scripts Fit Together", level=1)
    add_table(
        doc,
        ["Script", "What it does", "Status in the workflow"],
        [
            ["phase2/download_supreme_court_transcripts.py", "Downloads and extracts transcript text from the official SCOTUS oral argument transcript page.", "Mandatory for the legal-text corpus"],
            ["phase2/download_public_judicial_records.py", "Downloads the public tribunal records listed in a source CSV.", "Mandatory for the courtroom corpus"],
            ["phase2/build_weak_labels.py", "Assigns provisional emotion labels using transparent text rules.", "Mandatory until gold labels exist"],
            ["phase2/build_phase2_manifest.py", "Merges source indices, weak labels, and split assignments into labeled and unlabeled manifests.", "Mandatory"],
            ["phase2/train_phase2_from_checkpoint.py", "Warm-starts from the best MELD fold checkpoint and fine-tunes on the Phase 2 manifest.", "Mandatory for the first adaptation run"],
            ["phase2/evaluate_phase2_checkpoint.sh", "Runs the saved checkpoint through the evaluation script and saves metrics.", "Mandatory after training"],
        ],
    )
    add_para(
        doc,
        "The important design idea is that download, labeling, manifest creation, and training are separate. That separation makes the project easier to debug because each stage has a single purpose and a visible output file.",
    )

    doc.add_heading("8. Pipeline Diagram", level=1)
    doc.add_picture(str(ASSET_DIR / "phase2_pipeline.png"), width=Inches(6.5))
    cap = doc.add_paragraph("Figure 1. Phase 2 data preparation and warm-start fine-tuning pipeline.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        doc,
        "The diagram shows two parallel corpus paths. The text-only legal corpus helps with legal-domain adaptation, while the labeled multimodal courtroom corpus is what should drive the first Phase 2 emotion fine-tuning run. Both are organized by the same source-then-manifest logic.",
    )

    doc.add_heading("9. Fine-Tuning From The Best MELD Checkpoint", level=1)
    add_para(
        doc,
        "The recommended Phase 2 starting point is not a fresh model. It is the best MELD checkpoint from Phase 1, because it already knows how to represent conversational emotion using the current multimodal architecture. Fine-tuning starts from that representation and shifts it toward legal testimony.",
    )
    add_numbered(
        doc,
        [
            "Load the best MELD fold checkpoint.",
            "Reuse the pretrained text and audio backbones and the cross-modal fusion block.",
            "Reset the classifier head if the Phase 2 label space differs from MELD.",
            "Freeze the backbones at first if the Phase 2 dataset is small.",
            "Train the new head and fusion layers on the labeled courtroom manifest.",
            "Evaluate on the dev split and save the best checkpoint.",
            "Run the final test evaluation and confusion analysis.",
        ],
    )
    add_para(
        doc,
        "This is the safest adaptation strategy because it avoids throwing away what the model already learned on MELD. The MELD checkpoint is the learned prior; the Phase 2 data then adjusts the model toward courtroom language and courtroom emotion patterns.",
    )

    doc.add_heading("10. Fine-Tuning Diagram", level=1)
    doc.add_picture(str(ASSET_DIR / "phase2_finetune_path.png"), width=Inches(6.3))
    cap = doc.add_paragraph("Figure 2. Warm-start path from the best MELD fold checkpoint into Phase 2.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("11. What Is Mandatory And What Is Optional", level=1)
    add_table(
        doc,
        ["Item", "Mandatory or optional", "Reason"],
        [
            ["Build a labeled multimodal courtroom manifest", "Mandatory", "The model needs a supervised dataset with the correct columns."],
            ["Keep a separate text-only legal corpus", "Optional but recommended", "Useful for language adaptation and future text-only experiments."],
            ["Use video in the first Phase 2 run", "Optional", "Many public records will not provide stable video, so do not block the project on it."],
            ["Use the eyewitness paper as the structuring reference", "Mandatory as a concept", "It shapes the testimony organization idea, but not as a direct training source."],
            ["Start from the best MELD checkpoint", "Mandatory", "This is the cleanest and safest adaptation path."],
            ["Run full manual annotation immediately", "Optional", "Weak supervision is enough to start the technical pipeline."],
        ],
    )

    doc.add_heading("12. Practical Output You Should Expect", level=1)
    add_bullets(
        doc,
        [
            "A prepared legal dataset directory with raw downloads and extracted text.",
            "A tribunal source template that can be filled from public archive URLs.",
            "A labeled Phase 2 manifest for multimodal fine-tuning.",
            "A weak-label CSV that explains which rule produced each label.",
            "A checkpoint warm-start run starting from the best MELD fold.",
            "An evaluation JSON file and prediction CSV for the Phase 2 model.",
        ],
    )

    doc.add_heading("13. Known Limitations And Honest Scope Boundaries", level=1)
    add_para(
        doc,
        "The Phase 2 plan is technically sound, but it should be described honestly. Public legal records are not the same as a curated emotion dataset. Some records will be text-only, some will lack clean audio, and some will need manual source curation. That is normal for a legal-domain research project.",
    )
    add_bullets(
        doc,
        [
            "Do not claim guilt, innocence, or deception detection.",
            "Do not claim courtroom emotion labels are gold unless they have been annotated.",
            "Do not force every source into the same modality shape.",
            "Do not treat weak supervision as equivalent to human annotation.",
        ],
    )

    doc.add_heading("14. Student Checklist", level=1)
    add_numbered(
        doc,
        [
            "Read the Phase 2 blueprint before changing code.",
            "Download the public SCOTUS transcript corpus.",
            "Create the tribunal source CSV from public archive URLs.",
            "Download and index the tribunal records.",
            "Generate weak labels and inspect a small sample manually.",
            "Build the labeled multimodal manifest and keep the text-only corpus separate.",
            "Fine-tune from the best MELD fold checkpoint.",
            "Evaluate the saved Phase 2 checkpoint on held-out data.",
        ],
    )

    doc.add_heading("15. References", level=1)
    add_bullets(
        doc,
        [
            "U.S. Supreme Court oral argument transcript availability page: official public transcript access point.",
            "IRMCT archives and the public UCR access portal for ICTR/ICTY/IRMCT judicial records.",
            "Incongruence Identification in Eyewitness Testimony, arXiv 2025, as the structural reference for testimony organization.",
        ],
    )
    add_para(
        doc,
        "The main idea is simple: use public legal text to build the language and structure layer, use public tribunal records to build the courtroom emotion layer, and start from the best MELD checkpoint so the model does not have to learn multimodal emotion recognition from zero. That is the most realistic student-level Phase 2 path.",
    )

    return doc


def main() -> None:
    doc = build_document()
    DOC_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOC_PATH))
    print(f"Wrote {DOC_PATH}")


if __name__ == "__main__":
    main()

