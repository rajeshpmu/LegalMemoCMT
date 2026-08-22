from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION


ROOT = Path(__file__).resolve().parents[1]
SOP = ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx"


document = Document(SOP)
document.add_section(WD_SECTION.NEW_PAGE)
document.add_heading("Diarization Dependency Access Result", level=1)
document.add_paragraph(
    "The Hugging Face identity check passed for the configured account, and repository access to "
    "`pyannote/speaker-diarization-3.1` also passed. The pipeline then attempted to download its gated "
    "dependency `pyannote/segmentation-3.0` and failed. Therefore the token is recognized, but the model "
    "terms or access permission for the segmentation dependency are not yet available to this token."
)
document.add_paragraph(
    "The intermediate `TypeError: from_pretrained() got an unexpected keyword argument 'token'` is expected "
    "with pyannote.audio 3.4.0. The code catches it and retries with pyannote's supported `use_auth_token` "
    "argument. The final `NoneType` error is a secondary consequence of the missing segmentation model, not "
    "a new Python-version problem."
)
document.add_heading("Required Browser Action", level=2)
document.add_paragraph(
    "While signed into the Hugging Face account shown by the identity check, open the model page below and "
    "accept its conditions if prompted:"
)
document.add_paragraph("https://huggingface.co/pyannote/segmentation-3.0")
document.add_paragraph(
    "Also confirm that the conditions for `pyannote/speaker-diarization-3.1` and any linked gated speaker "
    "embedding dependency are accepted. The token must have read access."
)
document.add_heading("Retry Sequence", level=2)
for command in [
    "./.venv-diarization/bin/python phase2/check_clancy_diarization_prerequisites.py --load-model",
    "./.venv-diarization/bin/python -m pip check",
]:
    paragraph = document.add_paragraph(style="No Spacing")
    run = paragraph.add_run(command)
    run.font.name = "Courier New"
    run.font.size = 8
document.add_paragraph(
    "Proceed to the diarization pilot only when the checker prints both `model_repo_access=PASS` and "
    "`model_access=PASS`. The checker now converts dependency-loading failures into a concise diagnostic "
    "instead of exposing the secondary traceback."
)
document.save(SOP)
print(f"Updated {SOP}")
