from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "implementation_docments" / "LegalMemoCMT_Phase1_End_to_End_Code_Walkthrough.docx"
FIG_DIR = ROOT / "implementation_docments" / "figures" / "phase1_end_to_end_code_walkthrough"
FIG_DIR.mkdir(parents=True, exist_ok=True)
PIPELINE_MMD = FIG_DIR / "phase1_end_to_end_pipeline.mmd"
PIPELINE_SVG = FIG_DIR / "phase1_end_to_end_pipeline.svg"
PIPELINE_PNG = FIG_DIR / "phase1_end_to_end_pipeline.png"


def render_mermaid(code: str) -> None:
    PIPELINE_MMD.write_text(code, encoding="utf-8")
    try:
        subprocess.run(
            ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(PIPELINE_MMD), "-o", str(PIPELINE_SVG), "-b", "white"],
            check=True,
            cwd=str(ROOT),
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        subprocess.run(
            ["npx", "-y", "@mermaid-js/mermaid-cli", "-i", str(PIPELINE_MMD), "-o", str(PIPELINE_PNG), "-b", "white"],
            check=True,
            cwd=str(ROOT),
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
    except Exception:
        # The document can still be generated without the illustration if Mermaid is unavailable.
        pass


def configure(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Times New Roman"
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code(doc: Document, code: str) -> None:
    for line in code.rstrip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(9.0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def read(path: str) -> str:
    return (ROOT / path).read_text(encoding="utf-8")


def snippet(lines: list[str], start: int, end: int) -> str:
    return "\n".join(lines[start - 1 : end])


def build_doc() -> Document:
    render_mermaid(
        """flowchart LR
  A[Raw MELD annotation CSVs] --> B[Manifest builders]
  B --> C[Preprocessing: sample frames / load audio / normalize text]
  C --> D[Dataset + collate]
  D --> E[Model: encoders + fusion + classifier]
  E --> F[Train loop]
  F --> G[Best checkpoint]
  G --> H[Evaluate + confusion matrix]
  H --> I[Raw-MP4 demo]
""",
    )

    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LegalMemoCMT Phase 1 End-to-End Code Walkthrough")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "A code-level explanation of how the Phase 1 pipeline moves from raw MELD files to manifests, feature tensors, training, evaluation, and raw-video demo outputs."
    )
    r.italic = True
    r.font.size = Pt(13)
    r.font.name = "Times New Roman"

    add_para(
        doc,
        "This document is written as a technical reading guide. It is meant to help a reviewer trace a sample through the repository and understand which files are standard PyTorch / Hugging Face components and which files are custom to LegalMemoCMT Phase 1.",
    )

    doc.add_heading("1. The Whole Pipeline in One Sentence", level=1)
    add_para(
        doc,
        "Phase 1 takes a MELD utterance from the raw dataset, turns it into manifest rows, converts the raw media into model-ready features, feeds those features through the model, trains or evaluates the checkpoint, and then exports metrics, confusion matrices, and demo outputs that can be inspected sample by sample.",
    )
    add_table(
        doc,
        ["Stage", "Main files", "What the stage does"],
        [
            ["Raw data and manifests", "scripts/build_meld_manifest.py, scripts/build_meld_vit_facecrop_manifest.py, scripts/build_meld_cv_folds.py", "Turns MELD annotations into rows that point to text, audio, and visual feature paths."],
            ["Preprocessing", "src/data/preprocessing.py", "Samples frames, extracts audio, normalizes text, and prepares feature arrays."],
            ["Dataset", "src/data/dataset.py", "Converts a manifest row into padded tensors and masks."],
            ["Model", "src/models/model.py", "Encodes text, audio, and video and fuses them before classification."],
            ["Training", "src/train/train.py", "Runs epochs, applies loss, and saves the best checkpoint."],
            ["Evaluation", "src/train/evaluate.py", "Loads a checkpoint and writes metrics / prediction CSVs."],
            ["Demo", "scripts/run_phase1_raw_mp4_demo.sh, scripts/predict_phase1_raw_mp4_demo.py", "Processes a single raw mp4 end-to-end and prints probabilities."],
        ],
    )
    if PIPELINE_PNG.exists():
        doc.add_picture(str(PIPELINE_PNG), width=Inches(6.7))

    doc.add_heading("2. Raw MELD Input and Manifest Construction", level=1)
    add_para(
        doc,
        "The raw MELD annotations are not used directly by the training loop. Instead, the repository builds manifests that act like structured indexes. This is a very important software design choice: every sample gets an ID, a split, a label, a transcript, and one or more feature paths, so the rest of the pipeline no longer has to search through directories by hand.",
    )
    add_bullets(
        doc,
        [
            "scripts/build_meld_manifest.py creates a basic MELD manifest from the annotation CSVs.",
            "scripts/build_meld_vit_facecrop_manifest.py creates the visual feature manifest for the face-crop path.",
            "scripts/build_meld_cv_folds.py and scripts/build_meld_vit_facecrop_control_folds.py create the dialogue-safe train/val folds.",
            "These scripts preserve sample identity so the same utterance stays aligned across modalities and evaluation steps.",
        ],
    )
    add_code(
        doc,
        snippet(read("scripts/build_meld_vit_facecrop_manifest.py").splitlines(), 1, 120),
    )
    add_para(
        doc,
        "The key idea in the manifest builder is that the visual pipeline is not a generic image classifier. It is a per-utterance feature extraction step. For each utterance, the script locates the raw MP4 clip, samples a fixed number of frames, crops the face in each frame, runs a pretrained ViT, and saves the CLS embedding sequence to .npy. The manifest then stores the path to that .npy file.",
    )
    add_para(
        doc,
        "If a face is not detected, the code falls back to a center crop rather than crashing. That makes the pipeline robust enough for batch processing, although the fallback may weaken the quality of the visual cue.",
    )

    doc.add_heading("3. src/data/preprocessing.py", level=1)
    add_para(
        doc,
        "This module is the media-preprocessing layer. It is responsible for turning raw media files into numeric arrays the model can consume. Most of the file is standard data engineering logic, but it is crucial because the model cannot train on raw MP4 or raw WAV directly.",
    )
    add_code(
        doc,
        snippet(read("src/data/preprocessing.py").splitlines(), 1, 220),
    )
    add_bullets(
        doc,
        [
            "sample_video_frames(...) extracts a fixed number of RGB frames at evenly spaced indices.",
            "extract_video_features(...) is the legacy compact video path that turns frames into 128-dim vectors.",
            "load_audio_features(...) converts waveform audio into log-mel features when needed.",
            "load_audio_waveform(...) prepares raw waveform input for pretrained speech encoders such as HuBERT.",
            "normalize_text(...) removes whitespace noise so the transcript becomes consistent before tokenization.",
        ],
    )
    add_para(
        doc,
        "From a code perspective, the important thing is that preprocessing creates stable shapes. The model wants a sequence, not a variable-size video file. That is why the functions pad or clip to fixed lengths and replace invalid values with zeros when necessary.",
    )

    doc.add_heading("4. src/data/dataset.py", level=1)
    add_para(
        doc,
        "This module is the bridge between the manifest and the batch tensor. It takes one manifest row and converts it into the actual arrays returned by DataLoader. In plain terms, it translates a CSV record into the model's input dictionary.",
    )
    add_code(
        doc,
        snippet(read("src/data/dataset.py").splitlines(), 1, 260),
    )
    add_table(
        doc,
        ["Concept", "Legacy mode", "Pretrained / paper mode"],
        [
            ["Text", "Integer token sequence built by a lightweight tokenizer", "BERT input_ids and attention_mask"],
            ["Audio", "Feature arrays from log-mel preprocessing", "Raw waveform + attention mask for HuBERT-style encoders"],
            ["Video", "Feature arrays or cached embeddings", "Face/video embedding sequence and mask"],
        ],
    )
    add_para(
        doc,
        "The dataset class supports both legacy and pretrained encoder modes. That means the same manifest can feed either the older compact path or the paper-aligned Hugging Face encoder path. The loader also creates masks, which are important because the model must know which timesteps are real data and which are padding.",
    )
    add_para(
        doc,
        "The collate function pads all examples in a batch to the same length. This is standard PyTorch practice for variable-length multimodal data, and it is one of the places where the code is doing real engineering work rather than just calling a library function.",
    )

    doc.add_heading("5. src/models/model.py", level=1)
    add_para(
        doc,
        "This is the architecture file. It defines how the text, audio, and video branches are encoded and then fused into a single emotion prediction. If preprocessing is the data-engineering layer, model.py is the learning layer.",
    )
    add_code(
        doc,
        snippet(read("src/models/model.py").splitlines(), 1, 420),
    )
    add_bullets(
        doc,
        [
            "SequenceEncoder uses a linear projection, positional encoding, TransformerEncoder layers, and masked pooling.",
            "TextEncoder is the same idea, but it starts from token embeddings.",
            "BidirectionalCrossAttentionCMT is the paper-style text/audio cross-attention module.",
            "CrossModalFusion is the older modality-stacking fusion with CLS/MEAN/MAX/MIN pooling options.",
            "GatedModalFusion learns a soft weight for each modality and combines them conditionally.",
            "LegalMemoCMTPhase1 combines these building blocks into one classifier with an optional video auxiliary head.",
        ],
    )
    add_para(
        doc,
        "The important architectural decision is that the model keeps multiple fusion styles available. That allows the experiments to compare the paper-aligned CMT behavior, the gated fusion behavior, and the later auxiliary-loss behavior without rewriting the whole architecture each time.",
    )
    add_para(
        doc,
        "In the forward pass, the model first decides which modalities are present. Then it encodes text, audio, and video. Then it applies either the paper fusion or the gated fusion, and finally it passes the fused vector to the classifier. If the run uses only video, the video encoder can still produce a stand-alone prediction path.",
    )
    add_para(
        doc,
        "If the reviewer asks what is custom here, the answer is: the modality gating logic, the video auxiliary classifier, the flexible encoder-mode switching, and the support for different pooling strategies are all project-specific engineering decisions layered on top of standard PyTorch transformer building blocks.",
    )

    doc.add_heading("6. src/train/train.py", level=1)
    add_para(
        doc,
        "This is the training engine. It loads the manifest, builds the dataset, sets the model configuration, picks the loss function, loops through epochs, evaluates on validation data, and saves the best checkpoint. It is the part of the repository that actually learns the parameters.",
    )
    add_code(
        doc,
        snippet(read("src/train/train.py").splitlines(), 1, 340),
    )
    add_bullets(
        doc,
        [
            "parse_modalities(...) lets the run script switch between text, audio, and video ablations.",
            "parse_pooling(...) chooses CLS, mean, max, or min pooling.",
            "parse_fusion_mode(...) switches between legacy and gated fusion.",
            "FocalLoss is the custom imbalance-aware loss implementation.",
            "compute_class_weights(...) uses the training subset label counts to make weighted CE or weighted focal loss.",
            "run_epoch(...) handles both training and validation and computes accuracy, macro F1, and weighted F1.",
        ],
    )
    add_para(
        doc,
        "A student should notice that the validation checkpoint selection logic is simple but important. The script compares the validation metric and saves the best state_dict. That means the final evaluation is based on a checkpoint chosen by a validation criterion, not on the last epoch blindly.",
    )
    add_para(
        doc,
        "The training file also shows the difference between standard cross-entropy, class-weighted cross-entropy, and focal loss. That is where the class-imbalance experiments in Phase 1 originate.",
    )

    doc.add_heading("7. src/train/evaluate.py", level=1)
    add_para(
        doc,
        "The evaluation script is the read-only counterpart to train.py. It loads a checkpoint, runs inference on a chosen manifest split, and writes metrics and prediction rows. It does not update the model.",
    )
    add_code(
        doc,
        snippet(read("src/train/evaluate.py").splitlines(), 1, 220),
    )
    add_bullets(
        doc,
        [
            "The script reloads model_cfg from the checkpoint so evaluation matches training.",
            "It applies the same modality masking logic used during training.",
            "It computes accuracy, weighted accuracy, unweighted accuracy, macro F1, and weighted F1.",
            "It exports per-sample predictions so you can later build a confusion matrix and top-confusion table.",
        ],
    )
    add_para(
        doc,
        "This is the script that turns a checkpoint into numbers you can show in a report. The confusion matrix is especially important because it tells you not just whether the model is right overall, but which emotion pairs are being confused.",
    )

    doc.add_heading("8. Raw-MP4 Demo Path", level=1)
    add_para(
        doc,
        "The demo path is the same pipeline in single-sample form. Instead of using a manifest row that already points to precomputed features, the raw-mp4 demo starts from an actual video file, extracts the features first, and then feeds those features through the trained model.",
    )
    add_code(
        doc,
        snippet(read("scripts/run_phase1_raw_mp4_demo.sh").splitlines(), 1, 120),
    )
    add_para(
        doc,
        "The bash wrapper checks the manifest, checkpoint, and raw video path, then calls the Python inference script with the chosen modality and device settings. This separation between wrapper script and Python logic is deliberate: the wrapper makes the command easy to run, while the Python script contains the actual inference implementation.",
    )
    add_code(
        doc,
        snippet(read("scripts/predict_phase1_raw_mp4_demo.py").splitlines(), 1, 220),
    )
    add_para(
        doc,
        "When a live demo runs, the output you explain is not just the final label. You should also explain the confidence score and the top-3 probabilities. Those probabilities show whether the model is strongly confident, borderline, or uncertain about the class boundary.",
    )

    doc.add_heading("9. How a Single Sample Moves Through the System", level=1)
    add_numbered(
        doc,
        [
            "A MELD row is read from the manifest and matched to one sample_id.",
            "The dataset loads the transcript, audio, and video paths for that sample.",
            "Preprocessing turns the raw media into token IDs, waveform tensors, or visual features.",
            "The model encodes each modality and fuses them into one representation.",
            "The classifier outputs logits, softmax turns them into probabilities, and argmax gives the predicted emotion.",
            "Training compares the prediction to the label using the selected loss function; evaluation only measures the prediction without updating weights.",
            "The analysis scripts export predictions, metrics, and confusion matrices for later inspection.",
        ],
    )
    add_para(
        doc,
        "This sequence is the easiest way to explain the code under questioning. If you can move through the pipeline in this order, you can answer almost any 'where did this prediction come from?' question with confidence.",
    )

    doc.add_heading("10. Common Code-Level Failure Modes", level=1)
    add_table(
        doc,
        ["Symptom", "Likely code location", "Meaning"],
        [
            ["mat1 and mat2 shapes cannot be multiplied", "demo feature extraction or model input shape", "The feature dimension does not match the model's expected video dimension."],
            ["Missing keys during load", "checkpoint/model mismatch", "A new head or fusion branch exists in the architecture but not in the checkpoint, or vice versa."],
            ["No face found in preview", "face detection step in the preview script", "The detector fell back to center crop, or the frame was too hard to detect."],
            ["Low accuracy but high confidence", "fusion + calibration behavior", "The model has learned a strong but potentially biased association."],
            ["Neutral-heavy predictions", "training imbalance and fusion dynamics", "The model may be overfitting the majority class or underusing the visual branch."],
        ],
    )
    add_para(
        doc,
        "These failure modes are useful because they link the observed symptom to an actual file in the codebase. That is the right way to troubleshoot a research pipeline: do not guess globally; locate the exact stage where the behavior is introduced.",
    )

    doc.add_heading("11. Important Custom Functions by Pipeline Stage", level=1)
    add_para(
        doc,
        "The table below highlights the custom functions that matter most when you want to explain the code end to end. These are the functions a reviewer is most likely to ask about because they connect the raw data, the trained model, and the demo outputs.",
    )
    add_table(
        doc,
        ["Stage", "Custom function", "Why it matters"],
        [
            ["Face-crop manifest", "get_face_cascade()", "Loads the face detector used before ViT feature extraction."],
            ["Face-crop manifest", "crop_face_frame()", "Finds the face region, crops it, pads it, and resizes it to the ViT input size."],
            ["Face-crop manifest", "extract_vit_facecrop_embeddings()", "Samples frames, runs ViT, and saves the resulting facial embeddings as .npy."],
            ["Preprocessing", "sample_video_frames()", "Creates the fixed-length frame sequence used by the visual pipeline."],
            ["Preprocessing", "load_audio_waveform()", "Creates the waveform input used by the pretrained speech branch."],
            ["Dataset", "ManifestDataset.__getitem__()", "Converts one manifest row into a model-ready multimodal example."],
            ["Dataset", "collate_samples()", "Pads variable-length samples and creates masks for batching."],
            ["Model", "BidirectionalCrossAttentionCMT._masked_pool()", "Implements CLS / MEAN / MAX / MIN pooling after fusion."],
            ["Model", "GatedModalFusion.forward()", "Learns how much to trust each modality before classification."],
            ["Training", "compute_class_weights()", "Builds class weights from the training split for imbalance-aware losses."],
            ["Training", "FocalLoss.forward()", "Implements focal loss for hard-example emphasis."],
            ["Training", "run_epoch()", "Runs one full train or validation pass and computes metrics."],
            ["Evaluation", "main()", "Reloads the checkpoint and exports metrics and prediction tables."],
            ["Demo", "predict_phase1_raw_mp4_demo.py:main()", "Turns one raw video clip into final probabilities and top-k predictions."],
        ],
    )
    add_para(
        doc,
        "Student explanation: these functions are where the actual project logic lives. The library code gives you the building blocks, but these custom functions decide how MELD is turned into features, how the model sees those features, and how the outputs are written for analysis.",
    )

    doc.add_heading("11.1 Detailed Logic of the Custom Functions", level=2)
    add_para(
        doc,
        "The goal of this subsection is to explain each custom function in terms of inputs, internal steps, and outputs. This is the part you can use if an examiner asks you to walk through a function line by line.",
    )
    add_bullets(
        doc,
        [
            "get_face_cascade(): loads OpenCV's Haar face detector from the installed haarcascade XML file and raises an error if the cascade cannot be loaded. The function is a dependency setup step, because the later crop code cannot work unless face detection is available.",
            "crop_face_frame(): converts the sampled RGB frame back to uint8, runs face detection, chooses the largest detected face, pads the box slightly so the crop is not too tight, and resizes the crop to the configured ViT frame size. If no face is found, it falls back to a square center crop so processing can continue.",
            "extract_vit_facecrop_embeddings(): samples frames from the raw clip, applies face cropping to each sampled frame, batches those crop images through the pretrained ViT image processor and model, extracts the CLS token embedding from each frame, concatenates the per-frame embeddings, and saves them as a float32 array. This is the real bridge from raw video to reusable model input.",
            "sample_video_frames(): opens the raw video with OpenCV, counts the total frames, selects evenly spaced frame indices, reads only those frames, resizes them to the configured size, and pads with zeros if the clip is shorter than the requested frame count. The goal is to make every utterance look like a fixed-length sequence.",
            "load_audio_waveform(): loads the raw waveform through librosa or ffmpeg-based fallback logic, truncates or pads it to a fixed duration, removes invalid numeric values, and clips the final waveform to a safe range. This function prepares the speech branch input for the pretrained audio encoder path.",
            "ManifestDataset.__getitem__(): takes one manifest row and turns it into a single multimodal example. In pretrained mode it creates BERT input_ids, an attention mask, a HuBERT-style waveform, and the video feature array. In legacy mode it creates simpler token and feature tensors. The function is the row-to-example conversion step used by both training and evaluation.",
            "collate_samples(): batches multiple examples together by padding 1D and 2D arrays to the longest sequence length in the batch, building the corresponding masks, and returning a dictionary of batched tensors. This is essential because PyTorch cannot stack variable-length sequences directly.",
            "BidirectionalCrossAttentionCMT._masked_pool(): reduces a fused sequence back into a single vector. CLS returns the first token, MEAN averages valid positions, MAX selects the strongest valid activation, and MIN selects the smallest valid activation while masking padding tokens. This is the paper-aligned pooling logic that lets the model compare different aggregation strategies.",
            "GatedModalFusion.forward(): stacks the text, audio, and video vectors, sends them through a small gate network, converts the gate logits into softmax weights, and computes a weighted sum of the modalities. The function learns when to trust each branch instead of treating them equally.",
            "compute_class_weights(): counts the labels in the training subset, converts those counts into inverse-frequency weights, and returns a tensor that can be used by weighted cross-entropy or weighted focal loss. This is how the code reacts to MELD's class imbalance.",
            "FocalLoss.forward(): computes per-sample cross-entropy, converts each term into a confidence-like pt value, down-weights easy samples by (1-pt)^gamma, and averages the result. The intention is to make hard or minority-class examples matter more during optimization.",
            "run_epoch(): loops through a DataLoader batch by batch, applies modality masking, calls the model, computes the loss, optionally backpropagates and steps the optimizer, and aggregates prediction statistics. The same function is used for both train and validation, with the optimizer argument deciding whether learning happens.",
            "main() in evaluate.py: reloads the saved checkpoint, reconstructs the model configuration, creates the evaluation dataset, runs inference without gradient updates, and writes metrics plus per-sample CSV rows. This is what makes the evaluation reproducible and inspectable.",
            "main() in predict_phase1_raw_mp4_demo.py: loads a single sample_id, builds the one-example batch, optionally extracts raw-video features from the mp4 clip, applies the trained model, computes softmax probabilities, and prints the final label plus the top-k alternatives. This is the exact path you use during the ESA demo.",
        ],
    )
    add_para(
        doc,
        "A simple oral-exam strategy is to describe each function in three parts: what it receives, what it transforms, and what it produces. That keeps your explanation grounded in the code rather than in vague theory.",
    )

    doc.add_heading("11.2 Representative Code Snippets", level=2)
    add_para(
        doc,
        "The following excerpts show the exact code patterns behind the custom functions. They are abridged to keep the document readable, but they are enough to explain how each stage works in the repository.",
    )
    add_para(doc, "Face-crop manifest builder: face detection, crop selection, and ViT embedding extraction.", bold=True)
    add_code(
        doc,
        snippet(read("scripts/build_meld_vit_facecrop_manifest.py").splitlines(), 27, 97),
    )
    add_para(doc, "Line-by-line note for the crop helper inside this snippet:", bold=True)
    add_bullets(
        doc,
        [
            "h, w = frame_rgb.shape[:2] gets the height and width of the current frame.",
            "side = min(h, w) chooses the largest square that fits inside the frame.",
            "top = max((h - side) // 2, 0) computes the vertical start point for a centered crop.",
            "left = max((w - side) // 2, 0) computes the horizontal start point for a centered crop.",
            "crop = frame_rgb[top : top + side, left : left + side] extracts the square crop from the middle of the frame.",
            "if crop.size == 0: return frame_rgb protects against an invalid empty crop.",
            "return crop returns the centered square patch used when face detection is not available.",
        ],
    )
    add_para(doc, "Line-by-line note for the face-detection and embedding part:", bold=True)
    add_bullets(
        doc,
        [
            "get_face_cascade() loads the face detector XML and confirms that face detection is available.",
            "crop_face_frame(...) converts the frame to uint8, detects the largest face, pads the box slightly, and resizes the crop to the ViT input size.",
            "If no face is detected, the function falls back to center_crop_square so the pipeline keeps running.",
            "extract_vit_facecrop_embeddings(...) samples frames, converts them into face crops, batches them through ViT, and collects the CLS embeddings.",
            "outputs.last_hidden_state[:, 0, :] selects the CLS token vector from the ViT output for each frame.",
            "np.concatenate(embeddings, axis=0) joins the per-batch embeddings into one sequence for the utterance.",
            "np.save(video_feat_path, vit_embeddings) stores the final facial-cue representation as a reusable .npy file.",
        ],
    )
    add_para(doc, "Preprocessing: frame sampling and audio waveform loading.", bold=True)
    add_code(
        doc,
        snippet(read("src/data/preprocessing.py").splitlines(), 28, 166),
    )
    add_para(doc, "Line-by-line note for preprocessing:", bold=True)
    add_bullets(
        doc,
        [
            "sample_video_frames(...) opens the raw clip and decides which frames to keep.",
            "indices = np.linspace(...) spreads the selected frames across the full clip duration.",
            "cap.read() loads frames one by one from OpenCV.",
            "cv2.resize(...) makes every sampled frame the same spatial size.",
            "cv2.cvtColor(..., cv2.COLOR_BGR2RGB) converts OpenCV's default channel order into RGB for downstream use.",
            "If too few frames are available, zero frames are appended so the output length is still fixed.",
            "load_audio_waveform(...) loads raw speech, pads or truncates it, and returns a safe mono waveform.",
            "normalize_text(...) removes whitespace noise so the text branch sees a cleaned transcript.",
        ],
    )
    add_para(doc, "Dataset bridge: manifest row to model-ready example plus batch collation.", bold=True)
    add_code(
        doc,
        snippet(read("src/data/dataset.py").splitlines(), 121, 260),
    )
    add_para(doc, "Line-by-line note for the dataset bridge:", bold=True)
    add_bullets(
        doc,
        [
            "_encode_text_pretrained(...) uses the Hugging Face tokenizer to create BERT-style input_ids and attention_mask.",
            "_load_array(...) checks whether the path already points to cached .npy features or to raw media that needs preprocessing.",
            "The suffix branch decides whether audio or video features are loaded directly or derived on demand.",
            "In pretrained mode, __getitem__ returns text_input_ids, text_attention_mask, audio_waveform, and video_features.",
            "In legacy mode, __getitem__ returns simpler token and feature arrays instead.",
            "collate_samples() pads the batch so each tensor has a uniform shape and builds masks for the valid positions.",
        ],
    )
    add_para(doc, "Model fusion: pooling choices, gated fusion, and cross-attention pooling logic.", bold=True)
    add_code(
        doc,
        snippet(read("src/models/model.py").splitlines(), 101, 246),
    )
    add_para(doc, "Line-by-line note for the model fusion code:", bold=True)
    add_bullets(
        doc,
        [
            "CrossModalFusion stacks the modality vectors and adds learned modality tokens before Transformer fusion.",
            "The pooling branch chooses CLS, MAX, MIN, or MEAN depending on the run configuration.",
            "GatedModalFusion.forward(...) computes softmax weights over the stacked modality vectors and blends them.",
            "BidirectionalCrossAttentionCMT uses MultiheadAttention in both directions so text can attend to audio and audio can attend to text.",
            "_masked_pool(...) ignores padding positions when computing pooled representations.",
            "PretrainedBackboneEncoder wraps a Hugging Face backbone and projects its pooled output to the shared fusion dimension.",
            "LegalMemoCMTPhase1 orchestrates which encoders and fusion blocks are active for the current run.",
        ],
    )
    add_para(doc, "Training: imbalance-aware loss selection and the epoch loop.", bold=True)
    add_code(
        doc,
        snippet(read("src/train/train.py").splitlines(), 23, 204),
    )
    add_para(doc, "Line-by-line note for the training code:", bold=True)
    add_bullets(
        doc,
        [
            "parse_modalities() converts a comma-separated modality string into a clean Python set.",
            "parse_pooling(), parse_fusion_mode(), and parse_encoder_mode() sanitize the CLI arguments and prevent invalid settings.",
            "compute_class_weights() counts labels in the training subset and converts inverse frequency into weights.",
            "FocalLoss.forward() computes hard-example-focused loss by down-weighting easy samples.",
            "build_dataset() attaches the tokenizer only when the run is in pretrained/paper mode.",
            "apply_modality_mask() removes the modalities that the current ablation is not supposed to use.",
            "run_epoch() performs the actual train/validation pass, computes the loss, and aggregates predictions.",
            "main() parses the arguments, builds the model and loaders, and saves the best checkpoint by validation performance.",
        ],
    )
    add_para(doc, "Evaluation and raw-mp4 demo: loading the checkpoint, computing metrics, and printing probabilities.", bold=True)
    add_code(
        doc,
        snippet(read("src/train/evaluate.py").splitlines(), 23, 149),
    )
    add_para(doc, "Line-by-line note for evaluation:", bold=True)
    add_bullets(
        doc,
        [
            "The parser takes the manifest, checkpoint, split, modality selection, and device.",
            "torch.load(...) restores the saved model configuration from training.",
            "build_dataset(...) reconstructs the same input path used during training.",
            "DataLoader(..., shuffle=False) keeps the evaluation order fixed and reproducible.",
            "torch.no_grad() prevents gradient computation because evaluation does not learn.",
            "torch.softmax(logits, dim=-1) turns scores into probabilities for reporting.",
            "The rows list collects per-sample outputs for later confusion-matrix analysis.",
            "The metrics dictionary reports accuracy, weighted accuracy, unweighted accuracy, macro F1, and weighted F1.",
        ],
    )
    add_code(
        doc,
        snippet(read("scripts/predict_phase1_raw_mp4_demo.py").splitlines(), 47, 207),
    )
    add_para(doc, "Line-by-line note for the raw-mp4 demo:", bold=True)
    add_bullets(
        doc,
        [
            "The parser receives the manifest row, checkpoint, raw mp4 path, vision mode, and device.",
            "The checkpoint is loaded first so the demo recreates the trained configuration correctly.",
            "The matched manifest row provides the transcript and label for the live explanation.",
            "PreprocessConfig fixes the frame size, number of frames, and audio settings.",
            "The vision mode decides whether the demo uses face-crop embeddings or full-frame embeddings.",
            "The extracted embeddings are cached as .npy so the demo output can be inspected later.",
            "A one-sample ManifestDataset is created so the raw clip behaves like a normal batch item.",
            "apply_modality_mask() enforces the chosen modal combination before inference.",
            "softmax probabilities and top-k scores are printed so you can explain confidence during the ESA.",
        ],
    )

    doc.add_heading("12. What Is Standard Library vs What Is Custom", level=1)
    add_table(
        doc,
        ["Part", "Mostly standard / reused", "Project-specific work"],
        [
            ["Tokenization / backbones", "Hugging Face AutoTokenizer, AutoModel", "How the encoders are wired into the Phase 1 architecture"],
            ["Training loop", "PyTorch DataLoader, optimizer, loss APIs", "Loss selection, modality masking, checkpoint selection, and metrics reporting"],
            ["Preprocessing", "OpenCV, NumPy, librosa / soundfile", "MELD-specific manifest creation and face-crop logic"],
            ["Demo", "shell wrapper + Python CLI", "Raw-video extraction and per-sample explanation format"],
        ],
    )
    add_para(
        doc,
        "This distinction matters in a viva. It helps you explain what was reused from libraries and what you actually designed and implemented yourself.",
    )

    doc.add_heading("13. Final Traceability Summary", level=1)
    add_para(
        doc,
        "If you want to summarize the whole codebase in one sentence, say this: raw MELD annotations are converted into manifests, manifests are converted into padded feature batches, feature batches are encoded by text/audio/video modules in the model, train.py learns the parameters, evaluate.py measures them, and the raw-mp4 demo shows the same pipeline on one clip at a time.",
    )
    add_para(
        doc,
        "The most important lesson from the code is traceability. Every output file can be traced back to a manifest row, every manifest row can be traced back to a raw utterance, and every prediction can be traced to a specific preprocessing and fusion decision in the source code.",
    )

    doc.add_heading("14. Appendix: Main Files to Inspect First", level=1)
    add_table(
        doc,
        ["File", "Why to inspect it first"],
        [
            ["src/data/preprocessing.py", "Shows frame sampling, audio loading, and feature construction."],
            ["src/data/dataset.py", "Shows how manifest rows become tensors."],
            ["src/models/model.py", "Shows encoders, fusion, pooling, and classifier heads."],
            ["src/train/train.py", "Shows the learning loop, loss, and checkpoint selection."],
            ["src/train/evaluate.py", "Shows metrics and prediction export."],
            ["scripts/run_phase1_raw_mp4_demo.sh", "Shows how a raw clip is turned into a live demo command."],
            ["scripts/predict_phase1_raw_mp4_demo.py", "Shows the raw-video inference flow."],
        ],
    )

    doc.add_paragraph()
    foot = doc.add_paragraph()
    r = foot.add_run("End of document.")
    r.italic = True

    return doc


def main() -> None:
    doc = build_doc()
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
