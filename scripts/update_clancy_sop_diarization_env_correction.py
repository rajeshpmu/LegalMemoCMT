from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOP = ROOT / "implementation_docments/LegalMemoCMT_Clancy_Corpus_SOP.docx"


def code(document: Document, value: str) -> None:
    p = document.add_paragraph(style="No Spacing")
    r = p.add_run(value)
    r.font.name = "Courier New"
    r.font.size = Pt(8)


document = Document(SOP)
document.add_section(WD_SECTION.NEW_PAGE)
document.add_heading("Diarization Environment Compatibility Correction", level=1)
document.add_paragraph(
    "The first installation into `/opt/anaconda3/bin/python` installed pyannote.audio 3.4.0 alongside "
    "torchaudio 2.11.0. The preflight correctly reported that `torchaudio.AudioMetaData` was missing. "
    "Upgrading pyannote to 4.x was not retained as the operational solution because pyannote import "
    "segfaulted under the current Anaconda Python 3.13 runtime. The supported repository solution is an "
    "isolated Python 3.9 environment with a matched Torch audio stack."
)
document.add_heading("Working Environment", level=2)
for item in [
    "Environment: `.venv-diarization`",
    "Python: system `/usr/bin/python3` 3.9 used only to create the environment",
    "torch: 2.2.2",
    "torchaudio: 2.2.2",
    "pyannote.audio: 3.4.0",
    "torchaudio.AudioMetaData: available",
    "pip check: no broken requirements reported",
]:
    document.add_paragraph(item, style="List Bullet")
document.add_paragraph(
    "Keep this environment separate from `.venv` and `/opt/anaconda3`. The Phase 2 LegalMELD and ViT "
    "pipelines continue using their existing environments; only source-level speaker diarization uses "
    "`.venv-diarization`."
)
document.add_heading("Corrected Setup Command", level=2)
code(document, '''bash phase2/setup_clancy_diarization.sh''')
document.add_paragraph(
    "The setup script creates `.venv-diarization` when absent, upgrades pip, installs numpy below 2.0, "
    "pins torch and torchaudio to 2.2.2, and installs pyannote.audio 3.4.0. It then requests a Hugging Face "
    "token only if model access is to be checked."
)
document.add_heading("Corrected Preflight", level=2)
code(document, '''export HF_TOKEN="<your Hugging Face token>"
./.venv-diarization/bin/python phase2/check_clancy_diarization_prerequisites.py \\
  --model pyannote/speaker-diarization-3.1 \\
  --load-model''')
document.add_paragraph(
    "The checker now validates all three compatibility conditions: torch imports, torchaudio exposes "
    "`AudioMetaData`, and pyannote.audio imports. It then optionally loads the Hugging Face pipeline."
)
document.add_heading("Corrected Diarization Command", level=2)
code(document, '''PYTHON_BIN=$PWD/.venv-diarization/bin/python \\
HF_TOKEN="$HF_TOKEN" \\
bash phase2/run_clancy_diarization.sh \\
  --input-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_subtitle_evidence.csv \\
  --output-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_dataset_manifest_vit_200_diarized.csv \\
  --segments-csv data/processed/phase2/clancy/duration_0_8_to_20/clancy_diarization_segments_200.csv \\
  --device cpu \\
  --max-sources 1''')
document.add_heading("Observed Resolution", level=2)
document.add_paragraph(
    "The isolated environment was created and validated successfully. Its observed versions are torch "
    "2.2.2, torchaudio 2.2.2, and pyannote.audio 3.4.0, with `AudioMetaData=True` and no broken pip "
    "requirements. Model loading and diarization remain pending until the user supplies the Hugging Face "
    "token in the shell."
)
document.save(SOP)
print(f"Updated {SOP}")
